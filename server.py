"""
社团活动统计分析系统 - 后端服务
技术栈: Flask + SQLite + Pandas + MapReduce
"""
from dotenv import load_dotenv
load_dotenv()
from flask import Flask, request, jsonify, send_file, send_from_directory, session, redirect
import sqlite3
import os
import uuid
from datetime import datetime, timedelta
from io import BytesIO
import json
import re
import time
import math
import hashlib
import random
import string
from storage import storage, migrate_old_paths

def local_time(utc_str):
    if not utc_str:
        return utc_str
    try:
        dt = datetime.strptime(utc_str[:19], '%Y-%m-%d %H:%M:%S')
        local = dt + timedelta(hours=8)
        return local.strftime('%Y-%m-%d %H:%M:%S')
    except:
        return utc_str

def cn_now():
    return datetime.utcnow() + timedelta(hours=8)


class DataCleaner:
    @staticmethod
    def clean(data, headers, options=None):
        if options is None:
            options = {}
        report = {'original': len(data), 'dedupRemoved': 0, 'nullHandled': 0, 'normFixed': 0, 'after': 0}
        cleaned = list(data)
        seen = set()
        deduped = []
        for row in cleaned:
            key = '|||'.join(str(row.get(h, '')) for h in headers)
            if key not in seen:
                seen.add(key)
                deduped.append(row)
        report['dedupRemoved'] = len(cleaned) - len(deduped)
        cleaned = deduped
        valid = []
        for row in cleaned:
            has_valid = any(str(row.get(h, '')).strip() not in ('', 'undefined', 'null', 'NaN') for h in headers)
            if has_valid:
                valid.append(row)
        report['nullHandled'] = len(cleaned) - len(valid)
        cleaned = valid
        norm_fixed = 0
        normalized = []
        for row in cleaned:
            new_row = {}
            changed = False
            for h in headers:
                v = row.get(h, '')
                if isinstance(v, str):
                    nv = re.sub(r'\s+', ' ', v).strip()
                    if nv != v:
                        changed = True
                    new_row[h] = nv
                else:
                    new_row[h] = v
            if changed:
                norm_fixed += 1
            normalized.append(new_row)
        report['normFixed'] = norm_fixed
        report['after'] = len(normalized)
        return {'data': normalized, 'report': report}

    @staticmethod
    def auto_detect_column(headers, keywords, exclude=None):
        exclude = exclude or []
        best_h = None
        best_score = -1
        for h in headers:
            if h in exclude:
                continue
            h_lower = h.lower().replace(' ', '').replace('_', '')
            for kw in keywords:
                kw_lower = kw.lower().replace(' ', '').replace('_', '')
                if kw_lower == h_lower:
                    score = 1000
                elif kw_lower in h_lower or h_lower in kw_lower:
                    score = len(kw_lower) + 50
                else:
                    # 模糊匹配：检查关键词的每个字是否出现在表头中
                    common = sum(1 for ch in kw_lower if ch in h_lower)
                    if common >= len(kw_lower) * 0.6 and common >= 2:
                        score = common
                    else:
                        continue
                if score > best_score:
                    best_score = score
                    best_h = h
        return best_h


class LRUCache:
    def __init__(self, limit=50):
        self.limit = limit
        self.cache = {}
        self.order = []
        self.ttl = {}

    def get(self, key):
        if key not in self.cache:
            return None
        if key in self.ttl and time.time() > self.ttl[key]:
            self.delete(key)
            return None
        self.order.remove(key)
        self.order.append(key)
        return self.cache[key]

    def set(self, key, value, ttl=60):
        if key in self.cache:
            self.order.remove(key)
        elif len(self.cache) >= self.limit:
            oldest = self.order.pop(0)
            del self.cache[oldest]
            self.ttl.pop(oldest, None)
        self.cache[key] = value
        self.order.append(key)
        self.ttl[key] = time.time() + ttl

    def delete(self, key):
        if key in self.cache:
            del self.cache[key]
            self.order.remove(key)
            self.ttl.pop(key, None)

    def clear(self):
        self.cache.clear()
        self.order.clear()
        self.ttl.clear()


class StatsService:
    @staticmethod
    def descriptive_stats(values):
        if not values:
            return {'mean': 0, 'median': 0, 'stdDev': 0, 'mode': 0, 'min': 0, 'max': 0, 'count': 0}
        n = len(values)
        mean = sum(values) / n
        sorted_v = sorted(values)
        median = sorted_v[n // 2] if n % 2 else (sorted_v[n // 2 - 1] + sorted_v[n // 2]) / 2
        std_dev = (sum((v - mean) ** 2 for v in values) / (n - 1)) ** 0.5 if n > 1 else 0
        freq = {}
        for v in values:
            freq[v] = freq.get(v, 0) + 1
        max_freq = max(freq.values())
        mode = ', '.join(str(k) for k, v in freq.items() if v == max_freq)
        return {'mean': round(mean, 2), 'median': median, 'stdDev': round(std_dev, 2), 'mode': mode, 'min': min(values), 'max': max(values), 'count': n}


class Database:
    def __init__(self, db_path):
        self.db_path = db_path
        os.makedirs(os.path.dirname(db_path), exist_ok=True)
        self.init_db()

    def get_conn(self):
        conn = sqlite3.connect(self.db_path, timeout=10)
        conn.row_factory = sqlite3.Row
        return conn

    def init_db(self):
        conn = self.get_conn()
        try:
            conn.executescript('''
                CREATE TABLE IF NOT EXISTS activity_records (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_uuid TEXT NOT NULL,
                    club_name TEXT NOT NULL,
                    activity_date TEXT,
                    activity_content TEXT,
                    location TEXT,
                    participant_count INTEGER DEFAULT 0,
                    raw_data TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS guidance_records (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_uuid TEXT NOT NULL,
                    session_id TEXT NOT NULL,
                    club_name TEXT NOT NULL,
                    teacher_name TEXT,
                    guidance_date TEXT,
                    guidance_content TEXT,
                    raw_data TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS uploaded_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    uuid TEXT UNIQUE NOT NULL,
                    original_name TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    file_size INTEGER,
                    row_count INTEGER DEFAULT 0,
                    upload_time DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS base_data (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    uuid TEXT UNIQUE NOT NULL,
                    name TEXT NOT NULL,
                    original_name TEXT NOT NULL,
                    headers TEXT,
                    row_count INTEGER DEFAULT 0,
                    data_json TEXT,
                    upload_time DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS online_activity_data (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL,
                    has_photo INTEGER DEFAULT 0,
                    has_summary INTEGER DEFAULT 0,
                    activity_title TEXT,
                    activity_date TEXT,
                    raw_json TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_activity_club ON activity_records(club_name);
                CREATE INDEX IF NOT EXISTS idx_guidance_club ON guidance_records(club_name);
                CREATE INDEX IF NOT EXISTS idx_guidance_session ON guidance_records(session_id);
                CREATE TABLE IF NOT EXISTS club_uploads (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_token TEXT NOT NULL,
                    club_name TEXT NOT NULL,
                    file_name TEXT NOT NULL,
                    file_path TEXT NOT NULL,
                    file_type TEXT,
                    description TEXT,
                    upload_time DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_club_uploads_token ON club_uploads(club_token);
                CREATE TABLE IF NOT EXISTS club_tokens (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT UNIQUE NOT NULL,
                    token TEXT UNIQUE NOT NULL,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT UNIQUE NOT NULL,
                    password TEXT NOT NULL,
                    role TEXT DEFAULT 'user',
                    club_name TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_users_role ON users(role);
                CREATE INDEX IF NOT EXISTS idx_users_club ON users(club_name);
                CREATE TABLE IF NOT EXISTS notifications (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER,
                    title TEXT NOT NULL,
                    content TEXT,
                    type TEXT DEFAULT 'info',
                    is_read INTEGER DEFAULT 0,
                    link TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_notifications_user ON notifications(user_id);
                CREATE TABLE IF NOT EXISTS notices (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT NOT NULL,
                    content TEXT,
                    is_top INTEGER DEFAULT 0,
                    attachment_path TEXT,
                    attachment_name TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS club_showcase (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL,
                    description TEXT,
                    image_path TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS club_tools (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL,
                    tool_type TEXT NOT NULL,
                    title TEXT NOT NULL,
                    description TEXT,
                    options TEXT,
                    vote_mode TEXT DEFAULT 'single',
                    limit_count INTEGER DEFAULT 0,
                    format_hint TEXT,
                    deadline TEXT,
                    status TEXT DEFAULT 'active',
                    results TEXT DEFAULT '{}',
                    participants TEXT DEFAULT '[]',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS tree_hole (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    content TEXT NOT NULL,
                    scope TEXT DEFAULT 'public',
                    club_name TEXT DEFAULT '',
                    status TEXT DEFAULT 'active',
                    admin_note TEXT DEFAULT '',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS joint_activities (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL,
                    title TEXT NOT NULL,
                    description TEXT,
                    support_needed TEXT,
                    status TEXT DEFAULT 'open',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS joint_replies (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    activity_id INTEGER NOT NULL,
                    club_name TEXT NOT NULL,
                    reply_type TEXT DEFAULT 'message',
                    content TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (activity_id) REFERENCES joint_activities(id)
                );
                CREATE TABLE IF NOT EXISTS club_profiles (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL UNIQUE,
                    description TEXT DEFAULT '',
                    star_rating INTEGER DEFAULT 0,
                    show_star INTEGER DEFAULT 0,
                    registration_form TEXT DEFAULT '',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS club_teachers (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL,
                    teacher_name TEXT NOT NULL,
                    photo_path TEXT DEFAULT '',
                    introduction TEXT DEFAULT '',
                    user_id INTEGER DEFAULT 0,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS teacher_profiles (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL UNIQUE,
                    work_id TEXT DEFAULT '',
                    real_name TEXT DEFAULT '',
                    phone TEXT DEFAULT '',
                    email TEXT DEFAULT '',
                    avatar_path TEXT DEFAULT '',
                    introduction TEXT DEFAULT '',
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (user_id) REFERENCES users(id)
                );
                CREATE TABLE IF NOT EXISTS teacher_clubs (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL,
                    club_name TEXT NOT NULL,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (user_id) REFERENCES users(id),
                    UNIQUE(user_id, club_name)
                );
                CREATE TABLE IF NOT EXISTS club_registrations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL,
                    student_name TEXT DEFAULT '',
                    student_phone TEXT DEFAULT '',
                    form_data TEXT DEFAULT '',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS checkin_sessions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    club_name TEXT NOT NULL,
                    activity_name TEXT DEFAULT '',
                    checkin_code TEXT NOT NULL UNIQUE,
                    location_lat REAL DEFAULT 0,
                    location_lng REAL DEFAULT 0,
                    location_name TEXT DEFAULT '',
                    status TEXT DEFAULT 'open',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    closed_at DATETIME
                );
                CREATE INDEX IF NOT EXISTS idx_checkin_sessions_club ON checkin_sessions(club_name);
                CREATE TABLE IF NOT EXISTS checkin_records (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    session_id INTEGER NOT NULL,
                    club_name TEXT NOT NULL,
                    student_name TEXT NOT NULL,
                    student_class TEXT DEFAULT '',
                    student_id TEXT DEFAULT '',
                    checkin_method TEXT DEFAULT 'code',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (session_id) REFERENCES checkin_sessions(id)
                );
                CREATE INDEX IF NOT EXISTS idx_checkin_records_session ON checkin_records(session_id);
                CREATE TABLE IF NOT EXISTS teacher_checkin_checkout (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    session_id INTEGER DEFAULT 0,
                    teacher_user_id INTEGER NOT NULL,
                    club_name TEXT NOT NULL,
                    checkin_time DATETIME,
                    checkout_time DATETIME,
                    checkin_lat REAL DEFAULT 0,
                    checkin_lng REAL DEFAULT 0,
                    checkout_lat REAL DEFAULT 0,
                    checkout_lng REAL DEFAULT 0,
                    checkout_method TEXT DEFAULT '',
                    checkout_code TEXT DEFAULT '',
                    status TEXT DEFAULT 'checked_in',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    checkin_address TEXT DEFAULT '',
                    checkout_address TEXT DEFAULT ''
                );
                CREATE INDEX IF NOT EXISTS idx_teacher_checkin_teacher ON teacher_checkin_checkout(teacher_user_id);
                CREATE INDEX IF NOT EXISTS idx_teacher_checkin_session ON teacher_checkin_checkout(session_id);
                CREATE INDEX IF NOT EXISTS idx_teacher_checkin_status ON teacher_checkin_checkout(status);
                CREATE TABLE IF NOT EXISTS admin_notifications (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT NOT NULL,
                    content TEXT DEFAULT '',
                    target_type TEXT NOT NULL,
                    target_club TEXT DEFAULT '',
                    sender_id INTEGER NOT NULL,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
                CREATE TABLE IF NOT EXISTS admin_notification_reads (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    notification_id INTEGER NOT NULL,
                    user_id INTEGER NOT NULL,
                    read_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    cleared INTEGER DEFAULT 0,
                    UNIQUE(notification_id, user_id)
                );
            ''')
            conn.execute('''
                CREATE TABLE IF NOT EXISTS material_views (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    group_id TEXT NOT NULL,
                    user_id INTEGER NOT NULL,
                    viewed_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(group_id, user_id)
                );
            ''')
            conn.execute('''
                CREATE TABLE IF NOT EXISTS weekly_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    role TEXT NOT NULL,
                    club_name TEXT DEFAULT '',
                    week_start TEXT NOT NULL,
                    week_end TEXT NOT NULL,
                    content TEXT NOT NULL,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                );
            ''')
            conn.commit()
        finally:
            conn.close()
        for alter in [
            'ALTER TABLE club_uploads ADD COLUMN status TEXT DEFAULT "pending"',
            'ALTER TABLE club_uploads ADD COLUMN reject_reason TEXT DEFAULT ""',
            'ALTER TABLE club_uploads ADD COLUMN group_id TEXT DEFAULT ""',
            'ALTER TABLE club_uploads ADD COLUMN category TEXT DEFAULT ""',
            'ALTER TABLE club_uploads ADD COLUMN source TEXT DEFAULT "upload"',
            'ALTER TABLE teacher_checkin_checkout ADD COLUMN checkout_photo_path TEXT DEFAULT ""',
            'ALTER TABLE club_tools ADD COLUMN per_user_limit INTEGER DEFAULT 0',
            'ALTER TABLE club_tools ADD COLUMN anonymous INTEGER DEFAULT 0',
            'ALTER TABLE club_tools ADD COLUMN show_counts INTEGER DEFAULT 1',
            'ALTER TABLE club_tools ADD COLUMN results_visible INTEGER DEFAULT 1',
            'ALTER TABLE club_uploads ADD COLUMN summary_desc TEXT DEFAULT ""',
            'ALTER TABLE notifications ADD COLUMN cleared INTEGER DEFAULT 0',
            'ALTER TABLE admin_notification_reads ADD COLUMN cleared INTEGER DEFAULT 0',
        ]:
            try:
                conn = self.get_conn()
                conn.execute(alter)
                conn.commit()
                conn.close()
            except:
                pass
        try:
            conn = self.get_conn()
            cols = [r[1] for r in conn.execute("PRAGMA table_info(material_views)").fetchall()]
            if cols and 'upload_id' in cols:
                conn.execute("DROP TABLE material_views")
                conn.execute('''CREATE TABLE IF NOT EXISTS material_views (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    group_id TEXT NOT NULL,
                    user_id INTEGER NOT NULL,
                    viewed_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(group_id, user_id)
                )''')
                conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS feature_reads (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                feature TEXT NOT NULL,
                read_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(user_id, feature)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS user_profiles (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL UNIQUE,
                real_name TEXT DEFAULT '',
                student_id TEXT DEFAULT '',
                grade TEXT DEFAULT '',
                class_name TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                email TEXT DEFAULT '',
                avatar_path TEXT DEFAULT '',
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS club_departments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                dept_name TEXT NOT NULL,
                description TEXT DEFAULT '',
                image_path TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS club_members (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                user_id INTEGER DEFAULT 0,
                username TEXT DEFAULT '',
                real_name TEXT DEFAULT '',
                student_id_num TEXT DEFAULT '',
                class_name TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                department TEXT DEFAULT '',
                specialty TEXT DEFAULT '',
                source TEXT DEFAULT 'registration',
                joined_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS club_cadres (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                user_id INTEGER NOT NULL,
                real_name TEXT DEFAULT '',
                student_id_num TEXT DEFAULT '',
                added_by INTEGER DEFAULT 0,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(club_name, user_id)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS club_notices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                title TEXT NOT NULL,
                content TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS feedbacks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                user_id INTEGER DEFAULT 0,
                type TEXT DEFAULT 'other',
                title TEXT NOT NULL,
                body TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                result TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS club_favorites (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                club_name TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(user_id, club_name)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS quit_applications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                username TEXT DEFAULT '',
                club_name TEXT NOT NULL,
                reason TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                handler_note TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                handled_at DATETIME
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        for alter in [
            'ALTER TABLE club_registrations ADD COLUMN status TEXT DEFAULT "pending"',
            'ALTER TABLE club_registrations ADD COLUMN department TEXT DEFAULT ""',
            'ALTER TABLE club_registrations ADD COLUMN user_id INTEGER DEFAULT 0',
            'ALTER TABLE club_registrations ADD COLUMN student_class TEXT DEFAULT ""',
            'ALTER TABLE club_registrations ADD COLUMN student_id_num TEXT DEFAULT ""',
            'ALTER TABLE club_registrations ADD COLUMN specialty TEXT DEFAULT ""',
            'ALTER TABLE club_registrations ADD COLUMN reviewed_at DATETIME DEFAULT ""',
            'ALTER TABLE club_registrations ADD COLUMN attachment TEXT DEFAULT ""',
            'ALTER TABLE club_profiles ADD COLUMN emblem_url TEXT DEFAULT ""',
            'ALTER TABLE club_departments ADD COLUMN parent_id INTEGER DEFAULT 0',
            'ALTER TABLE checkin_sessions ADD COLUMN activity_time TEXT DEFAULT ""',
            'ALTER TABLE club_profiles ADD COLUMN president TEXT DEFAULT ""',
            'ALTER TABLE club_profiles ADD COLUMN category TEXT DEFAULT ""',
            'ALTER TABLE club_profiles ADD COLUMN guiding_unit TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN start_time TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN end_time TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN checkin_method TEXT DEFAULT "qrcode"',
            'ALTER TABLE checkin_sessions ADD COLUMN activity_content TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN plan_path TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN plan_text TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN teacher_ids TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN completion_photo TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN summary_path TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN summary_text TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN is_completed INTEGER DEFAULT 0',
            'ALTER TABLE checkin_sessions ADD COLUMN warning TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN warning_reason TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN checkout_code TEXT DEFAULT ""',
            'ALTER TABLE checkin_sessions ADD COLUMN checkout_method TEXT DEFAULT ""',
            'ALTER TABLE user_profiles ADD COLUMN college TEXT DEFAULT ""',
            'ALTER TABLE user_profiles ADD COLUMN grade TEXT DEFAULT ""',
            'ALTER TABLE club_members ADD COLUMN college TEXT DEFAULT ""',
            'ALTER TABLE club_registrations ADD COLUMN college TEXT DEFAULT ""',
            'ALTER TABLE checkin_records ADD COLUMN college TEXT DEFAULT ""',
            'ALTER TABLE recruitment_signups ADD COLUMN college TEXT DEFAULT ""',
            'ALTER TABLE recruitment_signups ADD COLUMN student_phone TEXT DEFAULT ""',
            'ALTER TABLE ai_pets ADD COLUMN mood TEXT DEFAULT "happy"',
            'ALTER TABLE ai_pets ADD COLUMN last_interact TEXT DEFAULT ""',
            'ALTER TABLE feedbacks ADD COLUMN activity_name TEXT DEFAULT ""',
            'ALTER TABLE feedbacks ADD COLUMN activity_time TEXT DEFAULT ""',
            'ALTER TABLE feedbacks ADD COLUMN file_path TEXT DEFAULT ""',
            'ALTER TABLE feedbacks ADD COLUMN file_name TEXT DEFAULT ""',
            'ALTER TABLE feedbacks ADD COLUMN files_json TEXT DEFAULT ""',
            "UPDATE feedbacks SET status='resolved' WHERE status='handled'",
            'ALTER TABLE scoring_rules ADD COLUMN date_start TEXT DEFAULT ""',
            'ALTER TABLE scoring_rules ADD COLUMN date_end TEXT DEFAULT ""',
            'ALTER TABLE final_credits ADD COLUMN semester TEXT DEFAULT ""',
            '''CREATE TABLE IF NOT EXISTS teacher_club_requests (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                teacher_user_id INTEGER NOT NULL,
                teacher_name TEXT DEFAULT '',
                introduction TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                requested_by INTEGER DEFAULT 0,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                reviewed_at DATETIME DEFAULT ''
            )''',
            '''CREATE TABLE IF NOT EXISTS sms_codes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                phone TEXT NOT NULL,
                code TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                expires_at DATETIME NOT NULL,
                used INTEGER DEFAULT 0
            )''',
            '''CREATE TABLE IF NOT EXISTS email_codes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT NOT NULL,
                code TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                expires_at DATETIME NOT NULL,
                used INTEGER DEFAULT 0
            )''',
            '''CREATE TABLE IF NOT EXISTS phone_change_requests (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                old_phone TEXT DEFAULT '',
                new_phone TEXT NOT NULL,
                status TEXT DEFAULT 'pending',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                reviewed_at DATETIME DEFAULT '',
                reviewed_by INTEGER DEFAULT 0
            )''',
            'ALTER TABLE scoring_submission_items ADD COLUMN activity_count REAL DEFAULT 0',
            'ALTER TABLE scoring_submission_items ADD COLUMN other_score REAL DEFAULT 0',
            'ALTER TABLE scoring_submissions ADD COLUMN date_start TEXT DEFAULT ""',
            'ALTER TABLE scoring_submissions ADD COLUMN date_end TEXT DEFAULT ""',
        ]:
            try:
                conn = self.get_conn()
                conn.execute(alter)
                conn.commit()
                conn.close()
            except:
                pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS ai_pets (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL UNIQUE,
                animal_type TEXT NOT NULL,
                stage TEXT DEFAULT 'egg',
                exp INTEGER DEFAULT 0,
                name TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS excellent_clubs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL UNIQUE,
                selected_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                star_rating INTEGER DEFAULT 0,
                activity_count INTEGER DEFAULT 0,
                teacher_guided_count INTEGER DEFAULT 0,
                material_approved_count INTEGER DEFAULT 0,
                excellent_activity_count INTEGER DEFAULT 0,
                combined_score REAL DEFAULT 0
            )''')
            # 兼容已有数据库：添加新列（已存在则跳过）
            for _col, _type in [('star_rating','INTEGER'),('activity_count','INTEGER'),('teacher_guided_count','INTEGER'),('material_approved_count','INTEGER'),('excellent_activity_count','INTEGER'),('combined_score','REAL')]:
                try:
                    conn.execute(f'ALTER TABLE excellent_clubs ADD COLUMN {_col} {_type} DEFAULT 0')
                except Exception:
                    pass
            conn.execute('''CREATE TABLE IF NOT EXISTS excellent_activities (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_id TEXT NOT NULL UNIQUE,
                selected_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.execute('''CREATE TABLE IF NOT EXISTS offcampus_requests (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                title TEXT NOT NULL,
                location TEXT DEFAULT '',
                activity_date TEXT DEFAULT '',
                description TEXT DEFAULT '',
                file_path TEXT DEFAULT '',
                file_name TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                reject_reason TEXT DEFAULT '',
                submitted_by TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_offcampus_status ON offcampus_requests(status)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_offcampus_club ON offcampus_requests(club_name)')
            conn.execute('''CREATE TABLE IF NOT EXISTS recruitments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                title TEXT NOT NULL,
                description TEXT DEFAULT '',
                recruit_type TEXT DEFAULT 'member',
                max_count INTEGER DEFAULT 0,
                current_count INTEGER DEFAULT 0,
                status TEXT DEFAULT 'pending',
                created_by INTEGER DEFAULT 0,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                approved_at DATETIME DEFAULT '',
                deadline TEXT DEFAULT ''
            )''')
            conn.execute('''CREATE TABLE IF NOT EXISTS recruitment_signups (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                recruitment_id INTEGER NOT NULL,
                user_id INTEGER DEFAULT 0,
                student_name TEXT NOT NULL,
                student_class TEXT DEFAULT '',
                student_id_num TEXT DEFAULT '',
                club_name TEXT DEFAULT '',
                signed_up_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (recruitment_id) REFERENCES recruitments(id)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS ai_chat_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                tool_calls TEXT DEFAULT '',
                FOREIGN KEY (user_id) REFERENCES users(id)
            )''')
            try:
                conn.execute('ALTER TABLE ai_chat_history ADD COLUMN tool_calls TEXT DEFAULT ""')
            except Exception:
                pass
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS finance_records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                type TEXT NOT NULL,
                category TEXT DEFAULT '',
                amount REAL NOT NULL,
                description TEXT DEFAULT '',
                record_date TEXT DEFAULT '',
                recorder TEXT DEFAULT '',
                attachment_path TEXT DEFAULT '',
                attachment_name TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS finance_permissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                user_id INTEGER NOT NULL,
                username TEXT DEFAULT '',
                real_name TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                reason TEXT DEFAULT '',
                reviewed_by TEXT DEFAULT '',
                reviewed_at DATETIME DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS finance_managers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                user_id INTEGER NOT NULL,
                username TEXT DEFAULT '',
                real_name TEXT DEFAULT '',
                granted_by TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(club_name, user_id)
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS scoring_rules (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                star_level INTEGER NOT NULL UNIQUE,
                collective_limit REAL DEFAULT 0,
                individual_limit REAL DEFAULT 0,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            for i in range(6):
                try:
                    conn.execute('INSERT OR IGNORE INTO scoring_rules (star_level, collective_limit, individual_limit) VALUES (?, 0, 0)', (i,))
                except:
                    pass
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS scoring_submissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL,
                status TEXT DEFAULT 'draft',
                submitted_by INTEGER DEFAULT 0,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.execute('''CREATE TABLE IF NOT EXISTS scoring_submission_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                submission_id INTEGER NOT NULL,
                student_user_id INTEGER DEFAULT 0,
                student_name TEXT DEFAULT '',
                student_id_num TEXT DEFAULT '',
                college TEXT DEFAULT '',
                class_name TEXT DEFAULT '',
                total_workload REAL DEFAULT 0,
                final_score REAL DEFAULT 0,
                FOREIGN KEY (submission_id) REFERENCES scoring_submissions(id)
            )''')
            conn.execute('''CREATE TABLE IF NOT EXISTS scoring_teacher_reviews (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                submission_id INTEGER NOT NULL,
                teacher_user_id INTEGER NOT NULL,
                teacher_name TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                review_note TEXT DEFAULT '',
                reviewed_at DATETIME DEFAULT '',
                UNIQUE(submission_id, teacher_user_id),
                FOREIGN KEY (submission_id) REFERENCES scoring_submissions(id)
            )''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_scoring_sub_club ON scoring_submissions(club_name)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_scoring_sub_status ON scoring_submissions(status)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_scoring_items_sub ON scoring_submission_items(submission_id)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_scoring_reviews_sub ON scoring_teacher_reviews(submission_id)')
            conn.execute('''CREATE TABLE IF NOT EXISTS final_credits (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_id_num TEXT NOT NULL,
                student_name TEXT DEFAULT '',
                college TEXT DEFAULT '',
                class_name TEXT DEFAULT '',
                club1 TEXT DEFAULT '',
                score1 REAL DEFAULT 0,
                club2 TEXT DEFAULT '',
                score2 REAL DEFAULT 0,
                club_scores TEXT DEFAULT '',
                final_credit REAL DEFAULT 0,
                calculated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_final_credits_sid ON final_credits(student_id_num)')
            conn.execute('''CREATE TABLE IF NOT EXISTS scoring_club_overrides (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                club_name TEXT NOT NULL UNIQUE,
                collective_limit REAL DEFAULT NULL,
                individual_limit REAL DEFAULT NULL,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS workload_submissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_user_id INTEGER NOT NULL,
                student_name TEXT DEFAULT '',
                club_name TEXT NOT NULL,
                item_name TEXT NOT NULL,
                score REAL NOT NULL,
                status TEXT DEFAULT 'pending',
                reviewer_id INTEGER DEFAULT 0,
                reviewer_name TEXT DEFAULT '',
                review_note TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                reviewed_at DATETIME DEFAULT '',
                FOREIGN KEY (student_user_id) REFERENCES users(id)
            )''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_workload_club ON workload_submissions(club_name)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_workload_student ON workload_submissions(student_user_id)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_workload_status ON workload_submissions(status)')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS doc_index (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                source_table TEXT NOT NULL,
                source_id INTEGER NOT NULL,
                club_name TEXT DEFAULT '',
                file_name TEXT DEFAULT '',
                file_path TEXT DEFAULT '',
                description TEXT DEFAULT '',
                tags TEXT DEFAULT '',
                category TEXT DEFAULT '',
                status TEXT DEFAULT '',
                source_type TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                indexed_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_doc_tags ON doc_index(tags)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_doc_source ON doc_index(source_table, source_id)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_doc_club ON doc_index(club_name)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_doc_category ON doc_index(category)')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS doc_embeddings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id INTEGER NOT NULL,
                doc_type TEXT NOT NULL,
                doc_text TEXT NOT NULL,
                embedding BLOB,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            conn.execute('''CREATE TABLE IF NOT EXISTS location_checkins (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER DEFAULT 0,
                username TEXT DEFAULT '',
                role TEXT DEFAULT '',
                club_name TEXT DEFAULT '',
                session_id INTEGER DEFAULT 0,
                lat REAL DEFAULT 0,
                lng REAL DEFAULT 0,
                address TEXT DEFAULT '',
                checkin_time DATETIME DEFAULT CURRENT_TIMESTAMP
            )''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_loc_checkin_user ON location_checkins(user_id)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_loc_checkin_session ON location_checkins(session_id)')
            conn.commit()
            conn.close()
        except:
            pass
        try:
            conn = self.get_conn()
            cursor = conn.execute('SELECT COUNT(*) as c FROM users WHERE username=?', ('0066',))
            if cursor.fetchone()['c'] == 0:
                conn.execute('INSERT INTO users (username, password, role, club_name) VALUES (?, ?, ?, ?)',
                           ('0066', '0000', 'admin', None))
                conn.commit()
            conn.close()
        except:
            pass
        print('  数据库初始化完成')
    try:
        conn = db.get_conn()
        cnt = migrate_old_paths(conn)
        conn.close()
        if cnt > 0:
            print(f'  文件路径迁移完成: {cnt} 个文件')
    except:
        pass


app = Flask(__name__, static_folder='public', static_url_path='')
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024
app.secret_key = 'club-stats-secret-key-2026'

@app.after_request
def add_no_cache_headers(response):
    if response.content_type and ('text/html' in response.content_type or 'javascript' in response.content_type or 'application/json' in response.content_type):
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

COLLEGES = ['数智科技产业学院', '信息商务学院', '学前教育学院', '育婴育幼学院', '人文教育学院', '艺术学院', '体育学院']

db = Database(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'club_stats.db'))
cache = LRUCache(limit=50)
cleaner = DataCleaner()


def get_current_user():
    user_id = session.get('user_id')
    if not user_id:
        return None
    conn = db.get_conn()
    try:
        user = conn.execute('SELECT id, username, role, club_name FROM users WHERE id=?', (user_id,)).fetchone()
        if not user:
            return None
        result = {'id': user['id'], 'username': user['username'], 'role': user['role'], 'club_name': user['club_name']}
        # 获取真实姓名
        if user['role'] == 'teacher':
            tp = conn.execute('SELECT real_name FROM teacher_profiles WHERE user_id=?', (user_id,)).fetchone()
            if tp and tp['real_name']:
                result['real_name'] = tp['real_name']
        else:
            up = conn.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (user_id,)).fetchone()
            if up and up['real_name']:
                result['real_name'] = up['real_name']
        if 'real_name' not in result:
            result['real_name'] = ''
        # 社团负责人如果 club_name 为空，尝试从 club_cadres/club_members 查找
        if user['role'] == 'user' and not result['club_name']:
            cr = conn.execute('SELECT DISTINCT club_name FROM club_cadres WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_members WHERE user_id=?', (user_id, user_id)).fetchall()
            if cr:
                result['club_name'] = cr[0]['club_name']
    finally:
        conn.close()
    return result


def is_cadre_of_club(user_id, club_name):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT id FROM club_cadres WHERE user_id=? AND club_name=?', (user_id, club_name)).fetchone()
    finally:
        conn.close()
    return row is not None


def get_cadre_clubs(user_id):
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT club_name FROM club_cadres WHERE user_id=?', (user_id,)).fetchall()
    finally:
        conn.close()
    return [r['club_name'] for r in rows]


@app.route('/api/register', methods=['POST'])
def register():
    data = request.json or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    club_name = data.get('clubName', '').strip()
    role = data.get('role', 'user').strip()
    student_id = data.get('studentId', '').strip()
    work_id = data.get('workId', '').strip()
    club_names = data.get('clubNames', [])
    if not username or not password:
        return jsonify({'error': '请输入姓名和密码'}), 400
    if len(username) < 2 or len(password) < 4:
        return jsonify({'error': '用户名至少2位，密码至少4位'}), 400
    if role not in ('user', 'student', 'teacher'):
        role = 'user'
    if role == 'teacher':
        if not work_id:
            return jsonify({'error': '请输入工号'}), 400
        if not club_names:
            return jsonify({'error': '请选择至少一个所属社团'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
        if existing:
            return jsonify({'error': '该姓名已注册'}), 400
        if role == 'user' and club_name:
            club_existing = conn.execute('SELECT id FROM users WHERE club_name=? AND role="user"', (club_name,)).fetchone()
            if club_existing:
                return jsonify({'error': f'社团「{club_name}」已注册账号，一个社团只能注册一个账号'}), 400
        first_club = club_name or (club_names[0] if club_names else '')
        conn.execute('INSERT INTO users (username, password, role, club_name) VALUES (?, ?, ?, ?)',
                    (username, password, role, first_club or None))
        new_club = False
        if role == 'user' and club_name:
            token_exists = conn.execute('SELECT id FROM club_tokens WHERE club_name=?', (club_name,)).fetchone()
            if not token_exists:
                import uuid
                token = uuid.uuid4().hex[:12]
                conn.execute('INSERT INTO club_tokens (club_name, token) VALUES (?, ?)', (club_name, token))
                new_club = True
        conn.commit()
        user = conn.execute('SELECT id, username, role, club_name FROM users WHERE username=?', (username,)).fetchone()
        if new_club:
            admins = conn.execute('SELECT id FROM users WHERE role="admin"').fetchall()
            for a in admins:
                send_notification(a['id'], '🏫 新社团注册', f'社团「{club_name}」已由 {username} 注册，已自动创建上传通道', 'info', '/dashboard.html', conn=conn)
            conn.commit()
        if student_id and role in ('user', 'student'):
            real_name_reg = data.get('realName', '').strip()
            class_name_reg = data.get('className', '').strip()
            college_reg = data.get('college', '').strip()
            phone_reg = data.get('phone', '').strip()
            conn.execute('INSERT OR IGNORE INTO user_profiles (user_id, student_id, real_name, class_name, college, phone) VALUES (?, ?, ?, ?, ?, ?)',
                (user['id'], student_id, real_name_reg, class_name_reg, college_reg, phone_reg))
        if role == 'student' and first_club:
            existing_member = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=?', (first_club, user['id'])).fetchone()
            if not existing_member:
                profile_data = {}
                try:
                    prow = conn.execute('SELECT real_name, student_id, class_name, college, phone FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
                    if prow:
                        profile_data = dict(prow)
                except:
                    pass
                conn.execute('''INSERT INTO club_members (club_name, user_id, username, real_name, student_id_num, class_name, phone, department, specialty, college, source)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'registration')''',
                    (first_club, user['id'], username, profile_data.get('real_name', ''), profile_data.get('student_id', student_id),
                     profile_data.get('class_name', ''), profile_data.get('phone', ''), '', '', profile_data.get('college', '')))
            conn.commit()
        if role == 'teacher':
            conn.execute('INSERT OR IGNORE INTO teacher_profiles (user_id, work_id) VALUES (?, ?)', (user['id'], work_id))
            for cn in club_names:
                cn = cn.strip()
                if not cn:
                    continue
                conn.execute('INSERT OR IGNORE INTO teacher_clubs (user_id, club_name) VALUES (?, ?)', (user['id'], cn))
                existing_teacher = conn.execute('SELECT id FROM club_teachers WHERE club_name=? AND user_id=?', (cn, user['id'])).fetchone()
                if not existing_teacher:
                    conn.execute('INSERT INTO club_teachers (club_name, teacher_name, user_id) VALUES (?, ?, ?)', (cn, username, user['id']))
            conn.commit()
            admins = conn.execute('SELECT id FROM users WHERE role="admin"').fetchall()
            for a in admins:
                send_notification(a['id'], '👨‍🏫 新指导老师注册', f'指导老师 {username}（工号：{work_id}）已注册，指导社团：{"、".join(club_names)}', 'info', '/dashboard.html', conn=conn)
            conn.commit()
    finally:
        conn.close()
    session['user_id'] = user['id']
    return jsonify({'success': True, 'user': {'id': user['id'], 'username': user['username'], 'role': user['role'], 'clubName': user['club_name']}})


@app.route('/api/login', methods=['POST'])
def login():
    data = request.json or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    if not username or not password:
        return jsonify({'error': '请输入姓名和密码'}), 400
    conn = db.get_conn()
    try:
        user = conn.execute('SELECT id, username, password, role, club_name FROM users WHERE username=?', (username,)).fetchone()
    finally:
        conn.close()
    if not user or user['password'] != password:
        return jsonify({'error': '姓名或密码错误'}), 400
    session['user_id'] = user['id']
    # 获取真实姓名
    real_name = ''
    conn2 = db.get_conn()
    try:
        if user['role'] == 'teacher':
            tp = conn2.execute('SELECT real_name FROM teacher_profiles WHERE user_id=?', (user['id'],)).fetchone()
            if tp and tp['real_name']:
                real_name = tp['real_name']
        else:
            up = conn2.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
            if up and up['real_name']:
                real_name = up['real_name']
    finally:
        conn2.close()
    return jsonify({'success': True, 'user': {'id': user['id'], 'username': user['username'], 'role': user['role'], 'clubName': user['club_name'], 'realName': real_name}})


@app.route('/api/logout', methods=['POST'])
def logout():
    session.pop('user_id', None)
    return jsonify({'success': True})


@app.route('/api/reset-password', methods=['POST'])
def reset_password():
    data = request.json or {}
    email = data.get('email', '').strip()
    code = data.get('code', '').strip()
    new_password = data.get('newPassword', '').strip()
    if not email or not code or not new_password:
        return jsonify({'error': '请填写完整信息'}), 400
    if len(new_password) < 4:
        return jsonify({'error': '新密码至少4个字符'}), 400
    conn = db.get_conn()
    try:
        # 验证码已在verify步骤标记为已使用，此处检查已使用则说明未走验证流程
        row = conn.execute('SELECT id, expires_at, used FROM email_codes WHERE email=? AND code=? ORDER BY id DESC LIMIT 1', (email, code)).fetchone()
        if not row:
            return jsonify({'error': '验证码错误'}), 400
        if row['used']:
            return jsonify({'error': '请先验证验证码'}), 400
        from datetime import datetime as dt
        try:
            expires = dt.strptime(row['expires_at'][:19], '%Y-%m-%d %H:%M:%S')
            if dt.now() > expires:
                return jsonify({'error': '验证码已过期'}), 400
        except:
            pass
        conn.execute('UPDATE email_codes SET used=1 WHERE id=?', (row['id'],))
        # 通过邮箱查找用户
        user = conn.execute('SELECT u.id, u.username, u.role FROM users u LEFT JOIN user_profiles up ON u.id=up.user_id LEFT JOIN teacher_profiles tp ON u.id=tp.user_id WHERE up.email=? OR tp.email=?', (email, email)).fetchone()
        if not user:
            return jsonify({'error': '该邮箱未注册'}), 400
        conn.execute('UPDATE users SET password=? WHERE id=?', (new_password, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '密码重置成功'})


@app.route('/api/club-departments/<club_name>')
def get_departments(club_name):
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT * FROM club_departments WHERE club_name=? ORDER BY id', (club_name,)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'name': r['dept_name'], 'description': r['description'], 'imagePath': r['image_path'], 'parentId': r['parent_id'] if 'parent_id' in r.keys() else 0} for r in rows]})


@app.route('/api/club-departments', methods=['POST'])
def add_department():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    club = data.get('clubName', '').strip()
    if not club:
        club = user['club_name'] or ''
    if not club:
        return jsonify({'error': '无社团信息'}), 400
    if user['role'] == 'student' and not is_cadre_of_club(user['id'], club):
        return jsonify({'error': '无权限操作该社团'}), 403
    if user['role'] == 'teacher':
        tc_conn2 = db.get_conn()
        try:
            tc_row2 = tc_conn2.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
        finally:
            tc_conn2.close()
        if not tc_row2:
            return jsonify({'error': '无权限操作该社团'}), 403
    dept_name = data.get('name', '').strip()
    if not dept_name:
        return jsonify({'error': '请输入部门名称'}), 400
    description = data.get('description', '').strip()
    parent_id = data.get('parentId', 0)
    conn = db.get_conn()
    try:
        conn.execute('INSERT INTO club_departments (club_name, dept_name, description, parent_id) VALUES (?, ?, ?, ?)', (club, dept_name, description, parent_id))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-departments/<int:did>', methods=['PUT'])
def update_department(did):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    dept_name = data.get('name', '').strip()
    description = data.get('description', '').strip()
    conn = db.get_conn()
    try:
        conn.execute('UPDATE club_departments SET dept_name=?, description=? WHERE id=?', (dept_name, description, did))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-departments/<int:did>', methods=['DELETE'])
def delete_department(did):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM club_departments WHERE id=?', (did,))
        conn.execute('UPDATE club_departments SET parent_id=0 WHERE parent_id=?', (did,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/department-image/<int:did>', methods=['POST'])
def upload_dept_image(did):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if 'image' not in request.files:
        return jsonify({'error': '未上传文件'}), 400
    file = request.files['image']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else 'png'
    img_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'dept_images')
    os.makedirs(img_dir, exist_ok=True)
    filename = f"dept_{did}_{int(time.time())}.{ext}"
    filepath = os.path.join(img_dir, filename)
    file.save(filepath)
    rel_path = f"/api/dept-image-file/{filename}"
    conn = db.get_conn()
    try:
        conn.execute('UPDATE club_departments SET image_path=? WHERE id=?', (rel_path, did))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'path': rel_path})


@app.route('/api/dept-image-file/<filename>')
def serve_dept_image(filename):
    img_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'dept_images')
    return send_from_directory(img_dir, filename)


@app.route('/api/my-club-memberships')
def my_club_memberships():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, club_name, real_name, class_name, college, department, joined_at FROM club_members WHERE user_id=?', (user['id'],)).fetchall()
        if not rows:
            rows_by_name = conn.execute('SELECT id, club_name, real_name, class_name, college, department, joined_at FROM club_members WHERE username=? AND user_id=0', (user['username'],)).fetchall()
            if not rows_by_name:
                profile = conn.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
                if profile and profile['real_name']:
                    rows_by_name = conn.execute('SELECT id, club_name, real_name, class_name, college, department, joined_at FROM club_members WHERE real_name=? AND user_id=0', (profile['real_name'],)).fetchall()
            for r in rows_by_name:
                conn.execute('UPDATE club_members SET user_id=?, username=? WHERE id=?', (user['id'], user['username'], r['id']))
            if rows_by_name:
                conn.commit()
            rows = rows_by_name
        result = []
        for r in rows:
            college_val = ''
            try:
                college_val = r['college'] if r['college'] else ''
            except:
                pass
            result.append({
                'clubName': r['club_name'],
                'realName': r['real_name'],
                'className': r['class_name'],
                'college': college_val,
                'department': r['department'] or '',
                'joinedAt': r['joined_at'] or ''
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/my-registrations')
def get_my_registrations():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    date_where = ''
    date_params = []
    if start_date:
        date_where += ' AND date(r.created_at)>=?'
        date_params.append(start_date)
    if end_date:
        date_where += ' AND date(r.created_at)<=?'
        date_params.append(end_date)
    conn = db.get_conn()
    try:
        rows = conn.execute(f'''SELECT r.id, r.club_name, r.status, r.department, r.student_name, r.student_class, 
            r.student_id_num, r.specialty, r.student_phone, r.created_at, r.reviewed_at
            FROM club_registrations r WHERE r.user_id=?{date_where} ORDER BY r.created_at DESC''', [user['id']] + date_params).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'clubName': r['club_name'], 'status': r['status'],
        'department': r['department'], 'studentName': r['student_name'], 'studentClass': r['student_class'],
        'studentIdNum': r['student_id_num'], 'specialty': r['specialty'], 'studentPhone': r['student_phone'],
        'time': local_time(r['created_at']), 'reviewedAt': local_time(r['reviewed_at'])} for r in rows]})


@app.route('/api/registration-approve', methods=['POST'])
def approve_registrations():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student'):
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    ids = data.get('ids', [])
    action = data.get('action', 'approve')
    department = data.get('department', '')
    if not ids:
        return jsonify({'error': '请选择要操作的记录'}), 400
    if action not in ('approve', 'reject'):
        return jsonify({'error': '无效操作'}), 400
    status = 'approved' if action == 'approve' else 'rejected'
    conn = db.get_conn()
    try:
        for rid in ids:
            reg = conn.execute('SELECT * FROM club_registrations WHERE id=?', (rid,)).fetchone()
            if not reg:
                continue
            if user['role'] == 'student' and not is_cadre_of_club(user['id'], reg['club_name']):
                continue
            if status == 'approved':
                if reg['user_id'] and reg['user_id'] != 0:
                    member_count = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE user_id=?', (reg['user_id'],)).fetchone()
                    if member_count and member_count['c'] >= 2:
                        conn.execute('DELETE FROM club_registrations WHERE id=?', (rid,))
                        continue
                existing = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=? AND user_id!=0', (reg['club_name'], reg['user_id'])).fetchone()
                if not existing:
                    conn.execute('''INSERT INTO club_members (club_name, user_id, username, real_name, student_id_num, class_name, phone, department, specialty, college, source)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'registration')''',
                        (reg['club_name'], reg['user_id'], '', reg['student_name'], reg['student_id_num'],
                         reg['student_class'], reg['student_phone'], department or reg['department'] or '', reg['specialty'], reg['college'] if 'college' in reg.keys() else ''))
                conn.execute('DELETE FROM club_registrations WHERE id=?', (rid,))
            else:
                conn.execute('DELETE FROM club_registrations WHERE id=?', (rid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/current-user')
def current_user_api():
    user = get_current_user()
    if not user:
        return jsonify({'loggedIn': False})
    result = {'id': user['id'], 'username': user['username'], 'role': user['role'], 'clubName': user['club_name'], 'realName': user.get('real_name', '')}
    if user['role'] == 'student':
        conn = db.get_conn()
        try:
            club_rows = conn.execute('SELECT club_name FROM club_members WHERE user_id=? ORDER BY club_name', (user['id'],)).fetchall()
            clubs = [cr['club_name'] for cr in club_rows if cr['club_name']]
            if clubs:
                result['clubName'] = clubs[0]
                result['clubs'] = clubs
            elif user['club_name']:
                result['clubs'] = [user['club_name']]
            cadre_rows = conn.execute('SELECT club_name FROM club_cadres WHERE user_id=?', (user['id'],)).fetchall()
            if cadre_rows:
                result['cadreClubs'] = [cr['club_name'] for cr in cadre_rows]
        finally:
            conn.close()
    elif user['role'] == 'user':
        conn = db.get_conn()
        try:
            if not result['clubName']:
                cr = conn.execute('SELECT DISTINCT club_name FROM club_cadres WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_members WHERE user_id=?', (user['id'], user['id'])).fetchall()
                if cr:
                    result['clubName'] = cr[0]['club_name']
            cadre_rows = conn.execute('SELECT club_name FROM club_cadres WHERE user_id=?', (user['id'],)).fetchall()
            if cadre_rows:
                result['cadreClubs'] = [cr['club_name'] for cr in cadre_rows]
        finally:
            conn.close()
    elif user['role'] == 'teacher':
        conn = db.get_conn()
        try:
            tc_rows = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=? ORDER BY club_name', (user['id'],)).fetchall()
            tc_clubs = [r['club_name'] for r in tc_rows if r['club_name']]
            if tc_clubs:
                result['clubs'] = tc_clubs
                result['clubName'] = tc_clubs[0]
        finally:
            conn.close()
    return jsonify({'loggedIn': True, 'user': result})


@app.route('/api/club-logo', methods=['POST'])
def upload_club_logo():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    club = request.form.get('clubName', '').strip()
    if not club:
        club = user['club_name'] or ''
    if not club:
        return jsonify({'error': '无社团信息'}), 400
    if user['role'] == 'student' and not is_cadre_of_club(user['id'], club):
        return jsonify({'error': '无权限操作该社团'}), 403
    if user['role'] == 'teacher':
        tc_conn = db.get_conn()
        try:
            tc_row = tc_conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
        finally:
            tc_conn.close()
        if not tc_row:
            return jsonify({'error': '无权限操作该社团'}), 403
    file = request.files.get('logo')
    if not file:
        return jsonify({'error': '请选择图片'}), 400
    import os
    upload_dir = os.path.join(os.path.dirname(__file__), 'public', 'uploads', 'logos')
    os.makedirs(upload_dir, exist_ok=True)
    ext = os.path.splitext(file.filename)[1] or '.png'
    filename = f'{club.replace("/", "_")}{ext}'
    filepath = os.path.join(upload_dir, filename)
    file.save(filepath)
    logo_url = f'/uploads/logos/{filename}'
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM club_profiles WHERE club_name=?', (club,)).fetchone()
        if existing:
            conn.execute('UPDATE club_profiles SET emblem_url=?, updated_at=CURRENT_TIMESTAMP WHERE club_name=?', (logo_url, club))
        else:
            conn.execute('INSERT INTO club_profiles (club_name, emblem_url) VALUES (?, ?)', (club, logo_url))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'logoUrl': logo_url})


@app.route('/api/club-leaders')
def get_club_leaders():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT u.id, u.username, u.club_name, up.real_name, up.student_id, up.class_name, up.college, up.phone FROM users u LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.role="user" ORDER BY u.club_name').fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'username': r['username'], 'clubName': r['club_name'] or '', 'realName': r['real_name'] or '', 'studentId': r['student_id'] or '', 'className': r['class_name'] or '', 'college': (r['college'] if 'college' in r.keys() else '') or '', 'phone': r['phone'] or ''} for r in rows]})


@app.route('/api/club-leaders/<int:uid>', methods=['PUT'])
def update_club_leader(uid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    real_name = data.get('realName', '').strip()
    student_id = data.get('studentId', '').strip()
    class_name = data.get('className', '').strip()
    college = data.get('college', '').strip() or data.get('studentCollege', '').strip()
    phone = data.get('phone', '').strip()
    club_name = data.get('clubName', '').strip()
    conn = db.get_conn()
    try:
        if club_name:
            conn.execute('UPDATE users SET club_name=? WHERE id=?', (club_name, uid))
        existing = conn.execute('SELECT id FROM user_profiles WHERE user_id=?', (uid,)).fetchone()
        if existing:
            conn.execute('UPDATE user_profiles SET real_name=?, student_id=?, class_name=?, college=?, phone=? WHERE user_id=?', (real_name, student_id, class_name, college, phone, uid))
        else:
            conn.execute('INSERT INTO user_profiles (user_id, real_name, student_id, class_name, college, phone) VALUES (?, ?, ?, ?, ?, ?)', (uid, real_name, student_id, class_name, college, phone))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-leaders/<int:uid>', methods=['DELETE'])
def delete_club_leader(uid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM user_profiles WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM users WHERE id=? AND role="user"', (uid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-leaders', methods=['POST'])
def add_club_leader():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    username = data.get('username', '').strip()
    password = data.get('password', '0000').strip()
    club_name = data.get('clubName', '').strip()
    real_name = data.get('realName', '').strip()
    student_id = data.get('studentId', '').strip()
    class_name = data.get('className', '').strip()
    phone = data.get('phone', '').strip()
    if not username or not club_name:
        return jsonify({'error': '请填写姓名和社团名称'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
        if existing:
            return jsonify({'error': '该姓名已注册'}), 400
        conn.execute('INSERT INTO users (username, password, role, club_name) VALUES (?, ?, "user", ?)', (username, password, club_name))
        new_user = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
        if new_user:
            conn.execute('INSERT INTO user_profiles (user_id, real_name, student_id, class_name, phone) VALUES (?, ?, ?, ?, ?)', (new_user['id'], real_name, student_id, class_name, phone))
            token_exists = conn.execute('SELECT id FROM club_tokens WHERE club_name=?', (club_name,)).fetchone()
            if not token_exists:
                token = uuid.uuid4().hex[:12]
                conn.execute('INSERT INTO club_tokens (club_name, token) VALUES (?, ?)', (club_name, token))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-leaders/batch-template')
def batch_replace_template():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = '社团负责人更换'
    headers = ['社团名称', '社长姓名', '学号', '学院', '班级', '联系电话']
    ws.append(headers)
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color='667eea', end_color='667eea', fill_type='solid')
    header_font_white = Font(bold=True, size=11, color='FFFFFF')
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    for col_idx in range(1, len(headers) + 1):
        ws.column_dimensions[chr(64 + col_idx)].width = 18
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name='社团负责人更换模板.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/club-leaders/batch-replace', methods=['POST'])
def batch_replace_leaders():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    if 'file' not in request.files:
        return jsonify({'error': '请上传Excel文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '请上传Excel文件'}), 400
    import openpyxl
    try:
        f.seek(0)
        wb = openpyxl.load_workbook(f)
        ws = wb.active
    except Exception as e:
        return jsonify({'error': 'Excel文件解析失败: ' + str(e)}), 400
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    success_count = 0
    fail_count = 0
    fail_details = []
    conn = db.get_conn()
    try:
        for idx, row in enumerate(rows, 2):
            if not row or not row[0]:
                continue
            club_name = str(row[0]).strip() if row[0] else ''
            real_name = str(row[1]).strip() if len(row) > 1 and row[1] else ''
            student_id = str(row[2]).strip() if len(row) > 2 and row[2] else ''
            college = str(row[3]).strip() if len(row) > 3 and row[3] else ''
            class_name = str(row[4]).strip() if len(row) > 4 and row[4] else ''
            phone = str(row[5]).strip() if len(row) > 5 and row[5] else ''
            if not club_name:
                fail_count += 1
                fail_details.append({'row': idx, 'error': '社团名称为空'})
                continue
            club_profile = conn.execute('SELECT club_name FROM club_profiles WHERE club_name=?', (club_name,)).fetchone()
            if not club_profile:
                fail_count += 1
                fail_details.append({'row': idx, 'clubName': club_name, 'error': '社团不存在'})
                continue
            existing_leader = conn.execute('SELECT id FROM users WHERE role="user" AND club_name=?', (club_name,)).fetchone()
            if existing_leader:
                uid = existing_leader['id']
                existing_profile = conn.execute('SELECT id FROM user_profiles WHERE user_id=?', (uid,)).fetchone()
                if existing_profile:
                    conn.execute('UPDATE user_profiles SET real_name=?, student_id=?, college=?, class_name=?, phone=? WHERE user_id=?', (real_name, student_id, college, class_name, phone, uid))
                else:
                    conn.execute('INSERT INTO user_profiles (user_id, real_name, student_id, college, class_name, phone) VALUES (?, ?, ?, ?, ?, ?)', (uid, real_name, student_id, college, class_name, phone))
                conn.execute('UPDATE club_profiles SET president=? WHERE club_name=?', (real_name, club_name))
            else:
                username = student_id if student_id else 'leader_' + uuid.uuid4().hex[:8]
                dup = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
                if dup:
                    username = username + '_' + uuid.uuid4().hex[:4]
                conn.execute('INSERT INTO users (username, password, role, club_name) VALUES (?, "0000", "user", ?)', (username, club_name))
                new_user = conn.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
                if new_user:
                    conn.execute('INSERT INTO user_profiles (user_id, real_name, student_id, college, class_name, phone) VALUES (?, ?, ?, ?, ?, ?)', (new_user['id'], real_name, student_id, college, class_name, phone))
                conn.execute('UPDATE club_profiles SET president=? WHERE club_name=?', (real_name, club_name))
                token_exists = conn.execute('SELECT id FROM club_tokens WHERE club_name=?', (club_name,)).fetchone()
                if not token_exists:
                    token = uuid.uuid4().hex[:12]
                    conn.execute('INSERT INTO club_tokens (club_name, token) VALUES (?, ?)', (club_name, token))
            success_count += 1
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({'error': '处理失败: ' + str(e)}), 500
    finally:
        conn.close()
    return jsonify({'success': True, 'data': {'successCount': success_count, 'failCount': fail_count, 'failDetails': fail_details}})


@app.route('/api/club-cadres', methods=['GET'])
def get_club_cadres():
    user = get_current_user()
    if not user or user['role'] != 'user':
        return jsonify({'error': '仅社团负责人可操作'}), 403
    club_name = user.get('club_name') or ''
    if not club_name:
        return jsonify({'error': '未关联社团'}), 400
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT cc.id, cc.user_id, cc.real_name, cc.student_id_num, cc.created_at FROM club_cadres cc WHERE cc.club_name=? ORDER BY cc.created_at DESC', (club_name,)).fetchall()
        cadres = [{'id': r['id'], 'user_id': r['user_id'], 'real_name': r['real_name'], 'student_id_num': r['student_id_num'], 'created_at': r['created_at']} for r in rows]
    finally:
        conn.close()
    return jsonify({'success': True, 'data': cadres})


@app.route('/api/club-cadres', methods=['POST'])
def add_club_cadre():
    user = get_current_user()
    if not user or user['role'] != 'user':
        return jsonify({'error': '仅社团负责人可操作'}), 403
    club_name = user.get('club_name') or ''
    if not club_name:
        return jsonify({'error': '未关联社团'}), 400
    data = request.get_json() or {}
    real_name = (data.get('real_name') or '').strip()
    student_id_num = (data.get('student_id_num') or '').strip()
    if not real_name or not student_id_num:
        return jsonify({'error': '请填写姓名和学号'}), 400
    conn = db.get_conn()
    try:
        member = conn.execute('SELECT cm.user_id, cm.real_name, cm.student_id_num FROM club_members cm WHERE cm.club_name=? AND cm.student_id_num=?', (club_name, student_id_num)).fetchone()
        if not member:
            member = conn.execute('SELECT cm.user_id, cm.real_name, cm.student_id_num FROM club_members cm WHERE cm.club_name=? AND cm.real_name=?', (club_name, real_name)).fetchone()
        if not member and student_id_num:
            profile_row = conn.execute('SELECT up.user_id, up.real_name, up.student_id FROM user_profiles up JOIN club_members cm ON up.user_id=cm.user_id WHERE cm.club_name=? AND up.student_id=?', (club_name, student_id_num)).fetchone()
            if profile_row:
                member = {'user_id': profile_row['user_id'], 'real_name': profile_row['real_name'] or real_name, 'student_id_num': profile_row['student_id'] or student_id_num}
        if not member and real_name:
            profile_row = conn.execute('SELECT up.user_id, up.real_name, up.student_id FROM user_profiles up JOIN club_members cm ON up.user_id=cm.user_id WHERE cm.club_name=? AND up.real_name=?', (club_name, real_name)).fetchone()
            if profile_row:
                member = {'user_id': profile_row['user_id'], 'real_name': profile_row['real_name'] or real_name, 'student_id_num': profile_row['student_id'] or student_id_num}
        if not member:
            return jsonify({'error': '该成员不在本社团中，请确认姓名和学号'}), 400
        if not member['user_id'] or member['user_id'] == 0:
            return jsonify({'error': '该成员未注册账号，无法设为骨干'}), 400
        existing = conn.execute('SELECT id FROM club_cadres WHERE club_name=? AND user_id=?', (club_name, member['user_id'])).fetchone()
        if existing:
            return jsonify({'error': '该成员已是社团骨干'}), 400
        target_user = conn.execute('SELECT role FROM users WHERE id=?', (member['user_id'],)).fetchone()
        if target_user and target_user['role'] == 'user':
            return jsonify({'error': '该成员是社团负责人，无需设为骨干'}), 400
        conn.execute('INSERT INTO club_cadres (club_name, user_id, real_name, student_id_num, added_by) VALUES (?, ?, ?, ?, ?)', (club_name, member['user_id'], member['real_name'] or real_name, member['student_id_num'] or student_id_num, user['id']))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '添加成功'})


@app.route('/api/club-cadres/<int:cadre_id>', methods=['DELETE'])
def delete_club_cadre(cadre_id):
    user = get_current_user()
    if not user or user['role'] != 'user':
        return jsonify({'error': '仅社团负责人可操作'}), 403
    club_name = user.get('club_name') or ''
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT club_name FROM club_cadres WHERE id=?', (cadre_id,)).fetchone()
        if not row:
            return jsonify({'error': '记录不存在'}), 404
        if row['club_name'] != club_name:
            return jsonify({'error': '无权操作'}), 403
        conn.execute('DELETE FROM club_cadres WHERE id=?', (cadre_id,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/my-cadres')
def my_cadres():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT cc.club_name, cc.real_name FROM club_cadres cc WHERE cc.user_id=?', (user['id'],)).fetchall()
        cadres = [{'club_name': r['club_name'], 'real_name': r['real_name']} for r in rows]
    finally:
        conn.close()
    return jsonify({'success': True, 'data': cadres})


@app.route('/api/my-profile')
def get_my_profile():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT * FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
    finally:
        conn.close()
    if row:
        return jsonify({'success': True, 'data': {
            'realName': row['real_name'], 'studentId': row['student_id'],
            'grade': row['grade'], 'className': row['class_name'],
            'college': row['college'] if 'college' in row.keys() else '',
            'phone': row['phone'], 'email': row['email'], 'avatarPath': row['avatar_path']
        }})
    return jsonify({'success': True, 'data': {
        'realName': '', 'studentId': '', 'grade': '', 'className': '', 'college': '',
        'phone': '', 'email': '', 'avatarPath': ''
    }})


@app.route('/api/my-profile', methods=['POST'])
def save_my_profile():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    real_name = data.get('realName', '').strip()
    student_id = data.get('studentId', '').strip()
    grade = data.get('grade', '').strip()
    class_name = data.get('className', '').strip()
    college = data.get('college', '').strip()
    phone = data.get('phone', '').strip()
    confirm_pass = data.get('confirmPassword', '').strip()
    email = data.get('email', '').strip()
    # 必填校验
    if not real_name:
        return jsonify({'error': '姓名不能为空'}), 400
    if not phone:
        return jsonify({'error': '联系电话不能为空'}), 400
    if not email:
        return jsonify({'error': '邮箱不能为空'}), 400
    if user['role'] == 'teacher':
        if not student_id:
            return jsonify({'error': '工号不能为空'}), 400
    if user['role'] in ('user', 'student'):
        if not student_id:
            return jsonify({'error': '学号不能为空'}), 400
        if not grade:
            return jsonify({'error': '年级不能为空'}), 400
        if not college:
            return jsonify({'error': '学院不能为空'}), 400
        if not class_name:
            return jsonify({'error': '班级不能为空'}), 400
    current_phone = ''
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
        current_profile = conn.execute('SELECT phone FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone() if existing else None
        current_phone = current_profile['phone'] if current_profile else ''
        if phone and phone != current_phone and current_phone:
            if not confirm_pass:
                conn.close()
                return jsonify({'error': '修改手机号需要输入当前密码确认', 'needPassword': True}), 400
            pass_row = conn.execute('SELECT password FROM users WHERE id=?', (user['id'],)).fetchone()
            if not pass_row or pass_row['password'] != confirm_pass:
                conn.close()
                return jsonify({'error': '密码错误，手机号未修改'}), 400
        if existing:
            conn.execute('''UPDATE user_profiles SET real_name=?, student_id=?, grade=?, class_name=?, college=?, phone=?, email=?, updated_at=CURRENT_TIMESTAMP WHERE user_id=?''',
                        (real_name, student_id, grade, class_name, college, phone, email, user['id']))
        else:
            conn.execute('''INSERT INTO user_profiles (user_id, real_name, student_id, grade, class_name, college, phone, email) VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                        (user['id'], real_name, student_id, grade, class_name, college, phone, email))
        conn.commit()
    except Exception as e:
        conn.close()
        return jsonify({'error': '保存失败：' + str(e)}), 500
    conn.close()
    result = {'success': True}
    if phone and current_phone and phone != current_phone:
        result['phoneChanged'] = True
    return jsonify(result)


@app.route('/api/upload-avatar', methods=['POST'])
def upload_avatar():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    if 'avatar' not in request.files:
        return jsonify({'error': '未上传文件'}), 400
    file = request.files['avatar']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    allowed_ext = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else 'png'
    if ext not in allowed_ext:
        return jsonify({'error': '仅支持 png/jpg/jpeg/gif/webp 格式'}), 400
    avatar_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'avatars')
    os.makedirs(avatar_dir, exist_ok=True)
    filename = f"avatar_{user['id']}_{int(time.time())}.{ext}"
    filepath = os.path.join(avatar_dir, filename)
    file.save(filepath)
    rel_path = f"/api/avatar-file/{filename}"
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
        if existing:
            conn.execute('UPDATE user_profiles SET avatar_path=?, updated_at=CURRENT_TIMESTAMP WHERE user_id=?', (rel_path, user['id']))
        else:
            conn.execute('INSERT INTO user_profiles (user_id, avatar_path) VALUES (?, ?)', (user['id'], rel_path))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'path': rel_path})


@app.route('/api/avatar-file/<filename>')
def serve_avatar(filename):
    avatar_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'avatars')
    return send_from_directory(avatar_dir, filename)


@app.route('/api/send-sms-code', methods=['POST'])
def send_sms_code():
    data = request.json or {}
    phone = data.get('phone', '').strip()
    if not phone or len(phone) < 11:
        return jsonify({'error': '请输入正确的手机号'}), 400
    import random, time
    code = str(random.randint(100000, 999999))
    now = time.time()
    expires_at = datetime.fromtimestamp(now + 300).strftime('%Y-%m-%d %H:%M:%S')
    conn = db.get_conn()
    try:
        recent = conn.execute('SELECT id FROM sms_codes WHERE phone=? AND created_at > datetime("now","-1 minute")', (phone,)).fetchone()
        if recent:
            return jsonify({'error': '发送太频繁，请1分钟后再试'}), 429
        conn.execute('INSERT INTO sms_codes (phone, code, expires_at) VALUES (?, ?, ?)', (phone, code, expires_at))
        conn.commit()
    finally:
        conn.close()
    print(f'[SMS] 验证码: {code} -> 手机: {phone}')
    return jsonify({'success': True, 'message': '验证码已发送'})

@app.route('/api/verify-sms-code', methods=['POST'])
def verify_sms_code():
    data = request.json or {}
    phone = data.get('phone', '').strip()
    code = data.get('code', '').strip()
    if not phone or not code:
        return jsonify({'error': '请输入手机号和验证码'}), 400
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT id, expires_at, used FROM sms_codes WHERE phone=? AND code=? ORDER BY id DESC LIMIT 1', (phone, code)).fetchone()
        if not row:
            return jsonify({'error': '验证码错误'}), 400
        if row['used']:
            return jsonify({'error': '验证码已使用'}), 400
        from datetime import datetime as dt
        try:
            expires = dt.strptime(row['expires_at'][:19], '%Y-%m-%d %H:%M:%S')
            if dt.now() > expires:
                return jsonify({'error': '验证码已过期'}), 400
        except:
            pass
        conn.execute('UPDATE sms_codes SET used=1 WHERE id=?', (row['id'],))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/send-email-code', methods=['POST'])
def send_email_code():
    data = request.json or {}
    email = data.get('email', '').strip()
    if not email or '@' not in email:
        return jsonify({'error': '请输入正确的邮箱地址'}), 400
    # 检查该邮箱是否已注册
    conn = db.get_conn()
    try:
        user_row = conn.execute('SELECT id FROM user_profiles WHERE email=? UNION SELECT id FROM teacher_profiles WHERE email=?', (email, email)).fetchone()
        if not user_row:
            return jsonify({'error': '该邮箱未注册'}), 400
        recent = conn.execute('SELECT id FROM email_codes WHERE email=? AND created_at > datetime("now","-1 minute")', (email,)).fetchone()
        if recent:
            return jsonify({'error': '发送太频繁，请1分钟后再试'}), 429
        code = str(random.randint(100000, 999999))
        now = time.time()
        expires_at = datetime.fromtimestamp(now + 300).strftime('%Y-%m-%d %H:%M:%S')
        conn.execute('INSERT INTO email_codes (email, code, expires_at) VALUES (?, ?, ?)', (email, code, expires_at))
        conn.commit()
    finally:
        conn.close()
    # 尝试发送邮件，如果未配置SMTP则打印到控制台
    smtp_host = os.environ.get('SMTP_HOST', '')
    smtp_port = int(os.environ.get('SMTP_PORT', '465'))
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_pass = os.environ.get('SMTP_PASS', '')
    smtp_ssl = os.environ.get('SMTP_SSL', 'true').lower() == 'true'
    if smtp_host and smtp_user and smtp_pass:
        try:
            import smtplib
            from email.mime.text import MIMEText
            msg = MIMEText(f'您的验证码为：{code}，5分钟内有效。如非本人操作请忽略。', 'plain', 'utf-8')
            msg['Subject'] = '密码重置验证码'
            msg['From'] = smtp_user
            msg['To'] = email
            if smtp_ssl:
                server = smtplib.SMTP_SSL(smtp_host, smtp_port)
            else:
                server = smtplib.SMTP(smtp_host, smtp_port)
                server.starttls()
            server.login(smtp_user, smtp_pass)
            server.sendmail(smtp_user, [email], msg.as_string())
            server.quit()
        except Exception as e:
            print(f'[EMAIL] 发送失败: {e}, 验证码: {code} -> 邮箱: {email}')
    else:
        print(f'[EMAIL] 验证码: {code} -> 邮箱: {email}')
    return jsonify({'success': True, 'message': '验证码已发送'})


@app.route('/api/verify-email-code', methods=['POST'])
def verify_email_code():
    data = request.json or {}
    email = data.get('email', '').strip()
    code = data.get('code', '').strip()
    if not email or not code:
        return jsonify({'error': '请输入邮箱和验证码'}), 400
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT id, expires_at, used FROM email_codes WHERE email=? AND code=? ORDER BY id DESC LIMIT 1', (email, code)).fetchone()
        if not row:
            return jsonify({'error': '验证码错误'}), 400
        if row['used']:
            return jsonify({'error': '验证码已使用'}), 400
        from datetime import datetime as dt
        try:
            expires = dt.strptime(row['expires_at'][:19], '%Y-%m-%d %H:%M:%S')
            if dt.now() > expires:
                return jsonify({'error': '验证码已过期'}), 400
        except:
            pass
        conn.execute('UPDATE email_codes SET used=1 WHERE id=?', (row['id'],))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/change-password', methods=['POST'])
def change_password():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    old_pass = data.get('oldPassword', '').strip()
    new_pass = data.get('newPassword', '').strip()
    if not old_pass or not new_pass:
        return jsonify({'error': '请输入当前密码和新密码'}), 400
    if len(new_pass) < 4:
        return jsonify({'error': '新密码至少4位'}), 400
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT password FROM users WHERE id=?', (user['id'],)).fetchone()
        if not row or row['password'] != old_pass:
            return jsonify({'error': '当前密码错误'}), 400
        conn.execute('UPDATE users SET password=? WHERE id=?', (new_pass, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-members/<club_name>')
def get_club_members(club_name):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT cm.* FROM club_members cm LEFT JOIN users u ON cm.user_id=u.id WHERE cm.club_name=? AND (cm.user_id=0 OR u.id IS NOT NULL) ORDER BY cm.joined_at DESC', (club_name,)).fetchall()
        result = []
        for r in rows:
            profile = {}
            if r['user_id']:
                try:
                    prow = conn.execute('SELECT real_name, student_id, grade, class_name, phone, email FROM user_profiles WHERE user_id=?', (r['user_id'],)).fetchone()
                    if prow:
                        profile = {'realName': prow['real_name'], 'studentId': prow['student_id'], 'grade': prow['grade'], 'className': prow['class_name'], 'phone': prow['phone'], 'email': prow['email']}
                except:
                    pass
            result.append({'id': r['id'], 'userId': r['user_id'], 'username': r['username'], 'realName': r['real_name'],
                'studentIdNum': r['student_id_num'], 'className': r['class_name'], 'phone': r['phone'],
                'department': r['department'], 'specialty': r['specialty'], 'college': (r['college'] if 'college' in r.keys() else '') or '', 'source': r['source'],
                'joinedAt': r['joined_at'], 'profile': profile})
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/club-members/<int:mid>', methods=['DELETE'])
def remove_club_member(mid):
    user = get_current_user()
    print(f'[DEBUG] remove_club_member: mid={mid}, user={user}, session={dict(session)}')
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn_chk = db.get_conn()
    try:
        mem = conn_chk.execute('SELECT club_name FROM club_members WHERE id=?', (mid,)).fetchone()
    finally:
        conn_chk.close()
    if not mem:
        return jsonify({'error': '成员不存在'}), 404
    if user['role'] == 'student':
        if not is_cadre_of_club(user['id'], mem['club_name']):
            return jsonify({'error': '无权限操作'}), 403
    elif user['role'] == 'teacher':
        tc_conn = db.get_conn()
        try:
            tc_row = tc_conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], mem['club_name'])).fetchone()
        finally:
            tc_conn.close()
        if not tc_row:
            return jsonify({'error': '无权限操作'}), 403
    elif user['role'] not in ('user', 'admin'):
        return jsonify({'error': '无权限操作'}), 403
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM club_members WHERE id=?', (mid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-members/batch-remove', methods=['POST'])
def batch_remove_club_members():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    ids = data.get('ids', [])
    if not ids:
        return jsonify({'error': '未选择成员'}), 400
    conn_chk = db.get_conn()
    try:
        placeholders = ','.join(['?'] * len(ids))
        rows = conn_chk.execute(f'SELECT DISTINCT club_name FROM club_members WHERE id IN ({placeholders})', ids).fetchall()
    finally:
        conn_chk.close()
    club_names = [r['club_name'] for r in rows if r['club_name']]
    if len(club_names) > 1:
        return jsonify({'error': '批量移除的成员必须属于同一社团'}), 400
    target_club = club_names[0] if club_names else ''
    if user['role'] == 'student':
        cadre_clubs = get_cadre_clubs(user['id'])
        if not cadre_clubs or target_club not in cadre_clubs:
            return jsonify({'error': '无权限操作'}), 403
    elif user['role'] == 'teacher':
        tc_conn = db.get_conn()
        try:
            tc_row = tc_conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], target_club)).fetchone()
        finally:
            tc_conn.close()
        if not tc_row:
            return jsonify({'error': '无权限操作'}), 403
    elif user['role'] not in ('user', 'admin'):
        return jsonify({'error': '无权限操作'}), 403
    conn = db.get_conn()
    try:
        placeholders = ','.join(['?'] * len(ids))
        conn.execute(f'DELETE FROM club_members WHERE id IN ({placeholders})', ids)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'removed': len(ids)})


@app.route('/api/member-analysis/<club_name>')
def member_analysis(club_name):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        members = conn.execute('SELECT cm.* FROM club_members cm LEFT JOIN users u ON cm.user_id=u.id WHERE cm.club_name=? AND (cm.user_id=0 OR u.id IS NOT NULL)', (club_name,)).fetchall()
        depts = conn.execute('SELECT dept_name FROM club_departments WHERE club_name=?', (club_name,)).fetchall()
    finally:
        conn.close()
    dept_stats = {}
    for d in depts:
        dept_stats[d['dept_name']] = 0
    class_stats = {}
    college_stats = {}
    specialty_stats = {}
    for m in members:
        dept = (m['department'] or '').strip() or '未分配'
        dept_stats[dept] = dept_stats.get(dept, 0) + 1
        cls = (m['class_name'] or '').strip() or '未知'
        class_stats[cls] = class_stats.get(cls, 0) + 1
        college = (m['college'] or '').strip() or '未知'
        college_stats[college] = college_stats.get(college, 0) + 1
        sp = (m['specialty'] or '').strip()
        if sp:
            specialty_stats[sp] = specialty_stats.get(sp, 0) + 1
    return jsonify({'success': True, 'data': {
        'total': len(members), 'deptCount': len(depts),
        'deptStats': dept_stats, 'classStats': class_stats, 'collegeStats': college_stats, 'specialtyStats': specialty_stats
    }})


@app.route('/api/club-notices', methods=['GET', 'POST'])
def handle_club_notices():
    if request.method == 'GET':
        user = get_current_user()
        if not user:
            return jsonify({'error': '未登录'}), 401
        club = request.args.get('club', '').strip()
        if not club and user['role'] == 'user':
            club = user['club_name'] or ''
        if not club:
            conn = db.get_conn()
            try:
                rows = conn.execute('SELECT * FROM club_notices ORDER BY created_at DESC').fetchall()
            finally:
                conn.close()
            return jsonify({'success': True, 'data': [{'id': r['id'], 'clubName': r['club_name'], 'title': r['title'], 'content': r['content'], 'createdBy': r['created_by'], 'time': local_time(r['created_at'])} for r in rows]})
        conn = db.get_conn()
        try:
            rows = conn.execute('SELECT * FROM club_notices WHERE club_name=? ORDER BY created_at DESC', (club,)).fetchall()
        finally:
            conn.close()
        return jsonify({'success': True, 'data': [{'id': r['id'], 'clubName': r['club_name'], 'title': r['title'], 'content': r['content'], 'createdBy': r['created_by'], 'time': local_time(r['created_at'])} for r in rows]})
    else:
        user = get_current_user()
        if not user or user['role'] not in ('user', 'admin', 'teacher'):
            return jsonify({'error': '请先登录'}), 401
        data = request.json or {}
        club = data.get('clubName', '').strip()
        if not club and user['role'] == 'user':
            club = user['club_name'] or ''
        if not club:
            return jsonify({'error': '无社团信息'}), 400
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        if not title:
            return jsonify({'error': '请输入标题'}), 400
        conn = db.get_conn()
        try:
            cursor = conn.execute('INSERT INTO club_notices (club_name, title, content, created_by) VALUES (?, ?, ?, ?)', (club, title, content, user['username']))
            notice_id = cursor.lastrowid
            members = conn.execute('SELECT user_id FROM club_members WHERE club_name=? AND user_id!=0', (club,)).fetchall()
            for m in members:
                send_notification(m['user_id'], '📢 社团通知', f'「{club}」发布了新通知：{title}', 'club_notice', f'club_notice:{notice_id}', conn=conn)
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True})

@app.route('/api/excellent-activity-detail/<group_id>')
def excellent_activity_detail(group_id):
    """公开API：获取优秀活动详情，无需登录"""
    conn = db.get_conn()
    try:
        if group_id.startswith('session_'):
            sid = group_id.replace('session_', '')
            try:
                sid_int = int(sid)
            except ValueError:
                return jsonify({'error': '无效的活动ID'}), 400
            sess = conn.execute('SELECT cs.*, COUNT(cr.id) as checkin_count FROM checkin_sessions cs LEFT JOIN checkin_records cr ON cs.id=cr.session_id WHERE cs.id=? GROUP BY cs.id', (sid_int,)).fetchone()
            if not sess:
                return jsonify({'error': '活动不存在'}), 404
            photos = conn.execute('SELECT id, club_token FROM club_uploads WHERE club_name=? AND source="activity" AND group_id=?', (sess['club_name'], group_id)).fetchall()
            photo_list = []
            token_row = conn.execute('SELECT token FROM club_tokens WHERE club_name=?', (sess['club_name'],)).fetchone()
            tk = token_row['token'] if token_row else ''
            for p in photos:
                photo_list.append({'id': p['id'], 'url': f'/api/club-download/{tk}/{p["id"]}'})
            return jsonify({'success': True, 'data': {
                'groupId': group_id,
                'type': 'session',
                'clubName': sess['club_name'],
                'title': sess['activity_name'] or sess['club_name'] + '的活动',
                'time': local_time(sess['created_at']),
                'activityName': sess['activity_name'],
                'locationName': sess['location_name'] if 'location_name' in sess.keys() else '',
                'activityContent': sess['activity_content'] if 'activity_content' in sess.keys() else '',
                'startTime': sess['start_time'] if 'start_time' in sess.keys() else '',
                'endTime': sess['end_time'] if 'end_time' in sess.keys() else '',
                'checkinCount': sess['checkin_count'],
                'photos': photo_list
            }})
        else:
            rows = conn.execute('SELECT id, club_name, file_name, file_type, description, upload_time, status, group_id, club_token, category FROM club_uploads WHERE group_id=? AND status="approved"', (group_id,)).fetchall()
            if not rows:
                return jsonify({'error': '活动不存在'}), 404
            title, body, photos = '', '', []
            tk = rows[0]['club_token']
            for r in rows:
                desc = r['description'] or ''
                if desc.startswith('[总结]'):
                    parts = desc.replace('[总结]', '').split('|||')
                    title = parts[0] if parts else ''
                    body = parts[1] if len(parts) > 1 else ''
                elif desc.startswith('[照片]'):
                    ft = (r['file_type'] or '').lower()
                    if ft.startswith('image') or ft in ('jpg', 'jpeg', 'png', 'gif', 'webp'):
                        photos.append({'id': r['id'], 'url': f'/api/club-download/{tk}/{r["id"]}'})
            return jsonify({'success': True, 'data': {
                'groupId': group_id,
                'type': 'material',
                'clubName': rows[0]['club_name'],
                'title': title or rows[0]['club_name'] + '的活动',
                'time': local_time(rows[0]['upload_time']),
                'body': body,
                'photos': photos
            }})
    finally:
        conn.close()


@app.route('/api/club-notices/<int:nid>', methods=['DELETE'])
def delete_club_notice(nid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM club_notices WHERE id=?', (nid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-notices/batch-delete', methods=['POST'])
def batch_delete_club_notices():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    ids = data.get('ids', [])
    if not ids:
        return jsonify({'error': '请选择要删除的通知'}), 400
    conn = db.get_conn()
    try:
        placeholders = ','.join(['?' for _ in ids])
        conn.execute(f'DELETE FROM club_notices WHERE id IN ({placeholders})', ids)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'deleted': len(ids)})


@app.route('/api/my-club-notices')
def get_my_club_notices():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        if user['role'] == 'student':
            clubs = conn.execute('SELECT DISTINCT club_name FROM club_members WHERE user_id=?', (user['id'],)).fetchall()
            club_names = [c['club_name'] for c in clubs]
        elif user['role'] == 'user':
            club_names = [user['club_name']] if user['club_name'] else []
        else:
            club_names = []
        if not club_names:
            return jsonify({'success': True, 'data': []})
        placeholders = ','.join(['?'] * len(club_names))
        rows = conn.execute(f'SELECT c.*, n.id as notif_id, n.is_read FROM club_notices c LEFT JOIN notifications n ON n.user_id=? AND n.type=\'club_notice\' AND n.link=\'club_notice:\'||CAST(c.id AS TEXT) WHERE c.club_name IN ({placeholders}) AND COALESCE(n.cleared,0)=0 ORDER BY c.created_at DESC', [user['id']] + club_names).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'clubName': r['club_name'], 'title': r['title'], 'content': r['content'], 'createdBy': r['created_by'], 'time': local_time(r['created_at']), 'notifId': r['notif_id'], 'isRead': bool(r['is_read']) if r['notif_id'] else False} for r in rows]})


@app.route('/api/club-notice/<int:nid>/read', methods=['POST'])
def mark_club_notice_read(nid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('UPDATE notifications SET is_read=1 WHERE id=? AND user_id=? AND type=\'club_notice\'', (nid, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-notice/clear-read', methods=['DELETE'])
def clear_read_club_notices():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('UPDATE notifications SET cleared=1 WHERE user_id=? AND type=\'club_notice\' AND is_read=1', (user['id'],))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-notice/batch-delete', methods=['POST'])
def batch_delete_club_notice_reads():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    ids = data.get('ids', [])
    if not ids:
        return jsonify({'error': '请选择要删除的通知'}), 400
    conn = db.get_conn()
    try:
        placeholders = ','.join(['?' for _ in ids])
        conn.execute(f'UPDATE notifications SET cleared=1 WHERE id IN ({placeholders}) AND user_id=? AND type=\'club_notice\'', ids + [user['id']])
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'deleted': len(ids)})


@app.route('/api/notifications')
def get_notifications():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    ntype = request.args.get('type', '')
    conn = db.get_conn()
    try:
        admin_exclude = " AND type NOT IN ('upload','offcampus','recruit','joint','feedback')" if user['role'] == 'admin' else ""
        user_exclude = " AND type NOT IN ('upload','approve','reject','recruit','joint')" if user['role'] == 'user' else ""
        student_exclude = " AND type NOT IN ('recruit')" if user['role'] == 'student' else ""
        cleared_exclude = " AND COALESCE(cleared,0)=0"
        if ntype:
            notifs = conn.execute(f'SELECT * FROM notifications WHERE user_id=? AND type=?{admin_exclude}{user_exclude}{student_exclude}{cleared_exclude} ORDER BY created_at DESC LIMIT 50', (user['id'], ntype)).fetchall()
            unread = conn.execute(f'SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND is_read=0 AND type=?{admin_exclude}{user_exclude}{student_exclude}{cleared_exclude}', (user['id'], ntype)).fetchone()['c']
        else:
            notifs = conn.execute(f'SELECT * FROM notifications WHERE user_id=?{admin_exclude}{user_exclude}{student_exclude}{cleared_exclude} ORDER BY created_at DESC LIMIT 50', (user['id'],)).fetchall()
            unread = conn.execute(f'SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND is_read=0{admin_exclude}{user_exclude}{student_exclude}{cleared_exclude}', (user['id'],)).fetchone()['c']
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'unread': unread,
        'notifications': [{'id': n['id'], 'title': n['title'], 'content': n['content'], 'type': n['type'], 'isRead': n['is_read'], 'link': n['link'], 'time': n['created_at']} for n in notifs]
    })


@app.route('/api/notifications/<int:nid>/read', methods=['POST'])
def mark_notification_read(nid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('UPDATE notifications SET is_read=1 WHERE id=? AND user_id=?', (nid, user['id']))
        conn.commit()
    finally:
        conn.close()


@app.route('/api/material-view-group/<group_id>', methods=['POST'])
def mark_material_group_viewed(group_id):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('INSERT OR IGNORE INTO material_views (group_id, user_id) VALUES (?, ?)', (group_id, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/material-views-batch', methods=['POST'])
def mark_materials_viewed_batch():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    group_ids = data.get('group_ids', [])
    if not group_ids:
        return jsonify({'success': True})
    conn = db.get_conn()
    try:
        for gid in group_ids:
            conn.execute('INSERT OR IGNORE INTO material_views (group_id, user_id) VALUES (?, ?)', (gid, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/feature-read/<feature>', methods=['POST'])
def mark_feature_read(feature):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    allowed = ['enroll', 'recruit', 'joint', 'volunteer', 'member', 'checkin', 'adminNotify', 'quitManage', 'scoring', 'scoringRules', 'scoringReview']
    if feature not in allowed:
        return jsonify({'error': '无效的功能'}), 400
    conn = db.get_conn()
    try:
        conn.execute('INSERT OR REPLACE INTO feature_reads (user_id, feature, read_at) VALUES (?, ?, datetime("now"))', (user['id'], feature))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})

@app.route('/api/nav-badges', methods=['GET'])
def nav_badges():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    badges = {}
    conn = db.get_conn()
    try:
        fr_rows = conn.execute('SELECT feature, read_at FROM feature_reads WHERE user_id=?', (user['id'],)).fetchall()
        feature_reads = {r['feature']: r['read_at'] for r in fr_rows}
        notif_unread = conn.execute('SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND is_read=0 AND COALESCE(cleared,0)=0' + (" AND type NOT IN ('upload','offcampus','recruit','joint','feedback')" if user['role'] == 'admin' else '') + (" AND type NOT IN ('upload','approve','reject','recruit','joint')" if user['role'] == 'user' else '') + (" AND type NOT IN ('recruit')" if user['role'] == 'student' else ''), (user['id'],)).fetchone()['c']
        badges['notif'] = notif_unread
        if user['role'] == 'admin':
            pending_offcampus = conn.execute("SELECT COUNT(*) as c FROM offcampus_requests WHERE status='pending'").fetchone()['c']
            badges['offcampus'] = pending_offcampus
            pending_feedback = conn.execute("SELECT COUNT(*) as c FROM feedbacks WHERE status='pending'").fetchone()['c']
            badges['feedback'] = pending_feedback
            pending_materials = conn.execute("SELECT COUNT(DISTINCT COALESCE(group_id, CAST(id AS TEXT))) as c FROM club_uploads WHERE status='pending' AND (source='upload' OR source='activity' OR source IS NULL)").fetchone()['c']
            badges['materials'] = pending_materials
            pending_scoring = conn.execute("SELECT COUNT(*) as c FROM scoring_submissions WHERE status='submitted_tuanwei'").fetchone()['c']
            badges['scoringRules'] = pending_scoring
            recruit_read_at = feature_reads.get('recruit')
            if recruit_read_at:
                pending_recruit = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE status='pending' AND created_at>?", (recruit_read_at,)).fetchone()['c']
            else:
                pending_recruit = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE status='pending'").fetchone()['c']
            badges['recruit'] = pending_recruit
            joint_read_at = feature_reads.get('joint')
            if joint_read_at:
                new_joints = conn.execute("SELECT COUNT(*) as c FROM joint_activities WHERE created_at>? AND club_name!=?", (joint_read_at, user.get('club_name') or '')).fetchone()['c']
                new_cooperate = conn.execute("SELECT COUNT(*) as c FROM joint_replies WHERE reply_type='cooperate' AND created_at>? AND activity_id IN (SELECT id FROM joint_activities WHERE status='open')", (joint_read_at,)).fetchone()['c']
                badges['joint'] = new_joints + new_cooperate
            else:
                joint_cooperate = conn.execute("SELECT COUNT(*) as c FROM joint_replies WHERE reply_type='cooperate' AND activity_id IN (SELECT id FROM joint_activities WHERE status='open')").fetchone()['c']
                new_joints_all = conn.execute("SELECT COUNT(*) as c FROM joint_activities WHERE club_name!=?", (user.get('club_name') or '',)).fetchone()['c']
                badges['joint'] = joint_cooperate + new_joints_all
            quit_read_at = feature_reads.get('quitManage')
            if quit_read_at:
                pending_quit = conn.execute("SELECT COUNT(*) as c FROM quit_applications WHERE status='pending' AND created_at>?", (quit_read_at,)).fetchone()['c']
            else:
                pending_quit = conn.execute("SELECT COUNT(*) as c FROM quit_applications WHERE status='pending'").fetchone()['c']
            badges['quitManage'] = pending_quit
        elif user['role'] == 'user':
            club_name = user.get('club_name') or ''
            if club_name:
                pending_workload = conn.execute("SELECT COUNT(*) as c FROM workload_submissions WHERE club_name=? AND status='pending'", (club_name,)).fetchone()['c']
                badges['workload'] = pending_workload
                recruit_read_at = feature_reads.get('recruit')
                if recruit_read_at:
                    recruit_reviewed = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE club_name=? AND status='approved' AND created_at>? AND (deadline='' OR deadline IS NULL OR datetime(deadline)>datetime('now','localtime'))", (club_name, recruit_read_at)).fetchone()['c']
                else:
                    recruit_reviewed = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE club_name=? AND status='approved' AND (deadline='' OR deadline IS NULL OR datetime(deadline)>datetime('now','localtime'))", (club_name,)).fetchone()['c']
                badges['recruit'] = recruit_reviewed
                pending_materials = conn.execute("SELECT COUNT(DISTINCT COALESCE(group_id, CAST(id AS TEXT))) as c FROM club_uploads WHERE club_name=? AND status IN ('approved','rejected') AND COALESCE(group_id, CAST(id AS TEXT)) NOT IN (SELECT group_id FROM material_views WHERE user_id=?)", (club_name, user['id'])).fetchone()['c']
                badges['materials'] = pending_materials
                joint_read_at = feature_reads.get('joint')
                if joint_read_at:
                    new_joints = conn.execute("SELECT COUNT(*) as c FROM joint_activities WHERE created_at>? AND club_name!=?", (joint_read_at, club_name)).fetchone()['c']
                    new_replies = conn.execute("SELECT COUNT(*) as c FROM joint_replies WHERE activity_id IN (SELECT id FROM joint_activities WHERE club_name=? AND status='open') AND created_at>?", (club_name, joint_read_at)).fetchone()['c']
                    badges['joint'] = new_joints + new_replies
                else:
                    joint_replies = conn.execute("SELECT COUNT(*) as c FROM joint_replies WHERE activity_id IN (SELECT id FROM joint_activities WHERE club_name=? AND status='open')", (club_name,)).fetchone()['c']
                    new_joints_all = conn.execute("SELECT COUNT(*) as c FROM joint_activities WHERE club_name!=?", (club_name,)).fetchone()['c']
                    badges['joint'] = joint_replies + new_joints_all
            club = user.get('club_name') or ''
            admin_notify_read_at = feature_reads.get('adminNotify')
            if admin_notify_read_at:
                total_admin_notifs = conn.execute("SELECT COUNT(*) as c FROM admin_notifications WHERE created_at>? AND (target_type=? OR (target_type=? AND (target_club=? OR ','||target_club||',' LIKE '%,'||?||',%')))", (admin_notify_read_at, 'all_leaders', 'specific_leader', club, club)).fetchone()['c']
                badges['adminNotify'] = max(0, total_admin_notifs)
            else:
                total_admin_notifs = conn.execute("SELECT COUNT(*) as c FROM admin_notifications WHERE target_type=? OR (target_type=? AND (target_club=? OR ','||target_club||',' LIKE '%,'||?||',%'))", ('all_leaders', 'specific_leader', club, club)).fetchone()['c']
                read_admin_notifs = conn.execute("SELECT COUNT(DISTINCT r.notification_id) as c FROM admin_notification_reads r JOIN admin_notifications n ON r.notification_id=n.id WHERE r.user_id=? AND (n.target_type=? OR (n.target_type=? AND (n.target_club=? OR ','||n.target_club||',' LIKE '%,'||?||',%')))", (user['id'], 'all_leaders', 'specific_leader', club, club)).fetchone()['c']
                badges['adminNotify'] = max(0, total_admin_notifs - read_admin_notifs)
            club_notice_unread = conn.execute("SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND is_read=0 AND type='club_notice' AND COALESCE(cleared,0)=0", (user['id'],)).fetchone()['c']
            badges['clubnotice'] = club_notice_unread
            if club_name:
                quit_read_at = feature_reads.get('quitManage')
                if quit_read_at:
                    pending_quit = conn.execute("SELECT COUNT(*) as c FROM quit_applications WHERE club_name=? AND status='pending' AND created_at>?", (club_name, quit_read_at)).fetchone()['c']
                else:
                    pending_quit = conn.execute("SELECT COUNT(*) as c FROM quit_applications WHERE club_name=? AND status='pending'", (club_name,)).fetchone()['c']
                badges['quitManage'] = pending_quit
                rejected_scoring = conn.execute("SELECT COUNT(*) as c FROM scoring_submissions WHERE club_name=? AND status='draft' AND id IN (SELECT submission_id FROM scoring_teacher_reviews WHERE status='rejected')", (club_name,)).fetchone()['c']
                approved_scoring = conn.execute("SELECT COUNT(*) as c FROM scoring_submissions WHERE club_name=? AND status='teacher_approved'", (club_name,)).fetchone()['c']
                scoring_badge = rejected_scoring + approved_scoring
                if scoring_badge > 0:
                    badges['scoring'] = scoring_badge
        elif user['role'] == 'student':
            admin_notify_read_at = feature_reads.get('adminNotify')
            if admin_notify_read_at:
                total_admin_notifs = conn.execute("SELECT COUNT(*) as c FROM admin_notifications WHERE created_at>? AND target_type=?", (admin_notify_read_at, 'all_students')).fetchone()['c']
                badges['adminNotify'] = max(0, total_admin_notifs)
            else:
                total_admin_notifs = conn.execute("SELECT COUNT(*) as c FROM admin_notifications WHERE target_type=?", ('all_students',)).fetchone()['c']
                read_admin_notifs = conn.execute("SELECT COUNT(DISTINCT r.notification_id) as c FROM admin_notification_reads r JOIN admin_notifications n ON r.notification_id=n.id WHERE r.user_id=? AND n.target_type=?", (user['id'], 'all_students')).fetchone()['c']
                badges['adminNotify'] = max(0, total_admin_notifs - read_admin_notifs)
            club_notice_unread = conn.execute("SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND is_read=0 AND type='club_notice' AND COALESCE(cleared,0)=0", (user['id'],)).fetchone()['c']
            badges['clubnotice'] = club_notice_unread
            enroll_read_at = feature_reads.get('enroll')
            if enroll_read_at:
                volunteer_count = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE status='approved' AND recruit_type='volunteer' AND created_at>? AND (deadline='' OR deadline IS NULL OR datetime(deadline)>datetime('now','localtime')) AND id NOT IN (SELECT recruitment_id FROM recruitment_signups WHERE user_id=?)", (enroll_read_at, user['id'])).fetchone()['c']
                member_count = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE status='approved' AND recruit_type='member' AND created_at>? AND (deadline='' OR deadline IS NULL OR datetime(deadline)>datetime('now','localtime')) AND id NOT IN (SELECT recruitment_id FROM recruitment_signups WHERE user_id=?)", (enroll_read_at, user['id'])).fetchone()['c']
            else:
                volunteer_count = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE status='approved' AND recruit_type='volunteer' AND (deadline='' OR deadline IS NULL OR datetime(deadline)>datetime('now','localtime')) AND id NOT IN (SELECT recruitment_id FROM recruitment_signups WHERE user_id=?)", (user['id'],)).fetchone()['c']
                member_count = conn.execute("SELECT COUNT(*) as c FROM recruitments WHERE status='approved' AND recruit_type='member' AND (deadline='' OR deadline IS NULL OR datetime(deadline)>datetime('now','localtime')) AND id NOT IN (SELECT recruitment_id FROM recruitment_signups WHERE user_id=?)", (user['id'],)).fetchone()['c']
            badges['volunteer'] = volunteer_count
            badges['member'] = member_count
            badges['enroll'] = volunteer_count + member_count
            checkin_read_at = feature_reads.get('checkin')
            if checkin_read_at:
                checkin_count = conn.execute("SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND type='checkin' AND is_read=0 AND COALESCE(cleared,0)=0 AND created_at>?", (user['id'], checkin_read_at)).fetchone()['c']
            else:
                checkin_count = conn.execute("SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND type='checkin' AND is_read=0 AND COALESCE(cleared,0)=0", (user['id'],)).fetchone()['c']
            badges['checkin'] = checkin_count
        elif user['role'] == 'teacher':
            teacher_name = user.get('username') or ''
            admin_notify_read_at = feature_reads.get('adminNotify')
            if admin_notify_read_at:
                total_admin_notifs = conn.execute("SELECT COUNT(*) as c FROM admin_notifications WHERE created_at>? AND (target_type=? OR (target_type=? AND (target_club=? OR ','||target_club||',' LIKE '%,'||?||',%')))", (admin_notify_read_at, 'all_teachers', 'specific_teacher', teacher_name, teacher_name)).fetchone()['c']
                badges['adminNotify'] = max(0, total_admin_notifs)
            else:
                total_admin_notifs = conn.execute("SELECT COUNT(*) as c FROM admin_notifications WHERE target_type=? OR (target_type=? AND (target_club=? OR ','||target_club||',' LIKE '%,'||?||',%'))", ('all_teachers', 'specific_teacher', teacher_name, teacher_name)).fetchone()['c']
                read_admin_notifs = conn.execute("SELECT COUNT(DISTINCT r.notification_id) as c FROM admin_notification_reads r JOIN admin_notifications n ON r.notification_id=n.id WHERE r.user_id=? AND (n.target_type=? OR (n.target_type=? AND (n.target_club=? OR ','||n.target_club||',' LIKE '%,'||?||',%')))", (user['id'], 'all_teachers', 'specific_teacher', teacher_name, teacher_name)).fetchone()['c']
                badges['adminNotify'] = max(0, total_admin_notifs - read_admin_notifs)
            try:
                teacher_club_list = [r['club_name'] for r in conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()]
                if not teacher_club_list:
                    teacher_club_list = [r['club_name'] for r in conn.execute('SELECT club_name FROM club_teachers WHERE user_id=?', (user['id'],)).fetchall()]
                if user.get('club_name') and user['club_name'] not in teacher_club_list:
                    teacher_club_list.append(user['club_name'])
                teacher_club_list = [c for c in teacher_club_list if c]
                if teacher_club_list:
                    placeholders_tc = ','.join(['?'] * len(teacher_club_list))
                    pending_scoring = conn.execute(f"SELECT COUNT(*) as c FROM scoring_submissions WHERE club_name IN ({placeholders_tc}) AND status='pending_teacher'", teacher_club_list).fetchone()['c']
                else:
                    pending_scoring = 0
                badges['scoring'] = pending_scoring
                badges['scoringReview'] = pending_scoring
            except Exception as e:
                badges['scoring'] = 0
                badges['scoringReview'] = 0
            teacher_clubs = [r['club_name'] for r in conn.execute('SELECT club_name FROM club_teachers WHERE user_id=?', (user['id'],)).fetchall()]
            if teacher_clubs:
                joint_read_at = feature_reads.get('joint')
                placeholders = ','.join(['?'] * len(teacher_clubs))
                if joint_read_at:
                    new_joints = conn.execute(f"SELECT COUNT(*) as c FROM joint_activities WHERE created_at>? AND club_name NOT IN ({placeholders})", [joint_read_at] + teacher_clubs).fetchone()['c']
                    new_replies = conn.execute(f"SELECT COUNT(*) as c FROM joint_replies WHERE activity_id IN (SELECT id FROM joint_activities WHERE club_name IN ({placeholders}) AND status='open') AND created_at>?", teacher_clubs + [joint_read_at]).fetchone()['c']
                    badges['joint'] = new_joints + new_replies
                else:
                    new_joints_all = conn.execute(f"SELECT COUNT(*) as c FROM joint_activities WHERE club_name NOT IN ({placeholders})", teacher_clubs).fetchone()['c']
                    joint_replies = conn.execute(f"SELECT COUNT(*) as c FROM joint_replies WHERE activity_id IN (SELECT id FROM joint_activities WHERE club_name IN ({placeholders}) AND status='open')", teacher_clubs).fetchone()['c']
                    badges['joint'] = new_joints_all + joint_replies
            teacher_quit_clubs = teacher_clubs if teacher_clubs else [user.get('club_name')] if user.get('club_name') else []
            if teacher_quit_clubs:
                placeholders = ','.join(['?'] * len(teacher_quit_clubs))
                pending_quit = conn.execute(f"SELECT COUNT(*) as c FROM quit_applications WHERE club_name IN ({placeholders}) AND status='pending'", teacher_quit_clubs).fetchone()['c']
                badges['quitManage'] = pending_quit
    finally:
        conn.close()
    return jsonify({'success': True, 'badges': badges})


@app.route('/api/notifications/<int:nid>', methods=['DELETE'])
def delete_notification(nid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM notifications WHERE id=? AND user_id=?', (nid, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/notifications/clear', methods=['DELETE'])
def clear_read_notifications():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM notifications WHERE user_id=? AND is_read=1', (user['id'],))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/notifications/batch-delete', methods=['POST'])
def batch_delete_notifications():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    ids = data.get('ids', [])
    if not ids:
        return jsonify({'error': '请选择要删除的通知'}), 400
    conn = db.get_conn()
    try:
        placeholders = ','.join(['?' for _ in ids])
        conn.execute(f'DELETE FROM notifications WHERE id IN ({placeholders}) AND user_id=?', ids + [user['id']])
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'deleted': len(ids)})


@app.route('/api/admin-notifications/batch-delete', methods=['POST'])
def batch_delete_admin_notifications():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可操作'}), 403
    data = request.json or {}
    ids = data.get('ids', [])
    if not ids:
        return jsonify({'error': '请选择要删除的通知'}), 400
    conn = db.get_conn()
    try:
        placeholders = ','.join(['?' for _ in ids])
        conn.execute(f'DELETE FROM admin_notification_reads WHERE notification_id IN ({placeholders})', ids)
        conn.execute(f'DELETE FROM admin_notifications WHERE id IN ({placeholders})', ids)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'deleted': len(ids)})


@app.route('/api/my-admin-notifications/clear', methods=['DELETE'])
def clear_read_admin_notifications():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    if user['role'] == 'admin':
        return jsonify({'error': '管理员不可清理'}), 403
    conn = db.get_conn()
    try:
        conn.execute('UPDATE admin_notification_reads SET cleared=1 WHERE user_id=? AND cleared=0', (user['id'],))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/my-admin-notifications/<int:nid>/read', methods=['DELETE'])
def delete_admin_notification_read(nid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM admin_notification_reads WHERE notification_id=? AND user_id=?', (nid, user['id'])).fetchone()
        if existing:
            conn.execute('UPDATE admin_notification_reads SET cleared=1 WHERE notification_id=? AND user_id=?', (nid, user['id']))
        else:
            conn.execute('INSERT OR IGNORE INTO admin_notification_reads (notification_id, user_id, cleared) VALUES (?, ?, 1)', (nid, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/notifications/batch-delete-admin-read', methods=['POST'])
def batch_delete_admin_notification_reads():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    ids = data.get('ids', [])
    if not ids:
        return jsonify({'error': '请选择要删除的通知'}), 400
    conn = db.get_conn()
    try:
        placeholders = ','.join(['?' for _ in ids])
        existing = conn.execute(f'SELECT notification_id FROM admin_notification_reads WHERE notification_id IN ({placeholders}) AND user_id=?', ids + [user['id']]).fetchall()
        existing_ids = [r['notification_id'] for r in existing]
        if existing_ids:
            ep = ','.join(['?' for _ in existing_ids])
            conn.execute(f'UPDATE admin_notification_reads SET cleared=1 WHERE notification_id IN ({ep}) AND user_id=?', existing_ids + [user['id']])
        missing_ids = [i for i in ids if i not in existing_ids]
        for mid in missing_ids:
            conn.execute('INSERT OR IGNORE INTO admin_notification_reads (notification_id, user_id, cleared) VALUES (?, ?, 1)', (mid, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'deleted': len(ids)})


def send_notification(user_id, title, content, ntype='info', link='', conn=None):
    own_conn = conn is None
    if own_conn:
        conn = db.get_conn()
    try:
        conn.execute('INSERT INTO notifications (user_id, title, content, type, link) VALUES (?, ?, ?, ?, ?)',
                    (user_id, title, content, ntype, link))
        if own_conn:
            conn.commit()
    finally:
        if own_conn:
            conn.close()


def normalize_name(name):
    if not name:
        return ''
    s = str(name)
    s = re.sub(r'[\u3000\xa0]', ' ', s)
    s = re.sub(r'\s+', '', s)
    s = re.sub(r'[（(].*?[）)]', '', s)
    for suffix in ['社团', '协会', '团队', '联盟', '中心', '俱乐部']:
        s = s.replace(suffix, '')
    s = s.strip()
    return s


def name_similarity(a, b):
    if not a or not b:
        return 0.0
    if a == b:
        return 1.0
    la, lb = len(a), len(b)
    if la == 0 or lb == 0:
        return 0.0
    common = sum(1 for c in a if c in b)
    sim = 2.0 * common / (la + lb)
    if a in b or b in a:
        sim = max(sim, 0.85)
    return sim


def find_similar_club(norm_name, existing_norms, threshold=0.75):
    best_match = None
    best_sim = 0.0
    for en in existing_norms:
        sim = name_similarity(norm_name, en)
        if sim > best_sim:
            best_sim = sim
            best_match = en
    if best_sim >= threshold:
        return best_match
    return None


def auto_find_field(row, headers, keywords, exclude=None):
    exclude = exclude or []
    best_h = None
    best_score = -1
    for h in headers:
        if h in exclude:
            continue
        h_lower = h.lower().replace(' ', '').replace('_', '')
        for kw in keywords:
            kw_lower = kw.lower().replace(' ', '').replace('_', '')
            if kw_lower == h_lower:
                score = 1000
            elif kw_lower in h_lower or h_lower in kw_lower:
                score = len(kw_lower) + 50
            else:
                common = sum(1 for ch in kw_lower if ch in h_lower)
                if common >= len(kw_lower) * 0.6 and common >= 2:
                    score = common
                else:
                    continue
            if score > best_score:
                best_score = score
                best_h = h
    if best_h:
        return str(row.get(best_h, '') or '')
    return ''


def smart_parse_excel(file_storage):
    import pandas as pd
    import io
    file_storage.seek(0)
    content = file_storage.read()
    filename = (file_storage.filename or '').lower()

    best_df = None
    best_score = -1
    best_header_row = 0

    is_csv = filename.endswith('.csv')

    csv_encodings = ['utf-8-sig', 'gbk', 'gb2312', 'gb18030', 'latin-1']

    import re as _re
    _cn_pat = _re.compile(r'[\u4e00-\u9fff]')

    for header_row in range(0, 15):
        try:
            if is_csv:
                df = None
                for enc in csv_encodings:
                    try:
                        df = pd.read_csv(io.BytesIO(content), header=header_row, encoding=enc)
                        break
                    except (UnicodeDecodeError, pd.errors.ParserError):
                        continue
                if df is None:
                    continue
            else:
                df = pd.read_excel(io.BytesIO(content), header=header_row)
            df = df.fillna('')

            valid_cols = 0
            cn_col_bonus = 0
            for c in df.columns:
                col_str = str(c).strip()
                if col_str and not col_str.startswith('Unnamed') and not col_str.startswith('附件') and len(col_str) > 0:
                    valid_cols += 1
                    if _cn_pat.search(col_str):
                        cn_col_bonus += 50

            valid_rows = sum(1 for _, row in df.iterrows() if any(str(v).strip() for v in row))
            score = valid_cols * 100 + cn_col_bonus + valid_rows

            if score > best_score:
                best_score = score
                best_df = df
                best_header_row = header_row
        except:
            continue

    if best_df is None:
        return None, [], []

    df = best_df.dropna(how='all').fillna('')

    clean_headers = []
    for h in df.columns:
        h_str = str(h).strip()
        if h_str.startswith('Unnamed') or h_str.startswith('附件'):
            h_str = ''
        clean_headers.append(h_str)

    rename_map = {}
    for i, h in enumerate(df.columns):
        if clean_headers[i]:
            rename_map[h] = clean_headers[i]
        else:
            rename_map[h] = f'列{i+1}'

    df = df.rename(columns=rename_map)
    headers = list(df.columns)
    data = df.to_dict('records')

    return df, headers, data



def get_all_base_clubs(base_data_list=None):
    club_info_map = {}
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT club_name, guiding_unit FROM club_profiles').fetchall()
    finally:
        conn.close()
    for r in rows:
        norm = normalize_name(r['club_name'])
        if not norm:
            continue
        if norm not in club_info_map:
            club_info_map[norm] = {'original_name': r['club_name'], 'teacher': '', 'unit': ''}
        if r['guiding_unit'] and r['guiding_unit'].strip():
            club_info_map[norm]['unit'] = r['guiding_unit'].strip()
    conn = db.get_conn()
    try:
        teacher_rows = conn.execute('SELECT club_name, teacher_name FROM club_teachers').fetchall()
    finally:
        conn.close()
    teacher_map = {}
    for tr in teacher_rows:
        if tr['teacher_name'] and tr['teacher_name'].strip():
            norm = normalize_name(tr['club_name'])
            if norm not in teacher_map:
                teacher_map[norm] = []
            teacher_map[norm].append(tr['teacher_name'].strip())
    for norm, names in teacher_map.items():
        if norm in club_info_map:
            club_info_map[norm]['teacher'] = '、'.join(names)
        else:
            club_info_map[norm] = {'original_name': norm, 'teacher': '、'.join(names), 'unit': ''}
    return club_info_map


def compute_club_stats():
    conn = db.get_conn()
    try:
        activities = conn.execute('SELECT club_name, COUNT(*) as cnt FROM activity_records GROUP BY club_name').fetchall()
        guidances = conn.execute('SELECT club_name, COUNT(DISTINCT session_id) as cnt FROM guidance_records GROUP BY club_name').fetchall()
        teachers = conn.execute('SELECT club_name, teacher_name, COUNT(*) as cnt FROM guidance_records WHERE teacher_name != "" AND teacher_name IS NOT NULL GROUP BY club_name, teacher_name').fetchall()
    finally:
        conn.close()

    club_map = {}
    for a in activities:
        club_map[a['club_name']] = {'club': a['club_name'], 'activityCount': a['cnt'], 'guidanceCount': 0, 'teachers': {}, 'fromBase': False}
    for g in guidances:
        if g['club_name'] not in club_map:
            club_map[g['club_name']] = {'club': g['club_name'], 'activityCount': 0, 'guidanceCount': g['cnt'], 'teachers': {}, 'fromBase': False}
        else:
            club_map[g['club_name']]['guidanceCount'] = g['cnt']
    for t in teachers:
        if t['club_name'] in club_map:
            club_map[t['club_name']]['teachers'][t['teacher_name']] = t['cnt']

    base_clubs = get_all_base_clubs()

    existing_norms = set()
    norm_to_original = {}
    for cn in club_map.keys():
        n = normalize_name(cn)
        existing_norms.add(n)
        norm_to_original[n] = cn

    for norm, info in base_clubs.items():
        if norm in existing_norms:
            continue
        similar = find_similar_club(norm, existing_norms)
        if similar:
            continue
        club_map[info['original_name']] = {
            'club': info['original_name'],
            'activityCount': 0,
            'guidanceCount': 0,
            'teachers': {},
            'fromBase': True
        }
        existing_norms.add(norm)

    clubs = sorted(club_map.values(), key=lambda x: (-x['activityCount'], x['club']))
    return clubs


def get_teacher_name_from_sources(club_name, base_data_list=None):
    teacher_name = ''
    conn = db.get_conn()
    try:
        rows = conn.execute("SELECT teacher_name FROM club_teachers WHERE club_name=?", (club_name,)).fetchall()
    finally:
        conn.close()
    if rows:
        names = [r['teacher_name'] for r in rows if r['teacher_name'] and r['teacher_name'].strip()]
        if names:
            teacher_name = '、'.join(names)
    if not teacher_name:
        conn = db.get_conn()
        try:
            all_clubs = conn.execute("SELECT DISTINCT club_name FROM club_teachers").fetchall()
        finally:
            conn.close()
        norm = normalize_name(club_name)
        for ac in all_clubs:
            if normalize_name(ac['club_name']) == norm:
                conn = db.get_conn()
                try:
                    rows = conn.execute("SELECT teacher_name FROM club_teachers WHERE club_name=?", (ac['club_name'],)).fetchall()
                finally:
                    conn.close()
                names = [r['teacher_name'] for r in rows if r['teacher_name'] and r['teacher_name'].strip()]
                if names:
                    teacher_name = '、'.join(names)
                break
    if not teacher_name:
        conn = db.get_conn()
        try:
            act_row = conn.execute("SELECT raw_data FROM activity_records WHERE club_name=? LIMIT 1", (club_name,)).fetchone()
        finally:
            conn.close()
        if act_row:
            try:
                raw = json.loads(act_row['raw_data'])
                for k, v in raw.items():
                    k_clean = k.replace('\n', '').replace(' ', '')
                    if any(kw in k_clean for kw in ['老师姓名', '教师姓名', '指导老师', '指导教师', '老师', '教师']):
                        if '政治面貌' in k_clean or '联系方式' in k_clean:
                            continue
                        teacher_name = str(v).strip()
                        break
            except:
                pass
    return teacher_name


def get_unit_from_base(club_name, base_data_list=None):
    unit = ''
    conn = db.get_conn()
    try:
        row = conn.execute("SELECT guiding_unit FROM club_profiles WHERE club_name=?", (club_name,)).fetchone()
    finally:
        conn.close()
    if row and row['guiding_unit'] and row['guiding_unit'].strip():
        unit = row['guiding_unit'].strip()
    if not unit:
        conn = db.get_conn()
        try:
            all_clubs = conn.execute("SELECT club_name, guiding_unit FROM club_profiles WHERE guiding_unit IS NOT NULL AND guiding_unit != ''").fetchall()
        finally:
            conn.close()
        norm = normalize_name(club_name)
        for ac in all_clubs:
            if normalize_name(ac['club_name']) == norm:
                unit = ac['guiding_unit'].strip()
                break
    return unit


def get_online_activity_status(club_name):
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT has_photo, has_summary, activity_date, club_name FROM online_activity_data WHERE club_name=?', (club_name,)).fetchall()
        if not rows:
            norm = normalize_name(club_name)
            all_online = conn.execute('SELECT has_photo, has_summary, activity_date, club_name FROM online_activity_data').fetchall()
            for r in all_online:
                if normalize_name(r['club_name']) == norm:
                    rows.append(r)
        act_rows = conn.execute('SELECT activity_date, activity_content FROM activity_records WHERE club_name=?', (club_name,)).fetchall()
    finally:
        conn.close()

    has_photo = any(r['has_photo'] for r in rows) if rows else False
    has_summary = any(r['has_summary'] for r in rows) if rows else False

    online_dates = set()
    for r in rows:
        if r['activity_date']:
            online_dates.add(r['activity_date'])

    details = []
    seen = set()

    if act_rows:
        for r in act_rows:
            date_str = r['activity_date'] if r['activity_date'] else '未知时间'
            key = date_str
            if key in seen:
                continue
            seen.add(key)
            if date_str in online_dates:
                for or_ in rows:
                    if or_['activity_date'] == date_str:
                        photo_str = '已上传照片' if or_['has_photo'] else '未上传照片'
                        summary_str = '已上传总结' if or_['has_summary'] else '未上传总结'
                        details.append(f'{date_str} {photo_str} {summary_str}')
                        break
            else:
                photo_str = '已上传照片' if has_photo else '未上传照片'
                summary_str = '已上传总结' if has_summary else '未上传总结'
                details.append(f'{date_str} {photo_str} {summary_str}')
    elif rows:
        for r in rows:
            date_str = r['activity_date'] if r['activity_date'] else '未知时间'
            key = (date_str, r['has_photo'], r['has_summary'])
            if key in seen:
                continue
            seen.add(key)
            photo_str = '已上传照片' if r['has_photo'] else '未上传照片'
            summary_str = '已上传总结' if r['has_summary'] else '未上传总结'
            details.append(f'{date_str} {photo_str} {summary_str}')

    return {'hasPhoto': has_photo, 'hasSummary': has_summary, 'details': details}


def fetch_and_store_online_data(session, start_date='', end_date=''):
    base_url = 'https://xsc.jnpec.edu.cn/syt/st/jnye/record/activityPageData.htm'
    all_records = []
    page = 1
    total_pages = 1

    while page <= total_pages:
        params = {'roleType': 'Tw', 'page': str(page), 'rows': '100'}
        if start_date:
            params['startDate'] = start_date
        if end_date:
            params['endDate'] = end_date

        try:
            resp = session.get(base_url, params=params, verify=False, timeout=20)
        except Exception as e:
            return {'error': f'请求失败: {str(e)}', 'fetched': len(all_records)}

        if resp.status_code == 501:
            return {'error': '会话已过期，请重新登录', 'fetched': len(all_records)}

        if resp.status_code != 200:
            return {'error': f'服务器返回状态码 {resp.status_code}', 'fetched': len(all_records)}

        try:
            result = resp.json()
        except:
            return {'error': f'返回数据不是有效JSON: {resp.text[:500]}', 'fetched': len(all_records)}

        rows = result.get('rows', result.get('data', result.get('list', [])))
        if not rows and page == 1:
            if isinstance(result, list):
                rows = result
            elif isinstance(result, dict):
                for v in result.values():
                    if isinstance(v, list) and len(v) > 0:
                        rows = v
                        break

        if not rows:
            break

        all_records.extend(rows)

        total = result.get('total', result.get('totalCount', 0))
        if total and isinstance(total, (int, float)):
            total_pages = max(1, (int(total) + 99) // 100)
        else:
            total_pages = 1

        page += 1
        if page > 50:
            break

    if not all_records:
        return {'error': '未获取到数据，请检查登录状态是否有效', 'fetched': 0}

    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM online_activity_data')
        inserted = 0
        for rec in all_records:
            club_name = ''
            for key in ['clubName', 'club_name', 'shetuanName', 'shetuan_name', 'organizationName', 'orgName', 'name', 'title', '社团名称', '社团名']:
                if key in rec and rec[key]:
                    club_name = str(rec[key]).strip()
                    break
            if not club_name:
                continue

            has_photo = 0
            has_summary = 0
            activity_title = ''
            activity_date = ''

            for key in rec:
                kl = key.lower() if isinstance(key, str) else str(key)

                val = rec[key]
                val_str = str(val) if val is not None else ''
                val_stripped = val_str.strip()

                file_id_fields = {'hdzjFileId', 'hdzjfileid'}
                if kl in file_id_fields:
                    if val_stripped and val_stripped not in ('', '0', 'null', 'undefined', '[]', '{}', 'false', 'False'):
                        has_photo = 1

                summary_fields = {'zdjhFileId', 'zdjhfileid', 'summaryFile', 'summaryfile', 'reportId', 'reportid', 'conclusionFile'}
                if kl in summary_fields:
                    if val_stripped and val_stripped not in ('', '0', 'null', 'undefined', '[]', '{}', 'false', 'False'):
                        has_summary = 1

                if not has_photo:
                    if kl in ('canuploadimage', 'canuploadhdzjfile') or 'canupload' in kl and 'image' in kl:
                        if val is True or val == 'true' or val == 'True':
                            has_photo = 1

                if not has_summary:
                    if kl in ('canuploadhdzjfile',) or 'canupload' in kl and ('hdzj' in kl or 'summary' in kl or 'report' in kl):
                        if val is True or val == 'true' or val == 'True':
                            has_summary = 1

                if any(kw in kl for kw in ['subject', '活动主题', '主题']):
                    if not activity_title and val:
                        tv = str(val).strip() if isinstance(val, (str, int, float)) else ''
                        if tv and tv not in ('null', 'undefined', '[]', '{}'):
                            activity_title = tv

                if any(kw in kl for kw in ['name', '标题', '活动名称', 'activityname']):
                    if not activity_title and val:
                        if not isinstance(val, dict):
                            tv = str(val).strip() if isinstance(val, (str, int, float)) else ''
                            if tv and tv not in ('null', 'undefined', '[]', '{}'):
                                activity_title = tv

                if any(kw in kl for kw in ['date', 'time', '日期', '时间']):
                    if not activity_date and val:
                        tv = str(val).strip() if isinstance(val, (str, int, float)) else ''
                        if tv and tv not in ('null', 'undefined', '[]', '{}'):
                            activity_date = tv

            conn.execute(
                'INSERT INTO online_activity_data (club_name, has_photo, has_summary, activity_title, activity_date, raw_json) VALUES (?, ?, ?, ?, ?, ?)',
                (club_name, has_photo, has_summary, activity_title, activity_date, json.dumps(rec, ensure_ascii=False, default=str)))
            inserted += 1
        conn.commit()
    finally:
        conn.close()
    cache.clear()
    return {'success': True, 'fetched': len(all_records), 'inserted': inserted}


@app.route('/api/fetch-online', methods=['POST'])
def fetch_online_data():
    try:
        import requests as req_lib
        import urllib3
        urllib3.disable_warnings()
    except ImportError:
        return jsonify({'error': '需要安装 requests 库: pip install requests'}), 500

    data = request.json or {}
    jsessionid = data.get('jsessionid', '').strip()
    cookie_str = data.get('cookie', '').strip()
    start_date = data.get('startDate', '')
    end_date = data.get('endDate', '')

    if not jsessionid and not cookie_str:
        return jsonify({'error': '请提供 JSESSIONID 或 Cookie'}), 400

    session = req_lib.Session()

    if cookie_str:
        for part in cookie_str.split(';'):
            part = part.strip()
            if '=' in part:
                k, v = part.split('=', 1)
                k, v = k.strip(), v.strip()
                if k:
                    session.cookies.set(k, v, domain='xsc.jnpec.edu.cn', path='/')
    else:
        session.cookies.set('JSESSIONID', jsessionid, domain='xsc.jnpec.edu.cn', path='/')

    session.headers.update({
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'application/json, text/javascript, */*; q=0.01',
        'X-Requested-With': 'XMLHttpRequest',
        'Referer': 'https://xsc.jnpec.edu.cn/syt/st/jnye/record/activityPageData.htm?roleType=Tw',
    })

    result = fetch_and_store_online_data(session, start_date, end_date)
    return jsonify(result)


@app.route('/api/mock-online', methods=['POST'])
def mock_online_data():
    import random
    conn = db.get_conn()
    try:
        act_clubs = conn.execute('SELECT DISTINCT club_name FROM activity_records').fetchall()
        club_names = [r['club_name'] for r in act_clubs]
    finally:
        conn.close()

    if not club_names:
        return jsonify({'error': '请先上传社团活动记录'}), 200

    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM online_activity_data')
        inserted = 0
        for club in club_names:
            act_rows = conn.execute('SELECT activity_date FROM activity_records WHERE club_name=?', (club,)).fetchall()
            for r in act_rows:
                date_str = r['activity_date'] if r['activity_date'] else ''
                has_photo = 1 if random.random() > 0.3 else 0
                has_summary = 1 if random.random() > 0.4 else 0
                conn.execute(
                    'INSERT INTO online_activity_data (club_name, has_photo, has_summary, activity_title, activity_date, raw_json) VALUES (?, ?, ?, ?, ?, ?)',
                    (club, has_photo, has_summary, f'{club}活动', date_str, '{}'))
                inserted += 1
        conn.commit()
    finally:
        conn.close()
    cache.clear()
    return jsonify({'success': True, 'fetched': inserted, 'inserted': inserted, 'mock': True})


@app.route('/dashboard.html')
def serve_dashboard():
    resp = send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'dashboard.html'))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp


@app.route('/')
def index():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'login.html'))


@app.route('/stats.html')
def stats_page():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'stats.html'))


@app.route('/upload.html')
def upload_page():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'upload.html'))


@app.route('/feedback.html')
def feedback_page():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'feedback.html'))


@app.route('/api/upload/activity', methods=['POST'])
def upload_activity():
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    try:
        df, headers, data = smart_parse_excel(file)
        if df is None:
            return jsonify({'error': '文件解析失败，无法识别有效表头'}), 400
    except Exception as e:
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 400

    result = cleaner.clean(data, headers)
    club_col = cleaner.auto_detect_column(result['data'][0].keys() if result['data'] else headers, ['社团名称', '社团名', '社团', '组织名称', '组织', 'club', 'name', '社团组织', '所属社团', '社团全称', '社名', '团队名称', '团队名', '团队', '协会名称', '协会', '俱乐部', '组织机构', '单位名称', '单位', 'group', 'organization', 'society'])
    if not club_col:
        return jsonify({'error': '未找到社团名称列，请确保表格中包含"社团"或"名称"相关列'}), 400

    def safe_int(val):
        try:
            if val is None or str(val).strip() == '':
                return 0
            s = str(val).strip()
            import re as _re
            nums = _re.findall(r'-?\d+\.?\d*', s)
            if nums:
                return int(float(nums[0]))
            return 0
        except (ValueError, TypeError):
            return 0

    file_uuid = str(uuid.uuid4())
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM activity_records')
        conn.execute('DELETE FROM uploaded_files WHERE file_type=?', ('activity',))
        conn.execute('INSERT INTO uploaded_files (uuid, original_name, file_type, file_size, row_count) VALUES (?, ?, ?, ?, ?)',
                     (file_uuid, file.filename, 'activity', 0, len(result['data'])))
        for row in result['data']:
            conn.execute(
                'INSERT INTO activity_records (file_uuid, club_name, activity_date, activity_content, location, participant_count, raw_data) VALUES (?, ?, ?, ?, ?, ?, ?)',
                (file_uuid, str(row.get(club_col, '')).strip(),
                 auto_find_field(row, headers, ['活动日期', '日期', '时间', 'date', '活动时间', '开展日期', '开展时间', '举办日期', '举办时间', '活动开展日期', '活动开展时间', '起始日期', '起始时间', '发生日期', '发生时间', '记录日期', '记录时间', '创建日期', '创建时间', 'time', 'datetime'], exclude=[club_col]),
                 auto_find_field(row, headers, ['活动内容', '活动名称', '内容', '活动', '主题', '活动主题', '活动标题', '活动简介', '活动描述', '活动概况', '活动事项', '名称', 'title', 'content', 'description', 'subject', 'event', '活动类型', '活动项目'], exclude=[club_col]),
                 auto_find_field(row, headers, ['活动地点', '地点', '位置', 'place', 'location', '开展地点', '举办地点', '场地', '场所', '地址', '活动地址', '活动场所', 'venue', 'address', 'site'], exclude=[club_col]),
                 safe_int(auto_find_field(row, headers, ['参与人数', '人数', '参与', 'participant', '参加人数', '出席人数', '到场人数', '应到人数', '实到人数', '报名人数', '活动人数', '出席', '参加', 'count', 'number', 'total', '人数（人）', '人数/人', '人数(人)'], exclude=[club_col])),
                 json.dumps(row, ensure_ascii=False)))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({'error': f'数据存储失败: {str(e)}'}), 500
    finally:
        conn.close()
    cache.clear()
    return jsonify({
        'success': True, 'fileUuid': file_uuid, 'fileName': file.filename,
        'originalRows': len(data), 'cleanedRows': len(result['data']),
        'removed': len(data) - len(result['data']), 'headers': headers,
        'clubColumn': club_col, 'cleanReport': result['report']
    })


@app.route('/api/upload/guidance', methods=['POST'])
def upload_guidance():
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    try:
        df, headers, data = smart_parse_excel(file)
        if df is None:
            return jsonify({'error': '文件解析失败，无法识别有效表头'}), 400
    except Exception as e:
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 400

    result = cleaner.clean(data, headers)
    club_col = cleaner.auto_detect_column(result['data'][0].keys() if result['data'] else headers, ['社团名称', '社团名', '社团', '组织名称', '组织', 'club', '社团组织', '所属社团', '社团全称', '社名', '团队名称', '团队名', '团队', '协会名称', '协会', '俱乐部', '组织机构', '单位名称', '单位', 'group', 'organization', 'society'])
    teacher_col = cleaner.auto_detect_column(result['data'][0].keys() if result['data'] else headers, ['指导老师', '社团指导老师', '指导教师', '辅导老师', '负责老师', '老师', '教师', '指导', 'teacher', 'advisor', 'instructor', '指导者', '指导人', '辅导教师', '辅导者', '负责教师', '负责教师', '带队老师', '带队教师', '指导老师姓名', '教师姓名', '老师姓名', '导师', 'mentor', 'supervisor', 'coach'], exclude=[club_col] if club_col else [])
    if not club_col:
        return jsonify({'error': '未找到社团名称列'}), 400

    file_uuid = str(uuid.uuid4())
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM guidance_records')
        conn.execute('DELETE FROM uploaded_files WHERE file_type=?', ('guidance',))
        conn.execute('INSERT INTO uploaded_files (uuid, original_name, file_type, file_size, row_count) VALUES (?, ?, ?, ?, ?)',
                     (file_uuid, file.filename, 'guidance', 0, len(result['data'])))
        for row in result['data']:
            session_id = str(uuid.uuid4())
            teacher_val = str(row.get(teacher_col, '') if teacher_col else '').strip()
            teacher_names = [t.strip() for t in teacher_val.replace('，', ',').split(',') if t.strip()]
            if not teacher_names:
                teacher_names = ['']
            for t_name in teacher_names:
                conn.execute(
                    'INSERT INTO guidance_records (file_uuid, session_id, club_name, teacher_name, guidance_date, guidance_content, raw_data) VALUES (?, ?, ?, ?, ?, ?, ?)',
                    (file_uuid, session_id, str(row.get(club_col, '')).strip(),
                     t_name,
                     auto_find_field(row, headers, ['指导日期', '日期', '时间', 'date', '指导时间', '辅导日期', '辅导时间', '活动日期', '活动时间', '开展日期', '开展时间', '记录日期', '记录时间', '创建日期', '创建时间', '起始日期', '起始时间', 'time', 'datetime'], exclude=[club_col, teacher_col] if teacher_col else [club_col]),
                     auto_find_field(row, headers, ['指导内容', '内容', '主题', '指导情况', '辅导内容', '指导活动', '指导事项', '指导记录', '活动内容', '活动名称', '活动', '指导简介', '指导描述', '指导概况', 'description', 'content', 'detail', '备注', '说明', '情况说明', '指导说明'], exclude=[club_col, teacher_col] if teacher_col else [club_col]),
                     json.dumps(row, ensure_ascii=False)))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({'error': f'数据存储失败: {str(e)}'}), 500
    finally:
        conn.close()
    cache.clear()
    return jsonify({
        'success': True, 'fileUuid': file_uuid, 'fileName': file.filename,
        'originalRows': len(data), 'cleanedRows': len(result['data']),
        'headers': headers, 'clubColumn': club_col, 'teacherColumn': teacher_col,
        'cleanReport': result['report']
    })


@app.route('/api/base-data', methods=['GET'])
def get_base_data():
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT uuid, name, original_name, headers, row_count, upload_time FROM base_data ORDER BY upload_time DESC').fetchall()
    finally:
        conn.close()
    result = []
    for r in rows:
        result.append({
            'uuid': r['uuid'], 'name': r['name'], 'originalName': r['original_name'],
            'headers': json.loads(r['headers']) if r['headers'] else [],
            'rowCount': r['row_count'], 'uploadTime': r['upload_time']
        })
    return jsonify({'success': True, 'data': result})


UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


@app.route('/api/club-tokens', methods=['GET'])
def get_club_tokens():
    clubs = compute_club_stats()
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT club_name, token FROM club_tokens').fetchall()
        token_map = {r['club_name']: r['token'] for r in existing}
    finally:
        conn.close()

    result = []
    for c in clubs:
        club_name = c['club']
        if club_name in token_map:
            token = token_map[club_name]
        else:
            import hashlib
            token = hashlib.md5((club_name + str(uuid.uuid4())).encode()).hexdigest()[:12]
            conn = db.get_conn()
            try:
                conn.execute('INSERT OR IGNORE INTO club_tokens (club_name, token) VALUES (?, ?)', (club_name, token))
                conn.commit()
            finally:
                conn.close()
        result.append({'club': club_name, 'token': token, 'activityCount': c['activityCount']})
    return jsonify({'success': True, 'data': result})


@app.route('/upload/<token>')
def club_upload_page(token):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT club_name FROM club_tokens WHERE token=?', (token,)).fetchone()
    finally:
        conn.close()
    if not row:
        return '<h1>链接无效</h1><p>该上传链接不存在或已失效</p>', 404
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'upload.html'))


@app.route('/api/club-upload-info/<token>')
def club_upload_info(token):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT club_name FROM club_tokens WHERE token=?', (token,)).fetchone()
        if not row:
            return jsonify({'error': '链接无效'}), 404
        club_name = row['club_name']
        uploads = conn.execute('SELECT id, file_name, file_type, description, upload_time, status, reject_reason, category, group_id FROM club_uploads WHERE club_token=? ORDER BY upload_time DESC', (token,)).fetchall()
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'clubName': club_name,
        'token': token,
        'uploads': [{'id': u['id'], 'fileName': u['file_name'], 'fileType': u['file_type'], 'description': u['description'], 'uploadTime': u['upload_time'], 'status': u['status'] or 'pending', 'rejectReason': u['reject_reason'] or '', 'category': u['category'] if 'category' in u.keys() else '', 'groupId': u['group_id'] or ''} for u in uploads]
    })


@app.route('/api/club-upload/<token>', methods=['POST'])
def club_upload_file(token):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT club_name FROM club_tokens WHERE token=?', (token,)).fetchone()
    finally:
        conn.close()
    if not row:
        return jsonify({'error': '链接无效'}), 404

    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    description = request.form.get('description', '')
    group_id = request.form.get('groupId', '')
    category = request.form.get('category', '')
    club_name = row['club_name']

    import hashlib as hl
    ext = os.path.splitext(file.filename)[1]
    safe_name = hl.md5((file.filename + str(uuid.uuid4())).encode()).hexdigest()[:16] + ext
    key = 'club_uploads/' + token + '/' + safe_name
    storage.save(file, key)

    file_type = ext.lstrip('.').lower() if ext else 'unknown'

    conn = db.get_conn()
    try:
        conn.execute('INSERT INTO club_uploads (club_token, club_name, file_name, file_path, file_type, description, group_id, status, category, source) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                     (token, club_name, file.filename, key, file_type, description, group_id, 'pending', category, 'upload'))
        conn.commit()
    finally:
        conn.close()

    if description.startswith('[总结]'):
        title = description.replace('[总结]', '').split('|||')[0]
        conn = db.get_conn()
        try:
            admins = conn.execute('SELECT id FROM users WHERE role="admin"').fetchall()
        finally:
            conn.close()
        for admin in admins:
            send_notification(admin['id'], f'📋 新材料待审批', f'{club_name} 提交了活动材料「{title}」，请及时审批', 'upload', '/upload.html')

    if description.startswith('[反馈]'):
        parts = description.replace('[反馈]', '').split('|||')
        fb_title = parts[1] if len(parts) > 1 else '问题反馈'
        conn = db.get_conn()
        try:
            admins = conn.execute('SELECT id FROM users WHERE role="admin"').fetchall()
        finally:
            conn.close()
        for admin in admins:
            send_notification(admin['id'], f'💬 新问题反馈', f'{club_name} 反馈了问题「{fb_title}」，请及时处理', 'feedback', '/feedback.html')

    return jsonify({'success': True, 'fileName': file.filename, 'clubName': club_name})


@app.route('/api/club-upload/<token>/<int:file_id>', methods=['DELETE'])
def club_delete_file(token, file_id):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT file_path FROM club_uploads WHERE id=? AND club_token=?', (file_id, token)).fetchone()
        if row and row['file_path']:
            storage.delete(row['file_path'])
        conn.execute('DELETE FROM club_uploads WHERE id=? AND club_token=?', (file_id, token))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/delete-rejected-group/<group_id>', methods=['DELETE'])
def delete_rejected_group(group_id):
    user = get_current_user()
    if not user or user['role'] != 'user':
        return jsonify({'error': '无权限，仅社团负责人可操作'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, club_name, file_path, status, upload_time FROM club_uploads WHERE group_id=?', (group_id,)).fetchall()
        if not rows:
            return jsonify({'error': '记录不存在'}), 404
        for r in rows:
            if r['club_name'] != user['club_name']:
                return jsonify({'error': '无权删除此材料'}), 403
            if r['status'] != 'rejected':
                return jsonify({'error': '只能删除被打回的材料'}), 400
        from datetime import datetime, timedelta
        upload_time_str = rows[0]['upload_time'] or ''
        try:
            upload_time = datetime.strptime(upload_time_str, '%Y-%m-%d %H:%M:%S')
        except Exception:
            upload_time = datetime.strptime(upload_time_str[:19], '%Y-%m-%d %H:%M:%S') if len(upload_time_str) >= 19 else datetime.now()
        if datetime.now() - upload_time < timedelta(hours=24):
            return jsonify({'error': '被打回的材料需满24小时后才可删除'}), 400
        for r in rows:
            if r['file_path']:
                storage.delete(r['file_path'])
        conn.execute('DELETE FROM club_uploads WHERE group_id=?', (group_id,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/batch-delete-materials', methods=['POST'])
def batch_delete_materials():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限，仅管理员可操作'}), 403
    data = request.json or {}
    group_ids = data.get('groupIds', [])
    if not group_ids:
        return jsonify({'error': '请选择要删除的材料'}), 400
    conn = db.get_conn()
    try:
        placeholders = ','.join(['?' for _ in group_ids])
        rows = conn.execute(f'SELECT id, file_path FROM club_uploads WHERE group_id IN ({placeholders})', group_ids).fetchall()
        for r in rows:
            if r['file_path']:
                storage.delete(r['file_path'])
        conn.execute(f'DELETE FROM club_uploads WHERE group_id IN ({placeholders})', group_ids)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'deleted': len(group_ids)})


@app.route('/api/club-download/<token>/<int:file_id>')
def club_download_file(token, file_id):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT file_name, file_path, file_type FROM club_uploads WHERE id=? AND club_token=?', (file_id, token)).fetchone()
    finally:
        conn.close()
    if not row or not row['file_path']:
        return jsonify({'error': '文件不存在'}), 404
    path = storage.get_path(row['file_path'])
    if not path and not storage.exists(row['file_path']):
        return jsonify({'error': '文件不存在'}), 404
    if path:
        ft = (row['file_type'] or '').lower()
        if ft.startswith('image') or ft in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp'):
            return send_file(path, as_attachment=False)
        return send_file(path, as_attachment=True, download_name=row['file_name'])
    url = storage.get_url(row['file_path'])
    if url:
        return jsonify({'url': url})


@app.route('/api/approve-group/<group_id>', methods=['POST'])
def approve_group(group_id):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    action = data.get('action', '')
    reject_reason = data.get('rejectReason', '')

    conn = db.get_conn()
    try:
        upload = conn.execute('SELECT club_name, description FROM club_uploads WHERE group_id=? LIMIT 1', (group_id,)).fetchone()
        if not upload:
            return jsonify({'error': '记录不存在'}), 404
        title = upload['description'].replace('[总结]', '').split('|||')[0] if upload['description'] and upload['description'].startswith('[总结]') else '材料'
        club_users = conn.execute('SELECT id FROM users WHERE club_name=?', (upload['club_name'],)).fetchall()
        if action == 'approve':
            conn.execute('UPDATE club_uploads SET status="approved" WHERE group_id=?', (group_id,))
            for cu in club_users:
                send_notification(cu['id'], '✅ 材料已通过', f'您提交的「{title}」已通过审批', 'approve', '/upload.html', conn=conn)
        elif action == 'reject':
            if not reject_reason:
                return jsonify({'error': '请填写驳回原因'}), 400
            conn.execute('UPDATE club_uploads SET status="rejected", reject_reason=? WHERE group_id=?', (reject_reason, group_id))
            for cu in club_users:
                send_notification(cu['id'], '❌ 材料已打回', f'您提交的「{title}」被打回，原因：{reject_reason}', 'reject', '/upload.html', conn=conn)
        else:
            return jsonify({'error': '无效操作'}), 400
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/approve/<int:upload_id>', methods=['POST'])
def approve_upload(upload_id):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    action = data.get('action', '')
    reject_reason = data.get('rejectReason', '')

    conn = db.get_conn()
    try:
        upload = conn.execute('SELECT club_name, description, club_token, group_id FROM club_uploads WHERE id=?', (upload_id,)).fetchone()
        if not upload:
            return jsonify({'error': '记录不存在'}), 404
        group_id = upload['group_id'] or str(upload_id)
        title = upload['description'].replace('[总结]', '').split('|||')[0] if upload['description'] and upload['description'].startswith('[总结]') else '材料'
        club_users = conn.execute('SELECT id FROM users WHERE club_name=?', (upload['club_name'],)).fetchall()
        if action == 'approve':
            conn.execute('UPDATE club_uploads SET status="approved" WHERE group_id=? OR id=?', (group_id, upload_id))
            for cu in club_users:
                send_notification(cu['id'], '✅ 材料已通过', f'您提交的「{title}」已通过审批', 'approve', '/upload.html', conn=conn)
        elif action == 'reject':
            if not reject_reason:
                return jsonify({'error': '请填写驳回原因'}), 400
            conn.execute('UPDATE club_uploads SET status="rejected", reject_reason=? WHERE group_id=? OR id=?', (reject_reason, group_id, upload_id))
            for cu in club_users:
                send_notification(cu['id'], '❌ 材料已打回', f'您提交的「{title}」被打回，原因：{reject_reason}', 'reject', '/upload.html', conn=conn)
        else:
            return jsonify({'error': '无效操作'}), 400
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/pending-approvals')
def pending_approvals():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, club_name, file_name, description, status, reject_reason, upload_time, group_id FROM club_uploads WHERE status="pending" AND (source="upload" OR source="activity" OR source IS NULL) ORDER BY upload_time DESC').fetchall()
    finally:
        conn.close()
    groups = {}
    for r in rows:
        gid = r['group_id'] or str(r['id'])
        if gid not in groups:
            groups[gid] = {'groupId': gid, 'clubName': r['club_name'], 'time': r['upload_time'], 'files': [], 'description': r['description']}
        groups[gid]['files'].append({'id': r['id'], 'fileName': r['file_name'], 'fileType': r['file_name'].rsplit('.', 1)[-1] if '.' in r['file_name'] else ''})
    return jsonify({'success': True, 'data': list(groups.values())})


@app.route('/api/submit-feedback', methods=['POST'])
def submit_feedback():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    data = request.get_json(force=True) if request.is_json else request.form
    fb_type = (data.get('type') or 'other').strip()
    title = (data.get('title') or '').strip()
    body = (data.get('body') or data.get('content') or '').strip()
    activity_name = (data.get('activityName') or '').strip()
    activity_time = (data.get('activityTime') or '').strip()
    if not title and not body:
        return jsonify({'error': '请填写完整信息'}), 400
    club_name = (data.get('clubName') or user.get('club_name') or '').strip()
    if not club_name:
        return jsonify({'error': '未关联社团'}), 400
    file_path = ''
    file_name = ''
    files_list = []
    file = request.files.get('file')
    if file and file.filename:
        import hashlib
        ext = os.path.splitext(file.filename)[1] or '.png'
        md5 = hashlib.md5(file.read()).hexdigest()
        file.seek(0)
        save_name = md5 + ext
        key = 'feedback_files/' + save_name
        storage.save(file, key)
        file_path = key
        file_name = file.filename
        files_list.append({'path': key, 'name': file.filename})
    multi_files = request.files.getlist('files')
    for f in multi_files:
        if f and f.filename:
            import hashlib
            ext = os.path.splitext(f.filename)[1] or '.png'
            md5 = hashlib.md5(f.read()).hexdigest()
            f.seek(0)
            save_name = md5 + ext
            key = 'feedback_files/' + save_name
            storage.save(f, key)
            files_list.append({'path': key, 'name': f.filename})
    files_json_str = json.dumps(files_list, ensure_ascii=False) if files_list else ''
    if not file_path and files_list:
        file_path = files_list[0]['path']
        file_name = files_list[0]['name']
    conn = db.get_conn()
    try:
        conn.execute('INSERT INTO feedbacks (club_name, user_id, type, title, body, activity_name, activity_time, file_path, file_name, files_json) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                     (club_name, user['id'], fb_type, title, body, activity_name, activity_time, file_path, file_name, files_json_str))
        conn.commit()
    finally:
        conn.close()
    admins = []
    conn2 = db.get_conn()
    try:
        admins = conn2.execute("SELECT id FROM users WHERE role='admin'").fetchall()
    finally:
        conn2.close()
    for a in admins:
        send_notification(a['id'], '💬 新问题反馈', f'{club_name} 反馈了问题「{title}」，请及时处理', 'feedback', '/feedback.html')
    return jsonify({'success': True})


@app.route('/api/feedback-file/<int:fbid>')
def get_feedback_file(fbid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT file_path, file_name FROM feedbacks WHERE id=?', (fbid,)).fetchone()
    finally:
        conn.close()
    if not row or not row['file_path']:
        return jsonify({'error': '文件不存在'}), 404
    path = storage.get_path(row['file_path'])
    if not path:
        url = storage.get_url(row['file_path'])
        if url:
            return redirect(url)
        return jsonify({'error': '文件不存在'}), 404
    ext = os.path.splitext(row['file_name'] or '')[1].lower()
    img_exts = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg']
    if ext in img_exts:
        return send_file(path, as_attachment=False)
    return send_file(path, as_attachment=True, download_name=row['file_name'])


@app.route('/api/feedback-file-by-key/<path:file_key>')
def get_feedback_file_by_key(file_key):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    full_key = 'feedback_files/' + file_key
    path = storage.get_path(full_key)
    if not path:
        url = storage.get_url(full_key)
        if url:
            return redirect(url)
        return jsonify({'error': '文件不存在'}), 404
    name = os.path.basename(file_key)
    ext = os.path.splitext(name)[1].lower()
    img_exts = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg']
    if ext in img_exts:
        return send_file(path, as_attachment=False)
    return send_file(path, as_attachment=True, download_name=name)


@app.route('/api/my-feedbacks')
def my_feedbacks():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, type, title, body, status, result, created_at, activity_name, activity_time, file_path, file_name, files_json FROM feedbacks WHERE user_id=? ORDER BY created_at DESC', (user['id'],)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'type': r['type'], 'title': r['title'],
        'content': r['body'], 'status': r['status'] or 'pending', 'reply': r['result'] or '', 'createdAt': local_time(r['created_at']),
        'activityName': r['activity_name'] if 'activity_name' in r.keys() else '',
        'activityTime': r['activity_time'] if 'activity_time' in r.keys() else '',
        'filePath': r['file_path'] if 'file_path' in r.keys() else '',
        'fileName': r['file_name'] if 'file_name' in r.keys() else '',
        'filesJson': r['files_json'] if 'files_json' in r.keys() else ''} for r in rows]})


@app.route('/api/all-feedbacks')
def all_feedbacks():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        if user['role'] == 'admin':
            rows = conn.execute('SELECT f.id, f.club_name, f.type, f.title, f.body, f.status, f.result, f.created_at, f.activity_name, f.activity_time, f.file_path, f.file_name, f.files_json FROM feedbacks f ORDER BY f.club_name, f.created_at DESC').fetchall()
        else:
            rows = conn.execute('SELECT f.id, f.club_name, f.type, f.title, f.body, f.status, f.result, f.created_at, f.activity_name, f.activity_time, f.file_path, f.file_name, f.files_json FROM feedbacks f WHERE f.user_id=? ORDER BY f.created_at DESC', (user['id'],)).fetchall()
    finally:
        conn.close()
    result = []
    for r in rows:
        result.append({
            'id': r['id'],
            'clubName': r['club_name'],
            'type': r['type'],
            'title': r['title'],
            'body': r['body'],
            'time': local_time(r['created_at']),
            'status': r['status'] or 'pending',
            'result': r['result'] or '',
            'activityName': r['activity_name'] if 'activity_name' in r.keys() else '',
            'activityTime': r['activity_time'] if 'activity_time' in r.keys() else '',
            'filePath': r['file_path'] if 'file_path' in r.keys() else '',
            'fileName': r['file_name'] if 'file_name' in r.keys() else '',
            'filesJson': r['files_json'] if 'files_json' in r.keys() else ''
        })
    return jsonify({'success': True, 'data': result})


@app.route('/api/handle-feedback/<int:fb_id>', methods=['POST'])
def handle_feedback(fb_id):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    result_text = data.get('result', '').strip()
    if not result_text:
        return jsonify({'error': '请填写处理结果'}), 400
    conn = db.get_conn()
    try:
        fb = conn.execute('SELECT club_name, title, user_id FROM feedbacks WHERE id=?', (fb_id,)).fetchone()
        if not fb:
            return jsonify({'error': '反馈不存在'}), 404
        conn.execute('UPDATE feedbacks SET status="resolved", result=? WHERE id=?', (result_text, fb_id))
        if fb['user_id']:
            send_notification(fb['user_id'], '✅ 问题已处理', f'您反馈的「{fb["title"]}」已处理，结果：{result_text}', 'feedback', '/dashboard.html', conn=conn)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/all-materials')
def all_materials():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    date_where = ''
    date_params = []
    if start_date:
        date_where += ' AND date(upload_time)>=?'
        date_params.append(start_date)
    if end_date:
        date_where += ' AND date(upload_time)<=?'
        date_params.append(end_date)
    conn = db.get_conn()
    try:
        if user['role'] == 'admin':
            rows = conn.execute(f'SELECT id, club_name, file_name, file_type, description, upload_time, status, reject_reason, group_id, club_token, category FROM club_uploads WHERE (source="upload" OR source="activity" OR source IS NULL) AND (description LIKE "[总结]%" OR description LIKE "[照片]%" OR description LIKE "[PDF]%" OR description LIKE "[Word]%" OR description LIKE "[文件]%"){date_where} ORDER BY club_name, upload_time DESC', date_params).fetchall()
        elif user['role'] == 'student':
            rows = conn.execute(f'SELECT id, club_name, file_name, file_type, description, upload_time, status, reject_reason, group_id, club_token, category FROM club_uploads WHERE status="approved" AND (source="upload" OR source="activity" OR source IS NULL) AND (description LIKE "[总结]%" OR description LIKE "[照片]%" OR description LIKE "[PDF]%" OR description LIKE "[Word]%" OR description LIKE "[文件]%"){date_where} ORDER BY upload_time DESC', date_params).fetchall()
        else:
            token_row = conn.execute('SELECT token FROM club_tokens WHERE club_name=?', (user['club_name'],)).fetchone()
            if not token_row:
                return jsonify({'success': True, 'data': []})
            rows = conn.execute(f'SELECT id, club_name, file_name, file_type, description, upload_time, status, reject_reason, group_id, club_token, category FROM club_uploads WHERE club_token=? AND (source="upload" OR source="activity" OR source IS NULL) AND (description LIKE "[总结]%" OR description LIKE "[照片]%" OR description LIKE "[PDF]%" OR description LIKE "[Word]%" OR description LIKE "[文件]%"){date_where} ORDER BY upload_time DESC', [token_row['token']] + date_params).fetchall()
    finally:
        conn.close()
    groups = {}
    for r in rows:
        try:
            desc = r['description'] or ''
            gid = r['group_id'] or str(r['id'])
            if gid not in groups:
                groups[gid] = {'groupId': gid, 'clubName': r['club_name'], 'time': r['upload_time'], 'status': r['status'] or 'pending', 'rejectReason': r['reject_reason'] or '', 'token': r['club_token'], 'title': '', 'body': '', 'photos': [], 'summaryId': None, 'category': r['category'] if 'category' in r.keys() else ''}
            if desc.startswith('[总结]'):
                parts = desc.replace('[总结]', '').split('|||')
                groups[gid]['title'] = parts[0] if parts else ''
                groups[gid]['body'] = parts[1] if len(parts) > 1 else ''
                groups[gid]['status'] = r['status'] or groups[gid]['status']
                groups[gid]['rejectReason'] = r['reject_reason'] or groups[gid]['rejectReason']
                groups[gid]['summaryId'] = r['id']
            elif desc.startswith('[照片]') or desc.startswith('[PDF]') or desc.startswith('[Word]') or desc.startswith('[文件]'):
                ft = r['file_type'] if 'file_type' in r.keys() else ''
                groups[gid]['photos'].append({'id': r['id'], 'fileName': r['file_name'], 'fileType': ft})
        except Exception:
            pass
    return jsonify({'success': True, 'data': list(groups.values())})


@app.route('/api/export-materials')
def export_materials():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    try:
        import zipfile
        import re as _re
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return jsonify({'error': '导出功能需要安装 python-docx 库，请运行: pip install python-docx'}), 500

        status_filter = request.args.get('status', 'approved')
        start_date = request.args.get('start_date', '')
        end_date = request.args.get('end_date', '')

        date_where = ''
        date_params = []
        if start_date:
            date_where += ' AND date(upload_time)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(upload_time)<=?'
            date_params.append(end_date)

        conn = db.get_conn()
        try:
            rows = conn.execute(
                f'SELECT id, club_name, file_name, file_path, file_type, description, upload_time, status, reject_reason, group_id, club_token, category FROM club_uploads WHERE (source="upload" OR source="activity" OR source IS NULL) AND (description LIKE "[总结]%" OR description LIKE "[照片]%" OR description LIKE "[PDF]%" OR description LIKE "[Word]%" OR description LIKE "[文件]%"){date_where} ORDER BY club_name, upload_time DESC',
                date_params
            ).fetchall()
        finally:
            conn.close()

        def strip_html(html_str):
            if not html_str:
                return ''
            text = _re.sub(r'<br\s*/?>', '\n', html_str)
            text = _re.sub(r'</p>', '\n', text)
            text = _re.sub(r'</div>', '\n', text)
            text = _re.sub(r'</h[1-6]>', '\n', text)
            text = _re.sub(r'</li>', '\n', text)
            text = _re.sub(r'<[^>]+>', '', text)
            text = text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"').replace('&#39;', "'")
            text = _re.sub(r'\n{3,}', '\n\n', text)
            return text.strip()

        groups = {}
        for r in rows:
            desc = r['description'] or ''
            gid = r['group_id'] or str(r['id'])
            if gid not in groups:
                groups[gid] = {
                    'groupId': gid, 'clubName': r['club_name'], 'time': r['upload_time'],
                    'status': r['status'] or 'pending', 'rejectReason': r['reject_reason'] or '',
                    'title': '', 'body': '', 'bodyHtml': '', 'photos': [], 'category': r['category'] or ''
                }
            if desc.startswith('[总结]'):
                parts = desc.replace('[总结]', '').split('|||')
                groups[gid]['title'] = parts[0] if parts else ''
                groups[gid]['bodyHtml'] = parts[1] if len(parts) > 1 else ''
                groups[gid]['body'] = strip_html(parts[1] if len(parts) > 1 else '')
                groups[gid]['status'] = r['status'] or groups[gid]['status']
                groups[gid]['rejectReason'] = r['reject_reason'] or groups[gid]['rejectReason']
                if r['category']:
                    groups[gid]['category'] = r['category']
            elif desc.startswith('[照片]') or desc.startswith('[PDF]') or desc.startswith('[Word]') or desc.startswith('[文件]'):
                groups[gid]['photos'].append({
                    'id': r['id'], 'fileName': r['file_name'], 'filePath': r['file_path'],
                    'fileType': r['file_type'], 'desc': _re.sub(r'^\[(照片|PDF|Word|文件)\]', '', desc)
                })
                if r['category'] and not groups[gid]['category']:
                    groups[gid]['category'] = r['category']

        if status_filter == 'approved':
            filtered = [g for g in groups.values() if g['status'] == 'approved']
            sheet_title = '已通过材料'
        elif status_filter == 'rejected':
            filtered = [g for g in groups.values() if g['status'] == 'rejected']
            sheet_title = '已打回材料'
        elif status_filter == 'pending':
            filtered = [g for g in groups.values() if g['status'] == 'pending']
            sheet_title = '待审批材料'
        else:
            filtered = list(groups.values())
            sheet_title = '全部材料'

        category_names = {'honor': '一类_荣誉奖项', 'activity': '二类_活动参与', 'innovation': '三类_创新亮点'}
        status_names = {'pending': '待审批', 'approved': '已通过', 'rejected': '已打回'}

        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            used_names = {}
            for idx, m in enumerate(filtered, 1):
                safe_club = ''.join(c for c in (m['clubName'] or '未知社团') if c.isalnum() or c in ('_', '-', ' ', '（', '）', '(', ')')) or '未知社团'
                safe_title = ''.join(c for c in m['title'] if c.isalnum() or c in ('_', '-', ' ', '（', '）', '(', ')', '、', '，', '。', '！', '？', '：', '；')) or f'材料{idx}'
                base_name = f'{idx}_{safe_club}_{safe_title}'
                if base_name in used_names:
                    used_names[base_name] += 1
                    base_name = f'{base_name}_{used_names[base_name]}'
                else:
                    used_names[base_name] = 0

                doc = Document()
                style = doc.styles['Normal']
                style.font.name = 'Microsoft YaHei'
                style.font.size = Pt(10.5)
                style.paragraph_format.space_after = Pt(4)

                doc.add_heading(m['title'] or '无标题', level=1)

                meta_table = doc.add_table(rows=4, cols=2)
                meta_table.style = 'Table Grid'
                meta_cells = [
                    ('社团名称', m['clubName']),
                    ('材料分类', category_names.get(m['category'], m['category'] or '未分类')),
                    ('提交时间', m['time']),
                    ('审批状态', status_names.get(m['status'], m['status'])),
                ]
                for ri, (label, val) in enumerate(meta_cells):
                    meta_table.cell(ri, 0).text = label
                    if meta_table.cell(ri, 0).paragraphs[0].runs:
                        meta_table.cell(ri, 0).paragraphs[0].runs[0].bold = True
                    meta_table.cell(ri, 1).text = str(val or '')
                if m['rejectReason']:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    run = p.add_run('驳回原因：' + m['rejectReason'])
                    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

                doc.add_heading('正文内容', level=2)
                body_text = m['body'] or '无正文内容'
                for para_text in body_text.split('\n'):
                    para_text = para_text.strip()
                    if para_text:
                        doc.add_paragraph(para_text)

                image_photos = [p for p in m['photos'] if p['fileType'] and p['fileType'].lower() in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp', 'image')]
                if image_photos:
                    doc.add_heading('活动照片', level=2)
                    for p in image_photos:
                        fp = p['filePath']
                        if fp and os.path.exists(fp):
                            try:
                                doc.add_picture(fp, width=Inches(4.5))
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception:
                                doc.add_paragraph(f'[图片加载失败: {p["fileName"]}]')

                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                zf.writestr(f'{base_name}.docx', doc_buffer.getvalue())

                all_attach = m['photos']
                if all_attach:
                    attach_count = {}
                    for p in all_attach:
                        fp = p['filePath']
                        if fp and os.path.exists(fp):
                            ext = os.path.splitext(p['fileName'])[1] or os.path.splitext(fp)[1]
                            safe_file_name = ''.join(c for c in os.path.splitext(p['fileName'])[0] if c.isalnum() or c in ('_', '-', ' ')) or f'file_{p["id"]}'
                            if safe_file_name in attach_count:
                                attach_count[safe_file_name] += 1
                                safe_file_name = f'{safe_file_name}_{attach_count[safe_file_name]}'
                            else:
                                attach_count[safe_file_name] = 0
                            arc_name = f'上传文件/{base_name}/{safe_file_name}{ext}'
                            try:
                                with open(fp, 'rb') as fobj:
                                    zf.writestr(arc_name, fobj.read())
                            except Exception:
                                pass

            wb = Workbook()
            ws = wb.active
            ws.title = sheet_title

            title_font = Font(name='微软雅黑', bold=True, size=14, color='1a1a2e')
            header_font = Font(name='微软雅黑', bold=True, size=10, color='FFFFFF')
            header_fill = PatternFill(start_color='4361EE', end_color='4361EE', fill_type='solid')
            header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell_font = Font(name='微软雅黑', size=9)
            cell_align = Alignment(vertical='center', wrap_text=True)
            center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            thin_side = Side(style='thin', color='C0C0C0')
            thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
            approved_fill = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
            rejected_fill = PatternFill(start_color='FFE5E5', end_color='FFE5E5', fill_type='solid')
            pending_fill = PatternFill(start_color='FFF8E1', end_color='FFF8E1', fill_type='solid')

            headers = ['序号', '社团名称', '材料分类', '活动标题', '正文内容', '附件数量', '提交时间', '审批状态', '驳回原因']
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
            title_cell = ws.cell(row=1, column=1, value=sheet_title)
            title_cell.font = title_font
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 32

            for i, h in enumerate(headers, 1):
                cell = ws.cell(row=2, column=i, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                cell.border = thin_border
            ws.row_dimensions[2].height = 26

            for idx, m in enumerate(filtered, 1):
                row_idx = idx + 2
                cat_display = category_names.get(m['category'], m['category'] or '未分类')
                values = [idx, m['clubName'], cat_display, m['title'], m['body'], len(m['photos']), m['time'], status_names.get(m['status'], m['status']), m['rejectReason']]
                for col, v in enumerate(values, 1):
                    cell = ws.cell(row=row_idx, column=col, value=v)
                    cell.font = cell_font
                    cell.border = thin_border
                    if col in (1, 6):
                        cell.alignment = center_align
                    else:
                        cell.alignment = cell_align
                    if m['status'] == 'approved':
                        cell.fill = approved_fill
                    elif m['status'] == 'rejected':
                        cell.fill = rejected_fill
                    else:
                        cell.fill = pending_fill
                ws.row_dimensions[row_idx].height = 22

            col_widths = [6, 22, 14, 20, 40, 10, 18, 10, 24]
            for i, w in enumerate(col_widths, 1):
                ws.column_dimensions[get_column_letter(i)].width = w
            ws.freeze_panes = 'A3'

            xlsx_buffer = BytesIO()
            wb.save(xlsx_buffer)
            xlsx_buffer.seek(0)
            zf.writestr('材料汇总统计.xlsx', xlsx_buffer.getvalue())

        buffer.seek(0)
        filename = f'{sheet_title}_完整导出_{datetime.now().strftime("%Y%m%d")}.zip'
        return send_file(buffer, mimetype='application/zip',
                        as_attachment=True, download_name=filename)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/base-data', methods=['POST'])
def add_base_data():
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400
    file = request.files['file']
    name = request.form.get('name', '') or os.path.splitext(file.filename)[0]
    try:
        df, headers, data = smart_parse_excel(file)
        if df is None:
            return jsonify({'error': '文件解析失败，无法识别有效表头'}), 400
        if len([h for h in headers if not h.startswith('列')]) < 2:
            return jsonify({'error': '未能识别有效列名，请检查文件格式'}), 400
    except Exception as e:
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 400

    file_uuid = str(uuid.uuid4())
    conn = db.get_conn()
    try:
        conn.execute('INSERT INTO base_data (uuid, name, original_name, headers, row_count, data_json) VALUES (?, ?, ?, ?, ?, ?)',
                     (file_uuid, name, file.filename, json.dumps(headers, ensure_ascii=False), len(data), json.dumps(data, ensure_ascii=False)))
        conn.commit()
    finally:
        conn.close()
    cache.clear()
    return jsonify({'success': True, 'uuid': file_uuid, 'name': name, 'rowCount': len(data), 'headers': headers})


@app.route('/api/base-data/<uuid_val>', methods=['PUT'])
def update_base_data(uuid_val):
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400
    file = request.files['file']
    name = request.form.get('name', '')
    try:
        df, headers, data = smart_parse_excel(file)
        if df is None:
            return jsonify({'error': '文件解析失败'}), 400
    except Exception as e:
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 400
    conn = db.get_conn()
    try:
        if name:
            conn.execute('UPDATE base_data SET name=?, original_name=?, headers=?, row_count=?, data_json=? WHERE uuid=?',
                         (name, file.filename, json.dumps(headers, ensure_ascii=False), len(data), json.dumps(data, ensure_ascii=False), uuid_val))
        else:
            conn.execute('UPDATE base_data SET original_name=?, headers=?, row_count=?, data_json=? WHERE uuid=?',
                         (file.filename, json.dumps(headers, ensure_ascii=False), len(data), json.dumps(data, ensure_ascii=False), uuid_val))
        conn.commit()
    finally:
        conn.close()
    cache.clear()
    return jsonify({'success': True})


@app.route('/api/base-data/<uuid_val>', methods=['DELETE'])
def delete_base_data(uuid_val):
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM base_data WHERE uuid=?', (uuid_val,))
        conn.commit()
    finally:
        conn.close()
    cache.clear()
    return jsonify({'success': True})


@app.route('/api/base-data/<uuid_val>/match', methods=['POST'])
def match_base_data(uuid_val):
    club_name = request.json.get('clubName', '')
    match_col = request.json.get('matchColumn', '')
    if not club_name or not match_col:
        return jsonify({'success': True, 'data': {}})
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT data_json, headers FROM base_data WHERE uuid=?', (uuid_val,)).fetchone()
    finally:
        conn.close()
    if not row:
        return jsonify({'success': True, 'data': {}})
    data = json.loads(row['data_json'])
    headers = json.loads(row['headers'])
    for r in data:
        if str(r.get(match_col, '')).strip() == club_name:
            info = {}
            for h in headers:
                if h != match_col:
                    v = str(r.get(h, '')).strip()
                    if v and v not in ('undefined', 'null', 'NaN'):
                        info[h] = v
            return jsonify({'success': True, 'data': info})
    return jsonify({'success': True, 'data': {}})


@app.route('/api/stats/club')
def get_club_stats():
    cached = cache.get('club_stats')
    if cached:
        return jsonify({'success': True, 'data': cached, 'source': 'cache'})
    clubs = compute_club_stats()
    cache.set('club_stats', clubs, ttl=30)
    return jsonify({'success': True, 'data': clubs, 'source': 'db'})


@app.route('/api/stats/summary')
def get_summary():
    cached = cache.get('summary')
    if cached:
        return jsonify({'success': True, 'data': cached})
    conn = db.get_conn()
    try:
        act_total = conn.execute('SELECT COUNT(*) as c FROM activity_records').fetchone()['c']
        gui_total = conn.execute('SELECT COUNT(DISTINCT session_id) as c FROM guidance_records').fetchone()['c']
        club_cnt = conn.execute('SELECT COUNT(DISTINCT club_name) as c FROM activity_records').fetchone()['c']
        most_active = conn.execute('SELECT COUNT(*) as c FROM activity_records GROUP BY club_name ORDER BY c DESC LIMIT 1').fetchone()
        low_cnt = conn.execute('SELECT COUNT(*) as c FROM (SELECT club_name, COUNT(*) as cnt FROM activity_records GROUP BY club_name HAVING cnt < 2)').fetchone()
    finally:
        conn.close()

    base_clubs = get_all_base_clubs()
    existing_norms = set()
    conn = db.get_conn()
    try:
        for r in conn.execute('SELECT DISTINCT club_name FROM activity_records').fetchall():
            existing_norms.add(normalize_name(r['club_name']))
        for r in conn.execute('SELECT DISTINCT club_name FROM guidance_records').fetchall():
            existing_norms.add(normalize_name(r['club_name']))
    finally:
        conn.close()
    base_only_count = 0
    for norm in base_clubs:
        if norm not in existing_norms and not find_similar_club(norm, existing_norms):
            base_only_count += 1

    club_cnt = club_cnt + base_only_count

    summary = {
        'totalActivities': act_total,
        'totalGuidances': gui_total,
        'clubCount': club_cnt,
        'mostActive': most_active['c'] if most_active else 0,
        'lowCount': low_cnt['c'] if low_cnt else 0
    }
    cache.set('summary', summary, ttl=30)
    return jsonify({'success': True, 'data': summary})


@app.route('/api/stats/partitions')
def get_partitions():
    clubs = compute_club_stats()
    return jsonify({'success': True, 'data': {'all': clubs, 'normal': [c for c in clubs if c['activityCount'] >= 2], 'low': [c for c in clubs if c['activityCount'] < 2], 'high': [c for c in clubs if c['activityCount'] >= 5]}})


@app.route('/api/descriptive')
def get_descriptive():
    cached = cache.get('descriptive')
    if cached:
        return jsonify({'success': True, 'data': cached})
    conn = db.get_conn()
    try:
        act_counts = [r['cnt'] for r in conn.execute('SELECT COUNT(*) as cnt FROM activity_records GROUP BY club_name').fetchall()]
        gui_counts = [r['cnt'] for r in conn.execute('SELECT COUNT(DISTINCT session_id) as cnt FROM guidance_records GROUP BY club_name').fetchall()]
    finally:
        conn.close()
    result = {'activities': StatsService.descriptive_stats(act_counts), 'guidances': StatsService.descriptive_stats(gui_counts)}
    cache.set('descriptive', result, ttl=30)
    return jsonify({'success': True, 'data': result})


@app.route('/api/export')
def export_excel():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers
        from openpyxl.utils import get_column_letter
        wb = Workbook()

        clubs = compute_club_stats()

        include_online = request.args.get('online', '1') == '1'
        sheets_param = request.args.get('sheets', '')
        selected_sheets = [s.strip() for s in sheets_param.split(',') if s.strip()] if sheets_param else []
        all_sheet_ids = ['1', '2', '3', '4', '5', '6']
        if selected_sheets:
            active_sheets = [s for s in selected_sheets if s in all_sheet_ids]
        else:
            active_sheets = all_sheet_ids

        title_font = Font(name='微软雅黑', bold=True, size=16, color='1a1a2e')
        subtitle_font = Font(name='微软雅黑', bold=True, size=10, color='666666')
        header_font = Font(name='微软雅黑', bold=True, size=10, color='FFFFFF')
        header_fill = PatternFill(start_color='4361EE', end_color='4361EE', fill_type='solid')
        header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell_font = Font(name='微软雅黑', size=9)
        cell_align = Alignment(vertical='center', wrap_text=True)
        center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        num_align = Alignment(horizontal='center', vertical='center')

        thin_side = Side(style='thin', color='C0C0C0')
        thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

        warn_fill = PatternFill(start_color='FFE5E5', end_color='FFE5E5', fill_type='solid')
        ok_fill = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
        alt_fill = PatternFill(start_color='F8F9FA', end_color='F8F9FA', fill_type='solid')
        white_fill = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')

        rank_gold = Font(name='微软雅黑', bold=True, size=10, color='FFD700')
        rank_silver = Font(name='微软雅黑', bold=True, size=10, color='C0C0C0')
        rank_bronze = Font(name='微软雅黑', bold=True, size=10, color='CD7F32')
        rank_normal = Font(name='微软雅黑', size=9, color='666666')

        def write_title(ws, title, subtitle, col_count):
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
            title_cell = ws.cell(row=1, column=1, value=title)
            title_cell.font = title_font
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 36

        def write_headers(ws, row, headers):
            ws.row_dimensions[row].height = 28
            for i, h in enumerate(headers, 1):
                cell = ws.cell(row=row, column=i, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                cell.border = thin_border

        def write_data_row(ws, row, values, col_count, is_warn=False, is_alt=False):
            ws.row_dimensions[row].height = 22
            for i, v in enumerate(values, 1):
                cell = ws.cell(row=row, column=i, value=v)
                cell.font = cell_font
                cell.border = thin_border
                if i == 1:
                    cell.alignment = num_align
                    if isinstance(v, int) and v <= 3:
                        cell.font = [rank_normal, rank_gold, rank_silver, rank_bronze][v]
                    else:
                        cell.font = rank_normal
                elif isinstance(v, (int, float)):
                    cell.alignment = num_align
                    cell.font = Font(name='微软雅黑', size=9, bold=True)
                else:
                    cell.alignment = cell_align
                if is_warn:
                    cell.fill = warn_fill
                elif is_alt:
                    cell.fill = alt_fill
                else:
                    cell.fill = white_fill

        now_str = datetime.now().strftime('%Y年%m月%d日')

        # ==== Sheet 1: 社团活动统计总表 ====
        ws1 = wb.active
        ws1.title = '社团活动统计'
        if '1' in active_sheets:
            if include_online:
                s1_headers = ['排名', '社团名称', '社团活动情况', '活动次数', '指导老师', '业务指导单位']
                s1_cols = 6
            else:
                s1_headers = ['排名', '社团名称', '活动次数', '指导老师', '业务指导单位']
                s1_cols = 5

            write_title(ws1, '社团活动统计总表', f'统计日期：{now_str}    数据来源：社团活动记录 + 指导记录 + 线上爬取数据', s1_cols)
            write_headers(ws1, 2, s1_headers)

            for i, c in enumerate(clubs, 1):
                teachers = get_teacher_name_from_sources(c['club']) or '无'
                unit = get_unit_from_base(c['club']) or '—'

                if include_online:
                    online_status = get_online_activity_status(c['club'])
                    activity_status = '\n'.join(online_status['details']) if online_status['details'] else '—'
                    row_data = [i, c['club'], activity_status, c['activityCount'], teachers, unit]
                else:
                    row_data = [i, c['club'], c['activityCount'], teachers, unit]

                row_idx = i + 2
                ws1.append(row_data)
                write_data_row(ws1, row_idx, row_data, s1_cols, is_warn=(c['activityCount'] < 2), is_alt=(i % 2 == 0))
                if c['activityCount'] < 2:
                    for col in range(1, s1_cols + 1):
                        ws1.cell(row=row_idx, column=col).fill = warn_fill

            col_widths_1 = [6, 20, 32, 10, 28, 20] if include_online else [6, 20, 10, 28, 20]
            for i, w in enumerate(col_widths_1, 1):
                ws1.column_dimensions[get_column_letter(i)].width = w
            ws1.freeze_panes = 'A3'

        # ==== Sheet 2: 社团指导老师未开展指导统计表 ====
        if '2' in active_sheets:
            ws2 = wb.create_sheet('指导老师未开展指导统计')
            s2_headers = ['序号', '社团名称', '业务指导单位', '指导老师姓名']
            s2_cols = 4
            write_title(ws2, '社团指导老师未开展指导统计表', f'统计日期：{now_str}    筛选条件：指导次数 = 0', s2_cols)
            write_headers(ws2, 2, s2_headers)

            no_guidance_clubs = [c for c in clubs if c['guidanceCount'] == 0]
            row_num = 0
            for c in no_guidance_clubs:
                teacher_name = get_teacher_name_from_sources(c['club'])
                if teacher_name:
                    row_num += 1
                    unit = get_unit_from_base(c['club']) or '—'
                    row_data = [row_num, c['club'], unit, teacher_name]
                    ws2.append(row_data)
                    write_data_row(ws2, row_num + 2, row_data, s2_cols, is_alt=(row_num % 2 == 0))

            if row_num == 0:
                ws2.append(['—', '无数据', '—', '—'])
                write_data_row(ws2, 3, ['—', '无数据', '—', '—'], s2_cols)

            for i, w in enumerate([6, 22, 22, 20], 1):
                ws2.column_dimensions[get_column_letter(i)].width = w
            ws2.freeze_panes = 'A3'

        # ==== Sheet 3: 社团活动次数小于两次统计表 ====
        if '3' in active_sheets:
            ws3 = wb.create_sheet('活动不足2次社团统计')
            s3_headers = ['序号', '社团名称', '业务指导单位', '指导老师', '活动次数']
            s3_cols = 5
            write_title(ws3, '社团活动次数不足2次统计表', f'统计日期：{now_str}    筛选条件：活动次数 < 2', s3_cols)
            write_headers(ws3, 2, s3_headers)

            low_clubs = [c for c in clubs if c['activityCount'] < 2]
            for i, c in enumerate(low_clubs, 1):
                teacher_name = get_teacher_name_from_sources(c['club'])
                unit = get_unit_from_base(c['club']) or '—'
                row_data = [i, c['club'], unit, teacher_name or '无', c['activityCount']]
                ws3.append(row_data)
                write_data_row(ws3, i + 2, row_data, s3_cols, is_warn=True)

            if not low_clubs:
                ws3.append(['—', '无数据', '—', '—', '—'])
                write_data_row(ws3, 3, ['—', '无数据', '—', '—', '—'], s3_cols, is_warn=True)

            for i, w in enumerate([6, 22, 22, 20, 10], 1):
                ws3.column_dimensions[get_column_letter(i)].width = w
            ws3.freeze_panes = 'A3'

        # ==== Sheet 4: 社团指导老师指导次数统计 ====
        if '4' in active_sheets:
            ws4 = wb.create_sheet('指导老师指导次数统计')
            s4_headers = ['序号', '社团名称', '指导老师', '指导次数']
            s4_cols = 4
            write_title(ws4, '社团指导老师指导次数统计表', f'统计日期：{now_str}    数据来源：指导老师指导记录', s4_cols)
            write_headers(ws4, 2, s4_headers)

            row_num = 0
            for c in clubs:
                teacher_dict = c.get('teachers', {})
                if not teacher_dict:
                    continue
                for t_name, t_cnt in teacher_dict.items():
                    row_num += 1
                    row_data = [row_num, c['club'], t_name, t_cnt]
                    ws4.append(row_data)
                    write_data_row(ws4, row_num + 2, row_data, s4_cols, is_alt=(row_num % 2 == 0))

            if row_num == 0:
                ws4.append(['—', '无数据', '—', '—'])
                write_data_row(ws4, 3, ['—', '无数据', '—', '—'], s4_cols)

            for i, w in enumerate([6, 22, 18, 10], 1):
                ws4.column_dimensions[get_column_letter(i)].width = w
            ws4.freeze_panes = 'A3'

        # ==== Sheet 5: 指导次数≤1次的老师统计 ====
        if '5' in active_sheets:
            ws5 = wb.create_sheet('指导次数不足老师统计')
            s5_headers = ['序号', '指导老师', '指导次数', '所属社团']
            s5_cols = 4
            write_title(ws5, '指导次数不足老师统计表', f'统计日期：{now_str}    筛选条件：指导次数 ≤ 1（含未开展指导）', s5_cols)
            write_headers(ws5, 2, s5_headers)

            low_teachers = []
            for c in clubs:
                teacher_dict = c.get('teachers', {})
                for t_name, t_cnt in teacher_dict.items():
                    if t_cnt <= 1:
                        low_teachers.append({'name': t_name, 'count': t_cnt, 'club': c['club']})

            for c in clubs:
                if c['guidanceCount'] == 0:
                    base_teachers = get_teacher_name_from_sources(c['club'])
                    if base_teachers:
                        for t in base_teachers.replace('，', ',').split(','):
                            t = t.strip()
                            if t:
                                already = any(lt['name'] == t and lt['club'] == c['club'] for lt in low_teachers)
                                if not already:
                                    low_teachers.append({'name': t, 'count': 0, 'club': c['club']})

            low_teachers.sort(key=lambda x: (x['count'], x['name']))

            for i, t in enumerate(low_teachers, 1):
                row_data = [i, t['name'], t['count'], t['club']]
                ws5.append(row_data)
                write_data_row(ws5, i + 2, row_data, s5_cols, is_warn=True, is_alt=(i % 2 == 0))

            if not low_teachers:
                ws5.append(['—', '无数据', '—', '—'])
                write_data_row(ws5, 3, ['—', '无数据', '—', '—'], s5_cols, is_warn=True)

            for i, w in enumerate([6, 18, 10, 22], 1):
                ws5.column_dimensions[get_column_letter(i)].width = w
            ws5.freeze_panes = 'A3'

        # ==== Sheet 6: 线上数据详情 ====
        if include_online and '6' in active_sheets:
            ws6 = wb.create_sheet('线上活动数据详情')
            conn = db.get_conn()
            try:
                online_rows = conn.execute('SELECT club_name, has_photo, has_summary, activity_date, activity_title FROM online_activity_data ORDER BY club_name, activity_date').fetchall()
            finally:
                conn.close()
            s6_headers = ['序号', '社团名称', '活动时间', '活动主题', '是否上传照片', '是否上传总结']
            s6_cols = 6
            write_title(ws6, '线上活动数据详情', f'数据来源：学校官网API爬取    共{len(online_rows)}条记录', s6_cols)
            write_headers(ws6, 2, s6_headers)

            for i, r in enumerate(online_rows, 1):
                row_data = [i, r['club_name'], r['activity_date'] or '—', r['activity_title'] or '—',
                           '✓ 是' if r['has_photo'] else '✗ 否', '✓ 是' if r['has_summary'] else '✗ 否']
                ws6.append(row_data)
                write_data_row(ws6, i + 2, row_data, s6_cols, is_alt=(i % 2 == 0))
                photo_cell = ws6.cell(row=i + 2, column=5)
                summary_cell = ws6.cell(row=i + 2, column=6)
                if r['has_photo']:
                    photo_cell.font = Font(name='微软雅黑', size=9, color='2E7D32')
                else:
                    photo_cell.font = Font(name='微软雅黑', size=9, color='C62828')
                if r['has_summary']:
                    summary_cell.font = Font(name='微软雅黑', size=9, color='2E7D32')
                else:
                    summary_cell.font = Font(name='微软雅黑', size=9, color='C62828')

            for i, w in enumerate([6, 22, 18, 30, 14, 14], 1):
                ws6.column_dimensions[get_column_letter(i)].width = w
            ws6.freeze_panes = 'A3'

        # Print settings
        for ws in wb.worksheets:
            ws.sheet_properties.pageSetUpPr = None
            ws.page_setup.orientation = 'landscape'
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        filename = f'社团活动统计结果_{datetime.now().strftime("%Y%m%d")}.xlsx'
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        as_attachment=True, download_name=filename)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/reset', methods=['POST'])
def reset_data():
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM activity_records')
        conn.execute('DELETE FROM guidance_records')
        conn.execute('DELETE FROM uploaded_files')
        conn.execute('DELETE FROM online_activity_data')
        conn.commit()
    finally:
        conn.close()
    cache.clear()
    return jsonify({'success': True})


@app.route('/api/notices')
def get_notices():
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, title, content, is_top, attachment_path, attachment_name, created_at FROM notices ORDER BY is_top DESC, created_at DESC').fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'title': r['title'], 'content': r['content'], 'isTop': r['is_top'], 'attachmentName': r['attachment_name'] or '', 'time': local_time(r['created_at'])} for r in rows]})


NOTICE_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'notices')
os.makedirs(NOTICE_FOLDER, exist_ok=True)


@app.route('/api/notices', methods=['POST'])
def add_notice():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    title = request.form.get('title', '').strip()
    content = request.form.get('content', '').strip()
    is_top = 1 if request.form.get('isTop') else 0
    if not title:
        return jsonify({'error': '请输入通告标题'}), 400
    attachment_path = ''
    attachment_name = ''
    if 'attachment' in request.files:
        att = request.files['attachment']
        if att.filename:
            attachment_name = att.filename
            ext = os.path.splitext(att.filename)[1]
            safe_name = str(uuid.uuid4())[:16] + ext
            key = 'notices/' + safe_name
            storage.save(att, key)
            attachment_path = key
    conn = db.get_conn()
    try:
        conn.execute('INSERT INTO notices (title, content, is_top, attachment_path, attachment_name) VALUES (?, ?, ?, ?, ?)', (title, content, is_top, attachment_path, attachment_name))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/notice-file/<int:nid>')
def get_notice_file(nid):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT attachment_path, attachment_name FROM notices WHERE id=?', (nid,)).fetchone()
    finally:
        conn.close()
    if not row or not row['attachment_path']:
        return jsonify({'error': '文件不存在'}), 404
    path = storage.get_path(row['attachment_path'])
    if not path:
        url = storage.get_url(row['attachment_path'])
        if url:
            return redirect(url)
        return jsonify({'error': '文件不存在'}), 404
    return send_file(path, as_attachment=True, download_name=row['attachment_name'])


@app.route('/api/notices/<int:nid>', methods=['DELETE'])
def delete_notice(nid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM notices WHERE id=?', (nid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/notices/<int:nid>', methods=['PUT'])
def update_notice(nid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    title = data.get('title', '').strip()
    content = data.get('content', '').strip()
    is_top = 1 if data.get('isTop') else 0
    if not title:
        return jsonify({'error': '标题不能为空'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM notices WHERE id=?', (nid,)).fetchone()
        if not existing:
            return jsonify({'error': '通知不存在'}), 404
        conn.execute('UPDATE notices SET title=?, content=?, is_top=? WHERE id=?', (title, content, is_top, nid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


SHOWCASE_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'showcase')
os.makedirs(SHOWCASE_FOLDER, exist_ok=True)


@app.route('/api/showcase')
def get_showcase():
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, club_name, description, image_path, created_at FROM club_showcase ORDER BY created_at DESC').fetchall()
    finally:
        conn.close()
    result = []
    for r in rows:
        item = {'id': r['id'], 'clubName': r['club_name'], 'description': r['description'], 'time': local_time(r['created_at']), 'imageUrl': ''}
        if r['image_path'] and os.path.exists(r['image_path']):
            item['imageUrl'] = '/api/showcase-image/' + str(r['id'])
        result.append(item)
    return jsonify({'success': True, 'data': result})


@app.route('/api/showcase-image/<int:sid>')
def get_showcase_image(sid):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT image_path FROM club_showcase WHERE id=?', (sid,)).fetchone()
    finally:
        conn.close()
    if not row or not row['image_path'] or not os.path.exists(row['image_path']):
        return jsonify({'error': '图片不存在'}), 404
    return send_file(row['image_path'])


@app.route('/api/showcase', methods=['POST'])
def add_showcase():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    club_name = request.form.get('clubName', '').strip()
    description = request.form.get('description', '').strip()
    if not club_name:
        return jsonify({'error': '请输入社团名称'}), 400
    image_path = ''
    if 'image' in request.files:
        img = request.files['image']
        if img.filename:
            ext = os.path.splitext(img.filename)[1]
            safe_name = str(uuid.uuid4())[:16] + ext
            image_path = os.path.join(SHOWCASE_FOLDER, safe_name)
            img.save(image_path)
    conn = db.get_conn()
    try:
        conn.execute('INSERT INTO club_showcase (club_name, description, image_path) VALUES (?, ?, ?)', (club_name, description, image_path))
        leaders = conn.execute('SELECT id FROM users WHERE club_name=? AND role="user"', (club_name,)).fetchall()
        for l in leaders:
            send_notification(l['id'], '🖼️ 社团风采已发布', '您的社团「' + club_name + '」的风采已发布为首页轮播图', 'showcase', '/dashboard.html', conn=conn)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/showcase/<int:sid>', methods=['DELETE'])
def delete_showcase(sid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT image_path FROM club_showcase WHERE id=?', (sid,)).fetchone()
        if row and row['image_path'] and os.path.exists(row['image_path']):
            os.remove(row['image_path'])
        conn.execute('DELETE FROM club_showcase WHERE id=?', (sid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/showcase/<int:sid>', methods=['PUT'])
def update_showcase(sid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    club_name = data.get('clubName', '').strip()
    description = data.get('description', '').strip()
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM club_showcase WHERE id=?', (sid,)).fetchone()
        if not existing:
            return jsonify({'error': '风采不存在'}), 404
        if club_name:
            conn.execute('UPDATE club_showcase SET club_name=?, description=? WHERE id=?', (club_name, description, sid))
        else:
            conn.execute('UPDATE club_showcase SET description=? WHERE id=?', (description, sid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/showcase-image/<int:sid>', methods=['POST'])
def update_showcase_image(sid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    if 'image' not in request.files:
        return jsonify({'error': '请选择图片'}), 400
    img = request.files['image']
    if not img.filename:
        return jsonify({'error': '请选择图片'}), 400
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT image_path FROM club_showcase WHERE id=?', (sid,)).fetchone()
        if not row:
            return jsonify({'error': '风采不存在'}), 404
        if row['image_path'] and os.path.exists(row['image_path']):
            os.remove(row['image_path'])
        ext = os.path.splitext(img.filename)[1]
        safe_name = str(uuid.uuid4())[:16] + ext
        image_path = os.path.join(SHOWCASE_FOLDER, safe_name)
        img.save(image_path)
        conn.execute('UPDATE club_showcase SET image_path=? WHERE id=?', (image_path, sid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/showcase-materials')
def get_showcase_materials():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT group_id, club_name, description, status, club_token FROM club_uploads WHERE status="approved" AND group_id IS NOT NULL AND group_id!="" GROUP BY group_id ORDER BY MAX(upload_time) DESC').fetchall()
        result = []
        for r in rows:
            gid = r['group_id']
            files = conn.execute('SELECT id, file_name, file_type, description FROM club_uploads WHERE group_id=?', (gid,)).fetchall()
            title = ''
            has_photo = False
            for f in files:
                desc = f['description'] or ''
                if desc.startswith('[总结]'):
                    parts = desc.replace('[总结]', '').split('|||')
                    title = parts[0] if parts else ''
                if f['file_type'] and f['file_type'].startswith('image'):
                    has_photo = True
            result.append({
                'groupId': gid,
                'clubName': r['club_name'],
                'title': title,
                'description': '',
                'fileCount': len(files),
                'hasPhoto': has_photo,
                'token': r['club_token']
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/showcase-material-detail/<group_id>')
def get_showcase_material_detail(group_id):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, club_name, file_name, file_type, description, club_token FROM club_uploads WHERE group_id=?', (group_id,)).fetchall()
    finally:
        conn.close()
    if not rows:
        return jsonify({'error': '材料不存在'}), 404
    title = ''
    body = ''
    photos = []
    token = rows[0]['club_token']
    club_name = rows[0]['club_name']
    for r in rows:
        desc = r['description'] or ''
        if desc.startswith('[总结]'):
            parts = desc.replace('[总结]', '').split('|||')
            title = parts[0] if parts else ''
            body = parts[1] if len(parts) > 1 else ''
        elif desc.startswith('[照片]'):
            photos.append({'id': r['id'], 'name': r['file_name']})
        if r['file_type'] and r['file_type'].startswith('image') and not desc.startswith('[照片]'):
            photos.append({'id': r['id'], 'name': r['file_name']})
    return jsonify({'success': True, 'data': {'groupId': group_id, 'clubName': club_name, 'title': title, 'body': body, 'photos': photos, 'token': token}})


@app.route('/api/showcase-from-material/<group_id>', methods=['POST'])
def showcase_from_material(group_id):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, club_name, file_name, description, club_token FROM club_uploads WHERE group_id=?', (group_id,)).fetchall()
        if not rows:
            return jsonify({'error': '材料不存在'}), 404
        club_name = rows[0]['club_name']
        title = ''
        body = ''
        image_path = ''
        for r in rows:
            desc = r['description'] or ''
            if desc.startswith('[总结]'):
                parts = desc.replace('[总结]', '').split('|||')
                title = parts[0] if parts else ''
                body = parts[1] if len(parts) > 1 else ''
            elif desc.startswith('[照片]') and not image_path:
                file_row = conn.execute('SELECT file_path FROM club_uploads WHERE id=?', (r['id'],)).fetchone()
                if file_row and file_row['file_path']:
                    actual_path = storage.get_path(file_row['file_path'])
                    if actual_path:
                        ext = os.path.splitext(file_row['file_path'])[1]
                        safe_name = str(uuid.uuid4())[:16] + ext
                        dest = os.path.join(SHOWCASE_FOLDER, safe_name)
                        import shutil
                        shutil.copy2(actual_path, dest)
                        image_path = dest
        description = body or title
        existing = conn.execute('SELECT id FROM club_showcase WHERE club_name=? AND description=?', (club_name, description)).fetchone()
        if existing:
            return jsonify({'error': '该社团风采已存在'}), 400
        conn.execute('INSERT INTO club_showcase (club_name, description, image_path) VALUES (?, ?, ?)', (club_name, description, image_path))
        leaders = conn.execute('SELECT id FROM users WHERE club_name=? AND role="user"', (club_name,)).fetchall()
        for l in leaders:
            send_notification(l['id'], '🖼️ 材料发布为轮播图', '您的社团「' + club_name + '」的活动材料「' + (title or '活动') + '」已被发布为首页轮播图', 'showcase', '/dashboard.html', conn=conn)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/active-clubs')
def get_active_clubs():
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        date_where = ''
        date_params = []
        if start_date:
            date_where += ' AND date(activity_date)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(activity_date)<=?'
            date_params.append(end_date)
        activities = conn.execute(f'SELECT club_name, COUNT(*) as cnt FROM activity_records WHERE 1=1{date_where} GROUP BY club_name', date_params).fetchall()
    finally:
        conn.close()
    result = []
    for c in sorted(activities, key=lambda x: -x['cnt']):
        if c['cnt'] > 0:
            result.append({'club': c['club_name'], 'activityCount': c['cnt']})
    return jsonify({'success': True, 'data': result})


@app.route('/api/low-activity-clubs')
def get_low_activity_clubs():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    try:
        start_date = request.args.get('start_date', '').strip()
        end_date = request.args.get('end_date', '').strip()
        threshold = int(request.args.get('threshold', 2))
        conn = db.get_conn()
        try:
            date_where = ''
            date_params = []
            if start_date:
                date_where += ' AND date(created_at)>=?'
                date_params.append(start_date)
            if end_date:
                date_where += ' AND date(created_at)<=?'
                date_params.append(end_date)
            activity_counts = conn.execute(
                f'SELECT club_name, COUNT(*) as cnt FROM checkin_sessions WHERE 1=1{date_where} GROUP BY club_name',
                date_params
            ).fetchall()
            all_clubs = conn.execute('SELECT club_name FROM club_profiles ORDER BY club_name').fetchall()
            teacher_rows = conn.execute('SELECT club_name, teacher_name FROM club_teachers').fetchall()
            try:
                profile_rows = conn.execute('SELECT club_name, guiding_unit, president, category FROM club_profiles').fetchall()
            except Exception:
                profile_rows = []
        finally:
            conn.close()
        act_map = {r['club_name']: r['cnt'] for r in activity_counts}
        teacher_map = {}
        for r in teacher_rows:
            teacher_map.setdefault(r['club_name'], set()).add(r['teacher_name'])
        profile_map = {}
        for r in profile_rows:
            profile_map[r['club_name']] = {
                'guidingUnit': r['guiding_unit'] if 'guiding_unit' in r.keys() else '',
                'president': r['president'] if 'president' in r.keys() else '',
                'category': r['category'] if 'category' in r.keys() else ''
            }
        result = []
        for c in all_clubs:
            cn = c['club_name']
            cnt = act_map.get(cn, 0)
            if cnt < threshold:
                p = profile_map.get(cn, {'guidingUnit': '', 'president': '', 'category': ''})
                tnames = '、'.join(sorted(teacher_map.get(cn, set()))) if cn in teacher_map else ''
                result.append({
                    'clubName': cn,
                    'activityCount': cnt,
                    'guidingUnit': p.get('guidingUnit', ''),
                    'president': p.get('president', ''),
                    'category': p.get('category', ''),
                    'teacherNames': tnames
                })
        result.sort(key=lambda x: x['activityCount'])
        return jsonify({'success': True, 'data': result, 'threshold': threshold, 'totalClubs': len(all_clubs)})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/export-low-activity-clubs')
def export_low_activity_clubs():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        return jsonify({'error': '导出功能需要安装 openpyxl 库'}), 500
    try:
        start_date = request.args.get('start_date', '').strip()
        end_date = request.args.get('end_date', '').strip()
        threshold = int(request.args.get('threshold', 2))
        conn = db.get_conn()
        try:
            date_where = ''
            date_params = []
            if start_date:
                date_where += ' AND date(created_at)>=?'
                date_params.append(start_date)
            if end_date:
                date_where += ' AND date(created_at)<=?'
                date_params.append(end_date)
            activity_counts = conn.execute(
                f'SELECT club_name, COUNT(*) as cnt FROM checkin_sessions WHERE 1=1{date_where} GROUP BY club_name',
                date_params
            ).fetchall()
            all_clubs = conn.execute('SELECT club_name FROM club_profiles ORDER BY club_name').fetchall()
            teacher_rows = conn.execute('SELECT club_name, teacher_name FROM club_teachers').fetchall()
            try:
                profile_rows = conn.execute('SELECT club_name, guiding_unit, president, category FROM club_profiles').fetchall()
            except Exception:
                profile_rows = []
        finally:
            conn.close()
        act_map = {r['club_name']: r['cnt'] for r in activity_counts}
        teacher_map = {}
        for r in teacher_rows:
            teacher_map.setdefault(r['club_name'], set()).add(r['teacher_name'])
        profile_map = {}
        for r in profile_rows:
            profile_map[r['club_name']] = {
                'guidingUnit': r['guiding_unit'] if 'guiding_unit' in r.keys() else '',
                'president': r['president'] if 'president' in r.keys() else '',
                'category': r['category'] if 'category' in r.keys() else ''
            }
        result = []
        for c in all_clubs:
            cn = c['club_name']
            cnt = act_map.get(cn, 0)
            if cnt < threshold:
                p = profile_map.get(cn, {'guidingUnit': '', 'president': '', 'category': ''})
                tnames = '、'.join(sorted(teacher_map.get(cn, set()))) if cn in teacher_map else ''
                result.append({
                    'clubName': cn,
                    'activityCount': cnt,
                    'guidingUnit': p.get('guidingUnit', ''),
                    'president': p.get('president', ''),
                    'category': p.get('category', ''),
                    'teacherNames': tnames
                })
        result.sort(key=lambda x: x['activityCount'])
        wb = Workbook()
        ws = wb.active
        date_range = ''
        if start_date and end_date:
            date_range = f'（{start_date} 至 {end_date}）'
        elif start_date:
            date_range = f'（{start_date} 起）'
        elif end_date:
            date_range = f'（至 {end_date}）'
        ws.title = '低活动社团'
        header_fill = PatternFill(start_color='e74c3c', end_color='e74c3c', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF', size=11, name='微软雅黑')
        cell_font = Font(size=10, name='微软雅黑')
        thin_border = Border(left=Side(style='thin', color='d0d0e8'), right=Side(style='thin', color='d0d0e8'), top=Side(style='thin', color='d0d0e8'), bottom=Side(style='thin', color='d0d0e8'))
        warn_fill = PatternFill(start_color='FFE5E5', end_color='FFE5E5', fill_type='solid')
        zero_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
        title_font = Font(bold=True, size=14, name='微软雅黑', color='e74c3c')
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
        title_cell = ws.cell(row=1, column=1, value=f'活动次数低于{threshold}次的社团{date_range}')
        title_cell.font = title_font
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30
        headers = ['序号', '社团名称', '活动次数', '社团类型', '业务指导单位', '指导老师', '社长']
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=2, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border
        for i, c in enumerate(result, 1):
            row_data = [i, c['clubName'], c['activityCount'], c.get('category', '-'), c.get('guidingUnit', '-'), c.get('teacherNames', '-'), c.get('president', '-')]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(row=i + 2, column=col, value=val)
                cell.font = cell_font
                cell.border = thin_border
                cell.alignment = Alignment(vertical='center', wrap_text=True)
                if c['activityCount'] == 0:
                    cell.fill = zero_fill
                else:
                    cell.fill = warn_fill
        col_widths = [6, 24, 10, 14, 20, 20, 14]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f'低活动社团{date_range}.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/all-members-distribution')
def all_members_distribution():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        date_where = ''
        date_params = []
        if start_date:
            date_where += ' AND date(joined_at)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(joined_at)<=?'
            date_params.append(end_date)
        members = conn.execute(f'SELECT club_name, class_name, department, college FROM club_members WHERE 1=1{date_where}', date_params).fetchall()
        clubs = conn.execute('SELECT club_name FROM club_tokens').fetchall()
        total_users = conn.execute('SELECT COUNT(*) as c FROM users').fetchone()['c']
        total_clubs = conn.execute('SELECT COUNT(*) as c FROM club_profiles').fetchone()['c']
    finally:
        conn.close()
    club_counts = {}
    class_stats = {}
    dept_stats = {}
    for m in members:
        cn = m['club_name'] or '未知'
        club_counts[cn] = club_counts.get(cn, 0) + 1
        cls = (m['class_name'] or '').strip() or '未知'
        class_stats[cls] = class_stats.get(cls, 0) + 1
        dept = (m['department'] or '').strip() or '未知'
        dept_stats[dept] = dept_stats.get(dept, 0) + 1
    for c in clubs:
        cn = c['club_name']
        if cn not in club_counts:
            club_counts[cn] = 0
    college_stats = {}
    for col in COLLEGES:
        college_stats[col] = 0
    for m in members:
        mc = (m['college'] or '').strip()
        if mc in college_stats:
            college_stats[mc] += 1
    return jsonify({'success': True, 'data': {
        'totalMembers': len(members),
        'totalUsers': total_users,
        'totalClubs': total_clubs,
        'clubCounts': club_counts,
        'classStats': class_stats,
        'deptStats': dept_stats,
        'collegeStats': college_stats
    }})


@app.route('/api/analyze-members', methods=['POST'])
def analyze_members():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '请选择文件'}), 400
    try:
        import pandas as pd
        ext = os.path.splitext(f.filename)[1].lower()
        if ext == '.csv':
            df = pd.read_csv(f, encoding='utf-8-sig', on_bad_lines='skip')
        else:
            df = pd.read_excel(f, engine='openpyxl')
        df = df.dropna(how='all')
        if df.empty:
            return jsonify({'error': '文件为空或格式不正确'}), 400
        df.columns = [str(c).strip() for c in df.columns]
        summary = {'total': len(df)}
        for col in df.columns:
            col_lower = col.lower()
            if any(k in col_lower for k in ['性别', 'sex', 'gender']):
                vc = df[col].value_counts().to_dict()
                summary['male'] = int(vc.get('男', vc.get('M', 0)))
                summary['female'] = int(vc.get('女', vc.get('F', 0)))
            elif any(k in col_lower for k in ['年龄', 'age']):
                try:
                    summary['avgAge'] = round(df[col].mean(), 1)
                except:
                    pass
            elif any(k in col_lower for k in ['学院', 'college', '院系', '系别', 'department']):
                vc = df[col].value_counts().head(8).to_dict()
                summary['colleges'] = {str(k): int(v) for k, v in vc.items()}
            elif any(k in col_lower for k in ['爱好', 'hobby', '兴趣', '特长', 'interest']):
                all_vals = []
                for v in df[col].dropna():
                    for item in str(v).replace('，', ',').replace('、', ',').replace('/', ',').split(','):
                        item = item.strip()
                        if item:
                            all_vals.append(item)
                from collections import Counter
                top = Counter(all_vals).most_common(10)
                summary['hobbies'] = {k: v for k, v in top}
        records = df.head(100).to_dict('records')
        cleaned = []
        for r in records:
            row = {}
            for k, v in r.items():
                row[str(k)] = str(v) if pd.notna(v) else ''
            cleaned.append(row)
        return jsonify({'success': True, 'summary': summary, 'data': cleaned})
    except Exception as e:
        return jsonify({'error': f'分析失败: {str(e)}'}), 500


@app.route('/api/ai-generate', methods=['POST'])
def ai_generate():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    gen_type = data.get('type', '')
    club = data.get('club', '')
    if not club:
        return jsonify({'error': '请输入社团名称'}), 400
    if gen_type == 'copy':
        club_type = data.get('clubType', '其他')
        highlights = data.get('highlights', '')
        type_desc = {'文艺': '文艺表演与创作', '体育': '体育运动与竞技', '科技': '科技创新与实践', '学术': '学术研究与探讨', '公益': '公益服务与奉献', '其他': '多元发展与成长'}
        content = f"""🌟 {club} 纳新啦！🌟

亲爱的同学们，你是否渴望在大学里找到志同道合的伙伴？
你是否期待在{type_desc.get(club_type, '多元发展')}的道路上绽放光芒？

🎓 {club}诚挚邀请你的加入！

✨ 我们是谁？
{club}是一个致力于{type_desc.get(club_type, '多元发展与成长')}的校级优秀社团。在这里，我们用热情点燃梦想，用行动诠释青春！
{f'🏆 我们的亮点：{highlights}' if highlights else ''}

🎯 加入我们，你将获得：
• 专业的技能培训与指导
• 丰富的校内外活动机会
• 志同道合的伙伴与友谊
• 展示自我、提升能力的舞台
• 综合素质的全面提升

📋 招新对象：全校在校生
📍 报名方式：扫描下方二维码 / 填写线上报名表
⏰ 报名截止：待定

💫 每一个伟大的故事，都始于一次勇敢的选择。
加入{club}，让我们一起书写属于你的精彩篇章！

#{"济幼" if True else ""}{club}纳新 #社团招新 #大学生活"""
        return jsonify({'success': True, 'content': content})
    elif gen_type == 'form':
        form_type = data.get('formType', 'simple')
        if form_type == 'simple':
            content = f"""{club} 纳新报名表

━━━━━━━━━━━━━━━━━━━━
基本信息
━━━━━━━━━━━━━━━━━━━━
姓    名：_______________
学    号：_______________
性    别：□ 男  □ 女
联系电话：_______________
微 信 号：_______________
学    院：_______________
专    业：_______________
年    级：_______________

━━━━━━━━━━━━━━━━━━━━
个人意向
━━━━━━━━━━━━━━━━━━━━
是否服从调剂：□ 是  □ 否
个人特长/爱好：_________________________________

报名日期：_______________"""
        elif form_type == 'detailed':
            content = f"""{club} 纳新报名表（详细版）

━━━━━━━━━━━━━━━━━━━━
一、基本信息
━━━━━━━━━━━━━━━━━━━━
姓    名：_______________    性    别：□ 男  □ 女
学    号：_______________    出生年月：_______________
学    院：_______________    专    业：_______________
年    级：_______________    联系电话：_______________
微 信 号：_______________    QQ号码：_______________
政治面貌：□ 群众  □ 共青团员  □ 中共党员

━━━━━━━━━━━━━━━━━━━━
二、个人经历
━━━━━━━━━━━━━━━━━━━━
1. 曾任职务：_________________________________
2. 获奖情况：_________________________________
3. 特长技能：_________________________________
4. 兴趣爱好：_________________________________

━━━━━━━━━━━━━━━━━━━━
三、加入意向
━━━━━━━━━━━━━━━━━━━━
1. 为什么想加入{club}？
   _______________________________________________
2. 你能为社团带来什么？
   _______________________________________________
3. 对社团发展的建议：
   _______________________________________________

是否服从调剂：□ 是  □ 否
报名日期：_______________    签名：_______________"""
        else:
            content = f"""{club} 纳新报名表（专业版）

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
一、基本信息
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
姓    名：_______________    性    别：□ 男  □ 女
学    号：_______________    出生年月：_______________
身份证号：_______________    民族：_______________
学    院：_______________    专    业：_______________
年    级：_______________    班    级：_______________
联系电话：_______________    微    信：_______________
QQ号码：_______________    电子邮箱：_______________
政治面貌：_______________    籍    贯：_______________

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
二、教育背景
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
高中毕业学校：_________________________________
高考成绩/排名：_________________________________

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
三、个人能力评估
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 组织协调能力：□ 优秀  □ 良好  □ 一般
2. 沟通表达能力：□ 优秀  □ 良好  □ 一般
3. 创新思维能力：□ 优秀  □ 良好  □ 一般
4. 团队协作能力：□ 优秀  □ 良好  □ 一般

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
四、经历与成就
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
曾任职务：_________________________________
获奖情况：_________________________________
证书技能：_________________________________
项目经验：_________________________________

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
五、加入意向
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
意向部门：□ 策划部  □ 宣传部  □ 组织部  □ 外联部  □ 其他
为什么选择{club}？
_______________________________________________
你的优势与贡献：
_______________________________________________
未来规划与期望：
_______________________________________________

是否服从调剂：□ 是  □ 否
报名日期：_______________    本人签名：_______________"""
        return jsonify({'success': True, 'content': content})
    return jsonify({'error': '无效类型'}), 400


ANIMAL_TYPES = ['🐱猫', '🐶狗', '🐰兔', '🐼熊猫', '🦊狐狸', '🐧企鹅', '🦄独角兽', '🐲小龙', '🐸青蛙', '🦉猫头鹰', '🐹仓鼠', '🦋蝴蝶']
STAGE_THRESHOLDS = {'egg': 0, 'hatching': 30, 'baby': 80, 'child': 200, 'adult': 500}
STAGE_EMOJIS = {'egg': '🥚', 'hatching': '🐣', 'baby': '🐥', 'child': '🐣', 'adult': ''}


def get_pet_stage(exp):
    stages = sorted(STAGE_THRESHOLDS.items(), key=lambda x: x[1], reverse=True)
    for stage, threshold in stages:
        if exp >= threshold:
            return stage
    return 'egg'


def add_pet_exp(user_id, amount):
    conn = db.get_conn()
    try:
        pet = conn.execute('SELECT * FROM ai_pets WHERE user_id=?', (user_id,)).fetchone()
        if not pet:
            return None
        new_exp = pet['exp'] + amount
        old_stage = pet['stage']
        new_stage = get_pet_stage(new_exp)
        conn.execute('UPDATE ai_pets SET exp=?, stage=?, updated_at=CURRENT_TIMESTAMP WHERE user_id=?', (new_exp, new_stage, user_id))
        conn.commit()
        return {'exp': new_exp, 'stage': new_stage, 'leveled_up': new_stage != old_stage, 'animal_type': pet['animal_type'], 'name': pet['name']}
    finally:
        conn.close()


@app.route('/api/ai-pet', methods=['GET'])
def get_ai_pet():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        pet = conn.execute('SELECT * FROM ai_pets WHERE user_id=?', (user['id'],)).fetchone()
    finally:
        conn.close()
    if not pet:
        import random
        animal = random.choice(ANIMAL_TYPES)
        conn = db.get_conn()
        try:
            conn.execute('INSERT INTO ai_pets (user_id, animal_type, stage, exp) VALUES (?, ?, ?, ?)', (user['id'], animal, 'egg', 0))
            conn.commit()
        finally:
            conn.close()
        stage_info = get_stage_info(animal, 'egg', 0)
        return jsonify({'success': True, 'pet': stage_info})
    stage_info = get_stage_info(pet['animal_type'], pet['stage'], pet['exp'])
    stage_info['name'] = pet['name'] or ''
    stage_info['mood'] = pet['mood'] if 'mood' in pet.keys() and pet['mood'] else '😊 开心'
    return jsonify({'success': True, 'pet': stage_info})


def get_stage_info(animal_type, stage, exp):
    next_stages = sorted(STAGE_THRESHOLDS.items(), key=lambda x: x[1])
    next_stage = None
    next_exp = None
    for i, (s, t) in enumerate(next_stages):
        if s == stage and i + 1 < len(next_stages):
            next_stage = next_stages[i + 1][0]
            next_exp = next_stages[i + 1][1]
            break
    emoji = STAGE_EMOJIS.get(stage, '')
    if stage == 'adult':
        emoji = animal_type[:2]
    stage_names = {'egg': '蛋', 'hatching': '破壳中', 'baby': '幼年', 'child': '少年', 'adult': '成年'}
    return {
        'animalType': animal_type,
        'animalEmoji': animal_type[:2],
        'stage': stage,
        'stageName': stage_names.get(stage, stage),
        'emoji': emoji,
        'exp': exp,
        'nextStage': next_stage,
        'nextExp': next_exp,
        'expPercent': min(100, int(exp / next_exp * 100)) if next_exp else 100
    }


@app.route('/api/ai-pet', methods=['POST'])
def update_ai_pet():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    name = data.get('name', '').strip()
    conn = db.get_conn()
    try:
        if name:
            conn.execute('UPDATE ai_pets SET name=?, updated_at=CURRENT_TIMESTAMP WHERE user_id=?', (name, user['id']))
            conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/ai-pet-feed', methods=['POST'])
def feed_ai_pet():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    action = data.get('action', 'chat')
    food_type = data.get('food', 'default')
    exp_map = {'chat': 3, 'checkin': 5, 'create_tool': 8, 'upload': 6, 'enroll': 4, 'profile': 2, 'play': 3}
    food_bonus = {'apple': 2, 'fish': 3, 'carrot': 2, 'bamboo': 3, 'berry': 2, 'honey': 3, 'candy': 1, 'magic_bean': 5, 'default': 1}
    amount = exp_map.get(action, 2) + food_bonus.get(food_type, 1)
    result = add_pet_exp(user['id'], amount)
    if not result:
        return jsonify({'error': '宠物不存在'}), 404
    moods = ['😊 开心', '😄 兴奋', '🥰 满足', '😋 吃撑了', '🤩 超开心']
    import random
    new_mood = random.choice(moods)
    from datetime import datetime
    conn = db.get_conn()
    try:
        conn.execute('UPDATE ai_pets SET mood=?, last_interact=? WHERE user_id=?', (new_mood, datetime.now().isoformat(), user['id']))
        conn.commit()
    finally:
        conn.close()
    stage_info = get_stage_info(result['animal_type'], result['stage'], result['exp'])
    stage_info['leveledUp'] = result.get('leveled_up', False)
    stage_info['name'] = result.get('name', '')
    stage_info['expGained'] = amount
    stage_info['mood'] = new_mood
    return jsonify({'success': True, 'pet': stage_info})

FOOD_TYPES = [{'id':'apple','name':'🍎 红苹果','desc':'+2经验','bonus':2},{'id':'fish','name':'🐟 小鱼干','desc':'+3经验','bonus':3},{'id':'carrot','name':'🥕 胡萝卜','desc':'+2经验','bonus':2},{'id':'bamboo','name':'🎋 嫩竹叶','desc':'+3经验','bonus':3},{'id':'berry','name':'🍓 野莓果','desc':'+2经验','bonus':2},{'id':'honey','name':'🍯 蜂蜜罐','desc':'+3经验','bonus':3},{'id':'candy','name':'🍬 糖果','desc':'+1经验','bonus':1},{'id':'magic_bean','name':'🫘 魔法豆','desc':'+5经验','bonus':5},{'id':'default','name':'🍪 普通饼干','desc':'+1经验','bonus':1}]

@app.route('/api/ai-pet-games', methods=['GET'])
def pet_food_list():
    return jsonify({'success':True,'foods':FOOD_TYPES})

@app.route('/api/ai-pet-play', methods=['POST'])
def pet_play_game():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    game = data.get('game', 'guess')
    choice = data.get('choice', '')
    import random
    if game == 'guess':
        answer = random.randint(1, 10)
        user_guess = int(choice) if choice.isdigit() else -1
        if user_guess == answer:
            result = add_pet_exp(user['id'], 7)
            stage_info = get_stage_info(result['animal_type'], result['stage'], result['exp']) if result else {}
            return jsonify({'success':True,'win':True,'answer':answer,'message':'🎉 太厉害了！猜对了！','pet':stage_info,'expGained':7})
        else:
            return jsonify({'success':True,'win':False,'answer':answer,'message':f'答案是{answer}，差一点点~ 再试一次吧！','expGained':0})
    elif game == 'rps':
        options = ['rock','paper','scissors']
        pet_choice = random.choice(options)
        emoji_map = {'rock':'🪨','paper':'📄','scissors':'✂️'}
        name_map = {'rock':'石头','paper':'布','scissors':'剪刀'}
        win_conds = {'rock':'scissors','paper':'rock','scissors':'paper'}
        if choice == pet_choice:
            result = None
            msg = f'平局！我们都出了{emoji_map[choice]}{name_map[choice]}'
        elif win_conds.get(choice) == pet_choice:
            result = add_pet_exp(user['id'], 5)
            msg = f'你赢了！你出了{emoji_map[choice]}{name_map[choice]}，宠物出了{emoji_map[pet_choice]}{name_map[pet_choice]}'
        else:
            result = None
            msg = f'宠物赢了！你出了{emoji_map[choice]}{name_map[choice]}，宠物出了{emoji_map[pet_choice]}{name_map[pet_choice]}'
        stage_info = get_stage_info(result['animal_type'], result['stage'], result['exp']) if result else None
        return jsonify({'success':True,'message':msg,'petChoice':pet_choice,'pet':stage_info,'expGained':5 if result else 0})
    elif game == 'trivia':
        questions = [
            {'q':'地球上最大的海洋是？','opts':['太平洋','大西洋','印度洋','北冰洋'],'ans':0},
            {'q':'光年是什么单位？','opts':['时间','距离','速度','亮度'],'ans':1},
            {'q':'中国最长的河流是？','opts':['黄河','长江','珠江','淮河'],'ans':1},
            {'q':'水的化学式是？','opts':['CO2','H2O','O2','NaCl'],'ans':1},
            {'q':'太阳系最大的行星是？','opts':['土星','木星','天王星','海王星'],'ans':1},
            {'q':'"床前明月光"的作者是？','opts':['杜甫','白居易','李白','王维'],'ans':2},
            {'q':'世界上最高的山峰是？','opts':['K2','珠穆朗玛峰','干城章嘉峰','洛子峰'],'ans':1},
            {'q':'人体最大的器官是？','opts':['心脏','肝脏','皮肤','大脑'],'ans':2},
            {'q':'一年有多少个星期？','opts':['48','50','52','54'],'ans':2},
            {'q':'DNA的全称是？','opts':['脱氧核糖核酸','核糖核酸','氨基酸','蛋白质'],'ans':0},
            {'q':'"三人行必有我师"出自？','opts':['孟子','论语','大学','中庸'],'ans':1},
            {'q':'地球自转一周大约需要？','opts':['12小时','24小时','36小时','48小时'],'ans':1},
            {'q':'世界上面积最大的国家是？','opts':['中国','美国','加拿大','俄罗斯'],'ans':3},
            {'q':'"千里之行始于足下"出自？','opts':['孔子','老子','庄子','孟子'],'ans':1},
            {'q':'人体有多少块骨骼？','opts':['106','206','306','186'],'ans':1},
            {'q':'月球绕地球一周约需？','opts':['7天','14天','27天','60天'],'ans':2},
            {'q':'世界上最深的海沟是？','opts':['马里亚纳海沟','汤加海沟','日本海沟','菲律宾海沟'],'ans':0},
            {'q':'"但愿人长久"的下一句是？','opts':['千里共婵娟','对影成三人','低头思故乡','每逢佳节倍思亲'],'ans':0},
            {'q':'世界上使用人数最多的语言是？','opts':['英语','西班牙语','中文','印地语'],'ans':2},
            {'q':'地球的天然卫星是？','opts':['太阳','火星','月球','金星'],'ans':2},
        ]
        q = random.choice(questions)
        return jsonify({'success':True,'game':'trivia','question':q['q'],'options':q['opts'],'answerIndex':q['ans']})
    elif game == 'trivia_answer':
        correct = data.get('correct', False)
        if correct:
            result = add_pet_exp(user['id'], 8)
            stage_info = get_stage_info(result['animal_type'], result['stage'], result['exp']) if result else {}
            return jsonify({'success':True,'win':True,'message':'🎉 回答正确！你真博学！','pet':stage_info,'expGained':8})
        else:
            return jsonify({'success':True,'win':False,'message':'😅 答错了，没关系，下次一定行！','expGained':0})
    elif game == 'lucky':
        prizes = [
            {'name':'🌟 传说之冠','rarity':'legendary','weight':2,'exp':25},
            {'name':'💎 史诗宝石','rarity':'epic','weight':5,'exp':18},
            {'name':'🔮 稀有水晶','rarity':'rare','weight':12,'exp':12},
            {'name':'✨ 精良之星','rarity':'uncommon','weight':25,'exp':7},
            {'name':'🍀 普通四叶草','rarity':'common','weight':56,'exp':3},
        ]
        total_w = sum(p['weight'] for p in prizes)
        roll = random.randint(1, total_w)
        cumulative = 0
        prize = prizes[-1]
        for p in prizes:
            cumulative += p['weight']
            if roll <= cumulative:
                prize = p
                break
        rarity_emoji = {'legendary':'🌈','epic':'💜','rare':'💙','uncommon':'💚','common':'🤍'}
        result = add_pet_exp(user['id'], prize['exp'])
        stage_info = get_stage_info(result['animal_type'], result['stage'], result['exp']) if result else {}
        msg = f'{rarity_emoji.get(prize["rarity"],"")} 抽到了【{prize["name"]}】！获得 {prize["exp"]} 经验'
        if prize['rarity'] in ('legendary','epic'):
            msg += '\n🎊 太幸运了！稀有奖品！'
        return jsonify({'success':True,'game':'lucky','message':msg,'prize':prize,'pet':stage_info,'expGained':prize['exp']})
    elif game == 'adventure':
        events = [
            {'text':'🐾 宠物在森林里发现了一棵发光的树，摘下果实获得了经验！','exp':10},
            {'text':'🌊 宠物来到河边，钓到了一条金色小鱼！','exp':8},
            {'text':'🏔️ 宠物爬上山顶，看到了美丽的日出！','exp':12},
            {'text':'🏠 宠物帮助迷路的小动物找到了家，获得了感谢！','exp':9},
            {'text':'🌈 雨后出现了彩虹，宠物在彩虹下许了个愿！','exp':6},
            {'text':'🎪 宠物参加了森林音乐会，表演获得掌声！','exp':11},
            {'text':'📚 宠物在古书里发现了隐藏的知识！','exp':14},
            {'text':'🎭 宠物遇到了一位神秘旅者，学到了新技能！','exp':15},
            {'text':'🌸 宠物在花园里发现了一朵七色花！','exp':7},
            {'text':'⚔️ 宠物勇敢地击败了小怪兽，守护了村庄！','exp':13},
            {'text':'🌙 夜晚宠物看到了流星，许下了美好愿望！','exp':8},
            {'text':'🎨 宠物画了一幅美丽的画，被大家称赞！','exp':10},
            {'text':'🧩 宠物解开了一个古老的谜题！','exp':16},
            {'text':'🏖️ 宠物在沙滩上找到了一个藏宝箱！','exp':12},
            {'text':'🎵 宠物学会了弹奏一首美妙的曲子！','exp':9},
        ]
        event = random.choice(events)
        result = add_pet_exp(user['id'], event['exp'])
        stage_info = get_stage_info(result['animal_type'], result['stage'], result['exp']) if result else {}
        return jsonify({'success':True,'game':'adventure','message':event['text'],'pet':stage_info,'expGained':event['exp']})
    elif game == 'dice':
        user_dice = [random.randint(1,6) for _ in range(2)]
        pet_dice = [random.randint(1,6) for _ in range(2)]
        user_total = sum(user_dice)
        pet_total = sum(pet_dice)
        if user_total > pet_total:
            result = add_pet_exp(user['id'], 6)
            stage_info = get_stage_info(result['animal_type'], result['stage'], result['exp']) if result else {}
            msg = f'你掷出 {user_dice[0]}+{user_dice[1]}={user_total}，宠物掷出 {pet_dice[0]}+{pet_dice[1]}={pet_total}\n🎉 你赢了！'
            return jsonify({'success':True,'win':True,'message':msg,'pet':stage_info,'expGained':6,'userDice':user_dice,'petDice':pet_dice})
        elif user_total == pet_total:
            msg = f'你掷出 {user_dice[0]}+{user_dice[1]}={user_total}，宠物掷出 {pet_dice[0]}+{pet_dice[1]}={pet_total}\n🤝 平局！'
            return jsonify({'success':True,'win':False,'message':msg,'expGained':0,'userDice':user_dice,'petDice':pet_dice})
        else:
            msg = f'你掷出 {user_dice[0]}+{user_dice[1]}={user_total}，宠物掷出 {pet_dice[0]}+{pet_dice[1]}={pet_total}\n😢 宠物赢了！'
            return jsonify({'success':True,'win':False,'message':msg,'expGained':0,'userDice':user_dice,'petDice':pet_dice})
    return jsonify({'success':False,'error':'未知游戏'})

@app.route('/api/ai-event-plan', methods=['POST'])
def ai_event_plan():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    event_type = data.get('eventType', '团建')
    club = data.get('club', user.get('club_name', '社团'))
    msg = [{'role':'system','content':f'你是一个专业的社团活动策划师。请为"{club}"策划一个{event_type}类活动方案。包含：活动名称、活动目的、活动流程（按时间线）、所需物资、预算估算、预期效果。用Markdown格式输出，生动有感染力。'},{'role':'user','content':f'请为{club}策划{event_type}活动'}]
    result = call_llm_api(msg, max_tokens=1200)
    if result:
        return jsonify({'success':True,'content':result.strip()})
    plans = {
        '团建': f'''## 🎯 {club} 团建活动方案

**活动名称：** 「凝心聚力」趣味团建日
**活动目的：** 增进成员相互了解，培养团队协作精神

### 📋 活动流程
| 时间 | 环节 | 内容 |
|------|------|------|
| 14:00-14:20 | 破冰签到 | 趣味签到墙+分组 |
| 14:20-14:50 | 团队挑战① | 穿越电网/信任背摔 |
| 14:50-15:20 | 团队挑战② | 盲人方阵/解手链 |
| 15:20-15:50 | 创意比拼 | 团队logo设计+口号创作 |
| 15:50-16:30 | 自由交流 | 茶歇+分享感悟 |

### 💰 预算估算：200-500元（道具+零食）
### ✨ 预期效果：提升团队凝聚力80%+''',
        '文艺': f'''## 🎨 {club} 文艺汇演方案

**活动名称：** 「青春闪耀」文艺汇演
**活动目的：** 展示才艺，丰富校园文化生活

### 📋 活动流程
| 时间 | 环节 | 内容 |
|------|------|------|
| 18:00-18:30 | 入场签到 | 红毯仪式+拍照打卡 |
| 18:30-19:30 | 上半场 | 歌舞、朗诵、乐器演奏 |
| 19:30-19:45 | 中场互动 | 抽奖+观众才艺挑战 |
| 19:45-20:45 | 下半场 | 话剧、合唱、压轴表演 |
| 20:45-21:00 | 颁奖闭幕 | 最佳节目评选+合影 |

### 💰 预算估算：500-1000元
### ✨ 预期效果：参与人数100+人''',
        '公益': f'''## 🤝 {club} 公益志愿活动方案

**活动名称：** 「温暖传递」爱心公益行
**活动目的：** 践行社会责任，传递青春能量

### 📋 活动流程
| 时间 | 环节 | 内容 |
|------|------|------|
| 8:00-8:30 | 集合出发 | 物资分发+安全培训 |
| 8:30-11:00 | 志愿服务 | 社区清扫/敬老院探访/支教 |
| 11:00-11:30 | 互动交流 | 与受助对象互动 |
| 11:30-12:00 | 总结分享 | 志愿者感悟+合影 |

### 💰 预算估算：100-300元（物资）
### ✨ 预期效果：服务50+人次，传递正能量''',
        '学术': f'''## 📚 {club} 学术沙龙方案

**活动名称：** 「智慧碰撞」学术分享沙龙
**活动目的：** 激发学术热情，促进知识交流

### 📋 活动流程
| 时间 | 环节 | 内容 |
|------|------|------|
| 15:00-15:15 | 开幕致辞 | 主题介绍+嘉宾介绍 |
| 15:15-16:00 | 主题分享 | 3位分享者各15分钟 |
| 16:00-16:30 | 圆桌讨论 | 分组深度研讨 |
| 16:30-17:00 | 总结交流 | 成果展示+自由社交 |

### 💰 预算估算：100元（茶歇）
### ✨ 预期效果：参与30+人，产出研讨成果''',
        '竞赛': f'''## 🏆 {club} 技能竞赛方案

**活动名称：** 「巅峰对决」技能挑战赛
**活动目的：** 以赛促学，发掘优秀人才

### 📋 活动流程
| 时间 | 环节 | 内容 |
|------|------|------|
| 14:00-14:20 | 选手签到 | 抽签分组+规则说明 |
| 14:20-15:30 | 初赛阶段 | 分组淘汰赛 |
| 15:30-16:00 | 中场休息 | 观众互动小游戏 |
| 16:00-17:00 | 决赛阶段 | 冠亚军争夺战 |
| 17:00-17:30 | 颁奖典礼 | 获奖感言+合影 |

### 💰 预算估算：300-600元（奖品+道具）
### ✨ 预期效果：参赛20+人，观众50+人'''
    }
    plan = plans.get(event_type, plans['团建'])
    return jsonify({'success':True,'content':plan})

@app.route('/api/ai-polish-announcement', methods=['POST'])
def ai_polish_announcement():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    content = data.get('content', '').strip()
    style = data.get('style', 'formal')
    if not content:
        return jsonify({'error': '请输入公告内容'}), 400
    msg = [{'role':'system','content':f'你是一个社团公告润色助手。将用户提供的公告内容优化为{style}风格，纠正错别字、优化表达、增强感染力。直接输出润色后的公告，不要添加额外说明。'},{'role':'user','content':content}]
    result = call_llm_api(msg, max_tokens=600)
    if result:
        return jsonify({'success':True,'content':result.strip()})
    improved = content
    if style == 'formal':
        improved = '【通知】\n\n' + content + '\n\n——' + (user.get('club_name','') or '社团') + ' 发布'
    elif style == 'warm':
        improved = '🌟 亲爱的同学们：\n\n' + content + '\n\n💫 期待与你相遇！'
    elif style == 'urgent':
        improved = '⚠️ 重要通知 ⚠️\n\n' + content + '\n\n请相互转告，切勿错过！'
    return jsonify({'success':True,'content':improved})

@app.route('/api/ai-optimize-recruit', methods=['POST'])
def ai_optimize_recruit():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    content = data.get('content', '').strip()
    recruit_type = data.get('recruitType', 'member')
    if not content:
        return jsonify({'error': '请输入招募描述'}), 400
    type_name = '志愿者' if recruit_type == 'volunteer' else '活动成员'
    msg = [{'role':'system','content':f'你是一个社团招募文案优化师。优化以下{type_name}招募描述，使其更加吸引人、清晰明了。用HTML富文本格式输出（支持p/b/i/ul/li标签），不要输出html和body标签。'},{'role':'user','content':content}]
    result = call_llm_api(msg, max_tokens=600)
    if result:
        return jsonify({'success':True,'content':result.strip()})
    return jsonify({'success':True,'content':'<p><strong>✨ 优化建议：</strong></p><ul><li>突出活动亮点和收获</li><li>明确参与条件和要求</li><li>添加时间地点等关键信息</li><li>使用富有感染力的语言</li></ul><p><em>你可以输入具体描述内容让我帮你优化~</em></p>'})

@app.route('/api/ai-club-health', methods=['GET'])
def ai_club_health():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    club = user.get('club_name', '')
    if not club:
        return jsonify({'error': '你还没有关联社团'}), 400
    conn = db.get_conn()
    try:
        member_count = conn.execute('SELECT COUNT(*) as c FROM users WHERE club_name=? AND role="student"',(club,)).fetchone()['c']
        reg_count = conn.execute('SELECT COUNT(*) as c FROM club_registrations WHERE club_name=? AND status="pending"',(club,)).fetchone()['c']
        activity_count = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?',(club,)).fetchone()['c']
        checkin_count = conn.execute('SELECT COUNT(*) as c FROM checkin_records WHERE club_name=?',(club,)).fetchone()['c']
        dept_count = conn.execute('SELECT COUNT(*) as c FROM club_departments WHERE club_name=?',(club,)).fetchone()['c']
        notice_count = conn.execute('SELECT COUNT(*) as c FROM club_notices WHERE club_name=?',(club,)).fetchone()['c']
        recruit_count = conn.execute('SELECT COUNT(*) as c FROM recruitments WHERE club_name=? AND status="approved"',(club,)).fetchone()['c']
    finally:
        conn.close()
    scores = {}
    scores['成员规模'] = min(100, member_count * 5) if member_count > 0 else 0
    scores['纳新活力'] = min(100, reg_count * 15) if reg_count > 0 else 0
    scores['活动频率'] = min(100, activity_count * 15) if activity_count > 0 else 0
    scores['签到参与'] = min(100, checkin_count * 3) if checkin_count > 0 else 0
    scores['组织架构'] = min(100, dept_count * 25) if dept_count > 0 else 0
    scores['公告活跃'] = min(100, notice_count * 20) if notice_count > 0 else 0
    scores['招募开拓'] = min(100, recruit_count * 20) if recruit_count > 0 else 0
    total = sum(scores.values())
    avg = round(total / 7)
    level = '🌟 优秀' if avg >= 70 else ('👍 良好' if avg >= 50 else ('📈 发展' if avg >= 30 else '🌱 起步'))
    return jsonify({'success':True,'club':club,'score':avg,'level':level,'details':scores,'stats':{'members':member_count,'pending':reg_count,'activities':activity_count,'checkins':checkin_count,'departments':dept_count,'notices':notice_count,'recruits':recruit_count}})

@app.route('/api/ai-activity-summary', methods=['POST'])
def ai_activity_summary():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    activity_name = data.get('activityName', '').strip()
    highlights = data.get('highlights', '').strip()
    participants = data.get('participants', '').strip()
    if not activity_name:
        return jsonify({'error': '请输入活动名称'}), 400
    msg = [{'role':'system','content':'你是一个活动总结撰写助手。根据活动信息生成一份专业的活动总结报告，包含：活动概述、参与情况、活动亮点、改进建议。用Markdown格式输出。'},{'role':'user','content':f'活动名称：{activity_name}\n亮点：{highlights}\n参与情况：{participants}'}]
    result = call_llm_api(msg, max_tokens=800)
    if result:
        return jsonify({'success':True,'content':result.strip()})
    summary = f'''## 📋 {activity_name} 活动总结

### 📊 活动概述
本次活动「{activity_name}」圆满成功！{'参与者达' + participants + '人' if participants else '参与热情高涨'}。

### ✨ 活动亮点
{highlights if highlights else '• 活动组织有序，流程顺畅\n• 成员参与度高，气氛活跃\n• 达到了预期的活动目标'}

### 🔧 改进建议
• 提前做好更充分的物资准备
• 增加更多的互动环节
• 做好活动后的反馈收集

### 🎉 总结
本次活动为社团积累了宝贵经验，期待下一次更精彩的活动！'''
    return jsonify({'success':True,'content':summary})

@app.route('/api/ai-semester-plan', methods=['POST'])
def ai_semester_plan():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    club = data.get('club', user.get('club_name', '社团'))
    semester = data.get('semester', '本学期')
    msg = [{'role':'system','content':f'你是一个社团管理顾问。为{club}制定一份{semester}学期规划，包含：月度活动安排、纳新计划、成员培训、对外交流。用Markdown格式输出，具体实用。'},{'role':'user','content':f'请为{club}制定{semester}学期规划'}]
    result = call_llm_api(msg, max_tokens=1000)
    if result:
        return jsonify({'success':True,'content':result.strip()})
    plan = f'''## 📅 {club} {semester}学期规划

### 🗓 月度活动安排
| 月份 | 主题活动 | 负责人 |
|------|---------|--------|
| 第1月 | 🎯 新学期纳新宣传 | 宣传部 |
| 第2月 | 🤝 新成员见面会 | 组织部 |
| 第3月 | 🎨 社团特色活动 | 策划部 |
| 第4月 | 🏆 成果展示/评比 | 全体 |

### 🌱 纳新计划
- 线上：公众号推文+朋友圈海报
- 线下：食堂门口设点宣传
- 目标：招募20-30名新成员

### 📚 成员培训
- 每月1次技能培训工作坊
- 老带新结对子培养机制
- 月度之星评选激励

### 🌐 对外交流
- 与兄弟社团联谊1-2次
- 参加校级社团展示活动
- 争取校外合作资源'''
    return jsonify({'success':True,'content':plan})


@app.route('/api/ai-generate-poster', methods=['POST'])
def ai_generate_poster():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    club_name = data.get('club_name', '').strip()
    style = data.get('style', 'vibrant')
    theme = data.get('theme', '').strip()

    if not club_name:
        return jsonify({'error': '请提供社团名称'}), 400

    style_prompts = {
        'vibrant': '色彩鲜艳、充满活力的海报风格',
        'minimal': '简约现代、留白设计的海报风格',
        'traditional': '中国传统文化风格的海报',
        'modern': '科技感、未来风格的海报'
    }
    style_desc = style_prompts.get(style, style_prompts['vibrant'])

    prompt = f'{style_desc}，社团招募海报，社团名称"{club_name}"'
    if theme:
        prompt += f'，主题：{theme}'
    prompt += '，包含"加入我们"文字，适合打印，高质量'

    image_url = generate_image(prompt)

    if image_url:
        return jsonify({'success': True, 'image_url': image_url, 'prompt': prompt})

    return jsonify({
        'success': True,
        'image_url': None,
        'prompt': prompt,
        'fallback': True,
        'message': f'海报生成服务暂时不可用（已尝试通义万相和智谱GLM）。提示词已生成：{prompt}'
    })


@app.route('/api/ai-generate-poster-prompt', methods=['POST'])
def ai_generate_poster_prompt():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    club_name = data.get('club_name', '').strip()
    style = data.get('style', 'vibrant')
    theme = data.get('theme', '').strip()
    description = data.get('description', '').strip()

    if not club_name:
        return jsonify({'error': '请提供社团名称'}), 400

    messages = [
        {'role': 'system', 'content': '你是一个专业的海报设计师，擅长生成AI绘画提示词。'},
        {'role': 'user', 'content': f'请为社团"{club_name}"生成一个招募海报的AI绘画提示词。风格：{style}。{"主题："+theme if theme else ""}{"社团简介："+description if description else ""}要求：1)适合竖版海报 2)包含社团名称 3)有"加入我们"元素 4)输出英文提示词'}
    ]

    result = call_llm_api(messages, max_tokens=300)

    if result:
        return jsonify({'success': True, 'prompt': result})

    style_map = {'vibrant': 'colorful vibrant', 'minimal': 'minimalist clean', 'traditional': 'Chinese traditional ink', 'modern': 'futuristic tech'}
    fallback = f'{style_map.get(style, "colorful")} recruitment poster for "{club_name}" club, "Join Us" text, vertical format, high quality'
    return jsonify({'success': True, 'prompt': fallback, 'fallback': True})


@app.route('/api/ai-analyze-photo', methods=['POST'])
def ai_analyze_photo():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    photo_path = data.get('photo_path', '').strip()
    analysis_type = data.get('type', 'describe')
    if not photo_path:
        return jsonify({'error': '请提供照片路径'}), 400
    prompts = {
        'describe': '请详细描述这张图片的内容，包括场景、人物、活动等信息，用中文回答。',
        'verify': '请判断这张图片是否为真实的活动现场照片。检查：1)是否有人物活动 2)是否为摆拍或合成 3)场景是否与社团活动相关。给出判断结果和理由。',
        'ocr': '请识别并提取图片中的所有文字内容，按原始排版输出。'
    }
    prompt = prompts.get(analysis_type, prompts['describe'])
    import base64
    full_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), photo_path.lstrip('/'))
    if not os.path.exists(full_path):
        result = call_vision_api(photo_path, prompt)
    else:
        with open(full_path, 'rb') as f:
            img_b64 = base64.b64encode(f.read()).decode('utf-8')
        result = call_vision_api(img_b64, prompt)
    if result:
        return jsonify({'success': True, 'analysis': result, 'type': analysis_type})
    fallback = {
        'describe': '图片分析需要配置视觉模型API Key（QWEN_API_KEY）。当前为离线模式，无法分析图片内容。',
        'verify': '照片验证需要配置视觉模型。当前无法自动验证照片真伪，请人工审核。',
        'ocr': 'OCR识别需要配置视觉模型。当前无法识别图片中的文字。'
    }
    return jsonify({'success': True, 'analysis': fallback.get(analysis_type, fallback['describe']), 'type': analysis_type, 'fallback': True})


@app.route('/api/ai-batch-verify-photos', methods=['POST'])
def ai_batch_verify_photos():
    user = get_current_user()
    if not user or user.get('role') not in ['admin', 'user']:
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    club_name = data.get('club_name', '').strip()
    if not club_name:
        return jsonify({'error': '请提供社团名称'}), 400
    conn = db.get_conn()
    try:
        photos = conn.execute('SELECT id, activity_name, completion_photo FROM checkin_sessions WHERE club_name=? AND completion_photo!="" LIMIT 10', (club_name,)).fetchall()
    finally:
        conn.close()
    if not photos:
        return jsonify({'success': True, 'results': [], 'message': '没有找到活动照片'})
    results = []
    for photo in photos:
        results.append({
            'id': photo['id'],
            'photo_path': photo['completion_photo'],
            'activity_name': photo['activity_name'] if 'activity_name' in photo.keys() else '',
            'status': 'pending_analysis'
        })
    return jsonify({'success': True, 'results': results, 'total': len(results)})


QWEN_API_KEY = os.environ.get('QWEN_API_KEY', '')

IMAGE_GEN_API_KEY = os.environ.get('QWEN_API_KEY', '')

ZHIPU_API_KEY = os.environ.get('ZHIPU_API_KEY', '9b47fb12d3e54deea30741efd258eedc.woF9W94BJaxvPnIn')

LLM_PROVIDERS = {
    'zhipu': {
        'name': '智谱 GLM',
        'url': 'https://open.bigmodel.cn/api/paas/v4/chat/completions',
        'key_env': 'ZHIPU_API_KEY',
        'models': ['glm-4-flash', 'glm-4-plus', 'glm-4'],
        'default_model': 'glm-4-flash',
        'icon': '🟢',
        'desc': '智谱清言，GLM-4-flash 免费额度大，响应快'
    },
    'qwen-vl': {
        'name': '通义千问VL',
        'url': 'https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions',
        'key_env': 'QWEN_API_KEY',
        'models': ['qwen-vl-plus', 'qwen-vl-max'],
        'default_model': 'qwen-vl-plus',
        'icon': '👁',
        'desc': '通义千问视觉模型，可识别图片内容、验证照片真伪'
    }
}

ACTIVE_LLM = os.environ.get('ACTIVE_LLM', 'zhipu')

EMBEDDING_API_KEY = os.environ.get('QWEN_API_KEY', os.environ.get('ZHIPU_API_KEY', ''))
EMBEDDING_API_URL = 'https://dashscope.aliyuncs.com/compatible-mode/v1/embeddings'
EMBEDDING_MODEL = 'text-embedding-v3'


def get_embedding(text):
    if not EMBEDDING_API_KEY:
        return None
    import requests as req
    headers = {'Authorization': f'Bearer {EMBEDDING_API_KEY}', 'Content-Type': 'application/json'}
    payload = {
        'model': EMBEDDING_MODEL,
        'input': text[:2000],
        'dimensions': 1024
    }
    try:
        resp = req.post(EMBEDDING_API_URL, headers=headers, json=payload, timeout=15)
        if resp.status_code == 200:
            data = resp.json()
            return data.get('data', [{}])[0].get('embedding', None)
    except Exception:
        pass
    return None


def cosine_similarity(a, b):
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def get_llm_config():
    provider = LLM_PROVIDERS.get(ACTIVE_LLM, LLM_PROVIDERS['zhipu'])
    key_env = provider['key_env']
    api_key = os.environ.get(key_env, '')
    if not api_key:
        if key_env == 'ZHIPU_API_KEY':
            api_key = ZHIPU_API_KEY
    return provider, api_key


def call_llm_api(messages, max_tokens=800, tools=None, model_override=None):
    provider, api_key = get_llm_config()
    if not api_key:
        return None
    import requests as req
    headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
    model = model_override or provider['default_model']
    payload = {
        'model': model,
        'messages': messages,
        'max_tokens': max_tokens,
        'temperature': 0.7,
        'top_p': 0.9
    }
    if tools:
        payload['tools'] = tools
        payload['tool_choice'] = 'auto'
    try:
        resp = req.post(provider['url'], headers=headers, json=payload, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            if tools:
                return data
            content = data.get('choices', [{}])[0].get('message', {}).get('content', '')
            if content:
                return content
    except Exception:
        pass
    if not tools:
        for fallback_model in provider['models']:
            if fallback_model == model:
                continue
            payload['model'] = fallback_model
            try:
                resp = req.post(provider['url'], headers=headers, json=payload, timeout=20)
                if resp.status_code == 200:
                    data = resp.json()
                    content = data.get('choices', [{}])[0].get('message', {}).get('content', '')
                    if content:
                        return content
            except Exception:
                continue
    return None


def call_vision_api(image_input, prompt="请描述这张图片的内容", max_tokens=500):
    api_key = os.environ.get('QWEN_API_KEY', '')
    if not api_key:
        return None
    import requests as req
    headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
    if image_input.startswith('http'):
        image_content = {'type': 'image_url', 'image_url': {'url': image_input}}
    else:
        image_content = {'type': 'image_url', 'image_url': {'url': f'data:image/jpeg;base64,{image_input}'}}
    payload = {
        'model': 'qwen-vl-plus',
        'messages': [{
            'role': 'user',
            'content': [
                image_content,
                {'type': 'text', 'text': prompt}
            ]
        }],
        'max_tokens': max_tokens
    }
    try:
        resp = req.post('https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions', headers=headers, json=payload, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            return data.get('choices', [{}])[0].get('message', {}).get('content', '')
    except Exception:
        pass
    return None


def generate_image(prompt, size='1024x1024'):
    import requests as req
    import time

    # 策略1：优先尝试通义万相（异步任务+轮询）
    if IMAGE_GEN_API_KEY:
        try:
            headers = {'Authorization': f'Bearer {IMAGE_GEN_API_KEY}', 'Content-Type': 'application/json'}
            payload = {
                'model': 'wanx-v1',
                'input': {'prompt': prompt},
                'parameters': {'size': size, 'n': 1}
            }
            resp = req.post('https://dashscope.aliyuncs.com/api/v1/services/aigc/text2image/image-synthesis',
                           headers=headers, json=payload, timeout=30)
            data = resp.json()
            if resp.status_code == 200:
                output = data.get('output', {})
                results = output.get('results', [])
                if results:
                    return results[0].get('url', None)
                url = output.get('url')
                if url:
                    return url
                task_id = output.get('task_id')
                if task_id:
                    # 轮询获取结果（最多等待90秒）
                    task_url = f'https://dashscope.aliyuncs.com/api/v1/tasks/{task_id}'
                    for _ in range(30):
                        time.sleep(3)
                        task_resp = req.get(task_url, headers={'Authorization': f'Bearer {IMAGE_GEN_API_KEY}'}, timeout=15)
                        if task_resp.status_code != 200:
                            continue
                        task_data = task_resp.json()
                        task_status = task_data.get('output', {}).get('task_status', '')
                        if task_status == 'SUCCEEDED':
                            results = task_data.get('output', {}).get('results', [])
                            if results:
                                return results[0].get('url', None)
                            url = task_data.get('output', {}).get('url')
                            if url:
                                return url
                            return None
                        elif task_status in ('FAILED', 'UNKNOWN'):
                            break
            elif resp.status_code == 400:
                task_id = data.get('output', {}).get('task_id')
                if task_id:
                    task_url = f'https://dashscope.aliyuncs.com/api/v1/tasks/{task_id}'
                    for _ in range(30):
                        time.sleep(3)
                        task_resp = req.get(task_url, headers={'Authorization': f'Bearer {IMAGE_GEN_API_KEY}'}, timeout=15)
                        if task_resp.status_code != 200:
                            continue
                        task_data = task_resp.json()
                        task_status = task_data.get('output', {}).get('task_status', '')
                        if task_status == 'SUCCEEDED':
                            results = task_data.get('output', {}).get('results', [])
                            if results:
                                return results[0].get('url', None)
                            url = task_data.get('output', {}).get('url')
                            if url:
                                return url
                            return None
                        elif task_status in ('FAILED', 'UNKNOWN'):
                            break
        except Exception:
            pass

    # 策略2：Fallback 到智谱GLM cogview-3（同步返回，更稳定）
    if ZHIPU_API_KEY:
        try:
            headers = {'Authorization': f'Bearer {ZHIPU_API_KEY}', 'Content-Type': 'application/json'}
            payload = {'model': 'cogview-3', 'prompt': prompt}
            resp = req.post('https://open.bigmodel.cn/api/paas/v4/images/generations',
                           headers=headers, json=payload, timeout=60)
            if resp.status_code == 200:
                data = resp.json()
                results = data.get('data', [])
                if results and len(results) > 0:
                    return results[0].get('url', None)
        except Exception:
            pass

    return None


AI_TOOLS = [
    {
        'type': 'function',
        'function': {
            'name': 'check_approval',
            'description': '查看待审批项目并获取智能审批建议。支持材料审批、工作量审核、报名审批、财务权限审批、校外活动审批。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'approval_type': {
                        'type': 'string',
                        'enum': ['material', 'workload', 'registration', 'finance_permission', 'offcampus'],
                        'description': '审批类型：material=材料审批，workload=工作量审核，registration=报名审批，finance_permission=财务权限，offcampus=校外活动'
                    },
                    'club_name': {
                        'type': 'string',
                        'description': '社团名称，可选，不填则查看所有社团'
                    }
                },
                'required': ['approval_type']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'search_documents',
            'description': '在资料库中搜索文件。输入关键词即可快速找到相关文件，如搜索"书法活动照片"、"财务报表"、"报名表"等。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'query': {
                        'type': 'string',
                        'description': '搜索关键词'
                    },
                    'club_name': {
                        'type': 'string',
                        'description': '社团名称筛选，可选'
                    }
                },
                'required': ['query']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'check_alerts',
            'description': '检查系统预警，包括审批超时、社团活跃度骤降、赋分流程停滞等异常情况。',
            'parameters': {
                'type': 'object',
                'properties': {}
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'generate_insight_report',
            'description': '生成数据分析报告，包含社团总体概况、活跃度排名、待办积压、活动趋势等。可针对单个社团或全校生成。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'club_name': {
                        'type': 'string',
                        'description': '社团名称，不填则生成全校报告'
                    }
                }
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'query_club_data',
            'description': '查询社团的业务数据，如成员数、活动数、签到数、财务数据等。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'club_name': {
                        'type': 'string',
                        'description': '社团名称'
                    },
                    'data_type': {
                        'type': 'string',
                        'enum': ['members', 'activities', 'checkins', 'finance', 'workload', 'scoring', 'recruitments'],
                        'description': '数据类型：members=成员，activities=活动，checkins=签到，finance=财务，workload=工作量，scoring=赋分，recruitments=招募'
                    }
                },
                'required': ['club_name', 'data_type']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'list_clubs',
            'description': '获取所有社团列表及其基本信息（星级、类别等）。',
            'parameters': {
                'type': 'object',
                'properties': {}
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'analyze_photo',
            'description': '分析活动照片内容，验证照片真伪，或识别图片中的文字',
            'parameters': {
                'type': 'object',
                'properties': {
                    'photo_path': {'type': 'string', 'description': '照片文件路径'},
                    'analysis_type': {'type': 'string', 'enum': ['describe', 'verify', 'ocr'], 'description': '分析类型：describe=描述内容, verify=验证真伪, ocr=文字识别'}
                },
                'required': ['photo_path']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'query_data',
            'description': '查询系统中的各类数据集。支持：活动列表、成员列表、学生列表、指导老师情况、工作量统计、财务登记、招募申请、赋分记录、学分、签到记录、树洞内容、反馈列表、轮播图/优秀社团/优秀活动/热点资讯、资料库文件等。根据用户角色自动过滤权限范围。管理员可查看所有数据。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'dataset': {
                        'type': 'string',
                        'enum': ['activities', 'members', 'students', 'teachers', 'workload', 'finance', 'recruitments', 'scoring', 'credits', 'checkins', 'treehole', 'feedback', 'carousel', 'excellent_clubs', 'excellent_activities', 'hot_news', 'documents', 'clubs', 'notifications', 'departments', 'votes', 'surveys', 'joint_activities', 'offcampus_requests'],
                        'description': '数据集名称。students=学生列表（支持has_club过滤），teachers=指导老师列表（支持has_club过滤）'
                    },
                    'filters': {
                        'type': 'object',
                        'description': '过滤条件，如 {"club_name":"书法社","status":"pending","has_club":"false"}',
                        'properties': {
                            'club_name': {'type': 'string', 'description': '社团名称'},
                            'status': {'type': 'string', 'description': '状态过滤'},
                            'date_from': {'type': 'string', 'description': '开始日期'},
                            'date_to': {'type': 'string', 'description': '结束日期'},
                            'keyword': {'type': 'string', 'description': '关键词搜索'},
                            'has_club': {'type': 'string', 'enum': ['true', 'false'], 'description': '是否已关联社团：true=已关联，false=未关联（仅students和teachers数据集支持）'}
                        }
                    },
                    'limit': {
                        'type': 'integer',
                        'description': '返回条数限制，默认10'
                    }
                },
                'required': ['dataset']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'create_entity',
            'description': '创建实体。支持创建：活动、招募、投票、问卷、报名表、通知、联合活动申请、材料审批单、校外活动审批单。需要提供实体类型和属性。创建前会检查用户权限。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'entity_type': {
                        'type': 'string',
                        'enum': ['activity', 'recruitment', 'vote', 'survey', 'registration_form', 'notification', 'joint_activity', 'material_approval', 'offcampus_application'],
                        'description': '实体类型'
                    },
                    'attributes': {
                        'type': 'object',
                        'description': '实体属性，如活动名称、时间、地点等',
                        'properties': {
                            'name': {'type': 'string', 'description': '名称'},
                            'club_name': {'type': 'string', 'description': '社团名称'},
                            'description': {'type': 'string', 'description': '描述'},
                            'start_time': {'type': 'string', 'description': '开始时间'},
                            'end_time': {'type': 'string', 'description': '结束时间'},
                            'location': {'type': 'string', 'description': '地点'},
                            'content': {'type': 'string', 'description': '内容'},
                            'options': {'type': 'array', 'items': {'type': 'string'}, 'description': '选项列表（投票/问卷）'},
                            'max_participants': {'type': 'integer', 'description': '最大参与人数'},
                            'checkin_method': {'type': 'string', 'enum': ['code', 'qrcode', 'gps'], 'description': '签到方式'}
                        }
                    }
                },
                'required': ['entity_type', 'attributes']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'update_entity',
            'description': '更新实体状态或属性。支持：修改活动时间、审批报名、赋分审核、签到/签退、成员加入/退社、审批材料等。涉及资源变更的操作会先确认。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'entity_type': {
                        'type': 'string',
                        'enum': ['activity', 'registration', 'scoring', 'checkin', 'member', 'material', 'workload', 'offcampus', 'finance_permission'],
                        'description': '实体类型'
                    },
                    'entity_id': {
                        'type': 'integer',
                        'description': '实体ID'
                    },
                    'updates': {
                        'type': 'object',
                        'description': '更新内容',
                        'properties': {
                            'status': {'type': 'string', 'description': '新状态：approved/rejected/completed等'},
                            'action': {'type': 'string', 'description': '操作：approve/reject/checkin/checkout/join/leave等'},
                            'reason': {'type': 'string', 'description': '操作理由'},
                            'gps_lat': {'type': 'number', 'description': 'GPS纬度'},
                            'gps_lng': {'type': 'number', 'description': 'GPS经度'}
                        }
                    }
                },
                'required': ['entity_type', 'entity_id', 'updates']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'send_notification',
            'description': '发送通知给特定角色或用户。支持发送给：特定成员、全体成员、社团负责人群、指导老师、管理员。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'target': {
                        'type': 'string',
                        'description': '接收方：角色名(admin/user/teacher/student)、用户ID、或"all_members"全体成员'
                    },
                    'club_name': {
                        'type': 'string',
                        'description': '社团名称，发送社团通知时需要'
                    },
                    'content': {
                        'type': 'string',
                        'description': '通知内容'
                    },
                    'channel': {
                        'type': 'string',
                        'enum': ['system', 'push'],
                        'description': '通知渠道：system=系统内通知，push=推送通知',
                        'default': 'system'
                    }
                },
                'required': ['target', 'content']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'generate_report',
            'description': '生成数据分析报告。支持：活动概览统计、指导老师指导情况、工作量统计、参与率趋势、财务汇总、赋分分布、预警信息。根据用户角色自动限制数据范围。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'report_type': {
                        'type': 'string',
                        'enum': ['activity_overview', 'teacher_guidance', 'workload_stats', 'participation_trend', 'finance_summary', 'scoring_distribution', 'alert_summary', 'club_health'],
                        'description': '报告类型'
                    },
                    'club_name': {
                        'type': 'string',
                        'description': '社团名称，不填则根据角色自动选择'
                    },
                    'date_from': {
                        'type': 'string',
                        'description': '开始日期'
                    },
                    'date_to': {
                        'type': 'string',
                        'description': '结束日期'
                    }
                },
                'required': ['report_type']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'ai_generate',
            'description': 'AI生成内容。支持生成活动文案、宣传语等文本内容。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'generate_type': {
                        'type': 'string',
                        'enum': ['copywriting'],
                        'description': '生成类型：copywriting=活动文案/宣传语'
                    },
                    'params': {
                        'type': 'object',
                        'description': '生成参数',
                        'properties': {
                            'theme': {'type': 'string', 'description': '主题/活动名称'},
                            'club_name': {'type': 'string', 'description': '社团名称'},
                            'activity_type': {'type': 'string', 'description': '活动类型'},
                            'target_audience': {'type': 'string', 'description': '目标受众'},
                            'tone': {'type': 'string', 'description': '文案语调（如：正式、活泼、温馨）'}
                        }
                    }
                },
                'required': ['generate_type', 'params']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'recommend',
            'description': '为用户生成个性化推荐。支持：推荐社团、推荐活动、推荐招募（进行中的招募活动）、推荐联合活动伙伴。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'recommend_type': {
                        'type': 'string',
                        'enum': ['club', 'activity', 'recruitment', 'joint_partner'],
                        'description': '推荐类型：club=社团推荐，activity=活动推荐，recruitment=招募活动推荐，joint_partner=联合活动伙伴推荐'
                    },
                    'club_name': {
                        'type': 'string',
                        'description': '当前社团名称（推荐联合伙伴时需要）'
                    }
                },
                'required': ['recommend_type']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'query_database',
            'description': '直接查询数据库，执行SQL SELECT语句获取数据。管理员可查询所有表，非管理员自动按所属社团过滤。仅允许SELECT查询，禁止任何写操作。适合query_data工具无法满足的复杂查询需求。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'sql': {
                        'type': 'string',
                        'description': 'SQL SELECT查询语句。仅允许SELECT，禁止INSERT/UPDATE/DELETE/DROP/ALTER等写操作。查询结果最多返回50条记录。'
                    },
                    'description': {
                        'type': 'string',
                        'description': '查询目的说明，如"查询书法协会近期活动参与人数"'
                    }
                },
                'required': ['sql', 'description']
            }
        }
    },
    {
        'type': 'function',
        'function': {
            'name': 'generate_dashboard',
            'description': '生成实时可视化看板。支持成员活跃度、工作量分布、财务收支、签到考勤、赋分分布、社团概览等多种看板类型。看板会以图表形式在聊天中展示。',
            'parameters': {
                'type': 'object',
                'properties': {
                    'dashboard_type': {
                        'type': 'string',
                        'enum': ['member_activity', 'workload_distribution', 'finance_overview', 'attendance', 'scoring_distribution', 'club_overview', 'activity_trend'],
                        'description': '看板类型：member_activity=成员活跃度, workload_distribution=工作量分布, finance_overview=财务收支, attendance=签到考勤, scoring_distribution=赋分分布, club_overview=社团概览, activity_trend=活动趋势'
                    },
                    'club_name': {
                        'type': 'string',
                        'description': '社团名称，可选。不填则查看全校或当前用户社团'
                    }
                },
                'required': ['dashboard_type']
            }
        }
    }
]


def execute_tool_call(function_name, arguments, current_user=None):
    conn = db.get_conn()
    try:
        # 获取当前用户角色信息
        user_role = current_user.get('role', 'student') if current_user else 'student'
        user_club = current_user.get('club_name', '') if current_user else ''
        user_id = current_user.get('id', 0) if current_user else 0
        # 学生加入的社团
        user_clubs = []
        if current_user and user_role == 'student':
            try:
                cr = conn.execute('SELECT DISTINCT club_name FROM club_members WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_cadres WHERE user_id=?', (user_id, user_id)).fetchall()
                user_clubs = [r['club_name'] for r in cr if r['club_name']]
            except:
                pass
        # 指导老师指导的社团
        teacher_clubs = []
        if current_user and user_role == 'teacher':
            try:
                tr = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user_id,)).fetchall()
                teacher_clubs = [r['club_name'] for r in tr if r['club_name']]
            except:
                pass
        # 社团负责人如果 users.club_name 为空，尝试从 club_cadres/club_members 查找
        if current_user and user_role == 'user' and not user_club:
            try:
                cr = conn.execute('SELECT DISTINCT club_name FROM club_cadres WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_members WHERE user_id=?', (user_id, user_id)).fetchall()
                if cr:
                    user_club = cr[0]['club_name']
            except:
                pass

        if function_name == 'check_approval':
            atype = arguments.get('approval_type', 'material')
            kwargs = {}
            cn = arguments.get('club_name', '')
            # 权限校验：学生不能查看审批
            if user_role == 'student':
                return json.dumps({'error': '学生角色无权查看审批信息'}, ensure_ascii=False)
            # 社团负责人只能查看自己社团的审批
            if user_role == 'user':
                cn = user_club
            # 指导老师只能查看所指导社团的审批
            if user_role == 'teacher' and not cn:
                if teacher_clubs:
                    cn = teacher_clubs[0]
            if cn:
                kwargs['club_name'] = cn
            result = approval_agent.batch_analyze(atype, **kwargs)
            return json.dumps(result, ensure_ascii=False, default=str)

        elif function_name == 'search_documents':
            query = arguments.get('query', '')
            cn = arguments.get('club_name', '')
            # 角色权限过滤
            if user_role == 'student':
                if cn and cn not in user_clubs:
                    return json.dumps({'error': f'你只能搜索自己加入的社团（{", ".join(user_clubs)}）的文件'}, ensure_ascii=False)
                if not cn and user_clubs:
                    cn = user_clubs[0] if len(user_clubs) == 1 else ''
            elif user_role == 'teacher':
                if cn and cn not in teacher_clubs:
                    return json.dumps({'error': f'你只能搜索所指导社团（{", ".join(teacher_clubs)}）的文件'}, ensure_ascii=False)
                if not cn and len(teacher_clubs) == 1:
                    cn = teacher_clubs[0]
            elif user_role == 'user':
                if cn and cn != user_club:
                    return json.dumps({'error': f'你只能搜索自己社团（{user_club}）的文件'}, ensure_ascii=False)
                cn = user_club
            doc_agent = DocIndexAgent(db)
            sem_results = doc_agent.semantic_search(query, top_k=10)
            if sem_results.get('search_type') == 'semantic' and sem_results.get('total', 0) > 0:
                # 过滤语义搜索结果，只返回用户有权查看的社团文件
                if user_role != 'admin' and cn:
                    filtered = [r for r in sem_results.get('items', []) if r.get('club_name') == cn]
                    sem_results['items'] = filtered
                    sem_results['total'] = len(filtered)
                return json.dumps(sem_results, ensure_ascii=False, default=str)
            results = doc_agent.search(query, club_name=cn)
            summary = f'找到{len(results)}个文件'
            if results:
                summary += '：' + '；'.join([f'{r["file_name"]}({r["club_name"]})' for r in results[:5]])
            return json.dumps({'summary': summary, 'total': len(results), 'items': results[:10]}, ensure_ascii=False, default=str)

        elif function_name == 'check_alerts':
            # 权限校验：只有管理员可以查看系统预警
            if user_role != 'admin':
                return json.dumps({'total': 0, 'alerts': [], 'message': '系统预警仅管理员可查看，你可以查看自己社团的数据'}, ensure_ascii=False)
            alerts = notification_agent.check_all()
            return json.dumps({'total': len(alerts), 'alerts': alerts}, ensure_ascii=False, default=str)

        elif function_name == 'generate_insight_report':
            cn = arguments.get('club_name', '')
            # 权限校验
            if user_role == 'student':
                if cn and cn not in user_clubs:
                    return json.dumps({'error': f'你只能查看自己加入的社团报告'}, ensure_ascii=False)
                if not cn and user_clubs:
                    cn = user_clubs[0]
            elif user_role == 'teacher':
                if cn and cn not in teacher_clubs:
                    return json.dumps({'error': f'你只能查看所指导社团的报告'}, ensure_ascii=False)
                if not cn and teacher_clubs:
                    cn = teacher_clubs[0]
            elif user_role == 'user':
                if not cn:
                    cn = user_club
                if cn and cn != user_club:
                    return json.dumps({'error': f'你只能查看自己社团的报告'}, ensure_ascii=False)
            report = data_insight_agent.generate_report(club_name=cn)
            return json.dumps(report, ensure_ascii=False, default=str)

        elif function_name == 'query_club_data':
            cn = arguments.get('club_name', '')
            dtype = arguments.get('data_type', 'members')
            # 权限校验
            if user_role == 'student':
                if cn and cn not in user_clubs:
                    return json.dumps({'error': f'你只能查看自己加入的社团数据'}, ensure_ascii=False)
                if not cn and user_clubs:
                    cn = user_clubs[0]
                if not cn:
                    return json.dumps({'error': '你尚未加入任何社团'}, ensure_ascii=False)
            elif user_role == 'teacher':
                if cn and cn not in teacher_clubs:
                    return json.dumps({'error': f'你只能查看所指导社团的数据'}, ensure_ascii=False)
                if not cn and teacher_clubs:
                    cn = teacher_clubs[0]
                if not cn:
                    return json.dumps({'error': '你尚未指导任何社团'}, ensure_ascii=False)
            elif user_role == 'user':
                if not cn:
                    cn = user_club
                if cn and cn != user_club:
                    return json.dumps({'error': f'你只能查看自己社团的数据'}, ensure_ascii=False)
            if not cn:
                return json.dumps({'error': '请提供社团名称'}, ensure_ascii=False)
            result = {}
            if dtype == 'members':
                members = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE club_name=?', (cn,)).fetchone()
                leaders = conn.execute("SELECT real_name, department FROM club_members WHERE club_name=? AND department!='' LIMIT 10", (cn,)).fetchall()
                result = {'total': members['c'] if members else 0, 'leaders': [dict(l) for l in leaders]}
            elif dtype == 'activities':
                total = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (cn,)).fetchone()
                completed = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=? AND is_completed=1', (cn,)).fetchone()
                recent = conn.execute('SELECT activity_name, created_at FROM checkin_sessions WHERE club_name=? ORDER BY created_at DESC LIMIT 5', (cn,)).fetchall()
                result = {'total': total['c'] if total else 0, 'completed': completed['c'] if completed else 0, 'recent': [dict(r) for r in recent]}
            elif dtype == 'checkins':
                total = conn.execute('SELECT COUNT(*) as c FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=?', (cn,)).fetchone()
                result = {'total_checkins': total['c'] if total else 0}
            elif dtype == 'finance':
                income = conn.execute("SELECT COALESCE(SUM(amount),0) as s FROM finance_records WHERE club_name=? AND type='income'", (cn,)).fetchone()
                expense = conn.execute("SELECT COALESCE(SUM(amount),0) as s FROM finance_records WHERE club_name=? AND type='expense'", (cn,)).fetchone()
                result = {'income': income['s'] if income else 0, 'expense': expense['s'] if expense else 0, 'balance': (income['s'] if income else 0) - (expense['s'] if expense else 0)}
            elif dtype == 'workload':
                pending = conn.execute('SELECT COUNT(*) as c FROM workload_submissions WHERE club_name=? AND status="pending"', (cn,)).fetchone()
                approved = conn.execute('SELECT COUNT(*) as c FROM workload_submissions WHERE club_name=? AND status="approved"', (cn,)).fetchone()
                result = {'pending': pending['c'] if pending else 0, 'approved': approved['c'] if approved else 0}
            elif dtype == 'scoring':
                subs = conn.execute('SELECT status, COUNT(*) as c FROM scoring_submissions WHERE club_name=? GROUP BY status', (cn,)).fetchall()
                result = {'submissions': [dict(r) for r in subs]}
            elif dtype == 'recruitments':
                total = conn.execute('SELECT COUNT(*) as c FROM recruitments WHERE club_name=?', (cn,)).fetchone()
                active = conn.execute('SELECT COUNT(*) as c FROM recruitments WHERE club_name=? AND status="approved"', (cn,)).fetchone()
                result = {'total': total['c'] if total else 0, 'active': active['c'] if active else 0}
            return json.dumps(result, ensure_ascii=False, default=str)

        elif function_name == 'list_clubs':
            clubs = conn.execute('SELECT club_name, star_rating, category FROM club_profiles ORDER BY star_rating DESC').fetchall()
            return json.dumps({'clubs': [dict(c) for c in clubs]}, ensure_ascii=False, default=str)

        elif function_name == 'analyze_photo':
            photo_path = arguments.get('photo_path', '')
            analysis_type = arguments.get('analysis_type', 'describe')
            prompts_for_type = {
                'describe': '请详细描述这张图片的内容，包括场景、人物、活动等信息，用中文回答。',
                'verify': '请判断这张图片是否为真实的活动现场照片。检查：1)是否有人物活动 2)是否为摆拍或合成 3)场景是否与社团活动相关。给出判断结果和理由。',
                'ocr': '请识别并提取图片中的所有文字内容，按原始排版输出。'
            }
            prompt = prompts_for_type.get(analysis_type, '请描述这张图片')
            result = call_vision_api(photo_path, prompt)
            if not result:
                result = json.dumps({'analysis': '视觉模型未配置，无法分析图片', 'fallback': True}, ensure_ascii=False)
            return result

        elif function_name == 'query_data':
            dataset = arguments.get('dataset', '')
            filters = arguments.get('filters', {}) or {}
            limit = arguments.get('limit', 10)
            cn = filters.get('club_name', '')
            status = filters.get('status', '')
            keyword = filters.get('keyword', '')
            date_from = filters.get('date_from', '')
            date_to = filters.get('date_to', '')
            has_club = filters.get('has_club', '')

            # 角色权限过滤：非管理员只能查看自己社团的数据
            if user_role == 'student':
                # 学生数据集：只能查看自己
                if dataset == 'students':
                    pass  # 在数据集查询中过滤
                elif dataset == 'teachers':
                    # 学生不能查看老师列表
                    return json.dumps({'error': '老师列表仅管理员和社团负责人可查看'}, ensure_ascii=False)
                else:
                    # 其他数据集：只能查看自己加入的社团数据
                    allowed_clubs = user_clubs
                    if not allowed_clubs:
                        return json.dumps({'error': '你尚未加入任何社团，无法查看数据'}, ensure_ascii=False)
                    if cn and cn not in allowed_clubs:
                        return json.dumps({'error': f'你只能查看自己加入的社团数据（{", ".join(allowed_clubs)}）'}, ensure_ascii=False)
                    if not cn:
                        cn = allowed_clubs[0] if allowed_clubs else ''
                    if dataset in ('carousel', 'hot_news', 'feedback'):
                        if dataset == 'feedback':
                            return json.dumps({'error': '反馈数据仅管理员可查看'}, ensure_ascii=False)
                        if dataset == 'carousel':
                            return json.dumps({'error': '轮播图管理仅管理员可操作'}, ensure_ascii=False)
            elif user_role == 'teacher':
                # 老师数据集：只能查看自己
                if dataset == 'teachers':
                    pass  # 在数据集查询中过滤
                elif dataset == 'students':
                    # 指导老师可以查看所指导社团的学生
                    pass
                else:
                    # 其他数据集：只能查看所指导社团的数据
                    allowed_clubs = teacher_clubs
                    if not allowed_clubs:
                        return json.dumps({'error': '你尚未指导任何社团，无法查看数据'}, ensure_ascii=False)
                    if cn and cn not in allowed_clubs:
                        return json.dumps({'error': f'你只能查看所指导社团的数据（{", ".join(allowed_clubs)}）'}, ensure_ascii=False)
                    if not cn:
                        cn = allowed_clubs[0] if len(allowed_clubs) == 1 else ''
            elif user_role == 'user':
                # 社团负责人只能查看自己社团的数据
                if not cn:
                    cn = user_club
                if cn and cn != user_club:
                    return json.dumps({'error': f'你只能查看自己社团（{user_club}）的数据'}, ensure_ascii=False)
                if not cn:
                    return json.dumps({'error': '请指定社团名称'}, ensure_ascii=False)

            result_data = {}
            if dataset == 'activities':
                q = 'SELECT id, activity_name, club_name, status, created_at, is_completed FROM checkin_sessions WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if status: q += ' AND status=?'; params.append(status)
                if keyword: q += ' AND activity_name LIKE ?'; params.append(f'%{keyword}%')
                if date_from: q += ' AND created_at>=?'; params.append(date_from)
                if date_to: q += ' AND created_at<=?'; params.append(date_to)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                total = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE 1=1' + ('' if not cn else ' AND club_name=?'), ([cn] if cn else [])).fetchone()['c']
                result_data = {'total': total, 'items': [dict(r) for r in rows]}
            elif dataset == 'members':
                q = 'SELECT id, username, real_name, club_name, department, specialty FROM club_members WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if keyword: q += ' AND (real_name LIKE ? OR username LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                q += ' ORDER BY id DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                total_q = 'SELECT COUNT(*) as c FROM club_members WHERE 1=1' + ('' if not cn else ' AND club_name=?')
                total = conn.execute(total_q, ([cn] if cn else [])).fetchone()['c']
                result_data = {'total': total, 'items': [dict(r) for r in rows]}
            elif dataset == 'students':
                # 学生列表：从 users 表查询 role=student 的用户
                if has_club == 'false':
                    # 查询没有加入任何社团的学生
                    q = 'SELECT u.id, u.username, up.real_name, u.club_name FROM users u LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.role="student" AND u.id NOT IN (SELECT DISTINCT user_id FROM club_members WHERE user_id IS NOT NULL) AND u.id NOT IN (SELECT DISTINCT user_id FROM club_cadres WHERE user_id IS NOT NULL) AND (u.club_name IS NULL OR u.club_name="")'
                    params = []
                    if keyword: q += ' AND (u.username LIKE ? OR up.real_name LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                    q += ' ORDER BY u.id LIMIT ?'; params.append(limit)
                    rows = conn.execute(q, params).fetchall()
                    total = conn.execute('SELECT COUNT(*) as c FROM users u WHERE u.role="student" AND u.id NOT IN (SELECT DISTINCT user_id FROM club_members WHERE user_id IS NOT NULL) AND u.id NOT IN (SELECT DISTINCT user_id FROM club_cadres WHERE user_id IS NOT NULL) AND (u.club_name IS NULL OR u.club_name="")').fetchone()['c']
                    result_data = {'total': total, 'items': [dict(r) for r in rows], 'description': '未加入任何社团的学生'}
                elif has_club == 'true':
                    # 查询已加入社团的学生
                    q = 'SELECT u.id, u.username, up.real_name, u.club_name FROM users u LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.role="student" AND (u.id IN (SELECT DISTINCT user_id FROM club_members WHERE user_id IS NOT NULL) OR u.id IN (SELECT DISTINCT user_id FROM club_cadres WHERE user_id IS NOT NULL) OR (u.club_name IS NOT NULL AND u.club_name!=""))'
                    params = []
                    if cn: q += ' AND u.club_name=?'; params.append(cn)
                    if keyword: q += ' AND (u.username LIKE ? OR up.real_name LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                    q += ' ORDER BY u.id LIMIT ?'; params.append(limit)
                    rows = conn.execute(q, params).fetchall()
                    total = conn.execute('SELECT COUNT(*) as c FROM users u WHERE u.role="student" AND (u.id IN (SELECT DISTINCT user_id FROM club_members WHERE user_id IS NOT NULL) OR u.id IN (SELECT DISTINCT user_id FROM club_cadres WHERE user_id IS NOT NULL) OR (u.club_name IS NOT NULL AND u.club_name!=""))' + ('' if not cn else ' AND u.club_name=?'), ([cn] if cn else [])).fetchone()['c']
                    result_data = {'total': total, 'items': [dict(r) for r in rows], 'description': '已加入社团的学生'}
                else:
                    # 查询所有学生
                    q = 'SELECT u.id, u.username, up.real_name, u.club_name FROM users u LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.role="student"'
                    params = []
                    if cn: q += ' AND u.club_name=?'; params.append(cn)
                    if keyword: q += ' AND (u.username LIKE ? OR up.real_name LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                    q += ' ORDER BY u.id LIMIT ?'; params.append(limit)
                    rows = conn.execute(q, params).fetchall()
                    total = conn.execute('SELECT COUNT(*) as c FROM users u WHERE u.role="student"' + ('' if not cn else ' AND u.club_name=?'), ([cn] if cn else [])).fetchone()['c']
                    result_data = {'total': total, 'items': [dict(r) for r in rows]}
                # 非管理员过滤：学生只能看自己，老师只能看所指导社团的
                if user_role == 'student':
                    result_data['items'] = [r for r in result_data['items'] if r['id'] == current_user['id']]
                    result_data['total'] = len(result_data['items'])
                elif user_role == 'teacher':
                    result_data['items'] = [r for r in result_data['items'] if r.get('club_name') in teacher_clubs]
                    result_data['total'] = len(result_data['items'])
            elif dataset == 'teachers':
                if has_club == 'false':
                    # 查询没有指导任何社团的老师
                    q = 'SELECT u.id, u.username, tp.real_name, tp.work_id FROM users u LEFT JOIN teacher_profiles tp ON u.id=tp.user_id WHERE u.role="teacher" AND u.id NOT IN (SELECT DISTINCT user_id FROM teacher_clubs)'
                    params = []
                    if keyword: q += ' AND (u.username LIKE ? OR tp.real_name LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                    q += ' ORDER BY u.id LIMIT ?'; params.append(limit)
                    rows = conn.execute(q, params).fetchall()
                    total = conn.execute('SELECT COUNT(*) as c FROM users u WHERE u.role="teacher" AND u.id NOT IN (SELECT DISTINCT user_id FROM teacher_clubs)').fetchone()['c']
                    result_data = {'total': total, 'items': [dict(r) for r in rows], 'description': '未指导任何社团的老师'}
                elif has_club == 'true':
                    # 查询已指导社团的老师
                    q = 'SELECT tc.id, tc.user_id, tc.club_name, u.username, tp.real_name FROM teacher_clubs tc LEFT JOIN users u ON tc.user_id=u.id LEFT JOIN teacher_profiles tp ON tc.user_id=tp.user_id WHERE 1=1'
                    params = []
                    if cn: q += ' AND tc.club_name=?'; params.append(cn)
                    if keyword: q += ' AND (u.username LIKE ? OR tp.real_name LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                    q += ' ORDER BY tc.id LIMIT ?'; params.append(limit)
                    rows = conn.execute(q, params).fetchall()
                    result_data = {'items': [dict(r) for r in rows], 'description': '已指导社团的老师'}
                else:
                    # 查询所有老师及其指导社团
                    q = 'SELECT u.id, u.username, tp.real_name, tp.work_id, GROUP_CONCAT(tc.club_name) as clubs FROM users u LEFT JOIN teacher_profiles tp ON u.id=tp.user_id LEFT JOIN teacher_clubs tc ON u.id=tc.user_id WHERE u.role="teacher" GROUP BY u.id'
                    params = []
                    if keyword: q += ' HAVING (u.username LIKE ? OR tp.real_name LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                    q += ' ORDER BY u.id LIMIT ?'; params.append(limit)
                    rows = conn.execute(q, params).fetchall()
                    result_data = {'items': [dict(r) for r in rows]}
                # 非管理员过滤：老师只能看自己
                if user_role == 'teacher':
                    result_data['items'] = [r for r in result_data['items'] if r.get('user_id') == current_user['id'] or r.get('id') == current_user['id']]
                    result_data['total'] = len(result_data['items'])
            elif dataset == 'workload':
                q = 'SELECT id, club_name, student_name, item_name, score, status, created_at FROM workload_submissions WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if status: q += ' AND status=?'; params.append(status)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'finance':
                q = 'SELECT id, club_name, type, amount, description, created_at FROM finance_records WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if date_from: q += ' AND created_at>=?'; params.append(date_from)
                if date_to: q += ' AND created_at<=?'; params.append(date_to)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                income = conn.execute("SELECT COALESCE(SUM(amount),0) as s FROM finance_records WHERE type='income'" + ('' if not cn else ' AND club_name=?'), ([cn] if cn else [])).fetchone()['s']
                expense = conn.execute("SELECT COALESCE(SUM(amount),0) as s FROM finance_records WHERE type='expense'" + ('' if not cn else ' AND club_name=?'), ([cn] if cn else [])).fetchone()['s']
                result_data = {'income': income, 'expense': expense, 'balance': income - expense, 'items': [dict(r) for r in rows]}
            elif dataset == 'recruitments':
                q = 'SELECT id, club_name, title, status, created_at FROM recruitments WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if status: q += ' AND status=?'; params.append(status)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'scoring':
                q = 'SELECT ss.id, ss.club_name, ss.status, ss.created_at, si.student_name, si.total_workload, si.final_score FROM scoring_submissions ss LEFT JOIN scoring_submission_items si ON ss.id=si.submission_id WHERE 1=1'
                params = []
                if cn: q += ' AND ss.club_name=?'; params.append(cn)
                if status: q += ' AND ss.status=?'; params.append(status)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'credits':
                q = 'SELECT id, student_name, club1, score1, club2, score2, final_credit, semester FROM final_credits WHERE 1=1'
                params = []
                if cn: q += ' AND (club1=? OR club2=?)'; params.extend([cn, cn])
                q += ' ORDER BY id DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'checkins':
                q = 'SELECT cr.id, cr.session_id, cr.student_name, cr.student_id, cr.created_at as checkin_time, cs.activity_name, cs.club_name FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE 1=1'
                params = []
                if cn: q += ' AND cs.club_name=?'; params.append(cn)
                q += ' ORDER BY cr.created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'treehole':
                q = 'SELECT id, content, scope, club_name, created_at, status FROM tree_hole WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if status: q += ' AND status=?'; params.append(status)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'feedback':
                q = 'SELECT id, club_name, user_id, type, title, body, status, created_at FROM feedbacks WHERE 1=1'
                params = []
                if status: q += ' AND status=?'; params.append(status)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'carousel':
                rows = conn.execute('SELECT id, club_name, description, image_path FROM club_showcase ORDER BY id DESC LIMIT ?', (limit,)).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'excellent_clubs':
                q = 'SELECT id, club_name, selected_at FROM excellent_clubs WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                q += ' ORDER BY selected_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'excellent_activities':
                rows = conn.execute('SELECT id, group_id, selected_at FROM excellent_activities ORDER BY selected_at DESC LIMIT ?', (limit,)).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'hot_news':
                rows = conn.execute('SELECT id, title, content, attachment_path, created_at FROM notices ORDER BY created_at DESC LIMIT ?', (limit,)).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'documents':
                q = 'SELECT id, file_name, club_name, category, created_at FROM doc_index WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if keyword: q += ' AND (file_name LIKE ? OR tags LIKE ?)'; params.extend([f'%{keyword}%', f'%{keyword}%'])
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'clubs':
                rows = conn.execute('SELECT club_name, star_rating, category, description FROM club_profiles ORDER BY star_rating DESC LIMIT ?', (limit,)).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'notifications':
                q = 'SELECT id, club_name, title, content, created_at FROM club_notices WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'departments':
                q = 'SELECT id, club_name, dept_name, description FROM club_departments WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'votes':
                q = 'SELECT id, club_name, title, status, created_at FROM club_tools WHERE tool_type="vote" AND 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'surveys':
                q = 'SELECT id, club_name, title, status, created_at FROM club_tools WHERE tool_type="survey" AND 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'joint_activities':
                q = 'SELECT id, club_name, title, status, created_at FROM joint_activities WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif dataset == 'offcampus_requests':
                q = 'SELECT id, club_name, title, status, created_at FROM offcampus_requests WHERE 1=1'
                params = []
                if cn: q += ' AND club_name=?'; params.append(cn)
                if status: q += ' AND status=?'; params.append(status)
                q += ' ORDER BY created_at DESC LIMIT ?'; params.append(limit)
                rows = conn.execute(q, params).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            else:
                result_data = {'error': f'未知数据集：{dataset}'}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'create_entity':
            entity_type = arguments.get('entity_type', '')
            attrs = arguments.get('attributes', {}) or {}
            result_data = {}
            # 权限校验：学生不能创建实体
            if user_role == 'student':
                return json.dumps({'error': '学生角色无权创建内容，如需操作请联系社团负责人'}, ensure_ascii=False)
            # 指导老师只能查看，不能创建
            if user_role == 'teacher':
                return json.dumps({'error': '指导老师角色无权创建内容，如需操作请联系社团负责人或管理员'}, ensure_ascii=False)
            # 社团负责人只能为自己社团创建
            if user_role == 'user':
                cn = attrs.get('club_name', '')
                if cn and cn != user_club:
                    return json.dumps({'error': f'你只能为自己社团（{user_club}）创建内容'}, ensure_ascii=False)
                if not cn:
                    attrs['club_name'] = user_club
            if entity_type == 'activity':
                name = attrs.get('name', '未命名活动')
                cn = attrs.get('club_name', '')
                desc = attrs.get('description', '')
                start = attrs.get('start_time', '')
                end = attrs.get('end_time', '')
                loc = attrs.get('location', '')
                method = attrs.get('checkin_method', 'code')
                if not cn:
                    return json.dumps({'error': '创建活动需要指定社团名称', 'hint': '请提供 club_name 参数'}, ensure_ascii=False)
                try:
                    import random, string
                    code = ''.join(random.choices(string.digits, k=6))
                    conn.execute('INSERT INTO checkin_sessions (club_name, activity_name, checkin_code, status, created_at) VALUES (?, ?, ?, "open", CURRENT_TIMESTAMP)', (cn, name, code))
                    conn.commit()
                    sid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                    result_data = {'success': True, 'id': sid, 'message': f'活动"{name}"已创建，签到码：{code}'}
                except Exception as e:
                    result_data = {'error': f'创建活动失败：{str(e)}'}
            elif entity_type == 'recruitment':
                cn = attrs.get('club_name', '')
                title = attrs.get('name', '招募')
                desc = attrs.get('description', '')
                if not cn:
                    return json.dumps({'error': '创建招募需要指定社团名称'}, ensure_ascii=False)
                try:
                    conn.execute('INSERT INTO recruitments (club_name, title, description, status, created_at) VALUES (?, ?, ?, "pending", CURRENT_TIMESTAMP)', (cn, title, desc))
                    conn.commit()
                    rid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                    result_data = {'success': True, 'id': rid, 'message': f'招募"{title}"已创建，等待审批'}
                except Exception as e:
                    result_data = {'error': f'创建招募失败：{str(e)}'}
            elif entity_type == 'notification':
                cn = attrs.get('club_name', '')
                title = attrs.get('name', '通知')
                content = attrs.get('content', '')
                if not cn:
                    return json.dumps({'error': '创建通知需要指定社团名称'}, ensure_ascii=False)
                try:
                    conn.execute('INSERT INTO club_notices (club_name, title, content, created_at) VALUES (?, ?, ?, CURRENT_TIMESTAMP)', (cn, title, content))
                    conn.commit()
                    nid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                    result_data = {'success': True, 'id': nid, 'message': f'通知"{title}"已发布'}
                except Exception as e:
                    result_data = {'error': f'创建通知失败：{str(e)}'}
            elif entity_type == 'vote':
                cn = attrs.get('club_name', '')
                title = attrs.get('name', '投票')
                options = attrs.get('options', [])
                if not cn:
                    return json.dumps({'error': '创建投票需要指定社团名称'}, ensure_ascii=False)
                try:
                    conn.execute('INSERT INTO club_tools (club_name, tool_type, title, options, status, created_at) VALUES (?, "vote", ?, ?, "active", CURRENT_TIMESTAMP)', (cn, title, json.dumps(options, ensure_ascii=False)))
                    conn.commit()
                    vid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                    result_data = {'success': True, 'id': vid, 'message': f'投票"{title}"已创建'}
                except Exception as e:
                    result_data = {'error': f'创建投票失败：{str(e)}'}
            elif entity_type == 'survey':
                cn = attrs.get('club_name', '')
                title = attrs.get('name', '问卷')
                options = attrs.get('options', [])
                if not cn:
                    return json.dumps({'error': '创建问卷需要指定社团名称'}, ensure_ascii=False)
                try:
                    conn.execute('INSERT INTO club_tools (club_name, tool_type, title, options, status, created_at) VALUES (?, "survey", ?, ?, "active", CURRENT_TIMESTAMP)', (cn, title, json.dumps(options, ensure_ascii=False)))
                    conn.commit()
                    sid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                    result_data = {'success': True, 'id': sid, 'message': f'问卷"{title}"已创建'}
                except Exception as e:
                    result_data = {'error': f'创建问卷失败：{str(e)}'}
            elif entity_type == 'joint_activity':
                cn = attrs.get('club_name', '')
                title = attrs.get('name', '联合活动')
                desc = attrs.get('description', '')
                support = attrs.get('support_needed', '')
                if not cn:
                    return json.dumps({'error': '创建联合活动需要指定社团名称'}, ensure_ascii=False)
                try:
                    conn.execute('INSERT INTO joint_activities (club_name, title, description, support_needed, status, created_at) VALUES (?, ?, ?, ?, "open", CURRENT_TIMESTAMP)', (cn, title, desc, support))
                    conn.commit()
                    jid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                    result_data = {'success': True, 'id': jid, 'message': f'联合活动"{title}"已创建'}
                except Exception as e:
                    result_data = {'error': f'创建联合活动失败：{str(e)}'}
            elif entity_type == 'registration_form':
                cn = attrs.get('club_name', '')
                title = attrs.get('name', '报名表')
                if not cn:
                    return json.dumps({'error': '创建报名表需要指定社团名称'}, ensure_ascii=False)
                try:
                    conn.execute('INSERT INTO club_tools (club_name, tool_type, title, status, created_at) VALUES (?, "registration", ?, "active", CURRENT_TIMESTAMP)', (cn, title))
                    conn.commit()
                    rid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                    result_data = {'success': True, 'id': rid, 'message': f'报名表"{title}"已创建'}
                except Exception as e:
                    result_data = {'error': f'创建报名表失败：{str(e)}'}
            else:
                result_data = {'error': f'暂不支持创建类型：{entity_type}', 'hint': '当前支持：activity, recruitment, notification, vote, survey, joint_activity, registration_form'}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'update_entity':
            entity_type = arguments.get('entity_type', '')
            entity_id = arguments.get('entity_id', 0)
            updates = arguments.get('updates', {}) or {}
            result_data = {}
            # 权限校验
            if user_role == 'student':
                # 学生只能签到（checkin类型）
                if entity_type != 'checkin':
                    return json.dumps({'error': '学生角色无权执行此操作'}, ensure_ascii=False)
            if user_role == 'teacher':
                # 指导老师只能审批赋分和退社
                if entity_type not in ('scoring', 'workload'):
                    return json.dumps({'error': '指导老师只能审核赋分和工作量'}, ensure_ascii=False)
            if user_role == 'user':
                # 社团负责人只能操作自己社团的数据
                if entity_type in ('material', 'offcampus'):
                    return json.dumps({'error': '材料和校外活动审批仅管理员可操作'}, ensure_ascii=False)
            if entity_type == 'registration':
                action = updates.get('action', '')
                reason = updates.get('reason', '')
                if action in ('approve', 'reject'):
                    new_status = 'approved' if action == 'approve' else 'rejected'
                    try:
                        # 校验社团归属：非管理员只能审批自己社团的报名
                        reg = conn.execute('SELECT club_name FROM club_registrations WHERE id=?', (entity_id,)).fetchone()
                        if not reg:
                            return json.dumps({'error': '报名记录不存在'}, ensure_ascii=False)
                        if user_role == 'user' and reg['club_name'] != user_club:
                            return json.dumps({'error': '你只能审批自己社团的报名'}, ensure_ascii=False)
                        if user_role == 'teacher' and reg['club_name'] not in teacher_clubs:
                            return json.dumps({'error': '你只能审批所指导社团的报名'}, ensure_ascii=False)
                        conn.execute('UPDATE club_registrations SET status=?, reviewed_at=CURRENT_TIMESTAMP WHERE id=?', (new_status, entity_id))
                        conn.commit()
                        result_data = {'success': True, 'message': f'报名已{("通过" if action=="approve" else "拒绝")}'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '不支持的操作，请使用 approve 或 reject'}
            elif entity_type == 'workload':
                action = updates.get('action', '')
                if action in ('approve', 'reject'):
                    new_status = 'approved' if action == 'approve' else 'rejected'
                    try:
                        conn.execute('UPDATE workload_submissions SET status=? WHERE id=?', (new_status, entity_id))
                        conn.commit()
                        result_data = {'success': True, 'message': f'工作量已{("通过" if action=="approve" else "驳回")}'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '不支持的操作'}
            elif entity_type == 'scoring':
                action = updates.get('action', '')
                if action in ('approve', 'reject'):
                    new_status = 'approved' if action == 'approve' else 'rejected'
                    try:
                        conn.execute('UPDATE scoring_submissions SET status=? WHERE id=?', (new_status, entity_id))
                        conn.commit()
                        result_data = {'success': True, 'message': f'赋分已{("通过" if action=="approve" else "驳回")}'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '不支持的操作'}
            elif entity_type == 'material':
                action = updates.get('action', '')
                if action in ('approve', 'reject'):
                    new_status = 'approved' if action == 'approve' else 'rejected'
                    try:
                        conn.execute('UPDATE club_uploads SET status=? WHERE id=?', (new_status, entity_id))
                        conn.commit()
                        result_data = {'success': True, 'message': f'材料已{("通过" if action=="approve" else "驳回")}'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '不支持的操作'}
            elif entity_type == 'offcampus':
                action = updates.get('action', '')
                if action in ('approve', 'reject'):
                    new_status = 'approved' if action == 'approve' else 'rejected'
                    try:
                        conn.execute('UPDATE offcampus_requests SET status=? WHERE id=?', (new_status, entity_id))
                        conn.commit()
                        result_data = {'success': True, 'message': f'校外活动申请已{("通过" if action=="approve" else "驳回")}'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '不支持的操作'}
            elif entity_type == 'finance_permission':
                action = updates.get('action', '')
                # 权限校验：只有社团负责人和管理员可以审批财务权限
                if user_role == 'student':
                    return json.dumps({'error': '学生无权审批财务权限'}, ensure_ascii=False)
                if user_role == 'teacher':
                    return json.dumps({'error': '指导老师无权审批财务权限，请联系社团负责人'}, ensure_ascii=False)
                if action in ('approve', 'reject'):
                    new_status = 'approved' if action == 'approve' else 'rejected'
                    try:
                        record = conn.execute('SELECT * FROM finance_permissions WHERE id=?', (entity_id,)).fetchone()
                        if record:
                            # 社团负责人只能审批自己社团的财务权限
                            if user_role == 'user' and record['club_name'] != user_club:
                                return json.dumps({'error': '你只能审批自己社团的财务权限'}, ensure_ascii=False)
                            conn.execute('UPDATE finance_permissions SET status=?, reviewed_at=CURRENT_TIMESTAMP WHERE id=?', (new_status, entity_id))
                            if action == 'approve':
                                conn.execute('INSERT OR IGNORE INTO finance_managers (club_name, user_id, username, real_name, granted_by) VALUES (?, ?, ?, ?, ?)',
                                    (record['club_name'], record['user_id'], record['username'], record['real_name'], 'ai_assistant'))
                            conn.commit()
                            result_data = {'success': True, 'message': f'财务权限已{("通过" if action=="approve" else "拒绝")}'}
                        else:
                            result_data = {'error': '申请不存在'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '不支持的操作'}
            elif entity_type == 'activity':
                new_status = updates.get('status', '')
                if new_status:
                    try:
                        conn.execute('UPDATE checkin_sessions SET status=? WHERE id=?', (new_status, entity_id))
                        conn.commit()
                        result_data = {'success': True, 'message': f'活动状态已更新为{new_status}'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '请指定要更新的状态'}
            elif entity_type == 'checkin':
                action = updates.get('action', '')
                if action == 'checkin':
                    try:
                        session = conn.execute('SELECT id, club_name, status, checkin_code FROM checkin_sessions WHERE id=?', (entity_id,)).fetchone()
                        if not session:
                            return json.dumps({'error': '签到会话不存在'}, ensure_ascii=False)
                        if session['status'] != 'open':
                            return json.dumps({'error': '该签到已关闭'}, ensure_ascii=False)
                        conn.execute('INSERT INTO checkin_records (session_id, club_name, student_name, student_id, created_at) VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)',
                            (entity_id, session['club_name'], current_user.get('real_name', current_user.get('username', '')) if current_user else '', current_user.get('id', 0) if current_user else 0))
                        conn.commit()
                        result_data = {'success': True, 'message': '签到成功'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                elif action == 'checkout':
                    try:
                        conn.execute('UPDATE checkin_sessions SET status="closed", closed_at=CURRENT_TIMESTAMP WHERE id=?', (entity_id,))
                        conn.commit()
                        result_data = {'success': True, 'message': '签退成功，活动已关闭'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '签到操作请使用 checkin 或 checkout'}
            elif entity_type == 'member':
                action = updates.get('action', '')
                if action == 'join':
                    try:
                        if not current_user:
                            return json.dumps({'error': '请先登录'}, ensure_ascii=False)
                        club = updates.get('club_name', user_club)
                        if not club:
                            return json.dumps({'error': '请指定社团名称'}, ensure_ascii=False)
                        existing = conn.execute('SELECT id FROM club_members WHERE user_id=? AND club_name=?', (current_user['id'], club)).fetchone()
                        if existing:
                            return json.dumps({'error': '已经是该社团成员'}, ensure_ascii=False)
                        conn.execute('INSERT INTO club_members (club_name, user_id, username, real_name, joined_at) VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)',
                            (club, current_user['id'], current_user.get('username', ''), current_user.get('real_name', '')))
                        conn.commit()
                        result_data = {'success': True, 'message': f'已加入{club}'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                elif action == 'leave':
                    try:
                        if not current_user:
                            return json.dumps({'error': '请先登录'}, ensure_ascii=False)
                        conn.execute('DELETE FROM club_members WHERE id=?', (entity_id,))
                        conn.commit()
                        result_data = {'success': True, 'message': '已退出社团'}
                    except Exception as e:
                        result_data = {'error': str(e)}
                else:
                    result_data = {'error': '成员操作请使用 join 或 leave'}
            else:
                result_data = {'error': f'暂不支持更新类型：{entity_type}', 'hint': '当前支持：activity, registration, workload, scoring, material, offcampus, finance_permission, checkin, member'}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'send_notification':
            target = arguments.get('target', '')
            cn = arguments.get('club_name', '')
            content = arguments.get('content', '')
            channel = arguments.get('channel', 'system')
            # 权限校验
            if user_role == 'student':
                return json.dumps({'error': '学生角色无权发送通知'}, ensure_ascii=False)
            if user_role == 'teacher':
                return json.dumps({'error': '指导老师角色无权发送通知，如需通知请联系社团负责人'}, ensure_ascii=False)
            if user_role == 'user':
                # 社团负责人只能给自己社团发通知
                if not cn:
                    cn = user_club
                if cn != user_club:
                    return json.dumps({'error': f'你只能给自己社团（{user_club}）发送通知'}, ensure_ascii=False)
            result_data = {}
            try:
                if target == 'all_members' and cn:
                    count = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE club_name=?', (cn,)).fetchone()['c']
                    conn.execute('INSERT INTO club_notices (club_name, title, content, created_at) VALUES (?, ?, ?, CURRENT_TIMESTAMP)', (cn, '系统通知', content))
                    conn.commit()
                    result_data = {'success': True, 'message': f'通知已发送给{cn}全体成员（{count}人）'}
                elif target in ('admin', 'user', 'teacher', 'student'):
                    count = conn.execute('SELECT COUNT(*) as c FROM users WHERE role=?', (target,)).fetchone()['c']
                    result_data = {'success': True, 'message': f'通知已发送给{count}位{target}角色用户', 'note': '系统内通知已记录'}
                elif cn:
                    conn.execute('INSERT INTO club_notices (club_name, title, content, created_at) VALUES (?, ?, ?, CURRENT_TIMESTAMP)', (cn, '系统通知', content))
                    conn.commit()
                    result_data = {'success': True, 'message': f'通知已发送到{cn}'}
                else:
                    result_data = {'error': '请指定通知目标（target）或社团名称（club_name）'}
            except Exception as e:
                result_data = {'error': str(e)}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'generate_report':
            report_type = arguments.get('report_type', '')
            cn = arguments.get('club_name', '')
            date_from = arguments.get('date_from', '')
            date_to = arguments.get('date_to', '')
            # 权限校验
            if user_role == 'student':
                # 学生只能看自己社团的报告
                if not cn and user_clubs:
                    cn = user_clubs[0]
                if cn and cn not in user_clubs:
                    return json.dumps({'error': f'你只能查看自己加入的社团（{", ".join(user_clubs)}）的报告'}, ensure_ascii=False)
                # 学生不能看预警和全校报告
                if report_type in ('alert_summary',):
                    return json.dumps({'error': '预警报告仅管理员可查看'}, ensure_ascii=False)
            elif user_role == 'teacher':
                if not cn and teacher_clubs:
                    cn = teacher_clubs[0]
                if cn and cn not in teacher_clubs:
                    return json.dumps({'error': f'你只能查看所指导社团（{", ".join(teacher_clubs)}）的报告'}, ensure_ascii=False)
                if report_type in ('alert_summary',):
                    return json.dumps({'error': '预警报告仅管理员可查看'}, ensure_ascii=False)
            elif user_role == 'user':
                if not cn:
                    cn = user_club
                if cn and cn != user_club:
                    return json.dumps({'error': f'你只能查看自己社团（{user_club}）的报告'}, ensure_ascii=False)
                if report_type in ('alert_summary',):
                    return json.dumps({'error': '预警报告仅管理员可查看'}, ensure_ascii=False)
            result_data = {}
            if report_type == 'activity_overview':
                if cn:
                    total = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (cn,)).fetchone()['c']
                    completed = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=? AND is_completed=1', (cn,)).fetchone()['c']
                else:
                    total = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions').fetchone()['c']
                    completed = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE is_completed=1').fetchone()['c']
                clubs_data = conn.execute('SELECT club_name, COUNT(*) as c FROM checkin_sessions GROUP BY club_name ORDER BY c DESC LIMIT 10').fetchall()
                result_data = {'total_activities': total, 'completed': completed, 'completion_rate': round(completed/total*100) if total > 0 else 0, 'by_club': [dict(r) for r in clubs_data]}
            elif report_type == 'teacher_guidance':
                # 指导老师指导情况报告：综合 teacher_clubs + teacher_profiles + teacher_checkin_checkout
                if cn:
                    teacher_rows = conn.execute('SELECT tc.user_id, tc.club_name, tp.real_name FROM teacher_clubs tc LEFT JOIN teacher_profiles tp ON tc.user_id=tp.user_id WHERE tc.club_name=?', (cn,)).fetchall()
                else:
                    teacher_rows = conn.execute('SELECT tc.user_id, tc.club_name, tp.real_name FROM teacher_clubs tc LEFT JOIN teacher_profiles tp ON tc.user_id=tp.user_id').fetchall()
                items = []
                for tr in teacher_rows:
                    tid = tr['user_id']
                    tcn = tr['club_name']
                    tname = tr['real_name'] or ''
                    if not tname:
                        u = conn.execute('SELECT username FROM users WHERE id=?', (tid,)).fetchone()
                        tname = u['username'] if u else '未知'
                    # 签到签退指导次数
                    checked_out = conn.execute('SELECT COUNT(*) as c FROM teacher_checkin_checkout WHERE teacher_user_id=? AND club_name=? AND status="checked_out"', (tid, tcn)).fetchone()['c']
                    checked_in = conn.execute('SELECT COUNT(*) as c FROM teacher_checkin_checkout WHERE teacher_user_id=? AND club_name=? AND status="checked_in"', (tid, tcn)).fetchone()['c']
                    # 指导时长
                    duration_rows = conn.execute('SELECT checkin_time, checkout_time FROM teacher_checkin_checkout WHERE teacher_user_id=? AND club_name=? AND status="checked_out" AND checkin_time IS NOT NULL AND checkout_time IS NOT NULL', (tid, tcn)).fetchall()
                    total_minutes = 0
                    for dr in duration_rows:
                        try:
                            ci = datetime.strptime(str(dr['checkin_time']), '%Y-%m-%d %H:%M:%S')
                            co = datetime.strptime(str(dr['checkout_time']), '%Y-%m-%d %H:%M:%S')
                            total_minutes += max(0, (co - ci).total_seconds() / 60)
                        except:
                            pass
                    # 活动数
                    activity_count = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (tcn,)).fetchone()['c']
                    items.append({
                        'teacher_name': tname,
                        'club_name': tcn,
                        'guidance_count': checked_out,
                        'in_progress': checked_in,
                        'total_hours': round(total_minutes / 60, 1),
                        'activity_count': activity_count
                    })
                result_data = {'items': items}
            elif report_type == 'workload_stats':
                q = 'SELECT club_name, COUNT(*) as total, SUM(CASE WHEN status="approved" THEN 1 ELSE 0 END) as approved, SUM(CASE WHEN status="pending" THEN 1 ELSE 0 END) as pending FROM workload_submissions GROUP BY club_name ORDER BY total DESC'
                rows = conn.execute(q).fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif report_type == 'finance_summary':
                q = 'SELECT club_name, SUM(CASE WHEN type="income" THEN amount ELSE 0 END) as income, SUM(CASE WHEN type="expense" THEN amount ELSE 0 END) as expense FROM finance_records GROUP BY club_name ORDER BY income DESC'
                rows = conn.execute(q).fetchall()
                items = []
                for r in rows:
                    d = dict(r)
                    d['balance'] = d.get('income', 0) - d.get('expense', 0)
                    items.append(d)
                result_data = {'items': items}
            elif report_type == 'participation_trend':
                rows = conn.execute('SELECT cs.club_name, COUNT(DISTINCT cr.id) as participants, COUNT(cr.id) as checkins FROM checkin_sessions cs LEFT JOIN checkin_records cr ON cs.id=cr.session_id GROUP BY cs.club_name ORDER BY participants DESC LIMIT 10').fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif report_type == 'scoring_distribution':
                rows = conn.execute('SELECT ss.club_name, ss.status, COUNT(*) as count FROM scoring_submissions ss GROUP BY ss.club_name, ss.status ORDER BY ss.club_name').fetchall()
                result_data = {'items': [dict(r) for r in rows]}
            elif report_type == 'alert_summary':
                alerts = notification_agent.check_all() if hasattr(notification_agent, 'check_all') else []
                result_data = {'total_alerts': len(alerts), 'alerts': alerts[:10]}
            elif report_type == 'club_health':
                if cn:
                    report = data_insight_agent.generate_report(club_name=cn) if hasattr(data_insight_agent, 'generate_report') else {}
                    result_data = report
                else:
                    result_data = {'error': '社团健康报告需要指定社团名称'}
            else:
                result_data = {'error': f'未知报告类型：{report_type}'}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'ai_generate':
            gen_type = arguments.get('generate_type', '')
            params = arguments.get('params', {}) or {}
            result_data = {}
            if gen_type == 'copywriting':
                theme = params.get('theme', '社团活动')
                tone = params.get('tone', '活泼')
                cn = params.get('club_name', '')
                activity_type = params.get('activity_type', '')
                target = params.get('target_audience', '全体同学')
                prompt_text = f'请为{cn + "的" if cn else ""}{activity_type + "——" if activity_type else ""}"{theme}"写一段{tone}风格的宣传文案，目标受众是{target}。要求：吸引人、有号召力、200字以内。'
                reply = call_llm_api([{'role': 'user', 'content': prompt_text}], max_tokens=500)
                if reply:
                    result_data = {'success': True, 'content': reply, 'message': '文案已生成'}
                else:
                    result_data = {'error': '文案生成失败，AI服务不可用'}
            else:
                result_data = {'error': f'未知生成类型：{gen_type}'}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'recommend':
            rec_type = arguments.get('recommend_type', '')
            cn = arguments.get('club_name', '')
            result_data = {}
            if rec_type == 'club':
                # 查询用户已加入的社团
                joined_clubs = []
                joined_count = 0
                pending_count = 0
                if current_user:
                    joined_rows = conn.execute('SELECT DISTINCT club_name FROM club_members WHERE user_id=?', (user_id,)).fetchall()
                    joined_clubs = [r['club_name'] for r in joined_rows]
                    cadre_rows = conn.execute('SELECT DISTINCT club_name FROM club_cadres WHERE user_id=?', (user_id,)).fetchall()
                    for r in cadre_rows:
                        if r['club_name'] not in joined_clubs:
                            joined_clubs.append(r['club_name'])
                    joined_count = len(joined_clubs)
                    pending_row = conn.execute('SELECT COUNT(*) as c FROM club_registrations WHERE user_id=? AND status="pending"', (user_id,)).fetchone()
                    pending_count = pending_row['c'] if pending_row else 0
                # 每人最多加入2个社团
                remaining = max(0, 2 - joined_count - pending_count)
                # 排除已加入的社团
                clubs = conn.execute('SELECT cp.club_name, cp.category, cp.star_rating, cp.description, COUNT(cm.id) as member_count FROM club_profiles cp LEFT JOIN club_members cm ON cp.club_name=cm.club_name GROUP BY cp.club_name ORDER BY cp.star_rating DESC, member_count DESC').fetchall()
                filtered = [dict(r) for r in clubs if r['club_name'] not in joined_clubs]
                result_data = {
                    'recommendations': filtered[:5],
                    'joined_clubs': joined_clubs,
                    'joined_count': joined_count,
                    'pending_count': pending_count,
                    'remaining_slots': remaining,
                    'max_clubs': 2,
                    'message': f'您已加入{joined_count}个社团' + (f'、有{pending_count}个待审批' if pending_count else '') + f'，还可加入{remaining}个社团' if remaining > 0 else '，已达到上限（每人最多2个）'
                }
            elif rec_type == 'activity':
                activities = conn.execute('SELECT cs.activity_name, cs.club_name, cs.created_at, COUNT(cr.id) as participant_count FROM checkin_sessions cs LEFT JOIN checkin_records cr ON cs.id=cr.session_id WHERE cs.status="open" OR cs.is_completed=0 GROUP BY cs.id ORDER BY participant_count DESC, cs.created_at DESC LIMIT 5').fetchall()
                result_data = {'recommendations': [dict(r) for r in activities], 'message': '基于参与热度的活动推荐'}
            elif rec_type == 'joint_partner':
                if not cn:
                    return json.dumps({'error': '推荐联合活动伙伴需要指定当前社团名称'}, ensure_ascii=False)
                other_clubs = conn.execute('SELECT cp.club_name, cp.category, COUNT(cs.id) as activity_count FROM club_profiles cp LEFT JOIN checkin_sessions cs ON cp.club_name=cs.club_name WHERE cp.club_name!=? GROUP BY cp.club_name ORDER BY activity_count DESC LIMIT 5', (cn,)).fetchall()
                result_data = {'recommendations': [dict(r) for r in other_clubs], 'message': f'适合与{cn}联合活动的社团推荐'}
            elif rec_type == 'recruitment':
                # 推荐进行中的招募活动，优先推荐用户未报名的
                recruitments = conn.execute('SELECT r.id, r.club_name, r.title, r.description, r.status, r.created_at, COUNT(rs.id) as signup_count FROM recruitments r LEFT JOIN recruitment_signups rs ON r.id=rs.recruitment_id WHERE r.status="open" OR r.status="active" GROUP BY r.id ORDER BY r.created_at DESC LIMIT 5').fetchall()
                items = [dict(r) for r in recruitments]
                # 如果是学生，标记是否已报名
                if user_role == 'student' and items:
                    for item in items:
                        signed = conn.execute('SELECT COUNT(*) as c FROM recruitment_signups WHERE recruitment_id=? AND user_id=?', (item['id'], user_id)).fetchone()['c']
                        item['already_signed'] = signed > 0
                result_data = {'recommendations': items, 'message': '进行中的招募活动推荐'}
            else:
                result_data = {'error': f'未知推荐类型：{rec_type}'}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'query_database':
            sql = arguments.get('sql', '').strip()
            query_desc = arguments.get('description', '')
            if not sql:
                return json.dumps({'error': '请提供SQL查询语句'}, ensure_ascii=False)
            # 安全检查：只允许SELECT
            sql_upper = sql.upper().strip()
            forbidden_keywords = ['INSERT', 'UPDATE', 'DELETE', 'DROP', 'ALTER', 'CREATE', 'ATTACH', 'DETACH', 'REPLACE', 'PRAGMA']
            for kw in forbidden_keywords:
                if kw in sql_upper.split():
                    return json.dumps({'error': f'安全限制：禁止执行{kw}操作，仅允许SELECT查询'}, ensure_ascii=False)
            if not sql_upper.startswith('SELECT') and not sql_upper.startswith('WITH'):
                return json.dumps({'error': '安全限制：仅允许SELECT查询语句'}, ensure_ascii=False)
            import re as _re
            sql_lower = sql.lower()
            # 禁止的表（敏感数据）
            forbidden_tables = ['users', 'sms_codes', 'ai_chat_history', 'ai_pets']
            for ft in forbidden_tables:
                if _re.search(r'\b' + ft + r'\b', sql_lower):
                    if user_role != 'admin':
                        return json.dumps({'error': f'安全限制：无权查询{ft}表'}, ensure_ascii=False)
            # 用户隔离：非管理员自动注入社团过滤
            # 判断SQL中涉及的表
            club_columns = {
                'club_members': 'club_name',
                'club_cadres': 'club_name',
                'club_notices': 'club_name',
                'club_profiles': 'club_name',
                'club_departments': 'club_name',
                'club_tools': 'club_name',
                'club_uploads': 'club_name',
                'club_registrations': 'club_name',
                'club_showcase': 'club_name',
                'club_favorites': 'club_name',
                'checkin_sessions': 'club_name',
                'checkin_records': 'club_name',
                'finance_records': 'club_name',
                'finance_permissions': 'club_name',
                'finance_managers': 'club_name',
                'workload_submissions': 'club_name',
                'scoring_submissions': 'club_name',
                'scoring_submission_items': None,  # 通过submission_id关联
                'recruitments': 'club_name',
                'recruitment_signups': None,
                'teacher_clubs': 'club_name',
                'teacher_checkin_checkout': 'club_name',
                'joint_activities': 'club_name',
                'joint_replies': None,
                'offcampus_requests': 'club_name',
                'excellent_clubs': 'club_name',
                'excellent_activities': None,
                'feedbacks': 'club_name',
                'tree_hole': 'club_name',
                'doc_index': 'club_name',
                'notices': None,
                'notifications': None,
                'user_profiles': None,
                'teacher_profiles': None,
                'final_credits': None,
                'scoring_rules': None,
                'scoring_club_overrides': 'club_name',
                'location_checkins': None,
                'quit_applications': None,
                'teacher_club_requests': None,
            }
            if user_role != 'admin':
                # 检测SQL中涉及的含club_name列的表
                tables_in_sql = []
                for tbl in club_columns:
                    if _re.search(r'\b' + tbl + r'\b', sql_lower):
                        tables_in_sql.append(tbl)
                # 非管理员：如果有含club_name的表，需要注入过滤
                if tables_in_sql:
                    # 确定用户可访问的社团
                    if user_role == 'student':
                        allowed_clubs = user_clubs
                    elif user_role == 'teacher':
                        allowed_clubs = teacher_clubs
                    else:
                        allowed_clubs = [user_club] if user_club else []
                    if not allowed_clubs:
                        return json.dumps({'error': '你尚未关联任何社团，无法查询数据'}, ensure_ascii=False)
                    # 检查SQL是否已有club_name过滤
                    has_club_filter = 'club_name' in sql_lower
                    if not has_club_filter:
                        # 注入WHERE条件
                        club_placeholders = ','.join(['?' for _ in allowed_clubs])
                        if 'WHERE' in sql_upper:
                            sql = sql.replace('WHERE', f'WHERE club_name IN ({club_placeholders}) AND', 1)
                        elif 'GROUP BY' in sql_upper:
                            sql = sql.replace('GROUP BY', f'WHERE club_name IN ({club_placeholders}) GROUP BY', 1)
                        elif 'ORDER BY' in sql_upper:
                            sql = sql.replace('ORDER BY', f'WHERE club_name IN ({club_placeholders}) ORDER BY', 1)
                        elif 'LIMIT' in sql_upper:
                            sql = sql.replace('LIMIT', f'WHERE club_name IN ({club_placeholders}) LIMIT', 1)
                        else:
                            sql += f' WHERE club_name IN ({club_placeholders})'
                        # 执行带参数
                        # 限制返回条数
                        if 'LIMIT' not in sql_upper:
                            sql += ' LIMIT 50'
                        try:
                            rows = conn.execute(sql, allowed_clubs).fetchall()
                        except Exception as e:
                            return json.dumps({'error': f'SQL执行错误：{str(e)}'}, ensure_ascii=False)
                        result_data = {'items': [dict(r) for r in rows], 'total': len(rows), 'description': query_desc, 'sql': sql}
                        return json.dumps(result_data, ensure_ascii=False, default=str)
            # 管理员或已有club_name过滤的查询：直接执行
            if 'LIMIT' not in sql_upper:
                sql += ' LIMIT 50'
            try:
                rows = conn.execute(sql).fetchall()
            except Exception as e:
                return json.dumps({'error': f'SQL执行错误：{str(e)}'}, ensure_ascii=False)
            result_data = {'items': [dict(r) for r in rows], 'total': len(rows), 'description': query_desc, 'sql': sql}
            return json.dumps(result_data, ensure_ascii=False, default=str)

        elif function_name == 'generate_dashboard':
            dashboard_type = arguments.get('dashboard_type', '')
            filter_club = arguments.get('club_name', '') or user_club
            if not dashboard_type:
                return json.dumps({'error': '请指定看板类型'}, ensure_ascii=False)
            
            dashboard = {'type': 'dashboard', 'dashboard_type': dashboard_type, 'charts': [], 'summary': ''}
            
            try:
                if dashboard_type == 'member_activity':
                    # 成员活跃度：各社团成员数 vs 活动数
                    if filter_club:
                        rows = conn.execute('SELECT cm.user_id, u.username, COUNT(DISTINCT cs.id) as activity_count FROM club_members cm LEFT JOIN users u ON cm.user_id=u.id LEFT JOIN checkin_sessions cs ON cs.club_name=cm.club_name WHERE cm.club_name=? GROUP BY cm.user_id ORDER BY activity_count DESC LIMIT 20', (filter_club,)).fetchall()
                        labels = [r['username'] or f'用户{r["user_id"]}' for r in rows]
                        data = [r['activity_count'] for r in rows]
                        dashboard['title'] = f'{filter_club}成员活跃度'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '成员活动参与次数',
                            'labels': labels,
                            'datasets': [{'label': '参与次数', 'data': data}]
                        })
                        dashboard['summary'] = f'{filter_club}共有{len(rows)}名成员，最活跃成员参与了{max(data) if data else 0}次活动'
                    else:
                        rows = conn.execute('SELECT club_name, COUNT(DISTINCT user_id) as member_count, 0 as activity_count FROM club_members GROUP BY club_name ORDER BY member_count DESC LIMIT 15').fetchall()
                        # Get activity count per club
                        for r in rows:
                            ac = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (r['club_name'],)).fetchone()
                            r = dict(r)
                            r['activity_count'] = ac['c'] if ac else 0
                        labels = [r['club_name'] for r in rows]
                        member_data = [r['member_count'] for r in rows]
                        activity_data = [r['activity_count'] for r in rows]
                        dashboard['title'] = '各社团成员活跃度'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '成员数 vs 活动数',
                            'labels': labels,
                            'datasets': [{'label': '成员数', 'data': member_data}, {'label': '活动数', 'data': activity_data}]
                        })
                        dashboard['summary'] = f'共{len(rows)}个社团，平均成员数{sum(member_data)//len(member_data) if member_data else 0}人'

                elif dashboard_type == 'workload_distribution':
                    if filter_club:
                        rows = conn.execute('SELECT ws.student_name, ws.activity_count, ws.other_score, ws.total_score FROM workload_submissions ws WHERE ws.club_name=? AND ws.status="approved" ORDER BY ws.total_score DESC LIMIT 20', (filter_club,)).fetchall()
                    else:
                        rows = conn.execute('SELECT club_name, COUNT(*) as submit_count, AVG(total_score) as avg_score FROM workload_submissions WHERE status="approved" GROUP BY club_name ORDER BY avg_score DESC LIMIT 15').fetchall()
                    if 'student_name' in (rows[0].keys() if rows else []):
                        labels = [r['student_name'] for r in rows]
                        scores = [r['total_score'] or 0 for r in rows]
                        activities = [r['activity_count'] or 0 for r in rows]
                        dashboard['title'] = f'{filter_club}工作量分布'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '成员工作量得分',
                            'labels': labels,
                            'datasets': [{'label': '活动次数', 'data': activities, 'yAxisID': 'y'}, {'label': '总分', 'data': scores, 'yAxisID': 'y1'}]
                        })
                    else:
                        labels = [r['club_name'] for r in rows]
                        counts = [r['submit_count'] for r in rows]
                        avgs = [round(r['avg_score'] or 0, 1) for r in rows]
                        dashboard['title'] = '各社团工作量分布'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '提交数 vs 平均分',
                            'labels': labels,
                            'datasets': [{'label': '提交数', 'data': counts}, {'label': '平均分', 'data': avgs}]
                        })
                    dashboard['summary'] = f'共{len(rows)}条记录'

                elif dashboard_type == 'finance_overview':
                    if filter_club:
                        rows = conn.execute("SELECT date, type, amount FROM finance_records WHERE club_name=? ORDER BY date DESC LIMIT 30", (filter_club,)).fetchall()
                        income = sum(r['amount'] for r in rows if r['type'] == 'income')
                        expense = sum(r['amount'] for r in rows if r['type'] == 'expense')
                        dashboard['title'] = f'{filter_club}财务概览'
                        dashboard['charts'].append({
                            'chart_type': 'doughnut',
                            'title': '收入 vs 支出',
                            'labels': ['收入', '支出'],
                            'datasets': [{'label': '金额', 'data': [income, expense]}]
                        })
                        dashboard['summary'] = f'收入¥{income}，支出¥{expense}，余额¥{income - expense}'
                    else:
                        rows = conn.execute("SELECT club_name, SUM(CASE WHEN type='income' THEN amount ELSE 0 END) as income, SUM(CASE WHEN type='expense' THEN amount ELSE 0 END) as expense FROM finance_records GROUP BY club_name ORDER BY income DESC LIMIT 15").fetchall()
                        labels = [r['club_name'] for r in rows]
                        incomes = [r['income'] or 0 for r in rows]
                        expenses = [r['expense'] or 0 for r in rows]
                        dashboard['title'] = '各社团财务概览'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '各社团收入vs支出',
                            'labels': labels,
                            'datasets': [{'label': '收入', 'data': incomes}, {'label': '支出', 'data': expenses}]
                        })
                        dashboard['summary'] = f'共{len(rows)}个社团有财务记录'

                elif dashboard_type == 'attendance':
                    if filter_club:
                        rows = conn.execute('SELECT cs.activity_name, COUNT(DISTINCT c.user_id) as checkin_count FROM checkin_sessions cs LEFT JOIN checkins c ON c.session_id=cs.id WHERE cs.club_name=? GROUP BY cs.id ORDER BY cs.id DESC LIMIT 15', (filter_club,)).fetchall()
                        labels = [r['activity_name'] for r in rows]
                        data = [r['checkin_count'] for r in rows]
                        dashboard['title'] = f'{filter_club}签到考勤'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '各活动签到人数',
                            'labels': labels,
                            'datasets': [{'label': '签到人数', 'data': data}]
                        })
                    else:
                        rows = conn.execute('SELECT club_name, COUNT(*) as total_checkins FROM checkins c JOIN checkin_sessions cs ON c.session_id=cs.id GROUP BY cs.club_name ORDER BY total_checkins DESC LIMIT 15').fetchall()
                        labels = [r['club_name'] for r in rows]
                        data = [r['total_checkins'] for r in rows]
                        dashboard['title'] = '各社团签到统计'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '签到总数',
                            'labels': labels,
                            'datasets': [{'label': '签到次数', 'data': data}]
                        })
                    dashboard['summary'] = f'共{len(rows)}条记录'

                elif dashboard_type == 'scoring_distribution':
                    if filter_club:
                        rows = conn.execute('SELECT student_name, final_score FROM scoring_submissions WHERE club_name=? AND status="approved" ORDER BY final_score DESC LIMIT 20', (filter_club,)).fetchall()
                        labels = [r['student_name'] for r in rows]
                        data = [r['final_score'] or 0 for r in rows]
                        dashboard['title'] = f'{filter_club}赋分分布'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '成员赋分',
                            'labels': labels,
                            'datasets': [{'label': '赋分', 'data': data}]
                        })
                    else:
                        rows = conn.execute('SELECT club_name, COUNT(*) as cnt, AVG(final_score) as avg_score FROM scoring_submissions WHERE status="approved" GROUP BY club_name ORDER BY avg_score DESC LIMIT 15').fetchall()
                        labels = [r['club_name'] for r in rows]
                        avgs = [round(r['avg_score'] or 0, 1) for r in rows]
                        dashboard['title'] = '各社团赋分分布'
                        dashboard['charts'].append({
                            'chart_type': 'bar',
                            'title': '平均赋分',
                            'labels': labels,
                            'datasets': [{'label': '平均赋分', 'data': avgs}]
                        })
                    dashboard['summary'] = f'共{len(rows)}条记录'

                elif dashboard_type == 'club_overview':
                    rows = conn.execute('SELECT u.club_name, COUNT(DISTINCT u.id) as member_count, COUNT(DISTINCT cs.id) as activity_count FROM users u LEFT JOIN checkin_sessions cs ON cs.club_name=u.club_name WHERE u.role="student" AND u.club_name!="" GROUP BY u.club_name ORDER BY member_count DESC LIMIT 15').fetchall()
                    labels = [r['club_name'] for r in rows]
                    members = [r['member_count'] for r in rows]
                    activities = [r['activity_count'] for r in rows]
                    dashboard['title'] = '社团概览'
                    dashboard['charts'].append({
                        'chart_type': 'bar',
                        'title': '成员数 vs 活动数',
                        'labels': labels,
                        'datasets': [{'label': '成员数', 'data': members}, {'label': '活动数', 'data': activities}]
                    })
                    # Also add a pie chart for category distribution
                    cat_rows = conn.execute('SELECT category, COUNT(*) as cnt FROM clubs WHERE category!="" GROUP BY category ORDER BY cnt DESC').fetchall()
                    if cat_rows:
                        dashboard['charts'].append({
                            'chart_type': 'doughnut',
                            'title': '社团类别分布',
                            'labels': [r['category'] for r in cat_rows],
                            'datasets': [{'label': '数量', 'data': [r['cnt'] for r in cat_rows]}]
                        })
                    dashboard['summary'] = f'共{len(rows)}个社团'

                elif dashboard_type == 'activity_trend':
                    rows = conn.execute("SELECT DATE(created_at) as date, COUNT(*) as cnt FROM checkin_sessions WHERE created_at >= date('now', '-30 days') GROUP BY DATE(created_at) ORDER BY date").fetchall()
                    if not rows:
                        rows = conn.execute("SELECT DATE(created_at) as date, COUNT(*) as cnt FROM checkin_sessions GROUP BY DATE(created_at) ORDER BY date DESC LIMIT 30").fetchall()
                        rows = list(reversed(rows))
                    labels = [r['date'] for r in rows]
                    data = [r['cnt'] for r in rows]
                    dashboard['title'] = '近30天活动趋势'
                    dashboard['charts'].append({
                        'chart_type': 'line',
                        'title': '每日活动数',
                        'labels': labels,
                        'datasets': [{'label': '活动数', 'data': data, 'borderColor': 'rgb(75, 192, 192)', 'tension': 0.3}]
                    })
                    dashboard['summary'] = f'近{len(rows)}天共{sum(data)}次活动'

                return json.dumps(dashboard, ensure_ascii=False, default=str)
            except Exception as e:
                return json.dumps({'error': f'生成看板失败：{str(e)}'}, ensure_ascii=False)

        return json.dumps({'error': f'未知工具：{function_name}'}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({'error': str(e)}, ensure_ascii=False)
    finally:
        conn.close()


class ApprovalAgent:
    APPROVAL_TYPES = ('material', 'workload', 'registration', 'finance_permission', 'offcampus')

    def __init__(self, db_instance):
        self.db = db_instance

    def _get_conn(self):
        return self.db.get_conn()

    def _extract_file_content(self, file_path, file_type):
        content = ''
        try:
            if not file_path:
                return ''
            actual_path = storage.get_path(file_path)
            if not actual_path:
                return ''
            file_path = actual_path
            ext = (file_type or '').lower()
            if ext in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp'):
                try:
                    from PIL import Image
                    img = Image.open(file_path)
                    img_text = getattr(img, 'text', {}) or {}
                    if img_text:
                        content = ' '.join(str(v) for v in img_text.values() if v)
                except Exception:
                    pass
                fname = os.path.basename(file_path).lower()
                img_scene_keywords = {
                    '招新': ['招新', '报名', '纳新', '招募'],
                    '活动': ['活动', '比赛', '表演', '展演', '会议', '培训'],
                    '签到': ['签到', '考勤', '出勤'],
                    '颁奖': ['颁奖', '获奖', '荣誉', '表彰'],
                    '合影': ['合影', '合照', '集体照', '团队'],
                    '现场': ['现场', '场景', '实况'],
                    '海报': ['海报', '宣传', '展板'],
                }
                for scene, kws in img_scene_keywords.items():
                    if any(kw in fname for kw in kws):
                        content += ' ' + scene
                return content.strip()
            elif ext == 'pdf':
                try:
                    import fitz
                    doc = fitz.open(file_path)
                    for page_num in range(min(len(doc), 10)):
                        content += doc[page_num].get_text() + '\n'
                    doc.close()
                except ImportError:
                    try:
                        import PyPDF2
                        with open(file_path, 'rb') as f:
                            reader = PyPDF2.PdfReader(f)
                            for page_num in range(min(len(reader.pages), 10)):
                                content += reader.pages[page_num].extract_text() + '\n'
                    except Exception:
                        pass
            elif ext in ('doc', 'docx'):
                try:
                    import docx
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        content += para.text + '\n'
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                content += cell.text + ' '
                        content += '\n'
                except ImportError:
                    try:
                        import subprocess
                        result = subprocess.run(['python', '-m', 'docx2txt', 'convert', file_path], capture_output=True, text=True, timeout=10)
                        if result.returncode == 0:
                            content = result.stdout
                    except Exception:
                        pass
            elif ext in ('xls', 'xlsx'):
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path, read_only=True)
                    for sheet in wb.worksheets:
                        for row in sheet.iter_rows(max_row=50, values_only=True):
                            content += ' '.join(str(c) for c in row if c) + '\n'
                    wb.close()
                except Exception:
                    pass
            elif ext == 'txt':
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read(5000)
                except Exception:
                    pass
        except Exception:
            pass
        return content.strip()

    RELEVANCE_KEYWORDS = {
        '招新': ['招新', '报名', '招募', '纳新', '登记', '申请', '加入', '新成员'],
        '活动方案': ['方案', '策划', '计划', '安排', '流程', '规划', '日程'],
        '活动总结': ['总结', '小结', '回顾', '成果', '汇报', '心得', '体会'],
        '财务': ['财务', '报销', '经费', '预算', '收支', '账单', '费用', '申请表'],
        '签到': ['签到', '考勤', '点名', '出勤', '签到表'],
        '五四': ['五四', '青年节', '五四青年'],
        '比赛': ['比赛', '竞赛', '评选', '参赛', '决赛', '初赛'],
        '培训': ['培训', '讲座', '教学', '课程', '辅导'],
        '校外': ['校外', '外出', '参观', '交流', '联谊'],
        '会议': ['会议', '纪要', '讨论', '决议'],
        '宣传': ['海报', '宣传', '推广', '文案'],
        '荣誉': ['荣誉', '获奖', '奖项', '证书', '表彰'],
        '社团章程': ['章程', '制度', '规章', '规范', '守则'],
        '指导': ['指导', '辅导', '点评', '建议'],
        '展演': ['展演', '表演', '演出', '展示', '风采'],
    }

    FILE_KEYWORDS = {
        '报名表': ['报名', '登记', '申请表', '注册'],
        '签到表': ['签到', '考勤', '出勤', '点名'],
        '策划书': ['策划', '方案', '计划', '安排'],
        '总结': ['总结', '小结', '汇报', '回顾'],
        '经费申请': ['经费', '报销', '财务', '预算', '申请表'],
        '照片': ['照片', '图片', '合影', '现场', '截图'],
        '简历': ['简历', '履历'],
        '章程': ['章程', '制度', '规章'],
        '流程': ['流程', '操作', '指南', '手册'],
        '报告': ['报告', '调研', '分析'],
    }

    def _simulate_relevance(self, body_text, title_text, file_names, file_types, category, file_contents=None):
        relevance_score = 10
        relevance_reasons = []
        body = (body_text or '').lower()
        title = (title_text or '').lower()
        import re as _re
        body_clean = _re.sub(r'<[^>]+>', '', body).strip()
        title_clean = _re.sub(r'<[^>]+>', '', title).strip()
        combined_text = body_clean + ' ' + title_clean
        combined_text_raw = body + ' ' + title
        has_image = any(ft in file_types for ft in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp'))
        has_doc = any(ft in file_types for ft in ('pdf', 'doc', 'docx', 'xls', 'xlsx'))
        file_name_text = ' '.join(file_names or []).lower()
        file_name_clean = _re.sub(r'[_\-\.]', ' ', file_name_text)
        all_text = combined_text + ' ' + file_name_text + ' ' + file_name_clean
        if not file_names:
            relevance_score -= 10
            relevance_reasons.append('无上传材料')
        if not combined_text.strip() or len(combined_text.strip()) < 5:
            relevance_score -= 5
            relevance_reasons.append('正文内容过少')
        elif len(combined_text.strip()) < 10:
            relevance_score -= 3
        title_in_file = False
        if title_clean and len(title_clean) >= 2 and file_name_text:
            for fn in (file_names or []):
                fn_lower = fn.lower()
                fn_clean = _re.sub(r'[_\-\.]', ' ', fn_lower)
                if title_clean in fn_clean or title_clean in fn_lower:
                    title_in_file = True
                    break
                title_parts = [p for p in _re.split(r'[\s,，、]', title_clean) if len(p) >= 2]
                part_hits = sum(1 for p in title_parts if p in fn_clean)
                if len(title_parts) > 0 and part_hits >= max(1, len(title_parts) // 2):
                    title_in_file = True
                    break
        if title_in_file:
            relevance_score += 20
            relevance_reasons.append('标题与上传材料名称高度匹配')
        matched_topics = []
        for topic, keywords in self.RELEVANCE_KEYWORDS.items():
            text_hits = sum(1 for kw in keywords if kw in combined_text or kw in combined_text_raw)
            file_hits = sum(1 for kw in keywords if kw in file_name_text or kw in file_name_clean)
            if text_hits > 0 and file_hits > 0:
                matched_topics.append(topic)
                relevance_score += min(40, 20 + text_hits * 5 + file_hits * 5)
                relevance_reasons.append(f'正文与材料均涉及「{topic}」')
                if has_image:
                    relevance_score += 5
                if has_doc:
                    relevance_score += 5
                if has_image and has_doc:
                    relevance_score += 5
                    relevance_reasons.append('材料类型丰富，佐证充分')
            elif text_hits > 0 and (has_image or has_doc):
                relevance_score += min(8, 3 + text_hits * 2)
            elif text_hits > 0:
                relevance_score += min(5, 2 + text_hits * 1)
        for doc_type, keywords in self.FILE_KEYWORDS.items():
            file_hits = sum(1 for kw in keywords if kw in file_name_text or kw in file_name_clean)
            text_hits = sum(1 for kw in keywords if kw in combined_text or kw in combined_text_raw)
            if file_hits > 0 and text_hits > 0:
                relevance_score += 10
                if f'正文与材料均涉及' not in str(relevance_reasons):
                    relevance_reasons.append(f'上传{doc_type}与正文内容匹配')
        file_content_text = ''
        if file_contents:
            file_content_text = ' '.join(file_contents).lower()
        if file_content_text and combined_text.strip():
            content_matched = False
            for topic, keywords in self.RELEVANCE_KEYWORDS.items():
                text_hits = sum(1 for kw in keywords if kw in combined_text or kw in combined_text_raw)
                content_hits = sum(1 for kw in keywords if kw in file_content_text)
                if text_hits > 0 and content_hits > 0:
                    content_matched = True
                    if topic not in matched_topics:
                        relevance_score += min(30, 12 + text_hits * 4 + content_hits * 4)
                        relevance_reasons.append(f'正文与材料内容均涉及「{topic}」')
            body_words = set(combined_text)
            content_words = set(file_content_text)
            common_words = body_words & content_words
            meaningful_common = [w for w in common_words if len(w) >= 2]
            if len(meaningful_common) > 15:
                relevance_score += min(15, len(meaningful_common))
                if not content_matched:
                    relevance_reasons.append('正文与材料提取内容高度匹配')
            elif len(meaningful_common) > 8:
                relevance_score += min(8, len(meaningful_common))
                if not content_matched:
                    relevance_reasons.append('正文与材料提取内容部分匹配')
            if has_image and file_content_text:
                img_content_parts = [fc for fc in (file_contents or []) if fc]
                img_matched = False
                for icp in img_content_parts:
                    icp_lower = icp.lower()
                    for topic, keywords in self.RELEVANCE_KEYWORDS.items():
                        if any(kw in combined_text for kw in keywords) and any(kw in icp_lower for kw in keywords):
                            img_matched = True
                            break
                    if img_matched:
                        break
                if img_matched:
                    relevance_score += 10
                    relevance_reasons.append('图片内容与正文主题相关')
        if not matched_topics and not title_in_file and not file_content_text:
            if combined_text.strip() and file_name_text.strip():
                text_chars = set(combined_text)
                file_chars = set(file_name_clean)
                overlap = len(text_chars & file_chars)
                if overlap > 30:
                    relevance_score += 3
                else:
                    relevance_score -= 5
                    relevance_reasons.append('正文与上传材料内容无明显关联')
            else:
                relevance_score -= 5
                if '正文内容过少' not in str(relevance_reasons):
                    relevance_reasons.append('缺乏有效内容进行关联分析')
        if category:
            cat_map = {'honor': '荣誉', 'activity': '活动', 'innovation': '创新'}
            cat_name = cat_map.get(category, category)
            if cat_name in combined_text or cat_name in file_name_text:
                relevance_score += 5
        if body_clean and len(body_clean) > 30:
            relevance_score += 5
        rule_score = max(0, min(100, relevance_score))
        _, api_key = get_llm_config()
        if api_key and combined_text.strip() and (file_name_text.strip() or file_content_text.strip()):
            try:
                ai_score = self._ai_relevance_analyze(body_text or '', title_text or '', file_names or [], file_content_text, category or '')
                if ai_score is not None:
                    relevance_score = round(rule_score * 0.35 + ai_score * 0.65, 1)
                    relevance_reasons.append(f'AI语义分析相关度{ai_score}分（规则{rule_score}分）')
                else:
                    relevance_score = rule_score
            except Exception:
                relevance_score = rule_score
        else:
            relevance_score = rule_score
        if not relevance_reasons:
            if relevance_score >= 70:
                relevance_reasons.append('正文与材料内容基本一致')
            elif relevance_score >= 50:
                relevance_reasons.append('正文与材料有部分关联')
            else:
                relevance_reasons.append('正文与材料内容关联较弱')
        return relevance_score, relevance_reasons

    def _ai_relevance_analyze(self, body_text, title_text, file_names, file_content_text, category):
        _, api_key = get_llm_config()
        if not api_key:
            return None
        file_desc = '文件名：' + '、'.join(file_names[:5])
        if file_content_text:
            file_desc += '\n文件内容摘要：' + file_content_text[:800]
        cat_map = {'honor': '荣誉奖项', 'activity': '活动参与', 'innovation': '创新亮点'}
        cat_name = cat_map.get(category, category)
        prompt = f'你是一个严格的社团材料审批助手。请分析以下"正文"与"上传材料"的语义相关度，给出0-100的分数。\n\n评分标准（请严格按此标准，不要给同情分）：\n- 90-100：正文内容与材料高度一致，材料完全支撑正文，内容详实匹配\n- 70-89：正文与材料大部分相关，材料基本支撑正文\n- 40-69：正文与材料部分相关，但有不匹配之处\n- 15-39：正文与材料关联较弱，材料不能有效支撑正文\n- 0-14：正文与材料几乎无关，或正文内容过少无法判断相关性\n\n特别注意：\n- 如果正文内容非常简短（少于10个有效字符），即使标题与文件名相似，相关度也不应超过40分\n- 如果正文与材料内容没有实质性的语义关联，必须给低分（0-30）\n- 只有正文内容与材料内容在语义上确实匹配时，才给高分\n\n材料类别：{cat_name}\n正文标题：{title_text}\n正文内容：{body_text[:1000]}\n上传材料：{file_desc}\n\n请只返回一个0-100的整数分数，不要返回其他内容。'
        messages = [
            {'role': 'system', 'content': '你是一个严格的专业社团材料审批助手，擅长分析文本与材料的相关度。评分要严格区分相关和不相关：高度相关给70-100分，部分相关给40-69分，弱相关给15-39分，几乎无关给0-14分。只返回分数，不要解释。'},
            {'role': 'user', 'content': prompt}
        ]
        result = call_llm_api(messages, max_tokens=20)
        if result:
            try:
                score = int(re.search(r'\d+', result.strip()).group())
                return max(0, min(100, score))
            except (AttributeError, ValueError):
                pass
        return None

    def _generate_approval_reason(self, relevance_score, star_score, combined_score, star_rating, recommendation, relevance_reasons):
        parts = []
        if relevance_score >= 90:
            parts.append(f'正文与材料相关度极高（{relevance_score}分）')
        elif relevance_score >= 70:
            parts.append(f'正文与材料相关度较高（{relevance_score}分）')
        elif relevance_score >= 50:
            parts.append(f'正文与材料相关度一般（{relevance_score}分）')
        elif relevance_score >= 30:
            parts.append(f'正文与材料相关度较低（{relevance_score}分）')
        else:
            parts.append(f'正文与材料相关度极低（{relevance_score}分）')
        if star_rating >= 4:
            parts.append(f'星级较高（{star_rating}星）')
        elif star_rating >= 3:
            parts.append(f'星级一般（{star_rating}星）')
        elif star_rating > 0:
            parts.append(f'星级较低（{star_rating}星）')
        else:
            parts.append('暂无星级评定')
        if recommendation == 'approve':
            if relevance_score >= 70 and star_rating >= 4:
                return '，'.join(parts) + '，建议直接通过'
            elif relevance_score >= 70:
                return '，'.join(parts) + '，建议通过'
            else:
                return '，'.join(parts) + '，综合质量达标，建议通过'
        else:
            if relevance_score < 50:
                return '，'.join(parts) + '，内容不匹配，建议驳回并说明原因'
            elif relevance_score < 70:
                return '，'.join(parts) + '，建议驳回并要求重新上传'
            else:
                return '，'.join(parts) + '，综合质量不达标，建议驳回'
        return '，'.join(parts)

    def analyze_material(self, group_id=None, club_name=None):
        conn = self._get_conn()
        try:
            source_filter = '(source="upload" OR source="activity" OR source IS NULL)'
            if group_id:
                rows = conn.execute(
                    f'SELECT id, club_name, file_name, file_path, description, status, reject_reason, upload_time, group_id, category, file_type FROM club_uploads WHERE group_id=? AND status="pending" AND {source_filter}',
                    (group_id,)).fetchall()
            elif club_name:
                rows = conn.execute(
                    f'SELECT id, club_name, file_name, file_path, description, status, reject_reason, upload_time, group_id, category, file_type FROM club_uploads WHERE club_name=? AND status="pending" AND {source_filter}',
                    (club_name,)).fetchall()
            else:
                rows = conn.execute(
                    f'SELECT id, club_name, file_name, file_path, description, status, reject_reason, upload_time, group_id, category, file_type FROM club_uploads WHERE status="pending" AND {source_filter}').fetchall()
            if not rows:
                return {'items': [], 'summary': '暂无待审批材料'}
            groups = {}
            for r in rows:
                gid = r['group_id'] or str(r['id'])
                if gid not in groups:
                    groups[gid] = {'group_id': gid, 'club_name': r['club_name'], 'upload_time': r['upload_time'],
                                   'description': r['description'], 'category': r['category'], 'files': [],
                                   'file_types': set(), 'file_names': [], 'summary_desc': '', 'file_contents': []}
                desc = r['description'] or ''
                if desc.startswith('[总结]') and not groups[gid]['summary_desc']:
                    groups[gid]['summary_desc'] = desc
                fp = r['file_path'] if 'file_path' in r.keys() else ''
                groups[gid]['files'].append({'id': r['id'], 'file_name': r['file_name'], 'file_type': r['file_type'], 'file_path': fp})
                ft = (r['file_name'] or '').rsplit('.', 1)[-1].lower() if '.' in (r['file_name'] or '') else ''
                if ft:
                    groups[gid]['file_types'].add(ft)
                if r['file_name']:
                    groups[gid]['file_names'].append(r['file_name'])
                if fp:
                    fc = self._extract_file_content(fp, ft)
                    if fc:
                        groups[gid]['file_contents'].append(fc)
            club_stats = {}
            for gid, g in groups.items():
                cn = g['club_name']
                if cn not in club_stats:
                    star = conn.execute('SELECT star_rating FROM club_profiles WHERE club_name=?', (cn,)).fetchone()
                    club_stats[cn] = {
                        'star_rating': star['star_rating'] if star else 0,
                    }
            results = []
            for gid, g in groups.items():
                cs = club_stats.get(g['club_name'], {})
                desc = g['description'] or ''
                summary_desc = g.get('summary_desc', '') or ''
                body_text = ''
                title_text = ''
                all_desc_texts = []
                for f in g['files']:
                    fd = f.get('file_name', '') or ''
                    if fd:
                        name_without_ext = fd.rsplit('.', 1)[0] if '.' in fd else fd
                        all_desc_texts.append(name_without_ext)
                if summary_desc:
                    parts = summary_desc.replace('[总结]', '').split('|||')
                    title_text = parts[0] if parts else ''
                    body_text = parts[1] if len(parts) > 1 else ''
                elif desc.startswith('[总结]'):
                    parts = desc.replace('[总结]', '').split('|||')
                    title_text = parts[0] if parts else ''
                    body_text = parts[1] if len(parts) > 1 else ''
                else:
                    for f in g['files']:
                        fd = f.get('file_name', '') or ''
                        if fd:
                            title_text = fd.rsplit('.', 1)[0] if '.' in fd else fd
                            break
                    if desc:
                        clean_desc = desc
                        for prefix in ['[照片]', '[PDF]', '[Word]', '[文件]']:
                            if clean_desc.startswith(prefix):
                                clean_desc = clean_desc[len(prefix):]
                                break
                        if clean_desc:
                            title_text = clean_desc
                file_names = g['file_names']
                file_types = g['file_types']
                category = g['category']
                file_contents = g.get('file_contents', [])
                import re as _re_body
                body_clean_check = _re_body.sub(r'<[^>]+>', '', body_text).strip() if body_text else ''
                if len(body_clean_check) < 5:
                    for f in g['files']:
                        fd = f.get('file_name', '') or ''
                        if fd:
                            name_without_ext = fd.rsplit('.', 1)[0] if '.' in fd else fd
                            body_text = body_text + ' ' + name_without_ext if body_text else name_without_ext
                relevance_score, relevance_reasons = self._simulate_relevance(body_text, title_text, file_names, file_types, category, file_contents)
                star_rating = cs.get('star_rating', 0) or 0
                star_score = (star_rating / 5) * 100
                combined_score = round(relevance_score * 0.7 + star_score * 0.3, 1)
                if combined_score >= 60:
                    recommendation = 'approve'
                else:
                    recommendation = 'reject'
                confidence = min(95, max(40, int(abs(combined_score - 60) * 2 + 30)))
                approval_reason = self._generate_approval_reason(
                    relevance_score, star_score, combined_score, star_rating, recommendation, relevance_reasons)
                risk_flags = []
                if relevance_score < 50:
                    risk_flags.append(f'正文与材料相关度仅{relevance_score}分，内容不匹配')
                if star_rating <= 1 and star_rating > 0:
                    risk_flags.append(f'仅{star_rating}星级社团')
                if not body_text.strip() and not title_text.strip():
                    risk_flags.append('缺少正文内容')
                if not file_names:
                    risk_flags.append('无上传材料')
                auto_approvable = combined_score >= 60 and relevance_score >= 70 and len(risk_flags) == 0
                results.append({
                    'group_id': gid,
                    'club_name': g['club_name'],
                    'upload_time': g['upload_time'],
                    'description': g['description'],
                    'file_count': len(g['files']),
                    'file_types': list(file_types),
                    'relevance_score': relevance_score,
                    'star_rating': star_rating,
                    'star_score': round(star_score, 1),
                    'combined_score': combined_score,
                    'recommendation': recommendation,
                    'confidence': confidence,
                    'reasons': relevance_reasons,
                    'risk_flags': risk_flags,
                    'approval_reason': approval_reason,
                    'club_stats': cs,
                    'auto_approvable': auto_approvable,
                    'suggested_reject_reason': '；'.join(risk_flags) if recommendation == 'reject' else '',
                })
            results.sort(key=lambda x: x['combined_score'], reverse=True)
            total = len(results)
            auto_count = sum(1 for r in results if r['auto_approvable'])
            approve_count = sum(1 for r in results if r['recommendation'] == 'approve')
            reject_count = sum(1 for r in results if r['recommendation'] == 'reject')
            summary = f'共{total}组待审批材料：建议通过{approve_count}组（其中{auto_count}组可自动通过），建议驳回{reject_count}组'
        finally:
            conn.close()
        return {'items': results, 'summary': summary}

    def analyze_workload(self, club_name=None, sub_id=None):
        conn = self._get_conn()
        try:
            if sub_id:
                rows = conn.execute(
                    'SELECT ws.id, ws.student_user_id, ws.student_name, ws.item_name, ws.score, ws.status, ws.club_name, ws.created_at FROM workload_submissions ws WHERE ws.id=? AND ws.status="pending"',
                    (sub_id,)).fetchall()
            elif club_name:
                rows = conn.execute(
                    'SELECT ws.id, ws.student_user_id, ws.student_name, ws.item_name, ws.score, ws.status, ws.club_name, ws.created_at FROM workload_submissions ws WHERE ws.club_name=? AND ws.status="pending"',
                    (club_name,)).fetchall()
            else:
                rows = conn.execute(
                    'SELECT ws.id, ws.student_user_id, ws.student_name, ws.item_name, ws.score, ws.status, ws.club_name, ws.created_at FROM workload_submissions ws WHERE ws.status="pending"').fetchall()
            if not rows:
                return {'items': [], 'summary': '暂无待审核工作量'}
            club_workload_stats = {}
            for r in rows:
                cn = r['club_name']
                if cn not in club_workload_stats:
                    avg_row = conn.execute(
                        'SELECT AVG(CAST(score AS REAL)) as avg_score, MAX(CAST(score AS REAL)) as max_score, COUNT(*) as total FROM workload_submissions WHERE club_name=? AND status="approved"',
                        (cn,)).fetchone()
                    club_workload_stats[cn] = {
                        'avg_score': round(avg_row['avg_score'], 1) if avg_row and avg_row['avg_score'] else 0,
                        'max_score': avg_row['max_score'] if avg_row else 0,
                        'total_approved': avg_row['total'] if avg_row else 0
                    }
            results = []
            for r in rows:
                cn = r['club_name']
                ws = club_workload_stats.get(cn, {})
                score_val = 50
                reasons = []
                risk_flags = []
                item_score = r['score'] or 0
                avg_score = ws.get('avg_score', 0)
                max_score = ws.get('max_score', 0)
                student_id = r['student_user_id']
                item_name = r['item_name'] or ''
                checkin_match = False
                if student_id and item_name:
                    matched = conn.execute(
                        'SELECT cs.id FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.student_id=? AND cs.club_name=? AND cs.activity_name LIKE ? AND cs.status="closed" LIMIT 1',
                        (student_id, cn, f'%{item_name[:6]}%')).fetchone()
                    if matched:
                        checkin_match = True
                if checkin_match:
                    score_val += 25
                    reasons.append('与已签到活动匹配')
                else:
                    if item_name:
                        score_val -= 5
                        risk_flags.append('未找到匹配的签到记录')
                if avg_score > 0:
                    if item_score <= avg_score * 1.5:
                        score_val += 15
                        reasons.append(f'分数在合理范围内（社团均值{avg_score}）')
                    elif item_score <= avg_score * 2:
                        score_val -= 5
                        risk_flags.append(f'分数偏高（社团均值{avg_score}，本次{item_score}）')
                    else:
                        score_val -= 20
                        risk_flags.append(f'分数异常偏高（社团均值{avg_score}，本次{item_score}）')
                if item_score <= 0:
                    score_val -= 30
                    risk_flags.append('工作量分数为0或负数')
                if student_id:
                    dup_count = conn.execute(
                        'SELECT COUNT(*) as c FROM workload_submissions WHERE student_user_id=? AND club_name=? AND item_name=? AND status="pending" AND id!=?',
                        (student_id, cn, item_name, r['id'])).fetchone()
                    if dup_count and dup_count['c'] > 0:
                        score_val -= 20
                        risk_flags.append(f'存在{dup_count["c"]}条相同待审核记录，疑似重复提交')
                if item_name:
                    score_val += 5
                    reasons.append(f'包含项目名称：{item_name[:20]}')
                else:
                    score_val -= 15
                    risk_flags.append('缺少项目名称')
                score_val = max(0, min(100, score_val))
                if score_val >= 60:
                    recommendation = 'approve'
                else:
                    recommendation = 'reject'
                confidence = min(95, max(30, score_val if recommendation == 'approve' else (100 - score_val) if recommendation == 'reject' else abs(50 - score_val) + 30))
                results.append({
                    'id': r['id'],
                    'club_name': cn,
                    'student_name': r['student_name'],
                    'item_name': r['item_name'],
                    'score': item_score,
                    'created_at': r['created_at'],
                    'recommendation': recommendation,
                    'confidence': confidence,
                    'agent_score': score_val,
                    'reasons': reasons,
                    'risk_flags': risk_flags,
                    'checkin_matched': checkin_match,
                    'auto_approvable': score_val >= 80 and checkin_match and len(risk_flags) == 0,
                    'suggested_reject_reason': '；'.join(risk_flags) if recommendation == 'reject' else '',
                    'suggested_review_note': f"AI建议：{'；'.join(reasons)}" if reasons else ''
                })
            results.sort(key=lambda x: x['agent_score'], reverse=True)
            total = len(results)
            auto_count = sum(1 for r in results if r['auto_approvable'])
            approve_count = sum(1 for r in results if r['recommendation'] == 'approve')
            review_count = sum(1 for r in results if r['recommendation'] == 'manual_review')
            reject_count = sum(1 for r in results if r['recommendation'] == 'reject')
            summary = f'共{total}条待审核工作量：建议通过{approve_count}条（其中{auto_count}条可自动通过），需人工复核{review_count}条，建议驳回{reject_count}条'
        finally:
            conn.close()
        return {'items': results, 'summary': summary}

    def analyze_registration(self, club_name=None, reg_id=None):
        conn = self._get_conn()
        try:
            if reg_id:
                rows = conn.execute(
                    'SELECT r.id, r.club_name, r.status, r.department, r.student_name, r.student_class, r.student_id_num, r.specialty, r.user_id, r.created_at FROM club_registrations r WHERE r.id=? AND r.status="pending"',
                    (reg_id,)).fetchall()
            elif club_name:
                rows = conn.execute(
                    'SELECT r.id, r.club_name, r.status, r.department, r.student_name, r.student_class, r.student_id_num, r.specialty, r.user_id, r.created_at FROM club_registrations r WHERE r.club_name=? AND r.status="pending"',
                    (club_name,)).fetchall()
            else:
                rows = conn.execute(
                    'SELECT r.id, r.club_name, r.status, r.department, r.student_name, r.student_class, r.student_id_num, r.specialty, r.user_id, r.created_at FROM club_registrations r WHERE r.status="pending"').fetchall()
            if not rows:
                return {'items': [], 'summary': '暂无待审批报名'}
            results = []
            for r in rows:
                score_val = 50
                reasons = []
                risk_flags = []
                suggested_dept = ''
                uid = r['user_id']
                cn = r['club_name']
                if uid and uid != 0:
                    member_count = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE user_id=?', (uid,)).fetchone()
                    if member_count and member_count['c'] >= 2:
                        score_val -= 40
                        risk_flags.append(f'已加入{member_count["c"]}个社团，达到上限')
                    elif member_count and member_count['c'] == 1:
                        score_val += 5
                        reasons.append('已加入1个社团，可再加1个')
                    else:
                        score_val += 10
                        reasons.append('尚未加入任何社团')
                    existing = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=? AND user_id!=0',
                                            (cn, uid)).fetchone()
                    if existing:
                        score_val -= 50
                        risk_flags.append('已是该社团成员')
                    dup_reg = conn.execute('SELECT COUNT(*) as c FROM club_registrations WHERE user_id=? AND club_name=? AND status="pending" AND id!=?',
                                           (uid, cn, r['id'])).fetchone()
                    if dup_reg and dup_reg['c'] > 0:
                        score_val -= 15
                        risk_flags.append('存在重复报名')
                depts = conn.execute('SELECT dept_name, description FROM club_departments WHERE club_name=?', (cn,)).fetchall()
                if depts:
                    specialty = r['specialty'] or ''
                    best_dept = None
                    best_match = 0
                    for d in depts:
                        dept_desc = (d['description'] or '').lower()
                        dept_name = (d['dept_name'] or '').lower()
                        match_score = 0
                        if specialty:
                            spec_lower = specialty.lower()
                            for kw in spec_lower:
                                if kw in dept_desc or kw in dept_name:
                                    match_score += 1
                        if specialty and any(kw in dept_desc or kw in dept_name for kw in specialty):
                            match_score += 5
                        if match_score > best_match:
                            best_match = match_score
                            best_dept = d['dept_name']
                    if best_dept:
                        suggested_dept = best_dept
                        score_val += 8
                        reasons.append(f'根据专业推荐部门：{best_dept}')
                if r['student_class']:
                    score_val += 3
                    reasons.append(f'班级：{r["student_class"]}')
                if r['specialty']:
                    score_val += 3
                    reasons.append(f'专业：{r["specialty"]}')
                club_member_count = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE club_name=?', (cn,)).fetchone()
                if club_member_count and club_member_count['c'] < 30:
                    score_val += 5
                    reasons.append(f'社团当前{club_member_count["c"]}人，名额充足')
                score_val = max(0, min(100, score_val))
                if score_val >= 60:
                    recommendation = 'approve'
                else:
                    recommendation = 'reject'
                confidence = min(95, max(30, score_val if recommendation == 'approve' else (100 - score_val) if recommendation == 'reject' else abs(50 - score_val) + 30))
                results.append({
                    'id': r['id'],
                    'club_name': cn,
                    'student_name': r['student_name'],
                    'student_class': r['student_class'],
                    'specialty': r['specialty'],
                    'created_at': r['created_at'],
                    'recommendation': recommendation,
                    'confidence': confidence,
                    'agent_score': score_val,
                    'reasons': reasons,
                    'risk_flags': risk_flags,
                    'suggested_department': suggested_dept,
                    'auto_approvable': score_val >= 80 and len(risk_flags) == 0,
                    'suggested_reject_reason': '；'.join(risk_flags) if recommendation == 'reject' else ''
                })
            results.sort(key=lambda x: x['agent_score'], reverse=True)
            total = len(results)
            auto_count = sum(1 for r in results if r['auto_approvable'])
            approve_count = sum(1 for r in results if r['recommendation'] == 'approve')
            review_count = sum(1 for r in results if r['recommendation'] == 'manual_review')
            reject_count = sum(1 for r in results if r['recommendation'] == 'reject')
            summary = f'共{total}条待审批报名：建议通过{approve_count}条（其中{auto_count}条可自动通过），需人工复核{review_count}条，建议驳回{reject_count}条'
        finally:
            conn.close()
        return {'items': results, 'summary': summary}

    def analyze_finance_permission(self, club_name=None, pid=None):
        conn = self._get_conn()
        try:
            if pid:
                rows = conn.execute('SELECT * FROM finance_permissions WHERE id=? AND status="pending"', (pid,)).fetchall()
            elif club_name:
                rows = conn.execute('SELECT * FROM finance_permissions WHERE club_name=? AND status="pending"', (club_name,)).fetchall()
            else:
                rows = conn.execute('SELECT * FROM finance_permissions WHERE status="pending"').fetchall()
            if not rows:
                return {'items': [], 'summary': '暂无待审批财务权限申请'}
            results = []
            for r in rows:
                score_val = 50
                reasons = []
                risk_flags = []
                uid = r['user_id']
                cn = r['club_name']
                is_member = conn.execute('SELECT id FROM club_members WHERE user_id=? AND club_name=?', (uid, cn)).fetchone()
                if is_member:
                    score_val += 20
                    reasons.append('是该社团成员')
                else:
                    score_val -= 15
                    risk_flags.append('非该社团成员')
                is_leader = conn.execute("SELECT id FROM users WHERE id=? AND club_name=? AND role='user'", (uid, cn)).fetchone()
                if is_leader:
                    score_val += 25
                    reasons.append('是社团负责人')
                other_finance = conn.execute('SELECT COUNT(*) as c FROM finance_managers WHERE user_id=? AND club_name!=?', (uid, cn)).fetchone()
                if other_finance and other_finance['c'] > 0:
                    score_val -= 15
                    risk_flags.append(f'已是{other_finance["c"]}个其他社团的财务管理员')
                reason_text = r['reason'] or ''
                if len(reason_text) > 10:
                    score_val += 10
                    reasons.append('申请理由详细')
                elif len(reason_text) > 0:
                    score_val += 3
                else:
                    score_val -= 10
                    risk_flags.append('未填写申请理由')
                score_val = max(0, min(100, score_val))
                if score_val >= 60:
                    recommendation = 'approve'
                else:
                    recommendation = 'reject'
                confidence = min(95, max(30, score_val if recommendation == 'approve' else (100 - score_val) if recommendation == 'reject' else abs(50 - score_val) + 30))
                results.append({
                    'id': r['id'],
                    'club_name': cn,
                    'username': r['username'],
                    'real_name': r['real_name'],
                    'reason': reason_text,
                    'created_at': r['created_at'],
                    'recommendation': recommendation,
                    'confidence': confidence,
                    'agent_score': score_val,
                    'reasons': reasons,
                    'risk_flags': risk_flags,
                    'auto_approvable': score_val >= 85 and len(risk_flags) == 0,
                    'suggested_reject_reason': '；'.join(risk_flags) if recommendation == 'reject' else ''
                })
            results.sort(key=lambda x: x['agent_score'], reverse=True)
            total = len(results)
            auto_count = sum(1 for r in results if r['auto_approvable'])
            approve_count = sum(1 for r in results if r['recommendation'] == 'approve')
            review_count = sum(1 for r in results if r['recommendation'] == 'manual_review')
            reject_count = sum(1 for r in results if r['recommendation'] == 'reject')
            summary = f'共{total}条待审批财务权限：建议通过{approve_count}条（其中{auto_count}条可自动通过），需人工复核{review_count}条，建议驳回{reject_count}条'
        finally:
            conn.close()
        return {'items': results, 'summary': summary}

    def analyze_offcampus(self, club_name=None, rid=None):
        conn = self._get_conn()
        try:
            if rid:
                rows = conn.execute('SELECT * FROM offcampus_requests WHERE id=? AND status="pending"', (rid,)).fetchall()
            elif club_name:
                rows = conn.execute('SELECT * FROM offcampus_requests WHERE club_name=? AND status="pending"', (club_name,)).fetchall()
            else:
                rows = conn.execute('SELECT * FROM offcampus_requests WHERE status="pending"').fetchall()
            if not rows:
                return {'items': [], 'summary': '暂无待审批校外活动申请'}
            results = []
            for r in rows:
                score_val = 50
                reasons = []
                risk_flags = []
                cn = r['club_name']
                title = r['title'] or ''
                location = r['location'] or ''
                description = r['description'] or ''
                activity_date = r['activity_date'] or ''
                has_file = bool(r['file_path'])
                if title:
                    score_val += 10
                    reasons.append(f'活动名称：{title[:20]}')
                else:
                    score_val -= 15
                    risk_flags.append('缺少活动名称')
                if location:
                    score_val += 8
                    reasons.append(f'活动地点：{location[:20]}')
                else:
                    score_val -= 10
                    risk_flags.append('缺少活动地点')
                if description and len(description) > 10:
                    score_val += 10
                    reasons.append('活动描述详细')
                elif description:
                    score_val += 3
                else:
                    score_val -= 10
                    risk_flags.append('缺少活动描述')
                if activity_date:
                    score_val += 5
                    reasons.append(f'活动日期：{activity_date}')
                else:
                    score_val -= 8
                    risk_flags.append('缺少活动日期')
                if has_file:
                    score_val += 8
                    reasons.append('包含附件材料')
                else:
                    risk_flags.append('缺少附件材料')
                past_approved = conn.execute('SELECT COUNT(*) as c FROM offcampus_requests WHERE club_name=? AND status="approved"',
                                             (cn,)).fetchone()
                past_rejected = conn.execute('SELECT COUNT(*) as c FROM offcampus_requests WHERE club_name=? AND status="rejected"',
                                             (cn,)).fetchone()
                if past_approved and past_approved['c'] > 0:
                    score_val += 5
                    reasons.append(f'历史已通过{past_approved["c"]}次校外活动')
                if past_rejected and past_rejected['c'] > 2:
                    score_val -= 5
                    risk_flags.append(f'历史被驳回{past_rejected["c"]}次')
                score_val = max(0, min(100, score_val))
                if score_val >= 60:
                    recommendation = 'approve'
                else:
                    recommendation = 'reject'
                confidence = min(95, max(30, score_val if recommendation == 'approve' else (100 - score_val) if recommendation == 'reject' else abs(50 - score_val) + 30))
                results.append({
                    'id': r['id'],
                    'club_name': cn,
                    'title': title,
                    'location': location,
                    'activity_date': activity_date,
                    'created_at': r['created_at'],
                    'recommendation': recommendation,
                    'confidence': confidence,
                    'agent_score': score_val,
                    'reasons': reasons,
                    'risk_flags': risk_flags,
                    'auto_approvable': False,
                    'suggested_reject_reason': '；'.join(risk_flags) if recommendation == 'reject' else ''
                })
            results.sort(key=lambda x: x['agent_score'], reverse=True)
            total = len(results)
            approve_count = sum(1 for r in results if r['recommendation'] == 'approve')
            review_count = sum(1 for r in results if r['recommendation'] == 'manual_review')
            reject_count = sum(1 for r in results if r['recommendation'] == 'reject')
            summary = f'共{total}条待审批校外活动：建议通过{approve_count}条，需人工复核{review_count}条，建议驳回{reject_count}条'
        finally:
            conn.close()
        return {'items': results, 'summary': summary}

    def ai_deep_analyze(self, approval_type, item_data):
        _, api_key = get_llm_config()
        if not api_key:
            return None
        type_prompts = {
            'material': '你是一个社团材料智能审批助手。评选标准：正文与材料相关度（权重70%）+ 社团星级（权重30%）。综合得分≥60分建议通过，<60分建议驳回。请根据以下材料评分数据，给出更深入的审批建议，重点分析正文内容与上传材料的语义相关度。',
            'workload': '你是一个社团工作量审核助手。请根据以下工作量提交信息，判断分数是否合理、是否与活动匹配、是否存在异常。',
            'registration': '你是一个社团报名审批助手。请根据以下报名信息，判断学生是否适合加入该社团，推荐合适的部门。',
            'finance_permission': '你是一个财务权限审批助手。请根据以下申请信息，判断是否应该授予财务管理权限。',
            'offcampus': '你是一个校外活动审批助手。请根据以下校外活动申请信息，判断活动是否安全合规、材料是否完整。'
        }
        system_prompt = type_prompts.get(approval_type, '你是一个审批助手。')
        system_prompt += '\n\n请以JSON格式返回分析结果，包含以下字段：\n- recommendation: "approve"或"reject"或"manual_review"\n- confidence: 0-100的置信度\n- analysis: 详细分析文字（100字以内）\n- suggestions: 具体建议列表\n\n只返回JSON，不要其他内容。'
        user_msg = f'审批类型：{approval_type}\n待审批项目数据：\n{json.dumps(item_data, ensure_ascii=False, default=str)}'
        messages = [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_msg}
        ]
        result = call_llm_api(messages, max_tokens=500)
        if result:
            try:
                cleaned = result.strip()
                if cleaned.startswith('```'):
                    cleaned = re.sub(r'^```\w*\n?', '', cleaned)
                    cleaned = re.sub(r'\n?```$', '', cleaned)
                return json.loads(cleaned)
            except json.JSONDecodeError:
                return {'analysis': result, 'recommendation': 'manual_review', 'confidence': 50, 'suggestions': []}
        return None

    def batch_analyze(self, approval_type, **kwargs):
        analyzers = {
            'material': self.analyze_material,
            'workload': self.analyze_workload,
            'registration': self.analyze_registration,
            'finance_permission': self.analyze_finance_permission,
            'offcampus': self.analyze_offcampus
        }
        analyzer = analyzers.get(approval_type)
        if not analyzer:
            return {'error': f'不支持的审批类型：{approval_type}'}
        return analyzer(**kwargs)


approval_agent = ApprovalAgent(db)


class DocIndexAgent:
    TAG_KEYWORDS = {
        '活动照片': ['照片', '图片', '截图', '合影', '现场', '拍照'],
        '活动总结': ['总结', '小结', '回顾', '成果', '汇报'],
        '活动方案': ['方案', '策划', '计划', '安排', '流程'],
        '签到表': ['签到', '考勤', '点名', '出勤'],
        '财务报表': ['财务', '收支', '报销', '经费', '预算', '账单'],
        '通知公告': ['通知', '公告', '公示', '声明', '通报'],
        '报名表': ['报名', '登记', '注册', '申请表'],
        '宣传海报': ['海报', '宣传', '推广', '文案', '招新'],
        '规章制度': ['制度', '规章', '章程', '规范', '守则'],
        '会议记录': ['会议', '纪要', '讨论', '决议'],
        '比赛材料': ['比赛', '竞赛', '评选', '参赛'],
        '培训材料': ['培训', '讲座', '教学', '课程'],
        '校外活动': ['校外', '外出', '参观', '交流', '联谊'],
        '指导记录': ['指导', '辅导', '点评', '建议'],
        '学期规划': ['规划', '计划', '学期', '年度'],
        '成员名册': ['名册', '名单', '成员', '干部'],
        '评分赋分': ['评分', '赋分', '学分', '考核'],
        '反馈投诉': ['反馈', '投诉', '建议', '问题'],
    }

    CATEGORY_MAP = {
        '活动照片': '活动材料',
        '活动总结': '活动材料',
        '活动方案': '活动材料',
        '签到表': '活动材料',
        '比赛材料': '活动材料',
        '培训材料': '活动材料',
        '校外活动': '活动材料',
        '财务报表': '财务材料',
        '通知公告': '通知材料',
        '报名表': '招募材料',
        '宣传海报': '宣传材料',
        '规章制度': '管理材料',
        '会议记录': '管理材料',
        '指导记录': '指导材料',
        '学期规划': '规划材料',
        '成员名册': '管理材料',
        '评分赋分': '考核材料',
        '反馈投诉': '反馈材料',
    }

    def __init__(self, db_instance):
        self.db = db_instance

    def _get_conn(self):
        return self.db.get_conn()

    def generate_tags(self, file_name, description, club_name, category, source_table):
        tags = set()
        text = f'{file_name} {description} {club_name} {category}'.lower()
        for tag, keywords in self.TAG_KEYWORDS.items():
            for kw in keywords:
                if kw in text:
                    tags.add(tag)
                    break
        ext = (file_name or '').rsplit('.', 1)[-1].lower() if '.' in (file_name or '') else ''
        ext_tags = {'jpg': '图片', 'jpeg': '图片', 'png': '图片', 'gif': '图片', 'webp': '图片', 'bmp': '图片',
                     'pdf': 'PDF文档', 'doc': 'Word文档', 'docx': 'Word文档',
                     'xls': 'Excel表格', 'xlsx': 'Excel表格', 'csv': 'Excel表格',
                     'ppt': 'PPT演示', 'pptx': 'PPT演示',
                     'txt': '文本文件', 'md': '文本文件',
                     'zip': '压缩包', 'rar': '压缩包', '7z': '压缩包',
                     'mp4': '视频', 'mp3': '音频', 'wav': '音频'}
        if ext in ext_tags:
            tags.add(ext_tags[ext])
        desc = description or ''
        if desc.startswith('[总结]') or '|||' in desc:
            tags.add('活动总结')
        if desc.startswith('[照片]'):
            tags.add('活动照片')
        if desc.startswith('[方案]'):
            tags.add('活动方案')
        if source_table == 'offcampus_requests':
            tags.add('校外活动')
        elif source_table == 'finance_records':
            tags.add('财务报表')
        elif source_table == 'notices':
            tags.add('通知公告')
        elif source_table == 'feedbacks':
            tags.add('反馈投诉')
        if not tags:
            tags.add('其他')
        return ','.join(sorted(tags))

    def generate_tags_with_llm(self, file_name, description, club_name, category):
        import time
        provider, api_key = get_llm_config()
        if not api_key:
            return set()
        try:
            import re
            prompt = f"""请为以下社团文档生成5-12个同义词和相关标签，用于语义搜索。
文档信息：
- 文件名: {file_name or ''}
- 描述: {description or ''}
- 社团名称: {club_name or ''}
- 类别: {category or ''}

请考虑以下方面：
1. 奖项相关：奖项、获奖、荣誉、表彰、一等奖、省级、校级、国家级等
2. 活动类型：比赛、培训、会议、展示、演出、志愿等
3. 社团名称拆解：将社团名称中的关键词提取出来作为标签
4. 文件类型和内容关键词

只返回标签，用逗号分隔，不要其他内容。"""
            messages = [
                {'role': 'system', 'content': '你是一个标签生成助手，只为文档生成搜索用的同义词标签。'},
                {'role': 'user', 'content': prompt}
            ]
            time.sleep(0.3)
            resp = call_llm_api(messages, max_tokens=150)
            if not resp:
                return set()
            tags = set()
            for part in re.split(r'[,，、\n;；\s]+', resp):
                part = part.strip().strip('。.')
                if part and len(part) <= 20 and part not in tags:
                    tags.add(part)
                if len(tags) >= 12:
                    break
            return tags
        except Exception:
            return set()

    def rebuild_index(self):
        conn = self._get_conn()
        try:
            conn.execute('DELETE FROM doc_index')
            conn.execute('DELETE FROM doc_embeddings')
            sources = [
                ('club_uploads', 'SELECT id, club_name, file_name, file_path, description, category, status, upload_time FROM club_uploads WHERE file_path!="" AND file_path IS NOT NULL'),
                ('offcampus_requests', 'SELECT id, club_name, COALESCE(file_name,"") as file_name, COALESCE(file_path,"") as file_path, COALESCE(title,"")||" "||COALESCE(description,"") as description, "校外活动" as category, status, created_at FROM offcampus_requests WHERE file_path!="" AND file_path IS NOT NULL'),
                ('finance_records', 'SELECT id, club_name, COALESCE(attachment_name,"") as file_name, COALESCE(attachment_path,"") as file_path, COALESCE(description,"") as description, "财务" as category, "" as status, created_at FROM finance_records WHERE attachment_path!="" AND attachment_path IS NOT NULL'),
                ('notices', 'SELECT id, "" as club_name, COALESCE(attachment_name,"") as file_name, COALESCE(attachment_path,"") as file_path, COALESCE(title,"")||" "||COALESCE(content,"") as description, "公告" as category, "" as status, created_at FROM notices WHERE attachment_path!="" AND attachment_path IS NOT NULL'),
                ('feedbacks', 'SELECT id, club_name, COALESCE(file_name,"") as file_name, COALESCE(file_path,"") as file_path, COALESCE(title,"")||" "||COALESCE(body,"") as description, "反馈" as category, status, created_at FROM feedbacks WHERE file_path!="" AND file_path IS NOT NULL'),
                ('checkin_sessions', 'SELECT id, club_name, "" as file_name, COALESCE(plan_path,"") as file_path, COALESCE(activity_name,"") as description, "签到" as category, "" as status, created_at FROM checkin_sessions WHERE plan_path!="" AND plan_path IS NOT NULL'),
                ('club_showcase', 'SELECT id, club_name, COALESCE(club_name,"")||"风采展示" as file_name, COALESCE(image_path,"") as file_path, COALESCE(description,"") as description, "社团风采" as category, "" as status, created_at FROM club_showcase WHERE image_path!="" AND image_path IS NOT NULL'),
                ('joint_activities', 'SELECT id, club_name, "" as file_name, "" as file_path, COALESCE(title,"")||" "||COALESCE(description,"") as description, "联合活动" as category, status, created_at FROM joint_activities'),
                ('club_notices', 'SELECT id, club_name, "" as file_name, "" as file_path, COALESCE(title,"")||" "||COALESCE(content,"") as description, "社团通知" as category, "" as status, created_at FROM club_notices'),
                ('workload_submissions', 'SELECT id, club_name, item_name as file_name, "" as file_path, item_name as description, "工作量" as category, status, created_at FROM workload_submissions'),
                ('guidance_records', 'SELECT id, club_name, "" as file_name, "" as file_path, guidance_content as description, "指导记录" as category, "" as status, created_at FROM guidance_records'),
                ('activity_records', 'SELECT id, club_name, "" as file_name, "" as file_path, activity_content as description, "活动记录" as category, "" as status, created_at FROM activity_records'),
                ('weekly_reports', 'SELECT id, club_name, "" as file_name, "" as file_path, content as description, "周报" as category, "" as status, created_at FROM weekly_reports'),
                ('recruitments', 'SELECT id, club_name, title as file_name, "" as file_path, description, "招募" as category, status, created_at FROM recruitments'),
                ('scoring_submissions', 'SELECT id, club_name, "" as file_name, "" as file_path, "赋分提交" as description, "赋分" as category, status, created_at FROM scoring_submissions'),
                ('excellent_clubs', 'SELECT id, club_name, "" as file_name, "" as file_path, "优秀社团评选" as description, "优秀社团" as category, "" as status, selected_at as created_at FROM excellent_clubs'),
                ('club_registrations', 'SELECT id, club_name, student_name as file_name, "" as file_path, COALESCE(student_name,"")||"报名" as description, "报名表" as category, status, created_at FROM club_registrations'),
            ]
            total = 0
            embedding_tasks = []
            for table, query in sources:
                rows = conn.execute(query).fetchall()
                for r in rows:
                    fname = r['file_name'] or ''
                    fpath = r['file_path'] or ''
                    desc = r['description'] or ''
                    cn = r['club_name'] or ''
                    cat = r['category'] or ''
                    tags = self.generate_tags(fname, desc, cn, cat, table)
                    # 只对有实际文件的记录调用 LLM 生成同义词标签，避免大量内容型数据拖慢索引
                    if fpath:
                        llm_tags = self.generate_tags_with_llm(fname, desc, cn, cat)
                    else:
                        llm_tags = set()
                    all_tags = set(tags.split(',')) | llm_tags
                    # 把原始 category 加入 tags，确保可按 category 搜索
                    if cat:
                        all_tags.add(cat)
                    tags = ','.join(sorted(all_tags))
                    # category 优先用原始值，只有为空时才从 tags 推断
                    auto_cat = cat
                    if not auto_cat:
                        for tag in all_tags:
                            if tag in self.CATEGORY_MAP:
                                auto_cat = self.CATEGORY_MAP[tag]
                                break
                    if not auto_cat:
                        auto_cat = '其他'
                    time_val = ''
                    for k in ('created_at', 'upload_time', 'date', 'record_date'):
                        try:
                            time_val = r[k] or ''
                            break
                        except (KeyError, IndexError):
                            pass
                    cur = conn.execute('''INSERT INTO doc_index (source_table, source_id, club_name, file_name, file_path, description, tags, category, status, source_type, created_at, indexed_at)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,datetime("now","localtime"))''',
                        (table, r['id'], cn, fname, fpath, (desc or '')[:500], tags, auto_cat or cat, r['status'] or '', table, time_val))
                    doc_id = cur.lastrowid
                    embedding_tasks.append((doc_id, table, fname + ' ' + desc[:500]))
                    total += 1
            conn.commit()
            embedded = 0
            for doc_id, doc_type, doc_text in embedding_tasks:
                if self.index_document_embedding(doc_id, doc_type, doc_text):
                    embedded += 1
            return {'success': True, 'indexed': total, 'embedded': embedded}
        finally:
            conn.close()

    def _do_search(self, query, club_name='', category='', limit=30):
        conn = self._get_conn()
        try:
            q = query.strip().lower()
            if not q and not club_name and not category:
                rows = conn.execute('SELECT * FROM doc_index ORDER BY created_at DESC LIMIT ?', (limit,)).fetchall()
            else:
                conditions = []
                params = []
                if q:
                    or_parts = []
                    for word in q.split():
                        or_parts.append('(tags LIKE ? OR file_name LIKE ? OR description LIKE ? OR club_name LIKE ? OR category LIKE ?)')
                        pw = f'%{word}%'
                        params.extend([pw, pw, pw, pw, pw])
                    conditions.append('(' + ' AND '.join(or_parts) + ')')
                if club_name:
                    conditions.append('club_name=?')
                    params.append(club_name)
                if category:
                    conditions.append('category=?')
                    params.append(category)
                where = ' AND '.join(conditions)
                params.append(limit)
                rows = conn.execute(f'SELECT * FROM doc_index WHERE {where} ORDER BY created_at DESC LIMIT ?', params).fetchall()
            results = []
            for r in rows:
                score = 0
                if q:
                    for word in q.split():
                        if word in (r['tags'] or '').lower():
                            score += 10
                        if word in (r['file_name'] or '').lower():
                            score += 5
                        if word in (r['description'] or '').lower():
                            score += 3
                        if word in (r['club_name'] or '').lower():
                            score += 2
                results.append({
                    'id': r['id'], 'source_table': r['source_table'], 'source_id': r['source_id'],
                    'club_name': r['club_name'], 'file_name': r['file_name'], 'file_path': r['file_path'],
                    'file_exists': bool(r['file_path'] and storage.exists(r['file_path'])),
                    'description': (r['description'] or '')[:200], 'tags': r['tags'], 'category': r['category'],
                    'status': r['status'], 'source_type': r['source_type'], 'created_at': r['created_at'],
                    'relevance': score
                })
            if q:
                results.sort(key=lambda x: x['relevance'], reverse=True)
            return results
        finally:
            conn.close()

    def _expand_query_with_llm(self, query):
        provider, api_key = get_llm_config()
        if not api_key:
            return []
        try:
            import re
            prompt = f"""请将以下搜索查询分解为3-8个用于数据库搜索的关键词和同义词。
查询：{query}

要求：
1. 提取核心概念词（如"奖项"而不是"社团获得的奖项"）
2. 生成同义词和相关词（如"获奖""荣誉""表彰""证书"）
3. 只返回关键词，用逗号分隔，不要其他内容

示例：
查询"社团获得的奖项" → 奖项,获奖,荣誉,表彰,证书,比赛,获奖材料
查询"省级奖项" → 省级,奖项,省级荣誉,省级表彰,省级获奖"""
            messages = [
                {'role': 'system', 'content': '你是一个搜索查询扩展助手，只返回关键词列表。'},
                {'role': 'user', 'content': prompt}
            ]
            resp = call_llm_api(messages, max_tokens=100)
            if not resp:
                return []
            terms = [t.strip() for t in re.split(r'[,，、\n;；\s]+', resp) if t.strip() and len(t.strip()) <= 20]
            original = query.strip()
            return [t for t in terms if t != original][:8]
        except Exception:
            return []

    def search(self, query, club_name='', category='', limit=30):
        results = self._do_search(query, club_name, category, limit)
        if query.strip() and len(results) < 5:
            expanded = self._expand_query_with_llm(query)
            if expanded:
                seen = {r['id'] for r in results}
                for term in expanded:
                    for r in self._do_search(term, club_name, category, limit):
                        if r['id'] not in seen:
                            results.append(r)
                            seen.add(r['id'])
                results.sort(key=lambda x: x.get('relevance', 0), reverse=True)
                results = results[:limit]
        return results

    def get_stats(self):
        conn = self._get_conn()
        try:
            total = conn.execute('SELECT COUNT(*) as c FROM doc_index').fetchone()['c']
            by_cat = conn.execute('SELECT category, COUNT(*) as c FROM doc_index GROUP BY category ORDER BY c DESC').fetchall()
            by_club = conn.execute('SELECT club_name, COUNT(*) as c FROM doc_index WHERE club_name!="" GROUP BY club_name ORDER BY c DESC LIMIT 10').fetchall()
            by_tag = conn.execute('SELECT tags, COUNT(*) as c FROM doc_index GROUP BY tags ORDER BY c DESC LIMIT 15').fetchall()
            all_tags = set()
            tag_counts = {}
            for r in conn.execute('SELECT tags FROM doc_index').fetchall():
                for t in (r['tags'] or '').split(','):
                    t = t.strip()
                    if t:
                        all_tags.add(t)
                        tag_counts[t] = tag_counts.get(t, 0) + 1
            sorted_tags = sorted(tag_counts.items(), key=lambda x: -x[1])
            return {
                'total': total,
                'by_category': [dict(r) for r in by_cat],
                'by_club': [dict(r) for r in by_club],
                'top_tags': sorted_tags[:20],
                'all_tags': sorted(all_tags)
            }
        finally:
            conn.close()

    def ai_smart_search(self, query):
        results = self.search(query, limit=20)
        if not QWEN_API_KEY or not results:
            return {'results': results, 'ai_summary': ''}
        top_items = results[:5]
        items_desc = '\n'.join([f"- {r['club_name']} | {r['file_name']} | {r['tags']} | {(r['description'] or '')[:80]}" for r in top_items])
        messages = [
            {'role': 'system', 'content': '你是一个文件检索助手。根据搜索结果，用1-2句话总结找到了什么，并指出最相关的文件。只返回总结文字。'},
            {'role': 'user', 'content': f'用户搜索：{query}\n搜索结果：\n{items_desc}'}
        ]
        ai_summary = call_llm_api(messages, max_tokens=200)
        return {'results': results, 'ai_summary': ai_summary or ''}

    def semantic_search(self, query, top_k=10):
        query_embedding = get_embedding(query)
        if not query_embedding:
            return self.search(query)
        import pickle
        conn = self.db.get_conn()
        try:
            rows = conn.execute('SELECT doc_id, doc_type, doc_text, embedding FROM doc_embeddings WHERE embedding IS NOT NULL LIMIT 500').fetchall()
        finally:
            conn.close()
        if not rows:
            return self.search(query)
        scored = []
        for row in rows:
            try:
                doc_embedding = pickle.loads(row['embedding']) if row['embedding'] else None
                if doc_embedding:
                    sim = cosine_similarity(query_embedding, doc_embedding)
                    scored.append({
                        'doc_id': row['doc_id'],
                        'doc_type': row['doc_type'],
                        'text_preview': row['doc_text'][:200],
                        'similarity': round(sim, 4)
                    })
            except:
                continue
        scored.sort(key=lambda x: x['similarity'], reverse=True)
        return {
            'total': len(scored),
            'items': scored[:top_k],
            'query': query,
            'search_type': 'semantic'
        }

    def index_document_embedding(self, doc_id, doc_type, doc_text):
        embedding = get_embedding(doc_text[:2000])
        if not embedding:
            return False
        import pickle
        embedding_blob = pickle.dumps(embedding)
        conn = self.db.get_conn()
        try:
            existing = conn.execute('SELECT id FROM doc_embeddings WHERE doc_id=? AND doc_type=?', (doc_id, doc_type)).fetchone()
            if existing:
                conn.execute('UPDATE doc_embeddings SET doc_text=?, embedding=? WHERE doc_id=? AND doc_type=?',
                            (doc_text[:2000], embedding_blob, doc_id, doc_type))
            else:
                conn.execute('INSERT INTO doc_embeddings (doc_id, doc_type, doc_text, embedding) VALUES (?, ?, ?, ?)',
                            (doc_id, doc_type, doc_text[:2000], embedding_blob))
            conn.commit()
        finally:
            conn.close()
        return True

    def add_document_to_index(self, source_table, source_id):
        """增量索引：当新文件上传时自动添加到索引"""
        queries = {
            'club_uploads': 'SELECT id, club_name, file_name, file_path, description, category, status, upload_time FROM club_uploads WHERE id=?',
            'offcampus_requests': 'SELECT id, club_name, COALESCE(file_name,"") as file_name, COALESCE(file_path,"") as file_path, COALESCE(title,"")||" "||COALESCE(description,"") as description, "校外活动" as category, status, created_at FROM offcampus_requests WHERE id=?',
            'finance_records': 'SELECT id, club_name, COALESCE(attachment_name,"") as file_name, COALESCE(attachment_path,"") as file_path, COALESCE(description,"") as description, "财务" as category, "" as status, created_at FROM finance_records WHERE id=?',
            'notices': 'SELECT id, "" as club_name, COALESCE(attachment_name,"") as file_name, COALESCE(attachment_path,"") as file_path, COALESCE(title,"")||" "||COALESCE(content,"") as description, "公告" as category, "" as status, created_at FROM notices WHERE id=?',
            'feedbacks': 'SELECT id, club_name, COALESCE(file_name,"") as file_name, COALESCE(file_path,"") as file_path, COALESCE(title,"")||" "||COALESCE(body,"") as description, "反馈" as category, status, created_at FROM feedbacks WHERE id=?',
            'checkin_sessions': 'SELECT id, club_name, "" as file_name, COALESCE(plan_path,"") as file_path, COALESCE(activity_name,"") as description, "签到" as category, "" as status, created_at FROM checkin_sessions WHERE id=?',
            'club_showcase': 'SELECT id, club_name, COALESCE(club_name,"")||"风采展示" as file_name, COALESCE(image_path,"") as file_path, COALESCE(description,"") as description, "社团风采" as category, "" as status, created_at FROM club_showcase WHERE id=?',
            'joint_activities': 'SELECT id, club_name, "" as file_name, "" as file_path, COALESCE(title,"")||" "||COALESCE(description,"") as description, "联合活动" as category, status, created_at FROM joint_activities WHERE id=?',
            'club_notices': 'SELECT id, club_name, "" as file_name, "" as file_path, COALESCE(title,"")||" "||COALESCE(content,"") as description, "社团通知" as category, "" as status, created_at FROM club_notices WHERE id=?',
            'workload_submissions': 'SELECT id, club_name, item_name as file_name, "" as file_path, item_name as description, "工作量" as category, status, created_at FROM workload_submissions WHERE id=?',
            'guidance_records': 'SELECT id, club_name, "" as file_name, "" as file_path, guidance_content as description, "指导记录" as category, "" as status, created_at FROM guidance_records WHERE id=?',
            'activity_records': 'SELECT id, club_name, "" as file_name, "" as file_path, activity_content as description, "活动记录" as category, "" as status, created_at FROM activity_records WHERE id=?',
            'weekly_reports': 'SELECT id, club_name, "" as file_name, "" as file_path, content as description, "周报" as category, "" as status, created_at FROM weekly_reports WHERE id=?',
            'recruitments': 'SELECT id, club_name, title as file_name, "" as file_path, description, "招募" as category, status, created_at FROM recruitments WHERE id=?',
            'scoring_submissions': 'SELECT id, club_name, "" as file_name, "" as file_path, "赋分提交" as description, "赋分" as category, status, created_at FROM scoring_submissions WHERE id=?',
            'excellent_clubs': 'SELECT id, club_name, "" as file_name, "" as file_path, "优秀社团评选" as description, "优秀社团" as category, "" as status, selected_at as created_at FROM excellent_clubs WHERE id=?',
            'club_registrations': 'SELECT id, club_name, student_name as file_name, "" as file_path, COALESCE(student_name,"")||"报名" as description, "报名表" as category, status, created_at FROM club_registrations WHERE id=?',
        }
        query = queries.get(source_table)
        if not query:
            return {'error': f'不支持的数据源：{source_table}'}
        conn = self._get_conn()
        try:
            r = conn.execute(query, (source_id,)).fetchone()
            if not r:
                return {'error': '记录不存在'}
            fname = r['file_name'] or ''
            fpath = r['file_path'] or ''
            desc = r['description'] or ''
            cn = r['club_name'] or ''
            cat = r['category'] or ''
            tags = self.generate_tags(fname, desc, cn, cat, source_table)
            llm_tags = self.generate_tags_with_llm(fname, desc, cn, cat)
            all_tags = set(tags.split(',')) | llm_tags
            if cat:
                all_tags.add(cat)
            tags = ','.join(sorted(all_tags))
            auto_cat = cat
            if not auto_cat:
                for tag in all_tags:
                    if tag in self.CATEGORY_MAP:
                        auto_cat = self.CATEGORY_MAP[tag]
                        break
            if not auto_cat:
                auto_cat = '其他'
            time_val = ''
            for k in ('created_at', 'upload_time', 'date', 'record_date'):
                try:
                    time_val = r[k] or ''
                    break
                except (KeyError, IndexError):
                    pass
            cur = conn.execute('''INSERT INTO doc_index (source_table, source_id, club_name, file_name, file_path, description, tags, category, status, source_type, created_at, indexed_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,datetime("now","localtime"))''',
                (source_table, r['id'], cn, fname, fpath, (desc or '')[:500], tags, auto_cat or cat, r['status'] or '', source_table, time_val))
            doc_id = cur.lastrowid
            conn.commit()
            self.index_document_embedding(doc_id, source_table, f'{fname} {desc[:500]}')
            return {'success': True, 'id': doc_id}
        finally:
            conn.close()


class NotificationAgent:
    RULES = [
        {'id': 'material_overdue', 'name': '材料审批超时', 'desc': '材料待审批超过3天', 'severity': 'warning'},
        {'id': 'workload_overdue', 'name': '工作量审核超时', 'desc': '工作量待审核超过7天', 'severity': 'warning'},
        {'id': 'registration_overdue', 'name': '报名审批超时', 'desc': '报名待审批超过5天', 'severity': 'warning'},
        {'id': 'recruitment_expiring', 'name': '招募即将截止', 'desc': '招募活动3天内截止但报名不足', 'severity': 'info'},
        {'id': 'club_inactive', 'name': '社团活跃度骤降', 'desc': '社团近2周无任何活动', 'severity': 'warning'},
        {'id': 'scoring_stuck', 'name': '赋分流程停滞', 'desc': '赋分已提交但审核停滞超过5天', 'severity': 'warning'},
        {'id': 'offcampus_overdue', 'name': '校外活动审批超时', 'desc': '校外活动申请待审批超过3天', 'severity': 'warning'},
    ]

    def __init__(self, db_instance):
        self.db = db_instance

    def _get_conn(self):
        return self.db.get_conn()

    def check_all(self):
        alerts = []
        alerts.extend(self._check_material_overdue())
        alerts.extend(self._check_workload_overdue())
        alerts.extend(self._check_registration_overdue())
        alerts.extend(self._check_recruitment_expiring())
        alerts.extend(self._check_club_inactive())
        alerts.extend(self._check_scoring_stuck())
        alerts.extend(self._check_offcampus_overdue())
        alerts.sort(key=lambda x: {'high': 0, 'warning': 1, 'info': 2}.get(x.get('severity', 'info'), 2))
        return alerts

    def _check_material_overdue(self):
        conn = self._get_conn()
        try:
            rows = conn.execute('''SELECT group_id, club_name, MIN(upload_time) as first_upload
                FROM club_uploads WHERE status="pending" AND (source="upload" OR source="activity" OR source IS NULL) AND group_id!=""
                GROUP BY group_id HAVING julianday("now","localtime")-julianday(MIN(upload_time))>3''').fetchall()
            return [{'rule_id': 'material_overdue', 'type': '材料审批超时', 'severity': 'warning',
                     'message': f'{r["club_name"]}有材料待审批超过3天（提交于{r["first_upload"]}）',
                     'club_name': r['club_name'], 'group_id': r['group_id']} for r in rows]
        finally:
            conn.close()

    def _check_workload_overdue(self):
        conn = self._get_conn()
        try:
            rows = conn.execute('''SELECT club_name, COUNT(*) as cnt, MIN(created_at) as earliest
                FROM workload_submissions WHERE status="pending"
                GROUP BY club_name HAVING julianday("now","localtime")-julianday(MIN(created_at))>7''').fetchall()
            return [{'rule_id': 'workload_overdue', 'type': '工作量审核超时', 'severity': 'warning',
                     'message': f'{r["club_name"]}有{r["cnt"]}条工作量待审核超过7天',
                     'club_name': r['club_name']} for r in rows]
        finally:
            conn.close()

    def _check_registration_overdue(self):
        conn = self._get_conn()
        try:
            rows = conn.execute('''SELECT club_name, COUNT(*) as cnt, MIN(created_at) as earliest
                FROM club_registrations WHERE status="pending"
                GROUP BY club_name HAVING julianday("now","localtime")-julianday(MIN(created_at))>5''').fetchall()
            return [{'rule_id': 'registration_overdue', 'type': '报名审批超时', 'severity': 'warning',
                     'message': f'{r["club_name"]}有{r["cnt"]}条报名待审批超过5天',
                     'club_name': r['club_name']} for r in rows]
        finally:
            conn.close()

    def _check_recruitment_expiring(self):
        conn = self._get_conn()
        try:
            rows = conn.execute('''SELECT r.club_name, r.title, r.max_count, r.current_count, r.deadline
                FROM recruitments r WHERE r.status="approved"
                AND julianday(r.deadline)-julianday("now","localtime") BETWEEN 0 AND 3
                AND r.current_count < r.max_count * 0.5''').fetchall()
            return [{'rule_id': 'recruitment_expiring', 'type': '招募即将截止', 'severity': 'info',
                     'message': f'{r["club_name"]}的"{r["title"]}"3天内截止，仅报名{r["current_count"]}/{r["max_count"]}人',
                     'club_name': r['club_name']} for r in rows]
        finally:
            conn.close()

    def _check_club_inactive(self):
        conn = self._get_conn()
        try:
            clubs = conn.execute('SELECT club_name FROM club_profiles').fetchall()
            alerts = []
            for c in clubs:
                recent = conn.execute('''SELECT COUNT(*) as c FROM checkin_sessions
                    WHERE club_name=? AND created_at>=datetime("now","localtime","-14 days")''', (c['club_name'],)).fetchone()
                if recent and recent['c'] == 0:
                    alerts.append({'rule_id': 'club_inactive', 'type': '社团活跃度骤降', 'severity': 'warning',
                                   'message': f'{c["club_name"]}近2周无任何签到活动',
                                   'club_name': c['club_name']})
            return alerts
        finally:
            conn.close()

    def _check_scoring_stuck(self):
        conn = self._get_conn()
        try:
            rows = conn.execute('''SELECT ss.club_name, ss.status, ss.updated_at
                FROM scoring_submissions ss
                WHERE ss.status IN ("pending_teacher","teacher_approved")
                AND julianday("now","localtime")-julianday(ss.updated_at)>5''').fetchall()
            return [{'rule_id': 'scoring_stuck', 'type': '赋分流程停滞', 'severity': 'warning',
                     'message': f'{r["club_name"]}的赋分流程停滞超过5天（当前状态：{r["status"]}）',
                     'club_name': r['club_name']} for r in rows]
        finally:
            conn.close()

    def _check_offcampus_overdue(self):
        conn = self._get_conn()
        try:
            rows = conn.execute('''SELECT id, club_name, title, created_at FROM offcampus_requests
                WHERE status="pending" AND julianday("now","localtime")-julianday(created_at)>3''').fetchall()
            return [{'rule_id': 'offcampus_overdue', 'type': '校外活动审批超时', 'severity': 'warning',
                     'message': f'{r["club_name"]}的校外活动"{r["title"]}"待审批超过3天',
                     'club_name': r['club_name']} for r in rows]
        finally:
            conn.close()

    def send_alerts(self):
        alerts = self.check_all()
        if not alerts:
            return {'sent': 0, 'message': '当前无异常预警'}
        conn = self.db.get_conn()
        try:
            sent = 0
            admin_users = conn.execute("SELECT id FROM users WHERE role='admin'").fetchall()
            for alert in alerts[:10]:
                for admin in admin_users:
                    existing = conn.execute('''SELECT id FROM notifications WHERE user_id=? AND title=? AND created_at>=datetime("now","localtime","-24 hours")''',
                                            (admin['id'], f'[预警] {alert["type"]}')).fetchone()
                    if not existing:
                        send_notification(admin['id'], f'[预警] {alert["type"]}', alert['message'], 'warning', '', conn=conn)
                        sent += 1
            conn.commit()
        finally:
            conn.close()
        return {'sent': sent, 'total_alerts': len(alerts), 'message': f'发现{len(alerts)}条预警，已发送{sent}条通知'}


class DataInsightAgent:
    def __init__(self, db_instance):
        self.db = db_instance

    def _get_conn(self):
        return self.db.get_conn()

    def generate_report(self, club_name=''):
        conn = self._get_conn()
        try:
            report = {'generated_at': '', 'sections': []}
            from datetime import datetime
            report['generated_at'] = datetime.now().strftime('%Y-%m-%d %H:%M')

            if club_name:
                clubs = [club_name]
            else:
                club_rows = conn.execute('SELECT club_name FROM club_profiles').fetchall()
                clubs = [c['club_name'] for c in club_rows]

            total_clubs = len(clubs)
            total_members = conn.execute('SELECT COUNT(DISTINCT id) as c FROM club_members').fetchone()['c']
            total_activities = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE is_completed=1').fetchone()['c']
            total_checkins = conn.execute('SELECT COUNT(*) as c FROM checkin_records').fetchone()['c']
            report['sections'].append({
                'title': '总体概况',
                'content': f'共{total_clubs}个社团，{total_members}名成员，{total_activities}场已完成活动，{total_checkins}次签到记录'
            })

            club_stats = []
            for cn in clubs:
                members = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE club_name=?', (cn,)).fetchone()['c']
                completed_activities = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=? AND is_completed=1', (cn,)).fetchone()['c']
                checkins = conn.execute('SELECT COUNT(*) as c FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=?', (cn,)).fetchone()['c']
                approved_uploads = conn.execute('SELECT COUNT(DISTINCT group_id) as c FROM club_uploads WHERE club_name=? AND status="approved" AND (source="upload" OR source="activity" OR source IS NULL)', (cn,)).fetchone()['c']
                star = conn.execute('SELECT star_rating FROM club_profiles WHERE club_name=?', (cn,)).fetchone()
                offcampus_approved = conn.execute('SELECT COUNT(*) as c FROM offcampus_requests WHERE club_name=? AND status="approved"', (cn,)).fetchone()['c']
                teacher_guided = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=? AND is_completed=1 AND teacher_ids IS NOT NULL AND teacher_ids!=""', (cn,)).fetchone()['c']
                club_stats.append({
                    'club_name': cn, 'members': members, 'activities': completed_activities,
                    'checkins': checkins, 'uploads': approved_uploads,
                    'star_rating': star['star_rating'] if star else 0,
                    'offcampus_approved': offcampus_approved,
                    'teacher_guided': teacher_guided,
                })

            max_offcampus = max((c['offcampus_approved'] for c in club_stats), default=0) or 0
            max_teacher = max((c['teacher_guided'] for c in club_stats), default=0) or 0
            max_activities = max((c['activities'] for c in club_stats), default=0) or 0
            max_uploads = max((c['uploads'] for c in club_stats), default=0) or 0
            for cs in club_stats:
                offcampus_score = (cs['offcampus_approved'] / max_offcampus * 100) if max_offcampus > 0 else 0
                teacher_score = (cs['teacher_guided'] / max_teacher * 100) if max_teacher > 0 else 0
                activity_score = (cs['activities'] / max_activities * 100) if max_activities > 0 else 0
                upload_score = (cs['uploads'] / max_uploads * 100) if max_uploads > 0 else 0
                cs['activity_score'] = round(activity_score, 1)
                cs['offcampus_score'] = round(offcampus_score, 1)
                cs['teacher_score'] = round(teacher_score, 1)
                cs['upload_score'] = round(upload_score, 1)
                cs['combined_score'] = round(offcampus_score * 0.20 + teacher_score * 0.15 + activity_score * 0.40 + upload_score * 0.25, 1)

            if not club_name:
                club_stats.sort(key=lambda x: x['combined_score'], reverse=True)
                top3 = club_stats[:3]
                bottom5 = [c for c in club_stats if c['combined_score'] == 0]
                content = '最活跃：' + '、'.join([c['club_name'] + '（' + str(c['combined_score']) + '分）' for c in top3])
                if bottom5:
                    content += '；零活跃：' + '、'.join([c['club_name'] for c in bottom5[:5]])
                report['sections'].append({
                    'title': '活跃度排名',
                    'content': content,
                })

            pending_materials = conn.execute('SELECT COUNT(DISTINCT group_id) as c FROM club_uploads WHERE status="pending" AND (source="upload" OR source="activity" OR source IS NULL)').fetchone()['c']
            pending_workload = conn.execute('SELECT COUNT(*) as c FROM workload_submissions WHERE status="pending"').fetchone()['c']
            pending_reg = conn.execute('SELECT COUNT(*) as c FROM club_registrations WHERE status="pending"').fetchone()['c']
            if pending_materials or pending_workload or pending_reg:
                report['sections'].append({
                    'title': '待办积压',
                    'content': f'待审批材料{pending_materials}组，待审核工作量{pending_workload}条，待审批报名{pending_reg}条'
                })

            recent_7d = conn.execute('''SELECT COUNT(*) as c FROM checkin_sessions WHERE created_at>=datetime("now","localtime","-7 days")''').fetchone()['c']
            prev_7d = conn.execute('''SELECT COUNT(*) as c FROM checkin_sessions WHERE created_at>=datetime("now","localtime","-14 days") AND created_at<datetime("now","localtime","-7 days")''').fetchone()['c']
            if prev_7d > 0:
                change = round((recent_7d - prev_7d) / prev_7d * 100, 1)
                direction = '上升' if change > 0 else '下降'
                report['sections'].append({
                    'title': '活动趋势',
                    'content': f'近7天{recent_7d}场活动，较前7天{direction}{abs(change)}%'
                })
            elif recent_7d > 0:
                report['sections'].append({
                    'title': '活动趋势',
                    'content': f'近7天{recent_7d}场活动'
                })

            if club_name and club_stats:
                cs = club_stats[0]
                report['sections'].append({
                    'title': f'{club_name}详情',
                    'content': f'{cs["star_rating"]}星级，{cs["members"]}名成员，{cs["activities"]}场活动，{cs["checkins"]}次签到，{cs["uploads"]}组已通过材料'
                })

            if QWEN_API_KEY:
                report_text = '\n'.join([f'【{s["title"]}】{s["content"]}' for s in report['sections']])
                messages = [
                    {'role': 'system', 'content': '你是一个数据分析助手。根据统计数据，给出2-3条具体的改进建议。只返回建议，每条一行。'},
                    {'role': 'user', 'content': report_text}
                ]
                ai_suggestions = call_llm_api(messages, max_tokens=300)
                if ai_suggestions:
                    report['sections'].append({'title': 'AI改进建议', 'content': ai_suggestions.strip()})

            return report
        finally:
            conn.close()


doc_index_agent = DocIndexAgent(db)
notification_agent = NotificationAgent(db)
data_insight_agent = DataInsightAgent(db)


@app.route('/api/agent/doc-index/rebuild', methods=['POST'])
def doc_index_rebuild():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可操作'}), 403
    result = doc_index_agent.rebuild_index()
    return jsonify(result)


@app.route('/api/agent/doc-index/download/<int:doc_id>')
def doc_index_download(doc_id):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT file_name, file_path, source_table, description, tags, category, club_name FROM doc_index WHERE id=?', (doc_id,)).fetchone()
    finally:
        conn.close()
    if not row or not row['file_path']:
        return jsonify({'error': '文件不存在或已被删除'}), 404
    actual_path = storage.get_path(row['file_path'])
    if not actual_path:
        url = storage.get_url(row['file_path'])
        if url:
            return jsonify({'url': url})
        return jsonify({'error': '文件不存在或已被删除'}), 404
    fname = row['file_name'] or os.path.basename(row['file_path'])
    import zipfile, io, tempfile
    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.write(actual_path, fname)
        info_lines = []
        info_lines.append('文件名称：' + (row['file_name'] or ''))
        info_lines.append('所属社团：' + (row['club_name'] or ''))
        info_lines.append('分类：' + (row['category'] or ''))
        info_lines.append('标签：' + (row['tags'] or ''))
        info_lines.append('主题/介绍：' + (row['description'] or ''))
        info_lines.append('来源表：' + (row['source_table'] or ''))
        zf.writestr('信息说明.txt', '\n'.join(info_lines).encode('utf-8'))
    mem_zip.seek(0)
    zip_name = os.path.splitext(fname)[0] + '_含信息.zip'
    return send_file(mem_zip, as_attachment=True, download_name=zip_name, mimetype='application/zip')


@app.route('/api/agent/doc-index/search')
def doc_index_search():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    query = request.args.get('q', '')
    club = request.args.get('club', '')
    category = request.args.get('category', '')
    results = doc_index_agent.search(query, club_name=club, category=category)
    return jsonify({'success': True, 'results': results})


@app.route('/api/agent/doc-index/smart-search')
def doc_index_smart_search():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    query = request.args.get('q', '')
    result = doc_index_agent.ai_smart_search(query)
    return jsonify({'success': True, 'data': result})


@app.route('/api/agent/doc-index/stats')
def doc_index_stats():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    stats = doc_index_agent.get_stats()
    return jsonify({'success': True, 'data': stats})


@app.route('/api/semantic-search', methods=['POST'])
def semantic_search():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    query = data.get('query', '').strip()
    top_k = data.get('top_k', 10)
    if not query:
        return jsonify({'error': '请输入搜索内容'}), 400
    doc_agent = DocIndexAgent(db)
    results = doc_agent.semantic_search(query, top_k)
    if isinstance(results, list):
        return jsonify({'success': True, 'total': len(results), 'items': results, 'query': query, 'search_type': 'keyword'})
    return jsonify({'success': True, **results})


@app.route('/api/rebuild-embeddings', methods=['POST'])
def rebuild_embeddings():
    user = get_current_user()
    if not user or user.get('role') != 'admin':
        return jsonify({'error': '仅管理员可操作'}), 403
    doc_agent = DocIndexAgent(db)
    conn = db.get_conn()
    try:
        docs = conn.execute('SELECT doc_id, doc_type, doc_text FROM doc_embeddings WHERE embedding IS NULL LIMIT 100').fetchall()
    finally:
        conn.close()
    indexed = 0
    for doc in docs:
        if doc_agent.index_document_embedding(doc['doc_id'], doc['doc_type'], doc['doc_text']):
            indexed += 1
    return jsonify({'success': True, 'indexed': indexed, 'total': len(docs)})


@app.route('/api/agent/notification/check')
def agent_notification_check():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    alerts = notification_agent.check_all()
    return jsonify({'success': True, 'alerts': alerts, 'total': len(alerts)})


@app.route('/api/agent/notification/send', methods=['POST'])
def agent_notification_send():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    result = notification_agent.send_alerts()
    return jsonify({'success': True, 'data': result})


@app.route('/api/agent/insight/report')
def agent_insight_report():
    user = get_current_user()
    if not user or user['role'] not in ('admin', 'user'):
        return jsonify({'error': '无权限'}), 403
    club = request.args.get('club', '')
    if user['role'] == 'user':
        club = user.get('clubName', '')
    report = data_insight_agent.generate_report(club_name=club)
    return jsonify({'success': True, 'data': report})


@app.route('/api/agent/approval-preview/<approval_type>')
def agent_approval_preview(approval_type):
    user = get_current_user()
    if not user or user['role'] not in ('admin', 'user'):
        return jsonify({'error': '无权限'}), 403
    if approval_type not in ApprovalAgent.APPROVAL_TYPES:
        return jsonify({'error': f'不支持的审批类型，可选：{",".join(ApprovalAgent.APPROVAL_TYPES)}'}), 400
    club_name = request.args.get('club', '')
    kwargs = {}
    if club_name:
        kwargs['club_name'] = club_name
    try:
        result = approval_agent.batch_analyze(approval_type, **kwargs)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 200
    return jsonify({'success': True, 'data': result})


@app.route('/api/agent/approval-analyze/<approval_type>/<item_id>')
def agent_approval_analyze(approval_type, item_id):
    user = get_current_user()
    if not user or user['role'] not in ('admin', 'user'):
        return jsonify({'error': '无权限'}), 403
    if approval_type not in ApprovalAgent.APPROVAL_TYPES:
        return jsonify({'error': f'不支持的审批类型'}), 400
    kwargs = {}
    if approval_type == 'material':
        kwargs['group_id'] = item_id
    elif approval_type == 'workload':
        kwargs['sub_id'] = int(item_id)
    elif approval_type == 'registration':
        kwargs['reg_id'] = int(item_id)
    elif approval_type == 'finance_permission':
        kwargs['pid'] = int(item_id)
    elif approval_type == 'offcampus':
        kwargs['rid'] = int(item_id)
    result = approval_agent.batch_analyze(approval_type, **kwargs)
    items = result.get('items', [])
    item = items[0] if items else None
    ai_result = None
    if item:
        ai_result = approval_agent.ai_deep_analyze(approval_type, item)
    return jsonify({'success': True, 'data': {'rule_based': item, 'ai_analysis': ai_result}})


@app.route('/api/agent/auto-approve/<approval_type>', methods=['POST'])
def agent_auto_approve(approval_type):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可执行自动审批'}), 403
    if approval_type not in ApprovalAgent.APPROVAL_TYPES:
        return jsonify({'error': '不支持的审批类型'}), 400
    result = approval_agent.batch_analyze(approval_type)
    auto_items = [item for item in result.get('items', []) if item.get('auto_approvable')]
    if not auto_items:
        return jsonify({'success': True, 'message': '没有可自动通过的项', 'approved': 0})
    approved_count = 0
    conn = db.get_conn()
    try:
        for item in auto_items:
            try:
                if approval_type == 'material':
                    gid = item['group_id']
                    conn.execute('UPDATE club_uploads SET status="approved" WHERE group_id=?', (gid,))
                    club_users = conn.execute('SELECT id FROM users WHERE club_name=?', (item['club_name'],)).fetchall()
                    for cu in club_users:
                        send_notification(cu['id'], '✅ 材料已通过（AI自动审批）', f'您提交的材料已通过AI智能审批', 'approve', '/upload.html', conn=conn)
                elif approval_type == 'workload':
                    conn.execute('UPDATE workload_submissions SET status="approved", reviewer_id=?, reviewer_name=?, review_note=?, reviewed_at=datetime("now","localtime") WHERE id=?',
                                 (user['id'], 'AI智能审批', 'AI自动审批：符合通过条件', item['id']))
                elif approval_type == 'registration':
                    reg = conn.execute('SELECT * FROM club_registrations WHERE id=?', (item['id'],)).fetchone()
                    if reg:
                        if reg['user_id'] and reg['user_id'] != 0:
                            existing = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=? AND user_id!=0',
                                                    (reg['club_name'], reg['user_id'])).fetchone()
                            if not existing:
                                conn.execute('''INSERT INTO club_members (club_name, user_id, username, real_name, student_id_num, class_name, phone, department, specialty, college, source)
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'ai_approval')''',
                                    (reg['club_name'], reg['user_id'], '', reg['student_name'], reg['student_id_num'],
                                     reg['student_class'], reg.get('student_phone', ''), item.get('suggested_department', '') or reg.get('department', ''),
                                     reg['specialty'], reg.get('college', '') if 'college' in reg.keys() else ''))
                        conn.execute('DELETE FROM club_registrations WHERE id=?', (item['id'],))
                elif approval_type == 'finance_permission':
                    record = conn.execute('SELECT * FROM finance_permissions WHERE id=?', (item['id'],)).fetchone()
                    if record:
                        conn.execute('UPDATE finance_permissions SET status="approved", reviewed_by=?, reviewed_at=CURRENT_TIMESTAMP WHERE id=?',
                                     ('ai_approval', item['id']))
                        conn.execute('INSERT OR IGNORE INTO finance_managers (club_name, user_id, username, real_name, granted_by) VALUES (?, ?, ?, ?, ?)',
                                     (record['club_name'], record['user_id'], record['username'], record['real_name'], 'ai_approval'))
                approved_count += 1
            except Exception:
                continue
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'approved': approved_count, 'message': f'AI自动审批完成，已通过{approved_count}项'})


def smart_intent_detect(message, user):
    import re
    msg = message.lower().strip()
    original = message.strip()
    role = user.get('role', 'student') if user else 'student'
    club_name = user.get('club_name', '') if user else ''
    intents = []

    def match(keywords):
        for kw in keywords:
            if kw in msg:
                return True
        return False

    def fuzzy_match(patterns):
        for p in patterns:
            if re.search(p, msg):
                return True
        return False

    approval_kw = ['审批', '待审', '审核', '待批', '待办', '待处理', '待审核', '待审批',
                   '审批情况', '审批状态', '审批建议', '帮我审', '有没有审批', '需要审批',
                   '看看审批', '审批了没', '批了没', '批一下', '通过一下', '驳回']
    approval_fuzzy = [r'有.*审', r'还.*审', r'等.*批', r'批.*吗', r'审.*情']
    if match(approval_kw) or fuzzy_match(approval_fuzzy):
        atype = 'material'
        if match(['工作量', '赋分', '工时', '课时']):
            atype = 'workload'
        elif match(['报名', '纳新', '招新', '招募', '入社']):
            atype = 'registration'
        elif match(['财务', '权限', '资金', '经费', '报销']):
            atype = 'finance_permission'
        elif match(['校外', '出校', '外出', '校外活动']):
            atype = 'offcampus'
        intents.append({'tool': 'check_approval', 'args': {'approval_type': atype}})

    search_kw = ['搜索', '查找', '找文件', '找资料', '搜文件', '搜索文件', '找一下',
                 '有没有文件', '有没有资料', '帮我找', '文件在哪', '资料在哪', '查找文件',
                 '检索', '搜索一下', '照片', '图片', '文档', '附件', '报表', '总结', '方案']
    search_fuzzy = [r'找.*文件', r'找.*资料', r'搜.*文件', r'有.*文件', r'有.*资料', r'哪.*文件']
    if match(search_kw) or fuzzy_match(search_fuzzy):
        query = original
        strip_prefixes = ['帮我搜索', '帮我找', '帮我查找', '搜索', '查找', '找一下',
                          '找', '搜一下', '搜索一下', '检索', '有没有', '帮我检索']
        for prefix in sorted(strip_prefixes, key=len, reverse=True):
            if query.startswith(prefix):
                query = query[len(prefix):].strip()
                break
        if not query or len(query) < 2:
            query = original
        args = {'query': query}
        detected_club = _extract_club_name(original, '')
        if detected_club:
            args['club_name'] = detected_club
        elif club_name:
            args['club_name'] = club_name
        intents.append({'tool': 'search_documents', 'args': args})

    alert_kw = ['预警', '异常', '超时', '警报', '风险', '告警', '系统异常',
                '有没有问题', '有没有异常', '系统状态', '运行情况', '健康状态',
                '有什么问题', '出什么问题了', '系统有问题', '系统怎么样']
    alert_fuzzy = [r'系统.*问题', r'系统.*异常', r'有.*问题.*吗', r'有.*异常.*吗', r'运行.*情况']
    if match(alert_kw) or fuzzy_match(alert_fuzzy):
        intents.append({'tool': 'check_alerts', 'args': {}})

    report_kw = ['分析报告', '数据报告', '洞察', '数据分析', '总体概况', '整体情况',
                 '全局分析', '生成报告', '出一份报告', '运营报告', '运营分析',
                 '综合分析', '全面分析', '整体分析', '数据总览', '数据概况', '运营情况']
    report_fuzzy = [r'分析.*报告', r'数据.*报告', r'报告.*吗', r'整体.*怎么样']
    if match(report_kw) or fuzzy_match(report_fuzzy):
        args = {}
        cn = _extract_club_name(original, club_name)
        if cn:
            args['club_name'] = cn
        intents.append({'tool': 'generate_insight_report', 'args': args})

    club_data_kw = ['成员', '社员', '活动数', '签到', '财务', '收支', '工作量', '赋分', '招募',
                    '多少人', '有多少', '人数', '活动情况', '签到情况', '财务情况', '收支情况',
                    '社团数据', '数据查询', '社团信息', '社团详情', '社团怎么样', '社团状况',
                    '成员数', '社员数', '活动几', '签到几', '收入', '支出', '余额', '经费']
    club_data_fuzzy = [r'有.*成员', r'有.*社员', r'多少.*人', r'几.*人', r'社团.*数据',
                       r'社团.*信息', r'社团.*详情', r'社团.*情况', r'社团.*怎么样',
                       r'财务.*怎么样', r'活动.*怎么样', r'签到.*怎么样']
    if match(club_data_kw) or fuzzy_match(club_data_fuzzy):
        cn = _extract_club_name(original, club_name)
        if cn:
            dtype = 'members'
            if match(['活动', '活动数', '活动情况', '活动几', '活动怎么样']):
                dtype = 'activities'
            elif match(['签到', '打卡', '签到情况', '签到几', '签到怎么样']):
                dtype = 'checkins'
            elif match(['财务', '收支', '收入', '支出', '经费', '余额', '财务情况', '财务怎么样']):
                dtype = 'finance'
            elif match(['工作量', '工时', '课时']):
                dtype = 'workload'
            elif match(['赋分', '评分', '打分', '学分']):
                dtype = 'scoring'
            elif match(['招募', '纳新', '招新']):
                dtype = 'recruitments'
            intents.append({'tool': 'query_club_data', 'args': {'club_name': cn, 'data_type': dtype}})
        else:
            if not any(i['tool'] == 'list_clubs' for i in intents):
                intents.append({'tool': 'list_clubs', 'args': {}})

    list_club_kw = ['有哪些社团', '社团列表', '所有社团', '社团有哪些', '社团名单',
                    '社团总览', '全部社团', '看看社团', '几个社团', '多少社团']
    list_club_fuzzy = [r'社团.*有哪些', r'社团.*列表', r'社团.*多少', r'有.*社团.*吗']
    if match(list_club_kw) or fuzzy_match(list_club_fuzzy):
        if not any(i['tool'] == 'list_clubs' for i in intents):
            intents.append({'tool': 'list_clubs', 'args': {}})

    # 优秀社团/优秀活动查询意图
    excellent_kw = ['优秀社团', '被评为优秀', '评优', '优秀名单', '评选优秀']
    excellent_fuzzy = [r'.*优秀.*社团.*', r'社团.*优秀.*', r'.*评优.*', r'.*优秀.*名单.*']
    if match(excellent_kw) or fuzzy_match(excellent_fuzzy):
        cn = _extract_club_name(original, club_name)
        filters = {}
        if cn: filters['club_name'] = cn
        intents.append({'tool': 'query_data', 'args': {'dataset': 'excellent_clubs', 'filters': filters}})

    # 学生未加入社团 / 老师未指导社团 查询意图
    no_club_student_kw = ['没有加入社团', '未加入社团', '没加社团', '哪些学生没有', '没有社团的学生',
                          '未入社', '没入社', '学生没有社团', '学生没加入']
    no_club_teacher_kw = ['没有指导社团', '未指导社团', '哪个老师没有', '没有社团的老师',
                          '老师没有指导', '老师未指导', '哪些老师没有']
    no_club_fuzzy = [r'学生.*没有.*社团', r'没有.*社团.*学生', r'学生.*未.*社团',
                     r'老师.*没有.*社团', r'没有.*社团.*老师', r'老师.*未.*社团']
    if match(no_club_student_kw) or (fuzzy_match(no_club_fuzzy) and match(['学生'])):
        intents.append({'tool': 'query_data', 'args': {'dataset': 'students', 'filters': {'has_club': 'false'}}})
    if match(no_club_teacher_kw) or (fuzzy_match(no_club_fuzzy) and match(['老师', '教师'])):
        intents.append({'tool': 'query_data', 'args': {'dataset': 'teachers', 'filters': {'has_club': 'false'}}})
    # 学生列表 / 老师列表 通用查询
    if match(['学生列表', '所有学生', '学生名单', '学生总数', '多少学生', '学生概况']):
        intents.append({'tool': 'query_data', 'args': {'dataset': 'students', 'filters': {}}})
    if match(['老师列表', '所有老师', '老师名单', '老师总数', '多少老师', '指导老师列表']):
        intents.append({'tool': 'query_data', 'args': {'dataset': 'teachers', 'filters': {}}})

    photo_kw = ['分析照片', '识别图片', '照片分析', '图片分析', '验证照片', '照片真伪',
                '照片验证', '识别文字', 'OCR', '图片内容', '看图', '分析图片',
                '照片内容', '图片识别', '活动照片', '现场照片', '照片审核']
    photo_fuzzy = [r'照片.*分析', r'图片.*分析', r'照片.*验证', r'验证.*照片',
                   r'识别.*图片', r'图片.*内容', r'照片.*真伪', r'分析.*图片']
    if match(photo_kw) or fuzzy_match(photo_fuzzy):
        analysis_type = 'describe'
        if match(['验证', '真伪', '审核', '是否真实', '摆拍']):
            analysis_type = 'verify'
        elif match(['OCR', '文字', '识别文字', '提取文字']):
            analysis_type = 'ocr'
        intents.append({'tool': 'analyze_photo', 'args': {'analysis_type': analysis_type}})

    # 创建活动意图
    create_activity_kw = ['创建活动', '办活动', '发起活动', '新建活动', '添加活动', '举办活动', '开展活动',
                          '办个活动', '搞个活动', '组织活动', '策划活动', '办一场', '搞一场',
                          '办个比赛', '办比赛', '篮球赛', '足球赛', '运动会', '文艺演出', '晚会']
    create_activity_fuzzy = [r'办.*活动', r'搞.*活动', r'创建.*活动', r'发起.*活动', r'举办.*活动', r'组织.*活动', r'想办', r'想搞']
    if match(create_activity_kw) or fuzzy_match(create_activity_fuzzy):
        cn = _extract_club_name(original, club_name)
        attrs = {'name': original}
        if cn:
            attrs['club_name'] = cn
        intents.append({'tool': 'create_entity', 'args': {'entity_type': 'activity', 'attributes': attrs}})

    # 创建招募意图
    create_recruit_kw = ['发布招募', '发起招募', '创建招募', '纳新', '招新', '招募新成员', '发招募令']
    if match(create_recruit_kw):
        cn = _extract_club_name(original, club_name)
        attrs = {'name': original}
        if cn:
            attrs['club_name'] = cn
        intents.append({'tool': 'create_entity', 'args': {'entity_type': 'recruitment', 'attributes': attrs}})

    # 发送通知意图
    notify_kw = ['发通知', '发送通知', '通知成员', '通知大家', '提醒成员', '发公告', '发布公告',
                 '通知全体', '通知所有人', '群发通知', '发个通知', '提醒一下']
    notify_fuzzy = [r'通知.*成员', r'通知.*大家', r'发.*通知', r'发.*公告', r'提醒.*成员']
    if match(notify_kw) or fuzzy_match(notify_fuzzy):
        cn = _extract_club_name(original, club_name)
        args = {'target': 'all_members', 'content': original}
        if cn:
            args['club_name'] = cn
        intents.append({'tool': 'send_notification', 'args': args})

    # 生成文案意图
    copywriting_kw = ['生成文案', '写文案', '写宣传语', '宣传文案', '活动文案', '招募文案',
                      '帮我写文案', '帮我写宣传', '写个文案', '写一段文案']
    if match(copywriting_kw):
        cn = _extract_club_name(original, club_name)
        params = {'theme': original}
        if cn:
            params['club_name'] = cn
        intents.append({'tool': 'ai_generate', 'args': {'generate_type': 'copywriting', 'params': params}})

    # 推荐意图
    recommend_club_kw = ['推荐社团', '推荐一个社团', '有什么社团推荐', '适合我的社团', '感兴趣的社团',
                         '还可以加入', '还能加入', '可以加入', '加入什么社团', '还能报什么', '还能参加什么社团']
    recommend_activity_kw = ['推荐活动', '推荐一个活动', '有什么活动推荐', '适合我的活动', '感兴趣的活动']
    recommend_recruitment_kw = ['推荐招募', '报名招募', '想报名', '招募活动', '招募推荐', '有什么招募', '推荐报名', '想报一个招募', '招募报名']
    recommend_partner_kw = ['推荐联合', '联合活动伙伴', '合作社团', '一起办活动', '联合举办']
    if match(recommend_club_kw):
        intents.append({'tool': 'recommend', 'args': {'recommend_type': 'club'}})
    if match(recommend_activity_kw):
        intents.append({'tool': 'recommend', 'args': {'recommend_type': 'activity'}})
    if match(recommend_recruitment_kw):
        intents.append({'tool': 'recommend', 'args': {'recommend_type': 'recruitment'}})
    if match(recommend_partner_kw):
        cn = _extract_club_name(original, club_name)
        args = {'recommend_type': 'joint_partner'}
        if cn:
            args['club_name'] = cn
        intents.append({'tool': 'recommend', 'args': args})

    # 生成报告意图
    report_kw = ['生成报告', '数据报告', '分析报告', '统计报告', '活动概览', '财务汇总',
                 '工作量统计', '赋分分布', '参与率', '预警报告', '社团健康', '健康度']
    report_fuzzy = [r'生成.*报告', r'出.*报告', r'统计.*报告', r'分析.*报告']
    if match(report_kw) or fuzzy_match(report_fuzzy):
        report_type = 'activity_overview'
        if match(['工作量', '工时']):
            report_type = 'workload_stats'
        elif match(['财务', '收支', '经费']):
            report_type = 'finance_summary'
        elif match(['赋分', '评分']):
            report_type = 'scoring_distribution'
        elif match(['指导', '老师']):
            report_type = 'teacher_guidance'
        elif match(['参与率', '出勤']):
            report_type = 'participation_trend'
        elif match(['预警', '异常']):
            report_type = 'alert_summary'
        elif match(['健康', '社团评估']):
            report_type = 'club_health'
        cn = _extract_club_name(original, club_name)
        args = {'report_type': report_type}
        if cn:
            args['club_name'] = cn
        intents.append({'tool': 'generate_report', 'args': args})

    # 审批操作意图
    approve_kw = ['通过审批', '批准', '审批通过', '同意审批', '驳回审批', '拒绝审批',
                  '帮我审批', '批量通过', '批量驳回']
    if match(approve_kw):
        action = 'approve' if match(['通过', '批准', '同意', '批准']) else 'reject'
        intents.append({'tool': 'update_entity', 'args': {'entity_type': 'material', 'entity_id': 0, 'updates': {'action': action}}})

    return intents


def _extract_club_name(message, default=''):
    known_clubs = []
    try:
        conn = db.get_conn()
        rows = conn.execute('SELECT DISTINCT club_name FROM club_profiles').fetchall()
        known_clubs = [r['club_name'] for r in rows if r['club_name']]
        conn.close()
    except:
        pass
    if not known_clubs:
        try:
            conn = db.get_conn()
            rows = conn.execute('SELECT DISTINCT club_name FROM users WHERE role="user" AND club_name!=""').fetchall()
            known_clubs = [r['club_name'] for r in rows if r['club_name']]
            conn.close()
        except:
            pass
    for cn in known_clubs:
        if cn in message:
            return cn
    return default


def format_tool_result_for_human(tool_name, tool_args, raw_result):
    try:
        data = json.loads(raw_result) if isinstance(raw_result, str) else raw_result
    except:
        return raw_result

    if tool_name == 'check_approval':
        items = data if isinstance(data, list) else data.get('items', data.get('results', [data]))
        if not items or (isinstance(items, list) and len(items) == 0):
            atype = tool_args.get('approval_type', 'material')
            type_names = {'material': '材料', 'workload': '工作量', 'registration': '报名', 'finance_permission': '财务权限', 'offcampus': '校外活动'}
            return f'✅ 当前没有待审批的{type_names.get(atype, "")}项目，一切正常！\n\n💡 系统运行良好，你可以去查看其他类型的审批或生成数据报告。'
        lines = ['📋 **审批分析结果：**\n']
        approve_count = 0
        reject_count = 0
        if isinstance(items, list):
            for item in items[:8]:
                if isinstance(item, dict):
                    name = item.get('club_name', item.get('name', ''))
                    suggestion = item.get('suggestion', item.get('recommendation', ''))
                    reason = item.get('reason', item.get('analysis', ''))
                    confidence = item.get('confidence', item.get('score', 0))
                    sug_icon = '✅' if 'approve' in str(suggestion).lower() or '通过' in str(suggestion) else '❌' if 'reject' in str(suggestion).lower() or '驳回' in str(suggestion) else '🔍'
                    if 'approve' in str(suggestion).lower() or '通过' in str(suggestion):
                        approve_count += 1
                    elif 'reject' in str(suggestion).lower() or '驳回' in str(suggestion):
                        reject_count += 1
                    line = f'{sug_icon} **{name}**'
                    if suggestion:
                        line += f' → {suggestion}'
                    if confidence and isinstance(confidence, (int, float)):
                        line += f'（置信度 {min(int(confidence),99) if confidence > 1 else min(int(confidence*100),99)}%）'
                    if reason:
                        line += f'\n   💬 {reason}'
                    lines.append(line)
            if len(items) > 8:
                lines.append(f'\n...还有 {len(items)-8} 项')
            lines.append(f'\n📊 **汇总：** 共 {len(items)} 项待审批')
            if approve_count:
                lines.append(f'  建议通过 {approve_count} 项')
            if reject_count:
                lines.append(f'  建议驳回 {reject_count} 项')
        return '\n'.join(lines)

    elif tool_name == 'search_documents':
        total = data.get('total', 0)
        items = data.get('items', [])
        summary = data.get('summary', '')
        if total == 0:
            return '🔍 没有找到相关文件。\n\n💡 试试换个关键词搜索，比如"活动照片"、"财务报表"、"报名表"等。'
        lines = [f'🔍 **找到 {total} 个相关文件：**\n']
        clubs_in_result = set()
        for item in items[:6]:
            fn = item.get('file_name', item.get('name', ''))
            cn = item.get('club_name', '')
            tags = item.get('tags', [])
            if cn:
                clubs_in_result.add(cn)
            line = f'• **{fn}**'
            if cn:
                line += f' ({cn})'
            if tags:
                valid_tags = [t for t in tags[:3] if t]
                if valid_tags:
                    line += f' `{"` `".join(valid_tags)}`'
            lines.append(line)
        if total > 6:
            lines.append(f'\n...还有 {total-6} 个文件')
        if clubs_in_result:
            lines.append(f'\n📂 涉及社团：{"、".join(clubs_in_result)}')
        return '\n'.join(lines)

    elif tool_name == 'check_alerts':
        total = data.get('total', 0)
        alerts = data.get('alerts', [])
        message = data.get('message', '')
        # 检查是否有权限错误信息
        if message and ('仅管理员可查看' in message or '无权限' in message):
            return f'⚠️ {message}\n\n💡 你可以询问关于自己社团的活动、成员、财务等数据，我会帮你查询。'
        if total == 0:
            return '✅ 系统运行正常，没有异常预警！\n\n💡 一切正常，你可以查看审批情况或生成数据报告。'
        lines = [f'⚠️ **发现 {total} 条预警：**\n']
        critical_count = 0
        for a in alerts[:6]:
            if isinstance(a, dict):
                level = a.get('level', a.get('severity', 'info'))
                title = a.get('title', a.get('type', ''))
                desc = a.get('description', a.get('message', ''))
                icon = '🔴' if level in ['critical', 'high', 'error'] else '🟡' if level in ['warning', 'medium'] else '🔵'
                if level in ['critical', 'high', 'error']:
                    critical_count += 1
                lines.append(f'{icon} **{title}**\n   {desc}')
            else:
                lines.append(f'• {a}')
        if critical_count > 0:
            lines.append(f'\n🚨 **有 {critical_count} 条紧急预警需要立即处理！**')
        return '\n'.join(lines)

    elif tool_name == 'generate_insight_report':
        if isinstance(data, dict):
            lines = ['📊 **数据分析报告**\n']
            sections = data.get('sections', [])
            if sections and isinstance(sections, list):
                for sec in sections:
                    if isinstance(sec, dict):
                        title = sec.get('title', sec.get('section_title', ''))
                        content = sec.get('content', sec.get('data', sec.get('items', '')))
                        if title:
                            lines.append(f'\n**{title}**')
                        if isinstance(content, str):
                            lines.append(f'  {content}')
                        elif isinstance(content, list):
                            for item in content[:8]:
                                if isinstance(item, dict):
                                    name = item.get('club_name', item.get('name', item.get('label', '')))
                                    val = item.get('value', item.get('score', item.get('count', item.get('total', ''))))
                                    extra = item.get('trend', item.get('status', ''))
                                    line = f'  • {name}: **{val}**'
                                    if extra:
                                        line += f' ({extra})'
                                    lines.append(line)
                                else:
                                    lines.append(f'  • {item}')
                            if len(content) > 8:
                                lines.append(f'  ...还有 {len(content)-8} 项')
                        elif isinstance(content, dict):
                            for k, v in list(content.items())[:8]:
                                lines.append(f'  • {k}: **{v}**')
                if len(sections) == 0:
                    for key, val in data.items():
                        if key == 'generated_at':
                            continue
                        if isinstance(val, (str, int, float)):
                            lines.append(f'• **{key}**: {val}')
                        elif isinstance(val, list) and val:
                            lines.append(f'\n**{key}：**')
                            for item in val[:5]:
                                if isinstance(item, dict):
                                    name = item.get('club_name', item.get('name', item.get('label', '')))
                                    val2 = item.get('value', item.get('score', item.get('count', '')))
                                    lines.append(f'  • {name}: {val2}')
                                else:
                                    lines.append(f'  • {item}')
                        elif isinstance(val, dict):
                            lines.append(f'\n**{key}：**')
                            for k2, v2 in list(val.items())[:5]:
                                lines.append(f'  • {k2}: {v2}')
            else:
                for key, val in data.items():
                    if key == 'generated_at':
                        continue
                    if isinstance(val, (str, int, float)):
                        lines.append(f'• **{key}**: {val}')
                    elif isinstance(val, list) and val:
                        lines.append(f'\n**{key}：**')
                        for item in val[:5]:
                            if isinstance(item, dict):
                                name = item.get('club_name', item.get('name', item.get('label', '')))
                                val2 = item.get('value', item.get('score', item.get('count', '')))
                                lines.append(f'  • {name}: {val2}')
                            else:
                                lines.append(f'  • {item}')
                    elif isinstance(val, dict):
                        lines.append(f'\n**{key}：**')
                        for k2, v2 in list(val.items())[:5]:
                            lines.append(f'  • {k2}: {v2}')
            return '\n'.join(lines)
        return str(data)

    elif tool_name == 'query_club_data':
        cn = tool_args.get('club_name', '')
        dtype = tool_args.get('data_type', 'members')
        if isinstance(data, dict) and data.get('error'):
            return f'❌ {data["error"]}\n\n💡 你可以先问"有哪些社团"查看社团列表。'
        lines = []
        if dtype == 'members':
            total = data.get('total', 0)
            leaders = data.get('leaders', [])
            lines.append(f'👥 **{cn} 成员数据：**\n')
            lines.append(f'• 成员总数：**{total}** 人')
            if total == 0:
                lines.append('\n⚠️ 该社团暂无成员数据，可能需要添加成员。')
            elif total < 5:
                lines.append(f'\n💡 成员数较少（{total}人），建议开展纳新活动扩大规模。')
            elif total >= 20:
                lines.append(f'\n💡 成员规模不错！可以考虑设置更多部门来提高管理效率。')
            if leaders:
                lines.append(f'\n**骨干成员：**')
                for l in leaders[:5]:
                    lines.append(f'  • {l.get("real_name", "")} - {l.get("department", "")}')
        elif dtype == 'activities':
            total = data.get('total', 0)
            completed = data.get('completed', 0)
            recent = data.get('recent', [])
            lines.append(f'🎯 **{cn} 活动数据：**\n')
            lines.append(f'• 活动总数：**{total}** 次')
            lines.append(f'• 已完成：**{completed}** 次')
            if total > 0:
                completion_rate = round(completed / total * 100)
                lines.append(f'• 完成率：**{completion_rate}%**')
                if completion_rate < 50:
                    lines.append(f'\n⚠️ 活动完成率偏低，建议关注未完成活动的推进。')
            if recent:
                lines.append(f'\n**近期活动：**')
                for r in recent[:5]:
                    lines.append(f'  • {r.get("activity_name", "未命名")} - {r.get("created_at", "")}')
        elif dtype == 'finance':
            income = data.get('income', 0)
            expense = data.get('expense', 0)
            balance = data.get('balance', 0)
            lines.append(f'💰 **{cn} 财务数据：**\n')
            lines.append(f'• 收入：**¥{income}**')
            lines.append(f'• 支出：**¥{expense}**')
            bal_icon = '🟢' if balance >= 0 else '🔴'
            lines.append(f'• 余额：{bal_icon} **¥{balance}**')
            if balance < 0:
                lines.append(f'\n🚨 **财务预警：余额为负！** 建议尽快核实收支情况。')
            elif income > 0 and expense == 0:
                lines.append(f'\n💡 有收入无支出记录，记得及时登记支出。')
            elif expense > income:
                lines.append(f'\n⚠️ 支出超过收入，建议控制开支。')
        elif dtype == 'checkins':
            total = data.get('total_checkins', 0)
            lines.append(f'✅ **{cn} 签到数据：**\n')
            lines.append(f'• 签到总人次：**{total}**')
            if total == 0:
                lines.append('\n💡 暂无签到记录，建议开展签到活动。')
        elif dtype == 'workload':
            pending = data.get('pending', 0)
            approved = data.get('approved', 0)
            lines.append(f'📋 **{cn} 工作量数据：**\n')
            lines.append(f'• 待审核：**{pending}** 项')
            lines.append(f'• 已通过：**{approved}** 项')
            if pending > 5:
                lines.append(f'\n⚠️ 有 {pending} 项待审核工作量，建议尽快处理。')
        elif dtype == 'scoring':
            subs = data.get('submissions', [])
            lines.append(f'📝 **{cn} 赋分数据：**\n')
            for s in subs:
                lines.append(f'• {s.get("status", "")}：**{s.get("c", 0)}** 项')
        elif dtype == 'recruitments':
            total = data.get('total', 0)
            active = data.get('active', 0)
            lines.append(f'📢 **{cn} 招募数据：**\n')
            lines.append(f'• 招募总数：**{total}** 个')
            lines.append(f'• 进行中：**{active}** 个')
            if active == 0 and total > 0:
                lines.append(f'\n💡 目前没有进行中的招募，可以考虑发起新一轮纳新。')
        else:
            for k, v in data.items():
                lines.append(f'• **{k}**: {v}')
        return '\n'.join(lines) if lines else str(data)

    elif tool_name == 'list_clubs':
        clubs = data.get('clubs', [])
        if not clubs:
            return '🏫 暂无社团数据'
        lines = [f'🏫 **共 {len(clubs)} 个社团：**\n']
        categories = {}
        for i, c in enumerate(clubs, 1):
            name = c.get('club_name', '')
            star = c.get('star_rating', '')
            cat = c.get('category', '')
            if cat:
                categories[cat] = categories.get(cat, 0) + 1
            star_str = '⭐' * min(int(star or 0), 5) if star else ''
            line = f'{i}. **{name}**'
            if star_str:
                line += f' {star_str}'
            if cat:
                line += f' · {cat}'
            lines.append(line)
        if categories:
            lines.append(f'\n📊 **分类统计：**')
            for cat, count in sorted(categories.items(), key=lambda x: -x[1]):
                lines.append(f'  • {cat}：{count}个')
        return '\n'.join(lines)

    elif tool_name == 'analyze_photo':
        analysis_type = tool_args.get('analysis_type', 'describe')
        type_labels = {'describe': '图片描述', 'verify': '照片验证', 'ocr': '文字识别'}
        type_icons = {'describe': '🖼️', 'verify': '🔍', 'ocr': '📝'}
        icon = type_icons.get(analysis_type, '👁')
        label = type_labels.get(analysis_type, '图片分析')
        if isinstance(data, dict):
            is_fallback = data.get('fallback', False)
            analysis = data.get('analysis', '')
        else:
            is_fallback = False
            analysis = str(data)
        lines = [f'{icon} **{label}结果：**\n']
        lines.append(analysis)
        if is_fallback:
            lines.append('\n💡 视觉模型未配置，以上为离线提示。配置 QWEN_API_KEY 后可使用AI分析。')
        return '\n'.join(lines)

    elif tool_name == 'query_data':
        dataset = tool_args.get('dataset', '')
        items = data.get('items', [])
        total = data.get('total', len(items))
        dataset_names = {'activities': '活动', 'members': '成员', 'students': '学生', 'teachers': '指导老师', 'workload': '工作量',
                        'finance': '财务', 'recruitments': '招募', 'scoring': '赋分', 'credits': '学分',
                        'checkins': '签到', 'treehole': '树洞', 'feedback': '反馈', 'carousel': '轮播图',
                        'excellent_clubs': '优秀社团', 'excellent_activities': '优秀活动', 'hot_news': '热点资讯', 'documents': '资料库',
                        'clubs': '社团', 'notifications': '通知', 'departments': '部门', 'votes': '投票',
                        'surveys': '问卷', 'joint_activities': '联合活动', 'offcampus_requests': '校外活动申请'}
        dname = dataset_names.get(dataset, dataset)
        desc = data.get('description', '')
        if not items and not data.get('income') and data.get('error'):
            return f'❌ {data.get("error", "查询失败")}'
        lines = [f'📊 **{desc or dname + "数据"}查询结果：**\n']
        if data.get('income') is not None or data.get('expense') is not None:
            lines.append(f'• 收入：**¥{data.get("income", 0)}**')
            lines.append(f'• 支出：**¥{data.get("expense", 0)}**')
            bal = data.get('balance', 0)
            lines.append(f'• 余额：{"🟢" if bal >= 0 else "🔴"} **¥{bal}**')
        if total:
            lines.append(f'• 共 **{total}** 条记录')
        for item in items[:6]:
            if isinstance(item, dict):
                name = item.get('activity_name', item.get('real_name', item.get('title', item.get('club_name', item.get('name', '')))))
                status = item.get('status', '')
                extra = item.get('club_name', '') or item.get('department', '') or item.get('category', '')
                line = f'  • {name}'
                if extra and extra != name:
                    line += f' ({extra})'
                if status:
                    line += f' [{status}]'
                lines.append(line)
        if len(items) > 6:
            lines.append(f'  ...还有 {len(items)-6} 条')
        return '\n'.join(lines)

    elif tool_name == 'create_entity':
        if data.get('success'):
            return f'✅ {data.get("message", "创建成功")}'
        return f'❌ {data.get("error", "创建失败")}'

    elif tool_name == 'update_entity':
        if data.get('success'):
            return f'✅ {data.get("message", "更新成功")}'
        return f'❌ {data.get("error", "更新失败")}'

    elif tool_name == 'send_notification':
        if data.get('success'):
            return f'📢 {data.get("message", "通知已发送")}'
        return f'❌ {data.get("error", "发送失败")}'

    elif tool_name == 'generate_report':
        items = data.get('items', [])
        if not items and not data.get('total_activities') and data.get('error'):
            return f'❌ {data["error"]}'
        lines = ['📑 **数据报告：**\n']
        if data.get('total_activities') is not None:
            lines.append(f'• 活动总数：**{data["total_activities"]}**')
            lines.append(f'• 已完成：**{data.get("completed", 0)}**')
            lines.append(f'• 完成率：**{data.get("completion_rate", 0)}%**')
        # 检查是否为 teacher_guidance 报告（有 teacher_name 字段）
        if items and isinstance(items[0], dict) and 'teacher_name' in items[0]:
            lines.append('**指导老师指导情况：**\n')
            for item in items[:10]:
                tname = item.get('teacher_name', '未知')
                cn = item.get('club_name', '')
                gc = item.get('guidance_count', 0)
                ip = item.get('in_progress', 0)
                th = item.get('total_hours', 0)
                ac = item.get('activity_count', 0)
                line = f'  • **{tname}** — 指导社团：{cn}'
                line += f'，已指导 **{gc}** 次'
                if ip > 0:
                    line += f'（进行中 {ip} 次）'
                if th > 0:
                    line += f'，累计 **{th}** 小时'
                line += f'，社团活动 **{ac}** 场'
                lines.append(line)
        else:
            for item in items[:8]:
                if isinstance(item, dict):
                    name = item.get('club_name', item.get('name', ''))
                    vals = [f'{k}={v}' for k, v in item.items() if k != 'club_name' and k != 'name' and v is not None]
                    line = f'  • **{name}**'
                    if vals:
                        line += f' — {", ".join(vals[:4])}'
                    lines.append(line)
        if len(items) > 10:
            lines.append(f'  ...还有 {len(items)-10} 条')
        return '\n'.join(lines)

    elif tool_name == 'ai_generate':
        if data.get('success'):
            content = data.get('content', data.get('message', ''))
            return f'🎨 **内容已生成：**\n\n{content}'
        error = data.get('error', '生成失败')
        hint = data.get('hint', '')
        prompt = data.get('prompt', '')
        result = f'❌ {error}'
        if hint:
            result += f'\n\n💡 {hint}'
        if prompt:
            result += f'\n\n📝 已生成提示词：{prompt}'
        return result

    elif tool_name == 'recommend':
        recs = data.get('recommendations', [])
        msg = data.get('message', '')
        remaining = data.get('remaining_slots')
        joined_clubs = data.get('joined_clubs', [])
        if not recs and remaining is not None and remaining == 0:
            return f'💡 {msg}\n\n您已加入：{"、".join(joined_clubs)}'
        if not recs:
            return f'💡 {msg}\n\n暂无推荐结果。'
        lines = [f'💡 **{msg}：**\n']
        if joined_clubs:
            lines.append(f'📋 已加入：**{"、".join(joined_clubs)}**')
        if remaining is not None:
            lines.append(f'📌 还可加入 **{remaining}** 个社团（每人最多2个）\n')
        for i, r in enumerate(recs, 1):
            if isinstance(r, dict):
                # 招募推荐
                title = r.get('title', '')
                if title:
                    club = r.get('club_name', '')
                    desc = r.get('description', '')
                    signup_count = r.get('signup_count', 0)
                    already = r.get('already_signed', False)
                    line = f'{i}. **{title}** — {club}'
                    if signup_count:
                        line += f'（{signup_count}人已报名）'
                    if already:
                        line += ' ✅已报名'
                    if desc:
                        line += f'\n   {desc[:60]}'
                    lines.append(line)
                    continue
                # 社团/活动推荐
                name = r.get('club_name', r.get('activity_name', ''))
                cat = r.get('category', '')
                star = r.get('star_rating', '')
                count = r.get('member_count', r.get('participant_count', r.get('activity_count', '')))
                line = f'{i}. **{name}**'
                if cat:
                    line += f' · {cat}'
                if star:
                    line += f' {"⭐" * min(int(star), 5)}'
                if count:
                    line += f' ({count}人)'
                lines.append(line)
        return '\n'.join(lines)

    elif tool_name == 'query_database':
        items = data.get('items', [])
        total = data.get('total', 0)
        desc = data.get('description', '')
        if data.get('error'):
            return f'❌ {data["error"]}'
        if not items:
            return f'🔍 **数据库查询**（{desc}）：无匹配结果'
        lines = [f'🔍 **数据库查询**（{desc}）：共 **{total}** 条记录\n']
        # 动态展示列名
        if items and isinstance(items[0], dict):
            cols = list(items[0].keys())
            for item in items[:8]:
                vals = [f'{k}={v}' for k, v in item.items() if v is not None and k != 'id']
                lines.append(f'  • {" | ".join(vals[:6])}')
            if total > 8:
                lines.append(f'  ...还有 {total-8} 条')
        return '\n'.join(lines)

    return json.dumps(data, ensure_ascii=False, default=str)[:500]


def generate_smart_suggestions(tool_name, tool_result_data, user):
    suggestions = []
    role = user.get('role', 'student') if user else 'student'
    club_name = user.get('club_name', '') if user else ''

    if tool_name == 'check_approval':
        try:
            items = tool_result_data if isinstance(tool_result_data, list) else tool_result_data.get('items', tool_result_data.get('results', []))
            if not items or len(items) == 0:
                suggestions.append('检查系统预警')
                suggestions.append('生成数据分析报告')
            else:
                suggestions.append('查看其他类型的审批')
                if role == 'admin':
                    suggestions.append('生成数据分析报告')
                suggestions.append('检查系统预警')
        except:
            suggestions.append('查看其他类型的审批')
    elif tool_name == 'search_documents':
        try:
            total = tool_result_data.get('total', 0)
            if total == 0:
                suggestions.append('换个关键词搜索')
            else:
                suggestions.append('查看审批情况')
            suggestions.append('检查系统预警')
        except:
            suggestions.append('查看审批情况')
    elif tool_name == 'check_alerts':
        try:
            total = tool_result_data.get('total', 0)
            if total > 0:
                suggestions.append('查看待审批项目')
                suggestions.append('生成数据分析报告')
            else:
                suggestions.append('查看待审批项目')
                suggestions.append('有哪些社团')
        except:
            suggestions.append('查看待审批项目')
    elif tool_name == 'generate_insight_report':
        suggestions.append('查看系统预警')
        suggestions.append('查看待审批项目')
        if club_name:
            suggestions.append(f'查看{club_name}成员数据')
    elif tool_name == 'query_club_data':
        dtype = 'members'
        try:
            if isinstance(tool_result_data, dict):
                keys = list(tool_result_data.keys())
                if 'items' in keys:
                    first_item = tool_result_data['items'][0] if tool_result_data['items'] else {}
                    item_keys = list(first_item.keys()) if first_item else []
                    if 'activity_name' in item_keys: dtype = 'activities'
                    elif 'amount' in item_keys or 'income' in item_keys: dtype = 'finance'
                    elif 'checkin_time' in item_keys or 'session_id' in item_keys: dtype = 'checkins'
                    elif 'item_name' in item_keys or 'score' in item_keys: dtype = 'workload'
        except:
            pass
        if club_name:
            other_types = {'members': '活动数据', 'activities': '财务数据', 'finance': '签到数据', 'checkins': '工作量数据', 'workload': '成员数据'}
            other = other_types.get(dtype, '其他数据')
            suggestions.append(f'查看{club_name}的{other}')
        suggestions.append('生成分析报告')
    elif tool_name == 'list_clubs':
        suggestions.append('查看某个社团的详细数据')
        suggestions.append('生成全校数据报告')
    elif tool_name == 'query_data':
        dataset = 'members'
        try:
            if isinstance(tool_result_data, dict):
                keys = list(tool_result_data.keys())
                if 'items' in keys:
                    first_item = tool_result_data['items'][0] if tool_result_data['items'] else {}
                    item_keys = list(first_item.keys()) if first_item else []
                    if 'activity_name' in item_keys: dataset = 'activities'
                    elif 'amount' in item_keys or 'income' in item_keys: dataset = 'finance'
                    elif 'checkin_time' in item_keys: dataset = 'checkins'
                    elif 'item_name' in item_keys: dataset = 'workload'
                    elif 'content' in item_keys and 'scope' in item_keys: dataset = 'treehole'
                    elif 'title' in item_keys and 'body' in item_keys: dataset = 'feedback'
                    elif 'selected_at' in item_keys and 'club_name' in item_keys: dataset = 'excellent_clubs'
                    elif 'selected_at' in item_keys and 'group_id' in item_keys: dataset = 'excellent_activities'
        except:
            pass
        if club_name:
            suggestions.append(f'查看{club_name}的其他数据')
        suggestions.append('生成分析报告')
        if role == 'student':
            suggestions.append('推荐我感兴趣的活动')
        elif role == 'admin':
            suggestions.append('检查系统预警')
    elif tool_name == 'create_entity':
        suggestions.append('查看创建结果')
        if role in ('user', 'admin'):
            suggestions.append('发通知告诉成员')
            suggestions.append('生成活动海报')
    elif tool_name == 'update_entity':
        suggestions.append('查看审批情况')
        suggestions.append('生成分析报告')
    elif tool_name == 'send_notification':
        suggestions.append('查看通知记录')
        suggestions.append('查看审批情况')
    elif tool_name == 'generate_report':
        suggestions.append('检查系统预警')
        if club_name:
            suggestions.append(f'查看{club_name}成员数据')
        suggestions.append('查看待审批项目')
    elif tool_name == 'ai_generate':
        if role in ('user', 'admin'):
            suggestions.append('创建活动')
            suggestions.append('发通知告诉成员')
        suggestions.append('查看社团数据')
    elif tool_name == 'recommend':
        if role == 'student':
            suggestions.append('推荐我感兴趣的社团')
            suggestions.append('查看我的工作量')
        elif role == 'user':
            suggestions.append('查看社团数据')
            suggestions.append('生成分析报告')

    # 角色通用建议补充
    if not suggestions:
        if role == 'admin':
            suggestions.extend(['查看待审批项目', '检查系统预警', '生成全校数据报告'])
        elif role == 'user':
            suggestions.extend(['查看社团数据', '创建活动', '生成分析报告'])
        elif role == 'teacher':
            suggestions.extend(['查看赋分审核', '查看指导情况', '检查系统预警'])
        elif role == 'student':
            suggestions.extend(['推荐社团', '推荐活动', '查看我的数据'])

    return list(dict.fromkeys(suggestions))[:3]


@app.route('/api/ai-chat', methods=['POST'])
def ai_chat():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    message = data.get('message', '').strip()
    context = data.get('context', 'student')
    page = data.get('page', '')
    page_section = data.get('pageSection', '')
    if not message:
        return jsonify({'error': '请输入消息'}), 400
    club_name = user.get('club_name', '') if user else ''
    role = user.get('role', context) if user else context
    username = user.get('username', '') if user else ''

    conn = db.get_conn()
    try:
        conn.execute('INSERT INTO ai_chat_history (user_id, role, content) VALUES (?, ?, ?)', (user['id'], 'user', message))
        conn.commit()
    finally:
        conn.close()

    role_label = "管理员" if role == "admin" else "社团负责人" if role == "user" else "指导老师" if role == "teacher" else "学生"

    # 获取学生加入的社团列表
    user_clubs = []
    if role == 'student':
        try:
            conn = db.get_conn()
            club_rows = conn.execute('SELECT DISTINCT club_name FROM club_members WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_cadres WHERE user_id=?', (user['id'], user['id'])).fetchall()
            user_clubs = [r['club_name'] for r in club_rows if r['club_name']]
            conn.close()
        except:
            pass

    # 获取指导老师指导的社团
    teacher_clubs = []
    if role == 'teacher':
        try:
            conn = db.get_conn()
            tc_rows = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
            teacher_clubs = [r['club_name'] for r in tc_rows if r['club_name']]
            conn.close()
        except:
            pass

    # 社团负责人如果 club_name 为空，尝试从 club_cadres/club_members 查找
    if role == 'user' and not club_name:
        try:
            conn = db.get_conn()
            cr = conn.execute('SELECT DISTINCT club_name FROM club_cadres WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_members WHERE user_id=?', (user['id'], user['id'])).fetchall()
            if cr:
                club_name = cr[0]['club_name']
                user['club_name'] = club_name
            conn.close()
        except:
            pass

    # 页面上下文映射
    page_context_map = {
        'dashboard.html': {'name': '管理中心', 'desc': '包含社团管理、活动开展、招募管理、财务登记、工作量统计、树洞等功能'},
        'club-tools.html': {'name': '社团工具', 'desc': '包含活动开展、招募管理、成员管理、财务登记、工作量统计等功能'},
        'club-teacher.html': {'name': '指导老师管理', 'desc': '包含活动签到签退、赋分审核、工作量统计、招募管理、指导情况等功能'},
        'index.html': {'name': '管理后台', 'desc': '包含社团管理、用户管理、材料审批、通知管理、热点资讯等功能'},
        'upload.html': {'name': '材料上传', 'desc': '材料上传与审批管理'},
        'workload.html': {'name': '工作量审核', 'desc': '工作量提交与审核'},
        'feedback.html': {'name': '问题反馈', 'desc': '问题反馈与处理'},
        'stats.html': {'name': '数据统计', 'desc': '全校数据统计分析'},
        'checkin.html': {'name': '签到管理', 'desc': '活动签到签退管理'},
        'club-detail.html': {'name': '社团详情', 'desc': '社团信息展示页面'}
    }
    page_info = page_context_map.get(page, {})
    page_name = page_info.get('name', '')
    page_desc = page_info.get('desc', '')
    # 子区域上下文
    section_context_map = {
        'home': '首页概览', 'notif': '消息通知', 'enroll': '报名审批', 'recruit': '招募管理',
        'member': '成员管理', 'checkin': '活动签到', 'finance': '财务登记', 'scoring': '赋分管理',
        'workload': '工作量统计', 'treehole': '树洞', 'notices': '管理员通知', 'showcase': '社团风采',
        'offcampus': '校外活动', 'joint': '联合活动', 'quitManage': '退社管理', 'scoringRules': '赋分规则',
        'feedback': '问题反馈', 'profile': '社团资料', 'tools': '协助工具'
    }
    if page_section and page_section in section_context_map:
        page_desc = f'当前在{section_context_map[page_section]}页面。{page_desc}'

    system_prompt = f'''你是"云智社联管理系统"的AI智能助手，名叫小通，是本系统的"智慧大脑"和"智能管家"。你可以调取所有相关数据，提供精准的信息查询、预测建议和执行操作。

## 你的身份
- 你是云智社联管理系统的AI助手，名叫小通，是系统的"智慧大脑"和"智能管家"
- 你拥有多种工具能力，可以查询真实数据、创建实体、更新状态、发送通知、生成报告、生成文案、个性化推荐
- 当用户询问社团数据、审批情况等问题时，请主动使用工具获取真实数据
- 你也是一个知识渊博的通用AI助手，可以回答任何领域的问题

## 当前用户信息
- 用户名：{username}
- 角色：{role_label}
- 所属社团：{club_name or "未关联社团"}
{"- 加入的社团：" + "、".join(user_clubs) if user_clubs else ""}
{"- 指导的社团：" + "、".join(teacher_clubs) if teacher_clubs else ""}
{"- 当前页面：" + page_name + " — " + page_desc if page_name else ""}

## 工具使用指南
### 数据查询类
- 用户问审批/待审核/待审批 → 使用 check_approval 工具
- 用户问找文件/搜索文件/找资料 → 使用 search_documents 工具
- 用户问预警/异常/超时/问题 → 使用 check_alerts 工具
- 用户问分析/报告/数据洞察 → 使用 generate_report 工具（比 generate_insight_report 更全面）
- 用户问社团数据/成员/活动/签到/财务 → 使用 query_data 工具（比 query_club_data 更全面）
- 用户问有哪些社团/社团列表 → 使用 list_clubs 或 query_data(dataset="clubs") 工具
- query_data 支持的数据集：activities(活动)、members(成员)、students(学生，支持has_club过滤)、teachers(指导老师，支持has_club过滤)、workload(工作量)、finance(财务)、recruitments(招募)、scoring(赋分)、credits(学分)、checkins(签到)、treehole(树洞)、feedback(反馈)、excellent_clubs(优秀社团)、excellent_activities(优秀活动)、hot_news(热点资讯)、documents(资料库)、clubs(社团)、notifications(通知)、departments(部门)、votes(投票)、surveys(问卷)、joint_activities(联合活动)、offcampus_requests(校外活动申请)
- 查询未加入社团的学生：query_data(dataset="students", filters={{"has_club":"false"}})
- 查询未指导社团的老师：query_data(dataset="teachers", filters={{"has_club":"false"}})
- 复杂查询需求（query_data无法满足时）→ 使用 query_database 工具直接执行SQL SELECT查询
- query_database 示例：query_database(sql="SELECT club_name, COUNT(*) as cnt FROM club_members GROUP BY club_name ORDER BY cnt DESC", description="查询各社团人数")

### 创建操作类
- 用户想创建活动/办活动 → 使用 create_entity(entity_type="activity") 工具
- 用户想发布招募/纳新 → 使用 create_entity(entity_type="recruitment") 工具
- 用户想发通知/公告 → 使用 create_entity(entity_type="notification") 或 send_notification 工具
- 用户想发起投票 → 使用 create_entity(entity_type="vote") 工具

### 更新操作类
- 用户想审批/通过/驳回 → 使用 update_entity 工具（需确认后再执行）
- 用户想修改活动状态 → 使用 update_entity(entity_type="activity") 工具

### 通知发送类
- 用户想通知/提醒某人 → 使用 send_notification 工具

### 报告生成类
- 用户要活动概览 → generate_report(report_type="activity_overview")
- 用户要指导情况 → generate_report(report_type="teacher_guidance")
- 用户要工作量统计 → generate_report(report_type="workload_stats")
- 用户要财务汇总 → generate_report(report_type="finance_summary")
- 用户要赋分分布 → generate_report(report_type="scoring_distribution")
- 用户要预警信息 → generate_report(report_type="alert_summary")
- 用户要社团健康度 → generate_report(report_type="club_health")

### AI生成类
- 用户要生成文案/宣传语 → ai_generate(generate_type="copywriting")

### 推荐类
- 学生问推荐社团 → recommend(recommend_type="club")
- 学生问推荐活动 → recommend(recommend_type="activity")
- 学生问推荐招募/想报名招募 → recommend(recommend_type="recruitment")
- 负责人问联合活动伙伴 → recommend(recommend_type="joint_partner")

### 可视化看板类
- 用户想看图表/看板/可视化/数据可视化 → 使用 generate_dashboard 工具
- 用户想看成员活跃度/活跃度排行 → generate_dashboard(dashboard_type="member_activity")
- 用户想看工作量分布/工作量图表 → generate_dashboard(dashboard_type="workload_distribution")
- 用户想看财务图表/收支可视化 → generate_dashboard(dashboard_type="finance_overview")
- 用户想看签到考勤/考勤图表 → generate_dashboard(dashboard_type="attendance")
- 用户想看赋分分布/赋分图表 → generate_dashboard(dashboard_type="scoring_distribution")
- 用户想看社团概览/社团对比图 → generate_dashboard(dashboard_type="club_overview")
- 用户想看活动趋势/趋势图 → generate_dashboard(dashboard_type="activity_trend")
- 看板会以实时图表形式在聊天中展示，用户可以直接看到可视化结果

## 角色权限规则
### 管理员（admin）
- 可查看全校所有社团数据、审批所有类型、管理用户、配置首页内容
- 可审批：材料审批、校外活动审批、赋分审核
- 可配置：轮播图、优秀活动、热点资讯、赋分规则
- 可查看预警中心，主动推送预警

### 社团负责人（user）
- 只能操作自己社团的数据
- 可创建：活动、招募、投票、问卷、通知
- 可审批：报名审批、财务权限审批
- 可管理：成员增删、部门设置、骨干设置
- 可查看：社团数据、工作量、财务、签到记录
- 可生成文案

### 指导老师（teacher）
- 只能查看所指导社团的数据
- 可查看：活动详情、签到签退记录（含GPS）、成员活跃度
- 可审批：赋分审核、退社审批
- 可查看：工作量统计、招募进展
- 可给出指导建议

### 学生（student）
- 只能查看自己加入的社团数据
- 可操作：活动签到、活动报名、查看自己的工作量和学分
- 可参与：投票、问卷
- 可使用树洞匿名社区
- 可获取个性化推荐（社团、活动、招募）
- 每人最多加入2个社团，推荐社团时需排除已加入的社团并提示剩余名额

## 工作原则
1. **主动预测**：检测到异常或机会时主动推送建议
2. **模糊理解**：将口语转化为精确操作，如"我想办个篮球赛" → 自动调用 create_entity
3. **多步确认**：涉及资源变更或审批的操作，先给出计划，征得同意后再执行
4. **可解释性**：每个推荐或决策附带依据
5. **不要编造数据**：所有社团相关数据必须通过工具获取
6. **数据隔离**：严格遵守角色权限，非管理员只能查看自己所属社团的数据，系统会自动过滤

## 消息格式说明
- 用户消息会自动带上身份标签，格式为 `[身份：角色，社团：xxx，页面：xxx] 实际消息内容`
- 你应根据身份标签中的角色和社团信息，自动限制数据查询和操作范围
- 当需要查询数据或执行操作时，请使用提供的工具（function calling），系统会自动执行并将结果返回给你
- 你可以连续调用多个工具（最多3轮），例如先查询数据，再根据结果调用另一个工具
- 如果工具调用结果不足以回答问题，可以继续调用更多工具

## 回答原则
- 社团相关问题：优先使用工具获取真实数据，结合数据给出专业准确的回答
- 通用知识问题：给出详细、准确、有深度的回答
- 回答简洁明了，重点突出
- 用中文回答'''

    if club_name:
        try:
            conn = db.get_conn()
            member_count = conn.execute('SELECT COUNT(*) as c FROM users WHERE club_name=? AND role="student"', (club_name,)).fetchone()['c']
            activity_count = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (club_name,)).fetchone()['c']
            notice_count = conn.execute('SELECT COUNT(*) as c FROM club_notices WHERE club_name=?', (club_name,)).fetchone()['c']
            dept_count = conn.execute('SELECT COUNT(*) as c FROM club_departments WHERE club_name=?', (club_name,)).fetchone()['c']
            recruit_count = conn.execute('SELECT COUNT(*) as c FROM recruitments WHERE club_name=? AND status="approved"', (club_name,)).fetchone()['c']
            # 额外数据
            pending_workload = conn.execute('SELECT COUNT(*) as c FROM workload_submissions WHERE club_name=? AND status="pending"', (club_name,)).fetchone()['c']
            pending_scoring = conn.execute('SELECT COUNT(*) as c FROM scoring_submissions WHERE club_name=? AND status="pending"', (club_name,)).fetchone()['c']
            income = conn.execute("SELECT COALESCE(SUM(amount),0) as s FROM finance_records WHERE club_name=? AND type='income'", (club_name,)).fetchone()['s']
            expense = conn.execute("SELECT COALESCE(SUM(amount),0) as s FROM finance_records WHERE club_name=? AND type='expense'", (club_name,)).fetchone()['s']
            conn.close()
            system_prompt += f'\n\n## {club_name} 实时数据\n- 成员数：{member_count}人\n- 活动数：{activity_count}次\n- 公告数：{notice_count}条\n- 部门数：{dept_count}个\n- 招募数：{recruit_count}个\n- 待审工作量：{pending_workload}项\n- 待审赋分：{pending_scoring}项\n- 财务余额：¥{income - expense}（收入¥{income} 支出¥{expense}）'
        except:
            pass
    elif role == 'admin':
        try:
            conn = db.get_conn()
            total_students = conn.execute('SELECT COUNT(*) as c FROM users WHERE role="student"').fetchone()['c']
            total_clubs = conn.execute('SELECT COUNT(DISTINCT club_name) as c FROM users WHERE role="user" AND club_name!=""').fetchone()['c']
            total_activities = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions').fetchone()['c']
            pending_materials = conn.execute('SELECT COUNT(*) as c FROM club_uploads WHERE status="pending"').fetchone()['c']
            pending_offcampus = conn.execute('SELECT COUNT(*) as c FROM offcampus_requests WHERE status="pending"').fetchone()['c']
            conn.close()
            system_prompt += f'\n\n## 全校数据\n- 学生总数：{total_students}人\n- 社团数：{total_clubs}个\n- 活动总数：{total_activities}次\n- 待审材料：{pending_materials}项\n- 待审校外活动：{pending_offcampus}项'
        except:
            pass
    elif role == 'teacher' and teacher_clubs:
        try:
            conn = db.get_conn()
            club_data_parts = []
            for tc in teacher_clubs[:5]:
                ac = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (tc,)).fetchone()['c']
                mc = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE club_name=?', (tc,)).fetchone()['c']
                club_data_parts.append(f'{tc}：{mc}人/{ac}次活动')
            conn.close()
            system_prompt += f'\n\n## 指导社团数据\n- ' + '\n- '.join(club_data_parts)
        except:
            pass
    elif role == 'student' and user_clubs:
        try:
            conn = db.get_conn()
            club_data_parts = []
            for uc in user_clubs[:5]:
                ac = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (uc,)).fetchone()['c']
                mc = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE club_name=?', (uc,)).fetchone()['c']
                club_data_parts.append(f'{uc}：{mc}人/{ac}次活动')
            conn.close()
            system_prompt += f'\n\n## 加入社团数据\n- ' + '\n- '.join(club_data_parts)
        except:
            pass

    qwen_messages = [{'role': 'system', 'content': system_prompt}]

    # 构造身份标签
    identity_tag = f'[身份：{role_label}'
    if club_name:
        identity_tag += f'，社团：{club_name}'
    elif user_clubs:
        identity_tag += f'，社团：{"、".join(user_clubs[:3])}'
    elif teacher_clubs:
        identity_tag += f'，指导：{"、".join(teacher_clubs[:3])}'
    if page_name:
        identity_tag += f'，页面：{page_name}'
    identity_tag += ']'

    try:
        conn = db.get_conn()
        history = conn.execute('SELECT role, content FROM ai_chat_history WHERE user_id=? ORDER BY created_at DESC LIMIT 10', (user['id'],)).fetchall()
        conn.close()
        for h in reversed(history[:-1]):
            qwen_messages.append({'role': h['role'], 'content': h['content']})
    except:
        pass

    # 用户消息前自动加上身份标签
    tagged_message = f'{identity_tag} {message}'
    qwen_messages.append({'role': 'user', 'content': tagged_message})

    tool_calls_made = []
    reply = None

    provider, api_key = get_llm_config()
    if api_key:
        api_response = call_llm_api(qwen_messages, tools=AI_TOOLS)
        if api_response:
            choice = api_response.get('choices', [{}])[0]
            assistant_msg = choice.get('message', {})
            tool_calls = assistant_msg.get('tool_calls')

            if tool_calls:
                # 支持多轮工具调用，最多3轮
                max_rounds = 3
                for round_idx in range(max_rounds):
                    qwen_messages.append(assistant_msg)
                    for tc in tool_calls:
                        fn = tc.get('function', {})
                        fn_name = fn.get('name', '')
                        fn_args_str = fn.get('arguments', '{}')
                        try:
                            fn_args = json.loads(fn_args_str) if isinstance(fn_args_str, str) else fn_args_str
                        except json.JSONDecodeError:
                            fn_args = {}
                        tool_result = execute_tool_call(fn_name, fn_args, current_user=user)
                        if fn_name == 'generate_dashboard':
                            try:
                                dashboard_data = json.loads(tool_result)
                                if dashboard_data.get('type') == 'dashboard':
                                    tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': '已生成可视化看板', 'dashboard': dashboard_data})
                                else:
                                    tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': tool_result[:200]})
                            except:
                                tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': tool_result[:200]})
                        else:
                            tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': tool_result[:200]})
                        qwen_messages.append({
                            'role': 'tool',
                            'tool_call_id': tc.get('id', ''),
                            'content': tool_result
                        })
                    # 把工具结果返回给AI，让AI继续对话或调用更多工具
                    next_response = call_llm_api(qwen_messages, tools=AI_TOOLS, max_tokens=1000)
                    if next_response and isinstance(next_response, dict):
                        next_choice = next_response.get('choices', [{}])[0]
                        next_msg = next_choice.get('message', {})
                        next_tool_calls = next_msg.get('tool_calls')
                        if next_tool_calls:
                            # AI还想继续调用工具
                            assistant_msg = next_msg
                            tool_calls = next_tool_calls
                            continue
                        else:
                            # AI给出了最终回复
                            reply = next_msg.get('content', '')
                            break
                    elif next_response and isinstance(next_response, str):
                        reply = next_response
                        break
                    else:
                        break

                if not reply:
                    # 多轮工具调用后仍无回复，做一次不带工具的调用
                    final_response = call_llm_api(qwen_messages, max_tokens=1000)
                    if final_response:
                        reply = final_response
                if not reply:
                    reply = '我已查询了相关数据，但生成回复时遇到了问题，请稍后再试。'
            else:
                reply = assistant_msg.get('content', '')

    if not reply:
        reply = call_llm_api(qwen_messages)
    if not reply:
        intents = smart_intent_detect(message, user)
        if intents:
            parts = []
            all_suggestions = []
            for intent in intents:
                fn_name = intent['tool']
                fn_args = intent['args']
                tool_result = execute_tool_call(fn_name, fn_args, current_user=user)
                if fn_name == 'generate_dashboard':
                    try:
                        dashboard_data = json.loads(tool_result)
                        if dashboard_data.get('type') == 'dashboard':
                            tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': '已生成可视化看板', 'dashboard': dashboard_data})
                        else:
                            tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': tool_result[:200]})
                    except:
                        tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': tool_result[:200]})
                else:
                    tool_calls_made.append({'name': fn_name, 'arguments': fn_args, 'result_preview': tool_result[:200]})
                human_result = format_tool_result_for_human(fn_name, fn_args, tool_result)
                parts.append(human_result)
                try:
                    result_data = json.loads(tool_result) if isinstance(tool_result, str) else tool_result
                except:
                    result_data = {}
                sug = generate_smart_suggestions(fn_name, result_data, user)
                all_suggestions.extend(sug)
            reply = '\n\n---\n\n'.join(parts)
            if all_suggestions:
                unique = list(dict.fromkeys(all_suggestions))[:3]
                reply += '\n\n💡 **你可能还想了解：** ' + ' | '.join(unique)
        else:
            reply = generate_ai_reply(message, user, context)

    try:
        conn = db.get_conn()
        import json as _json_mod
        _tc_json = _json_mod.dumps(tool_calls_made, ensure_ascii=False, default=str) if tool_calls_made else ''
        conn.execute('INSERT INTO ai_chat_history (user_id, role, content, tool_calls) VALUES (?, ?, ?, ?)', (user['id'], 'assistant', reply, _tc_json))
        conn.commit()
        conn.close()
    except:
        pass

    add_pet_exp(user['id'], 3)
    response_data = {'success': True, 'reply': reply}
    if tool_calls_made:
        response_data['tool_calls'] = tool_calls_made
    return jsonify(response_data)


@app.route('/api/chat-history', methods=['GET'])
def chat_history():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT role, content, created_at, tool_calls FROM ai_chat_history WHERE user_id=? ORDER BY created_at ASC LIMIT 100', (user['id'],)).fetchall()
        import json as _json_mod2
        history = []
        for r in rows:
            item = {'role': r['role'], 'content': r['content'], 'created_at': r['created_at']}
            tc_str = r['tool_calls'] if 'tool_calls' in r.keys() else ''
            if tc_str:
                try:
                    item['tool_calls'] = _json_mod2.loads(tc_str)
                except Exception:
                    item['tool_calls'] = None
            else:
                item['tool_calls'] = None
            history.append(item)
        return jsonify({'success': True, 'history': history})
    finally:
        conn.close()


@app.route('/api/clear-chat', methods=['POST'])
def clear_chat():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM ai_chat_history WHERE user_id=?', (user['id'],))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@app.route('/api/llm-providers', methods=['GET'])
def get_llm_providers():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    # 只展示聊天模型（排除视觉模型）
    chat_providers = {k: v for k, v in LLM_PROVIDERS.items() if k != 'qwen-vl'}
    providers = []
    for key, cfg in chat_providers.items():
        key_val = os.environ.get(cfg['key_env'], '')
        if not key_val:
            if cfg['key_env'] == 'ZHIPU_API_KEY': key_val = ZHIPU_API_KEY
        has_key = bool(key_val)
        providers.append({
            'id': key,
            'name': cfg['name'],
            'icon': cfg['icon'],
            'desc': cfg['desc'],
            'models': cfg['models'],
            'default_model': cfg['default_model'],
            'has_key': has_key,
            'is_active': key == ACTIVE_LLM
        })
    return jsonify({'success': True, 'providers': providers, 'active': ACTIVE_LLM})


@app.route('/api/llm-switch', methods=['POST'])
def switch_llm_provider():
    global ACTIVE_LLM
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    provider_id = data.get('provider', '').strip()
    if provider_id not in LLM_PROVIDERS:
        return jsonify({'error': '不支持的模型提供商'}), 400
    cfg = LLM_PROVIDERS[provider_id]
    key_val = os.environ.get(cfg['key_env'], '')
    if not key_val:
        if cfg['key_env'] == 'ZHIPU_API_KEY': key_val = ZHIPU_API_KEY
    if not key_val:
        return jsonify({'error': f'未配置 {cfg["name"]} 的 API Key（环境变量 {cfg["key_env"]}）'}), 400
    ACTIVE_LLM = provider_id
    return jsonify({'success': True, 'active': provider_id, 'name': cfg['name']})


def generate_ai_reply(message, user, context):
    import json as _json, random, re
    from datetime import datetime
    msg = message.lower().strip()
    original = message.strip()
    club_name = user.get('club_name', '') if user else ''
    username = user.get('username', '') if user else ''
    role = user.get('role', context) if user else context

    def has_any(text, keywords):
        return any(kw in text for kw in keywords)

    def has_pattern(text, patterns):
        return any(re.search(p, text) for p in patterns)

    def get_club_stats(club):
        try:
            conn = db.get_conn()
            mc = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE club_name=?', (club,)).fetchone()['c']
            ac = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions WHERE club_name=?', (club,)).fetchone()['c']
            cc = conn.execute('SELECT COUNT(*) as c FROM checkin_records WHERE club_name=?', (club,)).fetchone()['c']
            dc = conn.execute('SELECT COUNT(*) as c FROM club_departments WHERE club_name=?', (club,)).fetchone()['c']
            nc = conn.execute('SELECT COUNT(*) as c FROM club_notices WHERE club_name=?', (club,)).fetchone()['c']
            rc = conn.execute('SELECT COUNT(*) as c FROM recruitments WHERE club_name=? AND status="approved"', (club,)).fetchone()['c']
            recent = conn.execute('SELECT activity_name, created_at, status FROM checkin_sessions WHERE club_name=? ORDER BY created_at DESC LIMIT 3', (club,)).fetchall()
            conn.close()
            return {'members': mc, 'activities': ac, 'checkins': cc, 'depts': dc, 'notices': nc, 'recruits': rc, 'recent': recent}
        except:
            return None

    if has_any(msg, ['你好', '嗨', 'hello', 'hi', '在吗', '嘿', '哈喽', '早上好', '下午好', '晚上好']):
        greetings = [
            f'你好{username}！😊 我是小通，你的社团AI助手。有什么可以帮你的吗？',
            f'嗨！👋 很高兴见到你！今天想了解什么？',
            f'你好呀！🌟 我随时准备为你服务！可以问我任何关于社团的问题~',
            f'Hi！✨ 有什么需要帮忙的尽管说！'
        ]
        return random.choice(greetings)

    if has_any(msg, ['谢谢', '感谢', '多谢', '太好了', '厉害', '棒']):
        replies = [
            '不客气！😊 随时都可以找我~',
            '很高兴能帮到你！还有其他问题吗？',
            '客气啦！有问题随时问我就好 🌟',
        ]
        return random.choice(replies)

    if has_any(msg, ['你是谁', '你叫什么', '介绍自己', '你是什么', '自我介绍']):
        role = user.get('role', 'student') if user else 'student'
        role_label = {'admin': '管理员', 'user': '社团负责人', 'teacher': '指导老师', 'student': '学生'}.get(role, '')
        role_features = {
            'admin': '📊 全校数据分析、📋 材料审批、⚠️ 预警监控、📑 报告生成',
            'user': '✨ 创建活动/招募/投票、📋 报名审批、💰 财务管理、🎨 海报文案生成',
            'teacher': '📋 赋分审核、📊 指导情况分析、👥 成员活跃度分析',
            'student': '💡 社团/活动推荐、📊 工作量学分查询、🎯 活动报名签到'
        }
        return f'''🤖 我是小通，云智社联管理系统的"智慧大脑"！

作为{role_label}，你可以通过我：
{role_features.get(role, '📊 数据查询、✨ 创建操作、📢 发送通知、📑 报告生成、🎨 AI创作、💡 智能推荐')}

🔧 **核心能力**：
- 📊 **数据查询** — 查询活动、成员、财务、签到等所有数据
- ✨ **创建操作** — 创建活动、招募、投票、通知等
- 🔄 **状态更新** — 审批通过/驳回、修改活动状态等
- 📢 **发送通知** — 给成员、负责人、老师发通知
- 📑 **报告生成** — 活动概览、工作量统计、财务汇总等
- 🎨 **AI生成** — 海报图片、活动文案
- 💡 **智能推荐** — 社团推荐、活动推荐、联合伙伴推荐

试试问我点什么吧！'''

    if has_any(msg, ['帮助', '能做什么', '功能', '怎么用', '使用方法', '帮我']):
        return '''🤖 我是小通，系统的"智慧大脑"，可以帮你做这些：

📊 **数据查询**：
• "查看成员数据" — 查看社团成员概况
• "活动情况" — 查看活动签到统计
• "财务数据" — 查看收支情况
• "查询签到记录" — 查看签到详情

✨ **创建操作**：
• "帮我创建一个活动" — 创建新活动
• "发布招募" — 发起招募
• "发起投票" — 创建投票
• "发通知" — 给成员发通知

🔄 **审批操作**：
• "查看待审批" — 查看待审批项目
• "帮我审批" — 审批通过或驳回

📑 **报告生成**：
• "生成活动报告" — 活动概览统计
• "生成财务报告" — 财务汇总
• "生成工作量报告" — 工作量统计

🎨 **AI创作**：
• "写文案" — 生成宣传文案

💡 **智能推荐**：
• "推荐社团" — 推荐适合的社团
• "推荐活动" — 推荐感兴趣的活动
• "推荐招募" — 推荐可报名的招募活动
• "推荐联合伙伴" — 推荐合作社团

💬 也可以直接问我任何问题！'''

    if has_any(msg, ['时间', '几点', '日期', '今天', '现在']):
        now = datetime.now()
        weekdays = ['一','二','三','四','五','六','日']
        return f'现在是 {now.strftime("%Y年%m月%d日")} 星期{weekdays[now.weekday()]} {now.strftime("%H:%M")} 🕐'

    if has_any(msg, ['笑话', '搞笑', '开心', '逗我', '段子']):
        jokes = [
            '为什么程序员总是分不清万圣节和圣诞节？\n因为 Oct 31 = Dec 25 😄',
            '数学书为什么总是不开心？\n因为它有太多问题了 😂',
            '为什么电脑不会感冒？\n因为它有Windows但不会开 😆',
            '一只蜗牛爬上苹果树，树上的毛毛虫说：你也太慢了吧！蜗牛说：慢？我已经比昨天快了一倍了，我可是坐电梯上来的！🐌',
            '老师问小明：如果你有12块糖，我问你要3块，你还剩几块？小明：12块。老师：你不会算术吗？小明：你不会觉得我会给你吧？🍬',
            '社团招新时学长说：加入我们社团，保证你大学生活丰富多彩！学弟问：具体有什么活动？学长想了想：开会。😅',
        ]
        return random.choice(jokes)

    if has_any(msg, ['纳新文案', '招新文案', '纳新宣传', '招新宣传', '宣传文案']):
        club = club_name or '社团'
        return f'''🌟 {club} 纳新啦！🌟

亲爱的同学们，你是否渴望在大学里找到志同道合的伙伴？

🎓 {club}诚挚邀请你的加入！

✨ 我们是谁？
{club}是一个致力于多元发展与成长的校级优秀社团。在这里，我们用热情点燃梦想，用行动诠释青春！

🎯 加入我们，你将获得：
• 专业的技能培训与指导
• 丰富的校内外活动机会
• 志同道合的伙伴与友谊
• 展示自我、提升能力的舞台

📋 招新对象：全校在校生
📍 报名方式：填写线上报名表
⏰ 报名截止：待定

💫 每一个伟大的故事，都始于一次勇敢的选择。
加入{club}，让我们一起书写属于你的精彩篇章！'''

    if has_any(msg, ['报名表', '纳新表']):
        club = club_name or '社团'
        return f'''{club} 纳新报名表

━━━━━━━━━━━━━━━━━━━━
基本信息
━━━━━━━━━━━━━━━━━━━━
姓    名：_______________
学    号：_______________
性    别：□ 男  □ 女
联系电话：_______________
微 信 号：_______________
学    院：_______________
专    业：_______________
年    级：_______________

━━━━━━━━━━━━━━━━━━━━
个人意向
━━━━━━━━━━━━━━━━━━━━
是否服从调剂：□ 是  □ 否
个人特长/爱好：_________________________________

报名日期：_______________'''

    if has_any(msg, ['成员', '社员', '人数', '成员分析', '成员数据', '有多少人', '多少人']):
        if role == 'admin':
            try:
                conn = db.get_conn()
                total = conn.execute('SELECT COUNT(*) as c FROM users WHERE role="student"').fetchone()['c']
                by_club = conn.execute('SELECT club_name, COUNT(*) as c FROM club_members WHERE club_name!="" GROUP BY club_name ORDER BY c DESC LIMIT 10').fetchall()
                no_club = conn.execute('SELECT COUNT(*) as c FROM users WHERE role="student" AND id NOT IN (SELECT DISTINCT user_id FROM club_members)').fetchone()['c']
                conn.close()
                reply = f'📊 全校成员概况：\n\n注册学生总数：{total}人\n已加入社团：{total - no_club}人\n未加入社团：{no_club}人\n\n各社团人数：\n'
                for r in by_club:
                    reply += f'• {r["club_name"]}：{r["c"]}人\n'
                return reply
            except:
                return '抱歉，获取成员数据时出错。'
        elif club_name:
            stats = get_club_stats(club_name)
            if stats:
                reply = f'📊 {club_name} 成员概况：\n\n成员总数：{stats["members"]}人\n部门数：{stats["depts"]}个\n活动数：{stats["activities"]}次\n签到人次：{stats["checkins"]}次\n'
                if stats['depts'] > 0:
                    try:
                        conn = db.get_conn()
                        depts = conn.execute('SELECT dept_name, description FROM club_departments WHERE club_name=?', (club_name,)).fetchall()
                        conn.close()
                        reply += f'\n部门列表：\n'
                        for d in depts:
                            reply += f'• {d["dept_name"]}'
                            if d['description']:
                                reply += f' - {d["description"]}'
                            reply += '\n'
                    except:
                        pass
                return reply
            return '获取数据时出错，请稍后再试。'
        else:
            return '你还没有关联社团，无法查看成员数据。先去社团页面加入一个社团吧！'

    if has_any(msg, ['活动', '签到', '活动概况', '活动情况', '活动数据']):
        if club_name:
            stats = get_club_stats(club_name)
            if stats:
                reply = f'📊 {club_name} 活动概况：\n\n活动总数：{stats["activities"]}次\n签到总人次：{stats["checkins"]}次\n'
                if stats['recent']:
                    reply += '\n最近活动：\n'
                    for r in stats['recent']:
                        status = '🟢 进行中' if r['status'] == 'open' else '🔴 已结束'
                        reply += f'• {r["activity_name"] or "未命名"} {status} - {r["created_at"]}\n'
                return reply
        elif role == 'admin':
            try:
                conn = db.get_conn()
                ts = conn.execute('SELECT COUNT(*) as c FROM checkin_sessions').fetchone()['c']
                tc = conn.execute('SELECT COUNT(*) as c FROM checkin_records').fetchone()['c']
                tcl = conn.execute('SELECT COUNT(DISTINCT club_name) as c FROM checkin_sessions').fetchone()['c']
                conn.close()
                return f'📊 全校活动概况：\n\n开展活动的社团：{tcl}个\n活动总数：{ts}次\n签到总人次：{tc}次'
            except:
                pass
        return '你还没有关联社团。'

    if has_any(msg, ['活动策划', '活动方案', '策划活动', '策划方案', '怎么策划', '办什么活动']):
        club = club_name or '社团'
        return f'''🎯 {club} 活动策划建议：

要生成详细活动方案，请告诉我活动类型：
• 团建类 - 增进成员感情
• 文艺类 - 展示才艺风采
• 公益类 - 服务社会大众
• 学术类 - 知识交流碰撞
• 竞赛类 - 激发竞争热情

你也可以在"AI工具"面板中使用"活动策划"功能，生成完整的活动方案！'''

    if has_any(msg, ['公告润色', '优化公告', '润色通知', '美化公告', '写公告', '写通知']):
        return '''📝 公告润色助手：

我可以帮你优化公告文案！请提供原始公告内容，并选择风格：
• 正式(formal) - 规范专业的公告格式
• 温馨(warm) - 亲切温暖的表达
• 紧急(urgent) - 突出重要性和时效性

你可以在"AI工具"面板中使用"公告润色"功能哦！'''

    if has_any(msg, ['社团健康', '社团评估', '社团评分', '社团状态', '社团怎么样']):
        club = club_name or ''
        if not club:
            return '你还没有关联社团。请先在社团管理页面关联你的社团。'
        stats = get_club_stats(club)
        if stats:
            scores = {
                '成员规模': min(100, stats['members'] * 5),
                '活动频率': min(100, stats['activities'] * 15),
                '签到参与': min(100, stats['checkins'] * 3),
                '组织架构': min(100, stats['depts'] * 25),
                '公告活跃': min(100, stats['notices'] * 20),
                '招募开拓': min(100, stats['recruits'] * 20),
            }
            avg = round(sum(scores.values()) / len(scores))
            level = '🌟 优秀' if avg >= 70 else ('👍 良好' if avg >= 50 else ('📈 发展中' if avg >= 30 else '🌱 起步期'))
            reply = f'🏥 {club} 健康度评估：\n\n综合评分：{avg}分 ({level})\n\n'
            for k, v in scores.items():
                bar = '█' * (v // 10) + '░' * (10 - v // 10)
                reply += f'{k}：{bar} {v}分\n'
            reply += f'\n📊 统计：成员{stats["members"]}人 | 活动{stats["activities"]}次 | 签到{stats["checkins"]}人次 | 部门{stats["depts"]}个 | 公告{stats["notices"]}条'
            return reply
        return '获取社团数据时出错，请稍后再试。'

    if has_any(msg, ['学期规划', '学期计划', '年度规划', '工作计划', '下学期']):
        club = club_name or '社团'
        return f'''📅 {club} 学期规划建议：

### 🗓 月度活动安排
| 月份 | 主题活动 | 重点 |
|------|---------|------|
| 第1月 | 🎯 纳新宣传 | 扩大影响力 |
| 第2月 | 🤝 迎新见面 | 凝聚新成员 |
| 第3月 | 🎨 特色活动 | 展现社团特色 |
| 第4月 | 🏆 成果展示 | 总结与表彰 |

在"AI工具"面板中可生成更详细的学期规划方案！'''

    if has_any(msg, ['活动总结', '总结报告', '写总结']):
        return '''📋 活动总结助手：

请告诉我以下信息，我帮你生成专业的活动总结：
1. 活动名称
2. 活动亮点
3. 参与人数

或者在"AI工具"面板中使用"活动总结"功能！'''

    if has_any(msg, ['招募优化', '优化招募', '招募文案', '招募描述']):
        return '''✨ 招募文案优化助手：

把你的招募描述发给我，我帮你优化得更加吸引人！
优化方向包括：
• 突出活动亮点和收获
• 优化排版和格式
• 增强语言感染力
• 明确参与条件和要求

在"AI工具"面板中也可以使用"招募优化"功能~'''

    if has_any(msg, ['宠物', '我的宠物', '宠物状态', '宠物在哪']):
        return '🐾 你的AI宠物可通过右下角的浮动按钮查看和互动！点击宠物可以摸摸它，还可以玩猜数字、石头剪刀布、知识问答、幸运抽奖、宠物冒险、掷骰子等多种游戏，喂食不同食物获得额外经验值哦~'

    if has_any(msg, ['活动创意', '社团活动', '推荐活动', '活动建议', '有什么活动']):
        ideas = [
            '🎯 团建类：密室逃脱、真人CS、户外拓展',
            '🎨 文化类：书法比赛、诗词大会、传统文化体验',
            '🎵 艺术类：才艺展示、音乐节、话剧表演',
            '🏃 运动类：趣味运动会、晨跑打卡、球类联赛',
            '🤝 公益类：志愿服务、爱心义卖、环保行动',
            '📚 学术类：读书分享、知识竞赛、学术讲座',
            '🎮 娱乐类：桌游大赛、电影之夜、电竞友谊赛',
            '🌟 创新类：创意市集、创业沙龙、技能工作坊'
        ]
        return '💡 社团活动创意推荐：\n\n' + '\n'.join(ideas) + '\n\n你可以根据社团特色选择合适的活动类型，也可以组合多种元素创造独特的活动体验！'

    if has_any(msg, ['天气', '今天天气', '下雨']):
        return '抱歉，我暂时无法获取实时天气信息 🌤️ 建议你查看天气预报APP获取最新天气情况。不过不管天气如何，社团活动都可以精彩进行！'

    if has_pattern(msg, [r'怎么.{0,4}加入', r'如何.{0,4}加入', r'怎么.{0,4}报名', r'如何.{0,4}报名', r'想加入']):
        return '''📝 加入社团的方式：

1. 在首页浏览社团列表，点击感兴趣的社团
2. 在社团详情页点击"报名加入"按钮
3. 填写个人信息后提交报名
4. 等待社团负责人审批

💡 小提示：每人最多可以加入2个社团哦！'''

    if has_pattern(msg, [r'怎么.{0,4}签到', r'如何.{0,4}签到', r'签到码', r'扫码签到']):
        return '''📋 签到方式：

1. **签到码签到**：输入社团提供的6位签到码
2. **扫码签到**：扫描社团展示的签到二维码

💡 签到后你的AI宠物会获得额外经验值哦！'''

    if has_any(msg, ['社团', '有哪些社团', '社团列表', '什么社团']):
        try:
            conn = db.get_conn()
            clubs = conn.execute('SELECT DISTINCT club_name FROM users WHERE role="user" AND club_name!="" ORDER BY club_name').fetchall()
            conn.close()
            if clubs:
                reply = '🏫 目前系统中的社团：\n\n'
                for i, c in enumerate(clubs, 1):
                    reply += f'{i}. {c["club_name"]}\n'
                reply += '\n点击首页的社团卡片可以查看详情和报名哦！'
                return reply
            return '目前还没有注册的社团。'
        except:
            return '获取社团列表时出错。'

    if has_any(msg, ['通知', '消息', '我的通知', '有消息吗']):
        return '📬 你可以在左侧菜单的"通知中心"查看所有消息通知。有新的审批、报名、签到等操作时都会收到通知哦！'

    if has_any(msg, ['密码', '改密码', '修改密码', '忘记密码']):
        return '🔐 修改密码请前往左侧菜单的"个人设置"页面。如果忘记密码，请联系管理员重置。'

    if has_any(msg, ['资料', '个人信息', '我的信息', '修改资料']):
        return '👤 修改个人信息请前往左侧菜单的"个人设置"页面，可以更新姓名、学号、班级、联系方式等。更新资料后你的AI宠物也会获得经验值哦！'

    if has_any(msg, ['退出', '登出', '注销', '退出登录']):
        return '👋 退出登录请点击左上角的退出按钮。下次再见！'

    if has_pattern(msg, [r'什么.{0,2}社团', r'社团.{0,2}什么', r'介绍.{0,2}社团']):
        if club_name:
            stats = get_club_stats(club_name)
            if stats:
                return f'''🏫 关于{club_name}：

成员数：{stats["members"]}人
活动数：{stats["activities"]}次
部门数：{stats["depts"]}个
公告数：{stats["notices"]}条

想了解更多详情，可以在首页点击社团卡片查看！'''
        return '你还没有关联社团。可以在首页浏览社团列表，选择感兴趣的社团加入！'

    if has_pattern(msg, [r'推荐', r'建议', r'应该']):
        if has_any(msg, ['活动', '办', '做']):
            return '''💡 根据当前情况，我建议：

1. **近期**：组织一次轻松的破冰活动，增进成员了解
2. **中期**：策划一次有特色的品牌活动，提升社团影响力
3. **长期**：建立完善的部门分工和培训体系

需要我帮你生成具体的活动策划方案吗？告诉我活动类型就好！'''
        return '你可以告诉我具体想了解什么方面，我来给你更有针对性的建议！'

    if has_pattern(msg, [r'怎么', r'如何', r'怎样']):
        return f'''关于"{original}"，我来帮你分析一下：

你可以尝试以下方式：
1. 在系统中查看相关数据和统计
2. 和社团成员讨论交流想法
3. 参考其他优秀社团的做法

如果你能告诉我更具体的需求，我可以给出更精准的建议！也可以试试这些功能：
• "活动策划" — 生成活动方案
• "社团评估" — 查看社团健康度
• "成员分析" — 查看成员数据'''

    if has_pattern(msg, [r'为什么', r'原因', r'怎么回事']):
        return f'''关于"{original}"这个问题：

我理解你的疑惑。可能的原因有：
1. 数据尚未更新 — 系统数据可能有延迟
2. 权限限制 — 部分功能需要特定角色才能访问
3. 操作步骤 — 可能需要先完成前置操作

如果问题持续存在，建议联系管理员查看。你也可以告诉我更多细节，我来帮你分析！'''

    if has_pattern(msg, [r'能.{0,2}吗', r'可以.{0,2}吗', r'是否']):
        return f'关于"{original}"——可以的！😊\n\n如果你需要具体操作指导，告诉我你想做什么，我来帮你一步步完成。你也可以输入"帮助"查看我的全部功能列表。'

    if has_any(msg, ['数学', '计算', '等于', '加', '减', '乘', '除', '算术', '方程', '公式']):
        try:
            expr = re.sub(r'[^\d+\-*/().%\s]', '', original)
            if expr and len(expr) > 0:
                result = eval(expr)
                return f'🧮 计算结果：{expr} = {result}\n\n如果你需要更复杂的数学问题解答，请详细描述问题，我会尽力帮你分析解题思路！'
        except:
            pass
        return f'📐 关于数学问题"{original}"：\n\n请把具体的数学问题告诉我，我可以帮你：\n• 计算算术表达式\n• 分析解题思路\n• 讲解相关知识点\n• 提供类似例题'

    if has_any(msg, ['英语', '翻译', '英文', '单词', '语法', '怎么读', '什么意思', '什么意思']):
        return f'🌐 关于"{original}"：\n\n我可以在英语学习方面帮你：\n• 基础单词和短语翻译\n• 语法知识点讲解\n• 常用表达和搭配\n• 写作技巧和建议\n\n请告诉我你想了解的具体内容！'

    if has_any(msg, ['历史', '朝代', '古代', '近代', '战争', '皇帝', '什么时候']):
        return f'📜 关于"{original}"：\n\n历史是一门充满智慧的学科！我可以帮你：\n• 梳理历史事件的时间线\n• 分析历史事件的原因和影响\n• 介绍重要的历史人物\n• 对比不同时期的特点\n\n请告诉我你想了解哪个时期或事件？'

    if has_any(msg, ['科学', '物理', '化学', '生物', '实验', '原理', '定律']):
        return f'🔬 关于"{original}"：\n\n科学探索永无止境！我可以帮你：\n• 解释科学原理和概念\n• 分析物理/化学/生物现象\n• 介绍重要科学发现\n• 提供学习方法和建议\n\n请告诉我你想了解哪个方面？'

    if has_any(msg, ['编程', '代码', '程序', 'python', 'java', '前端', '后端', '开发', 'bug', '算法']):
        return f'💻 关于"{original}"：\n\n编程是一门实践性很强的技能！我可以帮你：\n• 解释编程概念和语法\n• 分析算法思路\n• 提供代码示例和框架\n• 推荐学习资源和路线\n\n请告诉我你想了解的具体技术问题！'

    if has_any(msg, ['心理', '焦虑', '压力', '抑郁', '情绪', '心情不好', '难过', '不开心', '烦', '累']):
        comfort_replies = [
            f'💙 我理解你的感受。关于"{original}"：\n\n每个人都会有情绪低落的时候，这很正常。以下是一些小建议：\n• 🌬️ 深呼吸，给自己一点时间\n• 🚶 适当运动，释放压力\n• 👫 和信任的朋友聊聊\n• 📝 写下你的感受\n\n如果情绪持续困扰你，建议寻求专业心理咨询师的帮助。你并不孤单！',
            f'🤗 听起来你现在不太好。关于"{original}"：\n\n记住，感到困难是暂时的。试试这些方法：\n• 🎵 听一些舒缓的音乐\n• 🌿 到户外走走，接触大自然\n• 💤 保证充足的睡眠\n• 📖 做一些让自己开心的事\n\n如果需要专业帮助，学校的心理咨询中心是很好的资源。一切都会好起来的！'
        ]
        return random.choice(comfort_replies)

    if has_any(msg, ['考试', '复习', '学习', '备考', '期末', '挂科', '成绩', '学分']):
        return f'📚 关于"{original}"：\n\n学习建议：\n• 📅 制定合理的复习计划，分块攻克\n• 📝 做笔记和思维导图，梳理知识框架\n• 🔄 间隔重复法，定期回顾已学内容\n• 🎯 重点突破薄弱环节\n• 😴 保证充足睡眠，效率比时间更重要\n\n加油！相信你可以的！💪'

    if has_any(msg, ['职业', '就业', '工作', '实习', '简历', '面试', '求职']):
        return f'💼 关于"{original}"：\n\n职业发展建议：\n• 🎯 明确自己的兴趣和优势方向\n• 📄 准备一份突出亮点的简历\n• 🤝 多参加实习和社团活动积累经验\n• 📚 持续学习提升专业技能\n• 🔗 拓展人脉，参加行业交流活动\n\n需要更具体的建议可以告诉我你的专业和方向！'

    if has_any(msg, ['人际', '社交', '朋友', '室友', '矛盾', '沟通', '相处']):
        return f'🤝 关于"{original}"：\n\n人际交往建议：\n• 👂 学会倾听，理解对方的立场\n• 💬 坦诚沟通，表达自己的感受而非指责\n• 🔄 换位思考，尝试理解对方的处境\n• 🎯 聚焦问题本身，而非人身攻击\n• 💪 保持边界感，尊重自己也尊重他人\n\n良好的人际关系需要双方共同努力！'

    if has_any(msg, ['健康', '运动', '减肥', '饮食', '睡眠', '养生', '锻炼']):
        return f'🏃 关于"{original}"：\n\n健康生活建议：\n• 🥗 均衡饮食，多吃蔬果\n• 🏋️ 每周至少3次中等强度运动\n• 😴 保证7-8小时睡眠\n• 💧 每天喝够8杯水\n• 🧘 适当放松，管理压力\n\n健康是一切的基础，从今天开始行动吧！'

    if has_any(msg, ['读书', '推荐书', '书单', '看什么书', '好书']):
        book_recs = [
            '📚 推荐书单：\n\n🎯 自我提升：《原子习惯》《深度工作》《心流》\n💡 思维拓展：《思考快与慢》《原则》《黑天鹅》\n📖 文学经典：《百年孤独》《活着》《小王子》\n🔬 科普读物：《人类简史》《时间简史》《基因传》\n💼 职场发展：《高效能人士的七个习惯》《非暴力沟通》\n\n告诉我你感兴趣的领域，我可以推荐更具体的书！',
        ]
        return random.choice(book_recs)

    if has_any(msg, ['电影', '看什么', '推荐电影', '好看的电影', '剧']):
        return '🎬 推荐影视：\n\n🎭 剧情片：《肖申克的救赎》《阿甘正传》《辛德勒的名单》\n😂 喜剧片：《当幸福来敲门》《三傻大闹宝莱坞》\n🧠 烧脑片：《盗梦空间》《星际穿越》《穆赫兰道》\n💕 爱情片：《怦然心动》《时空恋旅人》\n🎬 国产佳片：《霸王别姬》《让子弹飞》《流浪地球》\n\n告诉我你喜欢的类型，我可以推荐更精准！'

    if has_any(msg, ['脑筋急转弯', '谜语', '猜谜']):
        riddles = [
            ('什么东西越洗越脏？', '水'),
            ('什么动物最容易被贴在墙上？', '海豹（海报）'),
            ('什么路最窄？', '冤家路窄'),
            ('什么人一年只工作一天？', '圣诞老人'),
            ('什么东西有头无脚？', '砖头'),
            ('什么门永远关不上？', '球门'),
        ]
        r = random.choice(riddles)
        return f'🧩 脑筋急转弯：\n\n{r[0]}\n\n🤔 想好答案了吗？\n\n答案是：{r[1]} 😄\n\n还想玩更多？继续问我"脑筋急转弯"吧！'

    if has_any(msg, ['故事', '讲故事', '听故事']):
        stories = [
            '📖 小故事：\n\n一只小蜗牛问妈妈：为什么我们生来就要背这个又硬又重的壳呢？\n妈妈说：因为我们的身体没有骨骼支撑，只能爬，又爬不快，所以要这个壳保护！\n小蜗牛：毛虫姐姐没有骨头，也爬不快，为什么她不用背壳？\n妈妈：因为毛虫姐姐能变成蝴蝶，天空会保护她啊。\n小蜗牛：可是蚯蚓弟弟也没骨头爬不快，也不会变成蝴蝶，为什么不背壳？\n妈妈：因为蚯蚓弟弟会钻土，大地会保护他啊。\n小蜗牛哭了起来：我们好可怜，天空不保护，大地也不保护。\n蜗牛妈妈安慰他：所以我们有壳啊！我们不靠天，不靠地，我们靠自己。🐌',
            '📖 小故事：\n\n有一个人在山顶上养了一群羊。一天，一只小羊问老羊："山那边是什么？"\n老羊说："我没去过，但听说那里有更绿的草地。"\n小羊很想去看看，但它害怕山路陡峭。\n很多年过去了，小羊变成了老羊，它的孩子又问了同样的问题。\n它终于鼓起勇气，翻过了山顶，发现那边确实有更绿的草地，而且路并没有想象中那么难走。\n\n💡 有时候，阻碍我们的不是路途的遥远，而是迈出第一步的勇气。',
        ]
        return random.choice(stories)

    if has_any(msg, ['名言', '格言', '座右铭', '励志', '鸡汤', '鼓励']):
        quotes = [
            '✨ "千里之行，始于足下。" —— 老子\n\n每一步虽小，但方向对了，终会到达。',
            '✨ "学而不思则罔，思而不学则殆。" —— 孔子\n\n学习和思考缺一不可。',
            '✨ "天行健，君子以自强不息。" —— 《周易》\n\n保持进取，永不停歇。',
            '✨ "生活不是等待暴风雨过去，而是学会在雨中翩翩起舞。" \n\n拥抱困难，从中成长。',
            '✨ "你的时间有限，不要浪费在过别人的生活上。" —— 乔布斯\n\n做自己，活出精彩。',
            '✨ "世界上只有一种英雄主义，那就是认清生活的真相后依然热爱生活。" —— 罗曼·罗兰\n\n保持热爱，勇往直前。',
        ]
        return random.choice(quotes)

    if has_any(msg, ['音乐', '歌', '唱歌', '听歌', '推荐歌']):
        return '🎵 音乐推荐：\n\n🎶 放松舒缓：《晴天》《稻香》《平凡之路》\n🎸 励志向上：《倔强》《海阔天空》《追梦赤子心》\n🎹 安静治愈：《南山南》《成都》《起风了》\n🎤 经典老歌：《月亮代表我的心》《光辉岁月》\n\n告诉我你喜欢的风格，我可以推荐更多！'

    if has_any(msg, ['美食', '吃什么', '做饭', '菜谱', '好吃']):
        return '🍜 美食推荐：\n\n🍳 快手菜：番茄炒蛋、蒜蓉西兰花、酸辣土豆丝\n🍲 暖心汤：紫菜蛋花汤、番茄牛腩汤、银耳红枣汤\n🥗 健康轻食：鸡胸肉沙拉、牛油果三明治、燕麦酸奶\n🍜 宿舍神器：泡面加蛋加菜、电饭煲焖饭\n\n告诉我你的口味偏好，我可以推荐更具体的！'

    general_replies = [
        f'收到！关于"{original}"，我来想想... 🤔\n\n这是一个很好的问题！虽然我目前无法给出完整的答案，但我可以：\n• 帮你分析问题的不同角度\n• 提供相关的思路和建议\n• 推荐进一步了解的方向\n\n你可以更具体地描述你的问题，我会尽力帮你！也可以问我学习、生活、科技、文化等各方面的问题~',
        f'关于"{original}"，这个问题挺有意思的！💡\n\n我不仅能帮你管理社团，还可以：\n• 📚 回答学习和知识问题\n• 💡 提供生活建议和思路\n• 🎯 帮你分析和解决问题\n• 💬 陪你聊天解闷\n\n告诉我更多细节，我来帮你分析！',
        f'嗯，关于"{original}"，让我想想~\n\n我虽然是社团管理系统的AI，但我也能回答很多其他领域的问题：\n• 📖 学习辅导和知识问答\n• 💼 职业规划和发展建议\n• 🧠 心理健康和情绪支持\n• 🎨 创意写作和灵感激发\n\n把你的问题说得更具体一些吧！'
    ]
    return random.choice(general_replies)


@app.route('/api/tools', methods=['GET', 'POST'])
def handle_tools():
    if request.method == 'GET':
        user = get_current_user()
        if not user:
            return jsonify({'error': '请先登录'}), 401
        conn = db.get_conn()
        try:
            if user['role'] == 'admin':
                rows = conn.execute('SELECT * FROM club_tools ORDER BY created_at DESC').fetchall()
            elif user['role'] == 'student':
                student_clubs = conn.execute('SELECT DISTINCT club_name FROM club_members WHERE user_id=?', (user['id'],)).fetchall()
                club_names = [c['club_name'] for c in student_clubs]
                if user['club_name'] and user['club_name'] not in club_names:
                    club_names.append(user['club_name'])
                if club_names:
                    placeholders = ','.join(['?'] * len(club_names))
                    rows = conn.execute(f'SELECT * FROM club_tools WHERE club_name IN ({placeholders}) ORDER BY created_at DESC', club_names).fetchall()
                else:
                    rows = []
            else:
                rows = conn.execute('SELECT * FROM club_tools WHERE club_name=? ORDER BY created_at DESC', (user['club_name'],)).fetchall()
        finally:
            conn.close()
        import json as _json
        result = []
        for r in rows:
            item = {'id': r['id'], 'clubName': r['club_name'], 'toolType': r['tool_type'], 'title': r['title'], 'description': r['description'], 'deadline': r['deadline'], 'status': r['status'], 'limit': r['limit_count'], 'voteMode': r['vote_mode'], 'format': r['format_hint'], 'perUserLimit': r['per_user_limit'] if 'per_user_limit' in r.keys() else 0, 'anonymous': r['anonymous'] if 'anonymous' in r.keys() else 0, 'showCounts': r['show_counts'] if 'show_counts' in r.keys() else 1, 'resultsVisible': r['results_visible'] if 'results_visible' in r.keys() else 1}
            opts = r['options'] or ''
            item['options'] = _json.loads(opts) if opts else []
            results = r['results'] or '{}'
            item['results'] = _json.loads(results)
            parts = r['participants'] or '[]'
            plist = _json.loads(parts)
            is_anonymous = item['anonymous']
            is_creator = (user['role'] in ('user', 'admin') and user['club_name'] == r['club_name'])
            if r['tool_type'] == 'vote' and is_anonymous and not is_creator:
                plist_anon = []
                for p in plist:
                    anon_p = dict(p)
                    anon_p['username'] = '匿名用户'
                    plist_anon.append(anon_p)
                plist = plist_anon
            item['participantCount'] = len(plist)
            item['participants'] = plist
            item['totalVotes'] = sum(item['results'].values()) if isinstance(item['results'], dict) else 0
            if r['tool_type'] == 'vote' and not is_creator:
                if not item['showCounts'] and r['status'] == 'active':
                    item['results'] = {}
                    item['totalVotes'] = 0
                if not item['resultsVisible']:
                    item['results'] = {}
                    item['totalVotes'] = 0
            result.append(item)
        return jsonify({'success': True, 'data': result})
    else:
        user = get_current_user()
        if not user:
            return jsonify({'error': '请先登录'}), 401
        data = request.json or {}
        import json as _json
        tool = {
            'club_name': user['club_name'] or '管理员',
            'tool_type': data.get('toolType', ''),
            'title': data.get('title', ''),
            'description': data.get('description', ''),
            'options': _json.dumps(data.get('options', []), ensure_ascii=False),
            'vote_mode': data.get('voteMode', 'single'),
            'limit_count': data.get('limit', 0),
            'format_hint': data.get('format', ''),
            'deadline': data.get('deadline', ''),
            'per_user_limit': data.get('perUserLimit', 0),
            'anonymous': data.get('anonymous', 0),
            'show_counts': data.get('showCounts', 1),
            'results_visible': data.get('resultsVisible', 1),
        }
        survey_questions = data.get('surveyQuestions', [])
        if tool['tool_type'] == 'survey' and survey_questions:
            tool['options'] = _json.dumps(survey_questions, ensure_ascii=False)
        if not tool['title'] or not tool['tool_type']:
            return jsonify({'error': '请填写完整'}), 400
        conn = db.get_conn()
        try:
            conn.execute('INSERT INTO club_tools (club_name, tool_type, title, description, options, vote_mode, limit_count, format_hint, deadline, per_user_limit, anonymous, show_counts, results_visible) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                        (tool['club_name'], tool['tool_type'], tool['title'], tool['description'], tool['options'], tool['vote_mode'], tool['limit_count'], tool['format_hint'], tool['deadline'], tool['per_user_limit'], tool['anonymous'], tool['show_counts'], tool['results_visible']))
            conn.commit()
            tool_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
            if data.get('notify') and tool['club_name']:
                try:
                    type_names = {'chain': '接龙', 'signup': '报名', 'vote': '投票', 'survey': '问卷调查'}
                    member_rows = conn.execute('SELECT DISTINCT user_id FROM club_members WHERE club_name=? AND user_id!=0', (tool['club_name'],)).fetchall()
                    notified_ids = set()
                    for m in member_rows:
                        if m['user_id'] and m['user_id'] not in notified_ids:
                            conn.execute('INSERT INTO notifications (user_id, title, content, type, link) VALUES (?, ?, ?, ?, ?)',
                                        (m['user_id'], '📋 新的' + type_names.get(tool["tool_type"], "活动") + '：' + tool['title'],
                                         tool['club_name'] + '发布了新的' + type_names.get(tool["tool_type"], "活动") + '活动，快来参与吧！',
                                         'tool', '/api/tool-page/' + str(tool_id)))
                            notified_ids.add(m['user_id'])
                    fallback_rows = conn.execute('SELECT id FROM users WHERE club_name=? AND role="student"', (tool['club_name'],)).fetchall()
                    for m in fallback_rows:
                        if m['id'] not in notified_ids:
                            conn.execute('INSERT INTO notifications (user_id, title, content, type, link) VALUES (?, ?, ?, ?, ?)',
                                        (m['id'], '📋 新的' + type_names.get(tool["tool_type"], "活动") + '：' + tool['title'],
                                         tool['club_name'] + '发布了新的' + type_names.get(tool["tool_type"], "活动") + '活动，快来参与吧！',
                                         'tool', '/api/tool-page/' + str(tool_id)))
                            notified_ids.add(m['id'])
                    conn.commit()
                except:
                    pass
        finally:
            conn.close()
        return jsonify({'success': True})


@app.route('/api/tools/<int:tid>', methods=['DELETE'])
def delete_tool(tid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM club_tools WHERE id=? AND (club_name=? OR ?="admin")', (tid, user['club_name'], user['role']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/tool-page/<int:tid>')
def tool_page(tid):
    import json as _json
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT * FROM club_tools WHERE id=?', (tid,)).fetchone()
    finally:
        conn.close()
    if not row:
        return '<h3>工具不存在</h3>', 404
    opts = _json.loads(row['options'] or '[]')
    parts = _json.loads(row['participants'] or '[]')
    results = _json.loads(row['results'] or '{}')
    type_names = {'chain': '🔢 接龙', 'signup': '📝 报名', 'vote': '🗳️ 投票', 'survey': '📊 问卷调查'}
    tool_type = row['tool_type']
    h = f'''<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
    <title>{row["title"]}</title>
    <style>*{{margin:0;padding:0;box-sizing:border-box}}body{{font-family:'Microsoft YaHei',sans-serif;background:#f4f5fb;min-height:100vh;padding:20px}}
    .container{{max-width:600px;margin:0 auto}}.card{{background:#fff;border-radius:14px;padding:20px;margin-bottom:14px;border:1px solid #f0f0f8}}
    h2{{font-size:1.2em;color:#1a1a2e;margin-bottom:8px}}.meta{{font-size:0.8em;color:#a0a3bd;margin-bottom:12px}}
    .desc{{font-size:0.88em;color:#4a4d6a;line-height:1.6;margin-bottom:14px}}
    .option{{padding:12px;border-radius:10px;border:1.5px solid #f0f0f8;margin-bottom:8px;cursor:pointer;transition:all .2s;font-size:0.88em;color:#1a1a2e}}
    .option:hover{{border-color:#667eea;background:rgba(102,126,234,0.03)}}.option.selected{{border-color:#667eea;background:rgba(102,126,234,0.06)}}
    .btn{{padding:10px 20px;border-radius:10px;border:none;font-size:0.88em;cursor:pointer;font-weight:600;transition:all .2s}}
    .btn-primary{{background:linear-gradient(135deg,#667eea,#764ba2);color:#fff}}.btn:hover{{opacity:0.9}}
    input,textarea{{width:100%;padding:10px 12px;border-radius:8px;border:1.5px solid #eee;background:#fafaff;color:#1a1a2e;font-size:0.86em;outline:none;font-family:inherit;margin-bottom:8px}}
    input:focus,textarea:focus{{border-color:#667eea}}
    .result-bar{{height:24px;border-radius:6px;background:rgba(102,126,234,0.1);margin-top:4px;overflow:hidden}}
    .result-fill{{height:100%;border-radius:6px;background:linear-gradient(135deg,#667eea,#764ba2);transition:width .3s}}
    .part-list{{font-size:0.82em;color:#6b6d8a}}.part-item{{padding:6px 0;border-bottom:1px solid #fafaff}}
    .survey-q{{font-size:0.9em;font-weight:600;color:#1a1a2e;margin-bottom:8px;margin-top:14px}}.survey-a{{font-size:0.82em;color:#6b6d8a;padding:6px 0;border-bottom:1px solid #fafaff}}
    </style></head><body><div class="container">'''
    h += f'<div class="card"><h2>{type_names.get(tool_type, tool_type)} {row["title"]}</h2>'
    h += f'<div class="meta">{row["club_name"]} · {row["created_at"]}'
    if row['deadline']:
        h += f' · 截止：{row["deadline"]}'
    h += '</div>'
    if row['description']:
        h += f'<div class="desc">{row["description"]}</div>'
    if tool_type == 'vote':
        vote_mode = row['vote_mode'] if 'vote_mode' in row.keys() else 'single'
        is_anonymous = row['anonymous'] if 'anonymous' in row.keys() else 0
        show_counts = row['show_counts'] if 'show_counts' in row.keys() else 1
        results_visible = row['results_visible'] if 'results_visible' in row.keys() else 1
        mode_label = '多选' if vote_mode == 'multi' else '单选'
        anon_label = ' · 匿名投票' if is_anonymous else ''
        h += f'<div style="font-size:0.82em;color:#667eea;font-weight:600;margin-bottom:10px">{mode_label}{anon_label}</div>'
        total_votes = sum(results.values()) if isinstance(results, dict) else 0
        for i, opt in enumerate(opts):
            count = results.get(str(i), 0)
            pct = round(count / total_votes * 100) if total_votes > 0 else 0
            if show_counts:
                h += f'<div class="option" onclick="selectVote({i},this)">{opt} <span style="float:right;color:#667eea;font-weight:600">{count}票 ({pct}%)</span><div class="result-bar"><div class="result-fill" style="width:{pct}%"></div></div></div>'
            else:
                h += f'<div class="option" onclick="selectVote({i},this)">{opt}</div>'
        h += f'<div style="font-size:0.78em;color:#a0a3bd;margin-top:8px">共 {total_votes} 票 · {len(parts)} 人参与</div>'
        h += f'<button class="btn btn-primary" style="margin-top:12px" onclick="submitVote()">投票</button>'
        if not results_visible:
            h += '<div style="font-size:0.78em;color:#a0a3bd;margin-top:6px">投票结果仅创建者可见</div>'
    elif tool_type == 'signup':
        h += f'<div style="font-size:0.88em;color:#667eea;font-weight:600;margin-bottom:10px">已报名 {len(parts)}/{row["limit_count"] or "不限"}</div>'
        h += '<input type="text" id="signupInfo" placeholder="报名信息（如：姓名+联系方式）" />'
        h += '<button class="btn btn-primary" onclick="submitSignup()">报名</button>'
        if parts:
            h += '<div class="part-list" style="margin-top:12px">'
            for p in parts:
                h += f'<div class="part-item">{p.get("username","")} - {p.get("info","")} ({p.get("time","")})</div>'
            h += '</div>'
    elif tool_type == 'chain':
        h += f'<div style="font-size:0.88em;color:#667eea;font-weight:600;margin-bottom:10px">已接龙 {len(parts)}{("/"+str(row["limit_count"])) if row["limit_count"]>0 else ""} 人</div>'
        if row['format_hint']:
            h += f'<div style="font-size:0.78em;color:#a0a3bd;margin-bottom:8px">格式：{row["format_hint"]}</div>'
        h += '<input type="text" id="chainContent" placeholder="接龙内容" />'
        h += '<button class="btn btn-primary" onclick="submitChain()">参与接龙</button>'
        if parts:
            h += '<div class="part-list" style="margin-top:12px">'
            for i, p in enumerate(parts, 1):
                h += f'<div class="part-item">{i}. {p.get("username","")} - {p.get("content","")} ({p.get("time","")})</div>'
            h += '</div>'
    elif tool_type == 'survey':
        per_user_limit = row['per_user_limit'] if 'per_user_limit' in row.keys() else 0
        limit_info = f' · 每人限填{per_user_limit}次' if per_user_limit > 0 else ''
        h += f'<div style="font-size:0.88em;color:#667eea;font-weight:600;margin-bottom:10px">共 {len(opts)} 道题目 · {len(parts)} 人已填写{limit_info}</div>'
        for i, q in enumerate(opts):
            h += f'<div class="survey-q">{i+1}. {q}</div>'
            h += f'<textarea id="surveyQ{i}" placeholder="请输入你的回答" rows="2"></textarea>'
        h += '<button class="btn btn-primary" style="margin-top:10px" onclick="submitSurvey()">提交问卷</button>'
        if parts:
            h += '<div style="margin-top:16px;padding-top:12px;border-top:1px solid #f0f0f8"><div style="font-size:0.88em;font-weight:600;color:#1a1a2e;margin-bottom:8px">填写记录</div>'
            for p in parts:
                answers = p.get('answers', {})
                h += f'<div style="font-size:0.82em;color:#6b6d8a;margin-bottom:8px;padding:8px;background:#fafaff;border-radius:8px">'
                h += f'<b>{p.get("username","匿名")}</b> <span style="color:#c8cade">{p.get("time","")}</span><br>'
                for qi, q in enumerate(opts):
                    ans = answers.get(str(qi), '未回答')
                    h += f'Q{qi+1}: {ans}<br>'
                h += '</div>'
            h += '</div>'
    h += '</div></div>'
    opts_json = json.dumps(opts, ensure_ascii=False)
    vote_mode = row['vote_mode'] if 'vote_mode' in row.keys() else 'single'
    h += '<script>var tid=' + str(tid) + ';var selectedVote=-1;var selectedVotes=[];var voteMode="' + vote_mode + '";'
    h += 'function selectVote(i,el){if(voteMode==="multi"){if(el.classList.contains("selected")){el.classList.remove("selected");el.style.borderColor="#f0f0f8";el.style.background="";selectedVotes=selectedVotes.filter(function(v){return v!==i})}else{el.classList.add("selected");el.style.borderColor="#667eea";el.style.background="rgba(102,126,234,0.06)";selectedVotes.push(i)}}else{document.querySelectorAll(".option").forEach(function(o){o.classList.remove("selected");o.style.borderColor="#f0f0f8";o.style.background=""});el.classList.add("selected");el.style.borderColor="#667eea";el.style.background="rgba(102,126,234,0.06)";selectedVote=i}}'
    h += 'function submitVote(){if(voteMode==="multi"){if(selectedVotes.length===0){alert("请选择至少一个选项");return}fetch("/api/tools/"+tid+"/participate",{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify({choices:selectedVotes})}).then(function(r){return r.json()}).then(function(res){if(res.success){alert("投票成功！");location.reload()}else{alert(res.error||"失败")}})}else{if(selectedVote<0){alert("请选择选项");return}fetch("/api/tools/"+tid+"/participate",{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify({choice:selectedVote})}).then(function(r){return r.json()}).then(function(res){if(res.success){alert("投票成功！");location.reload()}else{alert(res.error||"失败")}})}}'
    h += 'function submitSignup(){var info=document.getElementById("signupInfo").value.trim();if(!info){alert("请输入报名信息");return}fetch("/api/tools/"+tid+"/participate",{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify({info:info})}).then(function(r){return r.json()}).then(function(res){if(res.success){alert("报名成功！");location.reload()}else{alert(res.error||"失败")}})}'
    h += 'function submitChain(){var content=document.getElementById("chainContent").value.trim();if(!content){alert("请输入接龙内容");return}fetch("/api/tools/"+tid+"/participate",{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify({content:content})}).then(function(r){return r.json()}).then(function(res){if(res.success){alert("接龙成功！");location.reload()}else{alert(res.error||"失败")}})}'
    h += 'function submitSurvey(){var answers={};var qs=' + opts_json + ';for(var i=0;i<qs.length;i++){var el=document.getElementById("surveyQ"+i);if(el)answers[i]=el.value.trim()}fetch("/api/tools/"+tid+"/participate",{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify({answers:answers})}).then(function(r){return r.json()}).then(function(res){if(res.success){alert("提交成功！");location.reload()}else{alert(res.error||"失败")}})}'
    h += '</script></body></html>'
    return h


@app.route('/api/export-survey/<int:tid>')
def export_survey(tid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '无权限'}), 403
    import json as _json
    from openpyxl import Workbook
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT * FROM club_tools WHERE id=?', (tid,)).fetchone()
        if not row:
            return jsonify({'error': '不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != row['club_name']:
            return jsonify({'error': '无权限'}), 403
    finally:
        conn.close()
    questions = _json.loads(row['options'] or '[]')
    parts = _json.loads(row['participants'] or '[]')
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = '问卷统计'
    headers = ['序号', '填写人', '填写时间']
    for q in questions:
        headers.append(q[:30])
    ws.append(headers)
    for i, p in enumerate(parts, 1):
        row_data = [i, p.get('username', ''), p.get('time', '')]
        answers = p.get('answers', {})
        for qi in range(len(questions)):
            row_data.append(answers.get(str(qi), ''))
        ws.append(row_data)
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'{row["title"]}_问卷统计.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/api/export-tool/<int:tid>')
def export_tool(tid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '无权限'}), 403
    import json as _json
    from openpyxl import Workbook
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT * FROM club_tools WHERE id=?', (tid,)).fetchone()
        if not row:
            return jsonify({'error': '不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != row['club_name']:
            return jsonify({'error': '无权限'}), 403
    finally:
        conn.close()
    tool_type = row['tool_type']
    title = row['title']
    opts = _json.loads(row['options'] or '[]')
    parts = _json.loads(row['participants'] or '[]')
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    type_labels = {'chain': '接龙', 'signup': '报名', 'vote': '投票', 'survey': '问卷'}
    ws.title = type_labels.get(tool_type, '工具') + '统计'
    if tool_type == 'chain':
        headers = ['序号', '填写人', '填写时间', '填写内容']
        ws.append(headers)
        for i, p in enumerate(parts, 1):
            ws.append([i, p.get('username', ''), p.get('time', ''), p.get('value', '')])
    elif tool_type == 'signup':
        headers = ['序号', '报名人', '报名时间', '报名选项']
        ws.append(headers)
        for i, p in enumerate(parts, 1):
            ws.append([i, p.get('username', ''), p.get('time', ''), p.get('option', '')])
    elif tool_type == 'vote':
        is_anonymous = row['anonymous'] if 'anonymous' in row.keys() else 0
        vote_mode = row['vote_mode'] if 'vote_mode' in row.keys() else 'single'
        if is_anonymous:
            headers = ['序号', '投票时间', '投票选项']
        else:
            headers = ['序号', '投票人', '投票时间', '投票选项']
        ws.append(headers)
        for i, p in enumerate(parts, 1):
            if is_anonymous:
                ws.append([i, p.get('time', ''), p.get('option', '')])
            else:
                ws.append([i, p.get('username', ''), p.get('time', ''), p.get('option', '')])
        ws.append([])
        ws.append(['选项统计'])
        ws.append(['选项', '票数'])
        opt_counts = {}
        for o in opts:
            opt_counts[o] = 0
        results_data = _json.loads(row['results'] or '{}')
        for i, o in enumerate(opts):
            opt_counts[o] = results_data.get(str(i), 0)
        for o, c in opt_counts.items():
            ws.append([o, c])
    elif tool_type == 'survey':
        headers = ['序号', '填写人', '填写时间']
        for q in opts:
            headers.append(q[:30])
        ws.append(headers)
        for i, p in enumerate(parts, 1):
            row_data = [i, p.get('username', ''), p.get('time', '')]
            answers = p.get('answers', {})
            for qi in range(len(opts)):
                row_data.append(answers.get(str(qi), ''))
            ws.append(row_data)
    else:
        headers = ['序号', '用户', '时间', '内容']
        ws.append(headers)
        for i, p in enumerate(parts, 1):
            ws.append([i, p.get('username', ''), p.get('time', ''), str(p)])
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'{title}_{type_labels.get(tool_type, "工具")}统计.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/tools/<int:tid>/participate', methods=['POST'])
def participate_tool(tid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    import json as _json
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT * FROM club_tools WHERE id=?', (tid,)).fetchone()
        if not row:
            return jsonify({'error': '不存在'}), 404
        if row['deadline']:
            try:
                dl = datetime.strptime(row['deadline'], '%Y-%m-%dT%H:%M')
                if datetime.now() > dl:
                    return jsonify({'error': '该活动已截止'}), 400
            except:
                try:
                    dl = datetime.strptime(row['deadline'], '%Y-%m-%d %H:%M')
                    if datetime.now() > dl:
                        return jsonify({'error': '该活动已截止'}), 400
                except:
                    pass
        parts = _json.loads(row['participants'] or '[]')
        results = _json.loads(row['results'] or '{}')
        opts = _json.loads(row['options'] or '[]')
        if row['tool_type'] in ('vote', 'signup', 'chain'):
            for p in parts:
                if p.get('username') == user['username']:
                    type_label = {'vote':'投票','signup':'报名','chain':'接龙'}.get(row['tool_type'], '参与')
                    return jsonify({'error': f'您已参与过该{type_label}，不能重复参与'}), 400
        elif row['tool_type'] == 'survey':
            per_user_limit = row['per_user_limit'] if 'per_user_limit' in row.keys() else 0
            if per_user_limit > 0:
                user_count = sum(1 for p in parts if p.get('username') == user['username'])
                if user_count >= per_user_limit:
                    return jsonify({'error': f'您已填写{user_count}次，每人最多填写{per_user_limit}次'}), 400
        entry = {'username': user['username'], 'time': datetime.now().strftime('%Y-%m-%d %H:%M')}
        if row['tool_type'] == 'vote':
            vote_mode = row['vote_mode'] if 'vote_mode' in row.keys() else 'single'
            if vote_mode == 'multi':
                choices = data.get('choices', [])
                if not choices or not isinstance(choices, list):
                    return jsonify({'error': '请选择投票选项'}), 400
                entry['choices'] = choices
                entry['option'] = '、'.join([opts[int(c)] for c in choices if int(c) < len(opts)])
                for c in choices:
                    ck = str(c)
                    results[ck] = results.get(ck, 0) + 1
            else:
                choice = data.get('choice')
                if choice is not None:
                    ck = str(choice)
                    results[ck] = results.get(ck, 0) + 1
                    entry['choice'] = choice
                    entry['option'] = opts[int(choice)] if int(choice) < len(opts) else ''
        elif row['tool_type'] == 'signup':
            if row['limit_count'] > 0 and len(parts) >= row['limit_count']:
                return jsonify({'error': '名额已满'}), 400
            entry['info'] = data.get('info', '')
        elif row['tool_type'] == 'chain':
            if row['limit_count'] > 0 and len(parts) >= row['limit_count']:
                return jsonify({'error': '名额已满'}), 400
            entry['content'] = data.get('content', '')
        elif row['tool_type'] == 'survey':
            entry['answers'] = data.get('answers', {})
        parts.append(entry)
        conn.execute('UPDATE club_tools SET participants=?, results=? WHERE id=?', (_json.dumps(parts, ensure_ascii=False), _json.dumps(results, ensure_ascii=False), tid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/tree-hole', methods=['GET', 'POST'])
def handle_tree_hole():
    if request.method == 'GET':
        scope = request.args.get('scope', 'public')
        club = request.args.get('club', '')
        conn = db.get_conn()
        try:
            if scope == 'club' and club:
                user = get_current_user()
                if not user or (user['role'] != 'admin' and user['club_name'] != club):
                    return jsonify({'error': '无权限查看该社团内部树洞'}), 403
                rows = conn.execute('SELECT id, content, scope, club_name, status, admin_note, created_at FROM tree_hole WHERE scope="club" AND club_name=? AND status!="hidden" ORDER BY created_at DESC LIMIT 50', (club,)).fetchall()
            else:
                rows = conn.execute('SELECT id, content, scope, club_name, status, admin_note, created_at FROM tree_hole WHERE scope="public" AND status!="hidden" ORDER BY created_at DESC LIMIT 50').fetchall()
        finally:
            conn.close()
        return jsonify({'success': True, 'data': [{'id': r['id'], 'content': r['content'], 'scope': r['scope'], 'clubName': r['club_name'], 'status': r['status'], 'adminNote': r['admin_note'], 'time': local_time(r['created_at'])} for r in rows]})
    else:
        data = request.json or {}
        content = data.get('content', '').strip()
        scope = data.get('scope', 'public')
        club = data.get('clubName', '')
        if not content:
            return jsonify({'error': '请输入内容'}), 400
        if scope == 'club':
            if not club:
                return jsonify({'error': '请选择社团'}), 400
        conn = db.get_conn()
        try:
            conn.execute('INSERT INTO tree_hole (content, scope, club_name) VALUES (?, ?, ?)', (content, scope, club))
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True})


@app.route('/api/tree-hole/admin')
def admin_tree_hole():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, content, scope, club_name, status, admin_note, created_at FROM tree_hole ORDER BY created_at DESC LIMIT 100').fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'content': r['content'], 'scope': r['scope'], 'clubName': r['club_name'], 'status': r['status'], 'adminNote': r['admin_note'], 'time': local_time(r['created_at'])} for r in rows]})


@app.route('/api/tree-hole/club')
def club_tree_hole():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    if user['role'] == 'admin':
        club = request.args.get('club', '')
    else:
        club = user['club_name'] or ''
    if not club:
        return jsonify({'success': True, 'data': []})
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, content, scope, club_name, status, admin_note, created_at FROM tree_hole WHERE scope="club" AND club_name=? AND status!="hidden" ORDER BY created_at DESC', (club,)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'content': r['content'], 'scope': r['scope'], 'clubName': r['club_name'], 'status': r['status'], 'adminNote': r['admin_note'], 'time': local_time(r['created_at'])} for r in rows]})


@app.route('/api/tree-hole/<int:hid>/manage', methods=['POST'])
def manage_tree_hole(hid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    action = data.get('action', '')
    if action not in ('hide', 'warn', 'restore'):
        return jsonify({'error': '无效操作'}), 400
    conn = db.get_conn()
    try:
        if action == 'hide':
            conn.execute('UPDATE tree_hole SET status="hidden" WHERE id=?', (hid,))
        elif action == 'warn':
            note = data.get('note', '')
            conn.execute('UPDATE tree_hole SET status="warned", admin_note=? WHERE id=?', (note, hid))
        elif action == 'restore':
            conn.execute('UPDATE tree_hole SET status="active", admin_note="" WHERE id=?', (hid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/joint-activities', methods=['GET', 'POST'])
def handle_joint_activities():
    if request.method == 'GET':
        user = get_current_user()
        show_all = request.args.get('all', '') == '1' and user and user['role'] == 'admin'
        conn = db.get_conn()
        try:
            if show_all:
                rows = conn.execute('SELECT ja.*, (SELECT COUNT(*) FROM joint_replies WHERE activity_id=ja.id AND reply_type="cooperate") as cooperate_count FROM joint_activities ja ORDER BY created_at DESC').fetchall()
            else:
                rows = conn.execute('SELECT ja.*, (SELECT COUNT(*) FROM joint_replies WHERE activity_id=ja.id AND reply_type="cooperate") as cooperate_count FROM joint_activities ja WHERE ja.status="open" ORDER BY created_at DESC').fetchall()
            result = []
            for r in rows:
                replies = conn.execute('SELECT id, club_name, reply_type, content, created_at FROM joint_replies WHERE activity_id=? ORDER BY created_at', (r['id'],)).fetchall()
                result.append({'id': r['id'], 'clubName': r['club_name'], 'title': r['title'], 'description': r['description'], 'supportNeeded': r['support_needed'], 'status': r['status'], 'cooperateCount': r['cooperate_count'], 'time': local_time(r['created_at']), 'replies': [{'id': rep['id'], 'clubName': rep['club_name'], 'type': rep['reply_type'], 'content': rep['content'], 'time': local_time(rep['created_at'])} for rep in replies]})
        finally:
            conn.close()
        return jsonify({'success': True, 'data': result})
    else:
        user = get_current_user()
        if not user or user['role'] not in ('user', 'admin', 'teacher'):
            return jsonify({'error': '请先登录'}), 401
        data = request.json or {}
        title = data.get('title', '').strip()
        description = data.get('description', '').strip()
        support = data.get('supportNeeded', '').strip()
        if not title:
            return jsonify({'error': '请输入活动主题'}), 400
        conn = db.get_conn()
        try:
            club = data.get('clubName', '').strip() or user['club_name'] or '管理员'
            if user['role'] == 'teacher':
                tc_rows = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
                teacher_clubs = [r['club_name'] for r in tc_rows if r['club_name']]
                if club not in teacher_clubs:
                    club = teacher_clubs[0] if teacher_clubs else (user['club_name'] or '管理员')
            conn.execute('INSERT INTO joint_activities (club_name, title, description, support_needed) VALUES (?, ?, ?, ?)', (club, title, description, support))
            conn.commit()
            all_users = conn.execute('SELECT id, role, club_name FROM users WHERE role IN ("user", "admin")').fetchall()
            for u in all_users:
                if u['club_name'] == club and u['role'] == 'user':
                    continue
                send_notification(u['id'], '🤝 新联合活动', f'「{club}」发布了联合活动「{title}」，快来看看吧', 'joint', '/club-tools.html?panel=joint', conn=conn)
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True})


@app.route('/api/joint-activities/<int:aid>/reply', methods=['POST'])
def reply_joint_activity(aid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    reply_type = data.get('type', 'message')
    content = data.get('content', '').strip()
    club = data.get('clubName', '').strip() or user['club_name'] or '管理员'
    conn = db.get_conn()
    try:
        activity = conn.execute('SELECT club_name, title FROM joint_activities WHERE id=?', (aid,)).fetchone()
        if not activity:
            return jsonify({'error': '活动不存在'}), 404
        conn.execute('INSERT INTO joint_replies (activity_id, club_name, reply_type, content) VALUES (?, ?, ?, ?)', (aid, club, reply_type, content))
        if reply_type == 'cooperate':
            club_users = conn.execute('SELECT id FROM users WHERE club_name=?', (activity['club_name'],)).fetchall()
            for cu in club_users:
                send_notification(cu['id'], '🤝 合作意向', f'「{club}」对您的联合活动「{activity["title"]}」表达了合作意向', 'joint', '/club-tools.html', conn=conn)
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/joint-activities/<int:aid>', methods=['DELETE'])
def delete_joint_activity(aid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        activity = conn.execute('SELECT club_name FROM joint_activities WHERE id=?', (aid,)).fetchone()
        if not activity:
            return jsonify({'error': '活动不存在'}), 404
        can_delete = False
        if user['role'] == 'admin':
            can_delete = True
        elif user['role'] == 'teacher':
            tc_rows = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
            teacher_clubs = [r['club_name'] for r in tc_rows if r['club_name']]
            if activity['club_name'] in teacher_clubs:
                can_delete = True
        elif user['club_name'] == activity['club_name']:
            can_delete = True
        if not can_delete:
            return jsonify({'error': '无权删除'}), 403
        conn.execute('DELETE FROM joint_replies WHERE activity_id=?', (aid,))
        conn.execute('DELETE FROM joint_activities WHERE id=?', (aid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/joint-activities/<int:aid>/close', methods=['POST'])
def close_joint_activity(aid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        conn.execute('UPDATE joint_activities SET status="closed" WHERE id=?', (aid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/joint-activities/<int:aid>/reopen', methods=['POST'])
def reopen_joint_activity(aid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        conn.execute('UPDATE joint_activities SET status="open" WHERE id=?', (aid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


def gen_checkin_code():
    return ''.join(random.choices(string.digits, k=6))


@app.route('/api/checkin-sessions', methods=['GET', 'POST'])
def handle_checkin_sessions():
    if request.method == 'GET':
        user = get_current_user()
        if not user or user['role'] not in ('user', 'admin', 'teacher', 'student'):
            return jsonify({'error': '请先登录'}), 401
        if user['role'] == 'admin':
            club = request.args.get('club', '')
        elif user['role'] == 'teacher':
            club = request.args.get('club', '')
            if club:
                tc_conn = db.get_conn()
                try:
                    tc_row = tc_conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
                finally:
                    tc_conn.close()
                if not tc_row:
                    return jsonify({'error': '无权限查看该社团'}), 403
            else:
                club = user['club_name'] or ''
                if not club:
                    tc_conn2 = db.get_conn()
                    try:
                        tc_first = tc_conn2.execute('SELECT club_name FROM teacher_clubs WHERE user_id=? ORDER BY club_name LIMIT 1', (user['id'],)).fetchone()
                    finally:
                        tc_conn2.close()
                    if tc_first:
                        club = tc_first['club_name']
        elif user['role'] == 'student':
            club = request.args.get('club', '') or user['club_name'] or ''
            if club and not is_cadre_of_club(user['id'], club):
                return jsonify({'error': '无权限查看该社团'}), 403
        else:
            club = user['club_name'] or ''
        if not club:
            return jsonify({'success': True, 'data': []})
        start_date = request.args.get('start_date', '').strip()
        end_date = request.args.get('end_date', '').strip()
        date_where = ''
        date_params = [club]
        if start_date:
            date_where += ' AND date(created_at)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(created_at)<=?'
            date_params.append(end_date)
        conn = db.get_conn()
        try:
            _now = cn_now()
            expired = conn.execute("SELECT id FROM checkin_sessions WHERE club_name=? AND status='open' AND end_time IS NOT NULL AND end_time!='' AND datetime(end_time)<?", (club, _now.strftime('%Y-%m-%d %H:%M'))).fetchall()
            for e in expired:
                conn.execute("UPDATE checkin_sessions SET status='closed', closed_at=CURRENT_TIMESTAMP WHERE id=?", (e['id'],))
            if expired:
                conn.commit()
            rows = conn.execute(f'SELECT * FROM checkin_sessions WHERE club_name=?{date_where} ORDER BY created_at DESC', date_params).fetchall()
        finally:
            conn.close()
        return jsonify({'success': True, 'data': [{'id': r['id'], 'clubName': r['club_name'], 'activityName': r['activity_name'], 'checkinCode': r['checkin_code'], 'locationName': r['location_name'], 'status': r['status'], 'activityTime': r['activity_time'] if 'activity_time' in r.keys() else '', 'startTime': r['start_time'] if 'start_time' in r.keys() else '', 'endTime': r['end_time'] if 'end_time' in r.keys() else '', 'checkinMethod': r['checkin_method'] if 'checkin_method' in r.keys() else 'qrcode', 'activityContent': r['activity_content'] if 'activity_content' in r.keys() else '', 'planText': r['plan_text'] if 'plan_text' in r.keys() else '', 'planPath': r['plan_path'] if 'plan_path' in r.keys() else '', 'summaryText': r['summary_text'] if 'summary_text' in r.keys() else '', 'summaryPath': r['summary_path'] if 'summary_path' in r.keys() else '', 'teacherIds': r['teacher_ids'] if 'teacher_ids' in r.keys() else '', 'isCompleted': r['is_completed'] if 'is_completed' in r.keys() else 0, 'completionPhoto': r['completion_photo'] if 'completion_photo' in r.keys() else '', 'warning': r['warning'] if 'warning' in r.keys() else '', 'warningReason': r['warning_reason'] if 'warning_reason' in r.keys() else '', 'checkoutCode': r['checkout_code'] if 'checkout_code' in r.keys() else '', 'checkoutMethod': r['checkout_method'] if 'checkout_method' in r.keys() else '', 'locationLat': r['location_lat'] if 'location_lat' in r.keys() else 0, 'locationLng': r['location_lng'] if 'location_lng' in r.keys() else 0, 'time': local_time(r['created_at'])} for r in rows]})
    else:
        user = get_current_user()
        if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
            return jsonify({'error': '请先登录'}), 401
        data = request.json or {}
        club = data.get('clubName', '').strip()
        if not club:
            club = user['club_name'] or ''
        if not club:
            return jsonify({'error': '无社团信息'}), 400
        if user['role'] == 'student' and not is_cadre_of_club(user['id'], club):
            return jsonify({'error': '无权限操作该社团'}), 403
        if user['role'] == 'teacher':
            tc_conn = db.get_conn()
            try:
                tc_row = tc_conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
            finally:
                tc_conn.close()
            if not tc_row:
                return jsonify({'error': '无权限操作该社团'}), 403
        activity_name = data.get('activityName', '').strip()
        location_name = data.get('locationName', '').strip()
        activity_time = data.get('activityTime', '').strip()
        start_time = data.get('startTime', '').strip()
        end_time = data.get('endTime', '').strip()
        checkin_method = data.get('checkinMethod', 'qrcode').strip()
        activity_content = data.get('activityContent', '').strip()
        plan_text = data.get('planText', '').strip()
        plan_path = data.get('planPath', '').strip()
        teacher_ids = data.get('teacherIds', '')
        location_lat = data.get('locationLat', 0)
        location_lng = data.get('locationLng', 0)
        if not plan_text and not plan_path:
            return jsonify({'error': '请上传活动计划文件'}), 400
        if isinstance(teacher_ids, list):
            teacher_ids = ','.join(str(t) for t in teacher_ids)
        code = gen_checkin_code()
        conn = db.get_conn()
        try:
            while conn.execute('SELECT id FROM checkin_sessions WHERE checkin_code=?', (code,)).fetchone():
                code = gen_checkin_code()
            conn.execute('INSERT INTO checkin_sessions (club_name, activity_name, checkin_code, location_name, location_lat, location_lng, activity_time, start_time, end_time, checkin_method, activity_content, plan_text, plan_path, teacher_ids) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (club, activity_name, code, location_name, location_lat, location_lng, activity_time, start_time, end_time, checkin_method, activity_content, plan_text, plan_path, teacher_ids))
            conn.commit()
            sid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
            try:
                import traceback as _tb
                member_rows = conn.execute('SELECT DISTINCT user_id FROM club_members WHERE club_name=? AND user_id!=0', (club,)).fetchall()
                notified_ids = set()
                link = f'/dashboard.html?page=checkin&sid={sid}'
                for m in member_rows:
                    if m['user_id'] and m['user_id'] not in notified_ids:
                        conn.execute('INSERT INTO notifications (user_id, title, content, type, link) VALUES (?, ?, ?, ?, ?)',
                                    (m['user_id'], '✅ '+club+'活动签到已发起', activity_name+'已发起，活动时间内可进行签到', 'checkin', link))
                        notified_ids.add(m['user_id'])
                fallback_rows = conn.execute('SELECT id FROM users WHERE club_name=? AND role="student"', (club,)).fetchall()
                for m in fallback_rows:
                    if m['id'] not in notified_ids:
                        conn.execute('INSERT INTO notifications (user_id, title, content, type, link) VALUES (?, ?, ?, ?, ?)',
                                    (m['id'], '✅ '+club+'活动签到已发起', activity_name+'已发起，活动时间内可进行签到', 'checkin', link))
                        notified_ids.add(m['id'])
                conn.commit()
                print(f'[checkin] 已向 {len(notified_ids)} 名社团成员发送签到通知 (session {sid})')
            except Exception as _e:
                print(f'[checkin] 发送签到通知失败: {_e}')
                import traceback
                traceback.print_exc()
        finally:
            conn.close()
        return jsonify({'success': True, 'id': sid, 'checkinCode': code, 'checkinMethod': checkin_method})


@app.route('/api/checkin-sessions/<int:sid>', methods=['GET', 'PUT', 'DELETE'])
def manage_checkin_session(sid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher', 'student'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (sid,)).fetchone()
        if not sess_row:
            return jsonify({'error': '签到会话不存在'}), 404
        if user['role'] not in ('admin',) and user['club_name'] != sess_row['club_name']:
            if user['role'] == 'teacher':
                tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], sess_row['club_name'])).fetchone()
                if not tc_row:
                    return jsonify({'error': '无权限'}), 403
            elif user['role'] == 'student':
                if not is_cadre_of_club(user['id'], sess_row['club_name']):
                    return jsonify({'error': '无权限'}), 403
            else:
                return jsonify({'error': '无权限'}), 403
        if request.method == 'GET':
            if sess_row['status'] == 'open':
                end_time = sess_row['end_time'] if 'end_time' in sess_row.keys() else ''
                if end_time:
                    try:
                        et = datetime.strptime(end_time.replace('T', ' ')[:16], '%Y-%m-%d %H:%M')
                        if cn_now() > et:
                            conn.execute("UPDATE checkin_sessions SET status='closed', closed_at=CURRENT_TIMESTAMP WHERE id=?", (sid,))
                            conn.commit()
                            sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (sid,)).fetchone()
                    except:
                        pass
            teacher_checkins = conn.execute('SELECT teacher_user_id, status FROM teacher_checkin_checkout WHERE session_id=?', (sid,)).fetchall()
            result = {'id': sess_row['id'], 'clubName': sess_row['club_name'], 'activityName': sess_row['activity_name'], 'checkinCode': sess_row['checkin_code'], 'status': sess_row['status'], 'startTime': sess_row['start_time'] if 'start_time' in sess_row.keys() else '', 'endTime': sess_row['end_time'] if 'end_time' in sess_row.keys() else '', 'activityTime': sess_row['activity_time'] if 'activity_time' in sess_row.keys() else '', 'locationName': sess_row['location_name'] if 'location_name' in sess_row.keys() else '', 'activityContent': sess_row['activity_content'] if 'activity_content' in sess_row.keys() else '', 'checkinMethod': sess_row['checkin_method'] if 'checkin_method' in sess_row.keys() else 'qrcode', 'checkoutCode': sess_row['checkout_code'] if 'checkout_code' in sess_row.keys() else '', 'checkoutMethod': sess_row['checkout_method'] if 'checkout_method' in sess_row.keys() else '', 'planText': sess_row['plan_text'] if 'plan_text' in sess_row.keys() else '', 'planPath': sess_row['plan_path'] if 'plan_path' in sess_row.keys() else '', 'summaryText': sess_row['summary_text'] if 'summary_text' in sess_row.keys() else '', 'summaryPath': sess_row['summary_path'] if 'summary_path' in sess_row.keys() else '', 'isCompleted': sess_row['is_completed'] if 'is_completed' in sess_row.keys() else 0, 'completionPhoto': sess_row['completion_photo'] if 'completion_photo' in sess_row.keys() else '', 'warning': sess_row['warning'] if 'warning' in sess_row.keys() else '', 'warningReason': sess_row['warning_reason'] if 'warning_reason' in sess_row.keys() else '', 'locationLat': sess_row['location_lat'] if 'location_lat' in sess_row.keys() else 0, 'locationLng': sess_row['location_lng'] if 'location_lng' in sess_row.keys() else 0, 'checkinCount': conn.execute('SELECT COUNT(*) as c FROM checkin_records WHERE session_id=?', (sid,)).fetchone()['c'], 'teacherCheckins': [{'teacherUserId': t['teacher_user_id'], 'status': t['status']} for t in teacher_checkins], 'time': local_time(sess_row['created_at'])}
            return jsonify({'success': True, 'data': result})
        if request.method == 'DELETE':
            conn.execute('DELETE FROM checkin_records WHERE session_id=?', (sid,))
            conn.execute('DELETE FROM checkin_sessions WHERE id=?', (sid,))
            conn.commit()
        else:
            action = (request.json or {}).get('action', 'close')
            if action == 'close':
                completion_photo = (request.json or {}).get('completionPhoto', '').strip()
                summary_text = (request.json or {}).get('summaryText', '').strip()
                summary_path = (request.json or {}).get('summaryPath', '').strip()
                if not completion_photo and not summary_text and not summary_path:
                    return jsonify({'error': '请至少上传活动照片或活动总结'}), 400
                if not summary_text and summary_path:
                    summary_text = '(活动总结文件已上传)'
                if not completion_photo and sess_row['completion_photo']:
                    completion_photo = sess_row['completion_photo']
                if not summary_text and not summary_path:
                    if sess_row['summary_text']:
                        summary_text = sess_row['summary_text']
                    if sess_row['summary_path']:
                        summary_path = sess_row['summary_path']
                is_completed = 1 if (completion_photo and (summary_text or summary_path)) else 0
                conn.execute('UPDATE checkin_sessions SET status="closed", closed_at=CURRENT_TIMESTAMP, completion_photo=?, summary_path=?, summary_text=?, is_completed=? WHERE id=?', (completion_photo, summary_path, summary_text, is_completed, sid))
            elif action == 'reopen':
                conn.execute('UPDATE checkin_sessions SET status="open", closed_at=NULL, is_completed=0 WHERE id=?', (sid,))
            conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/upload-activity-plan', methods=['POST'])
def upload_activity_plan():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '请选择文件'}), 400
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    if ext not in ('doc', 'docx', 'pdf'):
        return jsonify({'error': '仅支持Word(.doc/.docx)或PDF文件'}), 400
    import os, time, re
    filename = f'{user["id"]}_{int(time.time())}_{f.filename}'
    key = 'activity_plans/' + filename
    storage.save(f, key)
    filepath = storage.get_path(key) or ''
    text_content = ''
    extract_ok = False
    extract_method = ''
    warning = ''
    try:
        if ext == 'pdf':
            try:
                import fitz
                doc = fitz.open(filepath)
                for page in doc:
                    text_content += page.get_text() + '\n'
                doc.close()
                extract_ok = True
                extract_method = 'PyMuPDF'
            except ImportError:
                try:
                    import subprocess
                    result = subprocess.run(['python', '-m', 'pdfminer.tools.pdf2txt', filepath], capture_output=True, text=True, timeout=15)
                    text_content = result.stdout if result.stdout else ''
                    if text_content:
                        extract_ok = True
                        extract_method = 'pdfminer'
                except:
                    pass
            if extract_ok and not text_content.strip():
                warning = 'PDF文件可能为扫描件（图片型PDF），无法提取文字内容，请上传文字版PDF或Word文件'
                extract_ok = False
        elif ext in ('doc', 'docx'):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(filepath)
                for para in doc.paragraphs:
                    text_content += para.text + '\n'
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + '\n'
                for section in doc.sections:
                    header = section.header
                    if header and not header.is_linked_to_previous:
                        for para in header.paragraphs:
                            if para.text.strip():
                                text_content += para.text + '\n'
                    footer = section.footer
                    if footer and not footer.is_linked_to_previous:
                        for para in footer.paragraphs:
                            if para.text.strip():
                                text_content += para.text + '\n'
                extract_ok = True
                extract_method = 'python-docx'
            except ImportError:
                try:
                    import subprocess
                    result = subprocess.run(['python', '-c', 'import sys;from docx import Document;doc=Document(sys.argv[1]);print("\\n".join(p.text for p in doc.paragraphs))', filepath], capture_output=True, text=True, timeout=15)
                    text_content = result.stdout if result.stdout else ''
                    if text_content:
                        extract_ok = True
                        extract_method = 'python-docx(subprocess)'
                except:
                    pass
    except Exception as e:
        pass
    if text_content.startswith('[PDF') or text_content.startswith('[Word') or text_content.startswith('[文件'):
        text_content = ''
        extract_ok = False
    clean_text = re.sub(r'\s+', '', text_content)
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff\u3400-\u4dbf]', clean_text))
    english_words = len(re.findall(r'[a-zA-Z]+', text_content))
    char_count = chinese_chars + english_words
    is_valid = char_count >= 200
    if not is_valid and not warning:
        if ext == 'pdf' and extract_ok:
            warning = 'PDF文件提取到{}字，不足200字，活动计划无效。请上传内容更丰富的文件'.format(char_count)
        elif ext == 'pdf' and not extract_ok:
            warning = 'PDF文件无法提取文字，可能为扫描件，请上传文字版PDF或Word文件'
        elif ext in ('doc', 'docx') and extract_ok:
            warning = 'Word文件提取到{}字，不足200字，活动计划无效。请上传内容更丰富的文件'.format(char_count)
        elif ext in ('doc', 'docx') and not extract_ok:
            warning = 'Word文件解析失败，请确保文件格式正确'
    return jsonify({'success': True, 'text': text_content, 'charCount': char_count, 'filePath': filepath, 'fileName': f.filename, 'valid': is_valid, 'extractMethod': extract_method, 'warning': warning})


@app.route('/api/upload-activity-photo', methods=['POST'])
def upload_activity_photo():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if 'photo' not in request.files:
        return jsonify({'error': '请上传照片'}), 400
    f = request.files['photo']
    if not f.filename:
        return jsonify({'error': '请选择照片'}), 400
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    if ext not in ('jpg', 'jpeg', 'png', 'webp', 'gif'):
        return jsonify({'error': '仅支持图片文件'}), 400
    import os
    upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'activity_photos')
    os.makedirs(upload_dir, exist_ok=True)
    import time
    filename = f'{user["id"]}_{int(time.time())}.{ext}'
    filepath = os.path.join(upload_dir, filename)
    f.save(filepath)
    return jsonify({'success': True, 'path': filepath, 'url': '/api/activity-photo/' + filename})


@app.route('/api/upload-activity-summary', methods=['POST'])
def upload_activity_summary():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '请选择文件'}), 400
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    if ext not in ('doc', 'docx', 'pdf'):
        return jsonify({'error': '仅支持Word(.doc/.docx)或PDF文件'}), 400
    import os
    import time
    filename = f'{user["id"]}_{int(time.time())}_{f.filename}'
    key = 'activity_summaries/' + filename
    storage.save(f, key)
    filepath = storage.get_path(key) or ''
    text_content = ''
    try:
        if ext == 'pdf':
            try:
                import fitz
                doc = fitz.open(filepath)
                for page in doc:
                    text_content += page.get_text()
                doc.close()
            except ImportError:
                text_content = ''
        elif ext in ('doc', 'docx'):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(filepath)
                for para in doc.paragraphs:
                    text_content += para.text + '\n'
            except:
                text_content = ''
    except:
        text_content = ''
    return jsonify({'success': True, 'text': text_content, 'filePath': filepath, 'fileName': f.filename})


@app.route('/api/activity-photo/<filename>')
def serve_activity_photo(filename):
    import os
    upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'activity_photos')
    return send_from_directory(upload_dir, filename)


@app.route('/api/activity-summary-file/<path:storage_key>')
def serve_activity_summary_file(storage_key):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    if user['role'] not in ('admin', 'user'):
        return jsonify({'error': '无权限'}), 403
    path = storage.get_path(storage_key)
    if not path:
        url = storage.get_url(storage_key)
        if url:
            return jsonify({'url': url})
        return jsonify({'error': '文件不存在'}), 404
    return send_file(path, as_attachment=True)


@app.route('/api/activity-plan-file/<path:storage_key>')
def serve_activity_plan_file(storage_key):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    if user['role'] not in ('admin', 'user', 'teacher'):
        return jsonify({'error': '无权限'}), 403
    path = storage.get_path(storage_key)
    if not path:
        url = storage.get_url(storage_key)
        if url:
            return jsonify({'url': url})
        return jsonify({'error': '文件不存在'}), 404
    return send_file(path, as_attachment=True)


@app.route('/api/activity-plan-text/<int:session_id>')
def download_activity_plan_text(session_id):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT plan_text, activity_name FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
    finally:
        conn.close()
    if not sess or not sess['plan_text']:
        return jsonify({'error': '活动计划不存在'}), 404
    output = BytesIO()
    output.write(sess['plan_text'].encode('utf-8'))
    output.seek(0)
    fname = (sess['activity_name'] or '活动计划') + '.txt'
    return send_file(output, as_attachment=True, download_name=fname, mimetype='text/plain')


@app.route('/api/activity-summary-text/<int:session_id>')
def download_activity_summary_text(session_id):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT summary_text, activity_name FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
    finally:
        conn.close()
    if not sess or not sess['summary_text']:
        return jsonify({'error': '活动总结不存在'}), 404
    text = sess['summary_text']
    if text.startswith('(活动总结文件已上传)'):
        text = '活动总结已通过文件形式上传，请下载活动总结文件查看。'
    output = BytesIO()
    output.write(text.encode('utf-8'))
    output.seek(0)
    fname = (sess['activity_name'] or '活动总结') + '_总结.txt'
    return send_file(output, as_attachment=True, download_name=fname, mimetype='text/plain')


@app.route('/api/club-teachers-list')
def get_club_teachers_list():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    club = request.args.get('club', '')
    if not club and user:
        club = user.get('club_name', '')
    if not club:
        return jsonify({'success': True, 'data': []})
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, teacher_name, introduction FROM club_teachers WHERE club_name=?', (club,)).fetchall()
    finally:
        conn.close()
    teacher_list = [{'id': r['id'], 'name': r['teacher_name'], 'introduction': r['introduction'] or ''} for r in rows]
    base_teacher = get_teacher_name_from_sources(club)
    if base_teacher:
        for tn in base_teacher.split('、'):
            tn = tn.strip()
            if tn and not any(t['name'] == tn for t in teacher_list):
                teacher_list.append({'id': 0, 'name': tn, 'introduction': ''})
    return jsonify({'success': True, 'data': teacher_list})


@app.route('/api/checkin', methods=['POST'])
def student_checkin():
    data = request.json or {}
    code = data.get('code', '').strip()
    student_name = data.get('studentName', '').strip()
    student_class = data.get('studentClass', '').strip()
    student_id = data.get('studentId', '').strip()
    college = data.get('college', '').strip() or data.get('studentCollege', '').strip()
    method = data.get('method', 'code')
    if not code or not student_name:
        return jsonify({'error': '请输入签到码和姓名'}), 400
    conn = db.get_conn()
    try:
        sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE checkin_code=?', (code,)).fetchone()
        if not sess_row:
            return jsonify({'error': '签到码无效'}), 404
        if sess_row['status'] == 'open':
            end_time = sess_row['end_time'] if 'end_time' in sess_row.keys() else ''
            if end_time:
                try:
                    et = datetime.strptime(end_time.replace('T', ' ')[:16], '%Y-%m-%d %H:%M')
                    if cn_now() > et:
                        conn.execute("UPDATE checkin_sessions SET status='closed', closed_at=CURRENT_TIMESTAMP WHERE id=?", (sess_row['id'],))
                        conn.commit()
                        sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (sess_row['id'],)).fetchone()
                except:
                    pass
        if sess_row['status'] != 'open':
            return jsonify({'error': '签到已结束'}), 400
        start_time = sess_row['start_time'] if 'start_time' in sess_row.keys() else ''
        end_time = sess_row['end_time'] if 'end_time' in sess_row.keys() else ''
        if start_time and end_time:
            now = cn_now()
            try:
                st = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
                et = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')
                if now < st or now > et:
                    return jsonify({'error': '当前不在活动时间内，无法签到'}), 400
            except:
                pass
        existing = conn.execute('SELECT id FROM checkin_records WHERE session_id=? AND student_name=? AND student_id=? AND student_class=?', (sess_row['id'], student_name, student_id, student_class)).fetchone()
        if existing:
            return jsonify({'error': '您已签到过'}), 400
        conn.execute('INSERT INTO checkin_records (session_id, club_name, student_name, student_class, student_id, college, checkin_method) VALUES (?, ?, ?, ?, ?, ?, ?)', (sess_row['id'], sess_row['club_name'], student_name, student_class, student_id, college, method))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'activityName': sess_row['activity_name'], 'clubName': sess_row['club_name']})


@app.route('/api/checkin-records/<int:sid>')
def get_checkin_records(sid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (sid,)).fetchone()
        if not sess_row:
            return jsonify({'error': '签到会话不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != sess_row['club_name']:
            if user['role'] == 'student' and is_cadre_of_club(user['id'], sess_row['club_name']):
                pass
            else:
                return jsonify({'error': '无权限'}), 403
        rows = conn.execute('SELECT * FROM checkin_records WHERE session_id=? ORDER BY created_at', (sid,)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'studentName': r['student_name'], 'studentClass': r['student_class'], 'studentId': r['student_id'], 'method': r['checkin_method'], 'time': local_time(r['created_at'])} for r in rows]})


@app.route('/api/checkin-records-by-session/<int:sid>')
def get_checkin_by_session(sid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher', 'student'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (sid,)).fetchone()
        if not sess_row:
            return jsonify({'error': '签到会话不存在'}), 404
        if user['role'] == 'teacher':
            tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], sess_row['club_name'])).fetchone()
            if not tc_row:
                return jsonify({'error': '无权限'}), 403
        elif user['role'] == 'student':
            if not is_cadre_of_club(user['id'], sess_row['club_name']):
                return jsonify({'error': '无权限'}), 403
        elif user['role'] != 'admin' and user['club_name'] != sess_row['club_name']:
            return jsonify({'error': '无权限'}), 403
        rows = conn.execute('SELECT * FROM checkin_records WHERE session_id=? ORDER BY created_at', (sid,)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'studentName': r['student_name'], 'studentClass': r['student_class'], 'studentId': r['student_id'], 'studentCollege': r['college'] if 'college' in r.keys() else '', 'method': r['checkin_method'], 'activityName': sess_row['activity_name'], 'time': local_time(r['created_at'])} for r in rows]})


@app.route('/api/checkin-records-by-club/<club_name>')
def get_checkin_by_club(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher', 'student'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] == 'user' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    if user['role'] == 'student' and not is_cadre_of_club(user['id'], club_name):
        return jsonify({'error': '无权限'}), 403
    date = request.args.get('date', '')
    conn = db.get_conn()
    try:
        if date:
            rows = conn.execute("SELECT cr.*, cs.activity_name, cs.checkin_code FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.club_name=? AND DATE(cr.created_at)=? ORDER BY cr.created_at", (club_name, date)).fetchall()
        else:
            rows = conn.execute("SELECT cr.*, cs.activity_name, cs.checkin_code FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.club_name=? ORDER BY cr.created_at DESC", (club_name,)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'studentName': r['student_name'], 'studentClass': r['student_class'], 'studentId': r['student_id'], 'method': r['checkin_method'], 'activityName': r['activity_name'], 'time': local_time(r['created_at'])} for r in rows]})


@app.route('/api/export-checkin/<club_name>')
def export_checkin(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] != 'admin' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    date = request.args.get('date', '')
    from openpyxl import Workbook
    conn = db.get_conn()
    try:
        if date:
            rows = conn.execute("SELECT cr.*, cs.activity_name FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.club_name=? AND DATE(cr.created_at)=? ORDER BY cr.created_at", (club_name, date)).fetchall()
        else:
            rows = conn.execute("SELECT cr.*, cs.activity_name FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.club_name=? ORDER BY cr.created_at DESC", (club_name,)).fetchall()
    finally:
        conn.close()
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = '签到记录'
    ws.append(['序号', '姓名', '学院', '班级', '学号', '签到方式', '活动名称', '签到时间'])
    for i, r in enumerate(rows, 1):
        ws.append([i, r['student_name'], (r['college'] if 'college' in r.keys() else '') or '', r['student_class'], r['student_id'], r['checkin_method'], r['activity_name'], r['created_at']])
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'{club_name}_签到记录.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/checkin-stats/<club_name>')
def checkin_stats(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] == 'user' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute("SELECT cr.student_name, MAX(cr.student_class) as student_class, GROUP_CONCAT(DISTINCT cr.student_id) as student_id, COUNT(DISTINCT cr.session_id) as activity_count, GROUP_CONCAT(DISTINCT cs.activity_name) as activities FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.club_name=? GROUP BY cr.student_name ORDER BY activity_count DESC", (club_name,)).fetchall()
        total_sessions = conn.execute("SELECT COUNT(DISTINCT id) as c FROM checkin_sessions WHERE club_name=?", (club_name,)).fetchone()['c']
    finally:
        conn.close()
    return jsonify({'success': True, 'totalSessions': total_sessions, 'data': [{'studentName': r['student_name'], 'studentClass': r['student_class'], 'studentId': r['student_id'], 'activityCount': r['activity_count'], 'activities': r['activities'] or ''} for r in rows]})


@app.route('/api/checkin-stats-detail/<club_name>/<student_name>')
def checkin_stats_detail(club_name, student_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] != 'admin' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute("SELECT cr.*, cs.activity_name, cs.checkin_code FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.club_name=? AND cr.student_name=? ORDER BY cr.created_at DESC", (club_name, student_name)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'studentName': r['student_name'], 'studentClass': r['student_class'], 'studentId': r['student_id'], 'method': r['checkin_method'], 'activityName': r['activity_name'], 'time': local_time(r['created_at'])} for r in rows]})


@app.route('/api/export-checkin-stats/<club_name>')
def export_checkin_stats(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] != 'admin' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    from openpyxl import Workbook
    conn = db.get_conn()
    try:
        rows = conn.execute("SELECT cr.student_name, MAX(cr.student_class) as student_class, GROUP_CONCAT(DISTINCT cr.student_id) as student_id, MAX(cr.college) as college, COUNT(DISTINCT cr.session_id) as activity_count, GROUP_CONCAT(DISTINCT cs.activity_name) as activities FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.club_name=? GROUP BY cr.student_name ORDER BY activity_count DESC", (club_name,)).fetchall()
        total_sessions = conn.execute("SELECT COUNT(DISTINCT id) as c FROM checkin_sessions WHERE club_name=?", (club_name,)).fetchone()['c']
    finally:
        conn.close()
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = '社员活动统计'
    ws.append([f'社团：{club_name}', f'总活动场次：{total_sessions}'])
    ws.append(['序号', '姓名', '学院', '班级', '学号', '参与活动次数', '参与的活动'])
    for i, r in enumerate(rows, 1):
        ws.append([i, r['student_name'], (r['college'] if 'college' in r.keys() else '') or '', r['student_class'], r['student_id'], r['activity_count'], r['activities'] or ''])
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'{club_name}_社员活动统计.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/add-checkin-record', methods=['POST'])
def add_checkin_record():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    session_id = data.get('sessionId')
    student_name = data.get('studentName', '').strip()
    student_class = data.get('studentClass', '').strip()
    student_id = data.get('studentId', '').strip()
    college = data.get('studentCollege', '').strip() or data.get('college', '').strip()
    method = data.get('method', 'manual')
    if not session_id or not student_name:
        return jsonify({'error': '缺少必要参数'}), 400
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
        if not sess:
            return jsonify({'error': '签到会话不存在'}), 404
        if user['role'] == 'teacher':
            tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], sess['club_name'])).fetchone()
            if not tc_row:
                return jsonify({'error': '无权限'}), 403
        elif user['role'] != 'admin' and user['club_name'] != sess['club_name']:
            return jsonify({'error': '无权限'}), 403
        existing = conn.execute('SELECT id FROM checkin_records WHERE session_id=? AND student_name=? AND student_id=? AND student_class=?', (session_id, student_name, student_id, student_class)).fetchone()
        if existing:
            return jsonify({'error': '该学生已签到'}), 400
        conn.execute('INSERT INTO checkin_records (session_id, club_name, student_name, student_class, student_id, college, checkin_method) VALUES (?, ?, ?, ?, ?, ?, ?)', (session_id, sess['club_name'], student_name, student_class, student_id, college, method))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/checkin-record/<int:rid>', methods=['DELETE'])
def delete_checkin_record(rid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        record = conn.execute('SELECT * FROM checkin_records WHERE id=?', (rid,)).fetchone()
        if not record:
            return jsonify({'error': '记录不存在'}), 404
        sess = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (record['session_id'],)).fetchone()
        if not sess:
            return jsonify({'error': '会话不存在'}), 404
        if user['role'] == 'teacher':
            tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], sess['club_name'])).fetchone()
            if not tc_row:
                return jsonify({'error': '无权限'}), 403
        elif user['role'] != 'admin' and user['club_name'] != sess['club_name']:
            return jsonify({'error': '无权限'}), 403
        conn.execute('DELETE FROM checkin_records WHERE id=?', (rid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/export-checkin-session/<int:sid>')
def export_checkin_session(sid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student'):
        return jsonify({'error': '请先登录'}), 401
    from openpyxl import Workbook
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (sid,)).fetchone()
        if not sess:
            return jsonify({'error': '会话不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != sess['club_name']:
            if user['role'] == 'student' and is_cadre_of_club(user['id'], sess['club_name']):
                pass
            else:
                return jsonify({'error': '无权限'}), 403
        rows = conn.execute("SELECT * FROM checkin_records WHERE session_id=? ORDER BY created_at", (sid,)).fetchall()
    finally:
        conn.close()
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = '签到记录'
    ws.append([f'活动：{sess["activity_name"] or "未命名"}', f'签到码：{sess["checkin_code"]}', f'地点：{sess["location_name"] or ""}'])
    ws.append(['序号', '姓名', '学院', '班级', '学号', '签到方式', '签到时间'])
    for i, r in enumerate(rows, 1):
        method_map = {'code': '签到码', 'qrcode': '扫码', 'location': '定位', 'manual': '手动添加'}
        ws.append([i, r['student_name'], (r['college'] if 'college' in r.keys() else '') or '', r['student_class'], r['student_id'], method_map.get(r['checkin_method'], r['checkin_method']), r['created_at']])
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'{sess["activity_name"] or "活动"}_签到记录.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/reverse-geocode')
def reverse_geocode():
    lat = request.args.get('lat', '0')
    lng = request.args.get('lng', '0')
    try:
        lat_f = float(lat)
        lng_f = float(lng)
    except ValueError:
        return jsonify({'success': False, 'address': ''})
    if lat_f == 0 and lng_f == 0:
        return jsonify({'success': False, 'address': ''})
    # 优先使用高德逆地理编码（国内服务更稳定）
    try:
        import urllib.request
        url = f'https://restapi.amap.com/v3/geocode/regeo?key=2454fd77342fd22574cbe266eb5c647b&location={lng_f},{lat_f}&extensions=all&output=JSON'
        req = urllib.request.Request(url, headers={'User-Agent': 'ClubStats/1.0'})
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            if data.get('status') == '1' and data.get('regeocode'):
                addr = data['regeocode'].get('formatted_address', '')
                # 提取更详细的周边信息
                comp = data['regeocode'].get('addressComponent', {})
                neighborhood = comp.get('neighborhood', {})
                if isinstance(neighborhood, dict) and neighborhood.get('name'):
                    addr += '（附近：' + neighborhood['name'] + '）'
                elif isinstance(neighborhood, str) and neighborhood:
                    addr += '（附近：' + neighborhood + '）'
                return jsonify({'success': True, 'address': addr})
    except Exception:
        pass
    # 备用：OpenStreetMap
    try:
        import urllib.request
        url2 = f'https://nominatim.openstreetmap.org/reverse?format=json&lat={lat_f}&lon={lng_f}&accept-language=zh'
        req2 = urllib.request.Request(url2, headers={'User-Agent': 'ClubStats/1.0'})
        with urllib.request.urlopen(req2, timeout=8) as resp2:
            data2 = json.loads(resp2.read().decode('utf-8'))
            address2 = data2.get('display_name', '')
            return jsonify({'success': True, 'address': address2})
    except Exception:
        pass
    return jsonify({'success': True, 'address': f'纬度:{lat_f:.4f},经度:{lng_f:.4f}'})


@app.route('/api/location-checkin', methods=['POST'])
def location_checkin():
    import math
    data = request.json or {}
    lat = data.get('lat', 0)
    lng = data.get('lng', 0)
    session_id = data.get('sessionId', 0)
    code = data.get('code', '').strip()
    student_name = data.get('studentName', '').strip()
    student_class = data.get('studentClass', '').strip()
    student_id = data.get('studentId', '').strip()
    college = data.get('college', '').strip() or data.get('studentCollege', '').strip()
    user = get_current_user()
    try:
        lat_f = float(lat)
        lng_f = float(lng)
    except (ValueError, TypeError):
        return jsonify({'error': '经纬度格式错误'}), 400
    if lat_f == 0 and lng_f == 0:
        return jsonify({'error': '无法获取位置信息，请确保已开启定位权限'}), 400

    # 通过签到码查找session
    conn = db.get_conn()
    sess_row = None
    try:
        if code:
            sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE checkin_code=?', (code,)).fetchone()
        if not sess_row and session_id:
            sess_row = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
    except Exception:
        pass
    if sess_row:
        session_id = sess_row['id']
        # 检查签到是否已结束
        if sess_row['status'] != 'open':
            conn.close()
            return jsonify({'error': '签到已结束'}), 400
        # 检查是否在活动时间内
        start_time = sess_row['start_time'] if 'start_time' in sess_row.keys() else ''
        end_time = sess_row['end_time'] if 'end_time' in sess_row.keys() else ''
        if start_time and end_time:
            now = cn_now()
            try:
                st = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
                et = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')
                if now < st or now > et:
                    conn.close()
                    return jsonify({'error': '当前不在活动时间内，无法签到'}), 400
            except:
                pass
        # 距离校验：如果session设置了签到位置，检查距离是否在范围内
        sess_lat = sess_row['location_lat'] if 'location_lat' in sess_row.keys() else 0
        sess_lng = sess_row['location_lng'] if 'location_lng' in sess_row.keys() else 0
        if sess_lat and sess_lng:
            try:
                sess_lat_f = float(sess_lat)
                sess_lng_f = float(sess_lng)
                # Haversine公式计算距离（米）
                R = 6371000
                dlat = math.radians(lat_f - sess_lat_f)
                dlng = math.radians(lng_f - sess_lng_f)
                a = math.sin(dlat/2)**2 + math.cos(math.radians(sess_lat_f)) * math.cos(math.radians(lat_f)) * math.sin(dlng/2)**2
                c = 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))
                distance = R * c
                print(f'[定位签到] 学生位置: ({lat_f}, {lng_f}), 签到位置: ({sess_lat_f}, {sess_lng_f}), 距离: {distance:.1f}米')
                if distance > 500:
                    conn.close()
                    return jsonify({'error': f'距离活动地点{distance:.0f}米，超出500米签到范围'}), 400
            except (ValueError, TypeError):
                pass
        # 检查是否已签到
        sname = student_name or (user['username'] if user else '')
        sclas = student_class or ''
        ssid = student_id or ''
        existing = conn.execute('SELECT id FROM checkin_records WHERE session_id=? AND student_name=? AND student_id=? AND student_class=?',
                               (session_id, sname, ssid, sclas)).fetchone()
        if existing:
            conn.close()
            return jsonify({'error': '您已签到过'}), 400

    address = ''
    try:
        import urllib.request
        url = f'https://restapi.amap.com/v3/geocode/regeo?key=2454fd77342fd22574cbe266eb5c647b&location={lng_f},{lat_f}&extensions=base&output=JSON'
        req = urllib.request.Request(url, headers={'User-Agent': 'ClubStats/1.0'})
        with urllib.request.urlopen(req, timeout=5) as resp:
            rdata = json.loads(resp.read().decode('utf-8'))
            if rdata.get('status') == '1' and rdata.get('regeocode'):
                address = rdata['regeocode'].get('formatted_address', '')
    except Exception:
        address = f'纬度:{lat_f:.4f},经度:{lng_f:.4f}'
    try:
        if user:
            conn.execute('INSERT INTO location_checkins (user_id, username, role, club_name, session_id, lat, lng, address) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                         (user['id'], user['username'], user['role'], user.get('club_name', ''), session_id, lat_f, lng_f, address))
        else:
            conn.execute('INSERT INTO location_checkins (user_id, username, role, club_name, session_id, lat, lng, address) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                         (0, student_name, 'student', '', session_id, lat_f, lng_f, address))
        if session_id and sess_row:
            sname = student_name or (user['username'] if user else '')
            sclas = student_class or ''
            ssid = student_id or ''
            scol = college or ''
            conn.execute('INSERT INTO checkin_records (session_id, club_name, student_name, student_class, student_id, college, checkin_method) VALUES (?, ?, ?, ?, ?, ?, ?)',
                        (session_id, sess_row['club_name'], sname, sclas, ssid, scol, 'location'))
        conn.commit()
    finally:
        conn.close()
    result = {'success': True, 'address': address, 'lat': lat_f, 'lng': lng_f, 'activityName': sess_row['activity_name'] if sess_row else '', 'clubName': sess_row['club_name'] if sess_row else ''}
    if sess_row:
        sess_lat = sess_row['location_lat'] if 'location_lat' in sess_row.keys() else 0
        sess_lng = sess_row['location_lng'] if 'location_lng' in sess_row.keys() else 0
        if sess_lat and sess_lng:
            result['sessionLat'] = float(sess_lat)
            result['sessionLng'] = float(sess_lng)
            R = 6371000
            dlat = math.radians(lat_f - float(sess_lat))
            dlng = math.radians(lng_f - float(sess_lng))
            a = math.sin(dlat/2)**2 + math.cos(math.radians(float(sess_lat))) * math.cos(math.radians(lat_f)) * math.sin(dlng/2)**2
            c = 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))
            result['distance'] = round(R * c)
    return jsonify(result)


@app.route('/api/location-checkin-records')
def get_location_checkin_records():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        if user['role'] == 'admin':
            rows = conn.execute('SELECT * FROM location_checkins ORDER BY checkin_time DESC LIMIT 200').fetchall()
        elif user['role'] == 'teacher':
            club_names = [r['club_name'] for r in conn.execute('SELECT club_name FROM teacher_clubs WHERE teacher_user_id=?', (user['id'],)).fetchall()]
            if club_names:
                placeholders = ','.join(['?'] * len(club_names))
                rows = conn.execute(f'SELECT * FROM location_checkins WHERE club_name IN ({placeholders}) ORDER BY checkin_time DESC LIMIT 200', club_names).fetchall()
            else:
                rows = []
        elif user['role'] == 'user' and user.get('club_name'):
            rows = conn.execute('SELECT * FROM location_checkins WHERE club_name=? ORDER BY checkin_time DESC LIMIT 200', (user['club_name'],)).fetchall()
        else:
            rows = conn.execute('SELECT * FROM location_checkins WHERE user_id=? ORDER BY checkin_time DESC LIMIT 200', (user['id'],)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [dict(r) for r in rows]})


@app.route('/api/ip-location')
def ip_location():
    ip_addr = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr) or request.remote_addr
    if ip_addr and ',' in ip_addr:
        ip_addr = ip_addr.split(',')[0].strip()
    if not ip_addr or ip_addr.startswith('127.') or ip_addr.startswith('192.168.') or ip_addr.startswith('10.') or ip_addr == '::1':
        try:
            import urllib.request
            pub_url = 'https://api.ipify.org?format=text'
            pub_req = urllib.request.Request(pub_url, headers={'User-Agent': 'ClubStats/1.0'})
            with urllib.request.urlopen(pub_req, timeout=5) as pub_resp:
                public_ip = pub_resp.read().decode('utf-8').strip()
                if public_ip:
                    ip_addr = public_ip
        except Exception:
            return jsonify({'success': False, 'error': '局域网IP无法通过IP定位，请尝试使用GPS定位'})
    try:
        import urllib.request
        url = f'http://ip-api.com/json/{ip_addr}?lang=zh-CN&fields=lat,lon,city,regionName,country'
        req = urllib.request.Request(url, headers={'User-Agent': 'ClubStats/1.0'})
        with urllib.request.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            if data.get('lat') and data.get('lon'):
                addr_parts = [data.get('country', ''), data.get('regionName', ''), data.get('city', '')]
                addr = ' '.join([p for p in addr_parts if p]) or f"{data['lat']},{data['lon']}"
                return jsonify({'success': True, 'lat': data['lat'], 'lon': data['lon'], 'address': f'（IP定位 - 约）{addr}'})
    except Exception:
        pass
    return jsonify({'success': False, 'error': 'IP定位服务不可用'})


@app.route('/api/teacher-checkin', methods=['POST'])
def teacher_checkin():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '请先以指导老师身份登录'}), 401
    data = request.json or {}
    session_id = data.get('sessionId')
    try:
        session_id = int(session_id)
    except (ValueError, TypeError):
        return jsonify({'error': '活动会话ID格式错误'}), 400
    lat = data.get('lat', 0)
    lng = data.get('lng', 0)
    address = data.get('address', '')
    if not session_id:
        return jsonify({'error': '缺少活动会话ID'}), 400
    lat_val = float(lat) if lat else 0
    lng_val = float(lng) if lng else 0
    if lat_val == 0 and lng_val == 0:
        address = '手动签到'
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
        if not sess:
            return jsonify({'error': '活动会话不存在'}), 404
        if sess['status'] != 'open':
            return jsonify({'error': '活动已结束，无法签到'}), 400
        tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], sess['club_name'])).fetchone()
        if not tc_row:
            return jsonify({'error': '您不是该社团的指导老师'}), 403
        existing = conn.execute('SELECT id, status FROM teacher_checkin_checkout WHERE session_id=? AND teacher_user_id=?', (session_id, user['id'])).fetchone()
        if existing:
            return jsonify({'error': '您已签到过'}), 400
        conn.execute('INSERT INTO teacher_checkin_checkout (session_id, teacher_user_id, club_name, checkin_time, checkin_lat, checkin_lng, checkin_address, status) VALUES (?, ?, ?, CURRENT_TIMESTAMP, ?, ?, ?, ?)', (session_id, user['id'], sess['club_name'], lat_val, lng_val, address, 'checked_in'))
        conn.commit()
        record = conn.execute('SELECT * FROM teacher_checkin_checkout WHERE session_id=? AND teacher_user_id=?', (session_id, user['id'])).fetchone()
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'record': {
            'id': record['id'], 'sessionId': record['session_id'], 'clubName': record['club_name'],
            'checkinTime': local_time(record['checkin_time']), 'status': record['status'],
            'checkinLat': record['checkin_lat'] or '', 'checkinLng': record['checkin_lng'] or '',
            'checkinAddress': record['checkin_address'] or ''
        }
    })


@app.route('/api/teacher-checkout', methods=['POST'])
def teacher_checkout():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '请先以指导老师身份登录'}), 401
    session_id = request.form.get('sessionId', '')
    try:
        session_id = int(session_id)
    except (ValueError, TypeError):
        return jsonify({'error': '活动会话ID格式错误'}), 400
    lat = request.form.get('lat', 0)
    lng = request.form.get('lng', 0)
    address = request.form.get('address', '')
    checkout_code = request.form.get('checkoutCode', '')
    if not session_id:
        return jsonify({'error': '缺少活动会话ID'}), 400
    lat_val = float(lat) if lat else 0
    lng_val = float(lng) if lng else 0
    if lat_val == 0 and lng_val == 0:
        address = '手动签到'
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
        if not sess:
            return jsonify({'error': '活动会话不存在'}), 404
        if checkout_code:
            if not sess['checkout_code'] or sess['checkout_code'] != checkout_code:
                return jsonify({'error': '签退码不正确'}), 400
        record = conn.execute('SELECT * FROM teacher_checkin_checkout WHERE session_id=? AND teacher_user_id=?', (session_id, user['id'])).fetchone()
        if not record:
            return jsonify({'error': '您尚未签到，无法签退'}), 400
        if record['status'] == 'checked_out':
            return jsonify({'error': '您已签退过'}), 400
        conn.execute('UPDATE teacher_checkin_checkout SET checkout_time=CURRENT_TIMESTAMP, checkout_lat=?, checkout_lng=?, checkout_address=?, checkout_method=?, checkout_code=?, status=? WHERE id=?',
                     (lat_val, lng_val, address, 'code_location' if checkout_code else 'location', checkout_code, 'checked_out', record['id']))
        conn.commit()
        updated = conn.execute('SELECT * FROM teacher_checkin_checkout WHERE id=?', (record['id'],)).fetchone()
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'record': {
            'id': updated['id'], 'sessionId': updated['session_id'], 'clubName': updated['club_name'],
            'checkinTime': local_time(updated['checkin_time']), 'checkoutTime': local_time(updated['checkout_time']),
            'status': updated['status'],
            'checkinLat': updated['checkin_lat'] or '', 'checkinLng': updated['checkin_lng'] or '',
            'checkinAddress': updated['checkin_address'] or '',
            'checkoutLat': updated['checkout_lat'] or '', 'checkoutLng': updated['checkout_lng'] or '',
            'checkoutAddress': updated['checkout_address'] or ''
        }
    })


@app.route('/api/teacher-checkin-status')
def teacher_checkin_status():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '请先以指导老师身份登录'}), 401
    session_id = request.args.get('sessionId', '')
    if not session_id:
        return jsonify({'error': '缺少sessionId参数'}), 400
    conn = db.get_conn()
    try:
        record = conn.execute('SELECT * FROM teacher_checkin_checkout WHERE session_id=? AND teacher_user_id=?', (session_id, user['id'])).fetchone()
    finally:
        conn.close()
    if not record:
        return jsonify({'success': True, 'status': 'none', 'record': None})
    return jsonify({'success': True, 'status': record['status'], 'record': {'id': record['id'], 'sessionId': record['session_id'], 'clubName': record['club_name'], 'checkinTime': local_time(record['checkin_time']), 'checkoutTime': local_time(record['checkout_time']), 'checkinLat': record['checkin_lat'], 'checkinLng': record['checkin_lng'], 'checkoutLat': record['checkout_lat'], 'checkoutLng': record['checkout_lng'], 'checkoutMethod': record['checkout_method'], 'status': record['status']}})


@app.route('/api/teacher-guidance-stats')
def teacher_guidance_stats():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '请先以指导老师身份登录'}), 401
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        clubs = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
        by_club = {}
        by_month = {}
        total_guidance = 0
        this_month = 0
        total_hours = 0
        import datetime
        now = datetime.datetime.now()
        this_month_str = now.strftime('%Y-%m')
        for c in clubs:
            club_name = c['club_name']
            date_where = ' AND 1=1'
            date_params = [user['id'], club_name, 'checked_out']
            if start_date:
                date_where += ' AND date(tcc.checkin_time)>=?'
                date_params.append(start_date)
            if end_date:
                date_where += ' AND date(tcc.checkin_time)<=?'
                date_params.append(end_date)
            rows = conn.execute(f'SELECT tcc.*, cs.activity_name FROM teacher_checkin_checkout tcc LEFT JOIN checkin_sessions cs ON tcc.session_id=cs.id WHERE tcc.teacher_user_id=? AND tcc.club_name=? AND tcc.status=?{date_where}', date_params).fetchall()
            count = len(rows)
            by_club[club_name] = count
            total_guidance += count
            for r in rows:
                if r['checkin_time'] and r['checkout_time']:
                    try:
                        ci = datetime.datetime.strptime(r['checkin_time'], '%Y-%m-%d %H:%M:%S')
                        co = datetime.datetime.strptime(r['checkout_time'], '%Y-%m-%d %H:%M:%S')
                        hours = max(0, (co - ci).total_seconds() / 3600)
                        total_hours += hours
                    except:
                        pass
                    month_key = r['checkin_time'][:7] if r['checkin_time'] else ''
                    if month_key:
                        by_month[month_key] = by_month.get(month_key, 0) + 1
                    if month_key == this_month_str:
                        this_month += 1
        checked_in_only = conn.execute('SELECT COUNT(*) as c FROM teacher_checkin_checkout WHERE teacher_user_id=? AND status=?', (user['id'], 'checked_in')).fetchone()['c']
    finally:
        conn.close()
    return jsonify({'success': True, 'data': {'totalGuidance': total_guidance, 'byClub': by_club, 'byMonth': by_month, 'thisMonth': this_month, 'totalHours': round(total_hours, 1)}})


@app.route('/api/teacher-guidance-records')
def teacher_guidance_records():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '请先以指导老师身份登录'}), 401
    club = request.args.get('club', '')
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        date_where = ''
        date_params = [user['id']]
        if club:
            date_where += ' AND tcc.club_name=?'
            date_params.append(club)
        if start_date:
            date_where += ' AND date(tcc.checkin_time)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(tcc.checkin_time)<=?'
            date_params.append(end_date)
        rows = conn.execute(f'SELECT tcc.*, cs.activity_name, cs.location_name FROM teacher_checkin_checkout tcc LEFT JOIN checkin_sessions cs ON tcc.session_id=cs.id WHERE tcc.teacher_user_id=?{date_where} ORDER BY tcc.checkin_time DESC', date_params).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'sessionId': r['session_id'], 'clubName': r['club_name'], 'activityName': r['activity_name'] or '管理员添加', 'locationName': r['location_name'] or '', 'checkinTime': local_time(r['checkin_time']), 'checkoutTime': local_time(r['checkout_time']), 'checkinLat': r['checkin_lat'], 'checkinLng': r['checkin_lng'], 'checkoutLat': r['checkout_lat'], 'checkoutLng': r['checkout_lng'], 'checkoutMethod': r['checkout_method'], 'status': r['status'], 'createdAt': local_time(r['created_at'])} for r in rows]})


@app.route('/api/all-teacher-guidance-stats')
def all_teacher_guidance_stats():
    user = get_current_user()
    if not user or user['role'] not in ('admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    club_filter = request.args.get('club', '')
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    date_where = ''
    date_params = []
    if start_date:
        date_where += ' AND date(tcc.checkin_time)>=?'
        date_params.append(start_date)
    if end_date:
        date_where += ' AND date(tcc.checkin_time)<=?'
        date_params.append(end_date)
    conn = db.get_conn()
    try:
        if club_filter:
            teachers = conn.execute('SELECT DISTINCT u.id, u.username, tp.real_name, tp.avatar_path FROM users u LEFT JOIN teacher_profiles tp ON u.id=tp.user_id WHERE u.role="teacher" AND u.id IN (SELECT user_id FROM teacher_clubs WHERE club_name=?)', (club_filter,)).fetchall()
            result = []
            for t in teachers:
                total_count = conn.execute(f'SELECT COUNT(*) as c FROM teacher_checkin_checkout tcc WHERE tcc.teacher_user_id=? AND tcc.club_name=? AND tcc.status="checked_out"{date_where}', [t['id'], club_filter] + date_params).fetchone()['c']
                avatar = ''
                if 'avatar_path' in t.keys() and t['avatar_path']:
                    avatar = t['avatar_path']
                result.append({'userId': t['id'], 'name': t['real_name'] or t['username'], 'count': total_count, 'avatar': avatar})
            return jsonify({'success': True, 'data': result})
        else:
            all_clubs = conn.execute('SELECT DISTINCT club_name FROM teacher_clubs ORDER BY club_name').fetchall()
            result = {}
            for c in all_clubs:
                cn = c['club_name']
                guiding_unit = ''
                gp = conn.execute('SELECT guiding_unit FROM club_profiles WHERE club_name=?', (cn,)).fetchone()
                if gp and gp['guiding_unit']:
                    guiding_unit = gp['guiding_unit']
                teachers = conn.execute('SELECT DISTINCT u.id, u.username, tp.real_name FROM users u LEFT JOIN teacher_profiles tp ON u.id=tp.user_id WHERE u.role="teacher" AND u.id IN (SELECT user_id FROM teacher_clubs WHERE club_name=?)', (cn,)).fetchall()
                t_list = []
                for t in teachers:
                    total_count = conn.execute(f'SELECT COUNT(*) as c FROM teacher_checkin_checkout tcc WHERE tcc.teacher_user_id=? AND tcc.club_name=? AND tcc.status="checked_out"{date_where}', [t['id'], cn] + date_params).fetchone()['c']
                    records = conn.execute(f'SELECT tcc.id, tcc.teacher_user_id, tcc.checkin_time, tcc.checkout_time, tcc.status, tcc.checkin_lat, tcc.checkin_lng, tcc.checkout_lat, tcc.checkout_lng, tcc.checkin_address, tcc.checkout_address FROM teacher_checkin_checkout tcc WHERE tcc.teacher_user_id=? AND tcc.club_name=? AND tcc.status="checked_out"{date_where} ORDER BY tcc.checkin_time DESC', [t['id'], cn] + date_params).fetchall()
                    detail_list = []
                    for r in records:
                        duration = ''
                        if r['checkin_time'] and r['checkout_time']:
                            try:
                                import datetime as _dt
                                ci = _dt.datetime.strptime(str(r['checkin_time']), '%Y-%m-%d %H:%M:%S')
                                co = _dt.datetime.strptime(str(r['checkout_time']), '%Y-%m-%d %H:%M:%S')
                                total_sec = int((co - ci).total_seconds())
                                hours = total_sec // 3600
                                mins = (total_sec % 3600) // 60
                                secs = total_sec % 60
                                if hours > 0:
                                    duration = str(hours) + '小时' + (str(mins) + '分钟' if mins > 0 else '')
                                elif mins > 0:
                                    duration = str(mins) + '分钟'
                                else:
                                    duration = str(secs) + '秒'
                            except:
                                duration = ''
                        detail_list.append({
                            'id': r['id'],
                            'teacherUserId': r['teacher_user_id'],
                            'clubName': cn,
                            'checkinTime': local_time(r['checkin_time']),
                            'checkoutTime': local_time(r['checkout_time']),
                            'status': r['status'],
                            'duration': duration,
                            'checkinLat': r['checkin_lat'] if 'checkin_lat' in r.keys() else '',
                            'checkinLng': r['checkin_lng'] if 'checkin_lng' in r.keys() else '',
                            'checkinAddress': r['checkin_address'] if 'checkin_address' in r.keys() else '',
                            'checkoutLat': r['checkout_lat'] if 'checkout_lat' in r.keys() else '',
                            'checkoutLng': r['checkout_lng'] if 'checkout_lng' in r.keys() else '',
                            'checkoutAddress': r['checkout_address'] if 'checkout_address' in r.keys() else ''
                        })
                    t_list.append({'userId': t['id'], 'name': t['real_name'] or t['username'], 'count': total_count, 'records': detail_list})
                t_list.sort(key=lambda x: x['count'], reverse=True)
                result[cn] = {'guidingUnit': guiding_unit, 'teachers': t_list}
            return jsonify({'success': True, 'data': result})
    finally:
        conn.close()


@app.route('/api/admin/teacher-guidance', methods=['POST'])
def admin_add_teacher_guidance():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    teacher_user_id = data.get('teacherUserId')
    club_name = data.get('clubName', '').strip()
    checkin_time = data.get('checkinTime', '').strip()
    checkout_time = data.get('checkoutTime', '').strip()
    checkin_address = data.get('checkinAddress', '').strip()
    checkout_address = data.get('checkoutAddress', '').strip()
    if not teacher_user_id or not club_name or not checkin_time or not checkout_time:
        return jsonify({'error': '请填写完整信息'}), 400
    conn = db.get_conn()
    try:
        conn.execute('''INSERT INTO teacher_checkin_checkout (session_id, teacher_user_id, club_name, checkin_time, checkout_time, status, checkin_address, checkout_address)
            VALUES (0, ?, ?, ?, ?, 'checked_out', ?, ?)''',
            (teacher_user_id, club_name, checkin_time, checkout_time, checkin_address, checkout_address))
        conn.commit()
        new_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    except Exception as e:
        conn.close()
        return jsonify({'error': '添加失败：' + str(e)}), 500
    conn.close()
    return jsonify({'success': True, 'id': new_id})


@app.route('/api/admin/teacher-guidance/<int:gid>', methods=['PUT'])
def admin_edit_teacher_guidance(gid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    checkin_time = data.get('checkinTime', '').strip()
    checkout_time = data.get('checkoutTime', '').strip()
    checkin_address = data.get('checkinAddress', '').strip()
    checkout_address = data.get('checkoutAddress', '').strip()
    if not checkin_time or not checkout_time:
        return jsonify({'error': '请填写完整信息'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM teacher_checkin_checkout WHERE id=?', (gid,)).fetchone()
        if not existing:
            conn.close()
            return jsonify({'error': '记录不存在'}), 404
        conn.execute('UPDATE teacher_checkin_checkout SET checkin_time=?, checkout_time=?, checkin_address=?, checkout_address=? WHERE id=?',
                     (checkin_time, checkout_time, checkin_address, checkout_address, gid))
        conn.commit()
    except Exception as e:
        conn.close()
        return jsonify({'error': '修改失败：' + str(e)}), 500
    conn.close()
    return jsonify({'success': True})


@app.route('/api/admin/teacher-guidance/<int:gid>', methods=['DELETE'])
def admin_delete_teacher_guidance(gid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM teacher_checkin_checkout WHERE id=?', (gid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/export-teacher-guidance')
def export_teacher_guidance():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可访问'}), 403
    try:
        from openpyxl import Workbook
    except ImportError:
        return jsonify({'error': '导出功能需要安装 openpyxl 库，请在服务器执行: pip install openpyxl'}), 500
    conn = db.get_conn()
    try:
        all_clubs = conn.execute('SELECT DISTINCT club_name FROM teacher_clubs ORDER BY club_name').fetchall()
        all_records = []
        for c in all_clubs:
            club_name = c['club_name']
            guiding_unit = ''
            gp = conn.execute('SELECT guiding_unit FROM club_profiles WHERE club_name=?', (club_name,)).fetchone()
            if gp and gp['guiding_unit']:
                guiding_unit = gp['guiding_unit']
            teachers = conn.execute('SELECT DISTINCT u.id, u.username, tp.real_name FROM users u LEFT JOIN teacher_profiles tp ON u.id=tp.user_id WHERE u.role="teacher" AND u.id IN (SELECT user_id FROM teacher_clubs WHERE club_name=?)', (club_name,)).fetchall()
            for t in teachers:
                teacher_name = t['real_name'] or t['username']
                records = conn.execute('SELECT tcc.checkin_time, tcc.checkout_time, tcc.status, tcc.checkin_lat, tcc.checkin_lng, tcc.checkout_lat, tcc.checkout_lng, tcc.checkin_address, tcc.checkout_address, tcc.session_id FROM teacher_checkin_checkout tcc WHERE tcc.teacher_user_id=? AND tcc.club_name=? AND tcc.status="checked_out" ORDER BY tcc.checkin_time DESC', (t['id'], club_name)).fetchall()
                if not records:
                    all_records.append({
                        'clubName': club_name,
                        'guidingUnit': guiding_unit,
                        'teacherName': teacher_name,
                        'activityName': '',
                        'checkinTime': '',
                        'checkoutTime': '',
                        'duration': '',
                        'checkinLocation': '',
                        'checkoutLocation': ''
                    })
                for r in records:
                    duration = ''
                    if r['checkin_time'] and r['checkout_time']:
                        try:
                            import datetime as _dt
                            ci = _dt.datetime.strptime(str(r['checkin_time']), '%Y-%m-%d %H:%M:%S')
                            co = _dt.datetime.strptime(str(r['checkout_time']), '%Y-%m-%d %H:%M:%S')
                            total_sec = int((co - ci).total_seconds())
                            hours = total_sec // 3600
                            mins = (total_sec % 3600) // 60
                            secs = total_sec % 60
                            if hours > 0:
                                duration = str(hours) + '小时' + (str(mins) + '分钟' if mins > 0 else '')
                            elif mins > 0:
                                duration = str(mins) + '分钟'
                            else:
                                duration = str(secs) + '秒'
                        except:
                            duration = ''
                    guidance_time = local_time(r['checkin_time'])
                    if r['checkout_time'] and r['status'] == 'checked_out':
                        guidance_time = local_time(r['checkin_time']) + ' ~ ' + local_time(r['checkout_time'])
                    activity_name = ''
                    if 'session_id' in r.keys() and r['session_id']:
                        sess_row = conn.execute('SELECT activity_name FROM checkin_sessions WHERE id=?', (r['session_id'],)).fetchone()
                        if sess_row:
                            activity_name = sess_row['activity_name'] or ''
                    checkin_addr = r['checkin_address'] if 'checkin_address' in r.keys() and r['checkin_address'] else (f"{r['checkin_lat']},{r['checkin_lng']}" if 'checkin_lat' in r.keys() and r['checkin_lat'] else '')
                    checkout_addr = r['checkout_address'] if 'checkout_address' in r.keys() and r['checkout_address'] else (f"{r['checkout_lat']},{r['checkout_lng']}" if 'checkout_lat' in r.keys() and r['checkout_lat'] else '')
                    all_records.append({
                        'clubName': club_name,
                        'guidingUnit': guiding_unit,
                        'teacherName': teacher_name,
                        'activityName': activity_name,
                        'checkinTime': local_time(r['checkin_time']),
                        'checkoutTime': local_time(r['checkout_time']),
                        'duration': duration,
                        'checkinLocation': checkin_addr,
                        'checkoutLocation': checkout_addr
                    })
    finally:
        conn.close()
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = '指导老师指导情况'
    ws.append(['社团名称', '业务指导单位', '指导老师', '活动名称', '签到时间', '签退时间', '指导时长', '签到位置', '签退位置'])
    for r in all_records:
        ws.append([r['clubName'], r['guidingUnit'], r['teacherName'], r['activityName'], r['checkinTime'], r['checkoutTime'], r['duration'], r['checkinLocation'], r['checkoutLocation']])
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name='指导老师指导情况.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/generate-checkout-code', methods=['POST'])
def generate_checkout_code():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先以社团负责人身份登录'}), 401
    data = request.json or {}
    session_id = data.get('sessionId')
    if user['role'] == 'student':
        conn_chk = db.get_conn()
        try:
            sess_chk = conn_chk.execute('SELECT club_name FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
        finally:
            conn_chk.close()
        if not sess_chk or not is_cadre_of_club(user['id'], sess_chk['club_name']):
            return jsonify({'error': '无权限操作'}), 403
    method = data.get('method', '')
    if not session_id:
        return jsonify({'error': '缺少活动会话ID'}), 400
    if not method:
        method = random.choice(['qrcode', 'code'])
    if method not in ('qrcode', 'code'):
        return jsonify({'error': '签退方式无效，仅支持qrcode或code'}), 400
    code = ''.join(random.choices(string.digits, k=6))
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (session_id,)).fetchone()
        if not sess:
            return jsonify({'error': '活动会话不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != sess['club_name']:
            if user['role'] == 'teacher':
                tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], sess['club_name'])).fetchone()
                if not tc_row:
                    return jsonify({'error': '无权限'}), 403
            elif user['role'] == 'student':
                if not is_cadre_of_club(user['id'], sess['club_name']):
                    return jsonify({'error': '无权限'}), 403
            else:
                return jsonify({'error': '无权限'}), 403
        conn.execute('UPDATE checkin_sessions SET checkout_code=?, checkout_method=? WHERE id=?', (code, method, session_id))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'checkoutCode': code, 'checkoutMethod': method})


@app.route('/checkin.html')
def serve_checkin():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'checkin.html'))


@app.route('/api/club-detail/<club_name>')
def club_detail(club_name):
    conn = db.get_conn()
    try:
        profile = conn.execute('SELECT * FROM club_profiles WHERE club_name=?', (club_name,)).fetchone()
        teachers = conn.execute('SELECT id, teacher_name, photo_path, introduction, user_id FROM club_teachers WHERE club_name=?', (club_name,)).fetchall()
        activity_count = conn.execute('SELECT COUNT(*) as c FROM activity_records WHERE club_name=?', (club_name,)).fetchone()['c']
    finally:
        conn.close()
    profile_data = {}
    if profile:
        emblem_url = ''
        if 'emblem_url' in profile.keys():
            emblem_url = profile['emblem_url'] or ''
        president = ''
        if 'president' in profile.keys():
            president = profile['president'] or ''
        category = ''
        if 'category' in profile.keys():
            category = profile['category'] or ''
        guiding_unit = ''
        if 'guiding_unit' in profile.keys():
            guiding_unit = profile['guiding_unit'] or ''
        profile_data = {'description': profile['description'], 'starRating': profile['star_rating'], 'showStar': profile['show_star'], 'registrationForm': profile['registration_form'], 'emblemUrl': emblem_url, 'president': president, 'category': category, 'guidingUnit': guiding_unit}
    else:
        profile_data = {'description': '', 'starRating': 0, 'showStar': 0, 'registrationForm': '', 'emblemUrl': '', 'president': '', 'category': '', 'guidingUnit': ''}
    teacher_list = []
    for t in teachers:
        teacher_list.append({'id': t['id'], 'name': t['teacher_name'], 'photo': t['photo_path'], 'introduction': t['introduction'], 'userId': t['user_id'] if 'user_id' in t.keys() else 0})
    base_teacher = get_teacher_name_from_sources(club_name)
    if base_teacher:
        for tn in base_teacher.split('、'):
            tn = tn.strip()
            if tn and not any(t['name'] == tn for t in teacher_list):
                teacher_list.append({'id': 0, 'name': tn, 'photo': '', 'introduction': ''})
    return jsonify({'success': True, 'data': {'clubName': club_name, 'profile': profile_data, 'teachers': teacher_list, 'activityCount': activity_count}})


@app.route('/api/club-profile', methods=['POST'])
def update_club_profile():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    club = data.get('clubName', '').strip()
    if not club:
        if user['role'] == 'admin':
            return jsonify({'error': '请指定社团'}), 400
        club = user['club_name'] or ''
    if not club:
        return jsonify({'error': '无社团信息'}), 400
    if user['role'] == 'student' and not is_cadre_of_club(user['id'], club):
        return jsonify({'error': '无权限操作该社团'}), 403
    if user['role'] == 'teacher':
        tc_conn = db.get_conn()
        try:
            tc_row = tc_conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
        finally:
            tc_conn.close()
        if not tc_row:
            return jsonify({'error': '无权限操作该社团'}), 403
    description = data.get('description', '')
    star_rating = int(data.get('starRating', 0))
    show_star = int(data.get('showStar', 0))
    registration_form = data.get('registrationForm', '')
    president = data.get('president', '').strip()
    category = data.get('category', '').strip()
    guiding_unit = data.get('guidingUnit', '').strip()
    if not guiding_unit:
        return jsonify({'error': '业务指导单位为必填项'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM club_profiles WHERE club_name=?', (club,)).fetchone()
        if existing:
            conn.execute('UPDATE club_profiles SET description=?, star_rating=?, show_star=?, registration_form=?, president=?, category=?, guiding_unit=?, updated_at=CURRENT_TIMESTAMP WHERE club_name=?', (description, star_rating, show_star, registration_form, president, category, guiding_unit, club))
        else:
            conn.execute('INSERT INTO club_profiles (club_name, description, star_rating, show_star, registration_form, president, category, guiding_unit) VALUES (?, ?, ?, ?, ?, ?, ?, ?)', (club, description, star_rating, show_star, registration_form, president, category, guiding_unit))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-teachers', methods=['POST'])
def add_club_teacher():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    club = data.get('clubName', '').strip()
    if not club:
        club = user['club_name'] or ''
    if not club:
        return jsonify({'error': '无社团信息'}), 400
    if user['role'] == 'student' and not is_cadre_of_club(user['id'], club):
        return jsonify({'error': '无权限操作该社团'}), 403
    if user['role'] == 'teacher':
        tc_conn = db.get_conn()
        try:
            tc_row = tc_conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
        finally:
            tc_conn.close()
        if not tc_row:
            return jsonify({'error': '无权限操作该社团'}), 403
    teacher_name = data.get('teacherName', '').strip()
    teacher_work_id = data.get('teacherWorkId', '').strip()
    introduction = data.get('introduction', '')
    if not teacher_name:
        return jsonify({'error': '请输入老师姓名'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM club_teachers WHERE club_name=? AND teacher_name=?', (club, teacher_name)).fetchone()
        if existing:
            return jsonify({'error': '该老师已在本社团中'}), 400
        teacher_user = None
        if teacher_work_id:
            teacher_user = conn.execute("SELECT u.id, u.username FROM users u JOIN teacher_profiles tp ON u.id=tp.user_id WHERE u.role='teacher' AND tp.work_id=?", (teacher_work_id,)).fetchone()
        if not teacher_user:
            teacher_user = conn.execute("SELECT u.id, u.username FROM users u JOIN teacher_profiles tp ON u.id=tp.user_id LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.role='teacher' AND (u.username=? OR (tp.real_name IS NOT NULL AND tp.real_name=?) OR (up.real_name IS NOT NULL AND up.real_name=?))", (teacher_name, teacher_name, teacher_name)).fetchone()
        if teacher_user:
            already_in = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (teacher_user['id'], club)).fetchone()
            if already_in:
                conn.execute('INSERT INTO club_teachers (club_name, teacher_name, introduction, user_id) VALUES (?, ?, ?, ?)', (club, teacher_name, introduction, teacher_user['id']))
                conn.commit()
                tid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                return jsonify({'success': True, 'id': tid, 'auto': True})
            pending = conn.execute('SELECT id FROM teacher_club_requests WHERE teacher_user_id=? AND club_name=? AND status="pending"', (teacher_user['id'], club)).fetchone()
            if pending:
                return jsonify({'error': '已发送邀请，等待老师确认'}), 400
            conn.execute('INSERT INTO teacher_club_requests (club_name, teacher_user_id, teacher_name, introduction, requested_by) VALUES (?, ?, ?, ?, ?)', (club, teacher_user['id'], teacher_name, introduction, user['id']))
            conn.commit()
            return jsonify({'success': True, 'pending': True, 'message': '已发送邀请，等待老师确认'})
        conn.execute('INSERT INTO club_teachers (club_name, teacher_name, introduction) VALUES (?, ?, ?)', (club, teacher_name, introduction))
        conn.commit()
        tid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
    finally:
        conn.close()
    return jsonify({'success': True, 'id': tid})


@app.route('/api/teacher-search', methods=['GET'])
def teacher_search():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '无权限'}), 403
    name = request.args.get('name', '').strip()
    if not name:
        return jsonify({'success': True, 'data': []})
    conn = db.get_conn()
    try:
        rows = conn.execute("SELECT u.id, u.username, tp.work_id, COALESCE(tp.real_name, up.real_name, u.username) as real_name FROM users u JOIN teacher_profiles tp ON u.id=tp.user_id LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.role='teacher' AND (u.username LIKE ? OR (tp.real_name IS NOT NULL AND tp.real_name LIKE ?) OR (up.real_name IS NOT NULL AND up.real_name LIKE ?)) ORDER BY tp.work_id", ('%' + name + '%', '%' + name + '%', '%' + name + '%')).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'userId': r['id'], 'username': r['username'], 'workId': r['work_id'] or '', 'realName': r['real_name'] or ''} for r in rows]})


@app.route('/api/teacher-club-requests', methods=['GET'])
def get_teacher_club_requests():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '仅指导老师可查看'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT r.id, r.club_name, r.teacher_name, r.introduction, r.status, r.created_at FROM teacher_club_requests r WHERE r.teacher_user_id=? AND r.status="pending" ORDER BY r.created_at DESC', (user['id'],)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'clubName': r['club_name'], 'teacherName': r['teacher_name'], 'introduction': r['introduction'], 'status': r['status'], 'createdAt': local_time(r['created_at'])} for r in rows]})


@app.route('/api/teacher-club-requests/<int:req_id>', methods=['POST'])
def handle_teacher_club_request(req_id):
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '仅指导老师可操作'}), 403
    data = request.json or {}
    action = data.get('action', '')
    if action not in ('approve', 'reject'):
        return jsonify({'error': '无效操作'}), 400
    conn = db.get_conn()
    try:
        req_row = conn.execute('SELECT * FROM teacher_club_requests WHERE id=? AND teacher_user_id=? AND status="pending"', (req_id, user['id'])).fetchone()
        if not req_row:
            return jsonify({'error': '请求不存在或已处理'}), 404
        if action == 'approve':
            conn.execute('UPDATE teacher_club_requests SET status="approved", reviewed_at=CURRENT_TIMESTAMP WHERE id=?', (req_id,))
            existing = conn.execute('SELECT id FROM club_teachers WHERE club_name=? AND user_id=?', (req_row['club_name'], user['id'])).fetchone()
            if not existing:
                conn.execute('INSERT INTO club_teachers (club_name, teacher_name, introduction, user_id) VALUES (?, ?, ?, ?)', (req_row['club_name'], req_row['teacher_name'], req_row['introduction'], user['id']))
            conn.execute('INSERT OR IGNORE INTO teacher_clubs (user_id, club_name) VALUES (?, ?)', (user['id'], req_row['club_name']))
        else:
            conn.execute('UPDATE teacher_club_requests SET status="rejected", reviewed_at=CURRENT_TIMESTAMP WHERE id=?', (req_id,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '已同意' if action == 'approve' else '已拒绝'})


@app.route('/api/club-teachers/<int:tid>', methods=['PUT', 'DELETE'])
def manage_club_teacher(tid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        teacher = conn.execute('SELECT club_name FROM club_teachers WHERE id=?', (tid,)).fetchone()
        if not teacher:
            return jsonify({'error': '老师记录不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != teacher['club_name']:
            if user['role'] == 'student' and is_cadre_of_club(user['id'], teacher['club_name']):
                pass
            elif user['role'] == 'teacher':
                tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], teacher['club_name'])).fetchone()
                if not tc_row:
                    return jsonify({'error': '无权限操作'}), 403
            else:
                return jsonify({'error': '无权限操作'}), 403
        if request.method == 'DELETE':
            conn.execute('DELETE FROM club_teachers WHERE id=?', (tid,))
            conn.commit()
        else:
            data = request.json or {}
            introduction = data.get('introduction', '')
            conn.execute('UPDATE club_teachers SET introduction=? WHERE id=?', (introduction, tid))
            conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/teacher-photo/<int:tid>', methods=['POST'])
def upload_teacher_photo(tid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if 'photo' not in request.files:
        return jsonify({'error': '请选择图片'}), 400
    f = request.files['photo']
    if not f.filename:
        return jsonify({'error': '请选择图片'}), 400
    conn = db.get_conn()
    try:
        teacher = conn.execute('SELECT club_name FROM club_teachers WHERE id=?', (tid,)).fetchone()
        if not teacher:
            return jsonify({'error': '老师记录不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != teacher['club_name']:
            return jsonify({'error': '无权限操作'}), 403
    finally:
        conn.close()
    upload_dir = os.path.join(os.path.dirname(__file__), 'data', 'teacher_photos')
    os.makedirs(upload_dir, exist_ok=True)
    ext = os.path.splitext(f.filename)[1] or '.jpg'
    filename = f'teacher_{tid}_{int(time.time())}{ext}'
    filepath = os.path.join(upload_dir, filename)
    f.save(filepath)
    photo_url = f'/api/teacher-photo-file/{filename}'
    conn = db.get_conn()
    try:
        conn.execute('UPDATE club_teachers SET photo_path=? WHERE id=?', (photo_url, tid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'photoUrl': photo_url})


@app.route('/api/teacher-photo-file/<filename>')
def serve_teacher_photo(filename):
    upload_dir = os.path.join(os.path.dirname(__file__), 'data', 'teacher_photos')
    return send_from_directory(upload_dir, filename)


@app.route('/api/teacher-profile', methods=['GET', 'POST'])
def teacher_profile():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        if request.method == 'GET':
            row = conn.execute('SELECT * FROM teacher_profiles WHERE user_id=?', (user['id'],)).fetchone()
            clubs = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
        else:
            row = None
            clubs = []
    finally:
        conn.close()
    if request.method == 'GET':
        if not row:
            return jsonify({'success': True, 'data': {'workId': '', 'realName': '', 'phone': '', 'email': '', 'avatarPath': '', 'introduction': '', 'clubs': [c['club_name'] for c in clubs]}})
        return jsonify({'success': True, 'data': {'workId': row['work_id'] or '', 'realName': row['real_name'] or '', 'phone': row['phone'] or '', 'email': row['email'] or '', 'avatarPath': row['avatar_path'] or '', 'introduction': row['introduction'] or '', 'clubs': [c['club_name'] for c in clubs]}})
    data = request.json or {}
    work_id = data.get('workId', '').strip()
    real_name = data.get('realName', '').strip()
    phone = data.get('phone', '').strip()
    email = data.get('email', '').strip()
    introduction = data.get('introduction', '').strip()
    new_clubs = data.get('clubs', [])
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM teacher_profiles WHERE user_id=?', (user['id'],)).fetchone()
        if existing:
            conn.execute('UPDATE teacher_profiles SET work_id=?, real_name=?, phone=?, email=?, introduction=?, updated_at=CURRENT_TIMESTAMP WHERE user_id=?',
                        (work_id, real_name, phone, email, introduction, user['id']))
        else:
            conn.execute('INSERT INTO teacher_profiles (user_id, work_id, real_name, phone, email, introduction) VALUES (?, ?, ?, ?, ?, ?)',
                        (user['id'], work_id, real_name, phone, email, introduction))
        conn.commit()
        teacher_display_name = real_name or user['username']
        if new_clubs is not None:
            old_clubs = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
            old_club_names = [c['club_name'] for c in old_clubs]
            new_club_names = [c.strip() for c in new_clubs if c.strip()]
            removed_club_names = [c for c in old_club_names if c not in new_club_names]
            conn.execute('DELETE FROM teacher_clubs WHERE user_id=?', (user['id'],))
            for cn in new_club_names:
                if not cn:
                    continue
                conn.execute('INSERT OR IGNORE INTO teacher_clubs (user_id, club_name) VALUES (?, ?)', (user['id'], cn))
                existing_teacher = conn.execute('SELECT id FROM club_teachers WHERE club_name=? AND (user_id=? OR (user_id=0 AND teacher_name=?))', (cn, user['id'], teacher_display_name)).fetchone()
                if existing_teacher:
                    conn.execute('UPDATE club_teachers SET user_id=?, teacher_name=? WHERE id=?', (user['id'], teacher_display_name, existing_teacher['id']))
                else:
                    conn.execute('INSERT INTO club_teachers (club_name, teacher_name, user_id) VALUES (?, ?, ?)', (cn, teacher_display_name, user['id']))
            conn.execute('DELETE FROM club_teachers WHERE user_id=? AND club_name NOT IN ({})'.format(','.join(['?']*len(new_club_names)) if new_club_names else 'SELECT 0'), [user['id']] + new_club_names)
            for removed_cn in removed_club_names:
                conn.execute('DELETE FROM club_teachers WHERE club_name=? AND (user_id=? OR (user_id=0 AND teacher_name=?))', (removed_cn, user['id'], teacher_display_name))
                leader = conn.execute("SELECT id FROM users WHERE club_name=? AND role='user'", (removed_cn,)).fetchone()
                if leader:
                    send_notification(leader['id'], '指导老师变更', f'指导老师 {teacher_display_name} 已取消指导社团 {removed_cn}', 'warning', '', conn)
            primary_club = new_club_names[0] if new_club_names else ''
            conn.execute('UPDATE users SET club_name=? WHERE id=?', (primary_club, user['id']))
        conn.commit()
        if real_name:
            conn.execute('UPDATE club_teachers SET teacher_name=? WHERE user_id=?', (real_name, user['id']))
            conn.commit()
        if introduction:
            conn.execute('UPDATE club_teachers SET introduction=? WHERE user_id=?', (introduction, user['id']))
            conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/teacher-avatar', methods=['POST'])
def upload_teacher_avatar():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '无权限'}), 403
    if 'avatar' not in request.files:
        return jsonify({'error': '请选择图片'}), 400
    f = request.files['avatar']
    if not f.filename:
        return jsonify({'error': '请选择图片'}), 400
    upload_dir = os.path.join(os.path.dirname(__file__), 'data', 'teacher_photos')
    os.makedirs(upload_dir, exist_ok=True)
    ext = os.path.splitext(f.filename)[1] or '.jpg'
    filename = f'teacher_avatar_{user["id"]}_{int(time.time())}{ext}'
    filepath = os.path.join(upload_dir, filename)
    f.save(filepath)
    photo_url = f'/api/teacher-photo-file/{filename}'
    conn = db.get_conn()
    try:
        conn.execute('UPDATE teacher_profiles SET avatar_path=? WHERE user_id=?', (photo_url, user['id']))
        conn.commit()
        conn.execute('UPDATE club_teachers SET photo_path=? WHERE user_id=?', (photo_url, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'avatarUrl': photo_url})


@app.route('/api/teacher-clubs')
def get_teacher_clubs():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        clubs = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [c['club_name'] for c in clubs]})


@app.route('/api/all-clubs-list')
def get_all_clubs_list():
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT DISTINCT club_name FROM club_tokens ORDER BY club_name').fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [r['club_name'] for r in rows]})


@app.route('/api/colleges')
def get_colleges():
    return jsonify({'success': True, 'data': COLLEGES})


@app.route('/api/all-clubs')
def get_all_clubs():
    conn = db.get_conn()
    try:
        clubs = set()
        for r in conn.execute('SELECT DISTINCT club_name FROM club_tokens').fetchall():
            if r['club_name']: clubs.add(r['club_name'])
        for r in conn.execute("SELECT DISTINCT club_name FROM users WHERE role='user' AND club_name IS NOT NULL AND club_name!=''").fetchall():
            if r['club_name']: clubs.add(r['club_name'])
        for r in conn.execute('SELECT DISTINCT club_name FROM teacher_clubs').fetchall():
            if r['club_name']: clubs.add(r['club_name'])
        for r in conn.execute('SELECT DISTINCT club_name FROM club_profiles').fetchall():
            if r['club_name']: clubs.add(r['club_name'])
    finally:
        conn.close()
    return jsonify({'success': True, 'data': sorted(clubs)})


@app.route('/api/finance-records', methods=['GET', 'POST'])
def handle_finance_records():
    if request.method == 'GET':
        user = get_current_user()
        if not user:
            return jsonify({'error': '请先登录'}), 401
        club = request.args.get('club', '')
        if not club and user:
            club = user.get('club_name', '') or ''
        if not club:
            return jsonify({'success': True, 'data': []})
        rec_type = request.args.get('type', '')
        category = request.args.get('category', '')
        start_date = request.args.get('startDate', '')
        end_date = request.args.get('endDate', '')
        conn = db.get_conn()
        try:
            q = 'SELECT * FROM finance_records WHERE club_name=?'
            params = [club]
            if rec_type:
                q += ' AND type=?'
                params.append(rec_type)
            if category:
                q += ' AND category=?'
                params.append(category)
            if start_date:
                q += ' AND record_date>=?'
                params.append(start_date)
            if end_date:
                q += ' AND record_date<=?'
                params.append(end_date)
            q += ' ORDER BY record_date DESC, created_at DESC'
            rows = conn.execute(q, params).fetchall()
        finally:
            conn.close()
        total_income = sum(r['amount'] for r in rows if r['type'] == 'income')
        total_expense = sum(r['amount'] for r in rows if r['type'] == 'expense')
        return jsonify({
            'success': True,
            'data': [{'id': r['id'], 'type': r['type'], 'category': r['category'], 'amount': r['amount'], 'description': r['description'], 'recordDate': r['record_date'], 'recorder': r['recorder'], 'attachmentPath': r['attachment_path'], 'attachmentName': r['attachment_name'], 'createdAt': local_time(r['created_at'])} for r in rows],
            'summary': {'totalIncome': total_income, 'totalExpense': total_expense, 'balance': total_income - total_expense}
        })
    else:
        user = get_current_user()
        if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
            return jsonify({'error': '请先登录'}), 401
        data = request.json or {}
        club = data.get('clubName', '').strip()
        if not club:
            club = user['club_name'] or ''
        if not club:
            return jsonify({'error': '无社团信息'}), 400
        if user['role'] == 'student':
            if is_cadre_of_club(user['id'], club):
                pass
            else:
                conn_tmp = db.get_conn()
                try:
                    mgr_tmp = conn_tmp.execute('SELECT id FROM finance_managers WHERE club_name=? AND user_id=?', (club, user['id'])).fetchone()
                finally:
                    conn_tmp.close()
                if not mgr_tmp:
                    return jsonify({'error': '您没有该社团的财务管理权限'}), 403
        if user['role'] == 'teacher':
            conn_tmp = db.get_conn()
            try:
                tc_tmp = conn_tmp.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
            finally:
                conn_tmp.close()
            if not tc_tmp:
                return jsonify({'error': '您没有该社团的财务管理权限'}), 403
        rec_type = data.get('type', '').strip()
        category = data.get('category', '').strip()
        amount = data.get('amount', 0)
        description = data.get('description', '').strip()
        record_date = data.get('recordDate', '').strip()
        recorder = data.get('recorder', '').strip()
        if rec_type not in ('income', 'expense'):
            return jsonify({'error': '类型必须为收入或支出'}), 400
        if not category:
            return jsonify({'error': '请填写分类'}), 400
        try:
            amount = float(amount)
        except (ValueError, TypeError):
            return jsonify({'error': '金额必须为数字'}), 400
        if amount <= 0:
            return jsonify({'error': '金额必须大于0'}), 400
        if not record_date:
            return jsonify({'error': '请选择日期'}), 400
        conn = db.get_conn()
        try:
            conn.execute('INSERT INTO finance_records (club_name, type, category, amount, description, record_date, recorder) VALUES (?, ?, ?, ?, ?, ?, ?)', (club, rec_type, category, amount, description, record_date, recorder))
            conn.commit()
            fid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
        finally:
            conn.close()
        return jsonify({'success': True, 'id': fid})


@app.route('/api/finance-records/<int:fid>', methods=['PUT', 'DELETE'])
def manage_finance_record(fid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        record = conn.execute('SELECT * FROM finance_records WHERE id=?', (fid,)).fetchone()
        if not record:
            return jsonify({'error': '记录不存在'}), 404
        if user['role'] == 'student':
            if is_cadre_of_club(user['id'], record['club_name']):
                pass
            else:
                mgr = conn.execute('SELECT id FROM finance_managers WHERE club_name=? AND user_id=?', (record['club_name'], user['id'])).fetchone()
                if not mgr:
                    return jsonify({'error': '无权限'}), 403
        elif user['role'] == 'teacher':
            tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], record['club_name'])).fetchone()
            if not tc_row:
                return jsonify({'error': '无权限'}), 403
        elif user['role'] != 'admin' and user['club_name'] != record['club_name']:
            return jsonify({'error': '无权限'}), 403
        if request.method == 'DELETE':
            conn.execute('DELETE FROM finance_records WHERE id=?', (fid,))
            conn.commit()
        else:
            data = request.json or {}
            rec_type = data.get('type', record['type'])
            category = data.get('category', record['category'])
            amount = data.get('amount', record['amount'])
            description = data.get('description', record['description'])
            record_date = data.get('recordDate', record['record_date'])
            recorder = data.get('recorder', record['recorder'])
            if isinstance(amount, (int, float)):
                pass
            else:
                try:
                    amount = float(amount)
                except:
                    amount = record['amount']
            conn.execute('UPDATE finance_records SET type=?, category=?, amount=?, description=?, record_date=?, recorder=? WHERE id=?', (rec_type, category, amount, description, record_date, recorder, fid))
            conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/finance-attachment/<int:fid>', methods=['POST'])
def upload_finance_attachment(fid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401
    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '请选择文件'}), 400
    conn = db.get_conn()
    try:
        record = conn.execute('SELECT * FROM finance_records WHERE id=?', (fid,)).fetchone()
        if not record:
            return jsonify({'error': '记录不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != record['club_name']:
            return jsonify({'error': '无权限'}), 403
    finally:
        conn.close()
    import time as _t
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else 'bin'
    filename = f'{fid}_{int(_t.time())}_{f.filename}'
    key = 'finance_attachments/' + filename
    storage.save(f, key)
    conn = db.get_conn()
    try:
        conn.execute('UPDATE finance_records SET attachment_path=?, attachment_name=? WHERE id=?', (key, f.filename, fid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'path': key, 'name': f.filename})


@app.route('/api/finance-attachment-file/<path:storage_key>')
def serve_finance_attachment(storage_key):
    path = storage.get_path(storage_key)
    if path:
        return send_file(path)
    url = storage.get_url(storage_key)
    if url:
        return jsonify({'url': url})
    return jsonify({'error': '文件不存在'}), 404


@app.route('/api/all-completed-activities')
def all_completed_activities():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '请先登录'}), 401
    page = max(1, int(request.args.get('page', 1)))
    per_page = min(100, max(1, int(request.args.get('per_page', 30))))
    offset = (page - 1) * per_page
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        where_clauses = ['cs.is_completed=1']
        params = []
        if start_date:
            where_clauses.append('date(cs.created_at)>=?')
            params.append(start_date)
        if end_date:
            where_clauses.append('date(cs.created_at)<=?')
            params.append(end_date)
        where_sql = ' AND '.join(where_clauses)
        total = conn.execute(f'SELECT COUNT(*) as c FROM checkin_sessions cs WHERE {where_sql}', params).fetchone()['c']
        rows = conn.execute(f'SELECT cs.*, COUNT(cr.id) as checkin_count FROM checkin_sessions cs LEFT JOIN checkin_records cr ON cs.id=cr.session_id WHERE {where_sql} GROUP BY cs.id ORDER BY cs.created_at DESC LIMIT ? OFFSET ?', params + [per_page, offset]).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'total': total, 'page': page, 'per_page': per_page, 'data': [{'id': r['id'], 'clubName': r['club_name'], 'activityName': r['activity_name'], 'locationName': r['location_name'], 'activityTime': r['activity_time'] if 'activity_time' in r.keys() else '', 'startTime': r['start_time'] if 'start_time' in r.keys() else '', 'endTime': r['end_time'] if 'end_time' in r.keys() else '', 'activityContent': r['activity_content'] if 'activity_content' in r.keys() else '', 'teacherIds': r['teacher_ids'] if 'teacher_ids' in r.keys() else '', 'completionPhoto': r['completion_photo'] if 'completion_photo' in r.keys() else '', 'summaryText': r['summary_text'] if 'summary_text' in r.keys() else '', 'summaryPath': r['summary_path'] if 'summary_path' in r.keys() else '', 'planText': r['plan_text'] if 'plan_text' in r.keys() else '', 'planPath': r['plan_path'] if 'plan_path' in r.keys() else '', 'warning': r['warning'] if 'warning' in r.keys() else '', 'warningReason': r['warning_reason'] if 'warning_reason' in r.keys() else '', 'checkinCount': r['checkin_count'], 'createdAt': local_time(r['created_at'])} for r in rows]})


@app.route('/api/activity-warning/<int:sid>', methods=['POST'])
@app.route('/api/activity-warning', methods=['POST'], defaults={'sid': None})
def activity_warning(sid):
    if sid is None:
        sid = (request.json or {}).get('id', 0)
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    action = data.get('action', '')
    reason = data.get('reason', '').strip()
    conn = db.get_conn()
    try:
        sess = conn.execute('SELECT * FROM checkin_sessions WHERE id=?', (sid,)).fetchone()
        if not sess:
            return jsonify({'error': '活动不存在'}), 404
        if action == 'warn':
            if not reason:
                return jsonify({'error': '请填写警告原因'}), 400
            try:
                conn.execute('UPDATE checkin_sessions SET warning="warned", warning_reason=? WHERE id=?', (reason, sid))
            except:
                conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning TEXT DEFAULT ""')
                conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning_reason TEXT DEFAULT ""')
                conn.execute('UPDATE checkin_sessions SET warning="warned", warning_reason=? WHERE id=?', (reason, sid))
            leaders = conn.execute('SELECT id FROM users WHERE club_name=? AND role="user"', (sess['club_name'],)).fetchall()
            for l in leaders:
                send_notification(l['id'], '⚠️ 活动警告', '活动「' + (sess['activity_name'] or '') + '」收到警告：' + reason, 'warning', '/club-tools.html', conn=conn)
            conn.commit()
        elif action == 'unwarn':
            try:
                conn.execute('UPDATE checkin_sessions SET warning="", warning_reason="" WHERE id=?', (sid,))
            except:
                conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning TEXT DEFAULT ""')
                conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning_reason TEXT DEFAULT ""')
                conn.execute('UPDATE checkin_sessions SET warning="", warning_reason="" WHERE id=?', (sid,))
            conn.commit()
        else:
            return jsonify({'error': '无效操作'}), 400
    except Exception as e:
        return jsonify({'error': '操作失败：' + str(e)}), 500
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/finance-summary/<club_name>')
def finance_summary(club_name):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT type, amount, category, record_date FROM finance_records WHERE club_name=?', (club_name,)).fetchall()
    finally:
        conn.close()
    total_income = sum(r['amount'] for r in rows if r['type'] == 'income')
    total_expense = sum(r['amount'] for r in rows if r['type'] == 'expense')
    categories = {}
    for r in rows:
        cat = r['category'] or '其他'
        if cat not in categories:
            categories[cat] = {'income': 0, 'expense': 0}
        categories[cat][r['type']] = categories[cat].get(r['type'], 0) + r['amount']
    monthly = {}
    for r in rows:
        m = (r['record_date'] or '')[:7]
        if not m:
            continue
        if m not in monthly:
            monthly[m] = {'income': 0, 'expense': 0}
        monthly[m][r['type']] = monthly[m].get(r['type'], 0) + r['amount']
    return jsonify({
        'success': True,
        'totalIncome': total_income,
        'totalExpense': total_expense,
        'balance': total_income - total_expense,
        'categories': categories,
        'monthly': monthly,
        'recordCount': len(rows)
    })


@app.route('/api/export-finance/<club_name>')
def export_finance(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] in ('student', 'teacher'):
        if user['role'] == 'student' and is_cadre_of_club(user['id'], club_name):
            pass
        else:
            conn = db.get_conn()
            is_manager = conn.execute('SELECT 1 FROM finance_managers WHERE club_name=? AND user_id=?', (club_name, user['id'])).fetchone()
            is_member = conn.execute("SELECT 1 FROM users WHERE id=? AND club_name=?", (user['id'], club_name)).fetchone()
            conn.close()
            if not is_manager and not is_member:
                return jsonify({'error': '无权限'}), 403
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    except ImportError:
        return jsonify({'error': '导出功能需要安装 openpyxl 库，请在服务器执行: pip install openpyxl'}), 500
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT * FROM finance_records WHERE club_name=? ORDER BY record_date DESC', (club_name,)).fetchall()
    finally:
        conn.close()
    wb = Workbook()
    ws = wb.active
    ws.title = '财务记录'
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color='667eea', end_color='667eea', fill_type='solid')
    header_font_white = Font(bold=True, size=11, color='FFFFFF')
    headers = ['日期', '类型', '分类', '金额', '说明', '记录人']
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=i, value=h)
        c.font = header_font_white
        c.fill = header_fill
        c.alignment = Alignment(horizontal='center')
    for idx, r in enumerate(rows, 2):
        ws.cell(row=idx, column=1, value=r['record_date'])
        ws.cell(row=idx, column=2, value='收入' if r['type'] == 'income' else '支出')
        ws.cell(row=idx, column=3, value=r['category'])
        ws.cell(row=idx, column=4, value=r['amount'])
        ws.cell(row=idx, column=5, value=r['description'])
        ws.cell(row=idx, column=6, value=r['recorder'])
    total_row = len(rows) + 2
    ws.cell(row=total_row, column=1, value='合计')
    income_sum = sum(r['amount'] for r in rows if r['type'] == 'income')
    expense_sum = sum(r['amount'] for r in rows if r['type'] == 'expense')
    ws.cell(row=total_row, column=2, value=f'收入: {income_sum}')
    ws.cell(row=total_row, column=3, value=f'支出: {expense_sum}')
    ws.cell(row=total_row, column=4, value=f'余额: {income_sum - expense_sum}')
    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws.column_dimensions[col].width = 18
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=f'{club_name}_财务记录.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/finance-permission-apply', methods=['POST'])
def apply_finance_permission():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    club = data.get('clubName', '').strip()
    reason = data.get('reason', '').strip()
    if not club:
        return jsonify({'error': '请指定社团'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id, status FROM finance_permissions WHERE club_name=? AND user_id=? ORDER BY created_at DESC LIMIT 1', (club, user['id'])).fetchone()
        if existing and existing['status'] == 'pending':
            return jsonify({'error': '您已提交过申请，请等待审批'}), 400
        if existing and existing['status'] == 'approved':
            return jsonify({'error': '您已拥有该社团的财务管理权限'}), 400
        mgr = conn.execute('SELECT id FROM finance_managers WHERE club_name=? AND user_id=?', (club, user['id'])).fetchone()
        if mgr:
            return jsonify({'error': '您已拥有该社团的财务管理权限'}), 400
        real_name = ''
        try:
            profile = conn.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
            if profile:
                real_name = profile['real_name'] or ''
        except:
            pass
        conn.execute('INSERT INTO finance_permissions (club_name, user_id, username, real_name, reason) VALUES (?, ?, ?, ?, ?)', (club, user['id'], user['username'], real_name, reason))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/finance-permission-requests/<club_name>')
def get_finance_permission_requests(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] == 'user' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    if user['role'] == 'student' and not is_cadre_of_club(user['id'], club_name):
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT * FROM finance_permissions WHERE club_name=? ORDER BY created_at DESC', (club_name,)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'userId': r['user_id'], 'username': r['username'], 'realName': r['real_name'], 'reason': r['reason'], 'status': r['status'], 'reviewedBy': r['reviewed_by'], 'reviewedAt': local_time(r['reviewed_at']), 'createdAt': local_time(r['created_at'])} for r in rows]})


@app.route('/api/finance-permission-review/<int:pid>', methods=['POST'])
def review_finance_permission(pid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student'):
        return jsonify({'error': '请先登录'}), 401
    data = request.json or {}
    action = data.get('action', '')
    if action not in ('approve', 'reject'):
        return jsonify({'error': '无效操作'}), 400
    conn = db.get_conn()
    try:
        record = conn.execute('SELECT * FROM finance_permissions WHERE id=?', (pid,)).fetchone()
        if not record:
            return jsonify({'error': '申请不存在'}), 404
        if user['role'] == 'user' and user['club_name'] != record['club_name']:
            return jsonify({'error': '无权限'}), 403
        if user['role'] == 'student' and not is_cadre_of_club(user['id'], record['club_name']):
            return jsonify({'error': '无权限'}), 403
        new_status = 'approved' if action == 'approve' else 'rejected'
        conn.execute('UPDATE finance_permissions SET status=?, reviewed_by=?, reviewed_at=CURRENT_TIMESTAMP WHERE id=?', (new_status, user['username'], pid))
        if action == 'approve':
            try:
                conn.execute('INSERT OR IGNORE INTO finance_managers (club_name, user_id, username, real_name, granted_by) VALUES (?, ?, ?, ?, ?)', (record['club_name'], record['user_id'], record['username'], record['real_name'], user['username']))
            except:
                pass
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/finance-managers/<club_name>')
def get_finance_managers(club_name):
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT * FROM finance_managers WHERE club_name=?', (club_name,)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'userId': r['user_id'], 'username': r['username'], 'realName': r['real_name'], 'grantedBy': r['granted_by'], 'createdAt': local_time(r['created_at'])} for r in rows]})


@app.route('/api/finance-managers/<club_name>/<int:mid>', methods=['DELETE'])
def remove_finance_manager(club_name, mid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] == 'user' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    if user['role'] == 'student' and not is_cadre_of_club(user['id'], club_name):
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT club_name FROM finance_managers WHERE id=?', (mid,)).fetchone()
        if not row:
            return jsonify({'error': '管理员不存在'}), 404
        if row['club_name'] != club_name:
            return jsonify({'error': '无权限'}), 403
        conn.execute('DELETE FROM finance_managers WHERE id=?', (mid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/my-finance-permissions')
def my_finance_permissions():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        mgr_rows = conn.execute('SELECT club_name FROM finance_managers WHERE user_id=?', (user['id'],)).fetchall()
        app_rows = conn.execute('SELECT club_name, status, created_at FROM finance_permissions WHERE user_id=? ORDER BY created_at DESC', (user['id'],)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'managedClubs': [r['club_name'] for r in mgr_rows], 'applications': [{'clubName': r['club_name'], 'status': r['status'], 'createdAt': local_time(r['created_at'])} for r in app_rows]})


@app.route('/api/check-finance-manager/<club_name>')
def check_finance_manager(club_name):
    user = get_current_user()
    if not user:
        return jsonify({'isManager': False})
    is_leader = user['role'] in ('user', 'admin')
    conn = db.get_conn()
    try:
        mgr = conn.execute('SELECT id FROM finance_managers WHERE club_name=? AND user_id=?', (club_name, user['id'])).fetchone()
    finally:
        conn.close()
    return jsonify({'isManager': bool(mgr), 'isLeader': is_leader})


@app.route('/api/phone-change-requests', methods=['GET'])
def get_phone_change_requests():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT pcr.*, u.username FROM phone_change_requests pcr LEFT JOIN users u ON pcr.user_id=u.id WHERE pcr.status="pending" ORDER BY pcr.created_at DESC').fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [
        {'id': r['id'], 'userId': r['user_id'], 'username': r['username'], 'oldPhone': r['old_phone'], 'newPhone': r['new_phone'], 'status': r['status'], 'createdAt': local_time(r['created_at'])}
        for r in rows
    ]})

@app.route('/api/phone-change-requests/<int:rid>', methods=['PUT'])
def review_phone_change(rid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    action = data.get('action', '').strip()
    if action not in ('approve', 'reject'):
        return jsonify({'error': '无效操作'}), 400
    conn = db.get_conn()
    try:
        req = conn.execute('SELECT * FROM phone_change_requests WHERE id=? AND status="pending"', (rid,)).fetchone()
        if not req:
            return jsonify({'error': '申请不存在或已处理'}), 404
        if action == 'approve':
            conn.execute('UPDATE user_profiles SET phone=? WHERE user_id=?', (req['new_phone'], req['user_id']))
        conn.execute('UPDATE phone_change_requests SET status=?, reviewed_at=CURRENT_TIMESTAMP, reviewed_by=? WHERE id=?',
                    ('approved' if action == 'approve' else 'rejected', user['id'], rid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/admin/users', methods=['GET'])
def admin_get_users():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        search = request.args.get('search', '').strip()
        role_filter = request.args.get('role', '').strip()
        grade_filter = request.args.get('grade', '').strip()
        page = max(1, int(request.args.get('page', 1)))
        per_page = 50
        q = 'SELECT u.id, u.username, u.role, u.club_name, u.created_at FROM users u WHERE 1=1'
        params = []
        if search:
            q += ' AND (u.username LIKE ? OR u.club_name LIKE ? OR u.id IN (SELECT up.user_id FROM user_profiles up WHERE up.real_name LIKE ?) OR u.id IN (SELECT tp.user_id FROM teacher_profiles tp WHERE tp.work_id LIKE ?))'
            params.extend(['%' + search + '%', '%' + search + '%', '%' + search + '%', '%' + search + '%'])
        if role_filter:
            q += ' AND u.role=?'
            params.append(role_filter)
        if grade_filter:
            q += ' AND (u.id IN (SELECT up.user_id FROM user_profiles up WHERE up.student_id LIKE ?) OR u.id IN (SELECT tp.user_id FROM teacher_profiles tp WHERE tp.work_id LIKE ?))'
            params.extend([grade_filter + '%', grade_filter + '%'])
        total = conn.execute('SELECT COUNT(*) as c FROM (' + q + ')', params).fetchone()['c']
        q += ' ORDER BY u.id LIMIT ? OFFSET ?'
        params.extend([per_page, (page - 1) * per_page])
        rows = conn.execute(q, params).fetchall()
        result = []
        for r in rows:
            item = {'id': r['id'], 'username': r['username'], 'role': r['role'], 'clubName': r['club_name'], 'createdAt': local_time(r['created_at'])}
            try:
                profile = conn.execute('SELECT real_name, student_id, class_name, college, phone FROM user_profiles WHERE user_id=?', (r['id'],)).fetchone()
                if profile:
                    item['realName'] = profile['real_name'] or ''
                    item['studentId'] = profile['student_id'] or ''
                    item['className'] = profile['class_name'] or ''
                    item['college'] = profile['college'] if 'college' in profile.keys() else ''
                    item['phone'] = profile['phone'] or ''
            except:
                pass
            if r['role'] == 'teacher':
                tp = conn.execute('SELECT work_id FROM teacher_profiles WHERE user_id=?', (r['id'],)).fetchone()
                item['workId'] = tp['work_id'] if tp else ''
                tc_rows = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=? ORDER BY club_name', (r['id'],)).fetchall()
                item['clubs'] = [cr['club_name'] for cr in tc_rows if cr['club_name']]
                if not item['clubs'] and r['club_name']:
                    item['clubs'] = [r['club_name']]
            if r['role'] == 'student':
                club_rows = conn.execute('SELECT DISTINCT club_name FROM club_members WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_cadres WHERE user_id=?', (r['id'], r['id'])).fetchall()
                item['clubs'] = [cr['club_name'] for cr in club_rows if cr['club_name']]
                if not item['clubs'] and r['club_name']:
                    item['clubs'] = [r['club_name']]
            result.append(item)
        role_stats = {}
        for rr in conn.execute('SELECT role, COUNT(*) as c FROM users GROUP BY role').fetchall():
            role_stats[rr['role']] = rr['c']
        return jsonify({'success': True, 'data': result, 'total': total, 'page': page, 'perPage': per_page, 'totalPages': (total + per_page - 1) // per_page, 'roleStats': role_stats})
    finally:
        conn.close()


@app.route('/api/admin/users/<int:uid>', methods=['PUT'])
def admin_update_user(uid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    conn = db.get_conn()
    try:
        target = conn.execute('SELECT * FROM users WHERE id=?', (uid,)).fetchone()
        if not target:
            return jsonify({'error': '用户不存在'}), 404
        new_role = data.get('role', target['role'])
        new_club = data.get('clubName', target['club_name'])
        conn.execute('UPDATE users SET role=?, club_name=? WHERE id=?', (new_role, new_club, uid))
        if data.get('realName') or data.get('studentId') or data.get('className') or data.get('college') or data.get('phone') or data.get('grade') or data.get('email'):
            try:
                existing = conn.execute('SELECT id FROM user_profiles WHERE user_id=?', (uid,)).fetchone()
                if existing:
                    conn.execute('UPDATE user_profiles SET real_name=?, student_id=?, grade=?, class_name=?, college=?, phone=?, email=? WHERE user_id=?',
                                 (data.get('realName', ''), data.get('studentId', ''), data.get('grade', ''), data.get('className', ''), data.get('college', ''), data.get('phone', ''), data.get('email', ''), uid))
                else:
                    conn.execute('INSERT INTO user_profiles (user_id, real_name, student_id, grade, class_name, college, phone, email) VALUES (?,?,?,?,?,?,?,?)',
                                 (uid, data.get('realName', ''), data.get('studentId', ''), data.get('grade', ''), data.get('className', ''), data.get('college', ''), data.get('phone', ''), data.get('email', '')))
            except:
                pass
        if new_role == 'student' and new_club and new_club != target['club_name']:
            existing_member = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=?', (new_club, uid)).fetchone()
            if not existing_member:
                profile_data = {}
                try:
                    prow = conn.execute('SELECT real_name, student_id, class_name, college, phone FROM user_profiles WHERE user_id=?', (uid,)).fetchone()
                    if prow:
                        profile_data = dict(prow)
                except:
                    pass
                conn.execute('''INSERT INTO club_members (club_name, user_id, username, real_name, student_id_num, class_name, phone, department, specialty, college, source)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'admin')''',
                    (new_club, uid, target['username'], profile_data.get('real_name', ''), profile_data.get('student_id', ''),
                     profile_data.get('class_name', ''), profile_data.get('phone', ''), '', '', profile_data.get('college', '')))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/admin/users/<int:uid>', methods=['DELETE'])
def admin_delete_user(uid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    if uid == user['id']:
        return jsonify({'error': '不能删除自己'}), 400
    conn = db.get_conn()
    try:
        target = conn.execute('SELECT * FROM users WHERE id=?', (uid,)).fetchone()
        if not target:
            return jsonify({'error': '用户不存在'}), 404
        profile = conn.execute('SELECT real_name, student_id FROM user_profiles WHERE user_id=?', (uid,)).fetchone()
        del_name = profile['real_name'] if profile else ''
        del_sid = profile['student_id'] if profile else ''
        user_clubs = conn.execute('SELECT DISTINCT club_name FROM club_members WHERE user_id=? UNION SELECT DISTINCT club_name FROM club_cadres WHERE user_id=?', (uid, uid)).fetchall()
        club_names = [c['club_name'] for c in user_clubs]
        conn.execute('DELETE FROM user_profiles WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM teacher_profiles WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM ai_chat_history WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM ai_pets WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM notifications WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM admin_notification_reads WHERE user_id=?', (uid,))
        if del_name and club_names:
            placeholders = ','.join(['?'] * len(club_names))
            conn.execute('DELETE FROM club_members WHERE user_id=0 AND real_name=? AND club_name IN (' + placeholders + ')', [del_name] + club_names)
        conn.execute('DELETE FROM club_members WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM club_teachers WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM teacher_clubs WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM teacher_club_requests WHERE teacher_user_id=?', (uid,))
        if del_sid:
            conn.execute('DELETE FROM checkin_records WHERE student_id=?', (del_sid,))
        conn.execute('DELETE FROM checkin_records WHERE student_id=?', (uid,))
        conn.execute('DELETE FROM checkin_records WHERE student_id=?', (str(uid),))
        conn.execute('DELETE FROM teacher_checkin_checkout WHERE teacher_user_id=?', (uid,))
        if del_name:
            conn.execute('DELETE FROM finance_records WHERE recorder=?', (del_name,))
        conn.execute('DELETE FROM finance_managers WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM finance_permissions WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM scoring_submission_items WHERE student_user_id=?', (uid,))
        conn.execute('DELETE FROM scoring_teacher_reviews WHERE teacher_user_id=?', (uid,))
        if del_sid:
            conn.execute('DELETE FROM final_credits WHERE student_id_num=?', (del_sid,))
        conn.execute('DELETE FROM workload_submissions WHERE student_user_id=?', (uid,))
        conn.execute('DELETE FROM recruitment_signups WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM club_registrations WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM phone_change_requests WHERE user_id=?', (uid,))
        conn.execute('DELETE FROM users WHERE id=?', (uid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/admin/reset-password/<int:uid>', methods=['POST'])
def admin_reset_password(uid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    admin_pass = data.get('adminPassword', '').strip()
    new_pwd = data.get('password', '').strip()
    if not admin_pass:
        return jsonify({'error': '请输入您的管理员密码以确认操作'}), 400
    if not new_pwd or len(new_pwd) < 4:
        return jsonify({'error': '新密码至少4位'}), 400
    conn = db.get_conn()
    try:
        admin_row = conn.execute('SELECT password FROM users WHERE id=?', (user['id'],)).fetchone()
        if not admin_row or admin_row['password'] != admin_pass:
            return jsonify({'error': '管理员密码错误'}), 400
        conn.execute('UPDATE users SET password=? WHERE id=?', (new_pwd, uid))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/upload-enroll-attachment', methods=['POST'])
def upload_enroll_attachment():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '请选择文件'}), 400
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    if ext not in ('jpg', 'jpeg', 'png', 'webp', 'gif', 'pdf', 'doc', 'docx'):
        return jsonify({'error': '仅支持图片、PDF或Word文件'}), 400
    import os, time
    upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'enroll_attachments')
    os.makedirs(upload_dir, exist_ok=True)
    filename = f'{user["id"]}_{int(time.time())}_{f.filename}'
    filepath = os.path.join(upload_dir, filename)
    f.save(filepath)
    return jsonify({'success': True, 'filePath': filepath, 'fileName': f.filename})


@app.route('/api/club-registration', methods=['POST'])
def submit_registration():
    data = request.json or {}
    club = data.get('clubName', '').strip()
    if not club:
        return jsonify({'error': '请指定社团'}), 400
    student_name = data.get('studentName', '').strip()
    student_phone = data.get('studentPhone', '').strip()
    student_class = data.get('studentClass', '').strip()
    student_id_num = data.get('studentIdNum', '').strip()
    specialty = data.get('specialty', '').strip()
    department = data.get('department', '').strip()
    college = data.get('college', '').strip()
    attachment = data.get('attachment', '').strip()
    form_data = data.get('formData', '')
    if not student_name:
        return jsonify({'error': '请输入姓名'}), 400
    current_user = get_current_user()
    user_id = current_user['id'] if current_user else 0
    conn = db.get_conn()
    try:
        if user_id:
            existing = conn.execute('SELECT id FROM club_registrations WHERE club_name=? AND user_id=? AND status!="rejected"', (club, user_id)).fetchone()
            if existing:
                return jsonify({'error': '您已报名该社团，请勿重复报名'}), 400
            already_member = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=?', (club, user_id)).fetchone()
            if already_member:
                return jsonify({'error': '您已经是该社团成员，无需再次报名'}), 400
            member_count = conn.execute('SELECT COUNT(*) as c FROM club_members WHERE user_id=?', (user_id,)).fetchone()
            if member_count and member_count['c'] >= 2:
                return jsonify({'error': '每人最多加入2个社团，您已达到上限'}), 400
            pending_count = conn.execute('SELECT COUNT(*) as c FROM club_registrations WHERE user_id=? AND status="pending"', (user_id,)).fetchone()
            if pending_count and (member_count['c'] + pending_count['c']) >= 2:
                return jsonify({'error': '每人最多加入2个社团，您有待审批的报名'}), 400
        conn.execute('INSERT INTO club_registrations (club_name, student_name, student_phone, student_class, student_id_num, specialty, department, college, user_id, form_data, attachment) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                    (club, student_name, student_phone, student_class, student_id_num, specialty, department, college, user_id, json.dumps(form_data) if isinstance(form_data, dict) else form_data, attachment))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-registrations/<club_name>')
def get_registrations(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'student'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] != 'admin' and user['club_name'] != club_name:
        if user['role'] == 'student' and is_cadre_of_club(user['id'], club_name):
            pass
        else:
            return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, student_name, student_phone, student_class, student_id_num, specialty, department, status, form_data, created_at, reviewed_at FROM club_registrations WHERE club_name=? ORDER BY created_at DESC', (club_name,)).fetchall()
    finally:
        conn.close()
    result = []
    for r in rows:
        fd = r['form_data']
        try:
            if fd and fd.startswith('{'):
                fd = json.loads(fd)
        except:
            pass
        result.append({'id': r['id'], 'studentName': r['student_name'], 'studentPhone': r['student_phone'],
            'studentClass': r['student_class'], 'studentIdNum': r['student_id_num'], 'specialty': r['specialty'],
            'department': r['department'], 'status': r['status'], 'formData': fd,
            'time': local_time(r['created_at']), 'reviewedAt': local_time(r['reviewed_at'])})
    return jsonify({'success': True, 'data': result})


@app.route('/api/export-registrations/<club_name>')
def export_registrations(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] != 'admin' and user['club_name'] != club_name:
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, student_name, student_phone, form_data, created_at FROM club_registrations WHERE club_name=? ORDER BY created_at DESC', (club_name,)).fetchall()
    finally:
        conn.close()
    output = BytesIO()
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = '报名信息'
    ws.append(['序号', '姓名', '联系电话', '报名时间', '其他信息'])
    for i, r in enumerate(rows, 1):
        other = ''
        try:
            fd = json.loads(r['form_data']) if r['form_data'] else {}
            if isinstance(fd, dict):
                other = '; '.join([f'{k}: {v}' for k, v in fd.items()])
        except:
            other = str(r['form_data'] or '')
        ws.append([i, r['student_name'], r['student_phone'], r['created_at'], other])
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'{club_name}_报名信息.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/recruitments', methods=['GET', 'POST'])
def handle_recruitments():
    if request.method == 'GET':
        user = get_current_user()
        if not user:
            return jsonify({'error': '请先登录'}), 401
        start_date = request.args.get('start_date', '').strip()
        end_date = request.args.get('end_date', '').strip()
        date_where = ''
        date_params = []
        if start_date:
            date_where += ' AND date(created_at)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(created_at)<=?'
            date_params.append(end_date)
        conn = db.get_conn()
        try:
            if user['role'] == 'admin':
                rows = conn.execute(f'SELECT * FROM recruitments WHERE 1=1{date_where} ORDER BY created_at DESC', date_params).fetchall()
            elif user['role'] == 'user':
                rows = conn.execute(f'SELECT * FROM recruitments WHERE club_name=?{date_where} ORDER BY created_at DESC', [user['club_name']] + date_params).fetchall()
            elif user['role'] == 'teacher':
                tc_rows = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
                tc_clubs = [r['club_name'] for r in tc_rows]
                rows = conn.execute(f'SELECT * FROM recruitments WHERE 1=1{date_where} ORDER BY created_at DESC', date_params).fetchall()
            else:
                rows = conn.execute(f'SELECT * FROM recruitments WHERE status="approved"{date_where} ORDER BY created_at DESC', date_params).fetchall()
        finally:
            conn.close()
        result_data = [{'id': r['id'], 'clubName': r['club_name'], 'title': r['title'], 'description': r['description'], 'recruitType': r['recruit_type'], 'maxCount': r['max_count'], 'currentCount': r['current_count'], 'status': r['status'], 'deadline': r['deadline'], 'createdAt': local_time(r['created_at']), 'approvedAt': r['approved_at']} for r in rows]
        resp = {'success': True, 'data': result_data}
        if user['role'] == 'teacher':
            resp['managedClubs'] = tc_clubs
        return jsonify(resp)
    else:
        user = get_current_user()
        if not user or user['role'] not in ('user', 'admin', 'teacher'):
            return jsonify({'error': '请先登录'}), 401
        data = request.json or {}
        title = data.get('title', '').strip()
        description = data.get('description', '').strip()
        recruit_type = data.get('recruitType', 'member').strip()
        max_count = int(data.get('maxCount', 0))
        deadline = data.get('deadline', '').strip()
        if not title:
            return jsonify({'error': '请输入招募标题'}), 400
        club = user['club_name'] or ''
        if user['role'] == 'admin':
            club = data.get('clubName', '').strip() or ''
        elif user['role'] == 'teacher':
            club = data.get('clubName', '').strip() or user['club_name'] or ''
            if club:
                tc_chk = db.get_conn()
                try:
                    tc_row = tc_chk.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], club)).fetchone()
                finally:
                    tc_chk.close()
                if not tc_row:
                    return jsonify({'error': '您不是该社团的指导老师'}), 403
        if not club:
            return jsonify({'error': '无社团信息'}), 400
        conn = db.get_conn()
        try:
            type_name = '志愿者' if recruit_type == 'volunteer' else '活动成员'
            if user['role'] == 'admin':
                conn.execute('INSERT INTO recruitments (club_name, title, description, recruit_type, max_count, status, created_by, deadline) VALUES (?, ?, ?, ?, ?, "approved", ?, ?)', (club, title, description, recruit_type, max_count, user['id'], deadline))
                conn.commit()
                rid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                students = conn.execute('SELECT id FROM users WHERE role="student"').fetchall()
                for s in students:
                    send_notification(s['id'], f'📢 新招募通知', f'「{club}」发布了{type_name}招募「{title}」，快来报名吧', 'recruit', '/dashboard', conn=conn)
                club_leader = conn.execute('SELECT id FROM users WHERE role="user" AND club_name=?', (club,)).fetchone()
                if club_leader and club_leader['id'] != user['id']:
                    send_notification(club_leader['id'], f'📢 招募已发布', f'管理员为「{club}」发布了{type_name}招募「{title}」', 'recruit', '/dashboard', conn=conn)
                conn.commit()
            else:
                conn.execute('INSERT INTO recruitments (club_name, title, description, recruit_type, max_count, status, created_by, deadline) VALUES (?, ?, ?, ?, ?, "pending", ?, ?)', (club, title, description, recruit_type, max_count, user['id'], deadline))
                conn.commit()
                rid = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
                admins = conn.execute('SELECT id FROM users WHERE role="admin"').fetchall()
                for a in admins:
                    send_notification(a['id'], f'📢 新招募待审批', f'「{club}」发布了{type_name}招募「{title}」，请及时审批', 'recruit', '/dashboard', conn=conn)
                conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True, 'id': rid})


@app.route('/api/recruitments/<int:rid>', methods=['PUT', 'DELETE'])
def manage_recruitment(rid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rec = conn.execute('SELECT * FROM recruitments WHERE id=?', (rid,)).fetchone()
        if not rec:
            return jsonify({'error': '招募不存在'}), 404
        if request.method == 'DELETE':
            if user['role'] not in ('admin',) and user['club_name'] != rec['club_name']:
                if user['role'] == 'teacher':
                    tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], rec['club_name'])).fetchone()
                    if not tc_row:
                        return jsonify({'error': '无权限'}), 403
                else:
                    return jsonify({'error': '无权限'}), 403
            conn.execute('DELETE FROM recruitment_signups WHERE recruitment_id=?', (rid,))
            conn.execute('DELETE FROM recruitments WHERE id=?', (rid,))
            conn.commit()
        else:
            action = (request.json or {}).get('action', 'approve')
            if user['role'] == 'user':
                return jsonify({'error': '只有管理员可以审批招募'}), 403
            if user['role'] == 'teacher':
                tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], rec['club_name'])).fetchone()
                if not tc_row:
                    return jsonify({'error': '无权限审批该社团的招募'}), 403
            if action == 'approve':
                conn.execute('UPDATE recruitments SET status="approved", approved_at=CURRENT_TIMESTAMP WHERE id=?', (rid,))
                conn.commit()
                students = conn.execute('SELECT id FROM users WHERE role="student"').fetchall()
                type_name = '志愿者' if rec['recruit_type'] == 'volunteer' else '活动成员'
                for s in students:
                    send_notification(s['id'], f'📢 新招募通知', f'「{rec["club_name"]}」发布了{type_name}招募「{rec["title"]}」，快来报名吧', 'recruit', '/dashboard', conn=conn)
                club_users = conn.execute('SELECT id FROM users WHERE club_name=? AND role="user"', (rec['club_name'],)).fetchall()
                for cu in club_users:
                    send_notification(cu['id'], '✅ 招募已通过', f'您的招募「{rec["title"]}」已通过审批', 'recruit', '/dashboard', conn=conn)
                club_teachers = conn.execute('SELECT user_id FROM teacher_clubs WHERE club_name=?', (rec['club_name'],)).fetchall()
                for ct in club_teachers:
                    send_notification(ct['user_id'], '✅ 招募已通过', f'「{rec["club_name"]}」的招募「{rec["title"]}」已通过审批', 'recruit', '/dashboard', conn=conn)
                conn.commit()
            elif action == 'reject':
                conn.execute('UPDATE recruitments SET status="rejected" WHERE id=?', (rid,))
                conn.commit()
                club_users = conn.execute('SELECT id FROM users WHERE club_name=? AND role="user"', (rec['club_name'],)).fetchall()
                for cu in club_users:
                    send_notification(cu['id'], '❌ 招募未通过', f'您的招募「{rec["title"]}」未通过审批', 'recruit', '/dashboard', conn=conn)
                club_teachers = conn.execute('SELECT user_id FROM teacher_clubs WHERE club_name=?', (rec['club_name'],)).fetchall()
                for ct in club_teachers:
                    send_notification(ct['user_id'], '❌ 招募未通过', f'「{rec["club_name"]}」的招募「{rec["title"]}」未通过审批', 'recruit', '/dashboard', conn=conn)
                conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/recruitments/<int:rid>/signup', methods=['POST'])
def signup_recruitment(rid):
    user = get_current_user()
    if not user or user['role'] != 'student':
        return jsonify({'error': '仅学生可报名'}), 403
    conn = db.get_conn()
    try:
        rec = conn.execute('SELECT * FROM recruitments WHERE id=?', (rid,)).fetchone()
        if not rec:
            return jsonify({'error': '招募不存在'}), 404
        if rec['status'] != 'approved':
            return jsonify({'error': '该招募未通过审批'}), 400
        if rec['max_count'] > 0 and rec['current_count'] >= rec['max_count']:
            return jsonify({'error': '招募人数已满'}), 400
        if rec['deadline']:
            try:
                import datetime
                dl = datetime.datetime.strptime(rec['deadline'], '%Y-%m-%dT%H:%M')
                if datetime.datetime.now() > dl:
                    return jsonify({'error': '报名已截止'}), 400
            except:
                try:
                    dl = datetime.datetime.strptime(rec['deadline'], '%Y-%m-%d %H:%M')
                    if datetime.datetime.now() > dl:
                        return jsonify({'error': '报名已截止'}), 400
                except:
                    try:
                        dl = datetime.datetime.strptime(rec['deadline'], '%Y-%m-%d')
                        if datetime.datetime.now().date() > dl.date():
                            return jsonify({'error': '报名已截止'}), 400
                    except:
                        pass
        existing = conn.execute('SELECT id FROM recruitment_signups WHERE recruitment_id=? AND user_id=?', (rid, user['id'])).fetchone()
        if existing:
            return jsonify({'error': '您已报名，请勿重复报名'}), 400
        student_name = user['username']
        student_class = ''
        student_id_num = ''
        student_phone = ''
        student_college = ''
        profile = conn.execute('SELECT * FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
        if profile:
            student_name = profile['real_name'] or student_name
            student_class = profile['class_name'] or ''
            student_id_num = profile['student_id'] or ''
            student_phone = profile['phone'] or ''
            student_college = profile['college'] if 'college' in profile.keys() else ''
        conn.execute('INSERT INTO recruitment_signups (recruitment_id, user_id, student_name, student_class, student_id_num, student_phone, college, club_name) VALUES (?, ?, ?, ?, ?, ?, ?, ?)', (rid, user['id'], student_name, student_class, student_id_num, student_phone, student_college, rec['club_name']))
        conn.execute('UPDATE recruitments SET current_count=current_count+1 WHERE id=?', (rid,))
        conn.commit()
        if rec['created_by']:
            creator = conn.execute('SELECT id FROM users WHERE id=?', (rec['created_by'],)).fetchone()
            if creator:
                type_name = '志愿者' if rec['recruit_type'] == 'volunteer' else '活动成员'
                send_notification(creator['id'], f'📝 新报名通知', f'{student_name}报名了您的{type_name}招募「{rec["title"]}」', 'recruit', '/dashboard', conn=conn)
                conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/my-recruit-signups')
def my_recruit_signups():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    date_where = ''
    date_params = []
    if start_date:
        date_where += ' AND date(rs.signed_up_at)>=?'
        date_params.append(start_date)
    if end_date:
        date_where += ' AND date(rs.signed_up_at)<=?'
        date_params.append(end_date)
    conn = db.get_conn()
    try:
        rows = conn.execute(f'''SELECT rs.id, rs.recruitment_id, rs.student_name, rs.student_class, rs.signed_up_at,
            r.title, r.club_name, r.recruit_type, r.description, r.status as rec_status
            FROM recruitment_signups rs
            JOIN recruitments r ON rs.recruitment_id = r.id
            WHERE rs.user_id=?{date_where}
            ORDER BY rs.signed_up_at DESC''', [user['id']] + date_params).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [{'id': r['id'], 'recruitmentId': r['recruitment_id'],
        'studentName': r['student_name'], 'studentClass': r['student_class'], 'signedUpAt': r['signed_up_at'],
        'title': r['title'], 'clubName': r['club_name'], 'recruitType': r['recruit_type'],
        'description': r['description'], 'recStatus': r['rec_status']} for r in rows]})


@app.route('/api/recruitments/<int:rid>/signups')
def get_recruitment_signups(rid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rec = conn.execute('SELECT * FROM recruitments WHERE id=?', (rid,)).fetchone()
        if not rec:
            return jsonify({'error': '招募不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != rec['club_name']:
            if user['role'] == 'teacher':
                tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], rec['club_name'])).fetchone()
                if not tc_row:
                    pass
            elif user['role'] == 'student':
                pass
            else:
                return jsonify({'error': '无权限'}), 403
        rows = conn.execute('SELECT * FROM recruitment_signups WHERE recruitment_id=? ORDER BY signed_up_at', (rid,)).fetchall()
        result = []
        for r in rows:
            item = {'id': r['id'], 'studentName': r['student_name'], 'studentClass': r['student_class'], 'studentIdNum': r['student_id_num'], 'studentPhone': r['student_phone'] if 'student_phone' in r.keys() else '', 'signedUpAt': r['signed_up_at'], 'userId': r['user_id']}
            if r['user_id'] and r['user_id'] != 0:
                try:
                    prow = conn.execute('SELECT real_name, student_id, class_name, phone, email, grade FROM user_profiles WHERE user_id=?', (r['user_id'],)).fetchone()
                    if prow:
                        item['profileRealName'] = prow['real_name'] or ''
                        item['profileStudentId'] = prow['student_id'] or ''
                        item['profileClassName'] = prow['class_name'] or ''
                        item['profilePhone'] = prow['phone'] or ''
                        item['profileEmail'] = prow['email'] or ''
                        item['profileGrade'] = prow['grade'] or ''
                except:
                    pass
            result.append(item)
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/export-recruitment/<int:rid>')
def export_recruitment(rid):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rec = conn.execute('SELECT * FROM recruitments WHERE id=?', (rid,)).fetchone()
        if not rec:
            return jsonify({'error': '招募不存在'}), 404
        if user['role'] != 'admin' and user['club_name'] != rec['club_name']:
            if user['role'] == 'teacher':
                tc_row = conn.execute('SELECT id FROM teacher_clubs WHERE user_id=? AND club_name=?', (user['id'], rec['club_name'])).fetchone()
                if not tc_row:
                    pass
            elif user['role'] == 'student':
                pass
            else:
                return jsonify({'error': '无权限'}), 403
        rows = conn.execute('SELECT * FROM recruitment_signups WHERE recruitment_id=? ORDER BY signed_up_at', (rid,)).fetchall()
    finally:
        conn.close()
    try:
        from openpyxl import Workbook
    except ImportError:
        return jsonify({'error': '导出功能需要安装 openpyxl 库，请在服务器执行: pip install openpyxl'}), 500
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = '招募报名信息'
    ws.append(['序号', '姓名', '学院', '班级', '学号', '报名时间'])
    for i, r in enumerate(rows, 1):
        ws.append([i, r['student_name'], (r['college'] if 'college' in r.keys() else '') or '', r['student_class'], r['student_id_num'], r['signed_up_at']])
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'{rec["club_name"]}_{rec["title"]}_报名信息.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/all-club-details')
def all_club_details():
    club_info_map = get_all_base_clubs()
    conn = db.get_conn()
    try:
        profiles = {}
        for p in conn.execute('SELECT club_name, description, star_rating, show_star, emblem_url, president, category, guiding_unit FROM club_profiles').fetchall():
            profiles[p['club_name']] = {'description': p['description'], 'starRating': p['star_rating'], 'showStar': p['show_star'], 'emblem_url': p['emblem_url'] if 'emblem_url' in p.keys() else '', 'president': p['president'] if 'president' in p.keys() else '', 'category': p['category'] if 'category' in p.keys() else '', 'guidingUnit': p['guiding_unit'] if 'guiding_unit' in p.keys() else ''}
        teachers_map = {}
        for t in conn.execute('SELECT club_name, teacher_name, photo_path, introduction FROM club_teachers').fetchall():
            teachers_map.setdefault(t['club_name'], []).append({'name': t['teacher_name'], 'photo': t['photo_path'], 'introduction': t['introduction']})
        depts_map = {}
        for d in conn.execute('SELECT club_name, dept_name, description, parent_id FROM club_departments').fetchall():
            depts_map.setdefault(d['club_name'], []).append({'name': d['dept_name'], 'description': d['description'], 'parentId': d['parent_id'] if 'parent_id' in d.keys() else 0})
    finally:
        conn.close()
    result = []
    for norm, info in club_info_map.items():
        cn = info.get('original_name', norm)
        p = profiles.get(cn, profiles.get(norm, {'description': '', 'starRating': 0, 'showStar': 0, 'emblem_url': '', 'president': '', 'category': '', 'guidingUnit': ''}))
        tl = teachers_map.get(cn, teachers_map.get(norm, []))
        dl = depts_map.get(cn, depts_map.get(norm, []))
        base_teacher_str = info.get('teacher', '')
        if base_teacher_str:
            for tn in base_teacher_str.split('、'):
                tn = tn.strip()
                if tn and not any(t['name'] == tn for t in tl):
                    tl.append({'name': tn, 'photo': '', 'introduction': ''})
        result.append({'clubName': cn, 'description': p['description'], 'starRating': p['starRating'], 'showStar': p['showStar'], 'emblemUrl': p.get('emblem_url', ''), 'president': p.get('president', ''), 'category': p.get('category', ''), 'guidingUnit': p.get('guidingUnit', ''), 'teachers': tl, 'departments': dl})
    return jsonify({'success': True, 'data': result})


@app.route('/api/health')
def health():
    return jsonify({'status': 'ok', 'timestamp': datetime.now().isoformat()})


@app.route('/api/club-recommend', methods=['POST'])
def club_recommend():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 403
    data = request.get_json() or {}
    student_input = data.get('input', '').strip()
    if not student_input:
        return jsonify({'error': '请输入你的兴趣或特长'}), 400
    conn = db.get_conn()
    try:
        club_rows = conn.execute('SELECT club_name, description, star_rating, category, emblem_url FROM club_profiles').fetchall()
        dept_rows = conn.execute('SELECT club_name, dept_name, description FROM club_departments').fetchall()
        dept_map = {}
        for d in dept_rows:
            dept_map.setdefault(d['club_name'], []).append({'name': d['dept_name'], 'desc': d['description'] or ''})
        teacher_map = {}
        for t in conn.execute('SELECT club_name, teacher_name FROM club_teachers').fetchall():
            teacher_map.setdefault(t['club_name'], []).append(t['teacher_name'])
        session_rows = conn.execute('SELECT club_name, activity_name FROM checkin_sessions WHERE is_completed=1').fetchall()
        activity_map = {}
        for s in session_rows:
            activity_map.setdefault(s['club_name'], []).append(s['activity_name'] or '')
        upload_rows = conn.execute("SELECT club_name, description FROM club_uploads WHERE status='approved' AND description LIKE '[总结]%'").fetchall()
        summary_map = {}
        for u in upload_rows:
            summary_map.setdefault(u['club_name'], []).append(u['description'])
    finally:
        conn.close()
    SKILL_KEYWORDS = {
        '摄影': ['拍照', '摄影', '照片', '相机', '镜头', '拍摄', '图片', '视频', '剪辑', '后期', '修图', 'ps', 'pr', '剪映'],
        '视频': ['视频', '剪辑', '拍摄', '短片', 'vlog', '后期', 'pr', '剪映', '录制', '直播'],
        '新媒体': ['新媒体', '运营', '公众号', '微博', '抖音', '小红书', '推文', '文案', '排版'],
        '设计': ['设计', '海报', 'ps', 'ai', '美工', '视觉', 'ui', '平面设计', '排版'],
        '文案': ['文案', '写作', '文字', '编辑', '策划', '写作', '稿件', '文章', '报告'],
        '绘画': ['绘画', '画画', '手绘', '素描', '水彩', '国画', '油画', '漫画', '动漫', '板绘', '插画'],
        '书法': ['书法', '毛笔', '硬笔', '字帖', '楷书', '行书', '练字'],
        '音乐': ['音乐', '唱歌', '乐器', '吉他', '钢琴', '架子鼓', '古筝', '合唱', '声乐'],
        '舞蹈': ['舞蹈', '跳舞', '街舞', '民族舞', '拉丁', '芭蕾', '编舞'],
        '体育': ['体育', '运动', '篮球', '足球', '排球', '乒乓球', '羽毛球', '跑步', '健身', '武术', '太极'],
        '志愿': ['志愿', '公益', '支教', '社区', '服务', '奉献', '帮助', '爱心', '慈善'],
        '教育': ['教育', '教学', '辅导', '培训', '课程', '学前', '幼儿', '儿童', '耐心'],
        '手工': ['手工', 'diy', '制作', '编织', '剪纸', '陶艺', '布艺', '折纸'],
        '科技': ['科技', '编程', '代码', '计算机', '互联网', '人工智能', 'ai', '机器人', '电子', '创客'],
        '演讲': ['演讲', '辩论', '主持', '口才', '表达', '朗诵'],
        '组织': ['组织', '策划', '管理', '协调', '领导', '安排', '统筹', '外联'],
        '外语': ['英语', '日语', '韩语', '外语', '翻译', '口语'],
        '表演': ['表演', '话剧', '戏剧', '小品', '相声', '演出'],
    }
    CATEGORY_SKILLS = {
        '志愿公益类': ['志愿', '公益', '支教', '社区', '服务', '教育', '耐心', '帮助'],
        '志愿服务类': ['志愿', '公益', '支教', '社区', '服务', '教育', '耐心', '帮助'],
        '学术科技类': ['科技', '编程', '研究', '学术', '创新', '实验', '计算机'],
        '文化体育类': ['体育', '舞蹈', '音乐', '绘画', '书法', '表演', '手工'],
        '文化类': ['绘画', '书法', '手工', '表演', '动漫', '创作'],
        '艺术类': ['舞蹈', '音乐', '绘画', '表演', '设计', '创意'],
        '媒体类': ['摄影', '视频', '新媒体', '文案', '播音', '主持', '写作'],
        '体育类': ['体育', '运动', '篮球', '足球', '跑步', '健身', '武术'],
        '创新创业类': ['创新', '创业', '商业', '策划', '设计', '新媒体', '运营'],
        '思想政治类': ['演讲', '辩论', '组织', '理论', '研究'],
        '综合类': ['组织', '策划', '管理', '协调', '演讲', '新媒体', '志愿'],
        '自律互助类': ['组织', '管理', '协调', '服务'],
    }
    DEPT_SKILLS = {
        '宣传': ['摄影', '视频', '新媒体', '设计', '文案', '绘画'],
        '策划': ['文案', '组织', '策划', '创意'],
        '外联': ['组织', '演讲', '外语', '沟通'],
        '秘书': ['文案', '组织', '管理'],
        '技术': ['科技', '视频', '设计'],
        '文艺': ['音乐', '舞蹈', '表演', '绘画'],
        '体育': ['体育'],
        '教学': ['教育', '书法', '绘画', '音乐', '手工'],
        '活动': ['志愿', '组织', '策划', '教育'],
        '志愿': ['志愿', '公益', '教育', '服务'],
        '组织': ['组织', '管理', '策划', '协调'],
        '后勤': ['组织', '管理'],
        '文体': ['体育', '舞蹈', '音乐', '表演', '组织'],
        '权益': ['演讲', '组织', '沟通'],
        '主席': ['组织', '管理', '策划', '领导'],
        '仪仗': ['体育', '组织'],
        '绘画': ['绘画', '手工', '设计'],
        '播音': ['演讲', '新媒体', '文案'],
        '科创': ['科技', '编程', '创新'],
    }
    input_lower = student_input.lower()
    student_skills = set()
    skill_weights = {}
    for skill, keywords in SKILL_KEYWORDS.items():
        hit_count = sum(1 for kw in keywords if kw in input_lower)
        if hit_count > 0:
            student_skills.add(skill)
            skill_weights[skill] = hit_count
    results = []
    for club in club_rows:
        cn = club['club_name']
        desc = (club['description'] or '').replace('<br>', ' ').replace('<', ' ').replace('>', ' ')
        category = club['category'] or ''
        depts = dept_map.get(cn, [])
        activities = activity_map.get(cn, [])
        teachers = teacher_map.get(cn, [])
        summaries = summary_map.get(cn, [])
        club_name_only = cn.lower()
        dept_names_text = ' '.join(d['name'] for d in depts).lower()
        club_full_text = (cn + ' ' + desc[:200] + ' ' + category + ' ' + dept_names_text + ' ' + ' '.join(activities[:5])).lower()
        matched_depts = []
        skill_depths = {}
        for skill in student_skills:
            skill_kws = SKILL_KEYWORDS.get(skill, [])
            if any(kw in club_name_only for kw in skill_kws):
                skill_depths[skill] = 3
            elif any(kw in dept_names_text for kw in skill_kws):
                skill_depths[skill] = 2
            elif any(kw in club_full_text for kw in skill_kws):
                skill_depths[skill] = 0.5
            else:
                skill_depths[skill] = 0
            if category in CATEGORY_SKILLS and skill in CATEGORY_SKILLS[category]:
                skill_depths[skill] = max(skill_depths.get(skill, 0), 2)
        for dept in depts:
            dept_name_lower = dept['name'].lower()
            for skill in student_skills:
                for dk, dsk in DEPT_SKILLS.items():
                    if dk in dept_name_lower:
                        if skill in dsk:
                            skill_depths[skill] = max(skill_depths.get(skill, 0), 1.5)
                            if dept['name'] not in matched_depts:
                                matched_depts.append(dept['name'])
                        break
        if not student_skills:
            common_chars = set(input_lower) & set(club_full_text)
            if len(common_chars) > 5:
                skill_depths['generic'] = 0.5
        matched_skills = {s for s, d in skill_depths.items() if d > 0}
        if not matched_skills:
            continue
        match_reasons = list(matched_skills - {'generic'})[:4]
        if student_skills:
            total_weight = sum(skill_weights.values())
            weighted_depth = sum(skill_depths.get(s, 0) * skill_weights.get(s, 1) for s in student_skills) / max(total_weight, 1)
            name_matched = any(skill_depths.get(s, 0) >= 3 for s in student_skills)
            category_matched = any(skill_depths.get(s, 0) >= 2 for s in student_skills)
            match_pct = min(99, round(45 + (weighted_depth / 3.0) * 35 + (12 if category_matched else 0) + min(len(matched_depts), 3) * 4 + (10 if name_matched else 0)))
        else:
            match_pct = min(55, len(common_chars))
        if match_pct < 30:
            continue
        reason_parts = []
        if match_reasons:
            skill_names = match_reasons[:3]
            reason_parts.append('根据你的' + '、'.join(skill_names) + '特长')
        if matched_depts:
            dept_reasons = []
            for md in matched_depts[:2]:
                md_lower = md.lower()
                md_skills = []
                for dk, dsk in DEPT_SKILLS.items():
                    if dk in md_lower:
                        md_skills = [s for s in dsk if s in student_skills]
                        break
                if md_skills:
                    dept_reasons.append(md + '（' + '、'.join(md_skills[:2]) + '）')
                else:
                    dept_reasons.append(md)
            reason_parts.append('推荐加入' + '、'.join(dept_reasons))
        if category:
            cat_match = False
            for skill in student_skills:
                if category in CATEGORY_SKILLS and skill in CATEGORY_SKILLS[category]:
                    cat_match = True
                    break
            if cat_match:
                reason_parts.append(f'{category}社团与你的兴趣匹配')
        reason = '，'.join(reason_parts) if reason_parts else '与你的描述有一定关联'
        results.append({
            'clubName': cn,
            'category': category,
            'description': desc[:120],
            'starRating': club['star_rating'] or 0,
            'emblemUrl': club['emblem_url'] or '',
            'matchPct': match_pct,
            'reason': reason,
            'matchedDepts': matched_depts[:3],
            'teachers': teachers[:3],
            'activityCount': len(activities),
        })
    results.sort(key=lambda x: x['matchPct'], reverse=True)
    results = results[:5]
    if QWEN_API_KEY and student_input:
        try:
            club_brief = '\n'.join([f'- {c["clubName"]}({c["category"]}): {c["description"][:60]}' for c in results[:5]])
            messages = [
                {'role': 'system', 'content': '你是社团推荐助手。根据学生描述和候选社团，为每个社团生成一句简短的推荐理由（15字以内）。只返回JSON数组，格式：[{"club":"社团名","reason":"理由"}]'},
                {'role': 'user', 'content': f'学生描述：{student_input}\n候选社团：\n{club_brief}'}
            ]
            ai_result = call_llm_api(messages, max_tokens=300)
            if ai_result:
                import json as _json
                ai_list = _json.loads(ai_result.strip().removeprefix('```json').removeprefix('```').removesuffix('```'))
                ai_map = {item['club']: item['reason'] for item in ai_list if isinstance(item, dict) and 'club' in item}
                for r in results:
                    if r['clubName'] in ai_map:
                        r['reason'] = ai_map[r['clubName']]
        except Exception:
            pass
    return jsonify({'success': True, 'data': results})


@app.route('/api/online-count')
def online_count():
    conn = db.get_conn()
    try:
        cnt = conn.execute('SELECT COUNT(*) as c FROM online_activity_data').fetchone()['c']
    finally:
        conn.close()
    return jsonify({'count': cnt})


@app.route('/api/online-stats-all')
def online_stats_all():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        date_where = ''
        join_date_where = ''
        date_params = []
        if start_date:
            date_where += ' AND date(created_at)>=?'
            join_date_where += ' AND date(cs.created_at)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(created_at)<=?'
            join_date_where += ' AND date(cs.created_at)<=?'
            date_params.append(end_date)
        session_rows = conn.execute(f'SELECT club_name, COUNT(*) as cnt FROM checkin_sessions WHERE 1=1{date_where} GROUP BY club_name', date_params).fetchall()
        checkin_rows = conn.execute(f'SELECT cr.club_name, COUNT(*) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE 1=1{join_date_where} GROUP BY cr.club_name', date_params).fetchall()
        last_act_rows = conn.execute(f'SELECT club_name, MAX(created_at) as last_time FROM checkin_sessions WHERE 1=1{date_where} GROUP BY club_name', date_params).fetchall()
        total_sessions = conn.execute(f'SELECT COUNT(*) as c FROM checkin_sessions WHERE 1=1{date_where}', date_params).fetchone()['c']
        total_checkins = conn.execute(f'SELECT COUNT(*) as c FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE 1=1{join_date_where}', date_params).fetchone()['c']
        try:
            warning_rows = conn.execute(f'SELECT club_name, COUNT(*) as cnt FROM checkin_sessions WHERE warning="warned"{date_where} GROUP BY club_name', date_params).fetchall()
            warning_detail_rows = conn.execute(f'SELECT club_name, activity_name, warning_reason FROM checkin_sessions WHERE warning="warned"{date_where}', date_params).fetchall()
        except:
            conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning TEXT DEFAULT ""')
            conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning_reason TEXT DEFAULT ""')
            conn.commit()
            warning_rows = []
            warning_detail_rows = []
        try:
            profile_rows = conn.execute('SELECT club_name, guiding_unit, president, category FROM club_profiles').fetchall()
        except:
            conn.execute('ALTER TABLE club_profiles ADD COLUMN guiding_unit TEXT DEFAULT ""')
            conn.commit()
            profile_rows = conn.execute('SELECT club_name, guiding_unit, president, category FROM club_profiles').fetchall()
        teacher_rows = conn.execute('SELECT cs.club_name, cs.teacher_ids FROM checkin_sessions cs WHERE cs.teacher_ids IS NOT NULL AND cs.teacher_ids != ""').fetchall()
        teacher_info_rows = conn.execute('SELECT club_name, teacher_name FROM club_teachers').fetchall()
    finally:
        conn.close()
    session_map = {r['club_name']: r['cnt'] for r in session_rows}
    checkin_map = {r['club_name']: r['cnt'] for r in checkin_rows}
    last_map = {r['club_name']: r['last_time'] for r in last_act_rows}
    warning_map = {r['club_name']: r['cnt'] for r in warning_rows}
    warning_detail_map = {}
    for r in warning_detail_rows:
        warning_detail_map.setdefault(r['club_name'], []).append({'activityName': r['activity_name'] or '', 'reason': r['warning_reason'] or ''})
    profile_map = {}
    for r in profile_rows:
        profile_map[r['club_name']] = {'guidingUnit': r['guiding_unit'] if 'guiding_unit' in r.keys() else '', 'president': r['president'] if 'president' in r.keys() else '', 'category': r['category'] if 'category' in r.keys() else ''}
    teacher_name_map = {}
    for r in teacher_info_rows:
        teacher_name_map.setdefault(r['club_name'], set()).add(r['teacher_name'])
    session_teacher_map = {}
    for r in teacher_rows:
        tids = r['teacher_ids'].split(',') if r['teacher_ids'] else []
        for tid in tids:
            tid = tid.strip()
            if tid:
                session_teacher_map.setdefault(r['club_name'], set()).add(tid)
    all_clubs = set(list(session_map.keys()) + list(checkin_map.keys()))
    clubs = []
    for cn in all_clubs:
        p = profile_map.get(cn, {'guidingUnit': '', 'president': '', 'category': ''})
        tids_set = session_teacher_map.get(cn, set())
        tnames_set = teacher_name_map.get(cn, set())
        teacher_names_str = '、'.join(sorted(tnames_set)) if tnames_set else ''
        warnings = warning_detail_map.get(cn, [])
        warning_str = '; '.join([w['activityName'] + ': ' + w['reason'] for w in warnings]) if warnings else ''
        clubs.append({
            'clubName': cn,
            'sessionCount': session_map.get(cn, 0),
            'totalCheckins': checkin_map.get(cn, 0),
            'lastActivity': last_map.get(cn, ''),
            'warningCount': warning_map.get(cn, 0),
            'warningDetail': warning_str,
            'guidingUnit': p['guidingUnit'],
            'president': p['president'],
            'category': p['category'],
            'teacherNames': teacher_names_str
        })
    clubs.sort(key=lambda x: x['sessionCount'], reverse=True)
    return jsonify({
        'success': True,
        'data': {
            'totalSessions': total_sessions,
            'totalClubs': len(all_clubs),
            'totalCheckins': total_checkins,
            'clubs': clubs
        }
    })


@app.route('/api/export-online-stats')
def export_online_stats():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    except ImportError:
        return jsonify({'error': '导出功能需要安装 openpyxl 库，请运行: pip install openpyxl'}), 500
    conn = db.get_conn()
    try:
        date_where = ''
        join_date_where = ''
        date_params = []
        if start_date:
            date_where += ' AND date(created_at)>=?'
            join_date_where += ' AND date(cs.created_at)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(created_at)<=?'
            join_date_where += ' AND date(cs.created_at)<=?'
            date_params.append(end_date)
        session_rows = conn.execute(f'SELECT club_name, COUNT(*) as cnt FROM checkin_sessions WHERE 1=1{date_where} GROUP BY club_name', date_params).fetchall()
        checkin_rows = conn.execute(f'SELECT cr.club_name, COUNT(*) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE 1=1{join_date_where} GROUP BY cr.club_name', date_params).fetchall()
        try:
            warning_rows = conn.execute(f'SELECT club_name, COUNT(*) as cnt FROM checkin_sessions WHERE warning="warned"{date_where} GROUP BY club_name', date_params).fetchall()
            warning_detail_rows = conn.execute(f'SELECT club_name, activity_name, warning_reason FROM checkin_sessions WHERE warning="warned"{date_where}', date_params).fetchall()
        except:
            conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning TEXT DEFAULT ""')
            conn.execute('ALTER TABLE checkin_sessions ADD COLUMN warning_reason TEXT DEFAULT ""')
            conn.commit()
            warning_rows = []
            warning_detail_rows = []
        try:
            profile_rows = conn.execute('SELECT club_name, guiding_unit, president, category FROM club_profiles').fetchall()
        except:
            conn.execute('ALTER TABLE club_profiles ADD COLUMN guiding_unit TEXT DEFAULT ""')
            conn.commit()
            profile_rows = conn.execute('SELECT club_name, guiding_unit, president, category FROM club_profiles').fetchall()
        teacher_info_rows = conn.execute('SELECT club_name, teacher_name FROM club_teachers').fetchall()
    finally:
        conn.close()
    session_map = {r['club_name']: r['cnt'] for r in session_rows}
    checkin_map = {r['club_name']: r['cnt'] for r in checkin_rows}
    warning_map = {r['club_name']: r['cnt'] for r in warning_rows}
    warning_detail_map = {}
    for r in warning_detail_rows:
        warning_detail_map.setdefault(r['club_name'], []).append({'activityName': r['activity_name'] or '', 'reason': r['warning_reason'] or ''})
    profile_map = {}
    for r in profile_rows:
        profile_map[r['club_name']] = {'guidingUnit': r['guiding_unit'] if 'guiding_unit' in r.keys() else '', 'president': r['president'] if 'president' in r.keys() else '', 'category': r['category'] if 'category' in r.keys() else ''}
    teacher_name_map = {}
    for r in teacher_info_rows:
        teacher_name_map.setdefault(r['club_name'], set()).add(r['teacher_name'])
    all_clubs = set(list(session_map.keys()) + list(checkin_map.keys()))
    clubs = []
    for cn in all_clubs:
        p = profile_map.get(cn, {'guidingUnit': '', 'president': '', 'category': ''})
        tnames_set = teacher_name_map.get(cn, set())
        teacher_names_str = '、'.join(sorted(tnames_set)) if tnames_set else ''
        warnings = warning_detail_map.get(cn, [])
        warning_str = '; '.join([w['activityName'] + ': ' + w['reason'] for w in warnings]) if warnings else ''
        clubs.append({
            'clubName': cn,
            'sessionCount': session_map.get(cn, 0),
            'totalCheckins': checkin_map.get(cn, 0),
            'warningCount': warning_map.get(cn, 0),
            'warningDetail': warning_str,
            'guidingUnit': p['guidingUnit'],
            'president': p['president'],
            'category': p['category'],
            'teacherNames': teacher_names_str
        })
    clubs.sort(key=lambda x: x['sessionCount'], reverse=True)
    wb = Workbook()
    ws = wb.active
    ws.title = '社团活动数据分析'
    headers = ['社团名称', '活动次数', '业务指导单位', '指导老师', '社长', '社团类型', '签到人次', '平均签到', '警告次数', '警告详情']
    header_fill = PatternFill(start_color='667eea', end_color='667eea', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF', size=11)
    thin_border = Border(left=Side(style='thin', color='d0d0e8'), right=Side(style='thin', color='d0d0e8'), top=Side(style='thin', color='d0d0e8'), bottom=Side(style='thin', color='d0d0e8'))
    warn_fill = PatternFill(start_color='FFF3CD', end_color='FFF3CD', fill_type='solid')
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    for i, c in enumerate(clubs, 2):
        avg = round(c['totalCheckins'] / c['sessionCount'], 1) if c['sessionCount'] > 0 else 0
        row_data = [c['clubName'], c['sessionCount'], c['guidingUnit'] or '-', c['teacherNames'] or '-', c['president'] or '-', c['category'] or '-', c['totalCheckins'], avg, c['warningCount'], c['warningDetail'] or '-']
        for col, val in enumerate(row_data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.border = thin_border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
            if c['warningCount'] > 0:
                cell.fill = warn_fill
    for col in ws.columns:
        max_len = max(len(str(c.value or '')) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name='社团活动数据分析.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/excellent-club-scores')
def excellent_club_scores():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        date_where = ''
        date_params = []
        if start_date:
            date_where += ' AND date(created_at)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(created_at)<=?'
            date_params.append(end_date)
        all_clubs_rows = conn.execute('SELECT club_name FROM club_profiles').fetchall()
        all_club_names = [r['club_name'] for r in all_clubs_rows]
        star_map = {}
        for r in conn.execute('SELECT club_name, star_rating FROM club_profiles').fetchall():
            star_map[r['club_name']] = r['star_rating'] if r['star_rating'] else 0
        activity_count_map = {}
        for r in conn.execute(f'SELECT club_name, COUNT(*) as cnt FROM checkin_sessions WHERE 1=1{date_where} GROUP BY club_name', date_params).fetchall():
            activity_count_map[r['club_name']] = r['cnt']
        teacher_session_rows = conn.execute(f'SELECT club_name, id FROM checkin_sessions WHERE teacher_ids IS NOT NULL AND teacher_ids != ""{date_where}', date_params).fetchall()
        teacher_guided_map = {}
        for r in teacher_session_rows:
            cn = r['club_name']
            teacher_guided_map[cn] = teacher_guided_map.get(cn, 0) + 1
        material_approved_map = {}
        for r in conn.execute("SELECT club_name, COUNT(DISTINCT group_id) as cnt FROM club_uploads WHERE status='approved' AND (source='upload' OR source='activity' OR source IS NULL) GROUP BY club_name").fetchall():
            material_approved_map[r['club_name']] = r['cnt']
        excellent_activity_map = {}
        ea_rows = conn.execute('SELECT group_id FROM excellent_activities').fetchall()
        for ea in ea_rows:
            gid = ea['group_id'] or ''
            club_name = None
            if gid.startswith('session_'):
                try:
                    sid = int(gid.replace('session_', ''))
                    sess = conn.execute('SELECT club_name FROM checkin_sessions WHERE id=?', (sid,)).fetchone()
                    if sess:
                        club_name = sess['club_name']
                except ValueError:
                    pass
            else:
                mat = conn.execute('SELECT club_name FROM club_uploads WHERE group_id=? LIMIT 1', (gid,)).fetchone()
                if mat:
                    club_name = mat['club_name']
            if club_name:
                excellent_activity_map[club_name] = excellent_activity_map.get(club_name, 0) + 1
    finally:
        conn.close()
    results = []
    max_activity = max(activity_count_map.values()) if activity_count_map else 1
    max_teacher = max(teacher_guided_map.values()) if teacher_guided_map else 0
    max_material = max(material_approved_map.values()) if material_approved_map else 0
    max_excellent = max(excellent_activity_map.values()) if excellent_activity_map else 0
    for cn in all_club_names:
        star = star_map.get(cn, 0)
        star_score_map = {5: 100, 4: 70, 3: 40}
        star_score = star_score_map.get(star, 0)
        act_count = activity_count_map.get(cn, 0)
        act_score = (act_count / max_activity) * 100 if max_activity > 0 else 0
        teacher_count = teacher_guided_map.get(cn, 0)
        teacher_score = (teacher_count / max_teacher) * 100 if max_teacher > 0 else 0
        material_count = material_approved_map.get(cn, 0)
        material_score = (material_count / max_material) * 100 if max_material > 0 else 0
        excellent_count = excellent_activity_map.get(cn, 0)
        excellent_score = (excellent_count / max_excellent) * 100 if max_excellent > 0 else 0
        combined = round(star_score * 0.15 + act_score * 0.35 + teacher_score * 0.20 + material_score * 0.10 + excellent_score * 0.10, 1)
        results.append({
            'clubName': cn,
            'starRating': star,
            'starScore': round(star_score, 1),
            'activityCount': act_count,
            'activityScore': round(act_score, 1),
            'teacherGuidedCount': teacher_count,
            'teacherScore': round(teacher_score, 1),
            'materialApprovedCount': material_count,
            'materialScore': round(material_score, 1),
            'excellentActivityCount': excellent_count,
            'excellentScore': round(excellent_score, 1),
            'combinedScore': combined,
        })
    results.sort(key=lambda x: x['combinedScore'], reverse=True)
    return jsonify({'success': True, 'data': results, 'maxValues': {
        'maxActivity': max_activity,
        'maxTeacher': max_teacher,
        'maxMaterial': max_material,
        'maxExcellent': max_excellent,
    }})


@app.route('/api/excellent-clubs', methods=['GET', 'POST'])
def handle_excellent_clubs():
    if request.method == 'GET':
        conn = db.get_conn()
        try:
            rows = conn.execute('SELECT club_name, selected_at, star_rating, activity_count, teacher_guided_count, material_approved_count, excellent_activity_count, combined_score FROM excellent_clubs ORDER BY selected_at').fetchall()
        finally:
            conn.close()
        if not rows:
            return jsonify({'success': True, 'data': []})
        club_names = [r['club_name'] for r in rows]
        conn = db.get_conn()
        try:
            profiles = {}
            for p in conn.execute('SELECT club_name, emblem_url, president, star_rating FROM club_profiles').fetchall():
                profiles[p['club_name']] = {'emblem_url': p['emblem_url'] if 'emblem_url' in p.keys() else '', 'president': p['president'] if 'president' in p.keys() else '', 'star_rating': p['star_rating'] if 'star_rating' in p.keys() else 0}
        finally:
            conn.close()
        result = []
        for r in rows:
            cn = r['club_name']
            p = profiles.get(cn, {'emblem_url': '', 'president': '', 'star_rating': 0})
            result.append({
                'clubName': cn, 'emblemUrl': p.get('emblem_url', ''), 'president': p.get('president', ''), 'starRating': p.get('star_rating', 0),
                'selectedAt': r['selected_at'] if 'selected_at' in r.keys() else '',
                'evalStarRating': r['star_rating'] if 'star_rating' in r.keys() else 0,
                'activityCount': r['activity_count'] if 'activity_count' in r.keys() else 0,
                'teacherGuidedCount': r['teacher_guided_count'] if 'teacher_guided_count' in r.keys() else 0,
                'materialApprovedCount': r['material_approved_count'] if 'material_approved_count' in r.keys() else 0,
                'excellentActivityCount': r['excellent_activity_count'] if 'excellent_activity_count' in r.keys() else 0,
                'combinedScore': r['combined_score'] if 'combined_score' in r.keys() else 0,
            })
        return jsonify({'success': True, 'data': result})
    else:
        user = get_current_user()
        if not user or user['role'] != 'admin':
            return jsonify({'error': '无权限'}), 403
        data = request.json or {}
        clubs = data.get('clubs', [])
        if not isinstance(clubs, list):
            clubs = [clubs] if clubs else []
        if len(clubs) > 24:
            return jsonify({'error': '最多选择24个社团'}), 400
        conn = db.get_conn()
        try:
            conn.execute('DELETE FROM excellent_clubs')
            for item in clubs:
                # 兼容两种格式：字符串 或 对象（含评分数据）
                if isinstance(item, str):
                    cn = item.strip()
                    star_r = act_c = tea_c = mat_c = exc_c = 0
                    combined = 0
                elif isinstance(item, dict):
                    cn = (item.get('clubName') or item.get('club_name') or '').strip()
                    star_r = item.get('starRating', 0) or 0
                    act_c = item.get('activityCount', 0) or 0
                    tea_c = item.get('teacherGuidedCount', 0) or 0
                    mat_c = item.get('materialApprovedCount', 0) or 0
                    exc_c = item.get('excellentActivityCount', 0) or 0
                    combined = item.get('combinedScore', 0) or 0
                else:
                    continue
                if cn:
                    conn.execute('''INSERT INTO excellent_clubs
                        (club_name, star_rating, activity_count, teacher_guided_count, material_approved_count, excellent_activity_count, combined_score)
                        VALUES (?,?,?,?,?,?,?)''',
                        (cn, star_r, act_c, tea_c, mat_c, exc_c, combined))
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True})


@app.route('/api/excellent-activities', methods=['GET', 'POST'])
def handle_excellent_activities():
    if request.method == 'GET':
        conn = db.get_conn()
        try:
            rows = conn.execute('SELECT group_id FROM excellent_activities ORDER BY selected_at').fetchall()
        finally:
            conn.close()
        group_ids = [r['group_id'] for r in rows]
        if not group_ids:
            return jsonify({'success': True, 'data': []})
        result = []
        conn = db.get_conn()
        try:
            for gid in group_ids:
                if gid.startswith('session_'):
                    sid = gid.replace('session_', '')
                    try:
                        sid_int = int(sid)
                    except ValueError:
                        continue
                    sess = conn.execute('SELECT cs.*, COUNT(cr.id) as checkin_count FROM checkin_sessions cs LEFT JOIN checkin_records cr ON cs.id=cr.session_id WHERE cs.id=? GROUP BY cs.id', (sid_int,)).fetchone()
                    if sess:
                        result.append({'groupId': gid, 'clubName': sess['club_name'], 'title': sess['activity_name'] or sess['club_name'] + '的活动', 'time': local_time(sess['created_at']), 'activityName': sess['activity_name'], 'locationName': sess['location_name'] if 'location_name' in sess.keys() else '', 'checkinCount': sess['checkin_count']})
                else:
                    materials = conn.execute('SELECT group_id, club_name, description, created_at FROM club_uploads WHERE group_id=? LIMIT 1', (gid,)).fetchall()
                    if materials:
                        m = materials[0]
                        desc = m['description'] or ''
                        title = desc.replace('[总结]', '').split('|||')[0] if desc.startswith('[总结]') else m['club_name'] + '的活动'
                        result.append({'groupId': gid, 'clubName': m['club_name'], 'title': title, 'time': local_time(m['created_at'])})
        finally:
            conn.close()
        return jsonify({'success': True, 'data': result})
    else:
        user = get_current_user()
        if not user or user['role'] != 'admin':
            return jsonify({'error': '无权限'}), 403
        data = request.json or {}
        group_ids = data.get('groupIds', [])
        if not isinstance(group_ids, list):
            group_ids = [group_ids] if group_ids else []
        conn = db.get_conn()
        try:
            conn.execute('DELETE FROM excellent_activities')
            for gid in group_ids:
                gid = gid.strip()
                if gid:
                    conn.execute('INSERT INTO excellent_activities (group_id) VALUES (?)', (gid,))
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True})


@app.route('/api/excellent-activities/toggle', methods=['POST'])
def toggle_excellent_activity():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    group_id = data.get('group_id', '').strip()
    if not group_id:
        return jsonify({'error': '缺少活动标识'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM excellent_activities WHERE group_id=?', (group_id,)).fetchone()
        if existing:
            conn.execute('DELETE FROM excellent_activities WHERE group_id=?', (group_id,))
            conn.commit()
            return jsonify({'success': True, 'action': 'removed', 'message': '已取消优秀'})
        else:
            conn.execute('INSERT INTO excellent_activities (group_id) VALUES (?)', (group_id,))
            conn.commit()
            club_name = ''
            activity_name = ''
            if group_id.startswith('session_'):
                sid = group_id.replace('session_', '')
                try:
                    sid_int = int(sid)
                    sess = conn.execute('SELECT club_name, activity_name FROM checkin_sessions WHERE id=?', (sid_int,)).fetchone()
                    if sess:
                        club_name = sess['club_name']
                        activity_name = sess['activity_name'] or ''
                except ValueError:
                    pass
            else:
                mat = conn.execute('SELECT club_name, description FROM club_uploads WHERE group_id=? LIMIT 1', (group_id,)).fetchone()
                if mat:
                    club_name = mat['club_name']
                    desc = mat['description'] or ''
                    if desc.startswith('[总结]'):
                        activity_name = desc.replace('[总结]', '').split('|||')[0]
            if club_name:
                leaders = conn.execute('SELECT id FROM users WHERE club_name=? AND role="user"', (club_name,)).fetchall()
                for l in leaders:
                    send_notification(l['id'], '🌟 活动被评为优秀', '您的活动「' + (activity_name or club_name + '的活动') + '」被评为优秀社团活动，将展示在首页', 'excellent', '/dashboard.html', conn=conn)
                conn.commit()
            return jsonify({'success': True, 'action': 'added', 'message': '已选为优秀'})
    finally:
        conn.close()


OFFCAMPUS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'offcampus')
os.makedirs(OFFCAMPUS_FOLDER, exist_ok=True)

TEMPLATE_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'templates')
os.makedirs(TEMPLATE_FOLDER, exist_ok=True)


@app.route('/api/offcampus-template', methods=['GET', 'POST', 'DELETE'])
def offcampus_template():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    meta_path = os.path.join(TEMPLATE_FOLDER, 'offcampus_template.meta')
    def find_template_file():
        for f in os.listdir(TEMPLATE_FOLDER):
            if f.startswith('offcampus_template') and not f.endswith('.meta'):
                return os.path.join(TEMPLATE_FOLDER, f)
        return None
    if request.method == 'GET':
        if not find_template_file() or not os.path.exists(meta_path):
            return jsonify({'success': True, 'hasTemplate': False})
        with open(meta_path, 'r', encoding='utf-8') as f:
            meta = json.load(f)
        return jsonify({'success': True, 'hasTemplate': True, 'fileName': meta.get('fileName', ''), 'uploadTime': meta.get('uploadTime', '')})
    elif request.method == 'POST':
        if user['role'] != 'admin':
            return jsonify({'error': '仅管理员可上传模板'}), 403
        if 'file' not in request.files:
            return jsonify({'error': '请选择文件'}), 400
        f = request.files['file']
        if not f or not f.filename:
            return jsonify({'error': '文件无效'}), 400
        ext = os.path.splitext(f.filename)[1]
        fname = 'offcampus_template' + ext
        save_path = os.path.join(TEMPLATE_FOLDER, fname)
        for old in os.listdir(TEMPLATE_FOLDER):
            if old.startswith('offcampus_template'):
                os.remove(os.path.join(TEMPLATE_FOLDER, old))
        f.save(save_path)
        meta = {'fileName': f.filename, 'uploadTime': datetime.now().strftime('%Y-%m-%d %H:%M')}
        with open(meta_path, 'w', encoding='utf-8') as mf:
            json.dump(meta, mf, ensure_ascii=False)
        return jsonify({'success': True, 'fileName': f.filename})
    elif request.method == 'DELETE':
        if user['role'] != 'admin':
            return jsonify({'error': '仅管理员可删除模板'}), 403
        for old in os.listdir(TEMPLATE_FOLDER):
            if old.startswith('offcampus_template'):
                os.remove(os.path.join(TEMPLATE_FOLDER, old))
        if os.path.exists(meta_path):
            os.remove(meta_path)
        return jsonify({'success': True})


@app.route('/api/offcampus-template/download')
def download_offcampus_template():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    for fname in os.listdir(TEMPLATE_FOLDER):
        if fname.startswith('offcampus_template') and not fname.endswith('.meta'):
            meta_path = os.path.join(TEMPLATE_FOLDER, 'offcampus_template.meta')
            download_name = fname
            if os.path.exists(meta_path):
                with open(meta_path, 'r', encoding='utf-8') as f:
                    meta = json.load(f)
                    download_name = meta.get('fileName', fname)
            return send_file(os.path.join(TEMPLATE_FOLDER, fname), as_attachment=True, download_name=download_name)
    return jsonify({'error': '模板不存在'}), 404


@app.route('/api/offcampus', methods=['GET', 'POST'])
def handle_offcampus():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    if request.method == 'GET':
        conn = db.get_conn()
        try:
            if user['role'] == 'admin':
                rows = conn.execute('SELECT * FROM offcampus_requests ORDER BY CASE status WHEN "pending" THEN 0 ELSE 1 END, created_at DESC').fetchall()
            else:
                rows = conn.execute('SELECT * FROM offcampus_requests WHERE club_name=? ORDER BY created_at DESC', (user.get('club_name', '') or user.get('clubName', ''),)).fetchall()
        finally:
            conn.close()
        result = []
        for r in rows:
            item = {'id': r['id'], 'clubName': r['club_name'], 'title': r['title'], 'location': r['location'], 'activityDate': r['activity_date'], 'description': r['description'], 'fileName': r['file_name'], 'status': r['status'], 'rejectReason': r['reject_reason'], 'submittedBy': r['submitted_by'], 'time': local_time(r['created_at'])}
            if r['file_path']:
                item['fileUrl'] = '/api/offcampus/file/' + str(r['id'])
            result.append(item)
        return jsonify({'success': True, 'data': result})
    else:
        if user['role'] not in ('user', 'admin'):
            return jsonify({'error': '无权限'}), 403
        club_name = user.get('club_name', '') or user.get('clubName', '')
        if not club_name:
            return jsonify({'error': '未关联社团，请联系管理员绑定社团'}), 400
        title = request.form.get('title', '').strip()
        location = request.form.get('location', '').strip()
        activity_date = request.form.get('activityDate', '').strip()
        description = request.form.get('description', '').strip()
        if not title:
            return jsonify({'error': '请填写活动名称'}), 400
        file_path = ''
        file_name = ''
        if 'file' in request.files:
            f = request.files['file']
            if f and f.filename:
                file_name = f.filename
                ext = os.path.splitext(f.filename)[1]
                fname = hashlib.md5((file_name + str(uuid.uuid4())).encode()).hexdigest()[:16] + ext
                key = 'offcampus/' + fname
                storage.save(f, key)
                file_path = key
        conn = db.get_conn()
        try:
            conn.execute('INSERT INTO offcampus_requests (club_name, title, location, activity_date, description, file_path, file_name, status, submitted_by) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
                         (club_name, title, location, activity_date, description, file_path, file_name, 'pending', user.get('username', '')))
            conn.commit()
        finally:
            conn.close()
        conn = db.get_conn()
        try:
            admins = conn.execute("SELECT id FROM users WHERE role='admin'").fetchall()
        finally:
            conn.close()
        for a in admins:
            send_notification(a['id'], '🏕️ 新校外活动申请', f'{club_name} 提交了校外活动申请「{title}」，请及时审批', 'offcampus', '/dashboard.html')
        return jsonify({'success': True})


@app.route('/api/offcampus/file/<int:rid>')
def offcampus_file(rid):
    conn = db.get_conn()
    try:
        row = conn.execute('SELECT file_path, file_name FROM offcampus_requests WHERE id=?', (rid,)).fetchone()
    finally:
        conn.close()
    if not row or not row['file_path']:
        return '文件不存在', 404
    path = storage.get_path(row['file_path'])
    if not path:
        url = storage.get_url(row['file_path'])
        if url:
            return jsonify({'url': url})
        return '文件不存在', 404
    return send_file(path, as_attachment=True, download_name=row['file_name'])


@app.route('/api/offcampus/download-attachments', methods=['GET'])
def download_offcampus_attachments():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    import zipfile
    club = request.args.get('club', '').strip()
    date_start = request.args.get('date_start', '').strip()
    date_end = request.args.get('date_end', '').strip()
    status_filter = request.args.get('status', '').strip()
    conn = db.get_conn()
    try:
        q = 'SELECT * FROM offcampus_requests WHERE file_path IS NOT NULL AND file_path != ""'
        params = []
        if club:
            q += ' AND club_name=?'
            params.append(club)
        if date_start:
            q += ' AND created_at>=?'
            params.append(date_start + ' 00:00:00')
        if date_end:
            q += ' AND created_at<=?'
            params.append(date_end + ' 23:59:59')
        if status_filter:
            q += ' AND status=?'
            params.append(status_filter)
        q += ' ORDER BY created_at DESC'
        rows = conn.execute(q, params).fetchall()
    finally:
        conn.close()
    if not rows:
        return jsonify({'error': '没有可下载的附件'}), 404
    buf = BytesIO()
    status_map = {'pending': '待审批', 'approved': '已通过', 'rejected': '已驳回'}
    actual_count = 0
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for r in rows:
            fp = r['file_path']
            if not fp:
                continue
            actual_path = storage.get_path(fp)
            if not actual_path:
                continue
            actual_count += 1
            club_dir = r['club_name']
            fn = r['file_name'] or os.path.basename(fp)
            name, ext = os.path.splitext(fn)
            arc_name = club_dir + '/' + name + '_' + str(r['id']) + ext
            zf.write(actual_path, arc_name)
    if actual_count == 0:
        return jsonify({'error': '附件文件不存在，可能已被删除'}), 404
    buf.seek(0)
    zip_name = '校外活动附件'
    if club:
        zip_name += '_' + club
    if date_start or date_end:
        zip_name += '_' + (date_start or '起') + '-' + (date_end or '止')
    zip_name += '.zip'
    return send_file(buf, as_attachment=True, download_name=zip_name, mimetype='application/zip')


@app.route('/api/offcampus/<int:rid>', methods=['PUT', 'DELETE'])
def update_offcampus(rid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    if request.method == 'DELETE':
        if user['role'] not in ('admin', 'user'):
            return jsonify({'error': '无权限'}), 403
        conn = db.get_conn()
        try:
            row = conn.execute('SELECT club_name, file_path FROM offcampus_requests WHERE id=?', (rid,)).fetchone()
            if not row:
                return jsonify({'error': '不存在'}), 404
            if user['role'] != 'admin' and row['club_name'] != (user.get('club_name', '') or user.get('clubName', '')):
                return jsonify({'error': '无权限'}), 403
            if row['file_path']:
                storage.delete(row['file_path'])
            conn.execute('DELETE FROM offcampus_requests WHERE id=?', (rid,))
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True})
    else:
        if user['role'] != 'admin':
            return jsonify({'error': '无权限'}), 403
        data = request.json or {}
        action = data.get('action', '')
        if action not in ('approve', 'reject'):
            return jsonify({'error': '无效操作'}), 400
        conn = db.get_conn()
        try:
            row = conn.execute('SELECT club_name, title FROM offcampus_requests WHERE id=?', (rid,)).fetchone()
            if not row:
                return jsonify({'error': '不存在'}), 404
            if action == 'approve':
                conn.execute('UPDATE offcampus_requests SET status="approved" WHERE id=?', (rid,))
            else:
                reject_reason = data.get('rejectReason', '').strip()
                conn.execute('UPDATE offcampus_requests SET status="rejected", reject_reason=? WHERE id=?', (reject_reason, rid))
            conn.commit()
        finally:
            conn.close()
        club_users = []
        conn2 = db.get_conn()
        try:
            club_users = conn2.execute("SELECT id FROM users WHERE club_name=? AND role='user'", (row['club_name'],)).fetchall()
        finally:
            conn2.close()
        if action == 'approve':
            for cu in club_users:
                send_notification(cu['id'], '✅ 校外活动已通过', f'您提交的「{row["title"]}」已通过审批', 'approve', '/dashboard.html')
        else:
            reason = data.get('rejectReason', '')
            for cu in club_users:
                send_notification(cu['id'], '❌ 校外活动申请被驳回', f'您提交的「{row["title"]}」被驳回，原因：{reason}', 'reject', '/dashboard.html')
        return jsonify({'success': True})


@app.route('/api/debug/base-raw')
def debug_base_raw():
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT uuid, name, headers, data_json FROM base_data').fetchall()
    finally:
        conn.close()
    result = []
    for r in rows:
        headers = json.loads(r['headers']) if r['headers'] else []
        data = json.loads(r['data_json']) if r['data_json'] else []
        sample = data[:3] if data else []
        result.append({'uuid': r['uuid'], 'name': r['name'], 'headers': headers, 'sampleRows': sample, 'rowCount': len(data)})
    return jsonify({'success': True, 'data': result})


@app.route('/api/teachers')
def get_teachers_list():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT u.id, u.username, u.club_name, COALESCE(tp.real_name, up.real_name, u.username) as real_name FROM users u LEFT JOIN teacher_profiles tp ON u.id=tp.user_id LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.role="teacher" ORDER BY real_name').fetchall()
    finally:
        conn.close()
    result = [{'id': r['id'], 'username': r['username'], 'realName': r['real_name'], 'clubName': r['club_name'] or ''} for r in rows]
    return jsonify({'success': True, 'data': result})


@app.route('/api/admin-notifications', methods=['GET', 'POST'])
def handle_admin_notifications():
    if request.method == 'POST':
        user = get_current_user()
        if not user or user['role'] != 'admin':
            return jsonify({'error': '仅管理员可发送通知'}), 403
        data = request.json or {}
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        target_type = data.get('targetType', '').strip()
        target_clubs = data.get('targetClubs', [])
        target_teachers = data.get('targetTeachers', [])
        target_club_str = ','.join(target_clubs) if isinstance(target_clubs, list) else data.get('targetClub', '').strip()
        target_teacher_str = ','.join(target_teachers) if isinstance(target_teachers, list) else ''
        if not title:
            return jsonify({'error': '请输入通知标题'}), 400
        if target_type not in ('all_teachers', 'all_students', 'all_leaders', 'specific_leader', 'specific_teacher'):
            return jsonify({'error': '无效的通知对象类型'}), 400
        if target_type == 'specific_leader' and not target_club_str:
            return jsonify({'error': '请指定目标社团'}), 400
        if target_type == 'specific_teacher' and not target_teacher_str:
            return jsonify({'error': '请指定目标老师'}), 400
        target_store = target_club_str if target_type == 'specific_leader' else (target_teacher_str if target_type == 'specific_teacher' else '')
        conn = db.get_conn()
        try:
            conn.execute('INSERT INTO admin_notifications (title, content, target_type, target_club, sender_id) VALUES (?, ?, ?, ?, ?)',
                        (title, content, target_type, target_store, user['id']))
            conn.commit()
            notif_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        finally:
            conn.close()
        return jsonify({'success': True, 'id': notif_id})
    else:
        user = get_current_user()
        if not user or user['role'] != 'admin':
            return jsonify({'error': '仅管理员可查看'}), 403
        page = int(request.args.get('page', 1))
        page_size = int(request.args.get('pageSize', 20))
        target_type = request.args.get('targetType', '').strip()
        conn = db.get_conn()
        try:
            where = ''
            params = []
            if target_type:
                where = 'WHERE target_type=?'
                params.append(target_type)
            total = conn.execute(f'SELECT COUNT(*) as c FROM admin_notifications {where}', params).fetchone()['c']
            offset = (page - 1) * page_size
            rows = conn.execute(f'SELECT * FROM admin_notifications {where} ORDER BY created_at DESC LIMIT ? OFFSET ?', params + [page_size, offset]).fetchall()
        finally:
            conn.close()
        return jsonify({
            'success': True,
            'total': total,
            'page': page,
            'pageSize': page_size,
            'totalPages': math.ceil(total / page_size) if page_size > 0 else 0,
            'data': [{'id': r['id'], 'title': r['title'], 'content': r['content'], 'targetType': r['target_type'], 'targetClub': r['target_club'], 'senderId': r['sender_id'], 'time': local_time(r['created_at'])} for r in rows]
        })


@app.route('/api/my-admin-notifications')
def get_my_admin_notifications():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        if user['role'] == 'teacher':
            username = user['username'] or ''
            rows = conn.execute('SELECT * FROM admin_notifications WHERE target_type=? OR (target_type=? AND (","||target_club||"," LIKE "%,"||?||",%")) ORDER BY created_at DESC',
                              ('all_teachers', 'specific_teacher', username)).fetchall()
        elif user['role'] == 'student':
            rows = conn.execute('SELECT * FROM admin_notifications WHERE target_type=? ORDER BY created_at DESC', ('all_students',)).fetchall()
        elif user['role'] == 'user':
            club = user['club_name'] or ''
            rows = conn.execute('SELECT * FROM admin_notifications WHERE target_type=? OR (target_type=? AND (target_club=? OR ","||target_club||"," LIKE "%,"||?||",%")) ORDER BY created_at DESC',
                              ('all_leaders', 'specific_leader', club, club)).fetchall()
        else:
            rows = conn.execute('SELECT * FROM admin_notifications ORDER BY created_at DESC').fetchall()
        read_ids = set()
        cleared_ids = set()
        if rows:
            notif_ids = [r['id'] for r in rows]
            placeholders = ','.join(['?'] * len(notif_ids))
            read_rows = conn.execute(f'SELECT notification_id, cleared FROM admin_notification_reads WHERE user_id=? AND notification_id IN ({placeholders})', [user['id']] + notif_ids).fetchall()
            for rr in read_rows:
                if rr['cleared']:
                    cleared_ids.add(rr['notification_id'])
                else:
                    read_ids.add(rr['notification_id'])
        rows = [r for r in rows if r['id'] not in cleared_ids]
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'data': [{'id': r['id'], 'title': r['title'], 'content': r['content'], 'targetType': r['target_type'], 'targetClub': r['target_club'], 'time': local_time(r['created_at']), 'isRead': r['id'] in read_ids} for r in rows]
    })


@app.route('/api/admin-notifications/<int:nid>/read', methods=['POST'])
def mark_admin_notification_read(nid):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        conn.execute('INSERT OR IGNORE INTO admin_notification_reads (notification_id, user_id) VALUES (?, ?)', (nid, user['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/admin-notifications/<int:nid>', methods=['DELETE'])
def delete_admin_notification(nid):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可删除通知'}), 403
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM admin_notification_reads WHERE notification_id=?', (nid,))
        conn.execute('DELETE FROM admin_notifications WHERE id=?', (nid,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True})


@app.route('/api/club-structures')
def get_club_structures():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    club_name = request.args.get('clubName', '').strip()
    conn = db.get_conn()
    try:
        if club_name:
            clubs = conn.execute('SELECT club_name, description, star_rating, president, category, guiding_unit, emblem_url FROM club_profiles WHERE club_name=?', (club_name,)).fetchall()
        else:
            clubs = conn.execute('SELECT club_name, description, star_rating, president, category, guiding_unit, emblem_url FROM club_profiles ORDER BY club_name').fetchall()
        result = []
        for c in clubs:
            cn = c['club_name']
            depts = conn.execute('SELECT id, dept_name, description, parent_id FROM club_departments WHERE club_name=? ORDER BY id', (cn,)).fetchall()
            members = conn.execute('SELECT real_name, student_id_num, class_name, college, department, specialty, joined_at FROM club_members WHERE club_name=? ORDER BY joined_at', (cn,)).fetchall()
            teachers = conn.execute('SELECT teacher_name, introduction FROM club_teachers WHERE club_name=?', (cn,)).fetchall()
            leaders = conn.execute("SELECT u.username, up.real_name FROM users u LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.club_name=? AND u.role='user'", (cn,)).fetchall()
            dept_list = []
            for d in depts:
                dept_members = [m for m in members if m['department'] == d['dept_name']]
                dept_list.append({
                    'id': d['id'], 'name': d['dept_name'], 'description': d['description'],
                    'parentId': d['parent_id'] if 'parent_id' in d.keys() else 0,
                    'memberCount': len(dept_members),
                    'members': [{'realName': m['real_name'], 'studentIdNum': m['student_id_num'], 'className': m['class_name'], 'college': m['college'] if 'college' in m.keys() else '', 'specialty': m['specialty'], 'joinedAt': m['joined_at']} for m in dept_members]
                })
            unassigned = [m for m in members if not m['department'] or not any(d['dept_name'] == m['department'] for d in depts)]
            result.append({
                'clubName': cn, 'description': c['description'], 'starRating': c['star_rating'],
                'president': c['president'] if 'president' in c.keys() else '',
                'category': c['category'] if 'category' in c.keys() else '',
                'guidingUnit': c['guiding_unit'] if 'guiding_unit' in c.keys() else '',
                'emblemUrl': c['emblem_url'] if 'emblem_url' in c.keys() else '',
                'leaders': [{'username': l['username'], 'realName': l['real_name']} for l in leaders],
                'teachers': [{'name': t['teacher_name'], 'introduction': t['introduction']} for t in teachers],
                'departments': dept_list,
                'unassignedMembers': [{'realName': m['real_name'], 'studentIdNum': m['student_id_num'], 'className': m['class_name'], 'college': m['college'] if 'college' in m.keys() else '', 'specialty': m['specialty'], 'joinedAt': m['joined_at']} for m in unassigned],
                'totalMembers': len(members), 'totalDepts': len(depts)
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/export-club-structures')
def export_club_structures():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '无权限'}), 403
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return jsonify({'error': '导出功能需要安装 python-docx 库'}), 500
    club_name = request.args.get('clubName', '').strip()
    conn = db.get_conn()
    try:
        if club_name:
            clubs = conn.execute('SELECT club_name, description, star_rating, president, category, guiding_unit FROM club_profiles WHERE club_name=?', (club_name,)).fetchall()
        else:
            clubs = conn.execute('SELECT club_name, description, star_rating, president, category, guiding_unit FROM club_profiles ORDER BY club_name').fetchall()
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        for ci, club in enumerate(clubs):
            cn = club['club_name']
            if ci > 0:
                doc.add_page_break()
            doc.add_heading(cn, level=1)
            depts = conn.execute('SELECT id, dept_name, description, parent_id FROM club_departments WHERE club_name=? ORDER BY id', (cn,)).fetchall()
            doc.add_heading('部门结构', level=2)
            if depts:
                dept_map = {d['id']: d['dept_name'] for d in depts}
                for d in depts:
                    dept_name = d['dept_name'] or '未命名部门'
                    dept_desc = d['description'] or '暂无描述'
                    parent_name = dept_map.get(d['parent_id'] if 'parent_id' in d.keys() else 0, '')
                    p = doc.add_paragraph()
                    run = p.add_run('■ ' + dept_name)
                    run.bold = True
                    run.font.size = Pt(11)
                    if parent_name:
                        p.add_run('  上级部门：' + parent_name).font.size = Pt(9)
                    desc_p = doc.add_paragraph('  部门介绍：' + dept_desc)
                    desc_p.paragraph_format.space_after = Pt(2)
                    for run in desc_p.runs:
                        run.font.size = Pt(10)
            else:
                doc.add_paragraph('暂无部门信息')
    finally:
        conn.close()
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f'{club_name}_社团结构.docx' if club_name else '全部社团结构.docx'
    return send_file(buf, as_attachment=True, download_name=fname, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/workload/submit', methods=['POST'])
def workload_submit():
    user = get_current_user()
    if not user or user['role'] != 'student':
        return jsonify({'error': '仅学生可提交工作量'}), 403
    data = request.get_json(force=True)
    item_name = (data.get('item_name') or '').strip()
    score = data.get('score')
    if not item_name:
        return jsonify({'error': '请填写活动/事项名称'}), 400
    try:
        score = float(score)
        if score <= 0:
            raise ValueError
    except (TypeError, ValueError):
        return jsonify({'error': '工作量分数必须为正数'}), 400
    club_name = (data.get('club_name') or '').strip()
    if not club_name:
        club_name = user.get('club_name') or ''
    if not club_name:
        return jsonify({'error': '您未加入任何社团'}), 400
    conn = db.get_conn()
    try:
        member_row = conn.execute('SELECT 1 FROM club_members WHERE user_id=? AND club_name=?', (user['id'], club_name)).fetchone()
        if not member_row and club_name != (user.get('club_name') or ''):
            return jsonify({'error': '您不是该社团成员'}), 403
        profile = conn.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
        student_name = profile['real_name'] if profile and profile['real_name'] else user['username']
        conn.execute('INSERT INTO workload_submissions (student_user_id, student_name, club_name, item_name, score, status) VALUES (?,?,?,?,?,?)',
                     (user['id'], student_name, club_name, item_name, score, 'pending'))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '提交成功，等待审核'})


@app.route('/api/workload/my-submissions', methods=['GET'])
def workload_my_submissions():
    user = get_current_user()
    if not user or user['role'] != 'student':
        return jsonify({'error': '请先登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, item_name, score, status, review_note, created_at, reviewed_at FROM workload_submissions WHERE student_user_id=? ORDER BY created_at DESC', (user['id'],)).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [dict(r) for r in rows]})


@app.route('/api/workload/pending', methods=['GET'])
def workload_pending():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '仅社团负责人可查看'}), 403
    req_club = request.args.get('club', '').strip()
    if user['role'] == 'user':
        club_name = req_club if req_club else (user.get('club_name') or '')
    else:
        club_name = req_club
    if not club_name:
        return jsonify({'success': True, 'data': []})
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        date_where = ''
        date_params = [club_name, 'pending']
        if start_date:
            date_where += ' AND date(ws.created_at)>=?'
            date_params.append(start_date)
        if end_date:
            date_where += ' AND date(ws.created_at)<=?'
            date_params.append(end_date)
        rows = conn.execute(f'SELECT ws.id, ws.student_user_id, ws.student_name, ws.item_name, ws.score, ws.status, ws.created_at FROM workload_submissions ws WHERE ws.club_name=? AND ws.status=?{date_where} ORDER BY ws.created_at DESC', date_params).fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [dict(r) for r in rows]})


@app.route('/api/workload/review/<int:sub_id>', methods=['POST'])
def workload_review(sub_id):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '仅社团负责人可审核'}), 403
    data = request.get_json(force=True)
    action = data.get('action')
    review_note = (data.get('review_note') or '').strip()
    if action not in ('approve', 'reject'):
        return jsonify({'error': '无效操作'}), 400
    conn = db.get_conn()
    try:
        sub = conn.execute('SELECT id, club_name, status FROM workload_submissions WHERE id=?', (sub_id,)).fetchone()
        if not sub:
            return jsonify({'error': '记录不存在'}), 404
        if sub['status'] != 'pending':
            return jsonify({'error': '该记录已审核'}), 400
        if user['role'] == 'user' and sub['club_name'] != (user.get('club_name') or ''):
            return jsonify({'error': '无权审核'}), 403
        new_status = 'approved' if action == 'approve' else 'rejected'
        profile = conn.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
        reviewer_name = profile['real_name'] if profile and profile['real_name'] else user['username']
        conn.execute('UPDATE workload_submissions SET status=?, reviewer_id=?, reviewer_name=?, review_note=?, reviewed_at=datetime("now","localtime") WHERE id=?',
                     (new_status, user['id'], reviewer_name, review_note, sub_id))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '审核完成'})


@app.route('/api/workload/my-clubs', methods=['GET'])
def workload_my_clubs():
    user = get_current_user()
    if not user or user['role'] != 'student':
        return jsonify({'error': '仅学生可查看'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT club_name FROM club_members WHERE user_id=?', (user['id'],)).fetchall()
        clubs = [r['club_name'] for r in rows if r['club_name']]
        if not clubs and user.get('club_name'):
            clubs = [user['club_name']]
    finally:
        conn.close()
    return jsonify({'success': True, 'clubs': clubs})


@app.route('/api/workload/club-stats', methods=['GET'])
def workload_club_stats():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    if user['role'] == 'student':
        club_name = request.args.get('club', '') or user.get('club_name') or ''
    elif user['role'] == 'user':
        club_name = user.get('club_name') or ''
    elif user['role'] == 'admin':
        club_name = request.args.get('club', '')
    else:
        return jsonify({'error': '无权限'}), 403
    if not club_name:
        conn_pre = db.get_conn()
        try:
            mc = conn_pre.execute('SELECT club_name FROM club_members WHERE user_id=? LIMIT 1', (user['id'],)).fetchone()
            if mc:
                club_name = mc['club_name']
        finally:
            conn_pre.close()
    if not club_name:
        return jsonify({'success': True, 'data': []})
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    conn = db.get_conn()
    try:
        members = conn.execute('SELECT cm.user_id, cm.real_name, cm.student_id_num, cm.class_name FROM club_members cm LEFT JOIN users u ON cm.user_id=u.id WHERE cm.club_name=? AND (cm.user_id=0 OR u.id IS NOT NULL) ORDER BY cm.real_name', (club_name,)).fetchall()
        act_date_where = ''
        act_date_params = [club_name]
        wl_date_where = ''
        wl_date_params = [club_name, 'approved']
        if start_date:
            act_date_where += ' AND date(cs.created_at)>=?'
            act_date_params.append(start_date)
            wl_date_where += ' AND date(ws.created_at)>=?'
            wl_date_params.append(start_date)
        if end_date:
            act_date_where += ' AND date(cs.created_at)<=?'
            act_date_params.append(end_date)
            wl_date_where += ' AND date(ws.created_at)<=?'
            wl_date_params.append(end_date)
        activity_counts = {}
        act_rows = conn.execute(f'SELECT cr.student_id, COUNT(DISTINCT cr.session_id) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=? AND cs.is_completed=1 AND cr.student_id IS NOT NULL AND cr.student_id!="" AND cr.student_id!=0{act_date_where} GROUP BY cr.student_id', act_date_params).fetchall()
        for r in act_rows:
            activity_counts[r['student_id']] = r['cnt']
        act_name_rows = conn.execute(f'SELECT cr.student_name, COUNT(DISTINCT cr.session_id) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=? AND cs.is_completed=1{act_date_where} GROUP BY cr.student_name', act_date_params).fetchall()
        act_name_map = {}
        for r in act_name_rows:
            act_name_map[r['student_name']] = r['cnt']
        workload_sums = {}
        wl_rows = conn.execute(f'SELECT student_user_id, SUM(score) as total_score FROM workload_submissions ws WHERE ws.club_name=? AND ws.status=?{wl_date_where} GROUP BY student_user_id', wl_date_params).fetchall()
        for r in wl_rows:
            workload_sums[r['student_user_id']] = r['total_score']
        result = []
        for m in members:
            uid = m['user_id']
            name = m['real_name'] or ''
            act_count = activity_counts.get(uid, 0) or activity_counts.get(str(uid), 0) or act_name_map.get(name, 0)
            other_score = workload_sums.get(uid, 0)
            result.append({
                'user_id': uid,
                'real_name': name,
                'student_id_num': m['student_id_num'],
                'class_name': m['class_name'],
                'activity_count': act_count,
                'activity_score': act_count,
                'other_score': round(other_score, 1),
                'total_score': round(act_count + other_score, 1)
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result, 'club_name': club_name})


@app.route('/api/workload/student-detail', methods=['GET'])
def workload_student_detail():
    user = get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    req_club = request.args.get('club', '')
    student_user_id = request.args.get('user_id', '')
    if not student_user_id:
        return jsonify({'error': '缺少参数'}), 400
    student_user_id = int(student_user_id)
    conn = db.get_conn()
    try:
        student = conn.execute('SELECT u.id, u.username, u.club_name, up.real_name, up.student_id as sid, up.class_name FROM users u LEFT JOIN user_profiles up ON u.id=up.user_id WHERE u.id=?', (student_user_id,)).fetchone()
        if not student:
            return jsonify({'error': '学生不存在'}), 404
        club_name = ''
        filter_by_club = True
        if user['role'] == 'user':
            club_name = user.get('club_name') or ''
            member = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=?', (club_name, student_user_id)).fetchone()
            if not member:
                return jsonify({'error': '无权查看'}), 403
        elif user['role'] == 'student':
            club_name = req_club or user.get('club_name') or ''
            if not club_name:
                mc = conn.execute('SELECT club_name FROM club_members WHERE user_id=? LIMIT 1', (user['id'],)).fetchone()
                if mc:
                    club_name = mc['club_name']
            if not club_name:
                club_name = student['club_name'] or ''
            if not club_name:
                mc2 = conn.execute('SELECT club_name FROM club_members WHERE user_id=? LIMIT 1', (student_user_id,)).fetchone()
                if mc2:
                    club_name = mc2['club_name']
            if user['id'] == student_user_id:
                filter_by_club = False
            else:
                member = conn.execute('SELECT id FROM club_members WHERE club_name=? AND user_id=?', (club_name, student_user_id)).fetchone()
                if not member:
                    return jsonify({'error': '无权查看'}), 403
        elif user['role'] == 'admin':
            club_name = student['club_name'] or ''
            member_club = conn.execute('SELECT club_name FROM club_members WHERE user_id=? LIMIT 1', (student_user_id,)).fetchone()
            if member_club:
                club_name = member_club['club_name']
            filter_by_club = False
        activity_list = []
        if filter_by_club and club_name:
            act_rows = conn.execute('SELECT cs.activity_name, cs.created_at, cr.checkin_method, cs.club_name FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.student_id=? AND cs.club_name=? AND cs.is_completed=1 ORDER BY cs.created_at DESC', (student_user_id, club_name)).fetchall()
            if not act_rows:
                act_rows = conn.execute('SELECT cs.activity_name, cs.created_at, cr.checkin_method, cs.club_name FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.student_name=? AND cs.club_name=? AND cs.is_completed=1 ORDER BY cs.created_at DESC', (student['real_name'] or student['username'], club_name)).fetchall()
        else:
            act_rows = conn.execute('SELECT cs.activity_name, cs.created_at, cr.checkin_method, cs.club_name FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.student_id=? AND cs.is_completed=1 ORDER BY cs.created_at DESC', (student_user_id,)).fetchall()
            if not act_rows:
                act_rows = conn.execute('SELECT cs.activity_name, cs.created_at, cr.checkin_method, cs.club_name FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cr.student_name=? AND cs.is_completed=1 ORDER BY cs.created_at DESC', (student['real_name'] or student['username'],)).fetchall()
        for r in act_rows:
            activity_list.append({'name': r['activity_name'] or '未命名活动', 'date': r['created_at'], 'type': 'activity', 'club_name': r['club_name']})
        other_list = []
        if filter_by_club and club_name:
            wl_rows = conn.execute('SELECT item_name, score, status, created_at, review_note, club_name FROM workload_submissions WHERE student_user_id=? AND club_name=? ORDER BY created_at DESC', (student_user_id, club_name)).fetchall()
        else:
            wl_rows = conn.execute('SELECT item_name, score, status, created_at, review_note, club_name FROM workload_submissions WHERE student_user_id=? ORDER BY created_at DESC', (student_user_id,)).fetchall()
        for r in wl_rows:
            other_list.append({'name': r['item_name'], 'score': r['score'], 'status': r['status'], 'date': r['created_at'], 'review_note': r['review_note'], 'type': 'other', 'club_name': r['club_name']})
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'student': {'user_id': student['id'], 'real_name': student['real_name'] or student['username'], 'sid': student['sid'] or '', 'class_name': student['class_name'] or '', 'club_name': club_name},
        'activity_list': [dict(r) for r in activity_list],
        'other_list': [dict(r) for r in other_list],
        'activity_count': len(activity_list),
        'activity_score': len(activity_list),
        'other_score': round(sum(r['score'] for r in other_list if r['status'] == 'approved'), 1)
    })


@app.route('/workload.html')
def serve_workload():
    return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'workload.html'))


@app.route('/api/scoring-rules', methods=['GET', 'POST'])
def handle_scoring_rules():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可操作'}), 403
    if request.method == 'GET':
        conn = db.get_conn()
        try:
            rows = conn.execute('SELECT star_level, collective_limit, individual_limit, updated_at, date_start, date_end FROM scoring_rules ORDER BY star_level DESC').fetchall()
        finally:
            conn.close()
        return jsonify({'success': True, 'data': [dict(r) for r in rows]})
    else:
        data = request.get_json(force=True)
        rules = data.get('rules', [])
        date_start = (data.get('date_start') or '').strip()
        date_end = (data.get('date_end') or '').strip()
        if date_start and date_end and date_end < date_start:
            return jsonify({'error': '结束日期不能早于开始日期'}), 400
        conn = db.get_conn()
        try:
            for r in rules:
                star = r.get('star_level')
                coll = r.get('collective_limit', 0)
                indiv = r.get('individual_limit', 0)
                if star is None:
                    continue
                try:
                    coll = float(coll) if coll else 0
                    indiv = float(indiv) if indiv else 0
                except (TypeError, ValueError):
                    coll = 0
                    indiv = 0
                conn.execute('UPDATE scoring_rules SET collective_limit=?, individual_limit=?, updated_at=datetime("now","localtime") WHERE star_level=?', (coll, indiv, star))
            conn.execute('UPDATE scoring_rules SET date_start=?, date_end=?', (date_start, date_end))
            if date_start or date_end:
                conn.execute('UPDATE scoring_submissions SET date_start=?, date_end=? WHERE status=?', (date_start, date_end, 'draft'))
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True, 'message': '赋分规则已保存'})


@app.route('/api/scoring-club-overrides', methods=['GET', 'POST'])
def handle_scoring_club_overrides():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可操作'}), 403
    if request.method == 'GET':
        conn = db.get_conn()
        try:
            rows = conn.execute('SELECT sco.club_name, sco.collective_limit, sco.individual_limit, sco.updated_at, cp.star_rating FROM scoring_club_overrides sco LEFT JOIN club_profiles cp ON sco.club_name=cp.club_name ORDER BY cp.star_rating DESC, sco.club_name').fetchall()
        finally:
            conn.close()
        return jsonify({'success': True, 'data': [dict(r) for r in rows]})
    else:
        data = request.get_json(force=True)
        overrides = data.get('overrides', [])
        conn = db.get_conn()
        try:
            for o in overrides:
                club_name = (o.get('club_name') or '').strip()
                if not club_name:
                    continue
                coll = o.get('collective_limit')
                indiv = o.get('individual_limit')
                try:
                    coll = float(coll) if coll is not None and coll != '' else None
                except (TypeError, ValueError):
                    coll = None
                try:
                    indiv = float(indiv) if indiv is not None and indiv != '' else None
                except (TypeError, ValueError):
                    indiv = None
                existing = conn.execute('SELECT id FROM scoring_club_overrides WHERE club_name=?', (club_name,)).fetchone()
                if existing:
                    conn.execute('UPDATE scoring_club_overrides SET collective_limit=?, individual_limit=?, updated_at=datetime("now","localtime") WHERE club_name=?', (coll, indiv, club_name))
                else:
                    conn.execute('INSERT INTO scoring_club_overrides (club_name, collective_limit, individual_limit) VALUES (?,?,?)', (club_name, coll, indiv))
            conn.commit()
        finally:
            conn.close()
        return jsonify({'success': True, 'message': '社团单独赋分已保存'})


@app.route('/api/scoring-club-overrides/<club_name>', methods=['DELETE'])
def delete_scoring_club_override(club_name):
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可操作'}), 403
    conn = db.get_conn()
    try:
        conn.execute('DELETE FROM scoring_club_overrides WHERE club_name=?', (club_name,))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '已恢复为星级默认值'})


@app.route('/api/scoring/calculate', methods=['GET'])
def scoring_calculate():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '仅社团负责人、管理员或指导老师可查看'}), 403
    if user['role'] == 'user':
        club_name = user.get('club_name') or ''
    elif user['role'] == 'teacher':
        club_name = request.args.get('club', '') or user.get('club_name') or ''
    else:
        club_name = request.args.get('club', '')
    if not club_name:
        return jsonify({'success': True, 'data': [], 'club_name': '', 'star_level': 0, 'collective_limit': 0, 'individual_limit': 0})
    conn = db.get_conn()
    try:
        profile = conn.execute('SELECT star_rating FROM club_profiles WHERE club_name=?', (club_name,)).fetchone()
        star_level = profile['star_rating'] if profile else 0
        rule = conn.execute('SELECT collective_limit, individual_limit FROM scoring_rules WHERE star_level=?', (star_level,)).fetchone()
        collective_limit = rule['collective_limit'] if rule else 0
        individual_limit = rule['individual_limit'] if rule else 0
        date_row = conn.execute('SELECT date_start, date_end FROM scoring_rules LIMIT 1').fetchone()
        date_start = date_row['date_start'] if date_row else ''
        date_end = date_row['date_end'] if date_row else ''
        override = conn.execute('SELECT collective_limit, individual_limit FROM scoring_club_overrides WHERE club_name=?', (club_name,)).fetchone()
        has_override = False
        if override:
            if override['collective_limit'] is not None:
                collective_limit = override['collective_limit']
                has_override = True
            if override['individual_limit'] is not None:
                individual_limit = override['individual_limit']
                has_override = True
        members = conn.execute('SELECT cm.user_id, cm.real_name, cm.student_id_num, cm.class_name, cm.college FROM club_members cm LEFT JOIN users u ON cm.user_id=u.id WHERE cm.club_name=? AND (cm.user_id=0 OR u.id IS NOT NULL) ORDER BY cm.real_name', (club_name,)).fetchall()
        if date_start and date_end and date_end < date_start:
            date_start = ''
            date_end = ''
        activity_counts = {}
        date_params_act = [club_name]
        date_cond_act = ''
        if date_start:
            date_cond_act += ' AND date(cs.created_at)>=?'
            date_params_act.append(date_start)
        if date_end:
            date_cond_act += ' AND date(cs.created_at)<=?'
            date_params_act.append(date_end)
        act_rows = conn.execute('SELECT cr.student_id, COUNT(DISTINCT cr.session_id) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=? AND cs.is_completed=1 AND cr.student_id IS NOT NULL AND cr.student_id!="" AND cr.student_id!=0' + date_cond_act + ' GROUP BY cr.student_id', date_params_act).fetchall()
        for r in act_rows:
            activity_counts[r['student_id']] = r['cnt']
        date_params_name = [club_name]
        date_cond_name = ''
        if date_start:
            date_cond_name += ' AND date(cs.created_at)>=?'
            date_params_name.append(date_start)
        if date_end:
            date_cond_name += ' AND date(cs.created_at)<=?'
            date_params_name.append(date_end)
        act_name_rows = conn.execute('SELECT cr.student_name, COUNT(DISTINCT cr.session_id) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=? AND cs.is_completed=1' + date_cond_name + ' GROUP BY cr.student_name', date_params_name).fetchall()
        act_name_map = {}
        for r in act_name_rows:
            act_name_map[r['student_name']] = r['cnt']
        workload_sums = {}
        wl_date_params = [club_name, 'approved']
        wl_date_cond = ''
        if date_start:
            wl_date_cond += ' AND date(created_at)>=?'
            wl_date_params.append(date_start)
        if date_end:
            wl_date_cond += ' AND date(created_at)<=?'
            wl_date_params.append(date_end)
        wl_rows = conn.execute('SELECT student_user_id, SUM(score) as total_score FROM workload_submissions WHERE club_name=? AND status=?' + wl_date_cond + ' GROUP BY student_user_id', wl_date_params).fetchall()
        for r in wl_rows:
            workload_sums[r['student_user_id']] = r['total_score']
        member_workloads = []
        for m in members:
            uid = m['user_id']
            name = m['real_name'] or ''
            act_count = activity_counts.get(uid, 0) or activity_counts.get(str(uid), 0) or act_name_map.get(name, 0)
            other_score = workload_sums.get(uid, 0)
            total = round(act_count + other_score, 1)
            member_workloads.append({
                'user_id': uid,
                'real_name': name,
                'student_id_num': m['student_id_num'],
                'class_name': m['class_name'],
                'college': m['college'] if 'college' in m.keys() else '',
                'activity_count': act_count,
                'other_score': round(other_score, 1),
                'total_workload': total
            })
        total_workload_sum = sum(mw['total_workload'] for mw in member_workloads)
        results = []
        for mw in member_workloads:
            if total_workload_sum > 0 and collective_limit > 0:
                preliminary = (mw['total_workload'] / total_workload_sum) * collective_limit
            else:
                preliminary = 0
            if individual_limit > 0 and preliminary > individual_limit:
                final_score = round(individual_limit, 1)
            else:
                final_score = round(preliminary, 1)
            results.append({
                'user_id': mw['user_id'],
                'real_name': mw['real_name'],
                'student_id_num': mw['student_id_num'],
                'class_name': mw['class_name'],
                'college': mw['college'],
                'activity_count': mw['activity_count'],
                'other_score': mw['other_score'],
                'total_workload': mw['total_workload'],
                'preliminary_score': round(preliminary, 1),
                'final_score': final_score,
                'capped': preliminary > individual_limit if individual_limit > 0 else False
            })
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'data': results,
        'club_name': club_name,
        'star_level': star_level,
        'collective_limit': collective_limit,
        'individual_limit': individual_limit,
        'has_override': has_override,
        'total_workload_sum': round(total_workload_sum, 1),
        'member_count': len(results),
        'date_start': date_start,
        'date_end': date_end
    })


@app.route('/api/scoring/calculate-export', methods=['GET'])
def scoring_calculate_export():
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '仅社团负责人、管理员或指导老师可导出'}), 403
    if user['role'] == 'user':
        club_name = user.get('club_name') or ''
    elif user['role'] == 'teacher':
        club_name = request.args.get('club', '') or user.get('club_name') or ''
    else:
        club_name = request.args.get('club', '')
    if not club_name:
        return jsonify({'error': '请指定社团'}), 400
    conn = db.get_conn()
    try:
        profile = conn.execute('SELECT star_rating FROM club_profiles WHERE club_name=?', (club_name,)).fetchone()
        star_level = profile['star_rating'] if profile else 0
        rule = conn.execute('SELECT collective_limit, individual_limit FROM scoring_rules WHERE star_level=?', (star_level,)).fetchone()
        collective_limit = rule['collective_limit'] if rule else 0
        individual_limit = rule['individual_limit'] if rule else 0
        date_row = conn.execute('SELECT date_start, date_end FROM scoring_rules LIMIT 1').fetchone()
        date_start = date_row['date_start'] if date_row else ''
        date_end = date_row['date_end'] if date_row else ''
        override = conn.execute('SELECT collective_limit, individual_limit FROM scoring_club_overrides WHERE club_name=?', (club_name,)).fetchone()
        has_override = False
        if override:
            if override['collective_limit'] is not None:
                collective_limit = override['collective_limit']
                has_override = True
            if override['individual_limit'] is not None:
                individual_limit = override['individual_limit']
                has_override = True
        members = conn.execute('SELECT cm.user_id, cm.real_name, cm.student_id_num, cm.class_name, cm.college FROM club_members cm LEFT JOIN users u ON cm.user_id=u.id WHERE cm.club_name=? AND (cm.user_id=0 OR u.id IS NOT NULL) ORDER BY cm.real_name', (club_name,)).fetchall()
        if date_start and date_end and date_end < date_start:
            date_start = ''
            date_end = ''
        activity_counts = {}
        date_params_act = [club_name]
        date_cond_act = ''
        if date_start:
            date_cond_act += ' AND date(cs.created_at)>=?'
            date_params_act.append(date_start)
        if date_end:
            date_cond_act += ' AND date(cs.created_at)<=?'
            date_params_act.append(date_end)
        act_rows = conn.execute('SELECT cr.student_name, COUNT(DISTINCT cr.session_id) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=? AND cs.is_completed=1' + date_cond_act + ' GROUP BY cr.student_name', date_params_act).fetchall()
        for r in act_rows:
            activity_counts[r['student_name']] = r['cnt']
        workload_sums = {}
        wl_date_params = [club_name, 'approved']
        wl_date_cond = ''
        if date_start:
            wl_date_cond += ' AND date(created_at)>=?'
            wl_date_params.append(date_start)
        if date_end:
            wl_date_cond += ' AND date(created_at)<=?'
            wl_date_params.append(date_end)
        wl_rows = conn.execute('SELECT student_user_id, SUM(score) as total_score FROM workload_submissions WHERE club_name=? AND status=?' + wl_date_cond + ' GROUP BY student_user_id', wl_date_params).fetchall()
        for r in wl_rows:
            workload_sums[r['student_user_id']] = r['total_score']
        member_workloads = []
        for m in members:
            uid = m['user_id']
            name = m['real_name'] or ''
            act_count = activity_counts.get(name, 0)
            other_score = workload_sums.get(uid, 0)
            total = round(act_count + other_score, 1)
            member_workloads.append({
                'user_id': uid, 'real_name': name,
                'student_id_num': m['student_id_num'], 'class_name': m['class_name'],
                'college': m['college'] if 'college' in m.keys() else '',
                'activity_count': act_count, 'other_score': round(other_score, 1), 'total_workload': total
            })
        total_workload_sum = sum(mw['total_workload'] for mw in member_workloads)
        results = []
        for mw in member_workloads:
            if total_workload_sum > 0 and collective_limit > 0:
                preliminary = (mw['total_workload'] / total_workload_sum) * collective_limit
            else:
                preliminary = 0
            if individual_limit > 0 and preliminary > individual_limit:
                final_score = round(individual_limit, 1)
            else:
                final_score = round(preliminary, 1)
            results.append({
                'user_id': mw['user_id'], 'real_name': mw['real_name'],
                'student_id_num': mw['student_id_num'], 'class_name': mw['class_name'],
                'college': mw['college'], 'activity_count': mw['activity_count'],
                'other_score': mw['other_score'], 'total_workload': mw['total_workload'],
                'preliminary_score': round(preliminary, 1), 'final_score': final_score,
                'capped': preliminary > individual_limit if individual_limit > 0 else False
            })
    finally:
        conn.close()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '赋分计算结果'
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color='667EEA', end_color='667EEA', fill_type='solid')
    header_font_white = Font(bold=True, size=11, color='FFFFFF')
    headers = ['序号', '姓名', '学号', '班级', '学院', '活动次数', '其他工作量', '总工作量', '集体限额内得分', '最终赋分', '是否封顶']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
    info_row = 1
    ws.insert_rows(info_row)
    date_range_info = f' | 统计范围：{date_start}~{date_end}' if date_start or date_end else ''
    info_cell = ws.cell(row=info_row, column=1, value=f'社团：{club_name} | 星级：{star_level}星 | 集体限额：{collective_limit}分 | 个人封顶：{individual_limit}分 | 总工作量：{total_workload_sum} | 成员数：{len(results)}{" | 有单独赋分" if has_override else ""}{date_range_info}')
    info_cell.font = Font(bold=True, size=12, color='1A1A2E')
    ws.merge_cells(start_row=info_row, start_column=1, end_row=info_row, end_column=len(headers))
    for row_idx, r in enumerate(results, info_row + 1):
        ws.cell(row=row_idx, column=1, value=row_idx - info_row)
        ws.cell(row=row_idx, column=2, value=r['real_name'])
        ws.cell(row=row_idx, column=3, value=r['student_id_num'] or '')
        ws.cell(row=row_idx, column=4, value=r['class_name'] or '')
        ws.cell(row=row_idx, column=5, value=r['college'] or '')
        ws.cell(row=row_idx, column=6, value=r['activity_count'])
        ws.cell(row=row_idx, column=7, value=r['other_score'])
        ws.cell(row=row_idx, column=8, value=r['total_workload'])
        ws.cell(row=row_idx, column=9, value=r['preliminary_score'])
        final_cell = ws.cell(row=row_idx, column=10, value=r['final_score'])
        final_cell.font = Font(bold=True, color='00B894')
        capped_text = '是（达到封顶值）' if r['capped'] else '否'
        ws.cell(row=row_idx, column=11, value=capped_text)
    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K']:
        ws.column_dimensions[col_letter].width = 16 if col_letter in ('A', 'I', 'J', 'K') else 14
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    safe_name = club_name.replace('/', '_').replace('\\', '_')
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name=f'赋分计算_{safe_name}.xlsx')


@app.route('/api/scoring/my-score', methods=['GET'])
def scoring_my_score():
    user = get_current_user()
    if not user or user['role'] != 'student':
        return jsonify({'error': '仅学生可查看'}), 403
    club_name = (request.args.get('club') or '').strip()
    if not club_name:
        club_name = user.get('club_name') or ''
    if not club_name:
        return jsonify({'success': True, 'data': []})
    conn = db.get_conn()
    try:
        member_row = conn.execute('SELECT 1 FROM club_members WHERE user_id=? AND club_name=?', (user['id'], club_name)).fetchone()
        if not member_row and club_name != (user.get('club_name') or ''):
            return jsonify({'error': '您不是该社团成员'}), 403
        profile = conn.execute('SELECT star_rating FROM club_profiles WHERE club_name=?', (club_name,)).fetchone()
        star_level = profile['star_rating'] if profile else 0
        rule = conn.execute('SELECT collective_limit, individual_limit FROM scoring_rules WHERE star_level=?', (star_level,)).fetchone()
        collective_limit = rule['collective_limit'] if rule else 0
        individual_limit = rule['individual_limit'] if rule else 0
        override = conn.execute('SELECT collective_limit, individual_limit FROM scoring_club_overrides WHERE club_name=?', (club_name,)).fetchone()
        if override:
            if override['collective_limit'] is not None:
                collective_limit = override['collective_limit']
            if override['individual_limit'] is not None:
                individual_limit = override['individual_limit']
        members = conn.execute('SELECT cm.user_id, cm.real_name FROM club_members cm LEFT JOIN users u ON cm.user_id=u.id WHERE cm.club_name=? AND (cm.user_id=0 OR u.id IS NOT NULL)', (club_name,)).fetchall()
        activity_counts = {}
        act_rows = conn.execute('SELECT cr.student_name, COUNT(DISTINCT cr.session_id) as cnt FROM checkin_records cr JOIN checkin_sessions cs ON cr.session_id=cs.id WHERE cs.club_name=? AND cs.is_completed=1 GROUP BY cr.student_name', (club_name,)).fetchall()
        for r in act_rows:
            activity_counts[r['student_name']] = r['cnt']
        workload_sums = {}
        wl_rows = conn.execute('SELECT student_user_id, SUM(score) as total_score FROM workload_submissions WHERE club_name=? AND status=? GROUP BY student_user_id', (club_name, 'approved')).fetchall()
        for r in wl_rows:
            workload_sums[r['student_user_id']] = r['total_score']
        total_workload_sum = 0
        my_data = None
        for m in members:
            uid = m['user_id']
            name = m['real_name'] or ''
            act_count = activity_counts.get(name, 0)
            other_score = workload_sums.get(uid, 0)
            total = round(act_count + other_score, 1)
            total_workload_sum += total
            if uid == user['id']:
                my_data = {'activity_count': act_count, 'other_score': round(other_score, 1), 'total_workload': total}
        if my_data and total_workload_sum > 0 and collective_limit > 0:
            preliminary = (my_data['total_workload'] / total_workload_sum) * collective_limit
        else:
            preliminary = 0
        if individual_limit > 0 and preliminary > individual_limit:
            final_score = round(individual_limit, 1)
        else:
            final_score = round(preliminary, 1)
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'club_name': club_name,
        'star_level': star_level,
        'collective_limit': collective_limit,
        'individual_limit': individual_limit,
        'my_workload': my_data,
        'preliminary_score': round(preliminary, 1),
        'final_score': final_score,
        'capped': preliminary > individual_limit if individual_limit > 0 else False
    })


@app.route('/api/scoring/save', methods=['POST'])
def scoring_save():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '仅社团负责人可操作'}), 403
    club_name = user.get('club_name') or '' if user['role'] == 'user' else request.get_json(force=True).get('club_name', '') or request.args.get('club', '')
    if not club_name:
        return jsonify({'error': '未关联社团'}), 400
    data = request.get_json(force=True)
    items = data.get('items', [])
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id, status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (club_name,)).fetchone()
        if existing and existing['status'] in ('pending_teacher', 'teacher_approved', 'submitted_tuanwei', 'tuanwei_approved'):
            return jsonify({'error': '当前状态不可编辑，赋分已提交审核'}), 400
        date_row = conn.execute('SELECT date_start, date_end FROM scoring_rules LIMIT 1').fetchone()
        cur_ds = date_row['date_start'] if date_row else ''
        cur_de = date_row['date_end'] if date_row else ''
        if existing:
            sub_id = existing['id']
            conn.execute('UPDATE scoring_submissions SET updated_at=datetime("now","localtime"), date_start=?, date_end=? WHERE id=?', (cur_ds, cur_de, sub_id))
            conn.execute('DELETE FROM scoring_submission_items WHERE submission_id=?', (sub_id,))
        else:
            cursor = conn.execute('INSERT INTO scoring_submissions (club_name, status, submitted_by, date_start, date_end) VALUES (?,?,?,?,?)', (club_name, 'draft', user['id'], cur_ds, cur_de))
            sub_id = cursor.lastrowid
        for item in items:
            conn.execute('INSERT INTO scoring_submission_items (submission_id, student_user_id, student_name, student_id_num, college, class_name, total_workload, final_score, activity_count, other_score) VALUES (?,?,?,?,?,?,?,?,?,?)',
                         (sub_id, item.get('user_id') or item.get('student_user_id', 0), item.get('student_name', '') or item.get('real_name', ''), item.get('student_id_num', ''), item.get('college', ''), item.get('class_name', ''), item.get('total_workload', 0), item.get('final_score', 0), item.get('activity_count', 0), item.get('other_score', 0)))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '保存成功', 'submission_id': sub_id})


@app.route('/api/scoring/submit-to-teachers', methods=['POST'])
def scoring_submit_to_teachers():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '仅社团负责人可操作'}), 403
    club_name = user.get('club_name') or '' if user['role'] == 'user' else request.get_json(force=True).get('club_name', '')
    if not club_name:
        return jsonify({'error': '未关联社团'}), 400
    conn = db.get_conn()
    try:
        sub = conn.execute('SELECT id, status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (club_name,)).fetchone()
        if not sub:
            return jsonify({'error': '请先计算并保存赋分'}), 400
        if sub['status'] != 'draft':
            return jsonify({'error': '当前状态不可提交'}), 400
        sub_id = sub['id']
        conn.execute('DELETE FROM scoring_teacher_reviews WHERE submission_id=?', (sub_id,))
        teachers = conn.execute('SELECT user_id FROM teacher_clubs WHERE club_name=?', (club_name,)).fetchall()
        if not teachers:
            teachers = conn.execute('SELECT user_id FROM club_teachers WHERE club_name=? AND user_id IS NOT NULL AND user_id>0', (club_name,)).fetchall()
        teacher_ids = list(set([t['user_id'] for t in teachers if t['user_id']]))
        users_teachers = conn.execute('SELECT id as user_id FROM users WHERE role=? AND club_name=?', ('teacher', club_name)).fetchall()
        for ut in users_teachers:
            if ut['user_id'] not in teacher_ids:
                teacher_ids.append(ut['user_id'])
        if not teacher_ids:
            conn.execute('UPDATE scoring_submissions SET status=?, updated_at=datetime("now","localtime") WHERE id=?', ('teacher_approved', sub_id))
            conn.commit()
            return jsonify({'success': True, 'message': '无指导老师，已自动跳过审核', 'status': 'teacher_approved'})
        for tid in teacher_ids:
            profile = conn.execute('SELECT real_name FROM teacher_profiles WHERE user_id=?', (tid,)).fetchone()
            tname = profile['real_name'] if profile and profile['real_name'] else ''
            try:
                conn.execute('INSERT INTO scoring_teacher_reviews (submission_id, teacher_user_id, teacher_name, status) VALUES (?,?,?,?)', (sub_id, tid, tname, 'pending'))
            except:
                pass
        conn.execute('UPDATE scoring_submissions SET status=?, updated_at=datetime("now","localtime") WHERE id=?', ('pending_teacher', sub_id))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '已提交指导老师审核', 'status': 'pending_teacher'})


@app.route('/api/scoring/teacher-review/<int:review_id>', methods=['POST'])
def scoring_teacher_review(review_id):
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '仅指导老师可审核'}), 403
    data = request.get_json(force=True)
    action = data.get('action')
    review_note = (data.get('review_note') or '').strip()
    if action not in ('approve', 'reject'):
        return jsonify({'error': '无效操作'}), 400
    conn = db.get_conn()
    try:
        review = conn.execute('SELECT id, submission_id, teacher_user_id, status FROM scoring_teacher_reviews WHERE id=?', (review_id,)).fetchone()
        if not review:
            return jsonify({'error': '审核记录不存在'}), 404
        if review['teacher_user_id'] != user['id']:
            return jsonify({'error': '无权审核'}), 403
        if review['status'] != 'pending':
            return jsonify({'error': '已审核'}), 400
        new_status = 'approved' if action == 'approve' else 'rejected'
        conn.execute('UPDATE scoring_teacher_reviews SET status=?, review_note=?, reviewed_at=datetime("now","localtime") WHERE id=?', (new_status, review_note, review_id))
        if new_status == 'approved':
            pending_count = conn.execute('SELECT COUNT(*) as c FROM scoring_teacher_reviews WHERE submission_id=? AND status=?', (review['submission_id'], 'pending')).fetchone()
            if pending_count['c'] == 0:
                conn.execute('UPDATE scoring_submissions SET status=?, updated_at=datetime("now","localtime") WHERE id=?', ('teacher_approved', review['submission_id']))
        else:
            conn.execute('UPDATE scoring_submissions SET status=?, updated_at=datetime("now","localtime") WHERE id=?', ('draft', review['submission_id']))
            remaining = conn.execute('SELECT id FROM scoring_teacher_reviews WHERE submission_id=? AND status=? AND id!=?', (review['submission_id'], 'pending', review_id)).fetchall()
            for r in remaining:
                conn.execute('UPDATE scoring_teacher_reviews SET status=?, review_note=?, reviewed_at=datetime("now","localtime") WHERE id=?', ('cancelled', '其他老师已驳回，自动取消', r['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '审核完成'})


@app.route('/api/scoring/submit-to-tuanwei', methods=['POST'])
def scoring_submit_to_tuanwei():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '仅社团负责人可操作'}), 403
    club_name = user.get('club_name') or '' if user['role'] == 'user' else request.get_json(force=True).get('club_name', '')
    if not club_name:
        return jsonify({'error': '未关联社团'}), 400
    conn = db.get_conn()
    try:
        sub = conn.execute('SELECT id, status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (club_name,)).fetchone()
        if not sub:
            return jsonify({'error': '请先完成赋分'}), 400
        if sub['status'] != 'teacher_approved':
            return jsonify({'error': '需等待指导老师审核通过'}), 400
        conn.execute('UPDATE scoring_submissions SET status=?, updated_at=datetime("now","localtime") WHERE id=?', ('submitted_tuanwei', sub['id']))
        conn.commit()
    finally:
        conn.close()
    return jsonify({'success': True, 'message': '已提交到团委，等待团委审核'})


@app.route('/api/scoring/tuanwei-pending', methods=['GET'])
def scoring_tuanwei_pending():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可查看'}), 403
    conn = db.get_conn()
    try:
        clubs = conn.execute('SELECT club_name FROM club_profiles ORDER BY club_name').fetchall()
        result = []
        for c in clubs:
            sub = conn.execute('SELECT id, status, created_at, updated_at FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (c['club_name'],)).fetchone()
            if not sub or sub['status'] not in ('submitted_tuanwei', 'tuanwei_approved'):
                continue
            items = conn.execute('SELECT student_user_id, student_name, student_id_num, college, class_name, total_workload, final_score, activity_count, other_score FROM scoring_submission_items WHERE submission_id=?', (sub['id'],)).fetchall()
            result.append({
                'club_name': c['club_name'],
                'submission_id': sub['id'],
                'status': sub['status'],
                'created_at': sub['created_at'],
                'updated_at': sub['updated_at'],
                'member_count': len(items),
                'items': [dict(i) for i in items]
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/scoring/tuanwei-approved-list', methods=['GET'])
def scoring_tuanwei_approved_list():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可查看'}), 403
    semester = request.args.get('semester', '')
    conn = db.get_conn()
    try:
        rows = conn.execute("SELECT id, club_name, status, created_at, updated_at, date_start, date_end FROM scoring_submissions WHERE status='tuanwei_approved' ORDER BY updated_at DESC").fetchall()
        rules = conn.execute('SELECT star_level, collective_limit, individual_limit FROM scoring_rules ORDER BY star_level').fetchall()
        rules_map = {r['star_level']: {'collective_limit': r['collective_limit'], 'individual_limit': r['individual_limit']} for r in rules}
        overrides = conn.execute('SELECT club_name, collective_limit, individual_limit FROM scoring_club_overrides').fetchall()
        overrides_map = {o['club_name']: {'collective_limit': o['collective_limit'], 'individual_limit': o['individual_limit']} for o in overrides}
        clubs_star = {}
        for cp in conn.execute('SELECT club_name, star_rating FROM club_profiles').fetchall():
            clubs_star[cp['club_name']] = cp['star_rating']
        fallback_date = conn.execute('SELECT date_start, date_end FROM scoring_rules LIMIT 1').fetchone()
        fb_ds = fallback_date['date_start'] if fallback_date else ''
        fb_de = fallback_date['date_end'] if fallback_date else ''
        semesters = set()
        result = []
        for sub in rows:
            sub_ds = sub['date_start'] or fb_ds
            sub_de = sub['date_end'] or fb_de
            sem = ''
            if sub_ds and sub_de:
                sem = sub_ds + ' ~ ' + sub_de
            elif sub_ds:
                sem = sub_ds + ' ~ '
            elif sub_de:
                sem = ' ~ ' + sub_de
            if sem:
                semesters.add(sem)
            if semester and sem != semester:
                continue
            star = clubs_star.get(sub['club_name'], 0)
            rule = rules_map.get(star, {})
            ov = overrides_map.get(sub['club_name'])
            coll = rule.get('collective_limit', 0) if rule else 0
            indiv = rule.get('individual_limit', 0) if rule else 0
            if ov:
                if ov['collective_limit'] is not None:
                    coll = ov['collective_limit']
                if ov['individual_limit'] is not None:
                    indiv = ov['individual_limit']
            items = conn.execute('SELECT student_name, student_id_num, college, class_name, total_workload, final_score, activity_count, other_score FROM scoring_submission_items WHERE submission_id=?', (sub['id'],)).fetchall()
            result.append({
                'club_name': sub['club_name'],
                'submission_id': sub['id'],
                'star_level': star,
                'collective_limit': coll,
                'individual_limit': indiv,
                'created_at': sub['created_at'],
                'updated_at': sub['updated_at'],
                'semester': sem,
                'date_start': sub_ds,
                'date_end': sub_de,
                'member_count': len(items),
                'items': [dict(i) for i in items]
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result, 'semesters': sorted(semesters, reverse=True)})


@app.route('/api/scoring/tuanwei-approved-preview', methods=['GET'])
def scoring_tuanwei_approved_preview():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可查看'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute("SELECT id, club_name, status, created_at FROM scoring_submissions WHERE status='tuanwei_approved' ORDER BY created_at DESC").fetchall()
        rules = conn.execute('SELECT star_level, collective_limit, individual_limit FROM scoring_rules ORDER BY star_level').fetchall()
        rules_map = {r['star_level']: {'collective_limit': r['collective_limit'], 'individual_limit': r['individual_limit']} for r in rules}
        overrides = conn.execute('SELECT club_name, collective_limit, individual_limit FROM scoring_club_overrides').fetchall()
        overrides_map = {o['club_name']: {'collective_limit': o['collective_limit'], 'individual_limit': o['individual_limit']} for o in overrides}
        clubs_star = {}
        for cp in conn.execute('SELECT club_name, star_rating FROM club_profiles').fetchall():
            clubs_star[cp['club_name']] = cp['star_rating']
        result = []
        seen_clubs = set()
        for sub in rows:
            if sub['club_name'] in seen_clubs:
                continue
            seen_clubs.add(sub['club_name'])
            star = clubs_star.get(sub['club_name'], 0)
            rule = rules_map.get(star, {})
            ov = overrides_map.get(sub['club_name'])
            coll = rule.get('collective_limit', 0) if rule else 0
            indiv = rule.get('individual_limit', 0) if rule else 0
            if ov:
                if ov['collective_limit'] is not None:
                    coll = ov['collective_limit']
                if ov['individual_limit'] is not None:
                    indiv = ov['individual_limit']
            items = conn.execute('SELECT student_name, student_id_num, college, class_name, total_workload, final_score, activity_count, other_score FROM scoring_submission_items WHERE submission_id=?', (sub['id'],)).fetchall()
            result.append({
                'club_name': sub['club_name'],
                'submission_id': sub['id'],
                'star_level': star,
                'collective_limit': coll,
                'individual_limit': indiv,
                'created_at': sub['created_at'],
                'items': [dict(i) for i in items]
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/scoring/tuanwei-review', methods=['POST'])
def scoring_tuanwei_review():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员（团委）可操作'}), 403
    data = request.get_json(force=True)
    club_name = data.get('club_name', '')
    action = data.get('action', '')
    review_note = data.get('review_note', '')
    if not club_name:
        return jsonify({'error': '缺少社团名称'}), 400
    if action not in ('approve', 'reject'):
        return jsonify({'error': '操作类型无效'}), 400
    conn = db.get_conn()
    try:
        sub = conn.execute('SELECT id, status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (club_name,)).fetchone()
        if not sub:
            return jsonify({'error': '该社团无赋分提交记录'}), 400
        if sub['status'] != 'submitted_tuanwei':
            return jsonify({'error': f'当前状态为 {sub["status"]}，无法审核（需为 submitted_tuanwei）'}), 400
        if action == 'approve':
            conn.execute('UPDATE scoring_submissions SET status=?, updated_at=datetime("now","localtime") WHERE id=?', ('tuanwei_approved', sub['id']))
            conn.commit()
        else:
            conn.execute('UPDATE scoring_submissions SET status=?, updated_at=datetime("now","localtime") WHERE id=?', ('teacher_approved', sub['id']))
            conn.commit()
    finally:
        conn.close()
    if action == 'approve':
        return jsonify({'success': True, 'message': f'{club_name} 团委审核已通过'})
    else:
        return jsonify({'success': True, 'message': f'{club_name} 已驳回，退回指导老师确认阶段'})


@app.route('/api/scoring/submission-status', methods=['GET'])
def scoring_submission_status():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '请先登录'}), 401
    if user['role'] == 'user':
        club_name = user.get('club_name') or ''
    elif user['role'] == 'admin':
        club_name = request.args.get('club', '')
    elif user['role'] == 'teacher':
        club_name = request.args.get('club', '')
    else:
        return jsonify({'success': True, 'data': None})
    if not club_name:
        return jsonify({'success': True, 'data': None})
    conn = db.get_conn()
    try:
        sub = conn.execute('SELECT id, status, created_at, updated_at FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (club_name,)).fetchone()
        if not sub:
            return jsonify({'success': True, 'data': None})
        items = conn.execute('SELECT id, student_user_id, student_name, student_id_num, college, class_name, total_workload, final_score, activity_count, other_score FROM scoring_submission_items WHERE submission_id=?', (sub['id'],)).fetchall()
        reviews = conn.execute('SELECT id, teacher_user_id, teacher_name, status, review_note, reviewed_at FROM scoring_teacher_reviews WHERE submission_id=?', (sub['id'],)).fetchall()
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'data': {
            'id': sub['id'],
            'status': sub['status'],
            'created_at': sub['created_at'],
            'updated_at': sub['updated_at'],
            'items': [dict(i) for i in items],
            'reviews': [dict(r) for r in reviews]
        }
    })


@app.route('/api/scoring/teacher-reviews', methods=['GET'])
def scoring_teacher_reviews_list():
    user = get_current_user()
    if not user or user['role'] != 'teacher':
        return jsonify({'error': '仅指导老师可查看'}), 403
    conn = db.get_conn()
    try:
        clubs = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
        if not clubs:
            clubs = conn.execute('SELECT club_name FROM club_teachers WHERE user_id=?', (user['id'],)).fetchall()
        club_names = [c['club_name'] for c in clubs if c['club_name']]
        if user.get('club_name') and user['club_name'] not in club_names:
            club_names.append(user['club_name'])
        result = []
        for cn in club_names:
            sub = conn.execute('SELECT id, status, created_at FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (cn,)).fetchone()
            if not sub:
                continue
            review = conn.execute('SELECT id, status, review_note, reviewed_at FROM scoring_teacher_reviews WHERE submission_id=? AND teacher_user_id=?', (sub['id'], user['id'])).fetchone()
            items = conn.execute('SELECT student_name, student_id_num, college, class_name, total_workload, final_score, activity_count, other_score FROM scoring_submission_items WHERE submission_id=?', (sub['id'],)).fetchall()
            result.append({
                'club_name': cn,
                'submission_id': sub['id'],
                'submission_status': sub['status'],
                'created_at': sub['created_at'],
                'review': dict(review) if review else None,
                'items': [dict(i) for i in items]
            })
    finally:
        conn.close()
    return jsonify({'success': True, 'data': result})


@app.route('/api/scoring/export/<club_name>', methods=['GET'])
def scoring_export(club_name):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '无权限'}), 403
    if user['role'] == 'user' and user.get('club_name') != club_name:
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        sub = conn.execute('SELECT id, status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (club_name,)).fetchone()
        if not sub:
            return jsonify({'error': '暂无赋分数据'}), 404
        items = conn.execute('SELECT student_id_num, student_name, college, class_name, total_workload, final_score FROM scoring_submission_items WHERE submission_id=? ORDER BY student_id_num', (sub['id'],)).fetchall()
    finally:
        conn.close()
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '赋分结果'
    headers = ['学号', '姓名', '学院', '班级', '总工作量', '个人赋分']
    ws.append(headers)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = openpyxl.styles.Font(bold=True)
    for item in items:
        ws.append([item['student_id_num'], item['student_name'], item['college'], item['class_name'], item['total_workload'], item['final_score']])
    for col in ws.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[col[0].column_letter].width = max_length + 4
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f'{club_name}_赋分结果.xlsx'
    return send_file(buf, as_attachment=True, download_name=fname, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/final-credits/status', methods=['GET'])
def final_credits_status():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可查看'}), 403
    conn = db.get_conn()
    try:
        clubs = conn.execute('SELECT club_name FROM club_profiles ORDER BY club_name').fetchall()
        total_clubs = len(clubs)
        submitted = 0
        not_submitted = []
        for c in clubs:
            sub = conn.execute('SELECT status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (c['club_name'],)).fetchone()
            if sub and sub['status'] == 'tuanwei_approved':
                submitted += 1
            else:
                not_submitted.append(c['club_name'])
        has_final = conn.execute('SELECT COUNT(*) as c FROM final_credits').fetchone()['c']
    finally:
        conn.close()
    return jsonify({
        'success': True,
        'total_clubs': total_clubs,
        'submitted': submitted,
        'not_submitted': not_submitted,
        'all_submitted': submitted == total_clubs and total_clubs > 0,
        'has_final_credits': has_final > 0
    })


@app.route('/api/final-credits/calculate', methods=['POST'])
def final_credits_calculate():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可操作'}), 403
    conn = db.get_conn()
    try:
        date_row = conn.execute('SELECT date_start, date_end FROM scoring_rules LIMIT 1').fetchone()
        date_start = date_row['date_start'] if date_row else ''
        date_end = date_row['date_end'] if date_row else ''
        semester = ''
        if date_start:
            try:
                parts = date_start.split('-')
                y = int(parts[0])
                m = int(parts[1])
                if m >= 9:
                    semester = f'{y}-{y+1}学年第一学期'
                elif m >= 3:
                    semester = f'{y-1}-{y}学年第二学期'
                else:
                    semester = f'{y-1}-{y}学年第一学期'
            except:
                semester = ''
        clubs = conn.execute('SELECT club_name FROM club_profiles').fetchall()
        total_clubs = len(clubs)
        submitted = 0
        for c in clubs:
            sub = conn.execute('SELECT status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (c['club_name'],)).fetchone()
            if sub and sub['status'] == 'tuanwei_approved':
                submitted += 1
        if submitted < total_clubs or total_clubs == 0:
            return jsonify({'error': f'还有 {total_clubs - submitted} 个社团未通过团委审核，无法计算'}), 400
        student_map = {}
        for c in clubs:
            sub = conn.execute('SELECT id FROM scoring_submissions WHERE club_name=? AND status=? ORDER BY created_at DESC LIMIT 1', (c['club_name'], 'tuanwei_approved')).fetchone()
            if not sub:
                continue
            items = conn.execute('SELECT student_id_num, student_name, college, class_name, final_score FROM scoring_submission_items WHERE submission_id=?', (sub['id'],)).fetchall()
            for item in items:
                sid = item['student_id_num']
                if not sid:
                    continue
                if sid not in student_map:
                    sname = item['student_name']
                    if not sname:
                        profile = conn.execute('SELECT real_name FROM user_profiles WHERE student_id=?', (sid,)).fetchone()
                        sname = profile['real_name'] if profile and profile['real_name'] else sid
                    student_map[sid] = {
                        'student_id_num': sid,
                        'student_name': sname,
                        'college': item['college'] or '',
                        'class_name': item['class_name'] or '',
                        'clubs': []
                    }
                student_map[sid]['clubs'].append({
                    'club_name': c['club_name'],
                    'score': item['final_score']
                })
        if semester:
            conn.execute('DELETE FROM final_credits WHERE semester=?', (semester,))
        else:
            conn.execute('DELETE FROM final_credits WHERE semester="" OR semester IS NULL')
        for sid, info in student_map.items():
            clubs_list = info['clubs']
            sorted_clubs = sorted(clubs_list, key=lambda x: x['score'], reverse=True)
            if len(sorted_clubs) == 1:
                final_credit = sorted_clubs[0]['score']
            elif len(sorted_clubs) >= 2:
                final_credit = round(sorted_clubs[0]['score'] + 0.5 * sorted_clubs[1]['score'], 1)
            else:
                final_credit = 0
            club1 = sorted_clubs[0]['club_name'] if len(sorted_clubs) >= 1 else ''
            score1_val = sorted_clubs[0]['score'] if len(sorted_clubs) >= 1 else 0
            club2 = sorted_clubs[1]['club_name'] if len(sorted_clubs) >= 2 else ''
            score2_val = sorted_clubs[1]['score'] if len(sorted_clubs) >= 2 else 0
            club_scores_str = '; '.join([f"{c['club_name']}({c['score']}分)" for c in sorted_clubs])
            conn.execute('INSERT INTO final_credits (student_id_num, student_name, college, class_name, club1, score1, club2, score2, club_scores, final_credit, semester) VALUES (?,?,?,?,?,?,?,?,?,?,?)',
                         (sid, info['student_name'], info['college'], info['class_name'], club1, score1_val, club2, score2_val, club_scores_str, final_credit, semester))
        conn.commit()
        count = len(student_map)
    finally:
        conn.close()
    return jsonify({'success': True, 'message': f'计算完成，共 {count} 名学生', 'count': count})


@app.route('/api/my-final-credit', methods=['GET'])
def my_final_credit():
    user = get_current_user()
    if not user or user['role'] != 'student':
        return jsonify({'error': '仅学生可查看'}), 403
    conn = db.get_conn()
    try:
        profile = conn.execute('SELECT student_id FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
        if not profile or not profile['student_id']:
            return jsonify({'success': True, 'has_credit': False})
        rows = conn.execute('SELECT * FROM final_credits WHERE student_id_num=? ORDER BY semester DESC, calculated_at DESC', (profile['student_id'],)).fetchall()
    finally:
        conn.close()
    if not rows:
        return jsonify({'success': True, 'has_credit': False})
    semesters = []
    seen = set()
    for row in rows:
        sem = row['semester'] if 'semester' in row.keys() and row['semester'] else '未指定学期'
        if sem not in seen:
            seen.add(sem)
            semesters.append({
                'semester': sem,
                'student_name': row['student_name'],
                'college': row['college'],
                'class_name': row['class_name'],
                'club1': row['club1'],
                'score1': row['score1'],
                'club2': row['club2'],
                'score2': row['score2'],
                'club_scores': row['club_scores'],
                'final_credit': row['final_credit']
            })
    latest = semesters[0] if semesters else None
    return jsonify({
        'success': True,
        'has_credit': True,
        'student_name': latest['student_name'] if latest else '',
        'college': latest['college'] if latest else '',
        'class_name': latest['class_name'] if latest else '',
        'club1': latest['club1'] if latest else '',
        'score1': latest['score1'] if latest else 0,
        'club2': latest['club2'] if latest else '',
        'score2': latest['score2'] if latest else 0,
        'club_scores': latest['club_scores'] if latest else '',
        'final_credit': latest['final_credit'] if latest else 0,
        'semesters': semesters
    })


@app.route('/api/final-credits/list', methods=['GET'])
def final_credits_list():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可查看'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT id, student_id_num, student_name, college, class_name, club1, score1, club2, score2, club_scores, final_credit, semester, calculated_at FROM final_credits ORDER BY semester DESC, student_id_num').fetchall()
    finally:
        conn.close()
    return jsonify({'success': True, 'data': [dict(r) for r in rows]})


@app.route('/api/final-credits/export', methods=['GET'])
def final_credits_export():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可导出'}), 403
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT student_id_num, student_name, college, class_name, club1, score1, club2, score2, final_credit, semester FROM final_credits ORDER BY semester DESC, student_id_num').fetchall()
    finally:
        conn.close()
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '汇总'
    headers = ['学期', '学号', '姓名', '学院', '班级', '加入社团1', '社团1赋分', '加入社团2', '社团2赋分', '最终学分']
    ws.append(headers)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = openpyxl.styles.Font(bold=True)
    for r in rows:
        ws.append([r['semester'] or '', r['student_id_num'], r['student_name'], r['college'], r['class_name'], r['club1'], r['score1'], r['club2'] or '', r['score2'] if r['score2'] else '', r['final_credit']])
    for col in ws.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 4, 50)
    college_map = {}
    for r in rows:
        c = r['college'] or '未知学院'
        if c not in college_map:
            college_map[c] = []
        college_map[c].append(r)
    for college_name, college_rows in college_map.items():
        sheet_name = college_name[:31]
        ws_c = wb.create_sheet(title=sheet_name)
        ws_c.append(headers)
        for col_idx, header in enumerate(headers, 1):
            cell = ws_c.cell(row=1, column=col_idx)
            cell.font = openpyxl.styles.Font(bold=True)
        for r in college_rows:
            ws_c.append([r['semester'] or '', r['student_id_num'], r['student_name'], r['college'], r['class_name'], r['club1'], r['score1'], r['club2'] or '', r['score2'] if r['score2'] else '', r['final_credit']])
        for col in ws_c.columns:
            max_length = 0
            for cell in col:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws_c.column_dimensions[col[0].column_letter].width = min(max_length + 4, 50)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name='全校学生最终学分清单.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/scoring/export-all', methods=['GET'])
def scoring_export_all():
    user = get_current_user()
    if not user or user['role'] != 'admin':
        return jsonify({'error': '仅管理员可导出'}), 403
    conn = db.get_conn()
    try:
        clubs = conn.execute('SELECT club_name FROM club_profiles ORDER BY club_name').fetchall()
    finally:
        conn.close()
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '目录'
    ws.append(['社团名称', '状态'])
    ws.cell(row=1, column=1).font = openpyxl.styles.Font(bold=True)
    ws.cell(row=1, column=2).font = openpyxl.styles.Font(bold=True)
    conn = db.get_conn()
    try:
        for c in clubs:
            club_name = c['club_name']
            sub = conn.execute('SELECT id, status FROM scoring_submissions WHERE club_name=? ORDER BY created_at DESC LIMIT 1', (club_name,)).fetchone()
            status_text = {'draft': '草稿', 'pending_teacher': '待老师审核', 'teacher_approved': '老师已通过', 'submitted_tuanwei': '待团委审核', 'tuanwei_approved': '团委已通过'}
            st = status_text.get(sub['status'], '未提交') if sub else '未提交'
            ws.append([club_name, st])
            if not sub:
                continue
            items = conn.execute('SELECT student_id_num, student_name, college, class_name, total_workload, final_score FROM scoring_submission_items WHERE submission_id=? ORDER BY student_id_num', (sub['id'],)).fetchall()
            sheet_name = club_name[:31]
            ws_c = wb.create_sheet(title=sheet_name)
            headers = ['学号', '姓名', '学院', '班级', '总工作量', '个人赋分']
            ws_c.append(headers)
            for col_idx, header in enumerate(headers, 1):
                ws_c.cell(row=1, column=col_idx).font = openpyxl.styles.Font(bold=True)
            for item in items:
                ws_c.append([item['student_id_num'], item['student_name'], item['college'], item['class_name'], item['total_workload'], item['final_score']])
            for col in ws_c.columns:
                max_length = 0
                for cell in col:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                ws_c.column_dimensions[col[0].column_letter].width = min(max_length + 4, 50)
    finally:
        conn.close()
    for col in ws.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 4, 50)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name='各社团赋分明细.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/club-favorites/toggle', methods=['POST'])
def toggle_club_favorite():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    club_name = (data.get('club_name') or '').strip()
    if not club_name:
        return jsonify({'error': '社团名称不能为空'}), 400
    conn = db.get_conn()
    try:
        existing = conn.execute('SELECT id FROM club_favorites WHERE user_id=? AND club_name=?', (user['id'], club_name)).fetchone()
        if existing:
            conn.execute('DELETE FROM club_favorites WHERE id=?', (existing['id'],))
            conn.commit()
            return jsonify({'success': True, 'favorited': False})
        else:
            conn.execute('INSERT INTO club_favorites (user_id, club_name) VALUES (?, ?)', (user['id'], club_name))
            conn.commit()
            return jsonify({'success': True, 'favorited': True})
    finally:
        conn.close()


@app.route('/api/club-favorites')
def get_club_favorites():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT cf.club_name, cf.created_at, cp.description, cp.star_rating, cp.show_star, cp.emblem_url, cp.category FROM club_favorites cf LEFT JOIN club_profiles cp ON cf.club_name=cp.club_name WHERE cf.user_id=? ORDER BY cf.created_at DESC', (user['id'],)).fetchall()
        result = []
        for r in rows:
            result.append({
                'clubName': r['club_name'],
                'description': r['description'] or '',
                'starRating': r['star_rating'] or 0,
                'showStar': r['show_star'] or 0,
                'emblemUrl': r['emblem_url'] or '',
                'category': r['category'] or '',
                'favoritedAt': r['created_at']
            })
        return jsonify({'success': True, 'data': result})
    finally:
        conn.close()


@app.route('/api/quit-apply', methods=['POST'])
def submit_quit_application():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    data = request.json or {}
    club_name = (data.get('club_name') or '').strip()
    reason = (data.get('reason') or '').strip()
    if not club_name:
        return jsonify({'error': '社团名称不能为空'}), 400
    if not reason:
        return jsonify({'error': '请填写退社原因'}), 400
    conn = db.get_conn()
    try:
        member = conn.execute('SELECT id, user_id FROM club_members WHERE club_name=? AND user_id=?', (club_name, user['id'])).fetchone()
        if not member:
            member = conn.execute('SELECT id, user_id FROM club_members WHERE club_name=? AND username=?', (club_name, user['username'])).fetchone()
        if not member:
            profile = conn.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (user['id'],)).fetchone()
            if profile and profile['real_name']:
                member = conn.execute('SELECT id, user_id FROM club_members WHERE club_name=? AND real_name=?', (club_name, profile['real_name'])).fetchone()
        if not member:
            return jsonify({'error': '你不在该社团中'}), 400
        if member['user_id'] == 0 or member['user_id'] is None:
            conn.execute('UPDATE club_members SET user_id=?, username=? WHERE id=?', (user['id'], user['username'], member['id']))
            conn.commit()
        existing = conn.execute('SELECT id FROM quit_applications WHERE user_id=? AND club_name=? AND status=?', (user['id'], club_name, 'pending')).fetchone()
        if existing:
            return jsonify({'error': '你已提交过该社团的退社申请，请等待审批'}), 400
        conn.execute('INSERT INTO quit_applications (user_id, username, club_name, reason) VALUES (?, ?, ?, ?)', (user['id'], user.get('username', ''), club_name, reason))
        conn.commit()
        return jsonify({'success': True, 'message': '退社申请已提交，请等待社团负责人审批'})
    finally:
        conn.close()


@app.route('/api/my-quit-applications')
def get_my_quit_applications():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    conn = db.get_conn()
    try:
        rows = conn.execute('SELECT * FROM quit_applications WHERE user_id=? ORDER BY created_at DESC', (user['id'],)).fetchall()
        result = []
        for r in rows:
            result.append({
                'id': r['id'],
                'clubName': r['club_name'],
                'reason': r['reason'],
                'status': r['status'],
                'handlerNote': r['handler_note'] or '',
                'createdAt': r['created_at'],
                'handledAt': r['handled_at'] or ''
            })
        return jsonify({'success': True, 'data': result})
    finally:
        conn.close()


@app.route('/api/quit-applications/manage')
def get_quit_applications_manage():
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    if user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '无权限'}), 403
    conn = db.get_conn()
    try:
        if user['role'] == 'admin':
            rows = conn.execute('SELECT * FROM quit_applications ORDER BY created_at DESC').fetchall()
        elif user['role'] == 'teacher':
            teacher_clubs = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
            club_names = [r['club_name'] for r in teacher_clubs]
            if not club_names:
                club_name = user.get('club_name', '')
                if club_name:
                    club_names = [club_name]
            if not club_names:
                return jsonify({'success': True, 'data': []})
            placeholders = ','.join(['?' for _ in club_names])
            rows = conn.execute(f'SELECT * FROM quit_applications WHERE club_name IN ({placeholders}) ORDER BY created_at DESC', club_names).fetchall()
        else:
            club_name = user.get('club_name', '')
            if not club_name:
                return jsonify({'success': True, 'data': []})
            rows = conn.execute('SELECT * FROM quit_applications WHERE club_name=? ORDER BY created_at DESC', (club_name,)).fetchall()
        result = []
        for r in rows:
            real_name = r['username']
            if r['user_id'] and r['user_id'] != 0:
                profile = conn.execute('SELECT real_name FROM user_profiles WHERE user_id=?', (r['user_id'],)).fetchone()
                if profile and profile['real_name']:
                    real_name = profile['real_name']
            else:
                member = conn.execute('SELECT real_name FROM club_members WHERE club_name=? AND username=?', (r['club_name'], r['username'])).fetchone()
                if member and member['real_name']:
                    real_name = member['real_name']
            result.append({
                'id': r['id'],
                'userId': r['user_id'],
                'username': r['username'],
                'realName': real_name,
                'clubName': r['club_name'],
                'reason': r['reason'],
                'status': r['status'],
                'handlerNote': r['handler_note'] or '',
                'createdAt': r['created_at'],
                'handledAt': r['handled_at'] or ''
            })
        return jsonify({'success': True, 'data': result})
    finally:
        conn.close()


@app.route('/api/quit-applications/<int:app_id>/handle', methods=['POST'])
def handle_quit_application(app_id):
    user = get_current_user()
    if not user:
        return jsonify({'error': '未登录'}), 401
    if user['role'] not in ('user', 'admin', 'teacher'):
        return jsonify({'error': '无权限'}), 403
    data = request.json or {}
    action = data.get('action', '')
    handler_note = (data.get('handler_note') or '').strip()
    if action not in ('approve', 'reject'):
        return jsonify({'error': '无效操作'}), 400
    conn = db.get_conn()
    try:
        app = conn.execute('SELECT * FROM quit_applications WHERE id=?', (app_id,)).fetchone()
        if not app:
            return jsonify({'error': '申请不存在'}), 404
        if app['status'] != 'pending':
            return jsonify({'error': '该申请已处理'}), 400
        if user['role'] != 'admin' and app['club_name'] != user.get('club_name', ''):
            if user['role'] == 'teacher':
                teacher_clubs = conn.execute('SELECT club_name FROM teacher_clubs WHERE user_id=?', (user['id'],)).fetchall()
                teacher_club_names = [r['club_name'] for r in teacher_clubs]
                if app['club_name'] not in teacher_club_names:
                    return jsonify({'error': '无权限处理该申请'}), 403
            else:
                return jsonify({'error': '无权限处理该申请'}), 403
        if action == 'approve':
            conn.execute('UPDATE quit_applications SET status=?, handler_note=?, handled_at=CURRENT_TIMESTAMP WHERE id=?', ('approved', handler_note, app_id))
            conn.execute('DELETE FROM club_members WHERE club_name=? AND user_id=?', (app['club_name'], app['user_id']))
            conn.execute('DELETE FROM club_registrations WHERE club_name=? AND user_id=?', (app['club_name'], app['user_id']))
            conn.execute('DELETE FROM checkin_records WHERE club_name=? AND student_name=?', (app['club_name'], app['username']))
            conn.execute('DELETE FROM club_favorites WHERE user_id=? AND club_name=?', (app['user_id'], app['club_name']))
            conn.commit()
            return jsonify({'success': True, 'message': '已同意退社申请，该成员信息及活动记录已清除'})
        else:
            conn.execute('UPDATE quit_applications SET status=?, handler_note=?, handled_at=CURRENT_TIMESTAMP WHERE id=?', ('rejected', handler_note, app_id))
            conn.commit()
            return jsonify({'success': True, 'message': '已拒绝退社申请'})
    finally:
        conn.close()


# ==================== 智能周报 API ====================

@app.route('/api/weekly-report/generate', methods=['POST'])
def generate_weekly_report():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401

    from datetime import datetime, timedelta
    now = datetime.now()
    # 本周一到今天
    week_start_date = now - timedelta(days=now.weekday())
    week_start = week_start_date.strftime('%Y-%m-%d')
    week_end = now.strftime('%Y-%m-%d')

    conn = db.get_conn()
    try:
        if user['role'] == 'user':
            club_name = user.get('club_name', '')
            if not club_name:
                return jsonify({'error': '您未关联社团，无法生成周报'}), 400

            # 社团活动次数
            activity_count = conn.execute(
                'SELECT COUNT(*) as c FROM online_activity_data WHERE club_name=? AND activity_date>=? AND activity_date<=?',
                (club_name, week_start, week_end)
            ).fetchone()['c']

            # 各活动的参加人数（通过 checkin_sessions + checkin_records 统计）
            activities = conn.execute(
                'SELECT id, activity_title, activity_date FROM online_activity_data WHERE club_name=? AND activity_date>=? AND activity_date<=?',
                (club_name, week_start, week_end)
            ).fetchall()

            activity_participants = []
            for act in activities:
                # 通过 checkin_sessions 查找该活动对应的签到场次，统计参加人数
                sessions = conn.execute(
                    "SELECT id FROM checkin_sessions WHERE club_name=? AND activity_name=? AND date(created_at)>=? AND date(created_at)<=?",
                    (club_name, act['activity_title'] or '', week_start, week_end)
                ).fetchall()
                total = 0
                for s in sessions:
                    cnt = conn.execute(
                        'SELECT COUNT(*) as c FROM checkin_records WHERE session_id=?', (s['id'],)
                    ).fetchone()['c']
                    total += cnt
                activity_participants.append({
                    'title': act['activity_title'] or '未命名活动',
                    'count': total
                })

            max_activity = None
            min_activity = None
            if activity_participants:
                max_activity = max(activity_participants, key=lambda x: x['count'])
                min_activity = min(activity_participants, key=lambda x: x['count'])

            # 指导老师指导次数
            teacher_guidance_count = conn.execute(
                'SELECT COUNT(*) as c FROM teacher_checkin_checkout WHERE club_name=? AND date(checkin_time)>=? AND date(checkin_time)<=?',
                (club_name, week_start, week_end)
            ).fetchone()['c']

            # 组织周报内容
            content_lines = [
                f'【{club_name} 周报】',
                f'统计周期：{week_start} 至 {week_end}',
                '',
                f'一、社团活动概况',
                f'本周共开展活动 {activity_count} 次。',
            ]
            if max_activity:
                content_lines.append(f'参加人数最多的活动：「{max_activity["title"]}」，共 {max_activity["count"]} 人参加。')
            if min_activity:
                content_lines.append(f'参加人数最少的活动：「{min_activity["title"]}」，共 {min_activity["count"]} 人参加。')
            content_lines.append('')
            content_lines.append(f'二、指导老师指导情况')
            content_lines.append(f'本周指导老师到校指导 {teacher_guidance_count} 次。')

            content = '\n'.join(content_lines)

            # 检查本周是否已有该社团的周报
            existing = conn.execute(
                'SELECT id FROM weekly_reports WHERE role=? AND club_name=? AND week_start=? AND week_end=?',
                ('user', club_name, week_start, week_end)
            ).fetchone()
            if existing:
                conn.execute(
                    'UPDATE weekly_reports SET content=?, created_at=CURRENT_TIMESTAMP WHERE id=?',
                    (content, existing['id'])
                )
                report_id = existing['id']
            else:
                conn.execute(
                    'INSERT INTO weekly_reports (role, club_name, week_start, week_end, content) VALUES (?, ?, ?, ?, ?)',
                    ('user', club_name, week_start, week_end, content)
                )
                report_id = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
            conn.commit()

            report = conn.execute('SELECT id, week_start, week_end, content, created_at FROM weekly_reports WHERE id=?', (report_id,)).fetchone()
            return jsonify({'success': True, 'report': dict(report)})

        elif user['role'] == 'admin':
            # 活动的社团个数
            active_club_count = conn.execute(
                'SELECT COUNT(DISTINCT club_name) as c FROM online_activity_data WHERE activity_date>=? AND activity_date<=?',
                (week_start, week_end)
            ).fetchone()['c']

            # 活动次数最多的五个社团
            top5_activity_clubs = conn.execute(
                'SELECT club_name, COUNT(*) as c FROM online_activity_data WHERE activity_date>=? AND activity_date<=? GROUP BY club_name ORDER BY c DESC LIMIT 5',
                (week_start, week_end)
            ).fetchall()

            # 指导老师指导次数最多的五个老师
            top5_teachers = conn.execute(
                'SELECT tcc.teacher_user_id, tp.real_name, COUNT(*) as c FROM teacher_checkin_checkout tcc LEFT JOIN teacher_profiles tp ON tcc.teacher_user_id=tp.user_id WHERE date(tcc.checkin_time)>=? AND date(tcc.checkin_time)<=? GROUP BY tcc.teacher_user_id ORDER BY c DESC LIMIT 5',
                (week_start, week_end)
            ).fetchall()

            # 提交材料情况
            upload_rows = conn.execute(
                'SELECT club_name, COUNT(*) as c FROM club_uploads WHERE date(upload_time)>=? AND date(upload_time)<=? GROUP BY club_name',
                (week_start, week_end)
            ).fetchall()
            has_uploads = len(upload_rows) > 0

            # 外出情况
            offcampus_rows = conn.execute(
                'SELECT club_name, COUNT(*) as c FROM offcampus_requests WHERE date(created_at)>=? AND date(created_at)<=? GROUP BY club_name',
                (week_start, week_end)
            ).fetchall()
            has_offcampus = len(offcampus_rows) > 0

            # 联合活动发布情况
            joint_rows = conn.execute(
                'SELECT club_name, COUNT(*) as c FROM joint_activities WHERE date(created_at)>=? AND date(created_at)<=? GROUP BY club_name',
                (week_start, week_end)
            ).fetchall()
            has_joint = len(joint_rows) > 0

            # 招募情况
            recruit_rows = conn.execute(
                'SELECT club_name, COUNT(*) as c FROM recruitments WHERE date(created_at)>=? AND date(created_at)<=? GROUP BY club_name',
                (week_start, week_end)
            ).fetchall()
            has_recruit = len(recruit_rows) > 0

            # 组织周报内容
            content_lines = [
                '【管理员周报】',
                f'统计周期：{week_start} 至 {week_end}',
                '',
                f'一、活动概况',
                f'本周共有 {active_club_count} 个社团开展了活动。',
            ]
            if top5_activity_clubs:
                content_lines.append('活动次数最多的社团：')
                for i, row in enumerate(top5_activity_clubs, 1):
                    content_lines.append(f'  {i}. {row["club_name"]}：{row["c"]} 次')
            content_lines.append('')

            content_lines.append('二、指导老师指导情况')
            if top5_teachers:
                content_lines.append('指导次数最多的老师：')
                for i, row in enumerate(top5_teachers, 1):
                    name = row['real_name'] or f'老师(ID:{row["teacher_user_id"]})'
                    content_lines.append(f'  {i}. {name}：{row["c"]} 次')
            else:
                content_lines.append('本周暂无指导老师指导记录。')
            content_lines.append('')

            cn_list = ['一', '二', '三', '四', '五', '六']
            sec_idx = 2  # 一=活动概况, 二=指导老师, 从三开始

            if has_uploads:
                content_lines.append(f'{cn_list[sec_idx]}、材料提交情况')
                for row in upload_rows:
                    content_lines.append(f'  {row["club_name"]}：提交 {row["c"]} 份材料')
                content_lines.append('')
                sec_idx += 1

            if has_offcampus:
                content_lines.append(f'{cn_list[sec_idx]}、外出申请情况')
                for row in offcampus_rows:
                    content_lines.append(f'  {row["club_name"]}：{row["c"]} 条申请')
                content_lines.append('')
                sec_idx += 1

            if has_joint:
                content_lines.append(f'{cn_list[sec_idx]}、联合活动发布情况')
                for row in joint_rows:
                    content_lines.append(f'  {row["club_name"]}：发布 {row["c"]} 条联合活动')
                content_lines.append('')
                sec_idx += 1

            if has_recruit:
                content_lines.append(f'{cn_list[sec_idx]}、招募情况')
                for row in recruit_rows:
                    content_lines.append(f'  {row["club_name"]}：发布 {row["c"]} 条招募')
                content_lines.append('')

            content = '\n'.join(content_lines)

            # 检查本周是否已有管理员周报
            existing = conn.execute(
                'SELECT id FROM weekly_reports WHERE role=? AND week_start=? AND week_end=?',
                ('admin', week_start, week_end)
            ).fetchone()
            if existing:
                conn.execute(
                    'UPDATE weekly_reports SET content=?, created_at=CURRENT_TIMESTAMP WHERE id=?',
                    (content, existing['id'])
                )
                report_id = existing['id']
            else:
                conn.execute(
                    'INSERT INTO weekly_reports (role, club_name, week_start, week_end, content) VALUES (?, ?, ?, ?, ?)',
                    ('admin', '', week_start, week_end, content)
                )
                report_id = conn.execute('SELECT last_insert_rowid() as id').fetchone()['id']
            conn.commit()

            report = conn.execute('SELECT id, week_start, week_end, content, created_at FROM weekly_reports WHERE id=?', (report_id,)).fetchone()
            return jsonify({'success': True, 'report': dict(report)})
    finally:
        conn.close()


@app.route('/api/weekly-report/latest', methods=['GET'])
def get_latest_weekly_report():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401

    conn = db.get_conn()
    try:
        if user['role'] == 'user':
            club_name = user.get('club_name', '')
            report = conn.execute(
                'SELECT id, role, club_name, week_start, week_end, content, created_at FROM weekly_reports WHERE role=? AND club_name=? ORDER BY created_at DESC LIMIT 1',
                ('user', club_name)
            ).fetchone()
        else:
            report = conn.execute(
                'SELECT id, role, club_name, week_start, week_end, content, created_at FROM weekly_reports WHERE role=? ORDER BY created_at DESC LIMIT 1',
                ('admin',)
            ).fetchone()
        if report:
            return jsonify({'success': True, 'report': dict(report)})
        else:
            return jsonify({'success': True, 'report': None})
    finally:
        conn.close()


@app.route('/api/weekly-report/list', methods=['GET'])
def list_weekly_reports():
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401

    conn = db.get_conn()
    try:
        if user['role'] == 'user':
            club_name = user.get('club_name', '')
            rows = conn.execute(
                "SELECT id, role, club_name, week_start, week_end, content, created_at FROM weekly_reports WHERE role=? AND club_name=? AND created_at>=date('now','-30 days') ORDER BY created_at DESC",
                ('user', club_name)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT id, role, club_name, week_start, week_end, content, created_at FROM weekly_reports WHERE role=? AND created_at>=date('now','-30 days') ORDER BY created_at DESC",
                ('admin',)
            ).fetchall()
        return jsonify({'success': True, 'list': [dict(r) for r in rows]})
    finally:
        conn.close()


@app.route('/api/weekly-report/<int:report_id>', methods=['GET'])
def get_weekly_report(report_id):
    user = get_current_user()
    if not user or user['role'] not in ('user', 'admin'):
        return jsonify({'error': '请先登录'}), 401

    conn = db.get_conn()
    try:
        report = conn.execute(
            'SELECT id, role, club_name, week_start, week_end, content, created_at FROM weekly_reports WHERE id=?',
            (report_id,)
        ).fetchone()
        if not report:
            return jsonify({'error': '周报不存在'}), 404
        # 权限校验：社团负责人只能看自己社团的
        if user['role'] == 'user' and report['club_name'] != user.get('club_name', ''):
            return jsonify({'error': '无权限查看该周报'}), 403
        return jsonify({'success': True, 'report': dict(report)})
    finally:
        conn.close()


if __name__ == '__main__':
    import sys
    print('\n  ' + '=' * 50)
    print('  社团活动统计分析系统')
    print('  技术栈: Flask + SQLite + Pandas + MapReduce')
    print('  ' + '=' * 50 + '\n')
    # 生产环境使用 threaded=True 和 debug=False 提高稳定性
    is_debug = '--debug' in sys.argv
    use_ssl = '--ssl' in sys.argv
    ssl_ctx = None
    if use_ssl:
        cert_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'cert')
        cert_file = os.path.join(cert_dir, 'cert.pem')
        key_file = os.path.join(cert_dir, 'key.pem')
        if os.path.exists(cert_file) and os.path.exists(key_file):
            ssl_ctx = (cert_file, key_file)
            print('  🔒 HTTPS 已启用')
        else:
            print('  ⚠️ 未找到证书文件，正在自动生成自签名证书...')
            try:
                from cryptography import x509
                from cryptography.x509.oid import NameOID
                from cryptography.hazmat.primitives import hashes, serialization
                from cryptography.hazmat.primitives.asymmetric import rsa
                import datetime
                import ipaddress as _ip
                os.makedirs(cert_dir, exist_ok=True)
                key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
                subject = issuer = x509.Name([x509.NameAttribute(NameOID.COMMON_NAME, 'localhost')])
                cert = (x509.CertificateBuilder()
                    .subject_name(subject).issuer_name(issuer)
                    .public_key(key.public_key())
                    .serial_number(x509.random_serial_number())
                    .not_valid_before(datetime.datetime.utcnow())
                    .not_valid_after(datetime.datetime.utcnow() + datetime.timedelta(days=365))
                    .add_extension(x509.SubjectAlternativeName([x509.DNSName('localhost'), x509.IPAddress(_ip.IPv4Address('0.0.0.0'))]), critical=False)
                    .sign(key, hashes.SHA256()))
                with open(cert_file, 'wb') as f: f.write(cert.public_bytes(serialization.Encoding.PEM))
                with open(key_file, 'wb') as f: f.write(key.private_bytes(serialization.Encoding.PEM, serialization.PrivateFormat.TraditionalOpenSSL, serialization.NoEncryption()))
                ssl_ctx = (cert_file, key_file)
                print('  🔒 HTTPS 已启用（自签名证书）')
            except ImportError:
                print('  ⚠️ 缺少 cryptography 库，尝试 openssl 生成...')
                try:
                    from subprocess import run, PIPE
                    os.makedirs(cert_dir, exist_ok=True)
                    run(['openssl', 'req', '-x509', '-newkey', 'rsa:2048', '-keyout', key_file, '-out', cert_file, '-days', '365', '-nodes', '-subj', '/CN=localhost'], check=True, stdout=PIPE, stderr=PIPE)
                    ssl_ctx = (cert_file, key_file)
                    print('  🔒 HTTPS 已启用（自签名证书）')
                except Exception as e2:
                    print(f'  ⚠️ 证书生成失败: {e2}，将以 HTTP 模式运行')
            except Exception as e:
                print(f'  ⚠️ 证书生成失败: {e}，将以 HTTP 模式运行')
    if ssl_ctx:
        print(f'  📱 手机请访问: https://<本机IP>:5000')
        app.run(host='0.0.0.0', port=5000, debug=is_debug, threaded=True, ssl_context=ssl_ctx)
    else:
        app.run(host='0.0.0.0', port=5000, debug=is_debug, threaded=True)
